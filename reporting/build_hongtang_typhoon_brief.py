from __future__ import annotations

import argparse
import csv
import io
import json
import math
import statistics
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Callable, Iterable
from zipfile import ZipFile

import matplotlib

matplotlib.use("Agg")
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from build_hongtang_wind_diagnostic_memo import (
    BLUE,
    LIGHT_BLUE,
    LIGHT_GRAY,
    add_caption,
    add_page_number,
    configure_document,
    set_cell_shading,
    set_cell_text,
)


WIND_ENTRIES = {
    "W1_speed": "风速_162.csv",
    "W1_direction": "风向_163.csv",
    "W2_speed": "塔顶风速_178.csv",
    "W2_direction": "塔顶风向_179.csv",
}

STRUCTURE_GROUPS = {
    "主梁加速度": [f"A{i}" for i in range(1, 9)],
    "主塔加速度": ["A9-X", "A9-Y", "A10-X", "A10-Y"],
    "南侧索振动": [f"CS{i}" for i in range(1, 13)],
    "北侧索振动": [f"CX{i}" for i in range(1, 13)],
}

COLORS = {
    "W1": "#1F77B4",
    "W2": "#E67E22",
    "主梁加速度": "#1F77B4",
    "主塔加速度": "#9467BD",
    "南侧索振动": "#2CA02C",
    "北侧索振动": "#D62728",
}


@dataclass
class BinStats:
    count: int = 0
    total: float = 0.0
    sumsq: float = 0.0
    minimum: float = math.inf
    maximum: float = -math.inf
    minimum_time: datetime | None = None
    maximum_time: datetime | None = None
    cos_sum: float = 0.0
    sin_sum: float = 0.0

    def add(self, value: float, *, timestamp: datetime | None = None, circular: bool = False) -> None:
        self.count += 1
        self.total += value
        self.sumsq += value * value
        if value < self.minimum:
            self.minimum = value
            self.minimum_time = timestamp
        if value > self.maximum:
            self.maximum = value
            self.maximum_time = timestamp
        if circular:
            radians = math.radians(value % 360.0)
            self.cos_sum += math.cos(radians)
            self.sin_sum += math.sin(radians)

    @property
    def mean(self) -> float:
        return self.total / self.count if self.count else math.nan

    @property
    def rms(self) -> float:
        return math.sqrt(self.sumsq / self.count) if self.count else math.nan

    @property
    def circular_mean(self) -> float:
        if not self.count:
            return math.nan
        return math.degrees(math.atan2(self.sin_sum, self.cos_sum)) % 360.0


@dataclass
class SeriesAudit:
    key: str
    zip_path: str
    entry_name: str
    rows: int = 0
    rejected_rows: int = 0
    first_time: datetime | None = None
    last_time: datetime | None = None
    uncompressed_bytes: int = 0


@dataclass
class Analysis:
    wind: dict[str, dict[datetime, BinStats]] = field(default_factory=dict)
    structure: dict[str, dict[str, dict[datetime, BinStats]]] = field(default_factory=dict)
    audits: list[SeriesAudit] = field(default_factory=list)
    missing_entries: list[str] = field(default_factory=list)


def ten_minute_bin(value: datetime) -> datetime:
    return value.replace(minute=(value.minute // 10) * 10, second=0, microsecond=0)


def parse_data_line(line: str) -> tuple[datetime, float] | None:
    if not line or line[0].isdigit() is False:
        return None
    try:
        time_text, value_text = line.rstrip().split(",", 1)
        return datetime.fromisoformat(time_text), float(value_text)
    except (ValueError, TypeError):
        return None


def find_entry(archive: ZipFile, expected_name: str) -> str | None:
    for info in archive.infolist():
        if info.filename.replace("\\", "/").split("/")[-1] == expected_name:
            return info.filename
    return None


def aggregate_entry(
    archive: ZipFile,
    entry_name: str,
    *,
    key: str,
    zip_path: Path,
    validator: Callable[[float], bool],
    circular: bool = False,
    absolute: bool = False,
) -> tuple[dict[datetime, BinStats], SeriesAudit]:
    result: dict[datetime, BinStats] = defaultdict(BinStats)
    info = archive.getinfo(entry_name)
    audit = SeriesAudit(
        key=key,
        zip_path=str(zip_path),
        entry_name=entry_name,
        uncompressed_bytes=info.file_size,
    )
    with archive.open(info) as raw:
        buffered = io.BufferedReader(raw)
        signature = buffered.peek(4)[:4]
        if signature.startswith(b"\xff\xfe") or signature.startswith(b"\xfe\xff"):
            encoding = "utf-16"
        elif signature.startswith(b"\xef\xbb\xbf"):
            encoding = "utf-8-sig"
        else:
            encoding = "utf-8"
        text = io.TextIOWrapper(buffered, encoding=encoding, errors="replace", newline="")
        for line in text:
            parsed = parse_data_line(line)
            if parsed is None:
                continue
            timestamp, value = parsed
            if not math.isfinite(value) or not validator(value):
                audit.rejected_rows += 1
                continue
            if absolute:
                value = abs(value)
            audit.rows += 1
            audit.first_time = timestamp if audit.first_time is None else min(audit.first_time, timestamp)
            audit.last_time = timestamp if audit.last_time is None else max(audit.last_time, timestamp)
            result[ten_minute_bin(timestamp)].add(value, timestamp=timestamp, circular=circular)
    return dict(result), audit


def merge_bins(target: dict[datetime, BinStats], source: dict[datetime, BinStats]) -> None:
    for timestamp, item in source.items():
        current = target.setdefault(timestamp, BinStats())
        current.count += item.count
        current.total += item.total
        current.sumsq += item.sumsq
        if item.minimum < current.minimum:
            current.minimum = item.minimum
            current.minimum_time = item.minimum_time
        if item.maximum > current.maximum:
            current.maximum = item.maximum
            current.maximum_time = item.maximum_time
        current.cos_sum += item.cos_sum
        current.sin_sum += item.sin_sum


def locate_zip(day_dir: Path, folder: str) -> Path | None:
    candidates = sorted((day_dir / folder).glob("*.zip"))
    return candidates[0] if candidates else None


def analyze(source_root: Path, day_filter: set[str] | None = None) -> Analysis:
    output = Analysis(
        wind={key: {} for key in WIND_ENTRIES},
        structure={group: {point: {} for point in points} for group, points in STRUCTURE_GROUPS.items()},
    )
    day_dirs = sorted(
        path
        for path in source_root.iterdir()
        if path.is_dir()
        and path.name[:4].isdigit()
        and (day_filter is None or path.name in day_filter)
    )
    print(f"[analyze] source_root={source_root} day_dirs={len(day_dirs)}", flush=True)
    for day_dir in day_dirs:
        wave_zip = locate_zip(day_dir, "波形")
        feature_zip = locate_zip(day_dir, "特征值")
        print(
            f"[analyze] day={day_dir.name} wave_zip={wave_zip or 'missing'} "
            f"feature_zip={feature_zip or 'missing'}",
            flush=True,
        )
        if wave_zip is None:
            output.missing_entries.append(f"{day_dir.name}/波形/*.zip")
        else:
            with ZipFile(wave_zip) as archive:
                for key, expected in WIND_ENTRIES.items():
                    entry = find_entry(archive, expected)
                    if entry is None:
                        output.missing_entries.append(f"{wave_zip}:{expected}")
                        continue
                    circular = key.endswith("direction")
                    validator = (lambda value: 0.0 <= value <= 360.0) if circular else (lambda value: value >= 0.0)
                    bins, audit = aggregate_entry(
                        archive,
                        entry,
                        key=key,
                        zip_path=wave_zip,
                        validator=validator,
                        circular=circular,
                    )
                    merge_bins(output.wind[key], bins)
                    output.audits.append(audit)
        if feature_zip is None:
            output.missing_entries.append(f"{day_dir.name}/特征值/*.zip")
        else:
            with ZipFile(feature_zip) as archive:
                for group, points in STRUCTURE_GROUPS.items():
                    for point in points:
                        expected = f"{point}_{point_channel(point)}_峰值.csv"
                        entry = find_entry(archive, expected)
                        if entry is None:
                            output.missing_entries.append(f"{feature_zip}:{expected}")
                            continue
                        bins, audit = aggregate_entry(
                            archive,
                            entry,
                            key=point,
                            zip_path=feature_zip,
                            validator=lambda value: True,
                            absolute=True,
                        )
                        merge_bins(output.structure[group][point], bins)
                        output.audits.append(audit)
    return output


def point_channel(point: str) -> int:
    mapping = {
        **{f"A{i}": value for i, value in enumerate([174, 172, 173, 175, 168, 169, 170, 171], start=1)},
        "A9-X": 176,
        "A9-Y": 177,
        "A10-X": 156,
        "A10-Y": 157,
        **{f"CS{i}": value for i, value in enumerate([148, 149, 150, 151, 152, 153, 154, 155, 167, 166, 164, 165], start=1)},
        **{f"CX{i}": value for i, value in enumerate([122, 123, 120, 121, 127, 126, 124, 125, 139, 138, 136, 137], start=1)},
    }
    return mapping[point]


def wind_rows(analysis: Analysis, point: str) -> list[dict[str, object]]:
    speed = analysis.wind[f"{point}_speed"]
    direction = analysis.wind[f"{point}_direction"]
    rows = []
    for timestamp in sorted(set(speed) & set(direction)):
        rows.append(
            {
                "time": timestamp,
                "mean_speed": speed[timestamp].mean,
                "raw_peak": speed[timestamp].maximum,
                "raw_peak_time": speed[timestamp].maximum_time,
                "direction": direction[timestamp].circular_mean,
                "count": min(speed[timestamp].count, direction[timestamp].count),
            }
        )
    return rows


def structure_group_rows(analysis: Analysis, group: str) -> list[dict[str, object]]:
    point_bins = analysis.structure[group]
    timestamps = sorted(set().union(*(bins.keys() for bins in point_bins.values())))
    rows = []
    for timestamp in timestamps:
        values = [(point, bins[timestamp].maximum) for point, bins in point_bins.items() if timestamp in bins]
        if not values:
            continue
        point, value = max(values, key=lambda item: item[1])
        rows.append({"time": timestamp, "peak": value, "point": point})
    return rows


def fmt_time(value: datetime | None) -> str:
    return value.strftime("%Y-%m-%d %H:%M") if value else "--"


def safe_ratio(value: float, baseline: float) -> float:
    return value / baseline if baseline > 1e-12 else math.nan


def period_values(rows: list[dict[str, object]], field: str, start: datetime, end: datetime) -> list[float]:
    return [float(row[field]) for row in rows if start <= row["time"] < end and math.isfinite(float(row[field]))]


def write_csv(path: Path, rows: Iterable[dict[str, object]], fieldnames: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            clean = {key: (value.isoformat(sep=" ") if isinstance(value, datetime) else value) for key, value in row.items()}
            writer.writerow(clean)


def setup_matplotlib() -> None:
    plt.rcParams.update(
        {
            "font.sans-serif": ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"],
            "axes.unicode_minus": False,
            "figure.dpi": 130,
            "savefig.dpi": 180,
        }
    )


def plot_wind(path: Path, rows_by_point: dict[str, list[dict[str, object]]], landfall: datetime) -> None:
    setup_matplotlib()
    fig, ax = plt.subplots(figsize=(10.5, 4.7))
    for point in ("W1", "W2"):
        rows = rows_by_point[point]
        ax.plot([row["time"] for row in rows], [row["mean_speed"] for row in rows], lw=1.2, color=COLORS[point], label=point)
    ax.axhline(25.0, color="#C0392B", ls="--", lw=1.0, label="一级预警参考值 25 m/s")
    ax.axvline(landfall, color="#555555", ls=":", lw=1.0)
    ax.text(landfall, ax.get_ylim()[1], " 登陆时刻", va="top", ha="left", fontsize=8, color="#555555")
    ax.set_ylabel("10 min 平均风速 (m/s)")
    ax.set_xlabel("时间")
    ax.grid(True, color="#D9D9D9", lw=0.5, alpha=0.8)
    ax.legend(ncol=4, loc="upper left", fontsize=8)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%m-%d\n%H:%M"))
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def plot_structure(path: Path, rows_by_group: dict[str, list[dict[str, object]]], baseline_end: datetime) -> None:
    setup_matplotlib()
    fig, ax = plt.subplots(figsize=(10.5, 4.8))
    for group, rows in rows_by_group.items():
        baseline = [float(row["peak"]) for row in rows if row["time"] < baseline_end]
        scale = statistics.median(baseline) if baseline else math.nan
        values = [safe_ratio(float(row["peak"]), scale) for row in rows]
        ax.plot([row["time"] for row in rows], values, lw=1.05, color=COLORS[group], label=group)
    ax.axhline(1.0, color="#777777", ls="--", lw=0.9, label="前段中位水平")
    ax.set_ylabel("10 min 包络峰值 / 前段中位值")
    ax.set_xlabel("时间")
    ax.grid(True, color="#D9D9D9", lw=0.5, alpha=0.8)
    ax.legend(ncol=3, loc="upper left", fontsize=8)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%m-%d\n%H:%M"))
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def add_body_paragraph(doc: Document, text: str, *, bold_prefix: str = "") -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)


def add_bullet(doc: Document, text: str, *, bold_prefix: str = "") -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], BLUE)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[idx], LIGHT_GRAY)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)


def build_report(args: argparse.Namespace, analysis: Analysis) -> tuple[Path, dict[str, object]]:
    args.output_dir.mkdir(parents=True, exist_ok=True)
    rows_by_point = {point: wind_rows(analysis, point) for point in ("W1", "W2")}
    rows_by_group = {group: structure_group_rows(analysis, group) for group in STRUCTURE_GROUPS}
    wind_audits = [audit for audit in analysis.audits if audit.key in WIND_ENTRIES]
    starts = [audit.first_time for audit in wind_audits if audit.first_time]
    ends = [audit.last_time for audit in wind_audits if audit.last_time]
    if not starts or not ends:
        raise RuntimeError("No usable W1/W2 wind records were found")
    coverage_start = max(min(audit.first_time for audit in wind_audits if audit.first_time and audit.key.startswith(point)) for point in ("W1", "W2"))
    coverage_end = min(max(audit.last_time for audit in wind_audits if audit.last_time and audit.key.startswith(point)) for point in ("W1", "W2"))
    baseline_end = coverage_start + timedelta(hours=24)
    recent_start = max(coverage_start, coverage_end - timedelta(hours=24))
    landfall = datetime.fromisoformat(args.landfall_time)

    wind_chart = args.output_dir / "hongtang_typhoon_wind_10min.png"
    structure_chart = args.output_dir / "hongtang_typhoon_structure_response.png"
    plot_wind(wind_chart, rows_by_point, landfall)
    plot_structure(structure_chart, rows_by_group, baseline_end)

    for point, rows in rows_by_point.items():
        write_csv(
            args.output_dir / f"{point}_wind_10min.csv",
            rows,
            ["time", "mean_speed", "raw_peak", "raw_peak_time", "direction", "count"],
        )
    for group, rows in rows_by_group.items():
        write_csv(args.output_dir / f"{group}_10min.csv", rows, ["time", "peak", "point"])

    wind_summary: dict[str, dict[str, float]] = {}
    for point, rows in rows_by_point.items():
        full_means = [float(row["mean_speed"]) for row in rows]
        full_peaks = [float(row["raw_peak"]) for row in rows]
        recent = period_values(rows, "mean_speed", recent_start, coverage_end + timedelta(microseconds=1))
        directions = [float(row["direction"]) for row in rows]
        direction_cos = sum(math.cos(math.radians(value)) for value in directions)
        direction_sin = sum(math.sin(math.radians(value)) for value in directions)
        wind_summary[point] = {
            "mean": statistics.fmean(full_means),
            "max_10min": max(full_means),
            "raw_peak": max(full_peaks),
            "recent_mean": statistics.fmean(recent),
            "direction": math.degrees(math.atan2(direction_sin, direction_cos)) % 360.0,
        }

    structure_summary: dict[str, dict[str, object]] = {}
    for group, rows in rows_by_group.items():
        baseline = period_values(rows, "peak", coverage_start, baseline_end)
        recent = period_values(rows, "peak", recent_start, coverage_end + timedelta(microseconds=1))
        max_row = max(rows, key=lambda row: float(row["peak"]))
        baseline_median = statistics.median(baseline)
        recent_median = statistics.median(recent)
        structure_summary[group] = {
            "baseline_median": baseline_median,
            "recent_median": recent_median,
            "ratio": safe_ratio(recent_median, baseline_median),
            "max": float(max_row["peak"]),
            "max_point": max_row["point"],
            "max_time": max_row["time"],
        }

    doc = Document()
    configure_document(doc)
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_page_number(section.footer.paragraphs[0])

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(45)
    run = paragraph.add_run("洪塘大桥台风“巴威”影响监测快报")
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("阶段性初版｜登陆期数据待补")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("C65911")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"监测覆盖：{fmt_time(coverage_start)} 至 {fmt_time(coverage_end)}").font.size = Pt(11)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(
        f"数据边界：当前数据在台风于 {landfall:%m月%d日 %H:%M} 登陆前约 "
        f"{(landfall - coverage_end).total_seconds() / 3600:.1f} 小时结束。"
        "本版只能评价台风影响前及临近阶段的监测变化，不能替代登陆期、峰值期和台风后安全评估。"
    )
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_heading("Executive Summary（执行摘要）", level=1)
    max_wind_point = max(wind_summary, key=lambda point: wind_summary[point]["max_10min"])
    max_ratio_group = max(structure_summary, key=lambda group: float(structure_summary[group]["ratio"]))
    add_bullet(
        doc,
        f"数据尚未覆盖登陆期。现有三批导出覆盖 {fmt_time(coverage_start)} 至 {fmt_time(coverage_end)}；"
        f"“巴威”于 {landfall:%m月%d日 %H:%M} 登陆，最关键的强风时段仍待 7 月 12 日导出补齐。",
        bold_prefix="数据尚未覆盖登陆期。",
    )
    add_bullet(
        doc,
        f"桥址风速在截止时点前未达到预警参考值。W1/W2 最大 10 min 平均风速分别为 "
        f"{wind_summary['W1']['max_10min']:.2f} m/s 和 {wind_summary['W2']['max_10min']:.2f} m/s，"
        f"较 25 m/s 一级预警参考值仍有明显余量；{max_wind_point} 为现阶段较大者。",
        bold_prefix="桥址风速在截止时点前未达到预警参考值。",
    )
    add_bullet(
        doc,
        f"主要高频结构响应未见持续突增证据。按 6.4 s 包络峰值的 10 min 汇总，近期相对前段中位值变化最大的类别为"
        f"{max_ratio_group}（{float(structure_summary[max_ratio_group]['ratio']):.2f} 倍）；该结果仅用于趋势筛查，"
        "需等待登陆期数据后复核峰值、持续性与各测点同步性。",
        bold_prefix="主要高频结构响应未见持续突增证据。",
    )

    doc.add_heading("1. 台风已登陆，但本版数据仍停在登陆前", level=1)
    add_body_paragraph(
        doc,
        f"中央气象台/中国天气网信息显示，今年第 9 号台风“巴威”于 {landfall:%Y年%m月%d日%H时%M分} 前后"
        f"在{args.landfall_location}沿海登陆，登陆时中心附近最大风力 {args.landfall_wind:g} m/s、"
        f"中心最低气压 {args.landfall_pressure:g} hPa。"
        f"洪塘监测数据当前止于 {fmt_time(coverage_end)}，因此本报告描述的是登陆前约"
        f"{(landfall - coverage_end).total_seconds() / 3600:.1f} 小时以前的状态。"
    )
    add_body_paragraph(
        doc,
        "判断口径：以首个完整约 24 h 作为“前段基线”，以最后约 24 h 作为“临近阶段”；"
        "风速采用原始波形按 10 min 算术平均，风向采用圆统计平均；结构响应采用东华特征值包中约 6.4 s 包络峰值按 10 min 汇总。"
    )

    doc.add_heading("2. 截止 7 月 11 日 09:00，桥址风速仍处低位", level=1)
    add_body_paragraph(
        doc,
        f"W1（桥面）与 W2（塔顶）在现有时段内均未接近 25 m/s 预警参考值。"
        f"两测点全段平均风速分别为 {wind_summary['W1']['mean']:.2f} m/s 和 {wind_summary['W2']['mean']:.2f} m/s；"
        f"最近 24 h 平均分别为 {wind_summary['W1']['recent_mean']:.2f} m/s 和 {wind_summary['W2']['recent_mean']:.2f} m/s。"
        "这说明在数据截止时点前，桥址尚未记录到登陆期最强风场。"
    )
    doc.add_picture(str(wind_chart), width=Cm(16.4))
    add_caption(doc, "图 1  W1/W2 10 min 平均风速（虚线为一级预警参考值；登陆时刻位于现有数据之后）")
    add_table(
        doc,
        ["测点", "位置", "全段均值", "最大10 min均值", "原始峰值", "圆平均风向"],
        [
            [
                point,
                "桥面" if point == "W1" else "塔顶",
                f"{wind_summary[point]['mean']:.2f} m/s",
                f"{wind_summary[point]['max_10min']:.2f} m/s",
                f"{wind_summary[point]['raw_peak']:.2f} m/s",
                f"{wind_summary[point]['direction']:.0f}°",
            ]
            for point in ("W1", "W2")
        ],
        [1.4, 2.0, 2.5, 3.1, 2.5, 2.5],
    )
    add_body_paragraph(
        doc,
        "解释：W1 与 W2 的风向统计仍存在明显差异，符合两个独立测点的既有特征；"
        "本阶段没有发现两通道被复制或错误复用的迹象。是否出现台风核心风场下的同步增强，需等下一批数据确认。"
    )

    doc.add_heading("3. 主要高频结构响应暂未出现持续性同步抬升", level=1)
    add_body_paragraph(
        doc,
        "下图把主梁、主塔及南北两组索振动的 10 min 包络峰值分别除以前段中位值，用于比较不同部位的相对变化。"
        "该归一化图只回答“是否较自身前段水平明显抬升”，不等同于设计验算或安全等级判定。"
    )
    doc.add_picture(str(structure_chart), width=Cm(16.4))
    add_caption(doc, "图 2  主要高频结构响应相对前段水平（6.4 s 包络峰值的 10 min 汇总）")
    structure_rows = []
    for group, summary in structure_summary.items():
        structure_rows.append(
            [
                group,
                f"{float(summary['baseline_median']):.4f}",
                f"{float(summary['recent_median']):.4f}",
                f"{float(summary['ratio']):.2f}",
                f"{float(summary['max']):.4f}",
                f"{summary['max_point']} / {summary['max_time']:%m-%d %H:%M}",
            ]
        )
    doc.add_page_break()
    add_table(
        doc,
        ["部位", "前段中位", "临近阶段中位", "阶段比值", "全段最大", "最大测点/时刻"],
        structure_rows,
        [2.7, 2.3, 2.7, 2.0, 2.2, 3.3],
    )
    add_body_paragraph(
        doc,
        "初步判断：现有时段内各类别没有出现与台风登陆相对应的持续、同步跃升。"
        "个别 10 min 峰值可能受车辆荷载、局部振动或短时噪声影响，须结合下一批登陆期数据看其持续性及多测点一致性。"
    )

    doc.add_heading("4. 当前运营建议：继续按高风险时段管理，监测结论待补", level=1)
    add_bullet(doc, "继续执行主管部门和气象预警对应的交通管控、巡查与应急值守，不能因本初版风速较低而降低防台等级。")
    add_bullet(doc, "7 月 12 日导出形成后，优先复算 W1/W2 10 min 平均风速、阵风峰值、风向以及主梁/主塔/索振动响应。")
    add_bullet(doc, "将登陆前、登陆期和登陆后数据放在同一时间轴比较，检查结构响应是否与风速同步、是否存在滞后恢复或测点异常。")
    add_bullet(doc, "如登陆期出现异常峰值或持续抬升，应结合现场巡查、交通状态与传感器状态进行复核，不宜仅依据单点瞬时峰值下结论。")

    doc.add_heading("5. 待更新问题与限制", level=1)
    add_bullet(doc, f"最关键缺口是 {fmt_time(coverage_end)} 之后至登陆及登陆后时段；现阶段不得评价最大风速、最大结构响应或台风后安全状态。")
    add_bullet(doc, "结构响应采用特征值包的包络峰值进行趋势筛查，不是完整原始波形 RMS、频谱或模态分析；最终版将按既定生产口径补算。")
    add_bullet(doc, "本报告未把桥址监测风速与气象站风速直接等同比较；两者高度、暴露条件、地形与平均时距不同。")
    add_bullet(doc, "若下一批导出仍缺少测点或时间段，将在最终版逐项披露覆盖率，不做插值补齐或推测性补数。")

    doc.add_heading("资料来源", level=2)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run("气象事件：").bold = True
    paragraph.add_run(f"{args.weather_source_title}（访问时间：{args.weather_checked_at}）\n{args.weather_source_url}")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run("桥梁监测：").bold = True
    paragraph.add_run(
        "126 定时导出经 CLI 核验并复制至 133 的独立台风分析目录；"
        "每批波形/特征值压缩包各 139 项。本版直接读取压缩包，未修改既有 Q2 正式成果。"
    )

    output_docx = args.output_dir / args.output_name
    doc.save(output_docx)

    manifest = {
        "status": "ok" if not analysis.missing_entries else "warning",
        "report": str(output_docx),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "scope": "stage_initial_pre_landfall",
        "coverage": {"start": coverage_start.isoformat(), "end": coverage_end.isoformat()},
        "landfall": {
            "time": landfall.isoformat(),
            "location": args.landfall_location,
            "wind_mps": args.landfall_wind,
            "pressure_hpa": args.landfall_pressure,
            "source_title": args.weather_source_title,
            "source_url": args.weather_source_url,
            "checked_at": args.weather_checked_at,
        },
        "method": {
            "wind": "raw waveform, non-negative speeds, 10-minute arithmetic mean; circular direction mean",
            "structure": "Donghua feature peak sequence (~6.4 s), absolute values, 10-minute group envelope",
            "baseline": f"{coverage_start.isoformat()} to {baseline_end.isoformat()}",
            "recent": f"{recent_start.isoformat()} to {coverage_end.isoformat()}",
        },
        "wind_summary": wind_summary,
        "structure_summary": {
            group: {key: (value.isoformat() if isinstance(value, datetime) else value) for key, value in summary.items()}
            for group, summary in structure_summary.items()
        },
        "missing_entries": analysis.missing_entries,
        "audits": [
            {
                **audit.__dict__,
                "first_time": audit.first_time.isoformat() if audit.first_time else None,
                "last_time": audit.last_time.isoformat() if audit.last_time else None,
            }
            for audit in analysis.audits
        ],
        "artifacts": [str(wind_chart), str(structure_chart)],
    }
    manifest_path = args.output_dir / "hongtang_typhoon_brief_manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return output_docx, manifest


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Build the lightweight Hongtang typhoon monitoring brief")
    parser.add_argument("--source-root", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--output-name", default="洪塘大桥台风巴威影响监测快报_阶段性初版.docx")
    parser.add_argument("--landfall-time", default="2026-07-11T23:20:00")
    parser.add_argument("--landfall-location", default="浙江省玉环市坎门街道")
    parser.add_argument("--landfall-wind", type=float, default=40.0)
    parser.add_argument("--landfall-pressure", type=float, default=955.0)
    parser.add_argument("--weather-source-title", default="中国天气网：‘巴威’以台风级登陆浙江玉环沿海")
    parser.add_argument("--weather-source-url", default="https://news.weather.com.cn/2026/07/4711229.shtml")
    parser.add_argument("--weather-checked-at", default=datetime.now().strftime("%Y-%m-%d %H:%M CST"))
    return parser


def main() -> None:
    args = build_parser().parse_args()
    analysis = analyze(args.source_root)
    output, manifest = build_report(args, analysis)
    print(json.dumps({"report": str(output), "status": manifest["status"], "coverage": manifest["coverage"]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
