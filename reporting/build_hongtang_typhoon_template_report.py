from __future__ import annotations

import argparse
import copy
import json
import math
import statistics
from datetime import datetime, timedelta
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from build_hongtang_typhoon_brief import (
    Analysis,
    STRUCTURE_GROUPS,
    analyze,
    structure_group_rows,
    wind_rows,
)


BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
ORANGE = "E67E22"
COLORS = {
    "W1": "#1F77B4",
    "W2": "#E67E22",
    "主梁加速度": "#1F77B4",
    "主塔加速度": "#9467BD",
    "南侧索振动": "#2CA02C",
    "北侧索振动": "#D62728",
}
ALARM_LEVELS = (25.0, 29.92, 37.4)


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.findall(".//" + qn("w:t")))


def body_index_by_text(doc: Document, text: str, *, start: int = 0) -> int:
    body = doc._element.body
    for index, element in enumerate(body.iterchildren()):
        if index >= start and element.tag == qn("w:p") and paragraph_text(element).strip() == text.strip():
            return index
    raise ValueError(f"paragraph not found: {text}")


def remove_body_range(doc: Document, start: int, end: int) -> None:
    body = doc._element.body
    children = list(body.iterchildren())
    for element in children[start:end]:
        body.remove(element)


def remove_empty_section_break_before(doc: Document, body_index: int) -> None:
    """Drop a template-only empty section paragraph immediately before a chapter.

    The Q2 template places a next-page section break in an otherwise empty
    paragraph after the alarm-threshold table.  Once the much shorter typhoon
    chapter is substituted, Word can push that empty paragraph onto its own
    page and then start the following section on another page.  The result is
    a completely blank page.  The adjacent sections use the same page setup,
    header and footer, so the redundant section can safely be merged; the
    chapter itself receives an explicit page-break-before below.
    """
    body = doc._element.body
    children = list(body.iterchildren())
    if body_index <= 0:
        return
    previous = children[body_index - 1]
    if (
        previous.tag == qn("w:p")
        and not paragraph_text(previous).strip()
        and previous.find(".//" + qn("w:sectPr")) is not None
    ):
        body.remove(previous)


def add_numbered_heading(doc: Document, text: str, *, level: int):
    """Add a heading while preserving the template's direct multilevel list."""
    paragraph = doc.add_heading(text, level=level)
    source_num_pr = None
    for candidate in doc.paragraphs:
        if candidate is paragraph or candidate.style is None:
            continue
        if candidate.style.name != f"Heading {level}":
            continue
        num_pr = candidate._p.pPr.numPr if candidate._p.pPr is not None else None
        if num_pr is not None:
            source_num_pr = num_pr
            break
    if source_num_pr is not None:
        p_pr = paragraph._p.get_or_add_pPr()
        existing = p_pr.numPr
        if existing is not None:
            p_pr.remove(existing)
        p_pr.append(copy.deepcopy(source_num_pr))
    return paragraph


def replace_all_text(doc: Document, old: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        if old not in paragraph.text:
            continue
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                break
        else:
            paragraph.text = paragraph.text.replace(old, new)


def add_toc_field_before(doc: Document, before_element) -> None:
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    paragraph.append(p_pr)
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开文档时更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run.extend([begin, instr, separate, placeholder, end])
    paragraph.append(run)
    before_element.addprevious(paragraph)
    page_break_paragraph = OxmlElement("w:p")
    page_break_run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    page_break_run.append(page_break)
    page_break_paragraph.append(page_break_run)
    before_element.addprevious(page_break_paragraph)


def prepare_template(template: Path, output: Path, start: datetime, end: datetime) -> Document:
    doc = Document(template)
    # Remove the stale static Q2 TOC entries, then insert a live TOC field.
    toc_title = body_index_by_text(doc, "目  录")
    chapter_one = body_index_by_text(doc, "监测概况", start=toc_title + 1)
    body = doc._element.body
    chapter_one_element = list(body.iterchildren())[chapter_one]
    remove_body_range(doc, toc_title + 1, chapter_one)
    chapter_one = body_index_by_text(doc, "监测概况", start=toc_title + 1)
    chapter_one_element = list(doc._element.body.iterchildren())[chapter_one]
    add_toc_field_before(doc, chapter_one_element)

    # Remove Q2-only system operation and maintenance narrative/tables.
    system_status = body_index_by_text(doc, "健康监测系统运行状况")
    chapter_two = body_index_by_text(doc, "监测项目及内容", start=system_status + 1)
    remove_body_range(doc, system_status, chapter_two)

    # Keep the original chapter-4 heading so its direct numbering properties remain intact,
    # but remove every old Q2 result block below it.
    old_results = body_index_by_text(doc, "监测结果")
    remove_empty_section_break_before(doc, old_results)
    old_results = body_index_by_text(doc, "监测结果")
    children = list(doc._element.body.iterchildren())
    sect_index = next(i for i, element in enumerate(children) if element.tag == qn("w:sectPr"))
    remove_body_range(doc, old_results + 1, sect_index)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "监测结果":
            paragraph.paragraph_format.page_break_before = True
            break

    period = f"（监测时间：{start:%Y年%m月%d日%H:%M}~{end:%Y年%m月%d日%H:%M}）"
    replace_all_text(doc, "第二季度报告", "台风“巴威”影响监测专题报告")
    replace_all_text(doc, "（监测时间：2026年4月1日~2026年6月30日）", period)
    replace_all_text(doc, "报告编号：BG02FQJC2600002-J2", "报告编号：BG02FQJC2600002-J2（台风专题）")
    replace_all_text(doc, "报告日期：2026年07月12日", f"报告日期：{datetime.now():%Y年%m月%d日}")
    replace_all_text(doc, "监测结果", "台风影响监测结果")
    return doc


def filter_rows(rows: list[dict[str, object]], start: datetime, end: datetime) -> list[dict[str, object]]:
    return [row for row in rows if start <= row["time"] < end]


def circular_mean(values: list[float]) -> float:
    if not values:
        return math.nan
    return math.degrees(
        math.atan2(
            sum(math.sin(math.radians(value)) for value in values),
            sum(math.cos(math.radians(value)) for value in values),
        )
    ) % 360.0


def point_wind_summary(rows: list[dict[str, object]], landfall: datetime) -> dict[str, object]:
    if not rows:
        raise RuntimeError("wind rows are empty")
    raw_row = max(rows, key=lambda row: float(row["raw_peak"]))
    mean_row = max(rows, key=lambda row: float(row["mean_speed"]))
    pre = [row for row in rows if row["time"] < landfall]
    post = [row for row in rows if row["time"] >= landfall]
    return {
        "mean": statistics.fmean(float(row["mean_speed"]) for row in rows),
        "raw_max": float(raw_row["raw_peak"]),
        "raw_max_time": raw_row.get("raw_peak_time") or raw_row["time"],
        "raw_max_direction": float(raw_row["direction"]),
        "max_10min": float(mean_row["mean_speed"]),
        "max_10min_time": mean_row["time"],
        "max_10min_direction": float(mean_row["direction"]),
        "direction": circular_mean([float(row["direction"]) for row in rows]),
        "pre_mean": statistics.fmean(float(row["mean_speed"]) for row in pre) if pre else math.nan,
        "pre_max_10min": max((float(row["mean_speed"]) for row in pre), default=math.nan),
        "post_mean": statistics.fmean(float(row["mean_speed"]) for row in post) if post else math.nan,
        "post_max_10min": max((float(row["mean_speed"]) for row in post), default=math.nan),
        "bins": len(rows),
    }


def group_summary(rows: list[dict[str, object]], landfall: datetime) -> dict[str, object]:
    pre = [row for row in rows if row["time"] < landfall]
    post = [row for row in rows if row["time"] >= landfall]
    maximum = max(rows, key=lambda row: float(row["peak"]))
    pre_median = statistics.median(float(row["peak"]) for row in pre) if pre else math.nan
    post_median = statistics.median(float(row["peak"]) for row in post) if post else math.nan
    return {
        "pre_median": pre_median,
        "pre_max": max((float(row["peak"]) for row in pre), default=math.nan),
        "post_median": post_median,
        "post_max": max((float(row["peak"]) for row in post), default=math.nan),
        "median_ratio": post_median / pre_median if pre_median > 1e-12 and math.isfinite(post_median) else math.nan,
        "maximum": float(maximum["peak"]),
        "maximum_point": maximum["point"],
        "maximum_time": maximum["time"],
        "bins": len(rows),
    }


def wind_alarm_status(max_10min: float) -> str:
    if max_10min >= ALARM_LEVELS[2]:
        return "达到三级"
    if max_10min >= ALARM_LEVELS[1]:
        return "达到二级"
    if max_10min >= ALARM_LEVELS[0]:
        return "达到一级"
    return "未达一级"


def setup_plotting() -> None:
    plt.rcParams.update(
        {
            "font.sans-serif": ["Microsoft YaHei", "SimHei", "DejaVu Sans"],
            "axes.unicode_minus": False,
            "figure.dpi": 130,
            "savefig.dpi": 200,
            "axes.edgecolor": "#333333",
            "axes.labelcolor": "#333333",
            "text.color": "#333333",
        }
    )


def plot_wind_speed(path: Path, rows: dict[str, list[dict[str, object]]], landfall: datetime) -> None:
    setup_plotting()
    fig, ax = plt.subplots(figsize=(10.6, 4.9))
    max_value = 0.0
    for point in ("W1", "W2"):
        values = [float(row["mean_speed"]) for row in rows[point]]
        max_value = max(max_value, max(values))
        ax.plot([row["time"] for row in rows[point]], values, lw=1.35, color=COLORS[point], label=f"{point} 10 min平均")
        raw_row = max(rows[point], key=lambda row: float(row["raw_peak"]))
        ax.scatter(
            [raw_row.get("raw_peak_time") or raw_row["time"]],
            [float(raw_row["raw_peak"])],
            s=38,
            marker="o" if point == "W1" else "s",
            facecolors="none",
            edgecolors=COLORS[point],
            linewidths=1.2,
            label=f"{point} 原始最大值",
            zorder=4,
        )
        max_value = max(max_value, float(raw_row["raw_peak"]))
    ax.axvline(landfall, color="#555555", ls="--", lw=1.0, label="台风登陆")
    ax.set_ylim(0, max(8.0, max_value * 1.18))
    ax.set_ylabel("风速 (m/s)")
    ax.set_xlabel("时间")
    ax.grid(True, color="#D9D9D9", lw=0.5)
    ax.legend(ncol=3, fontsize=8, loc="upper left")
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%m-%d\n%H:%M"))
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def plot_wind_maxima(path: Path, summaries: dict[str, dict[str, object]]) -> None:
    setup_plotting()
    fig, ax = plt.subplots(figsize=(10.6, 4.4))
    labels = ["W1 桥面", "W2 塔顶"]
    y = [0, 1]
    raw = [float(summaries[point]["raw_max"]) for point in ("W1", "W2")]
    mean10 = [float(summaries[point]["max_10min"]) for point in ("W1", "W2")]
    ax.barh([value + 0.18 for value in y], raw, height=0.30, color="#1F77B4", label="原始最大值")
    ax.barh([value - 0.18 for value in y], mean10, height=0.30, color="#E5A04B", label="最大10 min平均")
    for level, label in zip(ALARM_LEVELS, ("一级25", "二级29.92", "三级37.4")):
        ax.axvline(level, color="#666666", ls="--" if level == 25 else ":", lw=0.9)
        ax.text(level + 0.2, 1.47, label, rotation=90, va="top", fontsize=7, color="#555555")
    for idx, value in enumerate(raw):
        ax.text(value + 0.25, idx + 0.18, f"{value:.2f}", va="center", fontsize=8)
    for idx, value in enumerate(mean10):
        ax.text(value + 0.25, idx - 0.18, f"{value:.2f}", va="center", fontsize=8)
    ax.set_yticks(y, labels)
    ax.set_xlim(0, 40.5)
    ax.set_xlabel("风速 (m/s)")
    ax.grid(True, axis="x", color="#D9D9D9", lw=0.5)
    ax.legend(ncol=2, loc="lower right", fontsize=8)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def plot_wind_direction(path: Path, rows: dict[str, list[dict[str, object]]], landfall: datetime) -> None:
    setup_plotting()
    fig, ax = plt.subplots(figsize=(10.6, 4.2))
    for point in ("W1", "W2"):
        ax.plot(
            [row["time"] for row in rows[point]],
            [row["direction"] for row in rows[point]],
            lw=0.95,
            color=COLORS[point],
            label=point,
        )
    ax.axvline(landfall, color="#555555", ls="--", lw=1.0, label="台风登陆")
    ax.set_ylim(0, 360)
    ax.set_yticks([0, 90, 180, 270, 360], ["0°/N", "90°/E", "180°/S", "270°/W", "360°/N"])
    ax.set_ylabel("10 min 圆平均风向")
    ax.set_xlabel("时间")
    ax.grid(True, color="#D9D9D9", lw=0.5)
    ax.legend(ncol=3, loc="upper left", fontsize=8)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%m-%d\n%H:%M"))
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def plot_structure(path: Path, rows: dict[str, list[dict[str, object]]], landfall: datetime) -> None:
    setup_plotting()
    fig, ax = plt.subplots(figsize=(10.6, 4.8))
    for group, items in rows.items():
        pre = [float(row["peak"]) for row in items if row["time"] < landfall]
        scale = statistics.median(pre) if pre else math.nan
        values = [float(row["peak"]) / scale if scale > 1e-12 else math.nan for row in items]
        ax.plot([row["time"] for row in items], values, lw=1.0, color=COLORS[group], label=group)
    ax.axhline(1.0, color="#777777", ls=":", lw=0.9, label="登陆前中位水平")
    ax.axvline(landfall, color="#555555", ls="--", lw=1.0, label="台风登陆")
    ax.set_ylabel("10 min包络峰值 / 登陆前中位值")
    ax.set_xlabel("时间")
    ax.grid(True, color="#D9D9D9", lw=0.5)
    ax.legend(ncol=3, fontsize=8, loc="upper left")
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%m-%d\n%H:%M"))
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = "000000") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def body_paragraph(doc: Document, text: str, *, bold_prefix: str = ""):
    style = "洪塘大桥月报正文" if "洪塘大桥月报正文" in [s.name for s in doc.styles] else "Normal"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Fließtext" if "Fließtext" in [s.name for s in doc.styles] else "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    add_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], BLUE)
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value)
            if row_index % 2:
                set_cell_shading(cells[index], LIGHT_GRAY)


def add_image(doc: Document, path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(15.8))
    add_caption(doc, caption)


def update_front_summary(
    doc: Document,
    start: datetime,
    end: datetime,
    wind_summary: dict[str, dict[str, object]],
    structure_summary: dict[str, dict[str, object]],
    alarm_text: str,
) -> None:
    if len(doc.tables) < 3:
        raise RuntimeError("period template front-summary tables were not found")
    metadata = doc.tables[1]
    metadata.cell(1, 4).text = f"{start:%Y年%m月%d日%H:%M}~{end:%Y年%m月%d日%H:%M}"
    max_ratio_group = max(structure_summary, key=lambda group: float(structure_summary[group]["median_ratio"]))
    results = (
        "本报告对台风‘巴威’登陆前24小时至最新实测时点的洪塘大桥监测数据进行专题分析，主要结论如下：\n"
        "1、风向风速监测\n"
        f"W1桥面、W2塔顶原始风速最大值分别为{float(wind_summary['W1']['raw_max']):.2f}m/s和"
        f"{float(wind_summary['W2']['raw_max']):.2f}m/s；最大10min平均风速分别为"
        f"{float(wind_summary['W1']['max_10min']):.2f}m/s和{float(wind_summary['W2']['max_10min']):.2f}m/s。"
        f"按10min平均风速阈值判断，{alarm_text}。\n"
        "2、主要结构振动监测\n"
        f"登陆后相对登陆前中位水平变化最大的类别为{max_ratio_group}，后/前比为"
        f"{float(structure_summary[max_ratio_group]['median_ratio']):.2f}。"
        "各类结构响应按持续性和多测点同步性进行筛查，单个短时尖峰不直接判定为结构异常。\n"
        "3、数据范围\n"
        f"实际监测覆盖为{start:%Y年%m月%d日%H:%M}至{end:%Y年%m月%d日%H:%M}，"
        "报告不对该时点之后的数据插值或推测。"
    )
    metadata.cell(3, 2).text = results

    continuation = doc.tables[2]
    continuation.cell(0, 1).text = (
        "4、综合判断\n"
        "本专题同时报告原始风速最大值和最大10min平均风速，并给出发生时间、同期风向及登陆前后阶段比较。"
        "主要结构振动响应采用特征包络峰值进行趋势筛查；若出现持续抬升或多测点同步异常，应结合原始波形、频谱、交通状态和现场巡查复核。"
    )
    continuation.cell(1, 1).text = (
        "针对本次台风影响，建议如下：\n"
        "1、运营管理以气象预警、交通管控指令及现场巡查为准，不因单项监测指标未达阈值而降低防台等级。\n"
        "2、持续关注风速峰值、10min平均风速及主梁、主塔、吊索振动的同步变化。\n"
        "3、后续如取得更新数据，按相同口径复算最大值、登陆前后变化和综合结论。\n"
        "（本栏以下空白）"
    )


def build(args: argparse.Namespace) -> tuple[Path, dict[str, object]]:
    args.output_dir.mkdir(parents=True, exist_ok=True)
    export_days = {part.strip() for part in args.export_days.split(",") if part.strip()}
    analysis: Analysis = analyze(args.source_root, day_filter=export_days)
    landfall = datetime.fromisoformat(args.landfall_time)
    requested_start = datetime.fromisoformat(args.window_start)
    wind_all = {point: wind_rows(analysis, point) for point in ("W1", "W2")}
    speed_audits = [audit for audit in analysis.audits if audit.key in {"W1_speed", "W2_speed"} and audit.last_time]
    available_end = min(
        max(audit.last_time for audit in speed_audits if audit.key == key)
        for key in ("W1_speed", "W2_speed")
    )
    requested_end = datetime.fromisoformat(args.window_end) if args.window_end else available_end
    end = min(requested_end, available_end)
    if end <= requested_start:
        raise RuntimeError(f"usable end {end} is not later than requested start {requested_start}")
    complete_bin_end = end.replace(minute=(end.minute // 10) * 10, second=0, microsecond=0)
    wind = {point: filter_rows(rows, requested_start, complete_bin_end) for point, rows in wind_all.items()}
    structure = {
        group: filter_rows(structure_group_rows(analysis, group), requested_start, complete_bin_end)
        for group in STRUCTURE_GROUPS
    }
    missing = [point for point, rows in wind.items() if not rows] + [group for group, rows in structure.items() if not rows]
    if missing:
        raise RuntimeError(f"missing report series: {missing}")

    expected_bins = int((complete_bin_end - requested_start).total_seconds() // 600)
    speed_audit_rows = sum(audit.rows for audit in speed_audits)
    speed_rejected = sum(audit.rejected_rows for audit in speed_audits)
    direction_audits = [audit for audit in analysis.audits if audit.key in {"W1_direction", "W2_direction"}]
    direction_rows = sum(audit.rows for audit in direction_audits)
    direction_rejected = sum(audit.rejected_rows for audit in direction_audits)
    direction_reject_rate = direction_rejected / max(1, direction_rows + direction_rejected)

    wind_summary = {point: point_wind_summary(rows, landfall) for point, rows in wind.items()}
    structure_summary = {group: group_summary(rows, landfall) for group, rows in structure.items()}

    wind_chart = args.output_dir / "wind_speed_window.png"
    wind_max_chart = args.output_dir / "wind_maximum_comparison.png"
    direction_chart = args.output_dir / "wind_direction_window.png"
    structure_chart = args.output_dir / "structure_response_window.png"
    plot_wind_speed(wind_chart, wind, landfall)
    plot_wind_maxima(wind_max_chart, wind_summary)
    plot_wind_direction(direction_chart, wind, landfall)
    plot_structure(structure_chart, structure, landfall)

    doc = prepare_template(args.template, args.output_dir / args.output_name, requested_start, end)
    add_numbered_heading(doc, "事件概况与数据覆盖", level=2)
    body_paragraph(
        doc,
        f"今年第9号台风“巴威”于{landfall:%Y年%m月%d日%H时%M分}前后在浙江省玉环市坎门街道沿海登陆，"
        "登陆时中心附近最大风力13级（40m/s），中心最低气压955hPa。"
        f"本专题报告按确认口径选取登陆前24小时起至最新可用实测时点，即{requested_start:%Y年%m月%d日%H:%M}"
        f"至{end:%Y年%m月%d日%H:%M}，共{(end-requested_start).total_seconds()/3600:.1f}小时。"
    )
    body_paragraph(
        doc,
        "数据范围说明：风速和风向采用东华原始波形按10min统计；主梁、主塔及吊索振动采用东华特征值包约6.4s包络峰值按10min汇总。"
        "所有结论仅针对实际覆盖时段，不对最新实测时点之后的数据作插值或推测。"
    )
    add_table(
        doc,
        "表 4-1 台风专题监测数据覆盖情况",
        ["数据类别", "测点范围", "开始时间", "结束时间", "10min时段数", "完整性"],
        [
            ["风向风速", "W1、W2", f"{requested_start:%m-%d %H:%M}", f"{end:%m-%d %H:%M}", str(min(len(wind['W1']), len(wind['W2']))), "完整" if min(len(wind['W1']), len(wind['W2'])) >= expected_bins else "部分"],
            ["主梁/主塔振动", "A1至A10", f"{requested_start:%m-%d %H:%M}", f"{end:%m-%d %H:%M}", str(min(len(structure['主梁加速度']), len(structure['主塔加速度']))), "完整" if min(len(structure['主梁加速度']), len(structure['主塔加速度'])) >= expected_bins else "部分"],
            ["吊索振动", "CS1至CS12、CX1至CX12", f"{requested_start:%m-%d %H:%M}", f"{end:%m-%d %H:%M}", str(min(len(structure['南侧索振动']), len(structure['北侧索振动']))), "完整" if min(len(structure['南侧索振动']), len(structure['北侧索振动'])) >= expected_bins else "部分"],
        ],
    )
    body_paragraph(
        doc,
        f"数据质量核验：两批源文件中W1/W2风速原始记录共读取{speed_audit_rows:,}条，负风速或非有限值剔除{speed_rejected:,}条；"
        f"风向超出0°~360°有效域的记录占{direction_reject_rate*100:.2f}%，已按既定口径剔除。"
        f"报告窗口应有{expected_bins}个完整10min时段，实际风速、主梁/主塔和吊索汇总时段数见表4-1。"
    )

    add_numbered_heading(doc, "风速最大值及登陆前后变化", level=2)
    dominant = max(wind_summary, key=lambda point: float(wind_summary[point]["raw_max"]))
    alarm_text = "、".join(
        f"{point}{wind_alarm_status(float(wind_summary[point]['max_10min']))}"
        for point in ("W1", "W2")
    )
    update_front_summary(doc, requested_start, end, wind_summary, structure_summary, alarm_text)
    body_paragraph(
        doc,
        f"结论先行：本时段原始风速最大值出现在{dominant}测点，为{float(wind_summary[dominant]['raw_max']):.2f}m/s；"
        f"W1、W2最大10min平均风速分别为{float(wind_summary['W1']['max_10min']):.2f}m/s和"
        f"{float(wind_summary['W2']['max_10min']):.2f}m/s；按10min平均风速阈值判断，{alarm_text}。"
        "原始最大值反映短时阵风，10min平均值更适合判断持续风荷载，两者不能混用。",
        bold_prefix="结论先行：",
    )
    add_image(doc, wind_chart, "图 4-1 W1/W2风速时程及原始最大值")
    add_image(doc, wind_max_chart, "图 4-2 W1/W2风速最大值与预警参考值比较")
    wind_table_rows = []
    for point, location in (("W1", "桥面"), ("W2", "塔顶")):
        summary = wind_summary[point]
        wind_table_rows.append(
            [
                point,
                location,
                f"{float(summary['raw_max']):.2f}",
                f"{summary['raw_max_time']:%m-%d %H:%M:%S}",
                f"{float(summary['max_10min']):.2f}",
                f"{summary['max_10min_time']:%m-%d %H:%M}",
                f"{float(summary['max_10min_direction']):.0f}°",
                wind_alarm_status(float(summary["max_10min"])),
            ]
        )
    add_table(
        doc,
        "表 4-2 W1/W2风速最大值分析",
        ["测点", "位置", "原始最大值(m/s)", "发生时间", "最大10min均值(m/s)", "发生时段", "同期风向", "预警判断"],
        wind_table_rows,
    )
    note = doc.add_paragraph(style="Fließtext" if "Fließtext" in [s.name for s in doc.styles] else "Normal")
    note.add_run("注：同期风向为最大10min平均风速所在时段的圆平均风向；预警判断按10min平均风速执行，原始最大值不直接套用该阈值。")
    phase_rows = []
    for point in ("W1", "W2"):
        summary = wind_summary[point]
        phase_rows.append(
            [
                point,
                f"{float(summary['pre_mean']):.2f}",
                f"{float(summary['pre_max_10min']):.2f}",
                f"{float(summary['post_mean']):.2f}",
                f"{float(summary['post_max_10min']):.2f}",
                f"{float(summary['post_mean'])/float(summary['pre_mean']):.2f}",
            ]
        )
    add_table(
        doc,
        "表 4-3 台风登陆前后风速阶段比较",
        ["测点", "登陆前均值", "登陆前最大10min", "登陆后均值", "登陆后最大10min", "登陆后/前均值比"],
        phase_rows,
    )
    body_paragraph(
        doc,
        f"W1登陆前后平均风速分别为{float(wind_summary['W1']['pre_mean']):.2f}m/s和{float(wind_summary['W1']['post_mean']):.2f}m/s；"
        f"W2分别为{float(wind_summary['W2']['pre_mean']):.2f}m/s和{float(wind_summary['W2']['post_mean']):.2f}m/s。"
        "桥面与塔顶测点处于不同绕流环境，风速水平和变化幅度不要求相同；应结合风向、构件遮挡和桥位地形解释，不宜套用单一高度增风关系。"
    )
    add_image(doc, direction_chart, "图 4-3 W1/W2 10min圆平均风向时程")

    add_numbered_heading(doc, "主要结构振动响应", level=2)
    body_paragraph(
        doc,
        "主梁、主塔及南北两组吊索振动采用各类别10min包络峰值除以登陆前中位值进行比较。"
        "比值为1表示与登陆前典型水平相当；短时尖峰可能同时受车辆荷载、局部振动和传感器噪声影响，需关注持续性及多测点同步性。"
    )
    add_image(doc, structure_chart, "图 4-4 主要结构振动响应相对登陆前水平")
    structure_rows = []
    for group, summary in structure_summary.items():
        structure_rows.append(
            [
                group,
                f"{float(summary['pre_median']):.4f}",
                f"{float(summary['post_median']):.4f}",
                f"{float(summary['median_ratio']):.2f}",
                f"{float(summary['maximum']):.4f}",
                f"{summary['maximum_point']} / {summary['maximum_time']:%m-%d %H:%M}",
            ]
        )
    add_table(
        doc,
        "表 4-4 主要结构振动响应登陆前后比较",
        ["部位", "登陆前中位", "登陆后中位", "后/前比值", "全段最大", "最大测点/时刻"],
        structure_rows,
    )
    max_ratio_group = max(structure_summary, key=lambda group: float(structure_summary[group]["median_ratio"]))
    body_paragraph(
        doc,
        f"登陆后相对登陆前中位水平变化最大的类别为{max_ratio_group}，比值为"
        f"{float(structure_summary[max_ratio_group]['median_ratio']):.2f}。"
        "本表用于台风窗口内的趋势筛查，不等同于设计承载能力验算；若出现明显持续抬升，应结合原始波形、频谱、交通状态及现场巡查进一步复核。"
    )

    add_numbered_heading(doc, "综合判断与运营建议", level=2)
    conclusions = [
        f"本专题窗口内按最大10min平均风速阈值判断：{alarm_text}；原始最大值作为短时阵风指标单独报告。",
        "风速最大值必须区分短时原始峰值和10min平均值；本报告已同时给出发生时间、同期风向和登陆前后阶段比较。",
        "主要结构振动响应需按持续性和多测点同步性判断，单个10min尖峰不足以直接认定结构异常。",
        "运营管理仍应以气象预警、交通管控指令和现场巡查为准，不能仅因监测风速未达阈值而降低防台等级。",
    ]
    for text in conclusions:
        paragraph = body_paragraph(doc, f"• {text}")
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.5)

    add_numbered_heading(doc, "数据限制与后续复核", level=2)
    body_paragraph(
        doc,
        f"本报告最新实测数据截至{end:%Y年%m月%d日%H:%M}。该时点之后的数据尚未纳入，"
        "不得据此评价更晚时段的最大响应。若后续导出补充新的台风影响数据，应按同一口径更新图表、最大值及结论。"
    )
    body_paragraph(
        doc,
        "资料来源：桥梁监测数据来自洪塘大桥东华定时导出；台风登陆时间、地点和强度来自中央气象台/中国天气网公开信息。"
    )
    tail = doc.add_paragraph(style="Fließtext" if "Fließtext" in [s.name for s in doc.styles] else "Normal")
    tail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tail.add_run("(以下无正文)")

    output = args.output_dir / args.output_name
    doc.save(output)
    manifest = {
        "status": "ok" if not analysis.missing_entries else "warning",
        "report": str(output),
        "template": str(args.template),
        "window": {"start": requested_start.isoformat(), "end": end.isoformat(), "landfall": landfall.isoformat()},
        "export_days": sorted(export_days),
        "wind_summary": {
            point: {key: (value.isoformat() if isinstance(value, datetime) else value) for key, value in summary.items()}
            for point, summary in wind_summary.items()
        },
        "structure_summary": {
            group: {key: (value.isoformat() if isinstance(value, datetime) else value) for key, value in summary.items()}
            for group, summary in structure_summary.items()
        },
        "missing_entries": analysis.missing_entries,
        "source_entry_audits": len(analysis.audits),
        "quality": {
            "expected_complete_10min_bins": expected_bins,
            "speed_rows": speed_audit_rows,
            "speed_rejected_rows": speed_rejected,
            "direction_rows": direction_rows,
            "direction_rejected_rows": direction_rejected,
            "direction_rejected_rate": direction_reject_rate,
        },
        "charts": [str(wind_chart), str(wind_max_chart), str(direction_chart), str(structure_chart)],
        "generated_at": datetime.now().isoformat(timespec="seconds"),
    }
    (args.output_dir / "hongtang_typhoon_template_report_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return output, manifest


def parser() -> argparse.ArgumentParser:
    value = argparse.ArgumentParser(description="Build Hongtang typhoon report with the Q2 period-report template")
    value.add_argument("--template", type=Path, required=True)
    value.add_argument("--source-root", type=Path, required=True)
    value.add_argument("--output-dir", type=Path, required=True)
    value.add_argument("--output-name", default="洪塘大桥台风巴威影响监测专题报告_季报模板版.docx")
    value.add_argument("--window-start", default="2026-07-10T23:20:00")
    value.add_argument("--window-end", default="")
    value.add_argument("--landfall-time", default="2026-07-11T23:20:00")
    value.add_argument("--export-days", default="2026-07-11,2026-07-12")
    return value


def main() -> None:
    args = parser().parse_args()
    output, manifest = build(args)
    print(json.dumps({"report": str(output), "status": manifest["status"], "window": manifest["window"]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
