from __future__ import annotations

import posixpath
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips


HONGTANG_HEADER_WIDTHS_TWIPS = (3150, 3900, 2203)
# Word restarts the visible body-page number at physical page 4.  The cover,
# intentional blank verso and approval page are the three unnumbered pages,
# so the body total is NUMPAGES - 3.
HONGTANG_FRONT_MATTER_PAGES = 3
HEADER_CJK_FONT = "楷体_GB2312"
HEADER_LATIN_FONT = "Times New Roman"
HEADER_FONT_SIZE_PT = 9
BODY_END_PAGE_BOOKMARK = "BMSBodyLastPage"
_FORMAT_FONT_ATTRIBUTES = ("ascii", "hAnsi", "cs", "eastAsia")
_FORMAT_SIZE_ELEMENTS = ("sz", "szCs")


@dataclass(frozen=True)
class HeaderPaginationAudit:
    valid: bool
    header_parts: int
    page_fields: int
    numpages_fields: int
    total_page_formula_fields: int
    front_matter_pages: tuple[int, ...]
    duplicate_page_phrases: int
    formatting_errors: tuple[str, ...]
    details: tuple[str, ...]


@dataclass(frozen=True)
class FooterPaginationAudit:
    valid: bool
    footer_parts: int
    pagination_paragraphs: int
    page_fields: int
    sectionpages_fields: int
    static_total_paragraphs: int
    formatting_errors: tuple[str, ...]
    details: tuple[str, ...]


@dataclass(frozen=True)
class MultiSectionFooterPaginationAudit:
    valid: bool
    body_start_section: int
    body_section_count: int
    footer_parts: int
    pagination_paragraphs: int
    page_fields: int
    pageref_fields: int
    bookmark_start_count: int
    bookmark_end_count: int
    static_total_paragraphs: int
    formatting_errors: tuple[str, ...]
    details: tuple[str, ...]


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _style_run_element(run_element) -> None:
    r_pr = run_element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), HEADER_LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), HEADER_LATIN_FONT)
    r_fonts.set(qn("w:cs"), HEADER_LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), HEADER_CJK_FONT)

    size = r_pr.find(qn("w:sz"))
    if size is None:
        size = OxmlElement("w:sz")
        r_pr.append(size)
    size.set(qn("w:val"), str(HEADER_FONT_SIZE_PT * 2))

    size_cs = r_pr.find(qn("w:szCs"))
    if size_cs is None:
        size_cs = OxmlElement("w:szCs")
        r_pr.append(size_cs)
    size_cs.set(qn("w:val"), str(HEADER_FONT_SIZE_PT * 2))


def _style_run(run) -> None:
    run.font.size = Pt(HEADER_FONT_SIZE_PT)
    _style_run_element(run._element)


def _append_xml_run(paragraph, child) -> None:
    run = OxmlElement("w:r")
    _style_run_element(run)
    run.append(child)
    paragraph._p.append(run)


def _field_char(field_type: str, *, dirty: bool = False):
    node = OxmlElement("w:fldChar")
    node.set(qn("w:fldCharType"), field_type)
    if dirty:
        node.set(qn("w:dirty"), "true")
    return node


def _instruction(text: str):
    node = OxmlElement("w:instrText")
    node.set(qn("xml:space"), "preserve")
    node.text = text
    return node


def _append_field(paragraph, instruction: str, result: str = "1") -> None:
    _append_xml_run(paragraph, _field_char("begin", dirty=True))
    _append_xml_run(paragraph, _instruction(f" {instruction} "))
    _append_xml_run(paragraph, _field_char("separate"))

    result_run = paragraph.add_run(result)
    _style_run(result_run)

    _append_xml_run(paragraph, _field_char("end"))


def _append_field_to_element(paragraph_element, instruction: str, result: str = "1") -> None:
    """Append a styled complex field to a raw ``w:p`` element."""

    _append_xml_run_to_element(paragraph_element, _field_char("begin", dirty=True))
    _append_xml_run_to_element(paragraph_element, _instruction(f" {instruction} "))
    _append_xml_run_to_element(paragraph_element, _field_char("separate"))
    _append_text_run_to_element(paragraph_element, result)
    _append_xml_run_to_element(paragraph_element, _field_char("end"))


def _append_xml_run_to_element(paragraph_element, child) -> None:
    run = OxmlElement("w:r")
    _style_run_element(run)
    run.append(child)
    paragraph_element.append(run)


def _append_text_run_to_element(paragraph_element, text: str) -> None:
    node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        node.set(qn("xml:space"), "preserve")
    node.text = text
    _append_xml_run_to_element(paragraph_element, node)


def _paragraph_visible_text(paragraph_element) -> str:
    return "".join(node.text or "" for node in paragraph_element.iter(qn("w:t")))


def _field_command_is(value: str, command: str) -> bool:
    normalized = _normalize_field_instruction(value)
    expected = command.upper()
    return normalized == expected or normalized.startswith(expected + " ")


def _field_command_count(values: list[str], command: str) -> int:
    return sum(_field_command_is(value, command) for value in values)


def _pageref_targets_bookmark(value: str, bookmark_name: str) -> bool:
    normalized = _normalize_field_instruction(value)
    tokens = normalized.split()
    return (
        len(tokens) >= 2
        and tokens[0] == "PAGEREF"
        and tokens[1].strip('"') == bookmark_name.upper()
    )


def _is_body_footer_pagination_paragraph(paragraph_element) -> bool:
    if any(
        nested is not paragraph_element
        for nested in paragraph_element.iter(qn("w:p"))
    ):
        return False
    fields = _field_instructions(paragraph_element)
    text = _paragraph_visible_text(paragraph_element)
    compact = re.sub(r"\s+", "", text)
    return (
        _field_command_count(fields, "PAGE") > 0
        and "\u7b2c" in compact
        and "\u5171" in compact
        and "\u9875" in compact
    )


def _is_body_footer_page_number_paragraph(paragraph_element) -> bool:
    """Return whether one leaf footer paragraph displays a visible page number.

    The final Zhishan body section historically displayed only ``第 {PAGE} 页``
    while preceding body sections displayed ``第 {PAGE} 页 / 共 {NUMPAGES} 页``.
    Both forms are pagination paragraphs and must be upgraded to the same
    body-total contract.
    """

    if any(
        nested is not paragraph_element
        for nested in paragraph_element.iter(qn("w:p"))
    ):
        return False
    fields = _field_instructions(paragraph_element)
    compact = re.sub(r"\s+", "", _paragraph_visible_text(paragraph_element))
    return (
        _field_command_count(fields, "PAGE") > 0
        and "\u7b2c" in compact
        and "\u9875" in compact
    )


def _literal_visible_text(paragraph_element) -> str:
    """Return visible text that is not a cached complex/simple field result."""

    simple_field_text = {
        node
        for field in paragraph_element.iter(qn("w:fldSimple"))
        for node in field.iter(qn("w:t"))
    }
    field_stack: list[bool] = []
    values: list[str] = []
    for node in paragraph_element.iter():
        if node.tag == qn("w:fldChar"):
            field_type = node.get(qn("w:fldCharType"), "")
            if field_type == "begin":
                field_stack.append(False)
            elif field_type == "separate" and field_stack:
                field_stack[-1] = True
            elif field_type == "end" and field_stack:
                field_stack.pop()
            continue
        if node.tag != qn("w:t") or node in simple_field_text or field_stack:
            continue
        values.append(node.text or "")
    return "".join(values)


def _has_static_page_number_text(paragraph_element) -> bool:
    literal = _literal_visible_text(paragraph_element)
    return bool(re.search(r"[0-9\u96f6\u3007\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u5343]", literal))


def _visible_run_has_required_direct_format(run_element) -> bool:
    properties = _run_properties(run_element.find(qn("w:rPr")))
    expected_size = str(HEADER_FONT_SIZE_PT * 2)
    return (
        all(properties.get(name) == HEADER_LATIN_FONT for name in ("ascii", "hAnsi", "cs"))
        and properties.get("eastAsia") == HEADER_CJK_FONT
        and all(properties.get(name) == expected_size for name in _FORMAT_SIZE_ELEMENTS)
    )


def _normalize_visible_run_format(paragraph_element) -> bool:
    changed = False
    for run in paragraph_element.findall(".//" + qn("w:r")):
        if run.find(".//" + qn("w:t")) is None:
            continue
        if not _visible_run_has_required_direct_format(run):
            _style_run_element(run)
            changed = True
    return changed


def ensure_section_footer_pagination_fields(document: Document) -> int:
    """Use real PAGE/SECTIONPAGES fields for body-section footer totals.

    Guanbing's legacy template stores the body total as static text inside a
    footer text box.  The visible PAGE number restarts at one for the body
    section, so SECTIONPAGES is the correct total and remains correct when the
    generated report gains or loses pages.  Both DrawingML and VML fallback
    copies are normalised because Word may choose either representation.
    """

    changed = 0
    seen_parts: set[str] = set()
    for section in document.sections:
        page_numbering = section._sectPr.find(qn("w:pgNumType"))
        if (
            page_numbering is None
            or page_numbering.get(qn("w:start")) != "1"
        ):
            continue
        footer = section.footer
        part_name = str(footer.part.partname)
        if part_name in seen_parts:
            continue
        seen_parts.add(part_name)
        for paragraph_element in footer._element.iter(qn("w:p")):
            if not _is_body_footer_pagination_paragraph(paragraph_element):
                continue
            fields = _field_instructions(paragraph_element)
            if (
                _field_command_count(fields, "PAGE") == 1
                and _field_command_count(fields, "SECTIONPAGES") == 1
                and not _has_static_page_number_text(paragraph_element)
            ):
                if _normalize_visible_run_format(paragraph_element):
                    changed += 1
                continue
            for child in list(paragraph_element):
                if child.tag != qn("w:pPr"):
                    paragraph_element.remove(child)
            _append_text_run_to_element(paragraph_element, "\u7b2c ")
            _append_field_to_element(paragraph_element, "PAGE")
            _append_text_run_to_element(paragraph_element, " \u9875 \u5171 ")
            _append_field_to_element(paragraph_element, "SECTIONPAGES")
            _append_text_run_to_element(paragraph_element, " \u9875")
            changed += 1
    return changed


def _last_page_restart_section_index(document: Document) -> int:
    """Return the last section that explicitly restarts visible pages at one.

    Reports may restart numbering for approval sheets and the table of
    contents before restarting it once more for the real body.  The final
    explicit restart is therefore the stable boundary between front matter
    and body sections.
    """

    restart_indexes = []
    for index, section in enumerate(document.sections):
        page_numbering = section._sectPr.find(qn("w:pgNumType"))
        if (
            page_numbering is not None
            and page_numbering.get(qn("w:start")) == "1"
        ):
            restart_indexes.append(index)
    if not restart_indexes:
        raise ValueError("document has no section that restarts page numbering at 1")
    return restart_indexes[-1]


def _ensure_body_end_page_bookmark(
    document: Document,
    bookmark_name: str,
) -> None:
    """Place one hidden bookmark on the final body paragraph.

    ``PAGEREF`` resolves the bookmark using the body's restarted page-number
    sequence.  It therefore remains correct when the body spans portrait and
    landscape sections, unlike ``SECTIONPAGES`` (one section only), and avoids
    a brittle hard-coded ``NUMPAGES - n`` front-matter offset.
    """

    if not bookmark_name or not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", bookmark_name):
        raise ValueError(f"invalid Word bookmark name: {bookmark_name!r}")

    root = document._element
    removed_ids: set[str] = set()
    for start in list(root.iter(qn("w:bookmarkStart"))):
        if start.get(qn("w:name")) != bookmark_name:
            continue
        removed_ids.add(start.get(qn("w:id"), ""))
        parent = start.getparent()
        if parent is not None:
            parent.remove(start)
    for end in list(root.iter(qn("w:bookmarkEnd"))):
        if end.get(qn("w:id"), "") not in removed_ids:
            continue
        parent = end.getparent()
        if parent is not None:
            parent.remove(end)

    # Work only in the main-document story and include paragraphs nested in
    # tables.  ``document.paragraphs`` omits table-cell paragraphs and could
    # therefore anchor PAGEREF on an earlier physical page when the report
    # ends with a table.
    body = document._element.body
    paragraphs = list(body.iter(qn("w:p")))
    if not paragraphs:
        target = document.add_paragraph()._p
    else:
        target = next(
            (
                paragraph
                for paragraph in reversed(paragraphs)
                if _paragraph_visible_text(paragraph).strip()
                or paragraph.find(".//" + qn("w:drawing")) is not None
                or paragraph.find(".//" + qn("w:pict")) is not None
                or paragraph.find(".//" + qn("w:object")) is not None
            ),
            paragraphs[-1],
        )
    existing_ids = [
        int(value)
        for node in root.iter(qn("w:bookmarkStart"))
        if (value := node.get(qn("w:id"), "")).isdigit()
    ]
    bookmark_id = str(max(existing_ids, default=-1) + 1)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    # A zero-width bookmark at the end of the final content paragraph is less
    # ambiguous than spanning the paragraph when it happens to cross a page.
    target.append(start)
    target.append(end)


def _body_default_footer_elements(document: Document, body_start_index: int):
    """Yield unique default footer elements used by the body section chain."""

    current_part = None
    seen_parts: set[str] = set()
    for index, section in enumerate(document.sections):
        references = [
            reference
            for reference in section._sectPr.findall(qn("w:footerReference"))
            if reference.get(qn("w:type"), "default") == "default"
        ]
        if references:
            relationship_id = references[-1].get(qn("r:id"), "")
            current_part = document.part.related_parts.get(relationship_id)
        if index < body_start_index or current_part is None:
            continue
        part_name = str(current_part.partname)
        if part_name in seen_parts:
            continue
        seen_parts.add(part_name)
        yield current_part.element


def ensure_multisection_body_footer_pagination_fields(
    document: Document,
    *,
    bookmark_name: str = BODY_END_PAGE_BOOKMARK,
) -> int:
    """Use ``PAGE`` plus a body-end ``PAGEREF`` total for multi-section bodies.

    This contract is intended for reports whose visible body numbering restarts
    at one and then continues through multiple Word sections.  It deliberately
    does not replace the simpler ``SECTIONPAGES`` contract used by a one-section
    report body.
    """

    body_start_index = _last_page_restart_section_index(document)
    if len(document.sections) - body_start_index < 2:
        raise ValueError(
            "multi-section pagination requires at least two body sections; "
            "use ensure_section_footer_pagination_fields for a single body section"
        )
    _ensure_body_end_page_bookmark(document, bookmark_name)

    changed = 0
    candidates = 0
    for footer_element in _body_default_footer_elements(document, body_start_index):
        for paragraph_element in footer_element.iter(qn("w:p")):
            if not _is_body_footer_page_number_paragraph(paragraph_element):
                continue
            candidates += 1
            fields = _field_instructions(paragraph_element)
            pageref_fields = [
                value
                for value in fields
                if _field_command_is(value, "PAGEREF")
            ]
            if (
                _field_command_count(fields, "PAGE") == 1
                and len(pageref_fields) == 1
                and _pageref_targets_bookmark(pageref_fields[0], bookmark_name)
                and _field_command_count(fields, "NUMPAGES") == 0
                and _field_command_count(fields, "SECTIONPAGES") == 0
                and not _has_static_page_number_text(paragraph_element)
            ):
                if _normalize_visible_run_format(paragraph_element):
                    changed += 1
                continue
            for child in list(paragraph_element):
                if child.tag != qn("w:pPr"):
                    paragraph_element.remove(child)
            _append_text_run_to_element(paragraph_element, "\u7b2c ")
            _append_field_to_element(paragraph_element, "PAGE")
            _append_text_run_to_element(paragraph_element, " \u9875 / \u5171 ")
            _append_field_to_element(
                paragraph_element,
                f"PAGEREF {bookmark_name}",
            )
            _append_text_run_to_element(paragraph_element, " \u9875")
            changed += 1
    if candidates == 0:
        raise ValueError("body sections do not contain a recognizable pagination footer")
    return changed


def _append_total_pages_formula(paragraph, front_matter_pages: int, result: str = "1") -> None:
    """Append the real nested Word field ``{ = { NUMPAGES } - n }``."""

    if front_matter_pages < 0:
        raise ValueError("front_matter_pages must be non-negative")

    _append_xml_run(paragraph, _field_char("begin", dirty=True))
    _append_xml_run(paragraph, _instruction(" = "))

    _append_xml_run(paragraph, _field_char("begin", dirty=True))
    _append_xml_run(paragraph, _instruction(" NUMPAGES "))
    _append_xml_run(paragraph, _field_char("separate"))
    nested_result = paragraph.add_run(str(front_matter_pages + 1))
    _style_run(nested_result)
    _append_xml_run(paragraph, _field_char("end"))

    _append_xml_run(paragraph, _instruction(f" - {front_matter_pages} "))
    _append_xml_run(paragraph, _field_char("separate"))
    result_run = paragraph.add_run(result)
    _style_run(result_run)
    _append_xml_run(paragraph, _field_char("end"))


def _replace_cell_pagination(cell, front_matter_pages: int) -> None:
    paragraph = cell.paragraphs[0]
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _style_run(run)
    _append_field(paragraph, "PAGE")
    run = paragraph.add_run(" 页 共 ")
    _style_run(run)
    _append_total_pages_formula(paragraph, front_matter_pages)
    run = paragraph.add_run(" 页")
    _style_run(run)


def _normalize_header_row_widths(table, row) -> None:
    """Keep the report number on one line while preserving the page cell width."""

    if len(row.cells) != len(HONGTANG_HEADER_WIDTHS_TWIPS):
        return
    for column_index, width_twips in enumerate(HONGTANG_HEADER_WIDTHS_TWIPS):
        width = Twips(width_twips)
        table.columns[column_index].width = width
        row.cells[column_index].width = width
        tc_width = row.cells[column_index]._tc.get_or_add_tcPr().get_or_add_tcW()
        tc_width.type = "dxa"
        tc_width.w = width_twips


def ensure_header_pagination_fields(
    document: Document,
    *,
    front_matter_pages: int = HONGTANG_FRONT_MATTER_PAGES,
) -> int:
    """Replace Hongtang's static page text with PAGE and adjusted NUMPAGES fields."""

    if front_matter_pages < 0:
        raise ValueError("front_matter_pages must be non-negative")

    changed = 0
    seen_parts: set[str] = set()
    for section in document.sections:
        header = section.header
        part_name = str(header.part.partname)
        if part_name in seen_parts:
            continue
        seen_parts.add(part_name)
        for table in header.tables:
            for row in table.rows:
                texts = [cell.text.strip() for cell in row.cells]
                if len(row.cells) < 3 or not any("报告编号" in text for text in texts):
                    continue
                _normalize_header_row_widths(table, row)
                _replace_cell_pagination(row.cells[-1], front_matter_pages)
                changed += 1
    return changed


def _normalize_field_instruction(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip().upper()


def _field_instructions(element) -> list[str]:
    """Return semantic field codes while ignoring Word's cached results.

    After Word updates a nested field, it can serialize the cached NUMPAGES
    result as ``w:instrText`` inside the nested field's result section.  A flat
    scan therefore sees ``= NUMPAGES 79 - 4`` and mistakes a valid formula for
    an invalid one.  Parsing begin/separate/end boundaries keeps only actual
    field-code text and still accepts legacy ``w:fldSimple`` fields.
    """

    instructions: list[str] = []
    stack: list[dict[str, object]] = []
    for node in element.iter():
        if node.tag == qn("w:fldSimple"):
            value = _normalize_field_instruction(node.get(qn("w:instr"), ""))
            if value:
                instructions.append(value)
            continue
        if node.tag == qn("w:fldChar"):
            field_type = node.get(qn("w:fldCharType"), "")
            if field_type == "begin":
                stack.append({"parts": [], "in_result": False})
            elif field_type == "separate" and stack:
                stack[-1]["in_result"] = True
            elif field_type == "end" and stack:
                frame = stack.pop()
                value = _normalize_field_instruction("".join(frame["parts"]))
                if value:
                    instructions.append(value)
            continue
        if node.tag != qn("w:instrText"):
            continue
        if not stack:
            value = _normalize_field_instruction(node.text or "")
            if value:
                instructions.append(value)
            continue
        # Cached text inside any active result section is not field code.
        if bool(stack[-1]["in_result"]):
            continue
        for frame in stack:
            if not bool(frame["in_result"]):
                frame["parts"].append(node.text or "")
    return instructions


def _paragraph_field_instructions(paragraph) -> list[str]:
    return _field_instructions(paragraph)


def _run_properties(r_pr) -> dict[str, str]:
    values: dict[str, str] = {}
    if r_pr is None:
        return values
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is not None:
        for attribute in _FORMAT_FONT_ATTRIBUTES:
            value = fonts.get(qn(f"w:{attribute}"))
            if value:
                values[attribute] = value
    for element_name in _FORMAT_SIZE_ELEMENTS:
        node = r_pr.find(qn(f"w:{element_name}"))
        if node is not None and node.get(qn("w:val")):
            values[element_name] = node.get(qn("w:val"))
    return values


def _style_context(styles_root) -> tuple[dict[str, str], str, dict[str, tuple[str, dict[str, str]]]]:
    defaults: dict[str, str] = {}
    default_r_pr = styles_root.find(
        f"./{qn('w:docDefaults')}/{qn('w:rPrDefault')}/{qn('w:rPr')}"
    )
    defaults.update(_run_properties(default_r_pr))

    default_paragraph_style = ""
    styles: dict[str, tuple[str, dict[str, str]]] = {}
    for style in styles_root.findall(qn("w:style")):
        style_id = style.get(qn("w:styleId"), "")
        if not style_id:
            continue
        based_on_node = style.find(qn("w:basedOn"))
        based_on = based_on_node.get(qn("w:val"), "") if based_on_node is not None else ""
        styles[style_id] = (based_on, _run_properties(style.find(qn("w:rPr"))))
        if (
            style.get(qn("w:type")) == "paragraph"
            and style.get(qn("w:default")) in {"1", "true", "on"}
        ):
            default_paragraph_style = style_id
    return defaults, default_paragraph_style, styles


def _resolved_style_properties(
    style_id: str,
    styles: dict[str, tuple[str, dict[str, str]]],
    seen: set[str] | None = None,
) -> dict[str, str]:
    if not style_id or style_id not in styles:
        return {}
    visited = set() if seen is None else set(seen)
    if style_id in visited:
        return {}
    visited.add(style_id)
    based_on, own = styles[style_id]
    values = _resolved_style_properties(based_on, styles, visited)
    values.update(own)
    return values


def _pagination_formatting_errors(
    paragraph,
    part_name: str,
    style_context: tuple[dict[str, str], str, dict[str, tuple[str, dict[str, str]]]] | None,
) -> list[str]:
    errors: list[str] = []
    expected_size = str(HEADER_FONT_SIZE_PT * 2)
    defaults: dict[str, str] = {}
    default_paragraph_style = ""
    styles: dict[str, tuple[str, dict[str, str]]] = {}
    if style_context is not None:
        defaults, default_paragraph_style, styles = style_context

    paragraph_properties = dict(defaults)
    p_pr = paragraph.find(qn("w:pPr"))
    style_id = default_paragraph_style
    if p_pr is not None:
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None and p_style.get(qn("w:val")):
            style_id = p_style.get(qn("w:val"))
    paragraph_properties.update(_resolved_style_properties(style_id, styles))
    if p_pr is not None:
        paragraph_properties.update(_run_properties(p_pr.find(qn("w:rPr"))))

    for run_index, run in enumerate(paragraph.findall(".//" + qn("w:r")), start=1):
        # Field-code runs are hidden implementation details.  Desktop Word and
        # WPS are free to rewrite or omit their direct formatting while keeping
        # the visible literal/result runs correctly formatted.
        if run.find(".//" + qn("w:t")) is None:
            continue
        effective = dict(paragraph_properties)
        r_pr = run.find(qn("w:rPr"))
        if r_pr is not None:
            r_style = r_pr.find(qn("w:rStyle"))
            if r_style is not None and r_style.get(qn("w:val")):
                effective.update(_resolved_style_properties(r_style.get(qn("w:val")), styles))
            effective.update(_run_properties(r_pr))
        for attr in ("ascii", "hAnsi", "cs"):
            if effective.get(attr) != HEADER_LATIN_FONT:
                errors.append(
                    f"{part_name}: pagination run {run_index} effective {attr} font is not {HEADER_LATIN_FONT}"
                )
        if effective.get("eastAsia") != HEADER_CJK_FONT:
            errors.append(
                f"{part_name}: pagination run {run_index} effective eastAsia font is not {HEADER_CJK_FONT}"
            )
        for size_tag in _FORMAT_SIZE_ELEMENTS:
            if effective.get(size_tag) != expected_size:
                errors.append(
                    f"{part_name}: pagination run {run_index} effective {size_tag} is not {expected_size} half-points"
                )
    return errors


def _resolve_package_relationship_target(source_part: str, target: str) -> str | None:
    """Resolve one internal OPC relationship target to an archive member name."""

    value = (target or "").strip().replace("\\", "/")
    if not value or "://" in value:
        return None
    if value.startswith("/"):
        resolved = posixpath.normpath(value.lstrip("/"))
    else:
        resolved = posixpath.normpath(
            posixpath.join(posixpath.dirname(source_part), value)
        )
    if resolved == ".." or resolved.startswith("../"):
        return None
    return resolved


def _body_section_default_footer_parts(
    archive: zipfile.ZipFile,
    document_root,
    namespace: dict[str, str],
) -> tuple[set[str], list[str]]:
    """Resolve restart-section default footers, including OOXML inheritance."""

    relationship_namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    relationships_root = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
    relationship_targets = {
        item.get("Id", ""): (item.get("Target", ""), item.get("TargetMode", ""))
        for item in relationships_root.findall(
            f"{{{relationship_namespace}}}Relationship"
        )
    }

    selected: set[str] = set()
    errors: list[str] = []
    inherited_default: str | None = None
    restart_sections = 0
    for section_index, section in enumerate(
        document_root.findall(".//w:sectPr", namespace),
        start=1,
    ):
        default_references = [
            reference
            for reference in section.findall("w:footerReference", namespace)
            if reference.get(qn("w:type"), "default") == "default"
        ]
        current_default = inherited_default
        if default_references:
            current_default = None
            if len(default_references) != 1:
                errors.append(
                    f"section {section_index}: expected one default footerReference, "
                    f"found {len(default_references)}"
                )
            else:
                relationship_id = default_references[0].get(qn("r:id"), "")
                relationship = relationship_targets.get(relationship_id)
                if relationship is None:
                    errors.append(
                        f"section {section_index}: unresolved default footer relationship "
                        f"{relationship_id or '<missing>'}"
                    )
                elif relationship[1].lower() == "external":
                    errors.append(
                        f"section {section_index}: external default footer relationship is invalid"
                    )
                else:
                    current_default = _resolve_package_relationship_target(
                        "word/document.xml",
                        relationship[0],
                    )
                    if current_default is None:
                        errors.append(
                            f"section {section_index}: invalid default footer target "
                            f"{relationship[0]!r}"
                        )
            inherited_default = current_default

        page_numbering = section.find("w:pgNumType", namespace)
        if page_numbering is None or page_numbering.get(qn("w:start")) != "1":
            continue
        restart_sections += 1
        if current_default is None:
            errors.append(
                f"section {section_index}: restart section has no resolvable inherited/default footer"
            )
            continue
        if current_default not in archive.namelist():
            errors.append(
                f"section {section_index}: default footer part is missing: {current_default}"
            )
            continue
        selected.add(current_default)

    if restart_sections == 0:
        errors.append("no section restarts visible page numbering at 1")
    if len(selected) != 1:
        errors.append(
            f"expected one unique restart-section default footer, found {len(selected)}"
        )
    return selected, errors


def _multisection_body_default_footer_parts(
    archive: zipfile.ZipFile,
    document_root,
    namespace: dict[str, str],
) -> tuple[int, int, set[str], list[str]]:
    """Resolve every default footer from the final page-number restart onward."""

    relationship_namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    relationships_root = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
    relationship_targets = {
        item.get("Id", ""): (item.get("Target", ""), item.get("TargetMode", ""))
        for item in relationships_root.findall(
            f"{{{relationship_namespace}}}Relationship"
        )
    }
    sections = document_root.findall(".//w:sectPr", namespace)
    restart_indexes = [
        index
        for index, section in enumerate(sections)
        if (
            (page_numbering := section.find("w:pgNumType", namespace)) is not None
            and page_numbering.get(qn("w:start")) == "1"
        )
    ]
    if not restart_indexes:
        return 0, 0, set(), ["no section restarts visible page numbering at 1"]
    body_start_index = restart_indexes[-1]
    body_section_count = len(sections) - body_start_index
    errors: list[str] = []
    if body_section_count < 2:
        errors.append(
            "final page-number restart does not begin a multi-section report body"
        )

    selected: set[str] = set()
    inherited_default: str | None = None
    for section_index, section in enumerate(sections):
        default_references = [
            reference
            for reference in section.findall("w:footerReference", namespace)
            if reference.get(qn("w:type"), "default") == "default"
        ]
        current_default = inherited_default
        if default_references:
            current_default = None
            if len(default_references) != 1:
                errors.append(
                    f"section {section_index + 1}: expected one default footerReference, "
                    f"found {len(default_references)}"
                )
            else:
                relationship_id = default_references[0].get(qn("r:id"), "")
                relationship = relationship_targets.get(relationship_id)
                if relationship is None:
                    errors.append(
                        f"section {section_index + 1}: unresolved default footer relationship "
                        f"{relationship_id or '<missing>'}"
                    )
                elif relationship[1].lower() == "external":
                    errors.append(
                        f"section {section_index + 1}: external default footer relationship is invalid"
                    )
                else:
                    current_default = _resolve_package_relationship_target(
                        "word/document.xml",
                        relationship[0],
                    )
                    if current_default is None:
                        errors.append(
                            f"section {section_index + 1}: invalid default footer target "
                            f"{relationship[0]!r}"
                        )
            inherited_default = current_default

        if section_index < body_start_index:
            continue
        if current_default is None:
            errors.append(
                f"section {section_index + 1}: body section has no resolvable inherited/default footer"
            )
            continue
        if current_default not in archive.namelist():
            errors.append(
                f"section {section_index + 1}: default footer part is missing: {current_default}"
            )
            continue
        selected.add(current_default)

    if not selected:
        errors.append("multi-section report body has no resolvable default footer parts")
    return body_start_index + 1, body_section_count, selected, errors


def audit_multisection_body_footer_pagination_fields(
    docx_path: Path,
    *,
    bookmark_name: str = BODY_END_PAGE_BOOKMARK,
) -> MultiSectionFooterPaginationAudit:
    """Audit the PAGE/PAGEREF contract used by a restarted multi-section body."""

    footer_parts = 0
    pagination_paragraphs = 0
    page_fields = 0
    pageref_fields = 0
    static_or_invalid = 0
    formatting_errors: list[str] = []
    details: list[str] = []
    with zipfile.ZipFile(docx_path) as archive:
        namespace = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        document_root = ET.fromstring(archive.read("word/document.xml"))
        body_start_section, body_section_count, body_footer_names, resolution_errors = (
            _multisection_body_default_footer_parts(
                archive,
                document_root,
                namespace,
            )
        )
        details.extend(resolution_errors)
        bookmark_starts = [
            node
            for node in document_root.findall(".//w:bookmarkStart", namespace)
            if node.get(qn("w:name")) == bookmark_name
        ]
        bookmark_ids = {node.get(qn("w:id"), "") for node in bookmark_starts}
        bookmark_ends = [
            node
            for node in document_root.findall(".//w:bookmarkEnd", namespace)
            if node.get(qn("w:id"), "") in bookmark_ids
        ]
        bookmark_start_count = len(bookmark_starts)
        bookmark_end_count = len(bookmark_ends)
        details.append(
            f"body starts at section {body_start_section}; sections={body_section_count}; "
            f"bookmark starts={bookmark_start_count}, ends={bookmark_end_count}"
        )

        style_context = None
        if "word/styles.xml" in archive.namelist():
            style_context = _style_context(ET.fromstring(archive.read("word/styles.xml")))
        for name in sorted(body_footer_names):
            root = ET.fromstring(archive.read(name))
            part_candidates = 0
            part_page = 0
            part_pageref = 0
            part_invalid = 0
            for paragraph in root.findall(".//" + qn("w:p")):
                if not _is_body_footer_page_number_paragraph(paragraph):
                    continue
                part_candidates += 1
                fields = _paragraph_field_instructions(paragraph)
                page_count = _field_command_count(fields, "PAGE")
                matching_pageref = sum(
                    _pageref_targets_bookmark(value, bookmark_name)
                    for value in fields
                    if _field_command_is(value, "PAGEREF")
                )
                part_page += page_count
                part_pageref += matching_pageref
                if (
                    page_count != 1
                    or matching_pageref != 1
                    or _field_command_count(fields, "PAGEREF") != 1
                    or _field_command_count(fields, "NUMPAGES") != 0
                    or _field_command_count(fields, "SECTIONPAGES") != 0
                    or _has_static_page_number_text(paragraph)
                ):
                    part_invalid += 1
                formatting_errors.extend(
                    _pagination_formatting_errors(paragraph, name, style_context)
                )
            if part_candidates == 0:
                details.append(f"{name}: no recognizable PAGE footer paragraph")
                static_or_invalid += 1
                continue
            footer_parts += 1
            pagination_paragraphs += part_candidates
            page_fields += part_page
            pageref_fields += part_pageref
            static_or_invalid += part_invalid
            details.append(
                f"{name}: paragraphs={part_candidates}, PAGE={part_page}, "
                f"PAGEREF({bookmark_name})={part_pageref}, "
                f"static_or_invalid={part_invalid}"
            )

    valid = (
        body_start_section > 0
        and body_section_count >= 2
        and footer_parts == len(body_footer_names)
        and footer_parts >= 1
        and pagination_paragraphs >= footer_parts
        and page_fields == pagination_paragraphs
        and pageref_fields == pagination_paragraphs
        and bookmark_start_count == 1
        and bookmark_end_count == 1
        and static_or_invalid == 0
        and not formatting_errors
        and not resolution_errors
    )
    return MultiSectionFooterPaginationAudit(
        valid=valid,
        body_start_section=body_start_section,
        body_section_count=body_section_count,
        footer_parts=footer_parts,
        pagination_paragraphs=pagination_paragraphs,
        page_fields=page_fields,
        pageref_fields=pageref_fields,
        bookmark_start_count=bookmark_start_count,
        bookmark_end_count=bookmark_end_count,
        static_total_paragraphs=static_or_invalid,
        formatting_errors=tuple(formatting_errors),
        details=tuple(details),
    )


def audit_section_footer_pagination_fields(docx_path: Path) -> FooterPaginationAudit:
    """Audit body footer PAGE/SECTIONPAGES fields and their visible styling."""

    footer_parts = 0
    pagination_paragraphs = 0
    page_fields = 0
    sectionpages_fields = 0
    static_total_paragraphs = 0
    formatting_errors: list[str] = []
    details: list[str] = []
    with zipfile.ZipFile(docx_path) as archive:
        namespace = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        document_root = ET.fromstring(archive.read("word/document.xml"))
        body_footer_names, resolution_errors = _body_section_default_footer_parts(
            archive,
            document_root,
            namespace,
        )
        details.extend(resolution_errors)
        style_context = None
        if "word/styles.xml" in archive.namelist():
            style_context = _style_context(ET.fromstring(archive.read("word/styles.xml")))
        for name in archive.namelist():
            if not re.fullmatch(r"word/footer\d+\.xml", name):
                continue
            if name not in body_footer_names:
                continue
            root = ET.fromstring(archive.read(name))
            part_candidates = 0
            part_page = 0
            part_sectionpages = 0
            part_static = 0
            for paragraph in root.findall(".//" + qn("w:p")):
                if not _is_body_footer_pagination_paragraph(paragraph):
                    continue
                part_candidates += 1
                fields = _paragraph_field_instructions(paragraph)
                page_count = _field_command_count(fields, "PAGE")
                section_count = _field_command_count(fields, "SECTIONPAGES")
                part_page += page_count
                part_sectionpages += section_count
                if (
                    page_count != 1
                    or section_count != 1
                    or _has_static_page_number_text(paragraph)
                ):
                    part_static += 1
                formatting_errors.extend(
                    _pagination_formatting_errors(paragraph, name, style_context)
                )
            if part_candidates:
                footer_parts += 1
                pagination_paragraphs += part_candidates
                page_fields += part_page
                sectionpages_fields += part_sectionpages
                static_total_paragraphs += part_static
                details.append(
                    f"{name}: paragraphs={part_candidates}, PAGE={part_page}, "
                    f"SECTIONPAGES={part_sectionpages}, static_or_invalid={part_static}"
                )
    valid = (
        footer_parts == 1
        and pagination_paragraphs >= 1
        and page_fields == pagination_paragraphs
        and sectionpages_fields == pagination_paragraphs
        and static_total_paragraphs == 0
        and not formatting_errors
        and not resolution_errors
    )
    return FooterPaginationAudit(
        valid=valid,
        footer_parts=footer_parts,
        pagination_paragraphs=pagination_paragraphs,
        page_fields=page_fields,
        sectionpages_fields=sectionpages_fields,
        static_total_paragraphs=static_total_paragraphs,
        formatting_errors=tuple(formatting_errors),
        details=tuple(details),
    )


def audit_header_pagination_fields(
    docx_path: Path,
    *,
    front_matter_pages: int = HONGTANG_FRONT_MATTER_PAGES,
) -> HeaderPaginationAudit:
    page_fields = 0
    numpages_fields = 0
    formula_fields = 0
    offsets: list[int] = []
    duplicate_phrases = 0
    formatting_errors: list[str] = []
    details: list[str] = []
    header_parts = 0
    with zipfile.ZipFile(docx_path) as archive:
        style_context = None
        if "word/styles.xml" in archive.namelist():
            style_context = _style_context(ET.fromstring(archive.read("word/styles.xml")))
        for name in archive.namelist():
            if not re.fullmatch(r"word/header\d+\.xml", name):
                continue
            xml = archive.read(name).decode("utf-8")
            if "报告编号" not in xml:
                continue
            header_parts += 1
            root = ET.fromstring(xml)
            namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            fields = _field_instructions(root)
            page_count = sum(value == "PAGE" for value in fields)
            numpages_count = sum(value == "NUMPAGES" for value in fields)
            found_offsets = [
                int(match.group(1))
                for value in fields
                if (match := re.fullmatch(r"=\s*NUMPAGES\s*-\s*(\d+)", value)) is not None
            ]
            formula_count = len(found_offsets)
            page_fields += page_count
            numpages_fields += numpages_count
            formula_fields += formula_count
            offsets.extend(found_offsets)
            plain_text = "".join(node.text or "" for node in root.findall(".//w:t", namespace))
            phrase_count = plain_text.count("第 ")
            duplicate_phrases += max(0, phrase_count - 1)

            for paragraph in root.findall(".//w:p", namespace):
                instructions = _paragraph_field_instructions(paragraph)
                if "PAGE" in instructions:
                    formatting_errors.extend(
                        _pagination_formatting_errors(paragraph, name, style_context)
                    )

            details.append(
                f"{name}: PAGE={page_count}, NUMPAGES={numpages_count}, "
                f"adjusted_total={found_offsets}, page_phrases={phrase_count}"
            )
    valid = (
        header_parts == 1
        and page_fields == 1
        and numpages_fields == 1
        and formula_fields == 1
        and offsets == [front_matter_pages]
        and duplicate_phrases == 0
        and not formatting_errors
    )
    return HeaderPaginationAudit(
        valid=valid,
        header_parts=header_parts,
        page_fields=page_fields,
        numpages_fields=numpages_fields,
        total_page_formula_fields=formula_fields,
        front_matter_pages=tuple(offsets),
        duplicate_page_phrases=duplicate_phrases,
        formatting_errors=tuple(formatting_errors),
        details=tuple(details),
    )
