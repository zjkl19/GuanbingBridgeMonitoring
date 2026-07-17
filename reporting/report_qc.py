from __future__ import annotations

import argparse
import hashlib
import json
import re
from zipfile import ZipFile
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table
from lxml import etree

try:
    from docx_header_fields import audit_section_footer_pagination_fields
    from docx_table_utils import find_first_row_index_by_first_cell, table_has_first_cell
    from image_block_utils import count_docx_images
except Exception:  # pragma: no cover - package import fallback
    from .docx_header_fields import audit_section_footer_pagination_fields
    from .docx_table_utils import find_first_row_index_by_first_cell, table_has_first_cell
    from .image_block_utils import count_docx_images


COMMON_FORBIDDEN_REVIEW_PHRASES = (
    "需结合原始数据和现场运维记录复核",
    "建议结合原始数据和传感器状态进一步复核",
    "建议结合原始数据进一步复核",
)

JLJ_FORBIDDEN_REVIEW_PHRASES = COMMON_FORBIDDEN_REVIEW_PHRASES + (
    "当前吊杆参数配置尚未完整校核",
    "索力换算结果暂仅用于时程展示",
)

REPORT_NUMBER_PATTERN = re.compile(
    r"报告编号\s*[:：]\s*([A-Za-z0-9]+(?:[-/][A-Za-z0-9]+)*)",
    re.IGNORECASE,
)
W_P = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
W_T = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"


@dataclass(frozen=True)
class ReportQcIssue:
    code: str
    severity: str
    message: str
    detail: str = ""


@dataclass(frozen=True)
class ReportQcResult:
    kind: str
    docx_path: str
    checked_at: str
    status: str
    issue_count: int
    warning_count: int
    issues: list[ReportQcIssue]
    summary: dict[str, Any]


def _all_doc_text(doc) -> str:
    parts: list[str] = []
    parts.extend(para.text for para in doc.paragraphs if para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _paragraph_own_text(paragraph) -> str:
    """Return text belonging to this paragraph, excluding nested text boxes."""
    nodes = []
    for text_node in paragraph.iter(W_T):
        owner = text_node.getparent()
        while owner is not None and owner.tag != W_P:
            owner = owner.getparent()
        if owner is paragraph:
            nodes.append(text_node.text or "")
    return "".join(nodes)


def _docx_report_numbers(path: Path) -> list[str]:
    """Read report numbers from body and every header/footer, including text boxes."""
    numbers: list[str] = []
    parser = etree.XMLParser(resolve_entities=False, no_network=True)
    with ZipFile(path) as archive:
        names = [
            name
            for name in archive.namelist()
            if name == "word/document.xml"
            or re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
        ]
        for name in names:
            root = etree.fromstring(archive.read(name), parser=parser)
            for paragraph in root.iter(W_P):
                text = _paragraph_own_text(paragraph)
                numbers.extend(match.group(1) for match in REPORT_NUMBER_PATTERN.finditer(text))
    return numbers


def _base_doc_summary(path: Path, doc) -> dict[str, Any]:
    return {
        "exists": path.exists(),
        "table_count": len(doc.tables),
        "image_count": count_docx_images(path),
        "paragraph_count": len(doc.paragraphs),
    }


def _check_common_report(
    *,
    kind: str,
    docx_path: Path | str,
    forbidden_phrases: tuple[str, ...] = COMMON_FORBIDDEN_REVIEW_PHRASES,
    require_images: bool = True,
) -> tuple[Path, Any | None, str, list[ReportQcIssue], dict[str, Any]]:
    path = Path(docx_path)
    issues: list[ReportQcIssue] = []
    summary: dict[str, Any] = {"exists": path.exists()}
    if not path.exists():
        issues.append(ReportQcIssue("missing-file", "error", f"报告文件不存在: {path}"))
        return path, None, "", issues, summary

    doc = Document(str(path))
    text = _all_doc_text(doc)
    summary.update(_base_doc_summary(path, doc))

    for phrase in forbidden_phrases:
        count = text.count(phrase)
        if count:
            issues.append(
                ReportQcIssue(
                    "forbidden-review-phrase",
                    "warning",
                    f"报告仍包含需删除的复核/临时说明: {phrase}",
                    f"count={count}",
                )
            )

    if require_images and summary["image_count"] <= 0:
        issues.append(ReportQcIssue("no-images", "warning", "报告中未检测到图片。"))
    return path, doc, text, issues, summary


def _find_front_cover_summary_table(doc) -> tuple[int | None, Table | None]:
    for idx, table in enumerate(doc.tables):
        if len(table.columns) < 2:
            continue
        if find_first_row_index_by_first_cell(table, "委托单位") == 0 and find_first_row_index_by_first_cell(table, "监测结果") is not None:
            return idx, table
    return None, None


def _front_summary_table_indices(doc, cover_idx: int | None) -> list[int]:
    if cover_idx is None:
        return []
    out = [cover_idx]
    for idx in range(cover_idx + 1, len(doc.tables)):
        table = doc.tables[idx]
        if table_has_first_cell(table, "监测结果"):
            out.append(idx)
            continue
        break
    return out


def _has_summary_table(table: Table) -> bool:
    return table_has_first_cell(table, "监测结果")


def check_jlj_report(
    docx_path: Path | str,
    *,
    expected_period_label: str | None = None,
    expected_image_paths: list[Path | str] | None = None,
) -> ReportQcResult:
    path, doc, text, issues, summary = _check_common_report(
        kind="jlj_monthly",
        docx_path=docx_path,
        forbidden_phrases=JLJ_FORBIDDEN_REVIEW_PHRASES,
    )
    summary.update({
        "summary_table_indices": [],
        "front_summary_table_indices": [],
        "transfer_marker_count": 0,
        "continue_marker_count": 0,
    })
    if doc is None:
        return _build_result("jlj_monthly", path, issues, summary)
    summary["transfer_marker_count"] = text.count("（转下页）")
    summary["continue_marker_count"] = text.count("（续上页）")

    cover_idx, _ = _find_front_cover_summary_table(doc)
    front_indices = _front_summary_table_indices(doc, cover_idx)
    summary_indices = [idx for idx, table in enumerate(doc.tables) if _has_summary_table(table)]
    summary["summary_table_indices"] = summary_indices
    summary["front_summary_table_indices"] = front_indices

    if cover_idx is None:
        issues.append(ReportQcIssue("missing-front-summary", "error", "未找到首页委托单位表格内的监测结果区域。"))
    else:
        stale_before = [idx for idx in summary_indices if idx < cover_idx]
        if stale_before:
            issues.append(
                ReportQcIssue(
                    "stale-summary-before-cover",
                    "warning",
                    "首页委托单位表格前存在独立监测结果表，可能是模板残留。",
                    ",".join(str(idx) for idx in stale_before),
                )
            )
        stale_later = [idx for idx in summary_indices if idx not in front_indices]
        if stale_later:
            issues.append(
                ReportQcIssue(
                    "summary-table-outside-front-block",
                    "warning",
                    "结论页监测结果表未连续贴在首页表格后，可能出现重复或错位。",
                    ",".join(str(idx) for idx in stale_later),
                )
            )

    if summary["transfer_marker_count"] != summary["continue_marker_count"]:
        issues.append(
            ReportQcIssue(
                "continuation-marker-mismatch",
                "info",
                "（转下页）与（续上页）数量不一致，仅作为分页提示复核信息。",
                f"transfer={summary['transfer_marker_count']}, continue={summary['continue_marker_count']}",
            )
        )

    bad_unit_matches = sorted(set(re.findall(r"(?:cm|m)/s[?？2]", text)))
    if bad_unit_matches:
        issues.append(
            ReportQcIssue(
                "unit-superscript-risk",
                "warning",
                "报告中仍存在疑似加速度单位上标错误。",
                ", ".join(bad_unit_matches),
            )
        )

    if expected_period_label:
        expected = _normalized_month_label(expected_period_label)
        found = [
            _normalized_month_label(f"{year}年{month}月")
            for year, month in re.findall(
                r"监测时间[：:]\s*[（(]?(\d{4})年\s*0?(\d{1,2})月", text
            )
        ]
        summary["expected_period_label"] = expected
        summary["document_period_labels"] = found
        if not found or expected not in found or any(value != expected for value in found):
            issues.append(
                ReportQcIssue(
                    "period-mismatch",
                    "error",
                    f"九龙江月报期次与任务不一致，应为{expected}。",
                    f"document_period_labels={found}",
                )
            )

    if expected_image_paths is not None:
        embedded_hashes = _docx_media_hashes(path)
        expected_hashes: list[str] = []
        for item in expected_image_paths:
            image_path = Path(item)
            if not image_path.is_file():
                issues.append(
                    ReportQcIssue(
                        "missing-expected-image",
                        "error",
                        "九龙江月报指定的结果图片源文件不存在。",
                        str(image_path),
                    )
                )
                continue
            digest = hashlib.sha256(image_path.read_bytes()).hexdigest().upper()
            expected_hashes.append(digest)
            if digest not in embedded_hashes:
                issues.append(
                    ReportQcIssue(
                        "missing-expected-image",
                        "error",
                        "九龙江月报未嵌入指定的本期结果图片。",
                        str(image_path),
                    )
                )
        summary["expected_result_image_count"] = len(expected_hashes)
        summary["matched_result_image_count"] = sum(
            digest in embedded_hashes for digest in expected_hashes
        )

    return _build_result("jlj_monthly", path, issues, summary)


def check_hongtang_report(docx_path: Path | str) -> ReportQcResult:
    path, doc, text, issues, summary = _check_common_report(kind="hongtang_period", docx_path=docx_path)
    if doc is None:
        return _build_result("hongtang_period", path, issues, summary)
    for required in ("监测结果", "交通状况监测", "结构应变监测"):
        if required not in text:
            issues.append(ReportQcIssue("missing-expected-text", "warning", f"洪塘周期报未检测到关键文本: {required}"))
    return _build_result("hongtang_period", path, issues, summary)


def check_guanbing_report(docx_path: Path | str) -> ReportQcResult:
    path, doc, text, issues, summary = _check_common_report(kind="guanbing_monthly", docx_path=docx_path)
    if doc is None:
        return _build_result("guanbing_monthly", path, issues, summary)
    for required in ("G104", "管柄大桥"):
        if required in text:
            break
    else:
        issues.append(ReportQcIssue("missing-expected-text", "warning", "管柄月报未检测到 G104/管柄大桥项目文本。"))
    if "m/s2" in text:
        issues.append(ReportQcIssue("unit-superscript-risk", "warning", "报告中仍存在 m/s2 文本，需检查是否应为 m/s²。"))
    report_numbers = _docx_report_numbers(path)
    unique_report_numbers = sorted(set(report_numbers))
    summary["report_numbers"] = unique_report_numbers
    summary["report_number_occurrence_count"] = len(report_numbers)
    if not unique_report_numbers:
        issues.append(
            ReportQcIssue(
                "missing-report-number",
                "error",
                "管柄月报正文及各类页眉/页脚中未检测到报告编号。",
            )
        )
    elif len(unique_report_numbers) > 1:
        issues.append(
            ReportQcIssue(
                "inconsistent-report-number",
                "error",
                "管柄月报正文及各类页眉/页脚中的报告编号不一致。",
                ", ".join(unique_report_numbers),
            )
        )
    pagination = audit_section_footer_pagination_fields(path)
    summary["footer_pagination"] = asdict(pagination)
    if not pagination.valid:
        issues.append(
            ReportQcIssue(
                "invalid-footer-pagination",
                "error",
                "Guanbing report body footer must use PAGE and SECTIONPAGES fields.",
                "; ".join((*pagination.details, *pagination.formatting_errors)),
            )
        )
    return _build_result("guanbing_monthly", path, issues, summary)


def _normalized_month_label(value: str) -> str:
    match = re.search(r"(\d{4})年\s*0?(\d{1,2})月", str(value or ""))
    return f"{int(match.group(1))}年{int(match.group(2))}月" if match else str(value or "").strip()


def _docx_media_hashes(path: Path) -> set[str]:
    with ZipFile(path) as archive:
        return {
            hashlib.sha256(archive.read(name)).hexdigest().upper()
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        }


def check_shuixianhua_report(
    docx_path: Path | str,
    *,
    expected_period_label: str | None = None,
    expected_image_paths: list[Path | str] | None = None,
) -> ReportQcResult:
    path, doc, text, issues, summary = _check_common_report(
        kind="shuixianhua_monthly", docx_path=docx_path
    )
    if doc is None:
        return _build_result("shuixianhua_monthly", path, issues, summary)
    for required in ("水仙花大桥", "监测结果"):
        if required not in text:
            issues.append(
                ReportQcIssue(
                    "missing-expected-text",
                    "warning",
                    f"水仙花月报未检测到关键文本: {required}",
                )
            )
    if expected_period_label:
        expected = _normalized_month_label(expected_period_label)
        found = [
            _normalized_month_label(f"{year}年{month}月")
            for year, month in re.findall(r"监测时间[：:]\s*[（(]?(\d{4})年\s*0?(\d{1,2})月", text)
        ]
        summary["expected_period_label"] = expected
        summary["document_period_labels"] = found
        if not found or expected not in found or any(value != expected for value in found):
            issues.append(
                ReportQcIssue(
                    "period-mismatch",
                    "error",
                    f"水仙花月报期次与任务不一致，应为{expected}。",
                    f"document_period_labels={found}",
                )
            )

    if expected_image_paths is not None:
        embedded_hashes = _docx_media_hashes(path)
        expected_hashes: list[str] = []
        for item in expected_image_paths:
            image_path = Path(item)
            if not image_path.is_file():
                issues.append(
                    ReportQcIssue(
                        "missing-expected-image",
                        "error",
                        "报告指定的结果图片源文件不存在。",
                        str(image_path),
                    )
                )
                continue
            digest = hashlib.sha256(image_path.read_bytes()).hexdigest().upper()
            expected_hashes.append(digest)
            if digest not in embedded_hashes:
                issues.append(
                    ReportQcIssue(
                        "missing-expected-image",
                        "error",
                        "报告未嵌入指定的本期结果图片。",
                        str(image_path),
                    )
                )
        summary["expected_result_image_count"] = len(expected_hashes)
        summary["matched_result_image_count"] = sum(
            digest in embedded_hashes for digest in expected_hashes
        )
    return _build_result("shuixianhua_monthly", path, issues, summary)


def check_report(kind: str, docx_path: Path | str) -> ReportQcResult:
    if kind == "jlj_monthly":
        return check_jlj_report(docx_path)
    if kind == "hongtang_period":
        return check_hongtang_report(docx_path)
    if kind == "guanbing_monthly":
        return check_guanbing_report(docx_path)
    if kind == "shuixianhua_monthly":
        return check_shuixianhua_report(docx_path)
    raise ValueError(f"Unsupported report kind: {kind}")


def _build_result(kind: str, path: Path, issues: list[ReportQcIssue], summary: dict[str, Any]) -> ReportQcResult:
    warning_count = sum(1 for issue in issues if issue.severity == "warning")
    error_count = sum(1 for issue in issues if issue.severity == "error")
    status = "failed" if error_count else ("warning" if warning_count else "ok")
    return ReportQcResult(
        kind=kind,
        docx_path=str(path),
        checked_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        status=status,
        issue_count=len(issues),
        warning_count=warning_count,
        issues=issues,
        summary=summary,
    )


def result_to_dict(result: ReportQcResult) -> dict[str, Any]:
    payload = asdict(result)
    payload["issues"] = [asdict(issue) for issue in result.issues]
    return payload


def write_report_qc_report(
    result: ReportQcResult,
    output_dir: Path | str,
    *,
    timestamp: str | None = None,
    prefix: str = "report_qc",
) -> tuple[Path, Path]:
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = timestamp or datetime.now().strftime("%Y%m%d_%H%M%S")
    json_path = out_dir / f"{prefix}_{result.kind}_{ts}.json"
    txt_path = out_dir / f"{prefix}_{result.kind}_{ts}.txt"
    payload = result_to_dict(result)
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str) + "\n", encoding="utf-8")

    lines = [
        f"检查时间: {result.checked_at}",
        f"报告类型: {result.kind}",
        f"报告文件: {result.docx_path}",
        f"状态: {result.status}",
        f"问题数量: {result.issue_count}",
        f"表格数量: {result.summary.get('table_count')}",
        f"图片数量: {result.summary.get('image_count')}",
        f"结论页表格索引: {result.summary.get('front_summary_table_indices')}",
    ]
    if result.issues:
        lines.append("")
        lines.append("问题清单:")
        for issue in result.issues:
            detail = f" ({issue.detail})" if issue.detail else ""
            lines.append(f"- [{issue.severity}/{issue.code}] {issue.message}{detail}")
    else:
        lines.append("")
        lines.append("检查通过：未发现已知格式或内容风险。")
    txt_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return txt_path, json_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Run report quality checks.")
    parser.add_argument("--docx", type=Path, required=True, help="Generated report DOCX path.")
    parser.add_argument(
        "--kind",
        choices=["jlj_monthly", "hongtang_period", "guanbing_monthly", "shuixianhua_monthly"],
        default="jlj_monthly",
    )
    parser.add_argument("--output-dir", type=Path, default=None, help="Directory for QC txt/json outputs.")
    parser.add_argument("--strict", action="store_true", help="Exit non-zero when QC status is warning/failed.")
    args = parser.parse_args()

    result = check_report(args.kind, args.docx)
    if args.output_dir:
        txt_path, json_path = write_report_qc_report(result, args.output_dir)
        print(f"QC report: {txt_path}")
        print(f"QC JSON:   {json_path}")
    print(f"QC {result.status}: {args.docx}")
    if result.issues:
        for issue in result.issues:
            print(f"- [{issue.severity}/{issue.code}] {issue.message}")
    if args.strict and result.status != "ok":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
