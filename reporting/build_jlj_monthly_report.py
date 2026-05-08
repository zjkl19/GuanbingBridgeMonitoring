from __future__ import annotations

import argparse
import json
import math
import os
import re
import subprocess
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Callable, Iterable
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from PIL import Image, ImageDraw, ImageFont, ImageOps

from analysis_manifest import analysis_manifest_context, missing_module_summary_items
from artifact_lookup import (
    filename_has_point_token,
    resolve_output_dirs as shared_resolve_output_dirs,
    should_skip_search_dir as shared_should_skip_search_dir,
)
from stats_lookup import resolve_from_analysis_manifest
from docx_utils import set_cell_text_preserve
from jlj_patrol import insert_docx_body_after_heading, period_label_month, resolve_patrol_report_source
from jlj_summary import clear_repeat_table_headers, is_summary_table, update_summary_table
from excel_utils import load_sheet_rows as load_xlsx_rows
from format_utils import (
    format_number,
    format_number_fixed,
    format_range,
    format_table_datetime,
    format_table_number,
    numeric_max,
    numeric_mean,
    numeric_min,
    parse_float,
    safe_text,
    table_cell_text,
)
from missing_summary import write_missing_summary
from report_build_manifest import write_report_build_manifest
from report_context import ReportBuildContext
from report_artifact_resolver import (
    find_latest_image_patterns as lookup_latest_image_patterns,
    find_latest_point_image_patterns as lookup_latest_point_image_patterns,
)
from report_qc import check_jlj_report, write_report_qc_report
from table_utils import (
    set_header_bold,
    set_table_auto_width as shared_set_table_auto_width,
    set_table_autofit as shared_set_table_autofit,
    set_table_column_widths as shared_set_table_column_widths,
    set_table_font_size as shared_set_table_font_size,
    set_table_outer_border as shared_set_table_outer_border,
    set_table_width as shared_set_table_width,
    style_table as shared_style_table,
)
from template_precheck import raise_for_template


JLJ_REPORT_REMOVE_PHRASES = (
    "\u5efa\u8bae\u7ed3\u5408\u539f\u59cb\u6570\u636e\u548c\u4f20\u611f\u5668\u72b6\u6001\u8fdb\u4e00\u6b65\u590d\u6838",
    "\u5efa\u8bae\u7ed3\u5408\u539f\u59cb\u6570\u636e\u8fdb\u4e00\u6b65\u590d\u6838",
    "\u9700\u7ed3\u5408\u539f\u59cb\u6570\u636e\u548c\u73b0\u573a\u8fd0\u7ef4\u8bb0\u5f55\u590d\u6838",
)

JLJ_REPORT_REMOVE_SENTENCES = (
    "\u5f53\u524d\u540a\u6746\u53c2\u6570\u914d\u7f6e\u5c1a\u672a\u5b8c\u6574\u6821\u6838\uff0c\u7d22\u529b\u6362\u7b97\u7ed3\u679c\u6682\u4ec5\u7528\u4e8e\u65f6\u7a0b\u5c55\u793a\u3002",
    "\u5f53\u524d\u540a\u6746\u53c2\u6570\u914d\u7f6e\u5c1a\u672a\u5b8c\u6574\u6821\u6838\uff0c\u7d22\u529b\u6362\u7b97\u7ed3\u679c\u6682\u4ec5\u7528\u4e8e\u65f6\u7a0b\u5c55\u793a",
)

@dataclass
class ImageItem:
    label: str
    path: Path | None


@dataclass
class ParagraphTemplate:
    style_name: str | None
    alignment: WD_ALIGN_PARAGRAPH | None
    left_indent: object
    right_indent: object
    first_line_indent: object
    space_before: object
    space_after: object
    line_spacing: object
    line_spacing_rule: object
    keep_together: object
    keep_with_next: object
    page_break_before: object
    widow_control: object
    run_bold: bool | None
    run_italic: bool | None
    run_underline: object
    run_font_name: str | None
    run_font_size: object


@dataclass
class CaptionTemplates:
    figure_paragraph: Paragraph
    table_paragraph: Paragraph


@dataclass
class SectionContent:
    narrative: str
    summary_sentence: str
    table_title: str | None = None
    table_columns: list[tuple[str, str, float | None]] | None = None
    table_rows: list[dict] | None = None
    figure_title: str | None = None
    image_items: list[ImageItem] | None = None
    table_width_mm: float | None = None
    table_font_size_pt: int | None = None
    available: bool = True
    blocks: list["SectionBlock"] | None = None


@dataclass
class SectionBlock:
    narrative: str = ""
    table_title: str | None = None
    table_columns: list[tuple[str, str, float | None]] | None = None
    table_rows: list[dict] | None = None
    figure_title: str | None = None
    image_items: list[ImageItem] | None = None
    table_width_mm: float | None = None
    table_font_size_pt: int | None = None


def clean_jlj_report_xml_text(docx_path: Path) -> None:
    """Remove review-reminder phrases left in manually adjusted templates."""
    path = Path(docx_path)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=str(path.parent)) as tmp:
        tmp_path = Path(tmp.name)
    try:
        changed = False
        with ZipFile(path, "r") as src, ZipFile(tmp_path, "w", ZIP_DEFLATED) as dst:
            for info in src.infolist():
                data = src.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    text = data.decode("utf-8", errors="ignore")
                    original = text
                    comma = "\uff0c"
                    semicolon = "\uff1b"
                    period = "\u3002"
                    for phrase in JLJ_REPORT_REMOVE_PHRASES:
                        text = text.replace(f"{comma}{phrase}{period}", period)
                        text = text.replace(f"{semicolon}{phrase}{period}", period)
                        text = text.replace(f"{phrase}{period}", "")
                        text = text.replace(phrase, "")
                    for sentence in JLJ_REPORT_REMOVE_SENTENCES:
                        text = text.replace(sentence, "")
                    if text != original:
                        data = text.encode("utf-8")
                        changed = True
                dst.writestr(info, data)
        if changed:
            tmp_path.replace(path)
        else:
            tmp_path.unlink(missing_ok=True)
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise


JLG_REPORT_LIMITS = {
    "wind_10min_avg_mps": 25.0,
    "eq_level2_mps2": 1.5,
    "eq_level3_mps2": 2.55,
    "accel_rms_level2_mps2": 0.315,
    "accel_rms_level3_mps2": 0.5,
    "cable_accel_rms_level2_mps2": 1.0,
    "cable_accel_rms_level3_mps2": 3.0,
}

JLG_FIRST_MODE_FREQ_HZ = 1.26
JLG_CONCRETE_ELASTIC_MODULUS_MPA = 34.5e3
JLG_STEEL_ELASTIC_MODULUS_MPA = 206e3
JLG_MAIN_GIRDER_STRESS_LIMITS_MPA = (-16.2, 17.6)
JLG_ARCH_RIB_STRESS_LIMITS_MPA = (-25.9, 20.9)

JLG_HEALTH_STATUS_MODULES: list[tuple[str, list[str]]] = [
    ("温度监测", ["temperature", "temp_humidity"]),
    ("湿度监测", ["temp_humidity"]),
    ("雨量监测", ["rainfall"]),
    ("主梁挠度监测", ["deflection"]),
    ("支座、梁段纵向位移监测", ["bearing_displacement"]),
    ("结构振动监测", ["acceleration"]),
    ("结构应变监测", ["strain"]),
    ("墩柱倾斜监测", ["tilt"]),
    ("裂缝监测", ["crack"]),
    ("吊杆索力监测", ["cable_force"]),
    ("风向风速监测", ["wind"]),
    ("地震动监测", ["eq"]),
    ("GNSS 位移监测", ["gnss"]),
]


JLG_MONTHLY_SECTIONS: list[tuple[str, str, str, str]] = [
    ("main_env", "主桥环境与作用监测", "温度监测", "4.1.1"),
    ("main_humidity", "主桥环境与作用监测", "湿度监测", "4.1.2"),
    ("main_rainfall", "主桥环境与作用监测", "雨量监测", "4.1.3"),
    ("main_wind", "主桥环境与作用监测", "风向风速监测", "4.1.4"),
    ("main_eq", "主桥环境与作用监测", "地震动监测", "4.1.5"),
    ("main_traffic", "主桥环境与作用监测", "车辆荷载监测", "4.1.6"),
    ("main_deflection", "主桥结构响应与结构变化监测", "主梁挠度监测", "4.2.1"),
    ("main_bearing", "主桥结构响应与结构变化监测", "支座、梁段纵向位移监测", "4.2.2"),
    ("main_gnss", "主桥结构响应与结构变化监测", "拱顶、拱脚位移监测（GNSS）", "4.2.3"),
    ("main_vibration", "主桥结构响应与结构变化监测", "结构振动监测", "4.2.4"),
    ("main_strain", "主桥结构响应与结构变化监测", "结构应变监测", "4.2.5"),
    ("main_crack", "主桥结构响应与结构变化监测", "裂缝监测", "4.2.6"),
    ("main_cable", "主桥结构响应与结构变化监测", "吊杆索力监测", "4.2.7"),
    ("north_strain", "北江滨匝道桥监测", "结构应变监测", "4.3.1"),
    ("north_bearing", "北江滨匝道桥监测", "支座位移监测", "4.3.2"),
    ("north_tilt", "北江滨匝道桥监测", "墩柱倾斜监测", "4.3.3"),
    ("south_strain", "南江滨匝道桥监测", "结构应变监测", "4.4.1"),
    ("south_bearing", "南江滨匝道桥监测", "支座位移监测", "4.4.2"),
    ("south_tilt", "南江滨匝道桥监测", "墩柱倾斜监测", "4.4.3"),
]


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    repo_root = Path(__file__).resolve().parents[1]
    default_template = repo_root / "reports" / "九龙江大桥健康监测2026年3月份月报_0508.docx"
    parser = argparse.ArgumentParser(description="Build Jiulongjiang monthly monitoring report.")
    parser.add_argument("--template", type=Path, default=default_template, help="DOCX template path.")
    parser.add_argument("--config", type=Path, default=repo_root / "config" / "jiulongjiang_config.json", help="Bridge config JSON path.")
    parser.add_argument("--result-root", type=Path, required=True, help="Data/result root containing stats, figures, and run outputs.")
    parser.add_argument("--output-dir", type=Path, default=None, help="Generated report output directory. Defaults to <result-root>/自动报告.")
    parser.add_argument("--period-label", default="2026年3月份", help="Report period label shown in the output filename and report text.")
    parser.add_argument("--monitoring-range", default="2026.03.23~2026.03.31", help="Monitoring range text shown in the report.")
    parser.add_argument("--report-date", default=datetime.now().strftime("%Y年%m月%d日"), help="Report date text shown in the report.")
    parser.add_argument("--image-root", type=Path, default=None, help="Figure lookup root. Defaults to result-root.")
    parser.add_argument("--wim-root", type=Path, default=None, help="Optional WIM result root.")
    parser.add_argument("--patrol-docx", type=Path, default=None, help="Optional patrol report source DOCX. Defaults to reports/九龙江大桥巡查报告.docx lookup.")
    parser.add_argument("--skip-template-precheck", action="store_true", help="Skip DOCX template anchor precheck.")
    return parser.parse_args(argv)


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_sheet_rows(path: Path, sheet: str | None = None) -> list[dict]:
    return load_xlsx_rows(path, sheet, strip_headers=False, skip_empty=False)


def resolve_existing_file(primary_root: Path | None, fallback_root: Path | None, filename: str) -> Path:
    manifest_path = resolve_from_analysis_manifest(primary_root, fallback_root, filename)
    if manifest_path is not None:
        return manifest_path

    candidates: list[Path] = []
    if primary_root is not None:
        candidates.append(primary_root / filename)
        candidates.append(primary_root / "stats" / filename)
    if fallback_root is not None:
        candidates.append(fallback_root / filename)
        candidates.append(fallback_root / "stats" / filename)
    for candidate in candidates:
        if candidate.exists():
            return candidate
    checked = ", ".join(str(p) for p in candidates)
    raise FileNotFoundError(f"Required file not found: {filename}. Checked: {checked}")


def should_skip_search_dir(path: Path) -> bool:
    return shared_should_skip_search_dir(path)


def resolve_output_dirs(root: Path, configured_dir: str) -> list[Path]:
    return shared_resolve_output_dirs(root, configured_dir)


def find_latest_point_image_patterns(root: Path, configured_dir: str, point_id: str, patterns: list[str]) -> Path | None:
    return lookup_latest_point_image_patterns(root, configured_dir, point_id, patterns).path


def find_latest_image_patterns(root: Path, configured_dir: str, patterns: list[str]) -> Path | None:
    return lookup_latest_image_patterns(root, configured_dir, patterns).path



def point_sort_key(point_id: str) -> tuple:
    idx = point_index(point_id)
    parts = re.split(r"(\d+)", point_id)
    normalized: list[object] = []
    for part in parts:
        if not part:
            continue
        if part.isdigit():
            normalized.append(int(part))
        else:
            normalized.append(part)
    return (idx is None, idx if idx is not None else 10**9, normalized)


def sorted_point_ids(point_ids: Iterable[str]) -> list[str]:
    clean = [safe_text(point_id) for point_id in point_ids if safe_text(point_id)]
    return sorted(dict.fromkeys(clean), key=point_sort_key)


def get_config_points(cfg: dict, key: str) -> list[str]:
    points = cfg.get("points", {}).get(key, [])
    if not isinstance(points, list):
        return []
    return sorted_point_ids(points)


def discover_csv_points(result_root: Path, prefixes: Iterable[str]) -> list[str]:
    prefix_tuple = tuple(prefixes)
    found: list[str] = []
    for daily_dir in sorted(result_root.glob("data_jlj_*")):
        csv_dir = daily_dir / "data" / "jlj" / "csv"
        if not csv_dir.exists():
            continue
        for path in csv_dir.glob("*.csv"):
            stem = path.stem
            if stem.startswith(prefix_tuple):
                found.append(stem)
    return sorted_point_ids(found)


def extract_dates_from_range(text: str) -> tuple[date, date] | None:
    pattern = re.compile(r"(\d{4})[年./-](\d{1,2})[月./-](\d{1,2})日?.*?(\d{4})[年./-](\d{1,2})[月./-](\d{1,2})日?")
    match = pattern.search(text)
    if not match:
        return None
    y1, m1, d1, y2, m2, d2 = map(int, match.groups())
    return date(y1, m1, d1), date(y2, m2, d2)


def iter_days(start_date: date, end_date: date) -> Iterable[date]:
    current = start_date
    while current <= end_date:
        yield current
        current += timedelta(days=1)


def format_day_range(start_day: date, end_day: date) -> str:
    if start_day == end_day:
        return start_day.strftime("%Y-%m-%d")
    return f"{start_day:%Y-%m-%d}~{end_day:%Y-%m-%d}"


def group_date_ranges(days: Iterable[date]) -> list[tuple[date, date]]:
    unique_days = sorted(set(days))
    if not unique_days:
        return []
    ranges: list[tuple[date, date]] = []
    start_day = unique_days[0]
    prev_day = unique_days[0]
    for current in unique_days[1:]:
        if current == prev_day + timedelta(days=1):
            prev_day = current
            continue
        ranges.append((start_day, prev_day))
        start_day = current
        prev_day = current
    ranges.append((start_day, prev_day))
    return ranges


def summarize_day_ranges(days: Iterable[date], max_ranges: int = 6) -> str:
    ranges = [format_day_range(start_day, end_day) for start_day, end_day in group_date_ranges(days)]
    if not ranges:
        return "/"
    if len(ranges) <= max_ranges:
        return "；".join(ranges)
    return "；".join(ranges[:max_ranges]) + "；等"


def resolve_jlj_monitoring_dates(monitoring_range: str, result_root: Path) -> tuple[date, date]:
    parsed = extract_dates_from_range(monitoring_range)
    if parsed is not None:
        return parsed
    days: list[date] = []
    for daily_dir in sorted(result_root.glob("data_jlj_*")):
        match = re.search(r"data_jlj_(\d{4})-(\d{2})-(\d{2})$", daily_dir.name)
        if match:
            y, m, d = map(int, match.groups())
            days.append(date(y, m, d))
    if not days:
        raise ValueError(f"Unable to resolve monitoring dates from range: {monitoring_range}")
    return min(days), max(days)


def normalize_jlj_raw_point_id(point_id: str) -> str:
    return re.sub(r"[-_][XYZ]$", "", point_id)


def locate_jlj_csv_dir(result_root: Path, current_day: date) -> Path | None:
    direct = result_root / f"data_jlj_{current_day:%Y-%m-%d}" / "data" / "jlj" / "csv"
    if direct.exists():
        return direct
    legacy = result_root / f"jljData{current_day:%Y%m%d}-{(current_day + timedelta(days=1)):%Y%m%d}" / "data" / "csv"
    if legacy.exists():
        return legacy
    return None


def find_jlj_csv_file(csv_dir: Path, point_id: str) -> Path | None:
    base = normalize_jlj_raw_point_id(point_id)
    candidate = csv_dir / f"{base}.csv"
    if candidate.exists():
        return candidate
    matches = sorted(csv_dir.glob(f"*{base}*.csv"))
    return matches[0] if matches else None


def csv_has_records(path: Path | None) -> bool:
    if path is None or not path.exists() or path.stat().st_size <= 0:
        return False
    with path.open("r", encoding="utf-8", errors="ignore") as fh:
        first_line = True
        for line in fh:
            if first_line:
                first_line = False
                continue
            if line.strip():
                return True
    return False


def collect_jlj_health_status_rows(cfg: dict, result_root: Path, start_date: date, end_date: date) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    points_cfg = cfg.get("points", {})

    for module_label, point_keys in JLG_HEALTH_STATUS_MODULES:
        module_points: list[str] = []
        for point_key in point_keys:
            module_points.extend(get_config_points(cfg, point_key))
        module_points = sorted_point_ids(module_points)
        if not module_points:
            continue

        all_missing_days: list[date] = []
        point_reason_days: dict[tuple[str, str], list[date]] = {}
        for current_day in iter_days(start_date, end_date):
            csv_dir = locate_jlj_csv_dir(result_root, current_day)
            if csv_dir is None:
                all_missing_days.append(current_day)
                continue
            for point_id in module_points:
                csv_path = find_jlj_csv_file(csv_dir, point_id)
                if csv_path is None:
                    point_reason_days.setdefault((point_id, "原始文件缺失"), []).append(current_day)
                    continue
                if not csv_has_records(csv_path):
                    point_reason_days.setdefault((point_id, "无原始记录"), []).append(current_day)

        if all_missing_days:
            rows.append(
                {
                    "module": module_label,
                    "points": "全测点",
                    "range": summarize_day_ranges(all_missing_days),
                    "reason": "原始文件缺失",
                }
            )

        grouped: dict[tuple[tuple[date, ...], str], set[str]] = {}
        for (point_id, reason), days in point_reason_days.items():
            grouped.setdefault((tuple(sorted(set(days))), reason), set()).add(point_id)

        for (days_key, reason), points in sorted(grouped.items(), key=lambda item: (item[0][0][0] if item[0][0] else start_date, point_sort_key(sorted_point_ids(item[1])[0]))):
            rows.append(
                {
                    "module": module_label,
                    "points": "、".join(sorted_point_ids(points)),
                    "range": summarize_day_ranges(days_key),
                    "reason": reason,
                }
            )
    return rows


def truncate_point_list(points: Iterable[str], limit: int = 12) -> str:
    items = sorted_point_ids([point for point in points if safe_text(point)])
    if not items:
        return "/"
    if len(items) <= limit:
        return "、".join(items)
    return "、".join(items[:limit]) + f"等{len(items)}个"


def collect_jlj_data_acquisition_rows(cfg: dict, result_root: Path, start_date: date, end_date: date) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for module_label, point_keys in JLG_HEALTH_STATUS_MODULES:
        design_points: list[str] = []
        for point_key in point_keys:
            design_points.extend(get_config_points(cfg, point_key))
        design_points = sorted_point_ids(design_points)
        if not design_points:
            continue

        acquired: set[str] = set()
        missing_days: list[date] = []
        for current_day in iter_days(start_date, end_date):
            csv_dir = locate_jlj_csv_dir(result_root, current_day)
            if csv_dir is None:
                missing_days.append(current_day)
                continue
            for point_id in design_points:
                if point_id in acquired:
                    continue
                if csv_has_records(find_jlj_csv_file(csv_dir, point_id)):
                    acquired.add(point_id)

        missing_points = [point_id for point_id in design_points if point_id not in acquired]
        design_count = len(design_points)
        acquired_count = len(acquired)
        rate = acquired_count / design_count * 100.0 if design_count else 0.0
        remarks: list[str] = []
        if missing_days:
            remarks.append(f"日期目录缺失：{summarize_day_ranges(missing_days)}")
        if missing_points:
            remarks.append(f"未获取测点：{truncate_point_list(missing_points)}")
        rows.append(
            {
                "module": module_label,
                "design_count": str(design_count),
                "acquired_count": str(acquired_count),
                "rate": f"{rate:.1f}%",
                "missing_points": truncate_point_list(missing_points),
                "remarks": "；".join(remarks) if remarks else "已获取",
            }
        )
    return rows


def summarize_jlj_data_acquisition(rows: list[dict[str, str]]) -> str:
    if not rows:
        return "本月未读取到设计测点配置，监测数据获取情况需人工复核。"
    total_design = sum(int(row.get("design_count", "0") or 0) for row in rows)
    total_acquired = sum(int(row.get("acquired_count", "0") or 0) for row in rows)
    rate = total_acquired / total_design * 100.0 if total_design else 0.0
    incomplete = [
        f"{row['module']}获取率{row['rate']}"
        for row in rows
        if int(row.get("acquired_count", "0") or 0) < int(row.get("design_count", "0") or 0)
    ]
    base = f"本月按监测项目统计设计测点共{total_design}项次，实际获取{total_acquired}项次，整体获取率约{rate:.1f}%。"
    if not incomplete:
        return base + "各监测项目均获取到有效原始记录。"
    return base + "其中" + "、".join(incomplete[:6]) + ("等项目未完全获取。" if len(incomplete) > 6 else "。")


def apply_health_status_section(doc: Document, cfg: dict, result_root: Path, monitoring_range: str) -> None:
    for heading_text in ("监测系统运行状况", "健康监测系统运行状况"):
        try:
            find_heading(doc, heading_text, 2)
            break
        except ValueError:
            continue
    else:
        raise ValueError("Heading not found: 监测系统运行状况 / 健康监测系统运行状况")
    # 2.1 用于人工填写系统运行质量说明；自动程序只校验标题存在，避免覆盖人工判断。
    return


def apply_monthly_data_status_section(
    doc: Document,
    cfg: dict,
    result_root: Path,
    monitoring_range: str,
    caption_templates: CaptionTemplates,
) -> None:
    start_date, end_date = resolve_jlj_monitoring_dates(monitoring_range, result_root)
    section_idx, heading_para = find_heading(doc, "本月监测数据情况", 2)
    next_heading = next_heading_at_or_above(doc, section_idx, 2)
    end_para = next_heading[1] if next_heading is not None else None
    text_template = capture_paragraph_template(find_body_template_paragraph(doc))

    clear_section_between(heading_para, end_para)
    anchor = end_para if end_para is not None else doc.add_paragraph()
    rows = collect_jlj_data_acquisition_rows(cfg, result_root, start_date, end_date)
    if not rows:
        add_text_paragraph_before(anchor, "未读取到设计测点配置，无法生成本月监测数据获取情况统计表。", text_template)
        return

    add_text_paragraph_before(anchor, "本月监测系统数据获取情况如下。设计测点以九龙江配置文件及竣工图测点布置为参考，实际获取情况按监测周期内原始 CSV 文件及有效记录统计。", text_template)
    insert_auto_caption_before(anchor, caption_templates.table_paragraph, "本月监测数据获取情况统计表")
    table = insert_table_before(anchor, rows=len(rows) + 1, cols=6)
    headers = ["监测项目", "设计测点数", "实际获取测点数", "获取率", "未获取测点", "备注"]
    for idx, header in enumerate(headers):
        set_cell_text_preserve(table.cell(0, idx), header)
    for ridx, row in enumerate(rows, start=1):
        values = [row["module"], row["design_count"], row["acquired_count"], row["rate"], row["missing_points"], row["remarks"]]
        for cidx, value in enumerate(values):
            set_cell_text_preserve(table.cell(ridx, cidx), value)
    style_table(table, left=True)
    set_header_bold(table)
    set_repeat_table_header(table)
    set_table_outer_border(table, size_eighth_pt=12)
    set_table_auto_width(table)
    set_table_column_widths(table, [30, 18, 22, 18, 58, 38])
    set_table_font_size(table, 9)


def summarize_first_mode_frequency(stats_root: Path, fallback_root: Path | None = None) -> tuple[str, str, list[str]]:
    try:
        workbook_path = resolve_existing_file(stats_root, fallback_root, "accel_spec_stats.xlsx")
    except FileNotFoundError:
        return "", "", []

    wb = load_workbook(workbook_path, read_only=True, data_only=True)
    point_means: list[float] = []
    all_values: list[float] = []
    point_ids: list[str] = []
    try:
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if len(rows) < 2:
                continue
            header = [safe_text(item) for item in rows[0]]
            first_freq_idx = next((idx for idx, name in enumerate(header) if name.startswith("Freq_")), None)
            if first_freq_idx is None:
                continue
            values = [parse_float(row[first_freq_idx]) for row in rows[1:] if len(row) > first_freq_idx]
            values = [value for value in values if value is not None]
            if not values:
                continue
            point_means.append(sum(values) / len(values))
            all_values.extend(values)
            point_ids.append(ws.title)
    finally:
        wb.close()

    if not all_values:
        return "", "", []

    min_freq = min(all_values)
    max_freq = max(all_values)
    mean_freq = sum(point_means) / len(point_means)
    detail = (
        f"已识别到有效一阶频率的{len(point_means)}个测点中，一阶频率范围约为"
        f"{format_number_fixed(min_freq, 3, 'Hz')}~{format_number_fixed(max_freq, 3, 'Hz')}，"
        f"平均约为{format_number_fixed(mean_freq, 3, 'Hz')}，与一阶自振频率"
        f"{format_number_fixed(JLG_FIRST_MODE_FREQ_HZ, 3, 'Hz')}总体接近。"
    )
    summary = (
        f"已识别测点一阶频率约为{format_number_fixed(min_freq, 3, 'Hz')}~"
        f"{format_number_fixed(max_freq, 3, 'Hz')}，与一阶自振频率"
        f"{format_number_fixed(JLG_FIRST_MODE_FREQ_HZ, 3, 'Hz')}总体接近。"
    )
    return detail, summary, sorted_point_ids(point_ids)


def select_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyhbd.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def build_tile(label: str, image_path: Path | None, tile_size: tuple[int, int]) -> Image.Image:
    canvas = Image.new("RGB", tile_size, "white")
    draw = ImageDraw.Draw(canvas)
    font_label = select_font(28)
    font_hint = select_font(22)
    margin = 16
    if image_path is not None and image_path.exists():
        with Image.open(image_path) as src:
            src = ImageOps.exif_transpose(src.convert("RGB"))
            bbox = (margin, margin, tile_size[0] - margin, tile_size[1] - 56)
            fitted = ImageOps.contain(src, (bbox[2] - bbox[0], bbox[3] - bbox[1]))
            x = bbox[0] + (bbox[2] - bbox[0] - fitted.width) // 2
            y = bbox[1] + (bbox[3] - bbox[1] - fitted.height) // 2
            canvas.paste(fitted, (x, y))
    else:
        draw.rectangle((margin, margin, tile_size[0] - margin, tile_size[1] - 56), outline="#B0B0B0", width=2)
        draw.text((tile_size[0] / 2, tile_size[1] / 2 - 12), "图片缺失", anchor="mm", fill="#808080", font=font_hint)

    draw.text((tile_size[0] / 2, tile_size[1] - 24), label, anchor="mm", fill="black", font=font_label)
    return canvas


def compose_grid(items: list[ImageItem], out_path: Path, cols: int = 2, tile_size: tuple[int, int] = (920, 560)) -> Path:
    rows = math.ceil(len(items) / cols)
    canvas = Image.new("RGB", (tile_size[0] * cols, tile_size[1] * rows), "white")
    for idx, item in enumerate(items):
        row = idx // cols
        col = idx % cols
        tile = build_tile(item.label, item.path, tile_size)
        canvas.paste(tile, (col * tile_size[0], row * tile_size[1]))
    ensure_dir(out_path.parent)
    canvas.save(out_path, quality=92)
    return out_path


def insert_paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_template_table_before(paragraph: Paragraph, template_tbl) -> Table:
    table_xml = deepcopy(template_tbl)
    paragraph._p.addprevious(table_xml)
    return Table(table_xml, paragraph._parent)


def insert_table_before(paragraph: Paragraph, rows: int, cols: int) -> Table:
    body = paragraph._parent
    table = body.add_table(rows=rows, cols=cols, width=Mm(160))
    paragraph._p.addprevious(table._tbl)
    return table


def clear_section_between(start_paragraph: Paragraph, end_paragraph: Paragraph | None) -> None:
    parent = start_paragraph._p.getparent()
    current = start_paragraph._p.getnext()
    end = end_paragraph._p if end_paragraph is not None else None
    while current is not None and current is not end:
        nxt = current.getnext()
        parent.remove(current)
        current = nxt



def element_has_section_break(element) -> bool:
    return any(child.tag == qn("w:sectPr") for child in element.iter())


def capture_section_break_before_heading(doc: Document, heading_text: str, level: int = 1):
    _, heading_para = find_heading(doc, heading_text, level)
    prev = heading_para._p.getprevious()
    if prev is not None and element_has_section_break(prev):
        return deepcopy(prev)
    return None


def ensure_section_break_before_heading(doc: Document, heading_text: str, template_element, level: int = 1) -> None:
    if template_element is None:
        return
    _, heading_para = find_heading(doc, heading_text, level)
    prev = heading_para._p.getprevious()
    if prev is not None and element_has_section_break(prev):
        return
    heading_para._p.addprevious(deepcopy(template_element))


def capture_paragraph_template(paragraph: Paragraph) -> ParagraphTemplate:
    pf = paragraph.paragraph_format
    run0 = paragraph.runs[0] if paragraph.runs else None
    return ParagraphTemplate(
        style_name=paragraph.style.name if paragraph.style is not None else None,
        alignment=paragraph.alignment,
        left_indent=pf.left_indent,
        right_indent=pf.right_indent,
        first_line_indent=pf.first_line_indent,
        space_before=pf.space_before,
        space_after=pf.space_after,
        line_spacing=pf.line_spacing,
        line_spacing_rule=pf.line_spacing_rule,
        keep_together=pf.keep_together,
        keep_with_next=pf.keep_with_next,
        page_break_before=pf.page_break_before,
        widow_control=pf.widow_control,
        run_bold=run0.bold if run0 else None,
        run_italic=run0.italic if run0 else None,
        run_underline=run0.underline if run0 else None,
        run_font_name=run0.font.name if run0 else None,
        run_font_size=run0.font.size if run0 else None,
    )


def apply_paragraph_template(paragraph: Paragraph, template: ParagraphTemplate) -> None:
    if template.style_name:
        paragraph.style = template.style_name
    paragraph.alignment = template.alignment
    pf = paragraph.paragraph_format
    pf.left_indent = template.left_indent
    pf.right_indent = template.right_indent
    pf.first_line_indent = template.first_line_indent
    pf.space_before = template.space_before
    pf.space_after = template.space_after
    pf.line_spacing = template.line_spacing
    pf.line_spacing_rule = template.line_spacing_rule
    pf.keep_together = template.keep_together
    pf.keep_with_next = template.keep_with_next
    pf.page_break_before = template.page_break_before
    pf.widow_control = template.widow_control
    for run in paragraph.runs:
        if template.run_bold is not None:
            run.bold = template.run_bold
        if template.run_italic is not None:
            run.italic = template.run_italic
        if template.run_underline is not None:
            run.underline = template.run_underline
        if template.run_font_name:
            run.font.name = template.run_font_name
        if template.run_font_size:
            run.font.size = template.run_font_size


def add_text_paragraph_before(anchor: Paragraph, text: str, template: ParagraphTemplate) -> Paragraph:
    para = insert_paragraph_before(anchor)
    para.add_run(text)
    apply_paragraph_template(para, template)
    return para



def style_table(table: Table, left: bool = False) -> None:
    shared_style_table(table, left=left, autofit=True, align_center=True)


def set_repeat_table_header(table: Table) -> None:
    if not table.rows:
        return
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")



def set_repeat_headers_for_all_tables(doc: Document) -> None:
    for table in doc.tables:
        if is_summary_table(table):
            clear_repeat_table_headers(table)
            continue
        set_repeat_table_header(table)



def set_table_autofit(table: Table, enabled: bool = True) -> None:
    shared_set_table_autofit(table, enabled)


def set_table_width(table: Table, width_mm: float) -> None:
    shared_set_table_width(table, width_mm)


def set_table_auto_width(table: Table) -> None:
    shared_set_table_auto_width(table)


def set_table_column_widths(table: Table, widths_mm: list[float]) -> None:
    shared_set_table_column_widths(table, widths_mm)


def set_table_outer_border(table: Table, size_eighth_pt: int = 12) -> None:
    shared_set_table_outer_border(table, size_eighth_pt=size_eighth_pt)


def set_table_font_size(table: Table, size_pt: int) -> None:
    shared_set_table_font_size(table, size_pt)


def remove_bookmarks(paragraph_element) -> None:
    for child in list(paragraph_element):
        if child.tag in {qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            paragraph_element.remove(child)


def append_text_run(paragraph_element, text: str) -> None:
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    paragraph_element.append(run)


def build_caption_paragraph(template_paragraph: Paragraph, title: str) -> Paragraph:
    para_xml = deepcopy(template_paragraph._p)
    remove_bookmarks(para_xml)
    end_count = 0
    last_end = None
    for child in para_xml:
        if child.tag != qn("w:r"):
            continue
        fld = child.find(qn("w:fldChar"))
        if fld is not None and fld.get(qn("w:fldCharType")) == "end":
            end_count += 1
            if end_count == 2:
                last_end = child
                break
    if last_end is None:
        raise ValueError("Caption template does not contain expected field sequence.")
    current = last_end.getnext()
    while current is not None:
        nxt = current.getnext()
        para_xml.remove(current)
        current = nxt
    append_text_run(para_xml, f" {title}")
    return Paragraph(para_xml, template_paragraph._parent)


def insert_auto_caption_before(anchor: Paragraph, template_paragraph: Paragraph, title: str) -> Paragraph:
    para = build_caption_paragraph(template_paragraph, title)
    anchor._p.addprevious(para._p)
    return para


def update_fields_with_word(docx_path: Path) -> None:
    if os.environ.get("BMS_NO_WORD") == "1":
        return
    try:
        import pythoncom
        import win32com.client  # type: ignore
    except ImportError:
        venv_python = Path(__file__).resolve().parent / ".venv" / "Scripts" / "python.exe"
        if not venv_python.exists():
            return
        script = f"""
import pythoncom
import win32com.client
p = r'''{str(docx_path)}'''
pythoncom.CoInitialize()
word = None
doc = None
try:
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc = word.Documents.Open(p)
    for story in doc.StoryRanges:
        try:
            story.Fields.Update()
        except Exception:
            pass
    try:
        doc.TablesOfContents(1).Update()
    except Exception:
        pass
    doc.Fields.Update()
    doc.Save()
finally:
    if doc is not None:
        doc.Close(SaveChanges=True)
    if word is not None:
        word.Quit()
    pythoncom.CoUninitialize()
"""
        subprocess.run([str(venv_python), "-c", script], check=False, timeout=180)
        return

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path))
        for story in doc.StoryRanges:
            try:
                story.Fields.Update()
            except Exception:
                pass
        try:
            doc.TablesOfContents(1).Update()
        except Exception:
            pass
        doc.Fields.Update()
        doc.Save()
    finally:
        if doc is not None:
            doc.Close(SaveChanges=True)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def heading_level(paragraph: Paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style else ""
    match = re.match(r"Heading (\d+)", style_name)
    if match:
        return int(match.group(1))
    return None


def find_heading(doc: Document, text: str, level: int, start_idx: int = 0, end_idx: int | None = None) -> tuple[int, Paragraph]:
    paragraphs = doc.paragraphs
    stop = len(paragraphs) if end_idx is None else end_idx
    for idx in range(start_idx, stop):
        para = paragraphs[idx]
        if heading_level(para) == level and para.text.strip() == text:
            return idx, para
    raise ValueError(f"Heading not found: level={level}, text={text}")


def next_heading_at_or_above(doc: Document, index: int, level: int) -> tuple[int, Paragraph] | None:
    paragraphs = doc.paragraphs
    for idx in range(index + 1, len(paragraphs)):
        para = paragraphs[idx]
        para_level = heading_level(para)
        if para_level is not None and para_level <= level:
            return idx, para
    return None


def find_section_anchor(doc: Document, parent_heading: str, child_heading: str) -> tuple[Paragraph, Paragraph | None]:
    parent_idx, _ = find_heading(doc, parent_heading, 2)
    next_parent = next_heading_at_or_above(doc, parent_idx, 2)
    child_end = next_parent[0] if next_parent is not None else len(doc.paragraphs)
    child_idx, child_para = find_heading(doc, child_heading, 3, start_idx=parent_idx + 1, end_idx=child_end)
    next_heading = next_heading_at_or_above(doc, child_idx, 3)
    return child_para, next_heading[1] if next_heading is not None else None


def pick_evenly(items: list[dict], limit: int) -> list[dict]:
    if len(items) <= limit:
        return items
    if limit <= 1:
        return [items[0]]
    seen: set[int] = set()
    selected: list[dict] = []
    for i in range(limit):
        idx = round(i * (len(items) - 1) / (limit - 1))
        if idx in seen:
            continue
        seen.add(idx)
        selected.append(items[idx])
    return selected


def point_index(point_id: str) -> int | None:
    match = re.match(r"^[A-Z]+(?:[A-Z]+)?-(\d+)-", point_id)
    if not match:
        return None
    return int(match.group(1))


def is_main_strain(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 1 <= idx <= 26


def is_main_girder_strain(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 1 <= idx <= 20


def is_arch_rib_strain(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 21 <= idx <= 26


def strain_to_stress_mpa(strain_micro: float | None, elastic_modulus_mpa: float) -> float | None:
    if strain_micro is None:
        return None
    return strain_micro * elastic_modulus_mpa * 1e-6


def summarize_stress_limit_status(
    rows: list[dict],
    predicate: Callable[[str], bool],
    elastic_modulus_mpa: float,
    lower_limit_mpa: float,
    upper_limit_mpa: float,
    subject: str,
    limit_name: str,
) -> tuple[str, str]:
    selected = [row for row in rows if predicate(safe_text(row.get("PointID")))]
    if not selected:
        return "", ""

    stress_mins: list[float] = []
    stress_maxs: list[float] = []
    exceeded: list[tuple[str, float]] = []
    for row in selected:
        point_id = safe_text(row.get("PointID"))
        min_stress = strain_to_stress_mpa(parse_float(row.get("Min")), elastic_modulus_mpa)
        max_stress = strain_to_stress_mpa(parse_float(row.get("Max")), elastic_modulus_mpa)
        if min_stress is not None:
            stress_mins.append(min_stress)
            if min_stress < lower_limit_mpa:
                exceeded.append((point_id, min_stress))
        if max_stress is not None:
            stress_maxs.append(max_stress)
            if max_stress > upper_limit_mpa:
                exceeded.append((point_id, max_stress))

    if not stress_mins and not stress_maxs:
        return "", ""

    stress_min = min(stress_mins) if stress_mins else None
    stress_max = max(stress_maxs) if stress_maxs else None
    stress_range = format_range(stress_min, stress_max, 3, "MPa")
    limit_text = f"{limit_name}（上限{format_number(upper_limit_mpa, 1, 'MPa')}、下限{format_number(lower_limit_mpa, 1, 'MPa')}）"
    if exceeded:
        point_id, stress_value = max(exceeded, key=lambda item: abs(item[1]))
        sentence = (
            f"{subject}按弹性模量换算后应力范围为{stress_range}，"
            f"其中{point_id}换算应力为{format_number(stress_value, 3, 'MPa')}，"
            f"超过{limit_text}。"
        )
        summary = f"{subject}换算应力范围为{stress_range}，超过{limit_name}。"
    else:
        sentence = (
            f"{subject}按弹性模量换算后应力范围为{stress_range}，"
            f"未超过{limit_text}，均处于预警阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        )
        summary = f"{subject}换算应力范围为{stress_range}，未超过{limit_name}。"
    return sentence, summary


def describe_two_level_upper_status(
    value: float | None,
    level2: float,
    level3: float,
    metric_name: str,
    unit: str,
    decimals: int = 3,
) -> str:
    if value is None:
        return "暂不具备预警判定条件。"
    level2_text = format_number_fixed(level2, decimals, unit)
    level3_text = format_number_fixed(level3, decimals, unit)
    if value <= level2:
        return f"未超过二级预警阈值{level2_text}，处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
    if value <= level3:
        return f"超过二级预警阈值{level2_text}，未达到三级预警阈值{level3_text}。"
    return f"超过三级预警阈值{level3_text}。"


def describe_threshold_status(
    value: float | None,
    level: float,
    metric_name: str,
    unit: str,
    *,
    level_name: str = "一级",
    decimals: int = 1,
    direction: str = "upper",
) -> str:
    if value is None:
        return "暂不具备预警判定条件。"
    limit_text = format_number_fixed(level, decimals, unit)
    exceeded = value >= level if direction == "upper" else value <= level
    if exceeded:
        return f"{metric_name}为{format_number_fixed(value, decimals, unit)}，超过{level_name}预警阈值{limit_text}。"
    return f"{metric_name}为{format_number_fixed(value, decimals, unit)}，未超过{level_name}预警阈值{limit_text}，处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"


def describe_two_level_range_status(
    lo: float | None,
    hi: float | None,
    lower1: float | None,
    upper1: float | None,
    lower2: float | None,
    upper2: float | None,
    subject: str,
    unit: str,
    decimals: int,
    first_level_name: str = "一级",
    second_level_name: str = "二级",
) -> str:
    if lo is None or hi is None or lower1 is None or upper1 is None:
        return ""
    if lower2 is not None and upper2 is not None and (lo < lower2 or hi > upper2):
        return (
            f"{subject}超过{second_level_name}预警阈值"
            f"（{format_number_fixed(lower2, decimals, unit)}~{format_number_fixed(upper2, decimals, unit)}）。"
            ""
        )
    if lo < lower1 or hi > upper1:
        return (
            f"{subject}超过{first_level_name}预警阈值"
            f"（{format_number_fixed(lower1, decimals, unit)}~{format_number_fixed(upper1, decimals, unit)}）。"
            ""
        )
    return (
        f"{subject}未超过{first_level_name}预警阈值"
        f"（{format_number_fixed(lower1, decimals, unit)}~{format_number_fixed(upper1, decimals, unit)}），"
        "处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
    )


def describe_two_level_extreme_status(
    lo: float | None,
    hi: float | None,
    lower1: float | None,
    upper1: float | None,
    lower2: float | None,
    upper2: float | None,
    subject: str,
    unit: str,
    decimals: int,
    first_level_name: str = "一级",
    second_level_name: str = "二级",
) -> str:
    if lo is None or hi is None or lower1 is None or upper1 is None:
        return ""

    def range_text(lower: float | None, upper: float | None) -> str:
        if lower is None or upper is None:
            return ""
        return f"（{format_number_fixed(lower, decimals, unit)}~{format_number_fixed(upper, decimals, unit)}）"

    if lower2 is not None and upper2 is not None:
        if hi > upper2:
            return (
                f"{subject}最大值为{format_number_fixed(hi, decimals, unit)}，超过{second_level_name}预警阈值上限"
                f"{format_number_fixed(upper2, decimals, unit)}{range_text(lower2, upper2)}。"
            )
        if lo < lower2:
            return (
                f"{subject}最小值为{format_number_fixed(lo, decimals, unit)}，低于{second_level_name}预警阈值下限"
                f"{format_number_fixed(lower2, decimals, unit)}{range_text(lower2, upper2)}。"
            )
    if hi > upper1:
        return (
            f"{subject}最大值为{format_number_fixed(hi, decimals, unit)}，超过{first_level_name}预警阈值上限"
            f"{format_number_fixed(upper1, decimals, unit)}{range_text(lower1, upper1)}。"
        )
    if lo < lower1:
        return (
            f"{subject}最小值为{format_number_fixed(lo, decimals, unit)}，低于{first_level_name}预警阈值下限"
            f"{format_number_fixed(lower1, decimals, unit)}{range_text(lower1, upper1)}。"
        )
    return (
        f"{subject}未超过{first_level_name}预警阈值{range_text(lower1, upper1)}，"
        "处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
    )


def extract_warn_line_bounds(cfg: dict, style_key: str) -> tuple[float | None, float | None, float | None, float | None]:
    lines = cfg.get("plot_styles", {}).get(style_key, {}).get("warn_lines", [])
    level2: list[float] = []
    level3: list[float] = []
    for item in lines if isinstance(lines, list) else []:
        if not isinstance(item, dict):
            continue
        value = parse_float(item.get("y"))
        label = safe_text(item.get("label"))
        if value is None:
            continue
        if "二级" in label:
            level2.append(value)
        elif "三级" in label:
            level3.append(value)
    return (
        min(level2) if level2 else None,
        max(level2) if level2 else None,
        min(level3) if level3 else None,
        max(level3) if level3 else None,
    )


def describe_range_warn_status(
    lo: float | None,
    hi: float | None,
    lower2: float | None,
    upper2: float | None,
    lower3: float | None,
    upper3: float | None,
    subject: str,
    unit: str,
    decimals: int,
) -> str:
    if lo is None or hi is None or lower2 is None or upper2 is None:
        return ""
    lower2_text = format_number_fixed(lower2, decimals, unit)
    upper2_text = format_number_fixed(upper2, decimals, unit)
    if lower3 is not None and upper3 is not None and (lo < lower3 or hi > upper3):
        return f"{subject}超过三级预警阈值（{format_number_fixed(lower3, decimals, unit)}~{format_number_fixed(upper3, decimals, unit)}）。"
    if lo < lower2 or hi > upper2:
        if lower3 is not None and upper3 is not None:
            return f"{subject}超过二级预警阈值（{lower2_text}~{upper2_text}），未达到三级预警阈值。"
        return f"{subject}超过二级预警阈值（{lower2_text}~{upper2_text}）。"
    return f"{subject}未超过二级预警阈值（{lower2_text}~{upper2_text}），处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"


def is_south_strain(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 27 <= idx <= 42


def is_north_strain(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 43 <= idx <= 50


def is_main_bearing(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 1 <= idx <= 8


def is_south_bearing(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 9 <= idx <= 18


def is_north_bearing(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 19 <= idx <= 24


def is_south_tilt(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 1 <= idx <= 2


def is_north_tilt(point_id: str) -> bool:
    idx = point_index(point_id)
    return idx is not None and 3 <= idx <= 6



def first_nonempty(rows: Iterable[dict], key: str) -> object:
    for row in rows:
        value = row.get(key)
        if value not in (None, ""):
            return value
    return None


def sort_rows_by_point(rows: list[dict]) -> list[dict]:
    return sorted(rows, key=lambda row: safe_text(row.get("PointID")))


def read_stats_rows(stats_root: Path, filename: str, fallback_root: Path | None = None) -> list[dict]:
    path = resolve_existing_file(stats_root, fallback_root, filename)
    return load_sheet_rows(path)


def load_section_rows(
    stats_root: Path,
    fallback_root: Path | None,
    filename: str,
    predicate: Callable[[dict], bool] | None = None,
) -> list[dict]:
    rows = read_stats_rows(stats_root, filename, fallback_root)
    if predicate is not None:
        rows = [row for row in rows if predicate(row)]
    return sort_rows_by_point(rows)


def choose_representative_points(rows: list[dict], limit: int = 6) -> list[dict]:
    compact = [row for row in rows if safe_text(row.get("PointID"))]
    return pick_evenly(compact, limit)


def resolve_expected_points(
    cfg: dict,
    result_root: Path,
    key: str,
    prefixes: Iterable[str],
    predicate: Callable[[str], bool] | None = None,
    fallback_rows: Iterable[dict] | None = None,
) -> list[str]:
    expected = get_config_points(cfg, key)
    expected.extend(discover_csv_points(result_root, prefixes))
    if fallback_rows is not None:
        expected.extend([safe_text(row.get("PointID")) for row in fallback_rows if safe_text(row.get("PointID"))])
    expected = sorted_point_ids(expected)
    if predicate is not None:
        expected = [point_id for point_id in expected if predicate(point_id)]
    return sorted_point_ids(expected)


def build_full_point_table_rows(
    expected_points: list[str],
    actual_rows: list[dict],
    columns: list[tuple[str, str, float | None]],
    formatters: dict[str, Callable[[object], str]] | None = None,
) -> list[dict]:
    formatters = formatters or {}
    row_map = {safe_text(row.get("PointID")): row for row in actual_rows if safe_text(row.get("PointID"))}
    output: list[dict] = []
    for point_id in expected_points:
        src = row_map.get(point_id, {})
        item: dict[str, object] = {}
        for key, _, _ in columns:
            if key == "PointID":
                item[key] = point_id
                continue
            value = src.get(key) if isinstance(src, dict) else None
            if key in formatters:
                item[key] = formatters[key](value)
            else:
                item[key] = "/" if value in (None, "") else value
        output.append(item)
    return output


def build_full_composite_table_rows(
    expected_keys: list[tuple[str, ...]],
    actual_rows: list[dict],
    key_fields: tuple[str, ...],
    columns: list[tuple[str, str, float | None]],
    formatters: dict[str, Callable[[object], str]] | None = None,
) -> list[dict]:
    formatters = formatters or {}
    row_map = {
        tuple(safe_text(row.get(field)) for field in key_fields): row
        for row in actual_rows
        if all(safe_text(row.get(field)) for field in key_fields)
    }
    output: list[dict] = []
    for key_values in expected_keys:
        src = row_map.get(tuple(key_values), {})
        item: dict[str, object] = {}
        for field, value in zip(key_fields, key_values):
            item[field] = value
        for column_key, _, _ in columns:
            if column_key in key_fields:
                continue
            value = src.get(column_key) if isinstance(src, dict) else None
            if column_key in formatters:
                item[column_key] = formatters[column_key](value)
            else:
                item[column_key] = "/" if value in (None, "") else value
        output.append(item)
    return output


def build_generic_table_rows(rows: list[dict], columns: list[tuple[str, str, float | None]], limit: int = 6) -> list[dict]:
    sampled = choose_representative_points(rows, limit)
    output: list[dict] = []
    for row in sampled:
        item = {}
        for key, _, _ in columns:
            item[key] = row.get(key)
        output.append(item)
    return output


def main_temp_rows(rows: list[dict]) -> tuple[list[dict], list[dict]]:
    env = [row for row in rows if safe_text(row.get("PointID")).startswith("JGWD-")]
    structure = [row for row in rows if safe_text(row.get("PointID")).startswith("WDCGQ-")]
    return env, structure


def classify_temperature_point(point_id: str) -> str:
    if point_id.startswith("JGWD-"):
        return "桥面温度"
    if point_id.startswith("WSDJ-"):
        if "-QM-" in point_id:
            return "桥址区环境温度"
        return "主梁内温度"
    if point_id.startswith("WDCGQ-") and "A20" in point_id:
        return "拱肋温度"
    return "主梁箱室温度"


def classify_humidity_point(point_id: str) -> str:
    if "-QM-" in point_id:
        return "桥址区环境相对湿度"
    return "主梁箱室相对湿度"


def numeric_table_formatters(keys: Iterable[str], decimals: int, keep_trailing: bool = False) -> dict[str, Callable[[object], str]]:
    return {key: (lambda value, d=decimals, keep=keep_trailing: format_table_number(value, d, keep)) for key in keys}


def build_missing_section(text: str) -> SectionContent:
    return SectionContent(
        narrative=text,
        summary_sentence=text,
        available=False,
        image_items=[],
        table_rows=[],
        table_columns=[],
    )


def build_numeric_summary_section(
    rows: list[dict],
    narrative_intro: str,
    summary_subject: str,
    min_key: str,
    max_key: str,
    mean_key: str,
    decimals: int,
    unit: str,
    table_title: str,
    table_columns: list[tuple[str, str, float | None]],
    figure_title: str,
    image_items: list[ImageItem],
    table_rows_override: list[dict] | None = None,
    warning_sentence: str | None = None,
) -> SectionContent:
    if not rows:
        return build_missing_section(f"本月未获取到{summary_subject}有效数据。")
    lo = numeric_min(rows, min_key)
    hi = numeric_max(rows, max_key)
    mean_v = numeric_mean(rows, mean_key)
    narrative = (
        f"{narrative_intro}共统计{len(rows)}个测点，监测值范围为"
        f"{format_range(lo, hi, decimals, unit)}，平均值约为{format_number(mean_v, decimals, unit)}。"
    )
    if warning_sentence:
        narrative += warning_sentence
    summary = f"{summary_subject}监测值范围为{format_range(lo, hi, decimals, unit)}。"
    if warning_sentence:
        summary += warning_sentence
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title=table_title,
        table_columns=table_columns,
        table_rows=table_rows_override if table_rows_override is not None else build_generic_table_rows(rows, table_columns),
        figure_title=figure_title,
        image_items=image_items,
    )


def image_for_point(image_root: Path, directory: str, point_id: str, patterns: list[str]) -> Path | None:
    return find_latest_point_image_patterns(image_root, directory, point_id, patterns)


def images_for_point(image_root: Path, directory: str, point_id: str, patterns: list[str]) -> list[Path]:
    resolved_dirs = resolve_output_dirs(image_root, directory)
    matched: list[Path] = []
    for folder in resolved_dirs:
        for pattern in patterns:
            for candidate in folder.glob(pattern):
                if filename_has_point_token(candidate, point_id):
                    matched.append(candidate.resolve())
    return sorted(set(matched), key=lambda p: p.stat().st_mtime, reverse=True)


def find_latest_two_deflection_images(image_root: Path, point_id: str) -> tuple[Path | None, Path | None]:
    matched = images_for_point(image_root, "时程曲线_挠度", point_id, [f"Defl_{point_id}_*.jpg"])
    if not matched:
        return None, None
    filt = matched[0]
    orig = matched[1] if len(matched) > 1 else None
    return orig, filt


def make_image_items(image_root: Path, directory: str, rows: list[dict], patterns_fn: Callable[[str], list[str]], label_fn: Callable[[dict], str], limit: int = 4) -> list[ImageItem]:
    items: list[ImageItem] = []
    for row in choose_representative_points(rows, limit):
        point_id = safe_text(row.get("PointID"))
        if not point_id:
            continue
        items.append(ImageItem(label_fn(row), image_for_point(image_root, directory, point_id, patterns_fn(point_id))))
    return items


def date_value_to_text(value: object) -> str:
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    return safe_text(value)


def build_temperature_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "temp_stats.xlsx")
    if not rows:
        return build_missing_section("本月未获取到主桥温度监测有效数据。")
    deck_rows = [row for row in rows if classify_temperature_point(safe_text(row.get("PointID"))) == "桥面温度"]
    arch_rows = [row for row in rows if classify_temperature_point(safe_text(row.get("PointID"))) == "拱肋温度"]
    box_rows = [row for row in rows if classify_temperature_point(safe_text(row.get("PointID"))) == "主梁箱室温度"]
    inner_rows = [row for row in rows if classify_temperature_point(safe_text(row.get("PointID"))) == "主梁内温度"]
    env_rows = [row for row in rows if classify_temperature_point(safe_text(row.get("PointID"))) == "桥址区环境温度"]
    deck_range = format_range(numeric_min(deck_rows, "Min"), numeric_max(deck_rows, "Max"), 1, "℃") if deck_rows else "--"
    arch_range = format_range(numeric_min(arch_rows, "Min"), numeric_max(arch_rows, "Max"), 1, "℃") if arch_rows else "--"
    box_range = format_range(numeric_min(box_rows, "Min"), numeric_max(box_rows, "Max"), 1, "℃") if box_rows else "--"
    inner_range = format_range(numeric_min(inner_rows, "Min"), numeric_max(inner_rows, "Max"), 1, "℃") if inner_rows else "--"
    env_range = format_range(numeric_min(env_rows, "Min"), numeric_max(env_rows, "Max"), 1, "℃") if env_rows else "--"
    component_rows = arch_rows + box_rows
    temp_status_parts: list[str] = []
    if deck_rows:
        temp_status_parts.append(
            describe_threshold_status(
                numeric_max(deck_rows, "Max"),
                60.0,
                "桥面温度最大值",
                "℃",
                level_name="一级",
                decimals=1,
            )
        )
    if component_rows:
        temp_status_parts.append(
            describe_two_level_extreme_status(
                numeric_min(component_rows, "Min"),
                numeric_max(component_rows, "Max"),
                -7.0,
                54.0,
                None,
                None,
                "混凝土、钢结构构件温度",
                "℃",
                1,
                first_level_name="一级",
            )
        )
    if inner_rows:
        temp_status_parts.append(
            describe_two_level_extreme_status(
                numeric_min(inner_rows, "Min"),
                numeric_max(inner_rows, "Max"),
                0.0,
                40.0,
                -2.0,
                46.0,
                "主梁内温度",
                "℃",
                1,
                first_level_name="一级",
                second_level_name="二级",
            )
        )
    if env_rows:
        temp_status_parts.append(
            describe_two_level_extreme_status(
                numeric_min(env_rows, "Min"),
                numeric_max(env_rows, "Max"),
                0.0,
                40.0,
                -2.0,
                46.0,
                "桥址区环境温度",
                "℃",
                1,
                first_level_name="一级",
                second_level_name="二级",
            )
        )
    temp_status = "".join(part for part in temp_status_parts if part)
    narrative = (
        f"选取典型监测数据进行分析。桥面温度监测范围为{deck_range}，"
        f"拱肋温度监测范围为{arch_range}，主梁箱室温度监测范围为{box_range}，"
        f"主梁内温度监测范围为{inner_range}，桥址区环境温度监测范围为{env_range}。{temp_status}"
    )
    summary = f"桥面、拱肋、主梁箱室、主梁内及桥址区环境温度监测范围分别为{deck_range}、{arch_range}、{box_range}、{inner_range}和{env_range}。{temp_status}"
    columns = [
        ("PointID", "测点编号", None),
        ("Category", "监测类型", None),
        ("Min", "最小值(℃)", None),
        ("Max", "最大值(℃)", None),
        ("Mean", "平均值(℃)", None),
    ]
    expected_points = resolve_expected_points(cfg, result_root, "temperature", ("JGWD-", "WDCGQ-", "WSDJ-"), fallback_rows=rows)
    table_rows = build_full_point_table_rows(
        expected_points,
        rows,
        columns,
        formatters={
            "Category": lambda value: value if value else "/",
            "Min": lambda value: format_table_number(value, 1, True),
            "Max": lambda value: format_table_number(value, 1, True),
            "Mean": lambda value: format_table_number(value, 1, True),
        },
    )
    for row in table_rows:
        row["Category"] = classify_temperature_point(safe_text(row.get("PointID")))
    image_items = make_image_items(
        image_root,
        "时程曲线_温度",
        rows,
        lambda pid: [f"{pid}_*.jpg"],
        lambda row: safe_text(row.get("PointID")),
        limit=4,
    )
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title="主桥温度监测统计表",
        table_columns=columns,
        table_rows=table_rows,
        figure_title="主桥温度监测典型时程曲线图",
        image_items=image_items,
    )


def build_humidity_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "humidity_stats.xlsx")
    columns = [
        ("PointID", "测点编号", None),
        ("Category", "监测类型", None),
        ("Min", "最小值(%)", None),
        ("Max", "最大值(%)", None),
        ("Mean", "平均值(%)", None),
    ]
    if not rows:
        return build_missing_section("本月未获取到主桥湿度监测有效数据。")
    env_rows = [row for row in rows if classify_humidity_point(safe_text(row.get("PointID"))) == "桥址区环境相对湿度"]
    box_rows = [row for row in rows if classify_humidity_point(safe_text(row.get("PointID"))) == "主梁箱室相对湿度"]
    env_range = format_range(numeric_min(env_rows, "Min"), numeric_max(env_rows, "Max"), 1, "%") if env_rows else "--"
    box_range = format_range(numeric_min(box_rows, "Min"), numeric_max(box_rows, "Max"), 1, "%") if box_rows else "--"
    humidity_status = describe_threshold_status(
        numeric_max(rows, "Max"),
        50.0,
        "相对湿度最大值",
        "%",
        level_name="一级",
        decimals=1,
    )
    image_items = make_image_items(
        image_root,
        "时程曲线_湿度",
        rows,
        lambda pid: [f"{pid}_*.jpg"],
        lambda row: safe_text(row.get("PointID")),
        limit=4,
    )
    expected_points = resolve_expected_points(cfg, result_root, "temp_humidity", ("WSDJ-",), fallback_rows=rows)
    table_rows = build_full_point_table_rows(
        expected_points,
        rows,
        columns,
        formatters={
            "Category": lambda value: value if value else "/",
            "Min": lambda value: format_table_number(value, 1, True),
            "Max": lambda value: format_table_number(value, 1, True),
            "Mean": lambda value: format_table_number(value, 1, True),
        },
    )
    for row in table_rows:
        row["Category"] = classify_humidity_point(safe_text(row.get("PointID")))
    return SectionContent(
        narrative=f"选取典型监测数据进行分析。主梁箱室相对湿度监测范围为{box_range}，桥址区环境相对湿度监测范围为{env_range}。{humidity_status}",
        summary_sentence=f"主梁箱室及桥址区环境相对湿度监测范围分别为{box_range}和{env_range}。{humidity_status}",
        table_title="主桥湿度监测统计表",
        table_columns=columns,
        table_rows=table_rows,
        figure_title="主桥湿度监测典型时程曲线图",
        image_items=image_items,
    )


def build_rainfall_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "rainfall_stats.xlsx")
    if not rows:
        return build_missing_section("本月未获取到主桥雨量监测有效数据。")
    row = rows[0]
    total_mm = parse_float(row.get("Total_mm"))
    max_mm_h = parse_float(row.get("Max_mm_h"))
    narrative = (
        f"本月主桥雨量累计约为{format_number(total_mm, 2, 'mm')}，"
        f"最大降雨强度约为{format_number(max_mm_h, 3, 'mm/h')}。"
    )
    summary = f"主桥本月累计降雨量约{format_number(total_mm, 2, 'mm')}，最大降雨强度约{format_number(max_mm_h, 3, 'mm/h')}。"
    columns = [
        ("PointID", "测点编号", None),
        ("StartTime", "起始时间", None),
        ("EndTime", "结束时间", None),
        ("Max_mm_h", "最大降雨强度(mm/h)", None),
        ("Total_mm", "累计降雨量(mm)", None),
    ]
    expected_points = resolve_expected_points(cfg, result_root, "rainfall", ("YLJ-",), fallback_rows=rows)
    table_rows = build_full_point_table_rows(
        expected_points,
        rows,
        columns,
        formatters={
            "StartTime": format_table_datetime,
            "EndTime": format_table_datetime,
            "Max_mm_h": lambda value: format_table_number(value, 3, False),
            "Total_mm": lambda value: format_table_number(value, 2, False),
        },
    )
    path = find_latest_image_patterns(image_root, "时程曲线_雨量", ["Rainfall_*.jpg"])
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title="主桥雨量监测统计表",
        table_columns=columns,
        table_rows=table_rows,
        figure_title="主桥降雨强度典型时程曲线图",
        image_items=[ImageItem("(a) 降雨强度时程", path)] if path else [],
    )


def build_wind_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "wind_stats.xlsx")
    if not rows:
        return build_missing_section("本月未获取到主桥风向风速监测有效数据。")
    summaries: list[dict[str, object]] = []
    for row in rows:
        pid = safe_text(row.get("PointID"))
        txt = image_root / "风速风向结果" / "风玫瑰"
        summary_files = sorted(txt.glob(f"{pid}_windrose_*_summary.txt"))
        summary = {"PointID": pid}
        if summary_files:
            text = summary_files[-1].read_text(encoding="utf-8")
            for line in text.splitlines():
                line = line.strip()
                if "平均风速" in line and "m/s" in line:
                    m = re.search(r"([0-9.]+)\s*m/s", line)
                    if m:
                        summary["MeanSpeed"] = float(m.group(1))
                elif "最大风速" in line and "m/s" in line:
                    m = re.search(r"([0-9.]+)\s*m/s", line)
                    if m:
                        summary["MaxSpeed"] = float(m.group(1))
                elif "主导风向" in line:
                    summary["DominantDir"] = line.split(":", 1)[-1].strip()
                elif "主要风速等级" in line:
                    summary["MainGrade"] = line.split(":", 1)[-1].strip()
        summary["Mean10minMax"] = parse_float(row.get("Mean10minMax"))
        summary["Mean10minTime"] = row.get("Mean10minTime")
        summaries.append(summary)

    deck_point_id = "CSFSY-02-K16-QM-G20"
    deck_summary = next((item for item in summaries if safe_text(item.get("PointID")) == deck_point_id), None)
    source_rows = [deck_summary] if deck_summary is not None else summaries
    max_speed = numeric_max(source_rows, "MaxSpeed")
    mean_speed = numeric_mean(source_rows, "MeanSpeed")
    max_10m = numeric_max(source_rows, "Mean10minMax")
    wind_limit = JLG_REPORT_LIMITS["wind_10min_avg_mps"]
    if max_10m is not None and max_10m <= wind_limit:
        limit_text = f"未超过{format_number_fixed(wind_limit, 0, 'm/s')}，处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
    elif max_10m is not None:
        limit_text = f"已超过{format_number_fixed(wind_limit, 0, 'm/s')}预警阈值，需重点复核。"
    else:
        limit_text = "暂不具备预警判定条件。"
    narrative = (
        f"选取典型监测数据进行分析。主桥桥面平均风速约为{format_number(mean_speed, 2, 'm/s')}，"
        f"瞬时最大风速为{format_number(max_speed, 2, 'm/s')}，"
        f"桥面10min平均风速最大值为{format_number(max_10m, 2, 'm/s')}，{limit_text}"
    )
    summary = f"主桥桥面风速监测中，桥面10min平均风速最大值为{format_number(max_10m, 2, 'm/s')}，{limit_text}"
    columns = [
        ("PointID", "测点编号", 22.0),
        ("DominantDir", "主导风向", 24.0),
        ("MainGrade", "主要风速等级", 22.0),
        ("MeanSpeed", "平均风速(m/s)", 18.0),
        ("MaxSpeed", "最大风速(m/s)", 18.0),
        ("Mean10minMax", "10min平均风速最大值(m/s)", 24.0),
        ("Mean10minTime", "对应时间", 26.0),
    ]
    expected_points = resolve_expected_points(cfg, result_root, "wind", ("CSFSY-",), fallback_rows=summaries)
    table_rows = build_full_point_table_rows(
        expected_points,
        summaries,
        columns,
        formatters={
            "DominantDir": lambda value: safe_text(value, "/") or "/",
            "MainGrade": lambda value: safe_text(value, "/") or "/",
            "MeanSpeed": lambda value: format_table_number(value, 2, False),
            "MaxSpeed": lambda value: format_table_number(value, 2, False),
            "Mean10minMax": lambda value: format_table_number(value, 2, False),
            "Mean10minTime": format_table_datetime,
        },
    )
    image_items: list[ImageItem] = []
    for row in choose_representative_points(summaries, 2):
        pid = safe_text(row.get("PointID"))
        rose = image_for_point(image_root, "风速风向结果/风玫瑰", pid, [f"{pid}_windrose_*.jpg"])
        speed10 = image_for_point(image_root, "风速风向结果/风速10min", pid, [f"{pid}_speed10min_*.jpg"])
        image_items.extend(
            [
                ImageItem(f"{pid} 风玫瑰图", rose),
                ImageItem(f"{pid} 10min平均风速时程", speed10),
            ]
        )
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title="主桥风向风速监测统计表",
        table_columns=columns,
        table_rows=table_rows,
        figure_title="主桥风向风速监测典型图",
        image_items=image_items,
        table_width_mm=154.0,
        table_font_size_pt=9,
    )


def iter_daily_eq_files(result_root: Path) -> list[Path]:
    return sorted(result_root.glob("data_jlj_*/*/jlj/csv/DZY-*.csv"))


def collect_eq_peak_rows(result_root: Path) -> list[dict]:
    import csv

    peaks: dict[tuple[str, str], tuple[float, str]] = {}
    files = iter_daily_eq_files(result_root)
    component_keys = [("X", "value_x"), ("Y", "value_y"), ("Z", "value_z")]
    for path in files:
        point_id = path.stem
        with path.open("r", encoding="utf-8-sig", newline="") as fh:
            reader = csv.DictReader(fh)
            for row in reader:
                ts = row.get("ts", "")
                for comp, key in component_keys:
                    value = parse_float(row.get(key))
                    if value is None:
                        continue
                    abs_value = abs(value)
                    prev = peaks.get((point_id, comp))
                    if prev is None or abs_value > prev[0]:
                        peaks[(point_id, comp)] = (abs_value, ts)
    output: list[dict] = []
    for (point_id, comp), (peak, ts) in sorted(peaks.items()):
        output.append({"PointID": point_id, "Component": comp, "Peak": peak, "PeakTime": ts})
    return output


def build_eq_section(cfg: dict, result_root: Path, image_root: Path) -> SectionContent:
    rows = collect_eq_peak_rows(result_root)
    if not rows:
        return build_missing_section("本月未获取到主桥地震动监测有效数据。")
    horizontal = [row for row in rows if row["Component"] in {"X", "Y"}]
    vertical = [row for row in rows if row["Component"] == "Z"]
    h_peak = numeric_max(horizontal, "Peak")
    v_peak = numeric_max(vertical, "Peak")
    eq_peak = max([value for value in (h_peak, v_peak) if value is not None], default=None)
    eq_status = describe_two_level_upper_status(
        eq_peak,
        JLG_REPORT_LIMITS["eq_level2_mps2"],
        JLG_REPORT_LIMITS["eq_level3_mps2"],
        "地震动加速度峰值",
        "m/s²",
        3,
    )
    narrative = (
        f"水平向地震动加速度峰值为{format_number_fixed(h_peak, 3, 'm/s²')}，"
        f"竖向地震动加速度峰值为{format_number_fixed(v_peak, 3, 'm/s²')}，{eq_status}典型时程见下图。"
    )
    summary = f"主桥地震动监测中，水平向峰值为{format_number_fixed(h_peak, 3, 'm/s²')}，竖向峰值为{format_number_fixed(v_peak, 3, 'm/s²')}，{eq_status}"
    columns = [
        ("PointID", "测点编号", None),
        ("Component", "分量", None),
        ("Peak", "峰值(m/s²)", None),
        ("PeakTime", "对应时间", None),
    ]
    expected_points = sorted_point_ids([normalize_jlj_raw_point_id(point_id) for point_id in get_config_points(cfg, "eq")])
    expected_points.extend(discover_csv_points(result_root, ("DZY-",)))
    expected_points = sorted_point_ids(expected_points)
    expected_keys = [(point_id, component) for point_id in expected_points for component in ("X", "Y", "Z")]
    table_rows = build_full_composite_table_rows(
        expected_keys,
        rows,
        ("PointID", "Component"),
        columns,
        formatters={
            "Peak": lambda value: format_table_number(value, 3, True),
            "PeakTime": format_table_datetime,
        },
    )
    image_items = [
        ImageItem("(a) X向地震动时程", find_latest_image_patterns(image_root, "地震动结果/地震动时程", ["EQ_X_*.jpg"])),
        ImageItem("(b) Y向地震动时程", find_latest_image_patterns(image_root, "地震动结果/地震动时程", ["EQ_Y_*.jpg"])),
        ImageItem("(c) Z向地震动时程", find_latest_image_patterns(image_root, "地震动结果/地震动时程", ["EQ_Z_*.jpg"])),
    ]
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title="主桥地震动监测统计表",
        table_columns=columns,
        table_rows=table_rows,
        figure_title="主桥地震动监测典型时程曲线图",
        image_items=image_items,
        table_width_mm=None,
    )


def build_traffic_section(wim_root: Path | None) -> SectionContent:
    if wim_root is None or not wim_root.exists():
        return build_missing_section("本月未获取到车辆荷载监测结果。")
    narrative = "本月已检测到车辆荷载监测结果目录，后续将接入九龙江月报的车辆荷载统计与图表。"
    return SectionContent(
        narrative=narrative,
        summary_sentence="车辆荷载监测结果目录已就绪，待接入九龙江月报专用统计逻辑。",
        available=False,
    )


def build_deflection_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "deflection_stats.xlsx")
    if not rows:
        return build_missing_section("本月未获取到主桥挠度监测有效数据。")
    expected_points = resolve_expected_points(cfg, result_root, "deflection", ("NDY-",), fallback_rows=rows)
    orig_columns = [
        ("PointID", "测点编号", None),
        ("OrigMin_mm", "最小值(mm)", None),
        ("OrigMax_mm", "最大值(mm)", None),
        ("OrigMean_mm", "平均值(mm)", None),
    ]
    filt_columns = [
        ("PointID", "测点编号", None),
        ("FiltMin_mm", "最小值(mm)", None),
        ("FiltMax_mm", "最大值(mm)", None),
        ("FiltMean_mm", "平均值(mm)", None),
    ]
    orig_table_rows = build_full_point_table_rows(
        expected_points,
        rows,
        orig_columns,
        formatters=numeric_table_formatters(("OrigMin_mm", "OrigMax_mm", "OrigMean_mm"), 1),
    )
    filt_table_rows = build_full_point_table_rows(
        expected_points,
        rows,
        filt_columns,
        formatters=numeric_table_formatters(("FiltMin_mm", "FiltMax_mm", "FiltMean_mm"), 1),
    )
    orig_lo = numeric_min(rows, "OrigMin_mm")
    orig_hi = numeric_max(rows, "OrigMax_mm")
    orig_mean = numeric_mean(rows, "OrigMean_mm")
    filt_lo = numeric_min(rows, "FiltMin_mm")
    filt_hi = numeric_max(rows, "FiltMax_mm")
    filt_mean = numeric_mean(rows, "FiltMean_mm")
    lower2, upper2, lower3, upper3 = extract_warn_line_bounds(cfg, "deflection")
    warn_sentence = describe_range_warn_status(
        filt_lo,
        filt_hi,
        lower2,
        upper2,
        lower3,
        upper3,
        "主桥挠度滤波后监测值",
        "mm",
        1,
    )
    orig_images: list[ImageItem] = []
    filt_images: list[ImageItem] = []
    for row in choose_representative_points(rows, 4):
        pid = safe_text(row.get("PointID"))
        orig_path, filt_path = find_latest_two_deflection_images(image_root, pid)
        orig_images.append(ImageItem(pid, orig_path))
        filt_images.append(ImageItem(pid, filt_path))
    narrative = (
        f"选取典型监测数据进行分析。主桥挠度原始数据监测值范围为{format_range(orig_lo, orig_hi, 1, 'mm')}，"
        f"平均值约为{format_number(orig_mean, 1, 'mm')}；"
        f"滤波后监测值范围为{format_range(filt_lo, filt_hi, 1, 'mm')}，"
        f"平均值约为{format_number(filt_mean, 1, 'mm')}。{warn_sentence}"
    )
    summary = (
        f"主桥挠度原始数据监测值范围为{format_range(orig_lo, orig_hi, 1, 'mm')}；"
        f"滤波后监测值范围为{format_range(filt_lo, filt_hi, 1, 'mm')}。{warn_sentence}"
    )
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        blocks=[
            SectionBlock(
                narrative="原始数据分析结果如下。",
                table_title="主桥挠度原始数据统计表",
                table_columns=orig_columns,
                table_rows=orig_table_rows,
                figure_title="主桥挠度原始数据典型时程曲线图",
                image_items=orig_images,
            ),
            SectionBlock(
                narrative="滤波后数据分析结果如下。",
                table_title="主桥挠度滤波后数据统计表",
                table_columns=filt_columns,
                table_rows=filt_table_rows,
                figure_title="主桥挠度滤波后数据典型时程曲线图",
                image_items=filt_images,
            ),
        ],
    )


def build_main_bearing_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "bearing_displacement_stats.xlsx", lambda row: is_main_bearing(safe_text(row.get("PointID"))))
    columns = [
        ("PointID", "测点编号", None),
        ("FiltMin_mm", "最小值(mm)", None),
        ("FiltMax_mm", "最大值(mm)", None),
        ("FiltMean_mm", "平均值(mm)", None),
    ]
    image_items = make_image_items(
        image_root,
        "时程曲线_支座位移",
        rows,
        lambda pid: [f"BearingDisp_{pid}_*Filt*.jpg", f"BearingDisp_{pid}_*.jpg"],
        lambda row: safe_text(row.get("PointID")),
        limit=4,
    )
    expected_points = resolve_expected_points(cfg, result_root, "bearing_displacement", ("WYJ-",), predicate=is_main_bearing, fallback_rows=rows)
    table_rows = build_full_point_table_rows(
        expected_points,
        rows,
        columns,
        formatters=numeric_table_formatters(("FiltMin_mm", "FiltMax_mm", "FiltMean_mm"), 1),
    )
    bearing_warning = describe_range_warn_status(
        numeric_min(rows, "FiltMin_mm"),
        numeric_max(rows, "FiltMax_mm"),
        -7.1,
        7.1,
        -8.9,
        8.9,
        "主桥支座及梁段纵向位移监测值",
        "mm",
        1,
    )
    return build_numeric_summary_section(
        rows,
        "选取典型监测数据进行分析。主桥支座及梁段纵向位移",
        "主桥支座及梁段纵向位移",
        "FiltMin_mm",
        "FiltMax_mm",
        "FiltMean_mm",
        1,
        "mm",
        "主桥支座、梁段纵向位移监测统计表",
        columns,
        "主桥支座、梁段纵向位移典型时程曲线图",
        image_items,
        table_rows_override=table_rows,
        warning_sentence=bearing_warning,
    )


def build_gnss_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "gnss_stats.xlsx")
    if not rows:
        return build_missing_section("本月未获取到主桥 GNSS 监测有效数据。")
    columns = [
        ("PointID", "测点编号", None),
        ("ComponentLabel", "分量", None),
        ("Min_mm", "最小值(mm)", None),
        ("Max_mm", "最大值(mm)", None),
        ("PeakToPeak_mm", "峰峰值(mm)", None),
    ]
    pp = numeric_max(rows, "PeakToPeak_mm")
    gnss_warning = describe_range_warn_status(
        numeric_min(rows, "Min_mm"),
        numeric_max(rows, "Max_mm"),
        -59.4,
        53.8,
        -74.2,
        67.2,
        "主桥拱顶、拱脚GNSS位移监测值",
        "mm",
        1,
    )
    narrative = f"选取典型监测数据进行分析。主桥拱顶、拱脚 GNSS 位移峰峰值最大约为{format_number(pp, 3, 'mm')}。{gnss_warning}"
    summary = f"主桥拱顶、拱脚 GNSS 位移峰峰值最大约为{format_number(pp, 3, 'mm')}。{gnss_warning}"
    image_items = make_image_items(
        image_root,
        "时程曲线_GNSS",
        [row for row in rows if safe_text(row.get('Component')) == 'X'],
        lambda pid: [f"GNSS_{pid}_*.jpg"],
        lambda row: safe_text(row.get("PointID")),
        limit=4,
    )
    expected_points = resolve_expected_points(cfg, result_root, "gnss", ("GNSS-",), fallback_rows=rows)
    expected_keys = [(point_id, component) for point_id in expected_points for component in ("X", "Y")]
    label_map = {"X": "X向位移", "Y": "Y向位移"}
    table_rows = build_full_composite_table_rows(
        expected_keys,
        rows,
        ("PointID", "Component"),
        [("PointID", "测点编号", None), ("Component", "分量代码", None), ("Min_mm", "最小值(mm)", None), ("Max_mm", "最大值(mm)", None), ("PeakToPeak_mm", "峰峰值(mm)", None)],
        formatters={
            "Min_mm": lambda value: format_table_number(value, 3, False),
            "Max_mm": lambda value: format_table_number(value, 3, False),
            "PeakToPeak_mm": lambda value: format_table_number(value, 3, False),
        },
    )
    normalized_rows = []
    for row in table_rows:
        normalized_rows.append(
            {
                "PointID": row["PointID"],
                "ComponentLabel": label_map.get(safe_text(row.get("Component")), safe_text(row.get("Component"), "/")),
                "Min_mm": row["Min_mm"],
                "Max_mm": row["Max_mm"],
                "PeakToPeak_mm": row["PeakToPeak_mm"],
            }
        )
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title="主桥 GNSS 位移监测统计表",
        table_columns=columns,
        table_rows=normalized_rows,
        figure_title="主桥 GNSS 位移监测典型时程曲线图",
        image_items=image_items,
    )


def build_vibration_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    try:
        rows = load_section_rows(stats_root, fallback_root, "accel_stats.xlsx", lambda row: safe_text(row.get("PointID")) not in {"", "None"})
    except FileNotFoundError:
        rows = []
    first_mode_detail, first_mode_summary, first_mode_points = summarize_first_mode_frequency(stats_root, fallback_root)
    if not rows and not first_mode_detail:
        return build_missing_section("本月未获取到主桥结构振动监测有效数据。")

    narrative_parts: list[str] = []
    summary_parts: list[str] = []
    if rows:
        abs_peak = max(abs(parse_float(row.get("Min")) or 0.0) if abs(parse_float(row.get("Min")) or 0.0) > abs(parse_float(row.get("Max")) or 0.0) else abs(parse_float(row.get("Max")) or 0.0) for row in rows)
        rms_peak = numeric_max(rows, "RMS10minMax")
        rms_status = describe_two_level_upper_status(
            rms_peak,
            JLG_REPORT_LIMITS["accel_rms_level2_mps2"],
            JLG_REPORT_LIMITS["accel_rms_level3_mps2"],
            "主桥振动10min均方根最大值",
            "m/s²",
            3,
        )
        narrative_parts.append(
            f"主桥振动加速度绝对峰值最大约为{format_number(abs_peak, 3, 'm/s²')}，"
            f"10min 均方根最大值约为{format_number(rms_peak, 3, 'm/s²')}，{rms_status}"
        )
        summary_parts.append(
            f"主桥振动加速度绝对峰值最大约为{format_number(abs_peak, 3, 'm/s²')}，"
            f"10min 均方根最大值约为{format_number(rms_peak, 3, 'm/s²')}，{rms_status}"
        )
    if first_mode_detail:
        narrative_parts.append(first_mode_detail)
    if first_mode_summary:
        summary_parts.append(first_mode_summary)
    narrative = "选取典型监测数据进行分析。" + "".join(narrative_parts)
    summary = "".join(summary_parts)
    columns = [
        ("PointID", "测点编号", None),
        ("Min", "最小值(m/s²)", None),
        ("Max", "最大值(m/s²)", None),
        ("RMS10minMax", "10min RMS最大值(m/s²)", None),
    ]
    expected_points = resolve_expected_points(cfg, result_root, "acceleration", ("ZDCQG-",), fallback_rows=rows or [{"PointID": point_id} for point_id in first_mode_points])
    table_rows = (
        build_full_point_table_rows(
            expected_points,
            rows,
            columns,
            formatters=numeric_table_formatters(("Min", "Max", "RMS10minMax"), 3),
        )
        if rows
        else None
    )
    image_items: list[ImageItem] = []
    representative_rows = choose_representative_points(rows, 2) if rows else [{"PointID": point_id} for point_id in pick_evenly([{"PointID": point_id} for point_id in first_mode_points], 2)]
    for row in representative_rows:
        pid = safe_text(row.get("PointID"))
        image_items.append(ImageItem(f"{pid} 振动时程", image_for_point(image_root, "时程曲线_加速度", pid, [f"{pid}_*.jpg"])))
        image_items.append(ImageItem(f"{pid} PSD", image_for_point(image_root, f"PSD_备查/{pid}", pid, [f"PSD_{pid}_*.jpg"])))
        image_items.append(ImageItem(f"{pid} 频谱峰值曲线", image_for_point(image_root, "频谱峰值曲线_加速度", pid, [f"SpecFreq_{pid}_*.jpg"])))
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title="主桥结构振动监测统计表" if table_rows else None,
        table_columns=columns if table_rows else None,
        table_rows=table_rows,
        figure_title="主桥结构振动监测典型图（含 PSD）",
        image_items=image_items,
    )


def build_main_strain_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "strain_stats.xlsx", lambda row: is_main_strain(safe_text(row.get("PointID"))))
    columns = [("PointID", "测点编号", None), ("Min", "最小值(με)", None), ("Max", "最大值(με)", None), ("Mean", "平均值(με)", None)]
    image_items = make_image_items(
        image_root,
        "时程曲线_应变",
        rows,
        lambda pid: [f"Strain_{pid}_*.jpg"],
        lambda row: safe_text(row.get("PointID")),
        limit=4,
    )
    expected_points = resolve_expected_points(cfg, result_root, "strain", ("DYBCGQ-", "DYBCQG-"), predicate=is_main_strain, fallback_rows=rows)
    table_rows = build_full_point_table_rows(
        expected_points,
        rows,
        columns,
        formatters=numeric_table_formatters(("Min", "Max", "Mean"), 3),
    )
    if not rows:
        return build_missing_section("本月未获取到主桥应变有效数据。")

    lo = numeric_min(rows, "Min")
    hi = numeric_max(rows, "Max")
    mean_v = numeric_mean(rows, "Mean")
    girder_sentence, girder_summary = summarize_stress_limit_status(
        rows,
        is_main_girder_strain,
        JLG_CONCRETE_ELASTIC_MODULUS_MPA,
        JLG_MAIN_GIRDER_STRESS_LIMITS_MPA[0],
        JLG_MAIN_GIRDER_STRESS_LIMITS_MPA[1],
        "主梁关键截面应变",
        "主梁跨中静应变二级预警阈值",
    )
    arch_sentence, arch_summary = summarize_stress_limit_status(
        rows,
        is_arch_rib_strain,
        JLG_STEEL_ELASTIC_MODULUS_MPA,
        JLG_ARCH_RIB_STRESS_LIMITS_MPA[0],
        JLG_ARCH_RIB_STRESS_LIMITS_MPA[1],
        "拱肋结构应变",
        "主拱拱顶静应变二级预警阈值",
    )
    stress_parts = [part for part in (girder_sentence, arch_sentence) if part]
    summary_parts = [part for part in (girder_summary, arch_summary) if part]
    narrative = (
        f"选取典型监测数据进行分析。主桥应变共统计{len(rows)}个测点，监测值范围为"
        f"{format_range(lo, hi, 3, 'με')}，平均值约为{format_number(mean_v, 3, 'με')}。"
        + "".join(stress_parts)
    )
    summary = f"主桥应变监测值范围为{format_range(lo, hi, 3, 'με')}。" + "".join(summary_parts)
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title="主桥结构应变监测统计表",
        table_columns=columns,
        table_rows=table_rows,
        figure_title="主桥结构应变监测典型时程曲线图",
        image_items=image_items,
    )


def build_crack_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "crack_stats.xlsx")
    columns = [("PointID", "测点编号", None), ("CrkMin", "最小值(mm)", None), ("CrkMax", "最大值(mm)", None), ("CrkMean", "平均值(mm)", None)]
    image_items = make_image_items(
        image_root,
        "时程曲线_裂缝宽度",
        rows,
        lambda pid: [f"裂缝宽度_{pid}_*.jpg"],
        lambda row: safe_text(row.get("PointID")),
        limit=4,
    )
    if not rows:
        return build_missing_section("本月未获取到主桥裂缝监测有效数据。")
    lo = numeric_min(rows, "CrkMin")
    hi = numeric_max(rows, "CrkMax")
    mean_v = numeric_mean(rows, "CrkMean")
    narrative = (
        f"选取典型监测数据进行分析。主桥裂缝宽度变化范围为{format_range(lo, hi, 3, 'mm')}，"
        f"平均值约为{format_number(mean_v, 3, 'mm')}。"
    )
    summary = f"主桥裂缝宽度变化范围为{format_range(lo, hi, 3, 'mm')}。"
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title="主桥裂缝监测统计表",
        table_columns=columns,
        table_rows=build_full_point_table_rows(
            resolve_expected_points(cfg, result_root, "crack", ("LFJ-",), fallback_rows=rows),
            rows,
            columns,
            formatters=numeric_table_formatters(("CrkMin", "CrkMax", "CrkMean"), 3),
        ),
        figure_title="主桥裂缝监测典型时程曲线图",
        image_items=image_items,
    )


def build_cable_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "cable_accel_stats.xlsx")
    if not rows:
        return build_missing_section("本月未获取到吊杆振动监测有效数据。")
    abs_peak = max(abs(parse_float(row.get("Min")) or 0.0) if abs(parse_float(row.get("Min")) or 0.0) > abs(parse_float(row.get("Max")) or 0.0) else abs(parse_float(row.get("Max")) or 0.0) for row in rows)
    rms_peak = numeric_max(rows, "RMS10minMax")
    rms_status = describe_two_level_upper_status(
        rms_peak,
        JLG_REPORT_LIMITS["cable_accel_rms_level2_mps2"],
        JLG_REPORT_LIMITS["cable_accel_rms_level3_mps2"],
        "吊杆振动10min均方根最大值",
        "m/s²",
        3,
    )
    narrative = (
        f"选取典型监测数据进行分析。吊杆振动加速度绝对峰值最大约为{format_number(abs_peak, 3, 'm/s²')}，"
        f"10min 均方根最大值约为{format_number(rms_peak, 3, 'm/s²')}，{rms_status}"
        ""
    )
    summary = f"吊杆振动加速度绝对峰值最大约为{format_number(abs_peak, 3, 'm/s²')}，10min 均方根最大值约为{format_number(rms_peak, 3, 'm/s²')}，{rms_status}"
    columns = [
        ("PointID", "测点编号", None),
        ("Min", "最小值(m/s²)", None),
        ("Max", "最大值(m/s²)", None),
        ("RMS10minMax", "10min RMS最大值(m/s²)", None),
    ]
    expected_points = resolve_expected_points(cfg, result_root, "cable_force", ("SLCGQ-",), fallback_rows=rows)
    table_rows = build_full_point_table_rows(
        expected_points,
        rows,
        columns,
        formatters=numeric_table_formatters(("Min", "Max", "RMS10minMax"), 3),
    )
    image_items: list[ImageItem] = []
    for row in choose_representative_points(rows, 2):
        pid = safe_text(row.get("PointID"))
        image_items.append(ImageItem(f"{pid} 吊杆加速度时程", image_for_point(image_root, "时程曲线_索力加速度", pid, [f"{pid}_*.jpg"])))
        image_items.append(ImageItem(f"{pid} 索力时程图", image_for_point(image_root, "索力时程图", pid, [f"CableForce_{pid}_*.jpg"])))
    return SectionContent(
        narrative=narrative,
        summary_sentence=summary,
        table_title="吊杆振动监测统计表",
        table_columns=columns,
        table_rows=table_rows,
        figure_title="吊杆振动与索力换算典型图",
        image_items=image_items,
    )


def build_north_strain_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "strain_stats.xlsx", lambda row: is_north_strain(safe_text(row.get("PointID"))))
    columns = [("PointID", "测点编号", None), ("Min", "最小值(με)", None), ("Max", "最大值(με)", None), ("Mean", "平均值(με)", None)]
    image_items = make_image_items(image_root, "时程曲线_应变", rows, lambda pid: [f"Strain_{pid}_*.jpg"], lambda row: safe_text(row.get("PointID")), limit=4)
    table_rows = build_full_point_table_rows(
        resolve_expected_points(cfg, result_root, "strain", ("DYBCGQ-", "DYBCQG-"), predicate=is_north_strain, fallback_rows=rows),
        rows,
        columns,
        formatters=numeric_table_formatters(("Min", "Max", "Mean"), 3),
    )
    return build_numeric_summary_section(
        rows,
        "选取典型监测数据进行分析。北江滨匝道桥结构应变",
        "北江滨匝道桥结构应变",
        "Min",
        "Max",
        "Mean",
        3,
        "με",
        "北江滨匝道桥结构应变监测统计表",
        columns,
        "北江滨匝道桥结构应变典型时程曲线图",
        image_items,
        table_rows_override=table_rows,
    )


def build_south_strain_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "strain_stats.xlsx", lambda row: is_south_strain(safe_text(row.get("PointID"))))
    columns = [("PointID", "测点编号", None), ("Min", "最小值(με)", None), ("Max", "最大值(με)", None), ("Mean", "平均值(με)", None)]
    image_items = make_image_items(image_root, "时程曲线_应变", rows, lambda pid: [f"Strain_{pid}_*.jpg"], lambda row: safe_text(row.get("PointID")), limit=4)
    table_rows = build_full_point_table_rows(
        resolve_expected_points(cfg, result_root, "strain", ("DYBCGQ-", "DYBCQG-"), predicate=is_south_strain, fallback_rows=rows),
        rows,
        columns,
        formatters=numeric_table_formatters(("Min", "Max", "Mean"), 3),
    )
    return build_numeric_summary_section(
        rows,
        "选取典型监测数据进行分析。南江滨匝道桥结构应变",
        "南江滨匝道桥结构应变",
        "Min",
        "Max",
        "Mean",
        3,
        "με",
        "南江滨匝道桥结构应变监测统计表",
        columns,
        "南江滨匝道桥结构应变典型时程曲线图",
        image_items,
        table_rows_override=table_rows,
    )


def build_north_bearing_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "bearing_displacement_stats.xlsx", lambda row: is_north_bearing(safe_text(row.get("PointID"))))
    columns = [("PointID", "测点编号", None), ("FiltMin_mm", "最小值(mm)", None), ("FiltMax_mm", "最大值(mm)", None), ("FiltMean_mm", "平均值(mm)", None)]
    image_items = make_image_items(image_root, "时程曲线_支座位移", rows, lambda pid: [f"BearingDisp_{pid}_*Filt*.jpg", f"BearingDisp_{pid}_*.jpg"], lambda row: safe_text(row.get("PointID")), limit=4)
    table_rows = build_full_point_table_rows(
        resolve_expected_points(cfg, result_root, "bearing_displacement", ("WYJ-",), predicate=is_north_bearing, fallback_rows=rows),
        rows,
        columns,
        formatters=numeric_table_formatters(("FiltMin_mm", "FiltMax_mm", "FiltMean_mm"), 1),
    )
    bearing_warning = describe_range_warn_status(
        numeric_min(rows, "FiltMin_mm"),
        numeric_max(rows, "FiltMax_mm"),
        -7.1,
        7.1,
        -8.9,
        8.9,
        "北江滨匝道桥支座位移监测值",
        "mm",
        1,
    )
    return build_numeric_summary_section(
        rows,
        "选取典型监测数据进行分析。北江滨匝道桥支座位移",
        "北江滨匝道桥支座位移",
        "FiltMin_mm",
        "FiltMax_mm",
        "FiltMean_mm",
        1,
        "mm",
        "北江滨匝道桥支座位移监测统计表",
        columns,
        "北江滨匝道桥支座位移典型时程曲线图",
        image_items,
        table_rows_override=table_rows,
        warning_sentence=bearing_warning,
    )


def build_south_bearing_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "bearing_displacement_stats.xlsx", lambda row: is_south_bearing(safe_text(row.get("PointID"))))
    columns = [("PointID", "测点编号", None), ("FiltMin_mm", "最小值(mm)", None), ("FiltMax_mm", "最大值(mm)", None), ("FiltMean_mm", "平均值(mm)", None)]
    image_items = make_image_items(image_root, "时程曲线_支座位移", rows, lambda pid: [f"BearingDisp_{pid}_*Filt*.jpg", f"BearingDisp_{pid}_*.jpg"], lambda row: safe_text(row.get("PointID")), limit=4)
    table_rows = build_full_point_table_rows(
        resolve_expected_points(cfg, result_root, "bearing_displacement", ("WYJ-",), predicate=is_south_bearing, fallback_rows=rows),
        rows,
        columns,
        formatters=numeric_table_formatters(("FiltMin_mm", "FiltMax_mm", "FiltMean_mm"), 1),
    )
    bearing_warning = describe_range_warn_status(
        numeric_min(rows, "FiltMin_mm"),
        numeric_max(rows, "FiltMax_mm"),
        -7.1,
        7.1,
        -8.9,
        8.9,
        "南江滨匝道桥支座位移监测值",
        "mm",
        1,
    )
    return build_numeric_summary_section(
        rows,
        "选取典型监测数据进行分析。南江滨匝道桥支座位移",
        "南江滨匝道桥支座位移",
        "FiltMin_mm",
        "FiltMax_mm",
        "FiltMean_mm",
        1,
        "mm",
        "南江滨匝道桥支座位移监测统计表",
        columns,
        "南江滨匝道桥支座位移典型时程曲线图",
        image_items,
        table_rows_override=table_rows,
        warning_sentence=bearing_warning,
    )


def build_north_tilt_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "tilt_stats.xlsx", lambda row: is_north_tilt(safe_text(row.get("PointID"))))
    columns = [("PointID", "测点编号", None), ("Min", "最小值(°)", None), ("Max", "最大值(°)", None), ("Mean", "平均值(°)", None)]
    image_items = make_image_items(image_root, "时程曲线_倾角", rows, lambda pid: [f"Tilt_{pid}_*.jpg"], lambda row: safe_text(row.get("PointID")), limit=4)
    table_rows = build_full_point_table_rows(
        resolve_expected_points(cfg, result_root, "tilt", (), predicate=is_north_tilt, fallback_rows=rows),
        rows,
        columns,
        formatters=numeric_table_formatters(("Min", "Max", "Mean"), 3),
    )
    lower2, upper2, lower3, upper3 = extract_warn_line_bounds(cfg, "tilt")
    warning_sentence = describe_range_warn_status(
        numeric_min(rows, "Min"),
        numeric_max(rows, "Max"),
        lower2,
        upper2,
        lower3,
        upper3,
        "北江滨匝道桥墩柱倾斜监测值",
        "°",
        3,
    )
    return build_numeric_summary_section(
        rows,
        "选取典型监测数据进行分析。北江滨匝道桥墩柱倾斜",
        "北江滨匝道桥墩柱倾斜",
        "Min",
        "Max",
        "Mean",
        3,
        "°",
        "北江滨匝道桥墩柱倾斜监测统计表",
        columns,
        "北江滨匝道桥墩柱倾斜典型时程曲线图",
        image_items,
        table_rows_override=table_rows,
        warning_sentence=warning_sentence,
    )


def build_south_tilt_section(cfg: dict, result_root: Path, stats_root: Path, fallback_root: Path | None, image_root: Path) -> SectionContent:
    rows = load_section_rows(stats_root, fallback_root, "tilt_stats.xlsx", lambda row: is_south_tilt(safe_text(row.get("PointID"))))
    columns = [("PointID", "测点编号", None), ("Min", "最小值(°)", None), ("Max", "最大值(°)", None), ("Mean", "平均值(°)", None)]
    image_items = make_image_items(image_root, "时程曲线_倾角", rows, lambda pid: [f"Tilt_{pid}_*.jpg"], lambda row: safe_text(row.get("PointID")), limit=4)
    table_rows = build_full_point_table_rows(
        resolve_expected_points(cfg, result_root, "tilt", (), predicate=is_south_tilt, fallback_rows=rows),
        rows,
        columns,
        formatters=numeric_table_formatters(("Min", "Max", "Mean"), 3),
    )
    lower2, upper2, lower3, upper3 = extract_warn_line_bounds(cfg, "tilt")
    warning_sentence = describe_range_warn_status(
        numeric_min(rows, "Min"),
        numeric_max(rows, "Max"),
        lower2,
        upper2,
        lower3,
        upper3,
        "南江滨匝道桥墩柱倾斜监测值",
        "°",
        3,
    )
    return build_numeric_summary_section(
        rows,
        "选取典型监测数据进行分析。南江滨匝道桥墩柱倾斜",
        "南江滨匝道桥墩柱倾斜",
        "Min",
        "Max",
        "Mean",
        3,
        "°",
        "南江滨匝道桥墩柱倾斜监测统计表",
        columns,
        "南江滨匝道桥墩柱倾斜典型时程曲线图",
        image_items,
        table_rows_override=table_rows,
        warning_sentence=warning_sentence,
    )


def build_section_map(cfg: dict, stats_root: Path, fallback_root: Path | None, result_root: Path, image_root: Path, wim_root: Path | None) -> dict[str, SectionContent]:
    return {
        "main_env": build_temperature_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_humidity": build_humidity_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_rainfall": build_rainfall_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_wind": build_wind_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_eq": build_eq_section(cfg, result_root, image_root),
        "main_traffic": build_traffic_section(wim_root),
        "main_deflection": build_deflection_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_bearing": build_main_bearing_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_gnss": build_gnss_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_vibration": build_vibration_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_strain": build_main_strain_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_crack": build_crack_section(cfg, result_root, stats_root, fallback_root, image_root),
        "main_cable": build_cable_section(cfg, result_root, stats_root, fallback_root, image_root),
        "north_strain": build_north_strain_section(cfg, result_root, stats_root, fallback_root, image_root),
        "north_bearing": build_north_bearing_section(cfg, result_root, stats_root, fallback_root, image_root),
        "north_tilt": build_north_tilt_section(cfg, result_root, stats_root, fallback_root, image_root),
        "south_strain": build_south_strain_section(cfg, result_root, stats_root, fallback_root, image_root),
        "south_bearing": build_south_bearing_section(cfg, result_root, stats_root, fallback_root, image_root),
        "south_tilt": build_south_tilt_section(cfg, result_root, stats_root, fallback_root, image_root),
    }


def find_caption_templates(doc: Document) -> CaptionTemplates:
    figure_para = None
    table_para = None
    figure_token = "SEQ " + "\u56fe"
    table_token = "SEQ " + "\u8868"
    for para in doc.paragraphs:
        fields = "".join(node.text or "" for node in para._p.iter() if node.tag.endswith("}instrText"))
        if figure_para is None and figure_token in fields:
            figure_para = para
        if table_para is None and table_token in fields:
            table_para = para
        if figure_para is not None and table_para is not None:
            break
    if figure_para is None or table_para is None:
        raise ValueError("Unable to locate chapter 1 auto-caption templates.")
    return CaptionTemplates(figure_paragraph=figure_para, table_paragraph=table_para)


def update_cover_metadata(doc: Document, monitoring_range: str, report_date: str) -> None:
    if len(doc.tables) > 1 and len(doc.tables[1].rows) > 1:
        if len(doc.tables[1].rows[1].cells) > 4:
            set_cell_text_preserve(doc.tables[1].rows[1].cells[4], monitoring_range)



def render_section_block(anchor: Paragraph, content: SectionBlock, body_template: ParagraphTemplate, caption_templates: CaptionTemplates) -> None:
    if content.narrative:
        add_text_paragraph_before(anchor, content.narrative, body_template)

    if content.table_title and content.table_columns and content.table_rows:
        insert_auto_caption_before(anchor, caption_templates.table_paragraph, content.table_title)
        table = insert_table_before(anchor, rows=len(content.table_rows) + 1, cols=len(content.table_columns))
        style_table(table)
        for col_idx, (_, label, _) in enumerate(content.table_columns):
            set_cell_text_preserve(table.cell(0, col_idx), label)
        for row_idx, row in enumerate(content.table_rows, start=1):
            for col_idx, (key, _, _) in enumerate(content.table_columns):
                set_cell_text_preserve(table.cell(row_idx, col_idx), table_cell_text(row.get(key)))
        set_header_bold(table)
        set_repeat_table_header(table)
        set_table_outer_border(table, size_eighth_pt=12)
        set_table_autofit(table, True)
        if content.table_width_mm is not None:
            set_table_width(table, content.table_width_mm)
        else:
            set_table_auto_width(table)
        if content.table_width_mm is not None and any(width is not None for _, _, width in content.table_columns):
            set_table_column_widths(table, [width or 26.0 for _, _, width in content.table_columns])
        if content.table_font_size_pt is not None:
            set_table_font_size(table, content.table_font_size_pt)

    valid_images = [item for item in (content.image_items or []) if item.path is not None and item.path.exists()]
    if content.figure_title and valid_images:
        for item in valid_images:
            pic_para = insert_paragraph_before(anchor)
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_para.add_run().add_picture(str(item.path), width=Mm(165))
        insert_auto_caption_before(anchor, caption_templates.figure_paragraph, content.figure_title)


def add_section_content(
    doc: Document,
    parent_heading: str,
    child_heading: str,
    content: SectionContent,
    body_template: ParagraphTemplate,
    caption_templates: CaptionTemplates,
    assets_dir: Path,
) -> None:
    heading_para, next_heading = find_section_anchor(doc, parent_heading, child_heading)
    clear_section_between(heading_para, next_heading)
    anchor = next_heading if next_heading is not None else doc.add_paragraph()

    if content.narrative:
        add_text_paragraph_before(anchor, content.narrative, body_template)

    if content.blocks:
        for block in content.blocks:
            render_section_block(anchor, block, body_template, caption_templates)
        return

    render_section_block(
        anchor,
        SectionBlock(
            table_title=content.table_title,
            table_columns=content.table_columns,
            table_rows=content.table_rows,
            figure_title=content.figure_title,
            image_items=content.image_items,
            table_width_mm=content.table_width_mm,
            table_font_size_pt=content.table_font_size_pt,
        ),
        body_template,
        caption_templates,
    )



def collect_missing_items(section_map: dict[str, SectionContent]) -> list[dict[str, str]]:
    items: list[dict[str, str]] = []
    labels = {
        key: f"{number} {parent} - {child}"
        for key, parent, child, number in JLG_MONTHLY_SECTIONS
    }
    for key, content in section_map.items():
        section_label = labels.get(key, key)
        if not content.available:
            items.append(
                {
                    "category": "章节内容缺失",
                    "section": section_label,
                    "item": key,
                    "detail": content.summary_sentence or content.narrative or "本月未获取到有效数据。",
                    "severity": "warning",
                    "source": key,
                }
            )
            continue

        image_items: list[ImageItem] = []
        if content.image_items:
            image_items.extend(content.image_items)
        for block in content.blocks or []:
            if block.image_items:
                image_items.extend(block.image_items)
        for image_item in image_items:
            if image_item.path is not None and image_item.path.exists():
                continue
            items.append(
                {
                    "category": "图表/资源缺失",
                    "section": section_label,
                    "item": image_item.label,
                    "detail": "报告生成时未找到对应图片。",
                    "severity": "warning",
                    "source": str(image_item.path or ""),
                }
            )
    return items


def find_body_template_paragraph(doc: Document) -> Paragraph:
    preferred_fragments = [
        "选取典型监测数据进行分析",
        "本月未获取到车辆荷载监测结果",
        "监测周期内原始数据缺失",
    ]
    for fragment in preferred_fragments:
        for para in doc.paragraphs:
            if fragment in para.text:
                return para
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and not text.startswith(("图 ", "表 ")) and heading_level(para) is None:
            return para
    return doc.paragraphs[0]


def build_report(
    template: Path,
    config_path: Path,
    result_root: Path,
    image_root: Path | None = None,
    output_dir: Path | None = None,
    wim_root: Path | None = None,
    period_label: str = "2026年3月份",
    monitoring_range: str = "2026.03.23~2026.03.31",
    report_date: str | None = None,
    patrol_docx: Path | None = None,
    precheck_template: bool = True,
) -> Path:
    if report_date is None:
        report_date = datetime.now().strftime("%Y年%m月%d日")

    ctx = ReportBuildContext.from_inputs(
        template=template,
        config_path=config_path,
        result_root=result_root,
        image_root=image_root,
        output_dir=output_dir,
        wim_root=wim_root,
        assets_subdir="generated_assets_jlj",
    )
    cfg = load_json(config_path)

    if precheck_template:
        raise_for_template("jlj_monthly", template)

    doc = Document(str(template))
    caption_templates = find_caption_templates(doc)
    body_template = capture_paragraph_template(find_body_template_paragraph(doc))
    chapter3_section_break = capture_section_break_before_heading(doc, "桥梁人工巡查结果", 1)

    update_cover_metadata(doc, monitoring_range, report_date)
    apply_health_status_section(doc, cfg, result_root, monitoring_range)
    apply_monthly_data_status_section(doc, cfg, result_root, monitoring_range, caption_templates)
    start_date, end_date = resolve_jlj_monitoring_dates(monitoring_range, result_root)
    data_acquisition_summary = summarize_jlj_data_acquisition(collect_jlj_data_acquisition_rows(cfg, result_root, start_date, end_date))
    section_map = build_section_map(cfg, result_root, None, result_root, ctx.image_root, wim_root)
    for key, parent_heading, child_heading, _ in JLG_MONTHLY_SECTIONS:
        add_section_content(doc, parent_heading, child_heading, section_map[key], body_template, caption_templates, ctx.assets_dir)
    ensure_section_break_before_heading(doc, "桥梁人工巡查结果", chapter3_section_break, 1)
    patrol_report_docx = resolve_patrol_report_source(template, patrol_docx)
    if patrol_report_docx is not None:
        insert_docx_body_after_heading(
            doc,
            "桥梁人工巡查结果",
            patrol_report_docx,
            target_month=period_label_month(period_label),
        )
    update_summary_table(doc, section_map, data_acquisition_summary)
    set_repeat_headers_for_all_tables(doc)
    if not any("以下无正文" in para.text for para in doc.paragraphs):
        ending_para = doc.add_paragraph()
        ending_para.add_run("（以下无正文）")
        apply_paragraph_template(ending_para, body_template)
        ending_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_docx = (ctx.output_dir / f"{template.stem}_{period_label}_自动生成_{timestamp}.docx").resolve()
    doc.save(str(output_docx))
    clean_jlj_report_xml_text(output_docx)
    update_fields_with_word(output_docx)
    qc_paths: dict[str, str] = {}
    qc_warnings: list[str] = []
    try:
        qc_result = check_jlj_report(output_docx)
        qc_txt, qc_json = write_report_qc_report(qc_result, ctx.output_dir, timestamp=timestamp)
        qc_paths = {"report_qc_txt": str(qc_txt), "report_qc_json": str(qc_json), "report_qc_status": qc_result.status}
        qc_warnings = [
            f"{issue.code}: {issue.message}"
            for issue in qc_result.issues
            if issue.severity == "warning"
        ]
    except Exception as exc:
        qc_warnings = [f"report_qc_failed: {exc}"]
    missing_items = collect_missing_items(section_map)
    analysis_context = ctx.analysis_context()
    missing_items.extend(missing_module_summary_items(analysis_context))
    manifest_path = write_report_build_manifest(
        context=ctx,
        report_type="jlj_monthly",
        output_docx=output_docx,
        timestamp=timestamp,
        missing=missing_items,
        warnings=qc_warnings,
        extra={"section_keys": list(section_map.keys()), **qc_paths},
        filename_prefix="jlj_report_build_manifest",
    )
    write_missing_summary(
        "九龙江月报",
        output_docx,
        missing_items,
        context={"result_root": str(result_root), "image_root": str(ctx.image_root), "wim_root": str(wim_root or ""), "analysis_manifest": analysis_context.get("path", ""), "report_build_manifest": str(manifest_path)},
    )
    return output_docx


def main() -> None:
    args = parse_args()
    if args.template is None or not args.template.exists():
        raise SystemExit("Template docx not found.")
    if not args.config.exists():
        raise SystemExit("Config file not found.")
    if args.result_root is None:
        raise SystemExit("--result-root is required for Jiulongjiang monthly report.")
    output = build_report(
        template=args.template,
        config_path=args.config,
        result_root=args.result_root,
        image_root=args.image_root,
        output_dir=args.output_dir,
        wim_root=args.wim_root,
        period_label=args.period_label,
        monitoring_range=args.monitoring_range,
        report_date=args.report_date,
        patrol_docx=args.patrol_docx,
        precheck_template=not args.skip_template_precheck,
    )
    print(f"Jiulongjiang monthly report generated: {output}")


if __name__ == "__main__":
    main()
