from __future__ import annotations

import hashlib
import re
import sys
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from PIL import Image, ImageOps


def heading_level(paragraph: Paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style else ""
    match = re.match(r"Heading (\d+)", style_name)
    if match:
        return int(match.group(1))
    return None


def find_heading(doc: Document, text: str, level: int, start_idx: int = 0, end_idx: int | None = None) -> tuple[int, Paragraph]:
    paragraphs = doc.paragraphs
    stop = len(paragraphs) if end_idx is None else end_idx
    for idx in range(start_idx, stop):
        para = paragraphs[idx]
        if heading_level(para) == level and para.text.strip() == text:
            return idx, para
    raise ValueError(f"Heading not found: level={level}, text={text}")


def next_heading_at_or_above(doc: Document, index: int, level: int) -> tuple[int, Paragraph] | None:
    paragraphs = doc.paragraphs
    for idx in range(index + 1, len(paragraphs)):
        para = paragraphs[idx]
        para_level = heading_level(para)
        if para_level is not None and para_level <= level:
            return idx, para
    return None


def clear_section_between(start_paragraph: Paragraph, end_paragraph: Paragraph | None) -> None:
    parent = start_paragraph._p.getparent()
    current = start_paragraph._p.getnext()
    end = end_paragraph._p if end_paragraph is not None else None
    while current is not None and current is not end:
        nxt = current.getnext()
        # The final patrol chapter normally runs to the end of document.xml.
        # Its body-level sectPr is not chapter content: it carries the final
        # section page setup and must remain the last child of w:body.
        if current.tag != qn("w:sectPr"):
            parent.remove(current)
        current = nxt


def replace_patrol_report_dates(text: str, target_month: int = 3) -> str:
    if not text:
        return text
    month_text = str(int(target_month))
    padded_month_text = f"{int(target_month):02d}"
    text = re.sub(r"(\d{4})年0?2月", rf"\1年{month_text}月", text)
    text = re.sub(r"(?<!\d)0?2月", f"{month_text}月", text)
    text = re.sub(
        r"(\d{4})([-/.])0?2([-/.])",
        lambda match: f"{match.group(1)}{match.group(2)}{padded_month_text}{match.group(3)}",
        text,
    )
    return text


def replace_text_nodes_in_element(element, target_month: int = 3) -> None:
    for node in element.iter():
        if node.tag == qn("w:t") and node.text:
            node.text = replace_patrol_report_dates(node.text, target_month)


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def add_page_break_to_paragraph_element(element) -> None:
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    element.insert(0, run)


def paragraph_has_page_break(element) -> bool:
    return any(node.tag == qn("w:br") and node.get(qn("w:type")) == "page" for node in element.iter(qn("w:br")))


def ensure_patrol_attachment_page_break(element) -> None:
    if element.tag != qn("w:p"):
        return
    text = paragraph_text(element).strip()
    if paragraph_has_page_break(element):
        return
    if text.startswith(("附件：", "附件", "附图")) or "缺损照片" in text:
        add_page_break_to_paragraph_element(element)


def ensure_subsequent_patrol_form_page_break(element, form_count: int) -> int:
    if element.tag != qn("w:p"):
        return form_count
    text = paragraph_text(element).strip()
    if not text.startswith("城市桥梁日常巡检报表"):
        return form_count
    if form_count > 0 and not paragraph_has_page_break(element):
        add_page_break_to_paragraph_element(element)
    return form_count + 1


def element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def element_image_count(element) -> int:
    return sum(1 for _ in element.iter(qn("w:drawing"))) + sum(1 for _ in element.iter(qn("w:pict")))


def should_skip_patrol_source_element(element) -> bool:
    if element.tag != qn("w:p"):
        return False
    return not element_text(element).strip() and element_image_count(element) == 0


def set_cell_borders_nil(tc_element) -> None:
    tc_pr = tc_element.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc_element.insert(0, tc_pr)
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def clean_patrol_photo_table_blanks(table_element) -> None:
    if table_element.tag != qn("w:tbl") or element_image_count(table_element) == 0:
        return
    for tc in table_element.iter(qn("w:tc")):
        if element_text(tc).strip() or element_image_count(tc) > 0:
            continue
        set_cell_borders_nil(tc)


def remap_copied_image_relationships(element, source_part, target_part) -> None:
    for node in element.iter():
        for attr_name in (qn("r:embed"), qn("r:link")):
            rel_id = node.get(attr_name)
            if not rel_id or rel_id not in source_part.related_parts:
                continue
            related_part = source_part.related_parts[rel_id]
            if not str(getattr(related_part, "content_type", "")).startswith("image/"):
                continue
            try:
                new_rel_id, _ = target_part.get_or_add_image(BytesIO(related_part.blob))
            except Exception:
                normalized = BytesIO()
                with Image.open(BytesIO(related_part.blob)) as src_img:
                    ImageOps.exif_transpose(src_img.convert("RGB")).save(normalized, format="JPEG", quality=92)
                normalized.seek(0)
                new_rel_id, _ = target_part.get_or_add_image(normalized)
            node.set(attr_name, new_rel_id)


def insert_docx_body_after_heading(
    target_doc: Document,
    heading_text: str,
    source_docx: Path,
    target_month: int = 3,
    *,
    rewrite_source_dates: bool = False,
) -> bool:
    """Replace the patrol chapter with a verified source document.

    Patrol records are factual field records.  Their dates must not be edited
    merely to make an older attachment look current, so date rewriting is off
    by default.  ``rewrite_source_dates`` remains only for explicit legacy
    tooling and is never enabled by the monthly report builder.
    """
    if not source_docx.exists():
        return False
    heading_idx, heading_para = find_heading(target_doc, heading_text, 1)
    next_heading = next_heading_at_or_above(target_doc, heading_idx, 1)
    clear_section_between(heading_para, next_heading[1] if next_heading is not None else None)

    source_doc = Document(str(source_docx))
    insert_after = heading_para._p
    patrol_form_count = 0
    for child in source_doc.element.body:
        if child.tag == qn("w:sectPr"):
            continue
        if should_skip_patrol_source_element(child):
            continue
        new_child = deepcopy(child)
        if rewrite_source_dates:
            replace_text_nodes_in_element(new_child, target_month)
        patrol_form_count = ensure_subsequent_patrol_form_page_break(new_child, patrol_form_count)
        ensure_patrol_attachment_page_break(new_child)
        remap_copied_image_relationships(new_child, source_doc.part, target_doc.part)
        clean_patrol_photo_table_blanks(new_child)
        insert_after.addnext(new_child)
        insert_after = new_child
    return True


def replace_patrol_section_with_note(
    target_doc: Document,
    heading_text: str,
    note: str = "本期巡查资料未提供。",
) -> None:
    """Clear template patrol content and leave an explicit current-period note."""
    heading_idx, heading_para = find_heading(target_doc, heading_text, 1)
    next_heading = next_heading_at_or_above(target_doc, heading_idx, 1)
    clear_section_between(heading_para, next_heading[1] if next_heading is not None else None)
    paragraph = target_doc.add_paragraph(note)
    heading_para._p.addnext(paragraph._p)


def patrol_report_periods(source_docx: Path) -> set[tuple[int, int]]:
    """Return every explicit year/month mentioned by a patrol source."""
    source = Document(str(source_docx))
    texts = [paragraph.text for paragraph in source.paragraphs]
    texts.extend(
        cell.text
        for table in source.tables
        for row in table.rows
        for cell in row.cells
    )
    joined = "\n".join(texts)
    periods = {
        (int(year), int(month))
        for year, month in re.findall(r"(\d{4})年\s*0?(\d{1,2})月", joined)
    }
    periods.update(
        (int(year), int(month))
        for year, month in re.findall(r"(\d{4})[-/.]0?(\d{1,2})[-/.]\d{1,2}", joined)
    )
    return {(year, month) for year, month in periods if 1 <= month <= 12}


def patrol_report_matches_period(source_docx: Path, target_year: int, target_month: int) -> bool:
    periods = patrol_report_periods(source_docx)
    return bool(periods) and periods == {(int(target_year), int(target_month))}


def patrol_source_availability_record(
    source_docx: Path | None,
    *,
    required: bool,
    target_year: int,
    target_month: int,
    action: str,
) -> dict[str, object]:
    """Build an auditable patrol-source availability record."""
    target_period = f"{int(target_year):04d}-{int(target_month):02d}"
    if source_docx is None:
        return {
            "required": bool(required),
            "status": "not_available",
            "target_period": target_period,
            "source": "",
            "source_sha256": "",
            "source_period": "",
            "action": action,
        }

    path = Path(source_docx).resolve()
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    periods = sorted(patrol_report_periods(path))
    source_period = ",".join(f"{year:04d}-{month:02d}" for year, month in periods)
    return {
        "required": bool(required),
        "status": "available",
        "target_period": target_period,
        "source": str(path),
        "source_sha256": digest.hexdigest().upper(),
        "source_period": source_period,
        "action": action,
    }


def resolve_jlj_patrol_report_docx(
    template: Path,
    *,
    target_year: int | None = None,
    target_month: int | None = None,
) -> Path | None:
    names = ["九龙江大桥巡查报告-2026年03月.docx", "九龙江大桥巡查报告.docx"]
    candidates: list[Path] = []
    search_dirs = [
        template.parent,
        Path.cwd() / "reports",
        Path(__file__).resolve().parents[1] / "reports",
    ]
    for name in names:
        candidates.extend(folder / name for folder in search_dirs)
    for folder in search_dirs:
        if folder.is_dir():
            candidates.extend(sorted(folder.glob("九龙江大桥巡查报告*.docx")))
    bundle_root = getattr(sys, "_MEIPASS", None)
    if bundle_root:
        bundle_reports = Path(bundle_root) / "reports"
        for name in names:
            candidates.append(bundle_reports / name)
        if bundle_reports.is_dir():
            candidates.extend(sorted(bundle_reports.glob("九龙江大桥巡查报告*.docx")))
    for candidate in dict.fromkeys(candidates):
        if not candidate.exists():
            continue
        if target_year is not None and target_month is not None and not patrol_report_matches_period(
            candidate, target_year, target_month
        ):
            continue
        return candidate
    return None


def resolve_patrol_report_source(
    template: Path,
    patrol_docx: Path | None = None,
    *,
    target_year: int | None = None,
    target_month: int | None = None,
) -> Path | None:
    if patrol_docx is not None:
        if not patrol_docx.exists():
            raise FileNotFoundError(f"Patrol report source docx not found: {patrol_docx}")
        if target_year is not None and target_month is not None and not patrol_report_matches_period(
            patrol_docx, target_year, target_month
        ):
            periods = sorted(patrol_report_periods(patrol_docx))
            raise ValueError(
                "Patrol report period does not match the report period: "
                f"expected={target_year:04d}-{target_month:02d}, source_periods={periods}, "
                f"source={patrol_docx}"
            )
        return patrol_docx
    return resolve_jlj_patrol_report_docx(
        template,
        target_year=target_year,
        target_month=target_month,
    )


def period_label_month(period_label: str, default_month: int = 3) -> int:
    match = re.search(r"(\d{1,2})月", period_label or "")
    if not match:
        return default_month
    month = int(match.group(1))
    return month if 1 <= month <= 12 else default_month
