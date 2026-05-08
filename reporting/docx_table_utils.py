from __future__ import annotations

import re
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table

from docx_utils import set_cell_text_preserve


def normalize_cell_text(text: Any) -> str:
    """Normalize Word cell text for robust anchor matching."""
    return re.sub(r"\s+", "", "" if text is None else str(text))


def cell_text_matches(text: Any, fragment: str, *, exact: bool = False) -> bool:
    value = normalize_cell_text(text)
    needle = normalize_cell_text(fragment)
    if not needle:
        return False
    return value == needle if exact else needle in value


def center_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def find_table_by_header(doc: Document, header_text: str) -> Table | None:
    """Find the first table containing a header/cell fragment."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell_text_matches(cell.text, header_text):
                    return table
    return None


def find_first_row_index_by_first_cell(table: Table, fragment: str) -> int | None:
    for idx, row in enumerate(table.rows):
        if row.cells and cell_text_matches(row.cells[0].text, fragment):
            return idx
    return None


def table_has_first_cell(table: Table, fragment: str) -> bool:
    return find_first_row_index_by_first_cell(table, fragment) is not None


def find_summary_table(
    doc: Document,
    *,
    result_label: str = "监测结果",
    advice_label: str = "建议",
) -> Table | None:
    """Find the report summary table by left-column semantic labels."""
    for table in doc.tables:
        if len(table.columns) < 2:
            continue
        if table_has_first_cell(table, result_label) and table_has_first_cell(table, advice_label):
            return table
    return find_table_by_header(doc, result_label)


def set_table_cell(table: Table, row_idx: int, col_idx: int, text: Any, *, preserve: bool = True, center: bool = True) -> None:
    cell = table.cell(row_idx, col_idx)
    value = "" if text is None else str(text)
    if preserve:
        set_cell_text_preserve(cell, value)
    else:
        cell.text = value
    if center:
        center_cell(cell)
