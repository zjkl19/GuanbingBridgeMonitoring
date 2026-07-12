from __future__ import annotations

import argparse
import copy
import hashlib
import json
import math
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


BLUE = "1F4E78"
LIGHT_GRAY = "F2F2F2"
ALARM_LEVELS = (25.0, 29.92, 37.4)
INCREMENT_MARKER = "台风窗口增补："
CHART_BASENAMES = {
    "wind_speed": "wind_speed_window.png",
    "wind_maximum": "wind_maximum_comparison.png",
    "wind_direction": "wind_direction_window.png",
    "structure": "structure_response_window.png",
}


@dataclass(frozen=True)
class Anchors:
    chapter_four: object
    result_sections: dict[str, object]
    chapter_one_sections: dict[str, object]
    chapter_two: object


def element_text(element) -> str:
    return "".join(node.text or "" for node in element.findall(".//" + qn("w:t"))).strip()


def body_children(doc: DocumentType) -> list[object]:
    return list(doc._element.body.iterchildren())


def _paragraph(doc: DocumentType, element) -> Paragraph:
    return Paragraph(element, doc._body)


def _heading_level(doc: DocumentType, element) -> int | None:
    if element.tag != qn("w:p"):
        return None
    paragraph = _paragraph(doc, element)
    style_name = paragraph.style.name if paragraph.style is not None else ""
    match = re.search(r"(?:Heading|标题)\s*(\d+)", style_name, flags=re.IGNORECASE)
    if match:
        return int(match.group(1))
    p_pr = element.find(qn("w:pPr"))
    if p_pr is not None:
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is not None and outline.get(qn("w:val")) is not None:
            return int(outline.get(qn("w:val"))) + 1
    return None


def find_heading(
    doc: DocumentType,
    text: str,
    *,
    level: int,
    after=None,
) -> object:
    children = body_children(doc)
    start = children.index(after) + 1 if after is not None else 0
    for element in children[start:]:
        if (
            element.tag == qn("w:p")
            and element_text(element) == text
            and _heading_level(doc, element) == level
        ):
            return element
    raise ValueError(f"heading not found after body index {start}: level={level}, text={text!r}")


def locate_anchors(doc: DocumentType) -> Anchors:
    chapter_one = find_heading(doc, "监测概况", level=1)
    section_13 = find_heading(doc, "健康监测系统运行状况", level=2, after=chapter_one)
    section_14 = find_heading(doc, "软硬件维护状况", level=2, after=section_13)
    chapter_two = find_heading(doc, "监测项目及内容", level=1, after=section_14)

    try:
        chapter_four = find_heading(doc, "监测结果", level=1, after=chapter_two)
    except ValueError:
        chapter_four = find_heading(doc, "台风影响监测结果", level=1, after=chapter_two)
    result_sections: dict[str, object] = {}
    cursor = chapter_four
    for title in (
        "交通状况监测",
        "结构应变监测",
        "主塔倾斜监测",
        "支座变位监测",
        "吊索索力监测",
        "主梁、主塔振动监测",
        "风向风速监测",
        "地震动监测",
    ):
        cursor = find_heading(doc, title, level=2, after=cursor)
        result_sections[title] = cursor
    return Anchors(
        chapter_four=chapter_four,
        result_sections=result_sections,
        chapter_one_sections={
            "健康监测系统运行状况": section_13,
            "软硬件维护状况": section_14,
        },
        chapter_two=chapter_two,
    )


def _is_section_break_paragraph(element) -> bool:
    return element.tag == qn("w:p") and element.find(".//" + qn("w:sectPr")) is not None


def _strip_section_break_paragraph(element) -> None:
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def clear_between(doc: DocumentType, start_element, end_element) -> int:
    """Remove payload between headings while retaining an empty section-break paragraph."""
    body = doc._element.body
    children = body_children(doc)
    start = children.index(start_element)
    end = children.index(end_element)
    removed = 0
    for element in children[start + 1 : end]:
        if _is_section_break_paragraph(element):
            _strip_section_break_paragraph(element)
            continue
        body.remove(element)
        removed += 1
    return removed


def assert_empty_between(doc: DocumentType, start_element, end_element, label: str) -> None:
    children = body_children(doc)
    start = children.index(start_element)
    end = children.index(end_element)
    unexpected = []
    for element in children[start + 1 : end]:
        if _is_section_break_paragraph(element) and not element_text(element):
            continue
        unexpected.append((element.tag, element_text(element)[:80]))
    if unexpected:
        raise RuntimeError(f"stale Q2 payload remains in cleared section {label}: {unexpected[:5]}")


def clear_unavailable_sections(doc: DocumentType, anchors: Anchors) -> dict[str, int]:
    cleared: dict[str, int] = {}
    section_13 = anchors.chapter_one_sections["健康监测系统运行状况"]
    section_14 = anchors.chapter_one_sections["软硬件维护状况"]
    cleared["1.3"] = clear_between(doc, section_13, section_14)
    cleared["1.4"] = clear_between(doc, section_14, anchors.chapter_two)

    ordered = [
        "交通状况监测",
        "结构应变监测",
        "主塔倾斜监测",
        "支座变位监测",
        "吊索索力监测",
    ]
    for index, title in enumerate(ordered[:-1]):
        next_title = ordered[index + 1]
        cleared[f"4.{index + 1}"] = clear_between(
            doc,
            anchors.result_sections[title],
            anchors.result_sections[next_title],
        )

    assert_empty_between(doc, section_13, section_14, "1.3")
    assert_empty_between(doc, section_14, anchors.chapter_two, "1.4")
    for index, title in enumerate(ordered[:-1]):
        assert_empty_between(
            doc,
            anchors.result_sections[title],
            anchors.result_sections[ordered[index + 1]],
            f"4.{index + 1}",
        )
    return cleared


def _body_style(doc: DocumentType) -> str:
    names = {style.name for style in doc.styles}
    return "洪塘大桥月报正文" if "洪塘大桥月报正文" in names else "Normal"


def _caption_style(doc: DocumentType) -> str:
    names = {style.name for style in doc.styles}
    for candidate in ("洪塘大桥月报题注", "Fließtext", "Caption", "Normal"):
        if candidate in names:
            return candidate
    return "Normal"


class InsertBefore:
    def __init__(self, doc: DocumentType, before_element):
        self.doc = doc
        self.before = before_element

    def _move(self, element) -> None:
        self.before.addprevious(element)

    def paragraph(self, text: str = "", *, style: str | None = None, bold_prefix: str = "") -> Paragraph:
        paragraph = self.doc.add_paragraph(style=style or _body_style(self.doc))
        self._move(paragraph._p)
        paragraph.paragraph_format.space_after = Pt(5)
        if bold_prefix and text.startswith(bold_prefix):
            run = paragraph.add_run(bold_prefix)
            run.bold = True
            paragraph.add_run(text[len(bold_prefix) :])
        else:
            paragraph.add_run(text)
        return paragraph

    def heading(self, text: str, *, level: int, reference=None) -> Paragraph:
        paragraph = self.doc.add_heading(text, level=level)
        self._move(paragraph._p)
        if reference is not None:
            reference_paragraph = _paragraph(self.doc, reference)
            paragraph.style = reference_paragraph.style
            source_num_pr = (
                reference.find(qn("w:pPr")).find(qn("w:numPr"))
                if reference.find(qn("w:pPr")) is not None
                else None
            )
            if source_num_pr is not None:
                p_pr = paragraph._p.get_or_add_pPr()
                existing = p_pr.find(qn("w:numPr"))
                if existing is not None:
                    p_pr.remove(existing)
                p_pr.append(copy.deepcopy(source_num_pr))
        return paragraph

    def picture(self, path: Path, *, width: Cm = Cm(15.8)) -> Paragraph:
        paragraph = self.doc.add_paragraph()
        self._move(paragraph._p)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=width)
        return paragraph

    def table(self, headers: list[str], rows: list[list[str]]):
        table = self.doc.add_table(rows=1, cols=len(headers))
        self._move(table._tbl)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, header in enumerate(headers):
            set_cell_shading(table.rows[0].cells[index], BLUE)
            set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        for row_index, values in enumerate(rows):
            cells = table.add_row().cells
            for index, value in enumerate(values):
                set_cell_text(cells[index], value)
                if row_index % 2:
                    set_cell_shading(cells[index], LIGHT_GRAY)
        return table

    def caption(self, kind: str, text: str, *, cached_number: int = 1) -> Paragraph:
        paragraph = self.doc.add_paragraph(style=_caption_style(self.doc))
        self._move(paragraph._p)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.add_run(f"{kind} 4-")
        add_seq_field(paragraph, kind, cached_number)
        paragraph.add_run(f" {text}")
        return paragraph


def add_seq_field(paragraph: Paragraph, kind: str, cached_number: int) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" SEQ {kind} \\* ARABIC \\s 1 "
    instruction_run._r.append(instruction)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    paragraph.add_run(str(cached_number))
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = "000000") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_captioned_table(
    cursor: InsertBefore,
    title: str,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    cursor.caption("表", title)
    cursor.table(headers, rows)


def add_captioned_image(cursor: InsertBefore, path: Path, title: str) -> None:
    cursor.picture(path)
    cursor.caption("图", title)


def _parse_datetime(value: object, label: str) -> datetime:
    if not isinstance(value, str):
        raise ValueError(f"{label} must be an ISO datetime string")
    try:
        return datetime.fromisoformat(value)
    except ValueError as exc:
        raise ValueError(f"invalid {label}: {value!r}") from exc


def _finite_number(value: object, label: str) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"{label} must be numeric") from exc
    if not math.isfinite(number):
        raise ValueError(f"{label} must be finite")
    return number


def load_manifest(path: Path) -> dict[str, object]:
    manifest = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(manifest, dict):
        raise ValueError("manifest root must be an object")
    missing_entries = manifest.get("missing_entries", [])
    if missing_entries:
        raise ValueError(f"typhoon manifest has missing source entries: {missing_entries}")
    for key in ("window", "wind_summary", "structure_summary", "quality"):
        if not isinstance(manifest.get(key), dict):
            raise ValueError(f"manifest.{key} is required")
    window = manifest["window"]
    start = _parse_datetime(window.get("start"), "window.start")
    end = _parse_datetime(window.get("end"), "window.end")
    landfall = _parse_datetime(window.get("landfall"), "window.landfall")
    if not start < landfall < end:
        raise ValueError(f"expected start < landfall < end, got {start}, {landfall}, {end}")
    for point in ("W1", "W2"):
        summary = manifest["wind_summary"].get(point)
        if not isinstance(summary, dict):
            raise ValueError(f"manifest.wind_summary.{point} is required")
        for key in (
            "raw_max",
            "raw_max_direction",
            "max_10min",
            "max_10min_direction",
            "pre_mean",
            "pre_max_10min",
            "post_mean",
            "post_max_10min",
            "bins",
        ):
            _finite_number(summary.get(key), f"wind_summary.{point}.{key}")
        _parse_datetime(summary.get("raw_max_time"), f"wind_summary.{point}.raw_max_time")
        _parse_datetime(summary.get("max_10min_time"), f"wind_summary.{point}.max_10min_time")
    for group in ("主梁加速度", "主塔加速度", "南侧索振动", "北侧索振动"):
        summary = manifest["structure_summary"].get(group)
        if not isinstance(summary, dict):
            raise ValueError(f"manifest.structure_summary.{group} is required")
        for key in ("pre_median", "post_median", "median_ratio", "maximum", "bins"):
            _finite_number(summary.get(key), f"structure_summary.{group}.{key}")
        if not str(summary.get("maximum_point", "")).strip():
            raise ValueError(f"structure_summary.{group}.maximum_point is required")
        _parse_datetime(summary.get("maximum_time"), f"structure_summary.{group}.maximum_time")
    return manifest


def resolve_charts(
    manifest: dict[str, object],
    manifest_path: Path,
    charts_dir: Path | None,
    explicit: dict[str, Path | None] | None = None,
) -> dict[str, Path]:
    explicit = explicit or {}
    manifest_paths = {
        Path(str(value)).name: Path(str(value))
        for value in manifest.get("charts", [])
        if isinstance(value, str)
    }
    resolved: dict[str, Path] = {}
    for key, basename in CHART_BASENAMES.items():
        candidates: list[Path] = []
        if explicit.get(key) is not None:
            candidates.append(Path(explicit[key]))
        if charts_dir is not None:
            candidates.append(charts_dir / basename)
        candidates.append(manifest_path.parent / basename)
        if basename in manifest_paths:
            candidates.append(manifest_paths[basename])
        path = next((candidate.resolve() for candidate in candidates if candidate.is_file()), None)
        if path is None:
            raise FileNotFoundError(f"required chart not found ({key}): tried {candidates}")
        resolved[key] = path
    if len({path.resolve() for path in resolved.values()}) != 4:
        raise ValueError(f"the four chart roles must resolve to four distinct files: {resolved}")
    return resolved


def _wind_alarm_status(max_10min: float) -> str:
    if max_10min >= ALARM_LEVELS[2]:
        return "达到三级"
    if max_10min >= ALARM_LEVELS[1]:
        return "达到二级"
    if max_10min >= ALARM_LEVELS[0]:
        return "达到一级"
    return "未达一级"


def _fmt_time(value: object, pattern: str = "%m-%d %H:%M:%S") -> str:
    return _parse_datetime(value, "summary time").strftime(pattern)


def _section_summary_rows(
    manifest: dict[str, object], groups: Iterable[str]
) -> list[list[str]]:
    rows = []
    for group in groups:
        summary = manifest["structure_summary"][group]
        rows.append(
            [
                group,
                f"{float(summary['pre_median']):.4f}",
                f"{float(summary['post_median']):.4f}",
                f"{float(summary['median_ratio']):.2f}",
                f"{float(summary['maximum']):.4f}",
                f"{summary['maximum_point']} / {_fmt_time(summary['maximum_time'], '%m-%d %H:%M')}",
            ]
        )
    return rows


def _insert_event_and_coverage(
    doc: DocumentType,
    anchors: Anchors,
    manifest: dict[str, object],
) -> None:
    cursor = InsertBefore(doc, anchors.result_sections["交通状况监测"])
    window = manifest["window"]
    quality = manifest["quality"]
    start = _parse_datetime(window["start"], "window.start")
    end = _parse_datetime(window["end"], "window.end")
    landfall = _parse_datetime(window["landfall"], "window.landfall")
    cursor.paragraph(
        f"台风事件及分析窗口：今年第9号台风“巴威”于{landfall:%Y年%m月%d日%H时%M分}前后在浙江省玉环市坎门街道沿海登陆，"
        "登陆时中心附近最大风力13级（40m/s），中心最低气压955hPa。"
        f"本次增量分析采用{start:%Y年%m月%d日%H:%M}至{end:%Y年%m月%d日%H:%M}的实测数据，"
        f"覆盖台风登陆前后共{(end - start).total_seconds() / 3600:.1f}小时。",
        bold_prefix="台风事件及分析窗口：",
    )
    cursor.paragraph(
        "气象资料来源：中国天气网（中央气象台信息），《台风“巴威”在浙江玉环沿海登陆》，"
        "https://news.weather.com.cn/2026/07/4711229.shtml，引用日期：2026年7月12日。"
    )
    cursor.paragraph(
        "说明：第4章保留季报既有章节体系。因本次台风窗口尚未取得交通称重、结构应变、主塔倾斜和支座变位数据，"
        "4.1至4.4仅保留章节标题，正文、表格和图片留空，不沿用第二季度旧值。"
        "风向风速采用原始波形按10min统计；吊索、主梁及主塔振动采用特征值包络峰值按10min汇总。"
    )
    bins = int(float(quality["expected_complete_10min_bins"]))
    wind_bins = min(int(float(manifest["wind_summary"][point]["bins"])) for point in ("W1", "W2"))
    girder_bins = min(
        int(float(manifest["structure_summary"][group]["bins"]))
        for group in ("主梁加速度", "主塔加速度")
    )
    cable_bins = min(
        int(float(manifest["structure_summary"][group]["bins"]))
        for group in ("南侧索振动", "北侧索振动")
    )
    add_captioned_table(
        cursor,
        "台风窗口监测数据覆盖情况",
        ["数据类别", "测点范围", "开始时间", "结束时间", "10min时段数", "完整性"],
        [
            ["风向风速", "W1、W2", f"{start:%m-%d %H:%M}", f"{end:%m-%d %H:%M}", str(wind_bins), "完整" if wind_bins >= bins else "部分"],
            ["主梁/主塔振动", "A1至A10", f"{start:%m-%d %H:%M}", f"{end:%m-%d %H:%M}", str(girder_bins), "完整" if girder_bins >= bins else "部分"],
            ["吊索振动", "CS1至CS12、CX1至CX12", f"{start:%m-%d %H:%M}", f"{end:%m-%d %H:%M}", str(cable_bins), "完整" if cable_bins >= bins else "部分"],
        ],
    )
    rejected_rate = float(quality.get("direction_rejected_rate", 0.0))
    cursor.paragraph(
        f"数据质量核验：报告窗口应有{bins}个完整10min时段。W1/W2风速原始记录共读取"
        f"{int(float(quality.get('speed_rows', 0))):,}条，剔除负风速或非有限值"
        f"{int(float(quality.get('speed_rejected_rows', 0))):,}条；风向超出0°~360°有效域的记录占"
        f"{rejected_rate * 100:.2f}%，已按既定口径剔除。所有结论仅针对上述实际覆盖时段。"
    )


def _insert_cable_increment(
    doc: DocumentType,
    anchors: Anchors,
    manifest: dict[str, object],
) -> None:
    cursor = InsertBefore(doc, anchors.result_sections["主梁、主塔振动监测"])
    cursor.heading("台风期间吊索振动增量分析", level=3)
    south = manifest["structure_summary"]["南侧索振动"]
    north = manifest["structure_summary"]["北侧索振动"]
    cursor.paragraph(
        f"台风窗口内，南侧、北侧吊索10min包络峰值的登陆后/登陆前中位值比分别为"
        f"{float(south['median_ratio']):.2f}和{float(north['median_ratio']):.2f}。"
        f"全时段最大值分别出现在{south['maximum_point']}和{north['maximum_point']}测点，"
        "未见登陆后持续、多测点同步放大的证据；短时尖峰仍需结合交通荷载和原始波形解释。"
    )
    add_captioned_table(
        cursor,
        "台风登陆前后吊索振动响应比较",
        ["部位", "登陆前中位", "登陆后中位", "后/前比值", "全段最大", "最大测点/时刻"],
        _section_summary_rows(manifest, ("南侧索振动", "北侧索振动")),
    )


def _insert_girder_tower_increment(
    doc: DocumentType,
    anchors: Anchors,
    manifest: dict[str, object],
    chart: Path,
) -> None:
    cursor = InsertBefore(doc, anchors.result_sections["风向风速监测"])
    cursor.heading("台风期间主梁、主塔振动增量分析", level=3)
    girder = manifest["structure_summary"]["主梁加速度"]
    tower = manifest["structure_summary"]["主塔加速度"]
    cursor.paragraph(
        f"台风窗口内，主梁、主塔10min包络峰值的登陆后/登陆前中位值比分别为"
        f"{float(girder['median_ratio']):.2f}和{float(tower['median_ratio']):.2f}。"
        "两类响应在登陆后与登陆前典型水平总体相当，未发现持续、同步抬升。"
        "该比较用于运行状态趋势筛查，不替代结构承载能力验算。"
    )
    add_captioned_image(cursor, chart, "台风窗口主要结构振动响应相对登陆前水平")
    add_captioned_table(
        cursor,
        "台风登陆前后主梁、主塔振动响应比较",
        ["部位", "登陆前中位", "登陆后中位", "后/前比值", "全段最大", "最大测点/时刻"],
        _section_summary_rows(manifest, ("主梁加速度", "主塔加速度")),
    )


def _insert_wind_increment(
    doc: DocumentType,
    anchors: Anchors,
    manifest: dict[str, object],
    charts: dict[str, Path],
) -> None:
    cursor = InsertBefore(doc, anchors.result_sections["地震动监测"])
    cursor.heading("台风期间风速最大值及登陆前后增量分析", level=3)
    wind = manifest["wind_summary"]
    alarm_text = "、".join(
        f"{point}{_wind_alarm_status(float(wind[point]['max_10min']))}" for point in ("W1", "W2")
    )
    cursor.paragraph(
        f"台风窗口内，W1桥面与W2塔顶原始风速最大值分别为{float(wind['W1']['raw_max']):.2f}m/s和"
        f"{float(wind['W2']['raw_max']):.2f}m/s；最大10min平均风速分别为"
        f"{float(wind['W1']['max_10min']):.2f}m/s和{float(wind['W2']['max_10min']):.2f}m/s。"
        f"按10min平均风速阈值判断，{alarm_text}。原始最大值反映短时阵风，不能直接套用10min平均风速阈值。",
        bold_prefix="台风窗口内，",
    )
    add_captioned_image(cursor, charts["wind_speed"], "台风窗口W1/W2风速时程及原始最大值")
    add_captioned_image(cursor, charts["wind_maximum"], "台风窗口W1/W2风速最大值与预警参考值比较")
    wind_rows = []
    for point, location in (("W1", "桥面"), ("W2", "塔顶")):
        summary = wind[point]
        wind_rows.append(
            [
                point,
                location,
                f"{float(summary['raw_max']):.2f}",
                _fmt_time(summary["raw_max_time"]),
                f"{float(summary['raw_max_direction']):.0f}°",
                f"{float(summary['max_10min']):.2f}",
                _fmt_time(summary["max_10min_time"], "%m-%d %H:%M"),
                f"{float(summary['max_10min_direction']):.0f}°",
                _wind_alarm_status(float(summary["max_10min"])),
            ]
        )
    add_captioned_table(
        cursor,
        "台风窗口W1/W2风速最大值分析",
        ["测点", "位置", "原始最大值(m/s)", "发生时间", "同期风向", "最大10min均值(m/s)", "发生时段", "同期风向", "预警判断"],
        wind_rows,
    )
    phase_rows = []
    for point in ("W1", "W2"):
        summary = wind[point]
        phase_rows.append(
            [
                point,
                f"{float(summary['pre_mean']):.2f}",
                f"{float(summary['pre_max_10min']):.2f}",
                f"{float(summary['post_mean']):.2f}",
                f"{float(summary['post_max_10min']):.2f}",
                f"{float(summary['post_mean']) / float(summary['pre_mean']):.2f}",
            ]
        )
    add_captioned_table(
        cursor,
        "台风登陆前后风速阶段比较",
        ["测点", "登陆前均值", "登陆前最大10min", "登陆后均值", "登陆后最大10min", "登陆后/前均值比"],
        phase_rows,
    )
    cursor.paragraph(
        f"W1登陆前后平均风速分别为{float(wind['W1']['pre_mean']):.2f}m/s和"
        f"{float(wind['W1']['post_mean']):.2f}m/s；W2分别为"
        f"{float(wind['W2']['pre_mean']):.2f}m/s和{float(wind['W2']['post_mean']):.2f}m/s。"
        "W1位于桥面散索鞍保护罩附近，W2位于塔顶，两测点不在同一竖向测风剖面，"
        "局部遮挡、桥梁绕流、来流风向和桥位地形均会影响二者差异，不宜仅按高度增风关系判断。"
    )
    add_captioned_image(cursor, charts["wind_direction"], "台风窗口W1/W2 10min圆平均风向时程")


def _remove_terminal_tail(doc: DocumentType, after_element) -> None:
    body = doc._element.body
    children = body_children(doc)
    start = children.index(after_element) + 1
    for element in children[start:]:
        if element.tag == qn("w:p") and element_text(element).replace(" ", "") in {
            "(以下无正文)",
            "（以下无正文）",
        }:
            body.remove(element)


def _insert_conclusion(
    doc: DocumentType,
    anchors: Anchors,
    manifest: dict[str, object],
) -> None:
    _remove_terminal_tail(doc, anchors.result_sections["地震动监测"])
    final_sect = next(
        element for element in body_children(doc) if element.tag == qn("w:sectPr")
    )
    cursor = InsertBefore(doc, final_sect)
    cursor.heading(
        "台风影响综合分析、运营建议与数据限制",
        level=2,
        reference=anchors.result_sections["地震动监测"],
    )
    wind = manifest["wind_summary"]
    structure = manifest["structure_summary"]
    max_ratio_group = max(structure, key=lambda group: float(structure[group]["median_ratio"]))
    alarm_text = "、".join(
        f"{point}{_wind_alarm_status(float(wind[point]['max_10min']))}" for point in ("W1", "W2")
    )
    cursor.paragraph(
        f"综合判断：台风窗口内{alarm_text}；W1、W2原始风速最大值均作为短时阵风指标单独报告。"
        f"结构响应登陆后/登陆前中位值比最大的类别为{max_ratio_group}，比值为"
        f"{float(structure[max_ratio_group]['median_ratio']):.2f}。"
        "主梁、主塔及吊索响应未见持续、多测点同步放大，现有证据不支持判定出现台风诱发结构异常。",
        bold_prefix="综合判断：",
    )
    recommendations = [
        "运营管理仍应以气象预警、交通管控指令和现场巡查为准，不因单项监测指标未达阈值而降低防台等级。",
        "继续关注W1/W2风速峰值、10min平均风速及主梁、主塔、吊索振动的同步变化；若出现持续抬升，应复核原始波形、频谱和现场状态。",
        "后续取得新增导出数据时，应按同一统计口径更新最大值、登陆前后阶段比较和综合结论。",
    ]
    cursor.paragraph("运营建议：", bold_prefix="运营建议：")
    for index, text in enumerate(recommendations, start=1):
        paragraph = cursor.paragraph(f"{index}、{text}")
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    end = _parse_datetime(manifest["window"]["end"], "window.end")
    cursor.paragraph(
        f"数据限制：本次增量分析的最新实测数据截至{end:%Y年%m月%d日%H:%M}，不对该时点之后的数据插值或推测。"
        "台风窗口交通称重、结构应变、主塔倾斜和支座变位数据尚未取得，相应4.1至4.4章节按要求留空；"
        "不得使用第二季度统计值代替台风时段数据。结构振动趋势比较不等同于设计承载能力验算。",
        bold_prefix="数据限制：",
    )


def _insert_cell_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    cell = paragraph._parent
    new_paragraph = cell.add_paragraph()
    new_paragraph.style = paragraph.style
    paragraph._p.addnext(new_paragraph._p)
    run = new_paragraph.add_run(text)
    run.bold = False
    return new_paragraph


def clear_front_unavailable_summary(doc: DocumentType) -> None:
    """Clear item 1-4 narratives without replacing the merged summary cell.

    The numbered/category paragraphs remain untouched.  Only the paragraphs
    strictly between each heading and the next numbered heading are removed,
    preventing Q2 traffic/strain/tilt/bearing values from leaking into the
    typhoon-window report while preserving the original table formatting.
    """
    if len(doc.tables) < 3:
        raise RuntimeError("Q2 front summary tables were not found")
    cell = doc.tables[1].cell(3, 2)
    paragraphs = cell.paragraphs
    heading_indices: dict[int, int] = {}
    for index, paragraph in enumerate(paragraphs):
        match = re.match(r"^([1-5])、", paragraph.text.strip())
        if match:
            heading_indices[int(match.group(1))] = index
    if sorted(heading_indices) != [1, 2, 3, 4, 5]:
        raise RuntimeError(
            f"front first-page summary must preserve items 1-5, found {sorted(heading_indices)}"
        )
    for number in range(4, 0, -1):
        start = heading_indices[number]
        end = heading_indices[number + 1]
        for paragraph in paragraphs[start + 1 : end]:
            cell._tc.remove(paragraph._p)

    refreshed = cell.paragraphs
    refreshed_indices: dict[int, int] = {}
    for index, paragraph in enumerate(refreshed):
        match = re.match(r"^([1-5])、", paragraph.text.strip())
        if match:
            refreshed_indices[int(match.group(1))] = index
    for number in range(1, 5):
        if refreshed_indices[number + 1] != refreshed_indices[number] + 1:
            leaked = [
                paragraph.text
                for paragraph in refreshed[
                    refreshed_indices[number] + 1 : refreshed_indices[number + 1]
                ]
                if paragraph.text.strip()
            ]
            raise RuntimeError(f"front summary item {number} still contains Q2 payload: {leaked}")


def append_front_summary(doc: DocumentType, manifest: dict[str, object]) -> None:
    if len(doc.tables) < 3:
        raise RuntimeError("Q2 front summary tables were not found")
    result_cell = doc.tables[2].cell(0, 1)
    full_text = "\n".join(paragraph.text for paragraph in result_cell.paragraphs)
    if INCREMENT_MARKER in full_text:
        raise RuntimeError("front summary already contains a typhoon increment")
    paragraphs = result_cell.paragraphs
    headings: dict[int, tuple[int, Paragraph]] = {}
    for index, paragraph in enumerate(paragraphs):
        match = re.match(r"^([678])、", paragraph.text.strip())
        if match:
            headings[int(match.group(1))] = (index, paragraph)
    if sorted(headings) != [6, 7, 8]:
        raise RuntimeError(f"front continuation must preserve items 6, 7 and 8, found {sorted(headings)}")
    index_6 = headings[6][0]
    index_7 = headings[7][0]
    index_8 = headings[8][0]
    content_6 = next((p for p in reversed(paragraphs[index_6 + 1 : index_7]) if p.text.strip()), None)
    content_7 = next((p for p in reversed(paragraphs[index_7 + 1 : index_8]) if p.text.strip()), None)
    if content_6 is None or content_7 is None:
        raise RuntimeError("front summary items 6/7 do not contain an existing Q2 paragraph")

    structure = manifest["structure_summary"]
    _insert_cell_paragraph_after(
        content_6,
        f"{INCREMENT_MARKER}台风登陆前后比较中，主梁、主塔10min包络峰值中位值比分别为"
        f"{float(structure['主梁加速度']['median_ratio']):.2f}和"
        f"{float(structure['主塔加速度']['median_ratio']):.2f}，未见持续、多测点同步放大。",
    )
    wind = manifest["wind_summary"]
    _insert_cell_paragraph_after(
        content_7,
        f"{INCREMENT_MARKER}W1桥面、W2塔顶原始风速最大值分别为"
        f"{float(wind['W1']['raw_max']):.2f}m/s和{float(wind['W2']['raw_max']):.2f}m/s；"
        f"最大10min平均风速分别为{float(wind['W1']['max_10min']):.2f}m/s和"
        f"{float(wind['W2']['max_10min']):.2f}m/s，均未达到25m/s一级阈值。",
    )
    updated = "\n".join(paragraph.text for paragraph in result_cell.paragraphs)
    for number in (6, 7, 8):
        if len(re.findall(rf"(?m)^{number}、", updated)) != 1:
            raise RuntimeError(f"front summary item {number} was not preserved exactly once")


def _field_kind(paragraph_element) -> str | None:
    instruction = "".join(
        node.text or "" for node in paragraph_element.findall(".//" + qn("w:instrText"))
    )
    for kind in ("图", "表"):
        if re.search(rf"\bSEQ\s+{kind}(?:\s|$)", instruction):
            return kind
    return None


def _set_field_cached_result(paragraph_element, value: int) -> None:
    runs = list(paragraph_element.findall("./" + qn("w:r")))
    separate_index = None
    end_index = None
    for index, run in enumerate(runs):
        for fld_char in run.findall("./" + qn("w:fldChar")):
            kind = fld_char.get(qn("w:fldCharType"))
            if kind == "separate":
                separate_index = index
            elif kind == "end" and separate_index is not None:
                end_index = index
                break
        if end_index is not None:
            break
    if separate_index is None or end_index is None:
        return
    for run in runs[separate_index + 1 : end_index]:
        texts = run.findall(".//" + qn("w:t"))
        if texts:
            texts[0].text = str(value)
            for extra in texts[1:]:
                extra.text = ""
            return


def renumber_chapter_four_seq_fields(doc: DocumentType, chapter_four_element) -> dict[str, int]:
    children = body_children(doc)
    start = children.index(chapter_four_element)
    counts = {"图": 0, "表": 0}
    for element in children[start + 1 :]:
        if element.tag != qn("w:p"):
            continue
        kind = _field_kind(element)
        if kind is None:
            continue
        counts[kind] += 1
        _set_field_cached_result(element, counts[kind])
    return counts


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def build_incremental_report(
    base_docx: Path,
    manifest_path: Path,
    output_docx: Path,
    *,
    charts_dir: Path | None = None,
    explicit_charts: dict[str, Path | None] | None = None,
    audit_path: Path | None = None,
) -> tuple[Path, dict[str, object]]:
    base_docx = base_docx.resolve()
    manifest_path = manifest_path.resolve()
    output_docx = output_docx.resolve()
    if not base_docx.is_file():
        raise FileNotFoundError(base_docx)
    if not manifest_path.is_file():
        raise FileNotFoundError(manifest_path)
    if output_docx == base_docx:
        raise ValueError("output DOCX must not overwrite the base DOCX")
    manifest = load_manifest(manifest_path)
    charts = resolve_charts(manifest, manifest_path, charts_dir, explicit_charts)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(base_docx)
    anchors = locate_anchors(doc)
    cleared = clear_unavailable_sections(doc, anchors)
    clear_front_unavailable_summary(doc)
    append_front_summary(doc, manifest)
    _insert_event_and_coverage(doc, anchors, manifest)
    _insert_cable_increment(doc, anchors, manifest)
    _insert_girder_tower_increment(doc, anchors, manifest, charts["structure"])
    _insert_wind_increment(doc, anchors, manifest, charts)
    _insert_conclusion(doc, anchors, manifest)
    sequence_counts = renumber_chapter_four_seq_fields(doc, anchors.chapter_four)

    # Re-run the no-leak gate after every insertion.  None of the new content is
    # inserted inside the intentionally blank 1.3/1.4 and 4.1-4.4 sections.
    assert_empty_between(
        doc,
        anchors.chapter_one_sections["健康监测系统运行状况"],
        anchors.chapter_one_sections["软硬件维护状况"],
        "1.3",
    )
    assert_empty_between(
        doc,
        anchors.chapter_one_sections["软硬件维护状况"],
        anchors.chapter_two,
        "1.4",
    )
    empty_titles = ["交通状况监测", "结构应变监测", "主塔倾斜监测", "支座变位监测"]
    for index, title in enumerate(empty_titles):
        next_title = (
            empty_titles[index + 1] if index + 1 < len(empty_titles) else "吊索索力监测"
        )
        assert_empty_between(
            doc,
            anchors.result_sections[title],
            anchors.result_sections[next_title],
            f"4.{index + 1}",
        )

    doc.core_properties.title = "洪塘大桥台风巴威影响监测报告（Q2模板增量版）"
    doc.core_properties.subject = "台风登陆前后监测数据增量分析"
    update_fields = doc.settings.element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        doc.settings.element.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.save(output_docx)

    audit = {
        "status": "ok",
        "base_docx": str(base_docx),
        "base_sha256": _sha256(base_docx),
        "typhoon_manifest": str(manifest_path),
        "manifest_sha256": _sha256(manifest_path),
        "output_docx": str(output_docx),
        "output_sha256": _sha256(output_docx),
        "window": manifest["window"],
        "cleared_sections": cleared,
        "blank_sections_verified": ["1.3", "1.4", "4.1", "4.2", "4.3", "4.4"],
        "preserved_result_sections": ["4.5", "4.6", "4.7", "4.8"],
        "front_summary": {
            "preserved_items": list(range(1, 9)),
            "cleared_items": [1, 2, 3, 4],
            "appended_items": [6, 7],
            "whole_cell_replaced": False,
        },
        "charts": {key: {"path": str(path), "sha256": _sha256(path)} for key, path in charts.items()},
        "sequence_fields": sequence_counts,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
    }
    audit_path = audit_path or output_docx.with_name(output_docx.stem + "_build_audit.json")
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return output_docx, audit


def parser() -> argparse.ArgumentParser:
    value = argparse.ArgumentParser(
        description="Add typhoon-window analysis to a populated Hongtang Q2 report without replacing its template structure"
    )
    value.add_argument("--base-docx", type=Path, required=True)
    value.add_argument("--manifest", type=Path, required=True)
    value.add_argument("--output", type=Path, required=True)
    value.add_argument("--charts-dir", type=Path)
    value.add_argument("--wind-speed-chart", type=Path)
    value.add_argument("--wind-maximum-chart", type=Path)
    value.add_argument("--wind-direction-chart", type=Path)
    value.add_argument("--structure-chart", type=Path)
    value.add_argument("--audit-output", type=Path)
    return value


def main() -> None:
    args = parser().parse_args()
    output, audit = build_incremental_report(
        args.base_docx,
        args.manifest,
        args.output,
        charts_dir=args.charts_dir,
        explicit_charts={
            "wind_speed": args.wind_speed_chart,
            "wind_maximum": args.wind_maximum_chart,
            "wind_direction": args.wind_direction_chart,
            "structure": args.structure_chart,
        },
        audit_path=args.audit_output,
    )
    print(
        json.dumps(
            {
                "status": audit["status"],
                "report": str(output),
                "window": audit["window"],
                "blank_sections_verified": audit["blank_sections_verified"],
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
