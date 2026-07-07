from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Any
from zipfile import ZipFile

from docx import Document
from docx.document import Document as DocxDocument

try:
    from analysis_manifest import manifest_precheck_warnings
except Exception:  # pragma: no cover - optional CLI helper path
    manifest_precheck_warnings = None

try:
    from docx_table_utils import find_summary_table
except Exception:  # pragma: no cover - package import path
    from .docx_table_utils import find_summary_table


@dataclass(frozen=True)
class TemplateIssue:
    code: str
    message: str


class TemplatePrecheckError(RuntimeError):
    def __init__(self, template: Path, issues: list[TemplateIssue]):
        self.template = template
        self.issues = issues
        details = "\n".join(f"- [{issue.code}] {issue.message}" for issue in issues)
        super().__init__(f"Template precheck failed: {template}\n{details}")


def issue_to_dict(issue: TemplateIssue) -> dict[str, str]:
    return {"code": issue.code, "message": issue.message}


def summarize_template(template: Path) -> dict[str, int | str]:
    if not template.exists():
        return {"exists": 0}
    doc = Document(str(template))
    image_count = 0
    try:
        with ZipFile(template) as zf:
            image_count = sum(1 for name in zf.namelist() if name.startswith("word/media/"))
    except Exception:
        image_count = 0
    return {
        "exists": 1,
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_count": len(_paragraph_texts(doc)),
        "table_count": len(doc.tables),
        "table_text_count": len(_table_texts(doc)),
        "field_count": len(_field_instr_texts(doc)),
        "image_count": image_count,
    }


def build_precheck_payload(
    kind: str,
    template: Path,
    issues: list[TemplateIssue],
    warnings: list[str] | None = None,
    context: dict[str, Any] | None = None,
) -> dict[str, Any]:
    warnings = warnings or []
    if issues:
        status = "failed"
    elif warnings:
        status = "warning"
    else:
        status = "ok"
    merged_context = dict(context or {})
    merged_context.setdefault("template_summary", summarize_template(template))
    return {
        "checked_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "kind": kind,
        "template": str(template),
        "status": status,
        "issue_count": len(issues),
        "warning_count": len(warnings),
        "issues": [issue_to_dict(issue) for issue in issues],
        "warnings": warnings,
        "context": merged_context,
    }


def write_precheck_report(
    kind: str,
    template: Path,
    issues: list[TemplateIssue],
    output_dir: Path,
    warnings: list[str] | None = None,
    context: dict[str, Any] | None = None,
    prefix: str = "template_precheck",
) -> tuple[Path, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    payload = build_precheck_payload(kind, template, issues, warnings=warnings, context=context)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_kind = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in kind)
    json_path = output_dir / f"{prefix}_{safe_kind}_{timestamp}.json"
    txt_path = output_dir / f"{prefix}_{safe_kind}_{timestamp}.txt"

    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        f"检查时间: {payload['checked_at']}",
        f"检查类型: {kind}",
        f"模板文件: {template}",
        f"状态: {payload['status']}",
        f"问题数量: {len(issues)}",
        f"提示数量: {len(payload['warnings'])}",
    ]
    if context:
        lines.append("")
        lines.append("上下文:")
        for key, value in context.items():
            lines.append(f"- {key}: {value}")
    if issues:
        lines.append("")
        lines.append("模板问题:")
        for issue in issues:
            lines.append(f"- [{issue.code}] {issue.message}")
    if payload["warnings"]:
        lines.append("")
        lines.append("目录/结果提示:")
        for warning in payload["warnings"]:
            lines.append(f"- {warning}")
    if not issues and not payload["warnings"]:
        lines.append("")
        lines.append("检查通过：未发现模板锚点或结果目录问题。")
    txt_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return txt_path, json_path


def _paragraph_texts(doc: DocxDocument) -> list[str]:
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]


def _table_texts(doc: DocxDocument) -> list[str]:
    texts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    texts.append(text)
    return texts


def _all_texts(doc: DocxDocument) -> list[str]:
    return _paragraph_texts(doc) + _table_texts(doc)


def _field_instr_texts(doc: DocxDocument) -> list[str]:
    fields: list[str] = []
    for para in doc.paragraphs:
        text = "".join(node.text or "" for node in para._p.iter() if node.tag.endswith("}instrText"))
        if text.strip():
            fields.append(text)
    return fields


def _contains_any(texts: Iterable[str], fragment: str) -> bool:
    return any(fragment in text for text in texts)


def _add_missing_fragment(issues: list[TemplateIssue], texts: list[str], fragment: object, note: str) -> None:
    if isinstance(fragment, (list, tuple, set)):
        options = [str(item) for item in fragment if item]
        _add_missing_any_fragment(issues, texts, options, note)
        return
    if fragment is None:
        return
    fragment_text = str(fragment)
    if not _contains_any(texts, fragment_text):
        issues.append(TemplateIssue("missing-text", f"{note}: {fragment_text}"))


def _add_missing_any_fragment(issues: list[TemplateIssue], texts: list[str], fragments: list[str], note: str) -> None:
    if not any(_contains_any(texts, fragment) for fragment in fragments):
        issues.append(TemplateIssue("missing-text", f"{note}: {' / '.join(fragments)}"))


def _add_missing_pattern(issues: list[TemplateIssue], texts: list[str], pattern: str, note: str) -> None:
    regex = re.compile(pattern)
    if not any(regex.search(text) for text in texts):
        issues.append(TemplateIssue("missing-text", f"{note}: /{pattern}/"))


def _section_available(section: dict | None) -> bool:
    if not isinstance(section, dict):
        return False
    if not section.get("enabled", True):
        return False
    return bool(section.get("available", True))


def check_hongtang_period_template(template: Path, manifest: dict | None = None) -> list[TemplateIssue]:
    doc = Document(str(template))
    texts = _all_texts(doc)
    issues: list[TemplateIssue] = []

    # These anchors are used unconditionally by the period report writer.
    required = [
        ("健康监测系统运行状况", "1.4 health-status section anchor"),
        ("交通状况监测", "WIM section start anchor"),
        ("结构应变监测", "WIM section end / next-section anchor"),
        ("桥梁共通过车辆", "WIM paragraph template"),
        ("季度交通状况分月统计表", "WIM quarterly table caption template"),
        ("桥梁交通流参数分析", "WIM figure caption template"),
    ]
    for fragment, note in required:
        _add_missing_fragment(issues, texts, fragment, note)
    _add_missing_pattern(issues, texts, r"\d{4}年\d{1,2}月交通状况监测", "WIM monthly heading template")
    _add_missing_pattern(issues, texts, r"续表\s*4-\d+", "WIM continuation caption template")

    if manifest:
        sections = manifest.get("sections", {})
        strain = sections.get("strain")
        if _section_available(strain):
            for key in (
                "girder_timeseries_caption",
                "girder_boxplot_caption",
                "tower_timeseries_caption",
                "tower_boxplot_caption",
            ):
                caption = strain.get(key)
                if caption:
                    _add_missing_fragment(issues, texts, caption, f"strain caption anchor ({key})")
            for fragment in ("主梁应变", "桥塔应变"):
                _add_missing_fragment(issues, texts, fragment, "strain section paragraph anchor")

        tilt = sections.get("tilt")
        if _section_available(tilt):
            _add_missing_fragment(issues, texts, "主塔倾角偏移的方向以闽侯上街-农林大学为纵桥向", "tilt narrative anchor")
            if tilt.get("caption"):
                _add_missing_fragment(issues, texts, tilt["caption"], "tilt caption anchor")

        bearing = sections.get("bearing_displacement")
        if _section_available(bearing):
            _add_missing_fragment(issues, texts, "支座变位的方向以闽侯上街-农林大学为纵桥向", "bearing narrative anchor")
            if bearing.get("caption"):
                _add_missing_fragment(issues, texts, bearing["caption"], "bearing caption anchor")

        cable = sections.get("cable_force")
        if _section_available(cable):
            if cable.get("accel_available"):
                _add_missing_fragment(issues, texts, "（1）索力加速度时程数据", "cable-acceleration subsection anchor")
                _add_missing_fragment(issues, texts, "（2）索力时程数据", "cable-acceleration subsection end anchor")
                if cable.get("accel_caption"):
                    _add_missing_fragment(issues, texts, cable["accel_caption"], "cable-acceleration caption anchor")
            if cable.get("force_available"):
                _add_missing_fragment(issues, texts, "（2）索力时程数据", "cable-force subsection anchor")
                _add_missing_fragment(issues, texts, "4.6 主梁、主塔振动监测", "cable-force subsection end anchor")
                if cable.get("force_caption"):
                    _add_missing_fragment(issues, texts, cable["force_caption"], "cable-force caption anchor")

        vibration = sections.get("vibration")
        if _section_available(vibration):
            for fragment in ("（1）振动时程数据", "（2）自振频率", "4.7 风向风速监测"):
                _add_missing_fragment(issues, texts, fragment, "vibration subsection anchor")
            for key in ("timeseries_caption", "freq_caption"):
                caption = vibration.get(key)
                if caption:
                    _add_missing_fragment(issues, texts, caption, f"vibration caption anchor ({key})")

        wind = sections.get("wind")
        if _section_available(wind):
            _add_missing_fragment(issues, texts, "风向风速监测", "wind section anchor")
            for key in ("speed_caption", "rose_caption"):
                caption = wind.get(key)
                if caption:
                    _add_missing_fragment(issues, texts, caption, f"wind caption anchor ({key})")

        eq = sections.get("eq")
        if _section_available(eq):
            _add_missing_fragment(issues, texts, "地震动监测", "earthquake section anchor")
            if eq.get("caption"):
                _add_missing_fragment(issues, texts, eq["caption"], "earthquake caption anchor")

    return issues


JLJ_SECTION_HEADINGS = [
    ("主桥环境与作用监测", "温度监测"),
    ("主桥环境与作用监测", "湿度监测"),
    ("主桥环境与作用监测", "雨量监测"),
    ("主桥环境与作用监测", "风向风速监测"),
    ("主桥环境与作用监测", "地震动监测"),
    ("主桥环境与作用监测", "车辆荷载监测"),
    ("主桥结构响应与结构变化监测", "主梁挠度监测"),
    ("主桥结构响应与结构变化监测", "支座、梁段纵向位移监测"),
    ("主桥结构响应与结构变化监测", "拱顶、拱脚位移监测（GNSS）"),
    ("主桥结构响应与结构变化监测", "结构振动监测"),
    ("主桥结构响应与结构变化监测", "结构应变监测"),
    ("主桥结构响应与结构变化监测", "裂缝监测"),
    ("主桥结构响应与结构变化监测", "吊杆索力监测"),
    ("北江滨匝道桥监测", "结构应变监测"),
    ("北江滨匝道桥监测", "支座位移监测"),
    ("北江滨匝道桥监测", "墩柱倾斜监测"),
    ("南江滨匝道桥监测", "结构应变监测"),
    ("南江滨匝道桥监测", "支座位移监测"),
    ("南江滨匝道桥监测", "墩柱倾斜监测"),
]


def check_jlj_monthly_template(template: Path) -> list[TemplateIssue]:
    doc = Document(str(template))
    texts = _all_texts(doc)
    fields = _field_instr_texts(doc)
    issues: list[TemplateIssue] = []

    if find_summary_table(doc) is None:
        issues.append(TemplateIssue("missing-table", "Jiulongjiang summary table with left-column labels '监测结果' and '建  议' was not found."))

    required = [
        ("监测结果", "summary table result cell"),
        ("温度监测", "body style fallback anchor"),
    ]
    for fragment, note in required:
        _add_missing_fragment(issues, texts, fragment, note)
    _add_missing_any_fragment(issues, texts, ["监测系统运行状况", "健康监测系统运行状况"], "health-status heading")
    _add_missing_any_fragment(issues, texts, ["建 议", "建议"], "summary table advice cell")

    for parent, child in JLJ_SECTION_HEADINGS:
        _add_missing_fragment(issues, texts, parent, "Jiulongjiang level-2 section heading")
        _add_missing_fragment(issues, texts, child, f"Jiulongjiang child heading under {parent}")

    if not any("SEQ 图" in field for field in fields):
        issues.append(TemplateIssue("missing-field", "No auto figure caption field (SEQ 图) found."))
    if not any("SEQ 表" in field for field in fields):
        issues.append(TemplateIssue("missing-field", "No auto table caption field (SEQ 表) found."))

    return issues


def check_guanbing_monthly_template(template: Path) -> list[TemplateIssue]:
    doc = Document(str(template))
    texts = _all_texts(doc)
    issues: list[TemplateIssue] = []

    required = [
        ("G104线管柄大桥", "Guanbing report title/project text"),
        ("桥面环境温度", "temperature section anchor"),
        ("桥面环境湿度", "humidity section anchor"),
        ("主梁挠度", "deflection section anchor"),
        ("主墩倾角", "tilt section anchor"),
        ("主梁关键截面应变", "strain section anchor"),
        ("主梁竖向加速度", "acceleration section anchor"),
        ("裂缝宽度", "crack section anchor"),
        ("综上所述", "conclusion summary anchor"),
        ("图 5 桥面环境温度测点时程图", "temperature figure anchor"),
        ("图 7 桥面环境湿度测点时程图", "humidity figure anchor"),
        ("图 13 第2跨主梁位移变化趋势", "deflection figure anchor"),
        ("图 14 第3跨主梁位移变化趋势", "deflection figure anchor"),
    ]
    for fragment, note in required:
        _add_missing_fragment(issues, texts, fragment, note)

    return issues


def check_shuixianhua_monthly_template(template: Path) -> list[TemplateIssue]:
    doc = Document(str(template))
    texts = _all_texts(doc)
    fields = _field_instr_texts(doc)
    issues: list[TemplateIssue] = []

    if find_summary_table(doc) is None:
        issues.append(TemplateIssue("missing-table", "Shuixianhua summary table with front-matter result/advice cells was not found."))

    required = [
        ("水仙花大桥桥梁健康监测", "Shuixianhua cover/project text"),
        ("桥梁监测报告", "front report summary title"),
        ("监测结果", "front summary result label"),
        ("建  议", "front summary advice label"),
        ("目  录", "TOC title"),
        ("主桥监测测点布置", "design layout chapter anchor"),
        (["桥梁环境（温湿度、风速风向）监测", "温度监测"], "environment layout subsection anchor"),
        (["结构温度监测", "温度监测"], "temperature layout subsection anchor"),
        ("地震动监测", "earthquake layout subsection anchor"),
        ("主梁挠度监测", "deflection layout subsection anchor"),
        ("支座及伸缩缝位移监测", "bearing/expansion layout subsection anchor"),
        ("拱顶、拱脚位移监测（GNSS）", "GNSS layout subsection anchor"),
        ("结构振动监测", "vibration layout subsection anchor"),
        ("吊杆及系杆索力监测", "cable-force layout subsection anchor"),
        (["结构应变及动应变监测", "结构应变监测"], "strain layout subsection anchor"),
        ("报警阈值设置", "threshold section anchor"),
        ("水仙花大桥监测阈值汇总表", "threshold table caption anchor"),
        ("测点布置图", "layout figure caption anchor"),
    ]
    for fragment, note in required:
        _add_missing_fragment(issues, texts, fragment, note)

    if not any("SEQ 图" in field for field in fields):
        issues.append(TemplateIssue("missing-field", "No auto figure caption field (SEQ 图) found."))
    if not any("SEQ 表" in field for field in fields):
        issues.append(TemplateIssue("missing-field", "No auto table caption field (SEQ 表) found."))

    return issues


def check_zhishan_monthly_template(template: Path) -> list[TemplateIssue]:
    doc = Document(str(template))
    texts = _all_texts(doc)
    fields = _field_instr_texts(doc)
    issues: list[TemplateIssue] = []

    if find_summary_table(doc) is None:
        issues.append(TemplateIssue("missing-table", "Zhishan summary table with front-matter result/advice cells was not found."))

    required = [
        ("芝山大桥", "Zhishan cover/project text"),
        ("监测结果", "front summary result label"),
        ("建  议", "front summary advice label"),
        ("梁端位移监测", "bearing displacement section anchor"),
        ("结构振动监测", "structure vibration section anchor"),
        ("结构应变监测", "strain section anchor"),
        ("斜拉索索力加速度监测", "cable acceleration section anchor"),
        ("索力监测结果", "cable force table anchor"),
        ("活载动应变（高通滤波后）", "dynamic strain high-pass section anchor"),
        ("低通滤波后", "dynamic strain low-pass section anchor"),
    ]
    for fragment, note in required:
        _add_missing_fragment(issues, texts, fragment, note)

    if not any("SEQ 图" in field for field in fields):
        issues.append(TemplateIssue("missing-field", "No auto figure caption field (SEQ 图) found."))
    if not any("SEQ 表" in field for field in fields):
        issues.append(TemplateIssue("missing-field", "No auto table caption field (SEQ 表) found."))

    return issues


def check_template(kind: str, template: Path, manifest: dict | None = None) -> list[TemplateIssue]:
    if not template.exists():
        return [TemplateIssue("missing-file", f"Template file does not exist: {template}")]
    if kind == "hongtang_period":
        return check_hongtang_period_template(template, manifest)
    if kind == "jlj_monthly":
        return check_jlj_monthly_template(template)
    if kind == "guanbing_monthly":
        return check_guanbing_monthly_template(template)
    if kind == "shuixianhua_monthly":
        return check_shuixianhua_monthly_template(template)
    if kind == "zhishan_monthly":
        return check_zhishan_monthly_template(template)
    raise ValueError(f"Unknown template kind: {kind}")


def raise_for_template(kind: str, template: Path, manifest: dict | None = None) -> None:
    issues = check_template(kind, template, manifest)
    if issues:
        raise TemplatePrecheckError(template, issues)


def main() -> None:
    parser = argparse.ArgumentParser(description="Precheck bridge report DOCX templates.")
    parser.add_argument("--kind", choices=["hongtang_period", "jlj_monthly", "guanbing_monthly", "shuixianhua_monthly", "zhishan_monthly"], required=True)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--manifest", type=Path, default=None, help="Optional analysis manifest for conditional anchor checks.")
    parser.add_argument("--output-dir", type=Path, default=None, help="Optional directory for txt/json precheck reports.")
    parser.add_argument("--result-root", type=Path, default=None, help="Optional analysis result root for manifest warnings.")
    args = parser.parse_args()

    manifest = None
    if args.manifest is not None and args.manifest.exists():
        manifest = json.loads(args.manifest.read_text(encoding="utf-8"))
    issues = check_template(args.kind, args.template, manifest=manifest)
    warnings = []
    if args.result_root is not None and manifest_precheck_warnings is not None:
        warnings = manifest_precheck_warnings(args.result_root)
    if args.output_dir is not None:
        context = {"manifest": str(args.manifest)} if args.manifest else None
        txt_path, json_path = write_precheck_report(args.kind, args.template, issues, args.output_dir, warnings=warnings, context=context)
        print(f"Precheck report: {txt_path}")
        print(f"Precheck JSON:   {json_path}")
    if issues:
        raise TemplatePrecheckError(args.template, issues)
    if warnings:
        print("Template precheck warnings:")
        for warning in warnings:
            print(f"- {warning}")
    print(f"Template precheck OK: {args.template}")


if __name__ == "__main__":
    main()
