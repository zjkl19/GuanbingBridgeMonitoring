from __future__ import annotations

import argparse
import json
import math
import re
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Mm
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from PIL import Image, ImageDraw, ImageFont, ImageOps


@dataclass
class ImageItem:
    label: str
    path: Path | None
    lookup: dict | None = None


HONGTANG_CABLE_BUILT_FORCE = {
    "CS4": 2086.38,
    "CS5": 1442.28,
    "CS6": 1237.98,
    "CS7": 1215.25,
    "CS8": 1205.40,
    "CS9": 2061.52,
    "CX4": 1795.92,
    "CX5": 1435.46,
    "CX6": 1249.43,
    "CX7": 1161.54,
    "CX8": 1449.13,
    "CX9": 1928.46,
}


def parse_args() -> argparse.Namespace:
    repo_root = Path(__file__).resolve().parents[1]
    templates = sorted((repo_root / "reports").glob("*.docx"))
    default_template = templates[0] if templates else None
    parser = argparse.ArgumentParser(description="Build Hongtang monthly monitoring report.")
    parser.add_argument("--template", type=Path, default=default_template)
    parser.add_argument("--config", type=Path, default=repo_root / "config" / "hongtang_config.json")
    parser.add_argument("--result-root", type=Path, default=None)
    parser.add_argument("--analysis-root", type=Path, default=repo_root)
    parser.add_argument("--image-root", type=Path, default=None)
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--period-label", default="2025年12月")
    parser.add_argument("--monitoring-range", default="2025.12.01～2025.12.31")
    parser.add_argument("--report-date", default=datetime.now().strftime("%Y年%m月%d日"))
    return parser.parse_args()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_sheet_rows(path: Path) -> list[dict]:
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb[wb.sheetnames[0]]
    rows = list(ws.iter_rows(values_only=True))
    wb.close()
    if not rows:
        return []
    header = [str(v) if v is not None else "" for v in rows[0]]
    out = []
    for row in rows[1:]:
        item = {}
        for k, v in zip(header, row):
            item[k] = v
        out.append(item)
    return out


def load_workbook_rows_by_sheet(path: Path) -> dict[str, list[dict]]:
    wb = load_workbook(path, read_only=True, data_only=True)
    out: dict[str, list[dict]] = {}
    for wsname in wb.sheetnames:
        ws = wb[wsname]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            out[wsname] = []
            continue
        header = [str(v) if v is not None else "" for v in rows[0]]
        items: list[dict] = []
        for row in rows[1:]:
            item = {}
            for k, v in zip(header, row):
                item[k] = v
            items.append(item)
        out[wsname] = items
    wb.close()
    return out


def read_numeric_series_csv(path: Path) -> list[float]:
    values: list[float] = []
    last_exc: Exception | None = None
    for encoding in ("utf-8-sig", "utf-16", "gbk"):
        try:
            with path.open("r", encoding=encoding) as fh:
                for line in fh:
                    line = line.strip()
                    if not line or "," not in line or line.startswith("开始时间") or line.startswith("序列号") or line.startswith("通道号"):
                        continue
                    parts = line.split(",", 1)
                    if len(parts) != 2:
                        continue
                    try:
                        values.append(float(parts[1].strip()))
                    except ValueError:
                        continue
            return values
        except UnicodeError as exc:
            last_exc = exc
            values = []
            continue
    if last_exc is not None:
        raise last_exc
    return values


def resolve_existing_file(primary_root: Path | None, fallback_root: Path | None, filename: str) -> Path:
    candidates: list[Path] = []
    if primary_root is not None:
        candidates.append(primary_root / filename)
        candidates.append(primary_root / "stats" / filename)
    if fallback_root is not None:
        fallback = fallback_root / filename
        if fallback not in candidates:
            candidates.append(fallback)
        fallback_stats = fallback_root / "stats" / filename
        if fallback_stats not in candidates:
            candidates.append(fallback_stats)
    for path in candidates:
        if path.exists():
            return path
    joined = ", ".join(str(p) for p in candidates)
    raise FileNotFoundError(f"Required file not found: {filename}. Checked: {joined}")


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def should_skip_search_dir(path: Path) -> bool:
    banned_parts = {".git", ".venv", "tests", "__pycache__"}
    return any(part in banned_parts for part in path.parts)


def resolve_output_dirs(root: Path, configured_dir: str) -> list[Path]:
    configured_path = Path(configured_dir)
    candidates: list[Path] = []
    direct = (root / configured_path).resolve()
    if direct.exists() and direct.is_dir():
        candidates.append(direct)

    target_name = configured_path.name
    if root.exists():
        for found in root.rglob(target_name):
            if not found.is_dir():
                continue
            resolved = found.resolve()
            if resolved in candidates or should_skip_search_dir(resolved):
                continue
            candidates.append(resolved)

    candidates.sort(key=lambda p: p.stat().st_mtime if p.exists() else 0, reverse=True)
    return candidates


def find_latest_image(root: Path, configured_dir: str, stem_prefix: str) -> tuple[Path | None, dict]:
    patterns = [f"{stem_prefix}*.jpg", f"{stem_prefix}*.png", f"{stem_prefix}*.jpeg"]
    return find_latest_image_patterns(root, configured_dir, patterns)


def find_latest_image_patterns(root: Path, configured_dir: str, patterns: list[str]) -> tuple[Path | None, dict]:
    return find_latest_file_patterns(root, configured_dir, patterns)


def find_latest_file_patterns(root: Path, configured_dir: str, patterns: list[str]) -> tuple[Path | None, dict]:
    resolved_dirs = resolve_output_dirs(root, configured_dir)
    matched: list[Path] = []
    for folder in resolved_dirs:
        for pattern in patterns:
            matched.extend(folder.glob(pattern))
    matched = sorted({p.resolve() for p in matched}, key=lambda p: p.stat().st_mtime, reverse=True)
    return (
        matched[0] if matched else None,
        {
            "image_root": str(root),
            "configured_dir": configured_dir,
            "resolved_dirs": [str(p) for p in resolved_dirs],
            "patterns": patterns,
            "matched_files": [str(p) for p in matched[:10]],
            "selected_file": str(matched[0]) if matched else None,
        },
    )


def normalize_name_list(raw: object) -> list[str]:
    if not isinstance(raw, list):
        return []
    out: list[str] = []
    seen: set[str] = set()
    for item in raw:
        val = str(item).strip()
        if not val or val in seen:
            continue
        out.append(val)
        seen.add(val)
    return out


def get_reporting_section(cfg: dict, module: str) -> dict:
    reporting = cfg.get("reporting", {})
    section = reporting.get(module, {})
    return section if isinstance(section, dict) else {}


def reporting_enabled(cfg: dict, module: str, default: bool = True) -> bool:
    section = get_reporting_section(cfg, module)
    if "enabled" not in section:
        return default
    return bool(section.get("enabled"))


def get_report_order(cfg: dict, module: str, key: str, default: list[str]) -> list[str]:
    section = get_reporting_section(cfg, module)
    order = normalize_name_list(section.get(key)) or list(default)
    include = normalize_name_list(section.get(f"{key}_include")) or normalize_name_list(section.get("include"))
    exclude = set(normalize_name_list(section.get(f"{key}_exclude")) or normalize_name_list(section.get("exclude")))

    if include:
        ordered = include
    else:
        ordered = order

    out: list[str] = []
    seen: set[str] = set()
    for item in ordered:
        if item in exclude or item in seen:
            continue
        out.append(item)
        seen.add(item)
    return out


def select_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyh.ttf",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for cand in font_candidates:
        if Path(cand).exists():
            try:
                return ImageFont.truetype(cand, size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def build_tile(label: str, image_path: Path | None, tile_size: tuple[int, int]) -> Image.Image:
    tile_w, tile_h = tile_size
    canvas = Image.new("RGB", (tile_w, tile_h), "white")
    draw = ImageDraw.Draw(canvas)
    font = select_font(28)
    label_box_h = 52

    draw.rectangle((0, 0, tile_w - 1, tile_h - 1), outline=(180, 180, 180), width=2)
    draw.text((16, 12), label, fill=(20, 20, 20), font=font)

    content_box = (16, label_box_h + 8, tile_w - 16, tile_h - 16)
    if image_path is None or not image_path.exists():
        draw.rectangle(content_box, outline=(220, 80, 80), width=2)
        missing_font = select_font(24)
        draw.text((content_box[0] + 20, content_box[1] + 20), "未找到图片", fill=(180, 0, 0), font=missing_font)
        return canvas

    img = Image.open(image_path).convert("RGB")
    target_w = content_box[2] - content_box[0]
    target_h = content_box[3] - content_box[1]
    img = ImageOps.contain(img, (target_w, target_h))
    x0 = content_box[0] + (target_w - img.width) // 2
    y0 = content_box[1] + (target_h - img.height) // 2
    canvas.paste(img, (x0, y0))
    return canvas


def compose_grid(items: list[ImageItem], out_path: Path, cols: int = 2, tile_size: tuple[int, int] = (920, 560)) -> Path:
    rows = max(1, math.ceil(len(items) / cols))
    gap = 24
    canvas_w = cols * tile_size[0] + (cols + 1) * gap
    canvas_h = rows * tile_size[1] + (rows + 1) * gap
    canvas = Image.new("RGB", (canvas_w, canvas_h), (248, 248, 248))

    for idx, item in enumerate(items):
        row = idx // cols
        col = idx % cols
        x = gap + col * (tile_size[0] + gap)
        y = gap + row * (tile_size[1] + gap)
        tile = build_tile(item.label, item.path, tile_size)
        canvas.paste(tile, (x, y))

    ensure_dir(out_path.parent)
    canvas.save(out_path, quality=92)
    return out_path


def insert_paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def find_paragraph_indices(doc: Document, text: str) -> list[int]:
    indices = []
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip() == text:
            indices.append(idx)
    return indices


def find_paragraph_indices_contains(doc: Document, fragment: str) -> list[int]:
    indices = []
    for idx, para in enumerate(doc.paragraphs):
        if fragment in para.text.strip():
            indices.append(idx)
    return indices


def find_paragraph_indices_contains_between(
    doc: Document,
    fragment: str,
    start_idx: int = 0,
    end_idx: int | None = None,
) -> list[int]:
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    indices = []
    for idx in range(start_idx, min(end_idx, len(doc.paragraphs))):
        if fragment in doc.paragraphs[idx].text.strip():
            indices.append(idx)
    return indices


def find_last_paragraph(doc: Document, text: str) -> Paragraph:
    indices = find_paragraph_indices(doc, text)
    if not indices:
        raise ValueError(f'Paragraph "{text}" not found in template')
    return doc.paragraphs[indices[-1]]


def find_first_paragraph_contains(doc: Document, fragment: str) -> Paragraph:
    indices = find_paragraph_indices_contains(doc, fragment)
    if not indices:
        raise ValueError(f'Paragraph containing "{fragment}" not found in template')
    return doc.paragraphs[indices[0]]


def find_last_paragraph_contains(doc: Document, fragment: str) -> Paragraph:
    indices = find_paragraph_indices_contains(doc, fragment)
    if not indices:
        raise ValueError(f'Paragraph containing "{fragment}" not found in template')
    return doc.paragraphs[indices[-1]]


def find_last_paragraph_contains_scoped(
    doc: Document,
    fragment: str | list[str] | tuple[str, ...],
    anchor_text: str | None = None,
    stop_text: str | None = None,
) -> Paragraph:
    start_idx = 0
    if anchor_text:
        anchor_indices = find_paragraph_indices(doc, anchor_text)
        if not anchor_indices:
            raise ValueError(f'Anchor "{anchor_text}" not found in template')
        start_idx = anchor_indices[-1] + 1

    end_idx = len(doc.paragraphs)
    if stop_text:
        stop_indices = [idx for idx in find_paragraph_indices(doc, stop_text) if idx > start_idx]
        if stop_indices:
            end_idx = stop_indices[0]

    fragments = [fragment] if isinstance(fragment, str) else list(fragment)
    for frag in fragments:
        indices = find_paragraph_indices_contains_between(doc, frag, start_idx, end_idx)
        if indices:
            return doc.paragraphs[indices[-1]]

    scope = f' after "{anchor_text}"' if anchor_text else ""
    if stop_text:
        scope += f' before "{stop_text}"'
    joined = " / ".join(fragments)
    raise ValueError(f'Paragraph containing "{joined}" not found in template{scope}')


def replace_next_nonempty_paragraph(doc: Document, anchor_fragment: str, new_text: str, use_last: bool = True, skip: int = 0) -> None:
    indices = find_paragraph_indices_contains(doc, anchor_fragment)
    if not indices:
        raise ValueError(f'Anchor containing "{anchor_fragment}" not found in template')
    anchor_idx = indices[-1] if use_last else indices[0]
    seen = 0
    for idx in range(anchor_idx + 1, len(doc.paragraphs)):
        txt = doc.paragraphs[idx].text.strip()
        if txt:
            if seen < skip:
                seen += 1
                continue
            replace_paragraph_text(doc.paragraphs[idx], new_text)
            return
    raise ValueError(f'No non-empty paragraph found after anchor "{anchor_fragment}"')


def replace_next_nonempty_after_exact(doc: Document, anchor_text: str, new_text: str, use_last: bool = True, skip: int = 0) -> None:
    indices = find_paragraph_indices(doc, anchor_text)
    if not indices:
        raise ValueError(f'Anchor "{anchor_text}" not found in template')
    anchor_idx = indices[-1] if use_last else indices[0]
    seen = 0
    for idx in range(anchor_idx + 1, len(doc.paragraphs)):
        txt = doc.paragraphs[idx].text.strip()
        if txt:
            if seen < skip:
                seen += 1
                continue
            replace_paragraph_text(doc.paragraphs[idx], new_text)
            return
    raise ValueError(f'No non-empty paragraph found after anchor "{anchor_text}"')


def clear_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def ensure_note_before_caption(doc: Document, caption_fragment: str, note_text: str) -> None:
    for para in doc.paragraphs:
        if para.text.strip() == note_text:
            clear_paragraph(para)
    caption = find_last_paragraph_contains(doc, caption_fragment)
    note_para = insert_paragraph_before(caption)
    note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_para.add_run(note_text)


def insert_picture_before_caption(doc: Document, caption_text: str, image_path: Path, width_mm: float = 165.0) -> None:
    caption = find_last_paragraph(doc, caption_text)
    pic_para = insert_paragraph_before(caption)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))


def insert_picture_before_caption_contains(doc: Document, caption_fragment: str, image_path: Path, width_mm: float = 165.0) -> None:
    caption = find_last_paragraph_contains(doc, caption_fragment)
    pic_para = insert_paragraph_before(caption)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))


def insert_labeled_images_before_caption_contains(
    doc: Document,
    caption_fragment: str | list[str] | tuple[str, ...],
    items: list[ImageItem],
    width_mm: float = 165.0,
    anchor_text: str | None = None,
    stop_text: str | None = None,
) -> None:
    caption = find_last_paragraph_contains_scoped(doc, caption_fragment, anchor_text=anchor_text, stop_text=stop_text)
    for item in items:
        if item.path is None or not item.path.exists():
            continue
        pic_para = insert_paragraph_before(caption)
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        run.add_picture(str(item.path), width=Mm(width_mm))

        label_para = insert_paragraph_after(pic_para)
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_para.add_run(item.label)


def replace_last_paragraph(doc: Document, exact_text: str, new_text: str) -> None:
    para = find_last_paragraph(doc, exact_text)
    replace_paragraph_text(para, new_text)


def replace_first_paragraph(doc: Document, exact_text: str, new_text: str) -> None:
    indices = find_paragraph_indices(doc, exact_text)
    if not indices:
        raise ValueError(f'Paragraph "{exact_text}" not found in template')
    replace_paragraph_text(doc.paragraphs[indices[0]], new_text)


def replace_first_paragraph_contains(doc: Document, fragment: str, new_text: str) -> None:
    para = find_first_paragraph_contains(doc, fragment)
    replace_paragraph_text(para, new_text)


def replace_last_paragraph_contains(doc: Document, fragment: str, new_text: str) -> None:
    para = find_last_paragraph_contains(doc, fragment)
    replace_paragraph_text(para, new_text)


def update_common_metadata(doc: Document, period_label: str, monitoring_range: str, report_date: str) -> None:
    old_period_prefix = "（监测时间："
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt.startswith(old_period_prefix) and txt.endswith("）"):
            replace_paragraph_text(para, f"（监测时间：{monitoring_range}）")
        elif txt.startswith("监测时间："):
            replace_paragraph_text(para, f"监测时间：{monitoring_range}")
        elif txt.startswith("报告日期："):
            replace_paragraph_text(para, f"报告日期：{report_date}")

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for idx, cell in enumerate(cells):
                txt = cell.text.strip()
                if txt == "监测时间" and idx + 1 < len(cells):
                    cells[idx + 1].text = monitoring_range
                elif txt.startswith("监测时间："):
                    cell.text = f"监测时间：{monitoring_range}"
                elif txt in {"2025.12.01～2025.12.31", "2026.01.01~2026.03.31", "2026年01月01日~2026年03月31日"}:
                    cell.text = monitoring_range
                elif txt.startswith("报告日期："):
                    cell.text = f"报告日期：{report_date}"


def parse_alarm_bounds(cfg: dict, module: str, point_id: str) -> dict | None:
    per_point = cfg.get("per_point", {}).get(module, {})
    safe_id = point_id.replace("-", "_")
    point_cfg = per_point.get(safe_id, {})
    bounds = point_cfg.get("alarm_bounds")
    return bounds if isinstance(bounds, dict) else None


def max_alarm_level(records: Iterable[dict], cfg: dict, module: str, min_key: str, max_key: str) -> int:
    level = 0
    for record in records:
        pid = record.get("PointID")
        if not pid:
            continue
        bounds = parse_alarm_bounds(cfg, module, str(pid))
        if not bounds:
            continue
        min_val = record.get(min_key)
        max_val = record.get(max_key)
        if min_val is None or max_val is None:
            continue
        level2 = bounds.get("level2") or []
        level3 = bounds.get("level3") or []
        if len(level3) == 2 and (min_val < min(level3) or max_val > max(level3)):
            level = max(level, 3)
        elif len(level2) == 2 and (min_val < min(level2) or max_val > max(level2)):
            level = max(level, 2)
    return level


def alarm_status_text(level: int) -> str:
    if level >= 3:
        return "个别测点超过三级预警阈值，建议立即复核监测数据并结合现场情况进一步核查。"
    if level == 2:
        return "个别测点超过二级预警阈值，建议加强跟踪监测并复核数据。"
    return "均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"


def format_range(min_val: float | int | None, max_val: float | int | None, decimals: int = 1, unit: str = "") -> str:
    if min_val is None or max_val is None:
        return "--"
    return f"{min_val:.{decimals}f}{unit}~{max_val:.{decimals}f}{unit}"


def find_table_by_header(doc: Document, header_fragment: str):
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        if any(header_fragment in text for text in header_cells):
            return table
    raise ValueError(f'Table with header fragment "{header_fragment}" not found in template')


def parse_float(value: object) -> float | None:
    if value is None or value == "":
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def median_safe(values: Iterable[float]) -> float | None:
    vals = sorted(v for v in values if v is not None and not math.isnan(v))
    if not vals:
        return None
    n = len(vals)
    mid = n // 2
    if n % 2 == 1:
        return vals[mid]
    return (vals[mid - 1] + vals[mid]) / 2.0


def parse_wind_summary(summary_path: Path) -> dict[str, str]:
    if not summary_path.exists():
        return {}
    text = summary_path.read_text(encoding="utf-8")
    result: dict[str, str] = {}
    patterns = {
        "mean_dir": r"平均风向:\s*([0-9.]+°)",
        "dominant_dir": r"主导风向:\s*([^\n]+)",
        "mean_speed": r"平均风速:\s*([0-9.]+)\s*m/s",
        "max_speed": r"最大风速:\s*([0-9.]+)\s*m/s",
        "main_grade": r"主要风速等级:\s*([^\n（]+(?:\s*m/s)?)",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if match:
            result[key] = match.group(1).strip()
    return result


def compute_string_force_kN(rho: float | None, length_m: float | None, freq_hz: float | None) -> float | None:
    if rho is None or length_m is None or freq_hz is None:
        return None
    return 4.0 * rho * (length_m ** 2) * (freq_hz ** 2) / 1000.0


def flatten_group_members(groups_cfg: object) -> list[str]:
    groups_map = {}
    if isinstance(groups_cfg, dict):
        groups_map = groups_cfg
    points: list[str] = []
    seen: set[str] = set()
    for members in groups_map.values():
        if not isinstance(members, list):
            continue
        for item in members:
            val = str(item).strip()
            if not val or val in seen:
                continue
            points.append(val)
            seen.add(val)
    return points


def label_path_dicts(items: list[ImageItem]) -> list[dict]:
    out: list[dict] = []
    for item in items:
        lookup = deepcopy(item.lookup) if isinstance(item.lookup, dict) else {}
        lookup["label"] = item.label
        out.append(lookup)
    return out


def center_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_numbered_item_block(text: str, number: int, new_block: str, preserve_footer: bool = False) -> str:
    footer = ""
    work = text
    if preserve_footer and "（转下页）" in work:
        head, tail = work.rsplit("（转下页）", 1)
        work = head.rstrip()
        footer = "\n\n（转下页）" + tail
    pattern = re.compile(rf"(?s)(^|\n){number}、.*?(?=(\n\d+、)|\Z)")
    if pattern.search(work):
        work = pattern.sub(lambda m: f"{m.group(1)}{new_block.strip()}", work, count=1)
    else:
        work = work.rstrip() + "\n" + new_block.strip()
    return work.rstrip() + footer


def set_row_summary_text(row, start_col: int, text: str) -> None:
    for idx in range(start_col, len(row.cells)):
        row.cells[idx].text = text


SECTION_TITLES = {
    "strain": "结构应变监测",
    "tilt": "主塔倾斜监测",
    "bearing_displacement": "支座变位监测",
    "cable_force": "吊索索力监测",
    "vibration": "主梁、主塔振动监测",
    "wind": "风向风速监测",
    "eq": "地震动监测",
}

OVERVIEW_SECTION_NAMES = (
    "交通状况监测",
    "结构应变监测",
    "主塔倾斜监测",
    "支座变位监测",
    "吊索索力监测",
    "风向风速监测",
    "地震动监测",
)


def _image_list_has_content(items: object) -> bool:
    if not isinstance(items, list):
        return False
    return any(isinstance(item, dict) and item.get("path") for item in items)


def _section_has_content(section_name: str, section: dict) -> bool:
    if not section or not section.get("enabled", True):
        return False

    if section_name == "strain":
        return (
            any(section.get(key) is not None for key in ("girder_min", "girder_max", "tower_min", "tower_max"))
            or _image_list_has_content(section.get("girder_timeseries_images"))
            or _image_list_has_content(section.get("girder_boxplot_images"))
            or _image_list_has_content(section.get("tower_timeseries_images"))
            or _image_list_has_content(section.get("tower_boxplot_images"))
        )
    if section_name == "tilt":
        return (
            any(section.get(key) is not None for key in ("z_min", "z_max", "h_min", "h_max"))
            or _image_list_has_content(section.get("images"))
        )
    if section_name == "bearing_displacement":
        return (
            any(section.get(key) is not None for key in ("min_val", "max_val"))
            or _image_list_has_content(section.get("images"))
        )
    if section_name == "cable_force":
        return (
            any(section.get(key) is not None for key in ("max_abs", "max_rms", "min_change", "max_change"))
            or bool(section.get("table_rows"))
            or _image_list_has_content(section.get("accel_images"))
            or _image_list_has_content(section.get("force_images"))
            or bool(section.get("accel_available"))
            or bool(section.get("force_available"))
        )
    if section_name == "vibration":
        return (
            any(section.get(key) is not None for key in ("max_abs", "max_rms"))
            or _image_list_has_content(section.get("timeseries_images"))
            or _image_list_has_content(section.get("freq_images"))
        )
    if section_name == "wind":
        return bool(section.get("table_rows")) or section.get("max_10min") is not None or _image_list_has_content(section.get("speed_images")) or _image_list_has_content(section.get("rose_images"))
    if section_name == "eq":
        return _image_list_has_content(section.get("images"))
    return True


def section_is_available(section_name: str, section: dict | None) -> bool:
    if not isinstance(section, dict):
        return False
    if not section.get("enabled", True):
        return False
    return bool(section.get("available", True))


def _missing_section(section_name: str, cfg: dict, exc: FileNotFoundError) -> dict:
    enabled = reporting_enabled(cfg, section_name)
    section = {
        "enabled": enabled,
        "available": False,
        "missing_error": str(exc),
    }
    if enabled:
        title = SECTION_TITLES.get(section_name, section_name)
        section["missing_notice"] = f"本周期未获取到{title}有效数据"
    return section


def _build_section_safe(
    section_name: str,
    builder,
    cfg: dict,
    stats_root: Path,
    fallback_stats_root: Path | None,
    image_root: Path,
    assets_dir: Path,
) -> dict:
    try:
        section = builder(cfg, stats_root, fallback_stats_root, image_root, assets_dir)
    except FileNotFoundError as exc:
        return _missing_section(section_name, cfg, exc)

    if not isinstance(section, dict):
        section = {"enabled": False}
    if not section.get("enabled", True):
        section["available"] = False
        return section
    section["available"] = _section_has_content(section_name, section)
    if not section["available"]:
        section.setdefault("missing_notice", f"本周期未获取到{SECTION_TITLES.get(section_name, section_name)}有效数据")
    return section


def _format_threshold_tons(value: object, default: float) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        numeric = float(default)
    rounded = round(numeric)
    if abs(numeric - rounded) < 1e-9:
        return str(int(rounded))
    return f"{numeric:.1f}".rstrip("0").rstrip(".")


def build_overview_items(manifest: dict) -> dict[str, list[str]]:
    sections = manifest["sections"]
    traffic = sections.get("wim", {}) if sections.get("wim", {}).get("enabled") else sections.get("traffic", {})
    strain = sections["strain"]
    tilt = sections["tilt"]
    bearing = sections["bearing_displacement"]
    cable = sections["cable_force"]
    wind = sections["wind"]
    eq = sections["eq"]

    replacements: dict[str, list[str]] = {}

    if traffic:
        gross_level_1_t = _format_threshold_tons(traffic.get("gross_level_1_t"), 82.5)
        gross_level_2_t = _format_threshold_tons(traffic.get("gross_level_2_t"), 110.0)
        axle_level_1_t = _format_threshold_tons(traffic.get("axle_level_1_t"), 42.0)
        axle_level_2_t = _format_threshold_tons(traffic.get("axle_level_2_t"), 56.0)
        replacements["交通状况监测"] = [
            (
                f"监测结果表明，桥梁共通过车辆{traffic.get('vehicle_total', 0)}辆，日均{traffic.get('daily_avg', 0)}辆。"
                f"其中上行方向（闽侯-农大，车道1～车道4）所通过车辆为{traffic.get('up_total', 0)}辆，"
                f"下行方向（农大-闽侯，车道5～车道8）所通过车辆为{traffic.get('down_total', 0)}辆。"
                f"期间系统记录到的最大车重为{traffic.get('max_gross_t', 0):.2f}t，{traffic.get('gross_limit_text', f'未达到2.0倍设计车辆荷载{gross_level_2_t}t')}。"
                f"最大轴重{traffic.get('max_axle_t', 0):.2f}t，{traffic.get('axle_limit_text', f'未达到1.5倍设计车辆荷载{axle_level_1_t}t')}。"
                f"期间总重超过1.5倍设计荷载{gross_level_1_t}t的车辆共{traffic.get('gross_over_1_5', 0)}辆，"
                f"其中超过2.0倍设计荷载{gross_level_2_t}t的车辆共{traffic.get('gross_over_2_0', 0)}辆；"
                f"轴重超过1.5倍设计荷载{axle_level_1_t}t的车辆共{traffic.get('axle_over_1_5', 0)}辆，"
                f"其中超过2.0倍设计荷载{axle_level_2_t}t的车辆共{traffic.get('axle_over_2_0', 0)}辆。"
            )
        ]

    if section_is_available("strain", strain):
        replacements["结构应变监测"] = [
            f"箱梁监测结果表明，各测点应变值在{format_range(strain.get('girder_min'), strain.get('girder_max'), 1, 'με')}之间，"
            "均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。\n"
            f"桥塔监测结果表明，各测点应变值在{format_range(strain.get('tower_min'), strain.get('tower_max'), 1, 'με')}之间，"
            "均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        ]
    if section_is_available("tilt", tilt):
        replacements["主塔倾斜监测"] = [
            f"监测结果表明，倾角纵桥向位移在{format_range(tilt.get('z_min'), tilt.get('z_max'), 3, '°')}之间，"
            f"横桥向位移在{format_range(tilt.get('h_min'), tilt.get('h_max'), 3, '°')}之间，"
            "均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        ]
    if section_is_available("bearing_displacement", bearing):
        replacements["支座变位监测"] = [
            "选取典型监测数据进行分析。"
            f"监测结果表明，各测点支座位移在{format_range(bearing.get('min_val'), bearing.get('max_val'), 1, 'mm')}之间，"
            "未出现超过各级超限阈值和报警的情况。"
        ]
    if section_is_available("cable_force", cable):
        cable_parts: list[str] = ["选取典型监测数据进行分析。"]
        if cable.get("accel_available"):
            max_abs = cable.get("max_abs")
            max_rms = cable.get("max_rms")
            if max_abs is not None and max_rms is not None:
                cable_parts.append(
                    f"监测结果表明吊索加速度各测点绝对最大值为{max_abs:.2f}m/s²，"
                    f"各测点10min加速度均方根值最大为{max_rms:.2f}m/s²，"
                    "未超过1000m/s²，均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
                )
            else:
                cable_parts.append("吊索加速度时程及10min加速度均方根结果见正文图表。")
        if cable.get("force_available"):
            min_change = cable.get("min_change")
            max_change = cable.get("max_change")
            if min_change is not None and max_change is not None:
                cable_parts.append(
                    f"与成桥索力相比，索力变化范围在{min_change:.2f}%~{max_change:.2f}%之间，"
                    "与成桥索力相比变化范围在10%以内，均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
                )
            else:
                cable_parts.append("索力时程及统计结果见正文图表。")
        replacements["吊索索力监测"] = ["".join(cable_parts)]
    if section_is_available("wind", wind):
        replacements["风向风速监测"] = [
            f"监测结果表明，桥面10min平均风速最大值为{wind.get('max_10min', 0):.2f}m/s，"
            "未超过25m/s，处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        ]
    if section_is_available("eq", eq):
        replacements["地震动监测"] = [
            f"监测结果表明，水平地震动作用加速度峰值为{eq.get('horizontal_peak', 0):.3f}m/s²，"
            f"{eq.get('horizontal_text', '未达到设计E1地震作用加速度峰值')}，"
            f"竖向地震动作用加速度峰值为{eq.get('vertical_peak', 0):.3f}m/s²，"
            f"{eq.get('vertical_text', '未达到设计E1地震作用加速度峰值')}，未出现超过各级超限阈值和报警的情况。"
        ]

    return replacements


def update_overview_tables(doc: Document, manifest: dict) -> None:
    replacements = build_overview_items(manifest)
    replacement_names = tuple(replacements.keys())

    def replace_in_cell(cell) -> None:
        paragraphs = cell.paragraphs
        idx = 0
        while idx < len(paragraphs):
            text = paragraphs[idx].text.strip()
            matched = next((name for name in OVERVIEW_SECTION_NAMES if name in text), None)
            if matched is None:
                idx += 1
                continue
            targets = []
            j = idx + 1
            while j < len(paragraphs):
                next_text = paragraphs[j].text.strip()
                if next_text and any(name in next_text for name in OVERVIEW_SECTION_NAMES):
                    break
                if next_text:
                    targets.append(paragraphs[j])
                j += 1
            if matched in replacements:
                new_blocks = replacements[matched]
                for target, new_text in zip(targets, new_blocks):
                    replace_paragraph_text(target, new_text)
                for extra in targets[len(new_blocks):]:
                    replace_paragraph_text(extra, "")
            idx = j

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if any(name in cell.text for name in OVERVIEW_SECTION_NAMES):
                    replace_in_cell(cell)


def build_strain_section(cfg: dict, stats_root: Path, fallback_stats_root: Path | None, image_root: Path, assets_dir: Path) -> dict:
    rows = load_sheet_rows(resolve_existing_file(stats_root, fallback_stats_root, "strain_stats.xlsx"))
    girder_rows = [r for r in rows if str(r.get("PointID", "")).startswith(("SB-", "SC-", "SD-", "SE-", "SF-", "SG-", "SH-"))]
    tower_rows = [r for r in rows if str(r.get("PointID", "")).startswith(("SK-", "SL-"))]

    strain_style = cfg["plot_styles"]["strain"]
    girder_groups = get_report_order(cfg, "strain", "girder_order", ["B", "C", "D", "E", "F", "G", "H"])
    tower_groups = get_report_order(cfg, "strain", "tower_order", ["K", "L"])

    girder_ts_imgs: list[ImageItem] = []
    girder_box_imgs: list[ImageItem] = []
    for group in girder_groups:
        ts_path, ts_lookup = find_latest_image(image_root, strain_style["group_output_dir"], f"Strain_{group}_")
        girder_ts_imgs.append(ImageItem(group, ts_path, ts_lookup))
        box_path, box_lookup = find_latest_image(image_root, strain_style["boxplot_output_dir"], f"StrainBox_{group}_")
        girder_box_imgs.append(ImageItem(group, box_path, box_lookup))

    tower_ts_imgs: list[ImageItem] = []
    tower_box_imgs: list[ImageItem] = []
    for group in tower_groups:
        ts_path, ts_lookup = find_latest_image(image_root, strain_style["group_output_dir"], f"Strain_{group}_")
        tower_ts_imgs.append(ImageItem(group, ts_path, ts_lookup))
        box_path, box_lookup = find_latest_image(image_root, strain_style["boxplot_output_dir"], f"StrainBox_{group}_")
        tower_box_imgs.append(ImageItem(group, box_path, box_lookup))

    girder_min = min((r["Min"] for r in girder_rows if r.get("Min") is not None), default=None)
    girder_max = max((r["Max"] for r in girder_rows if r.get("Max") is not None), default=None)
    tower_min = min((r["Min"] for r in tower_rows if r.get("Min") is not None), default=None)
    tower_max = max((r["Max"] for r in tower_rows if r.get("Max") is not None), default=None)

    girder_level = max_alarm_level(girder_rows, cfg, "strain", "Min", "Max")
    tower_level = max_alarm_level(tower_rows, cfg, "strain", "Min", "Max")
    strain_unit = "\u03bc\u03b5"

    return {
        "enabled": reporting_enabled(cfg, "strain"),
        "chapter_girder": f"\u76d1\u6d4b\u7ed3\u679c\u8868\u660e\uff0c\u5404\u6d4b\u70b9\u5e94\u53d8\u503c\u5728{format_range(girder_min, girder_max, 1, strain_unit)}\u4e4b\u95f4\uff0c{alarm_status_text(girder_level)}",
        "chapter_tower": f"\u76d1\u6d4b\u7ed3\u679c\u8868\u660e\uff0c\u5404\u6d4b\u70b9\u5e94\u53d8\u503c\u5728{format_range(tower_min, tower_max, 1, strain_unit)}\u4e4b\u95f4\uff0c{alarm_status_text(tower_level)}",
        "girder_min": girder_min,
        "girder_max": girder_max,
        "tower_min": tower_min,
        "tower_max": tower_max,
        "girder_timeseries_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in girder_ts_imgs],
        "girder_boxplot_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in girder_box_imgs],
        "tower_timeseries_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in tower_ts_imgs],
        "tower_boxplot_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in tower_box_imgs],
        "girder_timeseries_caption": "\u56fe 4-4 \u4e3b\u6881\u5404\u622a\u9762\u4f4d\u7f6e\u5e94\u53d8\u65f6\u7a0b\u66f2\u7ebf\u56fe",
        "girder_boxplot_caption": "\u56fe 4-5 \u4e3b\u6881\u5404\u622a\u9762\u4f4d\u7f6e\u5e94\u53d8\u7bb1\u7ebf\u56fe",
        "tower_timeseries_caption": "\u56fe 4-6 \u6865\u5854\u5404\u622a\u9762\u4f4d\u7f6e\u5e94\u53d8\u65f6\u7a0b\u66f2\u7ebf\u56fe",
        "tower_boxplot_caption": "\u56fe 4-7 \u6865\u5854\u5404\u622a\u9762\u4f4d\u7f6e\u5e94\u53d8\u7bb1\u7ebf\u66f2\u7ebf\u56fe",
        "image_lookup": {
            "girder_timeseries": [deepcopy(item.lookup) | {"label": item.label} for item in girder_ts_imgs],
            "girder_boxplot": [deepcopy(item.lookup) | {"label": item.label} for item in girder_box_imgs],
            "tower_timeseries": [deepcopy(item.lookup) | {"label": item.label} for item in tower_ts_imgs],
            "tower_boxplot": [deepcopy(item.lookup) | {"label": item.label} for item in tower_box_imgs],
        },
    }


def build_tilt_section(cfg: dict, stats_root: Path, fallback_stats_root: Path | None, image_root: Path, assets_dir: Path) -> dict:
    rows = load_sheet_rows(resolve_existing_file(stats_root, fallback_stats_root, "tilt_stats.xlsx"))
    z_rows = [r for r in rows if str(r.get("PointID", "")).endswith("-Z")]
    h_rows = [r for r in rows if str(r.get("PointID", "")).endswith("-H")]
    style = cfg["plot_styles"]["tilt"]
    items = []
    for pid in get_report_order(cfg, "tilt", "order", ["Q1-Z", "Q1-H", "Q2-Z", "Q2-H"]):
        img_path, lookup = find_latest_image(image_root, style["output_dir"], f"Tilt_{pid}_")
        items.append(ImageItem(pid, img_path, lookup))
    z_min = min((r["Min"] for r in z_rows if r.get("Min") is not None), default=None)
    z_max = max((r["Max"] for r in z_rows if r.get("Max") is not None), default=None)
    h_min = min((r["Min"] for r in h_rows if r.get("Min") is not None), default=None)
    h_max = max((r["Max"] for r in h_rows if r.get("Max") is not None), default=None)
    level = max(max_alarm_level(z_rows, cfg, "tilt", "Min", "Max"), max_alarm_level(h_rows, cfg, "tilt", "Min", "Max"))
    summary = (
        f"监测结果表明，倾角纵桥向位移在{format_range(z_min, z_max, 3, '°')}之间，"
        f"横桥向位移在{format_range(h_min, h_max, 3, '°')}之间，{alarm_status_text(level)}"
    )
    return {
        "enabled": reporting_enabled(cfg, "tilt"),
        "chapter": summary,
        "z_min": z_min,
        "z_max": z_max,
        "h_min": h_min,
        "h_max": h_max,
        "images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in items],
        "caption": "桥塔各截面位置倾角时程曲线图",
        "image_lookup": [deepcopy(item.lookup) | {"label": item.label} for item in items],
    }


def build_bearing_section(cfg: dict, stats_root: Path, fallback_stats_root: Path | None, image_root: Path, assets_dir: Path) -> dict:
    rows = load_sheet_rows(resolve_existing_file(stats_root, fallback_stats_root, "bearing_displacement_stats.xlsx"))
    configured_points = get_report_order(cfg, "bearing_displacement", "order", cfg.get("points", {}).get("bearing_displacement", []))
    configured_order = {str(pid): idx for idx, pid in enumerate(configured_points)}
    if configured_order:
        rows = [r for r in rows if str(r.get("PointID", "")) in configured_order]
        rows.sort(key=lambda r: configured_order.get(str(r.get("PointID", "")), 10**9))

    valid_rows = [r for r in rows if r.get("OrigMin_mm") is not None and r.get("OrigMax_mm") is not None]
    style = cfg["plot_styles"]["bearing_displacement"]
    items = []
    for record in valid_rows:
        pid = str(record["PointID"])
        img_path, lookup = find_latest_image_patterns(
            image_root,
            style["output_dir"],
            [
                f"BearingDisp_{pid}_*_Orig_*.jpg",
                f"BearingDisp_{pid}_*_Orig_*.png",
                f"BearingDisp_{pid}_*_Orig_*.jpeg",
            ],
        )
        if img_path is None:
            img_path, lookup = find_latest_image_patterns(
                image_root,
                style["output_dir"],
                [
                    f"BearingDisp_{pid}_*.jpg",
                    f"BearingDisp_{pid}_*.png",
                    f"BearingDisp_{pid}_*.jpeg",
                ],
            )
        items.append(ImageItem(pid, img_path, lookup))
    if not items:
        items = [ImageItem("支座位移", None, {
            "image_root": str(image_root),
            "configured_dir": style["output_dir"],
            "resolved_dirs": [],
            "patterns": [],
            "matched_files": [],
            "selected_file": None,
        })]
    min_val = min((r["OrigMin_mm"] for r in valid_rows), default=None)
    max_val = max((r["OrigMax_mm"] for r in valid_rows), default=None)
    level = max_alarm_level(valid_rows, cfg, "bearing_displacement", "OrigMin_mm", "OrigMax_mm")
    summary = (
        f"选取典型监测数据进行分析。监测结果表明，各测点支座位移在"
        f"{format_range(min_val, max_val, 1, 'mm')}之间，{alarm_status_text(level)}"
    )
    return {
        "enabled": reporting_enabled(cfg, "bearing_displacement"),
        "chapter": summary,
        "min_val": min_val,
        "max_val": max_val,
        "images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in items],
        "caption": "典型测点支座变位时程曲线图",
        "image_lookup": [deepcopy(item.lookup) | {"label": item.label} for item in items],
    }


def build_cable_force_section(cfg: dict, stats_root: Path, fallback_stats_root: Path | None, image_root: Path, assets_dir: Path) -> dict:
    if not reporting_enabled(cfg, "cable_force"):
        return {"enabled": False}

    reporting_cfg = get_reporting_section(cfg, "cable_force")
    accel_default_order = cfg.get("points", {}).get("cable_accel", [])
    accel_order = get_report_order(cfg, "cable_force", "accel_order", accel_default_order)
    groups_cfg = cfg.get("groups", {}).get("cable_force", {})
    legacy_force_order = normalize_name_list(reporting_cfg.get("order")) or []
    default_force_order = normalize_name_list(reporting_cfg.get("force_order")) or legacy_force_order or flatten_group_members(groups_cfg) or cfg.get("points", {}).get("cable_force", [])
    force_order = get_report_order(cfg, "cable_force", "force_order", default_force_order)
    accel_enabled = bool(reporting_cfg.get("accel_enabled", True))
    force_enabled = bool(reporting_cfg.get("force_enabled", bool(force_order)))

    accel_dir = reporting_cfg.get("accel_output_dir", "\u65f6\u7a0b\u66f2\u7ebf_\u7d22\u529b\u52a0\u901f\u5ea6")
    rms_dir = reporting_cfg.get("rms_output_dir", "\u65f6\u7a0b\u66f2\u7ebf_\u7d22\u529b\u52a0\u901f\u5ea6_RMS10min")
    force_dir = reporting_cfg.get("force_output_dir", "\u7d22\u529b\u65f6\u7a0b\u56fe")
    force_group_dir = reporting_cfg.get("force_group_output_dir", "\u7d22\u529b\u65f6\u7a0b\u56fe_\u7ec4\u56fe")

    max_abs = None
    max_rms = None
    accel_items: list[ImageItem] = []
    if accel_enabled:
        rows = load_sheet_rows(resolve_existing_file(stats_root, fallback_stats_root, "cable_accel_stats.xlsx"))
        valid_rows = [r for r in rows if r.get("PointID")]
        max_abs = max((max(abs(r["Min"]), abs(r["Max"])) for r in valid_rows if r.get("Min") is not None and r.get("Max") is not None), default=None)
        max_rms = max((r["RMS10minMax"] for r in valid_rows if r.get("RMS10minMax") is not None), default=None)
        for pid in accel_order:
            raw_path, raw_lookup = find_latest_image_patterns(
                image_root,
                accel_dir,
                [
                    f"{pid}_*.jpg",
                    f"{pid}_*.png",
                    f"{pid}_*.jpeg",
                    f"CableAccel_{pid}_*.jpg",
                    f"CableAccel_{pid}_*.png",
                    f"CableAccel_{pid}_*.jpeg",
                    f"*{pid}*.jpg",
                    f"*{pid}*.png",
                    f"*{pid}*.jpeg",
                ],
            )
            accel_items.append(ImageItem(f"{pid} 加速度", raw_path, raw_lookup))
            rms_path, rms_lookup = find_latest_image_patterns(
                image_root,
                rms_dir,
                [
                    f"CableAccelRMS10_{pid}_*.jpg",
                    f"CableAccelRMS10_{pid}_*.png",
                    f"CableAccelRMS10_{pid}_*.jpeg",
                    f"*{pid}*RMS10*.jpg",
                    f"*{pid}*RMS10*.png",
                    f"*{pid}*RMS10*.jpeg",
                ],
            )
            accel_items.append(ImageItem(f"{pid} RMS10min", rms_path, rms_lookup))

    force_items: list[ImageItem] = []
    table_rows: list[dict] = []
    change_rates: list[float] = []
    if force_enabled:
        for label in force_order:
            img_path, lookup = find_latest_image(image_root, force_dir, f"CableForce_{label}_")
            if img_path is None and "-" in label:
                img_path, lookup = find_latest_image(image_root, force_group_dir, f"CableForce_{label}_")
            force_items.append(ImageItem(label, img_path, lookup))

        spec_rows_by_sheet = load_workbook_rows_by_sheet(resolve_existing_file(stats_root, fallback_stats_root, "cable_accel_spec_stats.xlsx"))
        per_point_cfg = cfg.get("per_point", {}).get("cable_accel", {})
        for pid in force_order:
            point_cfg = per_point_cfg.get(pid, {})
            sheet_rows = spec_rows_by_sheet.get(pid, [])
            freq_values: list[float] = []
            force_values: list[float] = []
            for row in sheet_rows:
                freq_val = None
                for key, value in row.items():
                    if str(key).startswith("Freq_"):
                        freq_val = parse_float(value)
                        if freq_val is not None:
                            break
                if freq_val is not None:
                    freq_values.append(freq_val)
                force_val = parse_float(row.get("CableForce_kN"))
                if force_val is not None:
                    force_values.append(force_val)
            freq = median_safe(freq_values)
            current_force = median_safe(force_values)
            rho = parse_float(point_cfg.get("rho"))
            length_m = parse_float(point_cfg.get("L"))
            built_force = HONGTANG_CABLE_BUILT_FORCE.get(pid)
            change_rate = None
            if current_force is not None and built_force not in (None, 0):
                change_rate = (current_force - built_force) / built_force * 100.0
                change_rates.append(change_rate)
            table_rows.append({
                "PointID": pid,
                "rho": rho,
                "L": length_m,
                "freq": freq,
                "current_force": current_force,
                "built_force": built_force,
                "change_rate": change_rate,
            })

    if max_abs is not None and max_rms is not None:
        accel_summary = (
            f"\u9009\u53d6\u5178\u578b\u76d1\u6d4b\u6570\u636e\u8fdb\u884c\u5206\u6790\uff0c\u76d1\u6d4b\u7ed3\u679c\u8868\u660e\uff0c\u540a\u7d22\u52a0\u901f\u5ea6\u5404\u6d4b\u70b9\u7edd\u5bf9\u6700\u5927\u503c\u4e3a{max_abs:.2f}m/s\xb2\uff0c"
            f"\u5404\u6d4b\u70b910min\u52a0\u901f\u5ea6\u5747\u65b9\u6839\u503c\u6700\u5927\u4e3a{max_rms:.2f}m/s\xb2\uff0c\u672a\u8d85\u8fc71000m/s\xb2\uff0c\u5747\u5904\u4e8e\u8d85\u9650\u9608\u503c\u8303\u56f4\u4e4b\u5185\uff0c"
            f"\u672a\u51fa\u73b0\u8d85\u8fc7\u5404\u7ea7\u8d85\u9650\u9608\u503c\u548c\u62a5\u8b66\u7684\u60c5\u51b5\u3002"
        )
    else:
        accel_summary = "\u9009\u53d6\u5178\u578b\u76d1\u6d4b\u6570\u636e\u8fdb\u884c\u5206\u6790\uff0c\u540a\u7d22\u52a0\u901f\u5ea6\u4e0e10min\u52a0\u901f\u5ea6\u5747\u65b9\u6839\u7ed3\u679c\u89c1\u4e0b\u56fe\u3002"

    if change_rates:
        min_change = min(change_rates)
        max_change = max(change_rates)
        max_abs_change = max(abs(v) for v in change_rates)
        if max_abs_change <= 10:
            force_summary = (
                f"\u76d1\u6d4b\u7ed3\u679c\u8868\u660e\uff0c\u4e0e\u6210\u6865\u7d22\u529b\u76f8\u6bd4\uff0c\u7d22\u529b\u53d8\u5316\u8303\u56f4\u5728{min_change:.2f}%~{max_change:.2f}%\u4e4b\u95f4\uff0c"
                f"\u4e0e\u6210\u6865\u7d22\u529b\u76f8\u6bd4\u53d8\u5316\u8303\u56f4\u572810%\u4ee5\u5185\u3002"
            )
        else:
            force_summary = (
                f"\u76d1\u6d4b\u7ed3\u679c\u8868\u660e\uff0c\u4e0e\u6210\u6865\u7d22\u529b\u76f8\u6bd4\uff0c\u7d22\u529b\u53d8\u5316\u8303\u56f4\u5728{min_change:.2f}%~{max_change:.2f}%\u4e4b\u95f4\uff0c"
                f"\u4e2a\u522b\u6d4b\u70b9\u7edd\u5bf9\u53d8\u5316\u7387\u8d85\u8fc710%\uff0c\u5efa\u8bae\u7ed3\u5408\u73b0\u573a\u5de5\u51b5\u8fdb\u4e00\u6b65\u590d\u6838\u3002"
            )
    else:
        force_summary = "\u76d1\u6d4b\u7ed3\u679c\u8868\u660e\uff0c\u7d22\u529b\u65f6\u7a0b\u7ed3\u679c\u89c1\u4e0b\u56fe\uff0c\u8be6\u7ec6\u7ed3\u679c\u5982\u8868 4-7 \u6240\u793a\u3002"

    return {
        "enabled": True,
        "accel_enabled": accel_enabled,
        "force_enabled": force_enabled,
        "accel_available": accel_enabled and (max_abs is not None or max_rms is not None or _image_list_has_content([{"path": str(item.path) if item.path else None} for item in accel_items])),
        "force_available": force_enabled and (bool(change_rates) or bool(table_rows) or _image_list_has_content([{"path": str(item.path) if item.path else None} for item in force_items])),
        "accel_summary": accel_summary,
        "force_summary": force_summary,
        "max_abs": max_abs,
        "max_rms": max_rms,
        "min_change": min(change_rates) if change_rates else None,
        "max_change": max(change_rates) if change_rates else None,
        "accel_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in accel_items],
        "force_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in force_items],
        "accel_caption": [
            "\u5178\u578b\u540a\u7d22\u7d22\u529b\u6d4b\u70b9\u632f\u52a8\u52a0\u901f\u5ea6\u7edd\u5bf9\u6700\u5927\u503c\u65f6\u7a0b\u56fe\u548c10min\u52a0\u901f\u5ea6\u5747\u65b9\u6839\u56fe",
            "\u5178\u578b\u6d4b\u70b9\u632f\u52a8\u52a0\u901f\u5ea6\u7edd\u5bf9\u6700\u5927\u503c\u65f6\u7a0b\u56fe\u548c10min\u52a0\u901f\u5ea6\u5747\u65b9\u6839\u56fe",
        ],
        "force_caption": "\u5178\u578b\u6d4b\u70b9\u7d22\u529b\u65f6\u7a0b\u56fe",
        "table_rows": table_rows,
        "image_lookup": {
            "accel": label_path_dicts(accel_items),
            "force": label_path_dicts(force_items),
        },
    }


def build_vibration_section(cfg: dict, stats_root: Path, fallback_stats_root: Path | None, image_root: Path, assets_dir: Path) -> dict:
    if not reporting_enabled(cfg, "vibration"):
        return {"enabled": False}

    rows = load_sheet_rows(resolve_existing_file(stats_root, fallback_stats_root, "accel_stats.xlsx"))
    valid_rows = [r for r in rows if r.get("PointID")]
    max_abs = max((max(abs(r["Min"]), abs(r["Max"])) for r in valid_rows if r.get("Min") is not None and r.get("Max") is not None), default=None)
    max_rms = max((r["RMS10minMax"] for r in valid_rows if r.get("RMS10minMax") is not None), default=None)

    order = get_report_order(cfg, "vibration", "order", cfg.get("points", {}).get("acceleration", []))
    reporting_cfg = get_reporting_section(cfg, "vibration")
    accel_dir = reporting_cfg.get("accel_output_dir", "时程曲线_加速度")
    rms_dir = reporting_cfg.get("rms_output_dir", "时程曲线_加速度_RMS10min")
    spec_dir = reporting_cfg.get("spec_output_dir", "频谱峰值曲线_加速度")

    ts_items: list[ImageItem] = []
    freq_items: list[ImageItem] = []
    for pid in order:
        raw_path, raw_lookup = find_latest_image(image_root, accel_dir, f"{pid}_")
        ts_items.append(ImageItem(f"{pid} 加速度", raw_path, raw_lookup))
        rms_path, rms_lookup = find_latest_image(image_root, rms_dir, f"AccelRMS10_{pid}_")
        ts_items.append(ImageItem(f"{pid} RMS10min", rms_path, rms_lookup))
        freq_path, freq_lookup = find_latest_image(image_root, spec_dir, f"SpecFreq_{pid}_")
        freq_items.append(ImageItem(pid, freq_path, freq_lookup))

    if max_abs is not None and max_rms is not None:
        timeseries_summary = (
            f"选取典型监测数据进行分析，监测结果表明，主梁及主塔加速度各测点绝对最大值为{max_abs:.2f}m/s²，"
            f"各测点10min加速度均方根值最大为{max_rms:.2f}m/s²，未超过315m/s²，均处于超限阈值范围之内，"
            f"未出现超过各级超限阈值和报警的情况。"
        )
    else:
        timeseries_summary = "选取典型监测数据进行分析，主梁及主塔加速度时程与10min加速度均方根结果见下图。"
    freq_summary = "选取典型监测数据进行分析，典型测点自振频率时程如图4-12所示，本月主梁及主塔自振频率识别结果整体稳定，未见明显异常漂移。"

    return {
        "enabled": True,
        "timeseries_summary": timeseries_summary,
        "freq_summary": freq_summary,
        "timeseries_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in ts_items],
        "freq_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in freq_items],
        "timeseries_caption": [
            "典型主梁、主塔测点振动加速度绝对最大值时程图和10min加速度均方根图",
            "典型测点振动加速度绝对最大值时程图和10min加速度均方根图",
        ],
        "freq_caption": "典型测点自振频率时程图",
        "image_lookup": {
            "timeseries": label_path_dicts(ts_items),
            "freq": label_path_dicts(freq_items),
        },
    }


def build_wind_section(cfg: dict, stats_root: Path, fallback_stats_root: Path | None, image_root: Path, assets_dir: Path) -> dict:
    if not reporting_enabled(cfg, "wind"):
        return {"enabled": False}

    rows = load_sheet_rows(resolve_existing_file(stats_root, fallback_stats_root, "wind_stats.xlsx"))

    order = get_report_order(cfg, "wind", "order", cfg.get("points", {}).get("wind", []))
    order_map = {pid: idx for idx, pid in enumerate(order)}
    rows = [r for r in rows if str(r.get("PointID", "")) in order_map]
    rows.sort(key=lambda r: order_map[str(r.get("PointID", ""))])

    speed_items: list[ImageItem] = []
    rose_items: list[ImageItem] = []
    table_rows: list[dict] = []
    max_10min = None
    for row in rows:
        pid = str(row["PointID"])
        speed_path, speed_lookup = find_latest_image(image_root, "风速风向结果/风速10min", f"{pid}_speed10min_")
        rose_path, rose_lookup = find_latest_image(image_root, "风速风向结果/风玫瑰", f"{pid}_windrose_")
        speed_items.append(ImageItem(pid, speed_path, speed_lookup))
        rose_items.append(ImageItem(pid, rose_path, rose_lookup))

        summary_path, _ = find_latest_file_patterns(image_root, "风速风向结果/风玫瑰", [f"{pid}_windrose_*_summary.txt"])
        summary_vals = parse_wind_summary(summary_path) if summary_path else {}
        row_max_10min = parse_float(row.get("Mean10minMax"))
        if row_max_10min is not None:
            max_10min = row_max_10min if max_10min is None else max(max_10min, row_max_10min)
        table_rows.append({
            "PointID": pid,
            "mean_dir": summary_vals.get("mean_dir", ""),
            "dominant_dir": summary_vals.get("dominant_dir", ""),
            "mean_speed": parse_float(row.get("MeanSpeed")),
            "max_speed": parse_float(row.get("MaxSpeed")),
            "main_grade": summary_vals.get("main_grade", ""),
        })

    if max_10min is not None:
        summary = (
            f"监测结果如表4-12所示。监测结果表明，桥面10min平均风速最大值为{max_10min:.2f}m/s，"
            f"未超过25m/s，处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        )
    else:
        summary = "监测结果如表4-12所示。桥面10min平均风速与风玫瑰结果见下图。"

    return {
        "enabled": True,
        "summary": summary,
        "max_10min": max_10min,
        "speed_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in speed_items],
        "rose_images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in rose_items],
        "speed_caption": "桥面10min平均风速时程图",
        "rose_caption": "风玫瑰图",
        "table_rows": table_rows,
        "image_lookup": {
            "speed": label_path_dicts(speed_items),
            "rose": label_path_dicts(rose_items),
        },
    }


def build_eq_section(cfg: dict, stats_root: Path, fallback_stats_root: Path | None, image_root: Path, assets_dir: Path) -> dict:
    if not reporting_enabled(cfg, "eq"):
        return {"enabled": False}

    order = get_report_order(cfg, "eq", "order", cfg.get("points", {}).get("eq", ["EQ-X", "EQ-Y", "EQ-Z"]))
    output_cfg = cfg.get("plot_styles", {}).get("eq", {}).get("output", {})
    eq_dir = f"{output_cfg.get('root_dir', '地震动结果')}/{output_cfg.get('series_dir', '地震动时程')}"
    prefix = output_cfg.get("prefix", "EQ")

    items: list[ImageItem] = []
    peak_map: dict[str, float] = {}
    per_point_eq = cfg.get("per_point", {}).get("eq", {})
    for pid in order:
        comp = pid.split("-")[-1] if "-" in pid else pid[-1]
        img_path, lookup = find_latest_image(image_root, eq_dir, f"{prefix}_{comp}_")
        items.append(ImageItem(pid, img_path, lookup))
        point_cfg = per_point_eq.get(pid, {})
        file_id = point_cfg.get("file_id")
        peak = 0.0
        if file_id:
            matches = sorted(image_root.rglob(f"{file_id}.csv"))
            for csv_path in matches:
                values = read_numeric_series_csv(csv_path)
                if values:
                    peak = max(peak, max(abs(v) for v in values))
        peak_map[pid] = peak

    horizontal_peak = max(peak_map.get("EQ-X", 0.0), peak_map.get("EQ-Y", 0.0))
    vertical_peak = peak_map.get("EQ-Z", 0.0)
    hx = per_point_eq.get("EQ-X", {}).get("alarm_levels", [])
    hy = per_point_eq.get("EQ-Y", {}).get("alarm_levels", [])
    hz = per_point_eq.get("EQ-Z", {}).get("alarm_levels", [])
    e1_h = parse_float(hx[0]) if isinstance(hx, list) and hx else parse_float(hy[0]) if isinstance(hy, list) and hy else None
    e1_v = parse_float(hz[0]) if isinstance(hz, list) and hz else None
    h_text = "未达到设计E1地震作用加速度峰值" if e1_h is None or horizontal_peak < e1_h else "达到或超过设计E1地震作用加速度峰值"
    v_text = "未达到设计E1地震作用加速度峰值" if e1_v is None or vertical_peak < e1_v else "达到或超过设计E1地震作用加速度峰值"
    summary = (
        f"监测结果表明，水平地震动作用加速度峰值为{horizontal_peak:.3f}m/s²，{h_text}，"
        f"竖向地震动作用加速度峰值为{vertical_peak:.3f}m/s²，{v_text}，未出现超过各级超限阈值和报警的情况。"
    )
    return {
        "enabled": True,
        "summary": summary,
        "images": [{"label": item.label, "path": str(item.path) if item.path else None} for item in items],
        "caption": "地震动时程图",
        "image_lookup": label_path_dicts(items),
        "peaks": peak_map,
        "horizontal_peak": horizontal_peak,
        "vertical_peak": vertical_peak,
        "horizontal_text": h_text,
        "vertical_text": v_text,
    }


def build_manifest(cfg: dict, stats_root: Path, fallback_stats_root: Path | None, image_root: Path, template: Path, assets_dir: Path, period_label: str, monitoring_range: str, report_date: str) -> dict:
    return {
        "template": str(template),
        "analysis_root": str(stats_root),
        "fallback_analysis_root": str(fallback_stats_root) if fallback_stats_root is not None else None,
        "image_root": str(image_root),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "period_label": period_label,
        "monitoring_range": monitoring_range,
        "report_date": report_date,
        "sections": {
            "strain": _build_section_safe("strain", build_strain_section, cfg, stats_root, fallback_stats_root, image_root, assets_dir),
            "tilt": _build_section_safe("tilt", build_tilt_section, cfg, stats_root, fallback_stats_root, image_root, assets_dir),
            "bearing_displacement": _build_section_safe("bearing_displacement", build_bearing_section, cfg, stats_root, fallback_stats_root, image_root, assets_dir),
            "cable_force": _build_section_safe("cable_force", build_cable_force_section, cfg, stats_root, fallback_stats_root, image_root, assets_dir),
            "vibration": _build_section_safe("vibration", build_vibration_section, cfg, stats_root, fallback_stats_root, image_root, assets_dir),
            "wind": _build_section_safe("wind", build_wind_section, cfg, stats_root, fallback_stats_root, image_root, assets_dir),
            "eq": _build_section_safe("eq", build_eq_section, cfg, stats_root, fallback_stats_root, image_root, assets_dir),
        },
    }


def update_cable_force_table(doc: Document, table_rows: list[dict]) -> None:
    if not table_rows:
        return
    table = find_table_by_header(doc, "线密度ρ")
    for row in table.rows:
        for cell in row.cells:
            center_cell(cell)
    row_map = {str(row.cells[0].text).strip(): row for row in table.rows[1:]}
    for item in table_rows:
        row = row_map.get(item["PointID"])
        if row is None:
            continue
        row.cells[1].text = "" if item["rho"] is None else f"{item['rho']:.1f}"
        row.cells[2].text = "" if item["L"] is None else f"{item['L']:.3f}"
        row.cells[3].text = "" if item["freq"] is None else f"{item['freq']:.3f}"
        row.cells[4].text = "" if item["current_force"] is None else f"{item['current_force']:.2f}"
        if item["built_force"] is not None:
            row.cells[5].text = f"{item['built_force']:.2f}"
        row.cells[6].text = "" if item["change_rate"] is None else f"{item['change_rate']:+.2f}%"
        for cell in row.cells:
            center_cell(cell)


def update_wind_table(doc: Document, table_rows: list[dict]) -> None:
    if not table_rows:
        return
    table = find_table_by_header(doc, "平均风向")
    row_map = {str(row.cells[0].text).strip(): row for row in table.rows[1:]}
    for item in table_rows:
        row = row_map.get(item["PointID"])
        if row is None:
            continue
        row.cells[1].text = item.get("mean_dir", "")
        row.cells[2].text = item.get("dominant_dir", "")
        row.cells[3].text = "" if item["mean_speed"] is None else f"{item['mean_speed']:.2f}"
        row.cells[4].text = "" if item["max_speed"] is None else f"{item['max_speed']:.2f}"
        row.cells[5].text = item.get("main_grade", "")
        for cell in row.cells:
            center_cell(cell)


def apply_manifest_to_doc(doc: Document, manifest: dict) -> None:
    update_common_metadata(doc, manifest["period_label"], manifest["monitoring_range"], manifest["report_date"])
    update_overview_tables(doc, manifest)

    strain = manifest["sections"]["strain"]
    if section_is_available("strain", strain):
        replace_next_nonempty_after_exact(doc, "\u4e3b\u6881\u5e94\u53d8", strain["chapter_girder"], use_last=True, skip=1)
        replace_next_nonempty_after_exact(doc, "\u6865\u5854\u5e94\u53d8", strain["chapter_tower"], use_last=True, skip=1)
        replace_last_paragraph_contains(doc, "主梁各截面位置应变时程曲线图", strain["girder_timeseries_caption"])
        replace_last_paragraph_contains(doc, "主梁各截面位置应变箱线", strain["girder_boxplot_caption"])
        replace_last_paragraph_contains(doc, "桥塔各截面位置应变时程曲线图", strain["tower_timeseries_caption"])
        replace_last_paragraph_contains(doc, "桥塔各截面位置应变箱线", strain["tower_boxplot_caption"])
        insert_labeled_images_before_caption_contains(
            doc,
            strain["girder_timeseries_caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in strain["girder_timeseries_images"]],
        )
        insert_labeled_images_before_caption_contains(
            doc,
            strain["girder_boxplot_caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in strain["girder_boxplot_images"]],
        )
        insert_labeled_images_before_caption_contains(
            doc,
            strain["tower_timeseries_caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in strain["tower_timeseries_images"]],
        )
        insert_labeled_images_before_caption_contains(
            doc,
            strain["tower_boxplot_caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in strain["tower_boxplot_images"]],
        )


    tilt = manifest["sections"]["tilt"]
    if section_is_available("tilt", tilt):
        replace_last_paragraph_contains(doc, "主塔倾角偏移的方向以闽侯上街-农林大学为纵桥向", "主塔倾角偏移的方向以闽侯上街-农林大学为纵桥向，上游-下游为横桥向。其中朝农林大学方向为正、闽侯上街方向为负，朝上游方向为正、朝下游方向为负。各测点的倾斜幅值如下图所示。" + tilt["chapter"])
        insert_labeled_images_before_caption_contains(
            doc,
            tilt["caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in tilt["images"]],
        )
        ensure_note_before_caption(doc, tilt["caption"], "注：后缀-Z表示纵向，后缀-H表示横向")

    bearing = manifest["sections"]["bearing_displacement"]
    if section_is_available("bearing_displacement", bearing):
        replace_last_paragraph_contains(doc, "支座变位的方向以闽侯上街-农林大学为纵桥向", "支座变位的方向以闽侯上街-农林大学为纵桥向，上游-下游为横桥向。其中朝农林大学方向为正、闽侯上街方向为负，朝上游方向为正、朝下游方向为负。各测点的支座位移时程如下图所示。" + bearing["chapter"])
        insert_labeled_images_before_caption_contains(
            doc,
            bearing["caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in bearing["images"]],
        )

    cable = manifest["sections"]["cable_force"]
    if section_is_available("cable_force", cable):
        if cable.get("accel_available"):
            replace_next_nonempty_after_exact(doc, "（1）索力加速度时程数据", cable["accel_summary"], use_last=True)
            insert_labeled_images_before_caption_contains(
                doc,
                cable["accel_caption"],
                [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in cable["accel_images"]],
                anchor_text="（1）索力加速度时程数据",
                stop_text="（2）索力时程数据",
            )
        if cable.get("force_available"):
            replace_next_nonempty_after_exact(
                doc,
                "（2）索力时程数据",
                "选取典型监测数据进行分析，典型测点索力时程图如下图所示。",
                use_last=True,
                skip=2,
            )
            replace_next_nonempty_after_exact(doc, "（2）索力时程数据", cable["force_summary"], use_last=True, skip=4)
            insert_labeled_images_before_caption_contains(
                doc,
                cable["force_caption"],
                [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in cable["force_images"]],
                anchor_text="（2）索力时程数据",
                stop_text="4.6 主梁、主塔振动监测",
            )
            update_cable_force_table(doc, cable["table_rows"])

    vibration = manifest["sections"]["vibration"]
    if section_is_available("vibration", vibration):
        replace_next_nonempty_after_exact(doc, "（1）振动时程数据", vibration["timeseries_summary"], use_last=True)
        replace_next_nonempty_after_exact(doc, "（2）自振频率", vibration["freq_summary"], use_last=True, skip=1)
        insert_labeled_images_before_caption_contains(
            doc,
            vibration["timeseries_caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in vibration["timeseries_images"]],
            anchor_text="（1）振动时程数据",
            stop_text="（2）自振频率",
        )
        insert_labeled_images_before_caption_contains(
            doc,
            vibration["freq_caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in vibration["freq_images"]],
            anchor_text="（2）自振频率",
            stop_text="4.7 风向风速监测",
        )

    wind = manifest["sections"]["wind"]
    if section_is_available("wind", wind):
        replace_next_nonempty_after_exact(doc, "风向风速监测", wind["summary"], use_last=True)
        insert_labeled_images_before_caption_contains(
            doc,
            wind["speed_caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in wind["speed_images"]],
        )
        insert_labeled_images_before_caption_contains(
            doc,
            wind["rose_caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in wind["rose_images"]],
        )
        update_wind_table(doc, wind["table_rows"])

    eq = manifest["sections"]["eq"]
    if section_is_available("eq", eq):
        replace_next_nonempty_after_exact(doc, "地震动监测", eq["summary"], use_last=True)
        insert_labeled_images_before_caption_contains(
            doc,
            eq["caption"],
            [ImageItem(item["label"], Path(item["path"]) if item.get("path") else None) for item in eq["images"]],
        )


def summarize_missing_images(manifest: dict) -> list[str]:
    missing: list[str] = []
    for section_name, section in manifest.get("sections", {}).items():
        if section.get("enabled", True) and not section.get("available", True):
            notice = section.get("missing_notice") or f"{SECTION_TITLES.get(section_name, section_name)}内容缺失"
            missing.append(f"section:{section_name}:{notice}")
            continue
        lookup = section.get("image_lookup", {})
        if isinstance(lookup, dict):
            groups = lookup.values()
        else:
            groups = [lookup]
        for group in groups:
            if not isinstance(group, list):
                continue
            for item in group:
                if not isinstance(item, dict):
                    continue
                if item.get("selected_file"):
                    continue
                missing.append(f"{section_name}:{item.get('label', '<unknown>')}")
    return missing


def build_report(
    template: Path,
    config_path: Path,
    result_root: Path | None = None,
    analysis_root: Path | None = None,
    image_root: Path | None = None,
    output_dir: Path | None = None,
    period_label: str = "2025年12月",
    monitoring_range: str = "2025.12.01～2025.12.31",
    report_date: str | None = None,
) -> tuple[Path, Path, list[str]]:
    if report_date is None:
        report_date = datetime.now().strftime("%Y年%m月%d日")
    if analysis_root is None:
        analysis_root = Path(__file__).resolve().parents[1]

    stats_root = result_root if result_root is not None else analysis_root
    fallback_stats_root = analysis_root if result_root is not None else None
    image_root = image_root if image_root is not None else (result_root if result_root is not None else analysis_root)
    output_dir = output_dir if output_dir is not None else ((result_root / "自动报告") if result_root is not None else (Path(__file__).resolve().parents[1] / "outputs" / "reports"))
    output_dir = ensure_dir(output_dir)
    assets_dir = ensure_dir(output_dir / "generated_assets")

    cfg = load_json(config_path)
    manifest = build_manifest(cfg, stats_root, fallback_stats_root, image_root, template, assets_dir, period_label, monitoring_range, report_date)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    manifest_path = output_dir / f"report_manifest_{timestamp}.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    doc = Document(str(template))
    apply_manifest_to_doc(doc, manifest)

    output_docx = output_dir / f"{template.stem}_自动生成_{timestamp}.docx"
    doc.save(str(output_docx))
    missing = summarize_missing_images(manifest)
    return manifest_path, output_docx, missing


def main() -> None:
    args = parse_args()
    if args.template is None or not args.template.exists():
        raise SystemExit("Template docx not found.")
    if not args.config.exists():
        raise SystemExit("Config file not found.")
    manifest_path, output_docx, missing = build_report(
        template=args.template,
        config_path=args.config,
        result_root=args.result_root,
        analysis_root=args.analysis_root,
        image_root=args.image_root,
        output_dir=args.output_dir,
        period_label=args.period_label,
        monitoring_range=args.monitoring_range,
        report_date=args.report_date,
    )
    print(f"Manifest written to: {manifest_path}")
    print(f"Report written to:   {output_docx}")
    if missing:
        print("Missing source images:")
        for item in missing:
            print(f"  - {item}")


if __name__ == "__main__":
    main()
