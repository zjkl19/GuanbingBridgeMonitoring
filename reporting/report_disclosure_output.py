from __future__ import annotations

import hashlib
import json
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

from docx import Document

from workbench.report_disclosures import (
    DISCLOSURE_POLICY_VERSION,
    DisclosureItem,
    report_build_disclosure_item,
)


@dataclass(frozen=True)
class DisclosureReviewRequired(RuntimeError):
    candidates: tuple[DisclosureItem, ...]
    report_path: Path
    manifest_path: Path

    def __str__(self) -> str:
        return f"报告生成器发现{len(self.candidates)}项黄色缺项，需逐项确认后重新生成正式报告。"


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def _read_manifest(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8-sig"))
    if not isinstance(payload, dict):
        raise ValueError("报告生成清单必须是 JSON 对象")
    return payload


def _write_manifest(path: Path, payload: Mapping[str, Any]) -> None:
    path.write_text(
        json.dumps(dict(payload), ensure_ascii=False, indent=2, default=str) + "\n",
        encoding="utf-8",
    )


def _is_analysis_summary_item(
    raw: Mapping[str, Any],
    approved_analysis: Sequence[DisclosureItem],
) -> bool:
    label = str(raw.get("label") or raw.get("item") or "")
    if not label.startswith("analysis:"):
        return False
    folded = label.casefold()
    return any(
        item.source_kind == "analysis_module"
        and (
            (item.module_key and item.module_key.casefold() in folded)
            or (item.label and item.label.casefold() in folded)
        )
        for item in approved_analysis
    )


def _append_disclosure_section(report_path: Path, disclosures: Sequence[DisclosureItem]) -> None:
    if not disclosures:
        return
    document = Document(str(report_path))
    anchor = next(
        (paragraph for paragraph in document.paragraphs if "以下无正文" in paragraph.text),
        None,
    )
    created = []
    heading = document.add_paragraph()
    heading.add_run("缺项与数据完整性披露").bold = True
    created.append(heading)
    intro = document.add_paragraph(
        "本报告按实际有效数据生成。以下项目已经逐项审核确认，未对缺失时段进行插补，"
        "也未沿用模板中的旧期次结果媒体。"
    )
    created.append(intro)
    for index, item in enumerate(disclosures, start=1):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{index}. {item.label}：").bold = True
        reason = item.reason_zh.strip()
        label_prefix = f"{item.label}："
        if reason.startswith(label_prefix):
            reason = reason[len(label_prefix):].lstrip()
        paragraph.add_run(f"{reason} 处置：{item.action_zh}")
        created.append(paragraph)
    if anchor is not None:
        for paragraph in created:
            anchor._p.addprevious(paragraph._p)
    document.save(str(report_path))


def reconcile_and_apply_disclosures(
    *,
    report_type: str,
    report_path: Path,
    manifest_path: Path,
    analysis_manifest_path: Path | None,
    analysis_manifest_sha256: str,
    approved_disclosures: Iterable[Mapping[str, Any]],
) -> tuple[tuple[DisclosureItem, ...], dict[str, Any]]:
    """Fail closed on builder gaps, or apply an exact pre-approved plan."""

    payload = _read_manifest(manifest_path)
    approved: list[DisclosureItem] = []
    approval_records: dict[str, dict[str, Any]] = {}
    for raw in approved_disclosures:
        if not isinstance(raw, Mapping):
            continue
        try:
            item = DisclosureItem(**{
                name: str(raw.get(name) or "")
                for name in DisclosureItem.__dataclass_fields__
            })
        except TypeError:
            continue
        approved.append(item)
        approval_records[item.stable_id] = dict(raw)
    approved_analysis = [item for item in approved if item.source_kind != "report_build"]
    approved_report = {
        item.stable_id: item for item in approved if item.source_kind == "report_build"
    }

    missing_items = list(payload.get("missing_items") or [])
    warnings = list(payload.get("warnings") or [])
    candidates: list[DisclosureItem] = []
    hard: list[str] = []
    candidate_by_index: dict[int, DisclosureItem] = {}
    for index, raw in enumerate(missing_items):
        normalized = raw if isinstance(raw, Mapping) else {"label": str(raw)}
        if _is_analysis_summary_item(normalized, approved_analysis):
            continue
        candidate = report_build_disclosure_item(report_type, normalized)
        if candidate is None:
            hard.append(str(normalized.get("label") or normalized))
            continue
        candidates.append(candidate)
        candidate_by_index[index] = candidate
    if warnings:
        hard.extend(f"报告警告：{item}" for item in warnings)
    if hard:
        raise RuntimeError(
            "报告生成器发现不可人工放行的缺项或警告：" + "；".join(hard[:20])
        )

    actual_report = {item.stable_id: item for item in candidates}
    missing_approvals = sorted(set(actual_report) - set(approved_report))
    stale_approvals = sorted(set(approved_report) - set(actual_report))
    if missing_approvals:
        review_payload = dict(payload)
        review_payload.update({
            "status": "disclosure_review_required",
            "analysis_manifest": {
                "path": str(analysis_manifest_path or ""),
                "sha256": str(analysis_manifest_sha256).upper(),
            },
            "disclosure_policy_version": DISCLOSURE_POLICY_VERSION,
            "disclosure_candidates": [item.to_dict() for item in candidates],
            "disclosure_count": len(candidates),
        })
        _write_manifest(manifest_path, review_payload)
        raise DisclosureReviewRequired(tuple(candidates), report_path, manifest_path)
    if stale_approvals:
        raise RuntimeError(
            "已确认的报告缺项与本次生成结果不一致，必须重新预检查："
            + ", ".join(stale_approvals)
        )

    exact_disclosures = (*approved_analysis, *candidates)
    if exact_disclosures:
        expected_sha = str(analysis_manifest_sha256).upper()
        for item in exact_disclosures:
            record = approval_records.get(item.stable_id, {})
            if str(record.get("analysis_manifest_sha256") or "").upper() != expected_sha:
                raise RuntimeError(f"缺项确认未绑定当前分析清单：{item.label}")
            if int(record.get("policy_version") or 0) != DISCLOSURE_POLICY_VERSION:
                raise RuntimeError(f"缺项确认策略版本已失效：{item.label}")
        _append_disclosure_section(report_path, exact_disclosures)

    annotated_missing: list[Any] = []
    for index, raw in enumerate(missing_items):
        if not isinstance(raw, Mapping):
            annotated_missing.append(raw)
            continue
        enriched = dict(raw)
        candidate = candidate_by_index.get(index)
        if candidate is not None:
            enriched.update({
                "disclosure_stable_id": candidate.stable_id,
                "reason_code": candidate.reason_code,
                "reason_zh": candidate.reason_zh,
                "action_zh": candidate.action_zh,
                "confirmed_at": approval_records[candidate.stable_id].get("confirmed_at", ""),
            })
        annotated_missing.append(enriched)
    final_payload = dict(payload)
    final_payload.update({
        "status": "passed_with_disclosures" if exact_disclosures else "ok",
        "analysis_manifest": {
            "path": str(analysis_manifest_path or ""),
            "sha256": str(analysis_manifest_sha256).upper(),
        },
        "disclosure_policy_version": DISCLOSURE_POLICY_VERSION,
        "disclosure_count": len(exact_disclosures),
        "disclosures": [
            {
                **item.to_dict(),
                "confirmed_at": approval_records[item.stable_id].get("confirmed_at", ""),
            }
            for item in exact_disclosures
        ],
        "missing_items": annotated_missing,
        "missing_count": len(annotated_missing),
        "warnings": [],
        "output_docx": str(report_path.resolve()),
        "output_docx_sha256": _sha256(report_path),
        "finalized_at": datetime.now().astimezone().isoformat(timespec="seconds"),
    })
    _write_manifest(manifest_path, final_payload)
    return tuple(exact_disclosures), final_payload
