from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import subprocess
import tempfile
import zipfile
from xml.etree import ElementTree as ET
from dataclasses import asdict, is_dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm
from docx.table import Table
from openpyxl import load_workbook

from build_monthly_report import (
    SECTION_TITLES,
    apply_manifest_to_doc,
    build_manifest,
    ensure_dir,
    load_json,
    replace_paragraph_text,
    summarize_missing_images,
)
from build_quarterly_wim_sample import (
    T_NEXT,
    T_WIM,
    add_month_block,
    add_quarter_overview,
    add_text_paragraph_before,
    capture_paragraph_template,
    capture_wim_table_templates,
    clear_section_between,
    find_last_paragraph,
    insert_table_before,
    make_quarter_summary,
    parse_month_summary,
    resolve_wim_thresholds,
    set_summary_table,
    set_auto_caption_paragraph,
    set_header_bold,
    set_table_column_widths,
    style_table,
)
from template_precheck import raise_for_template
from analysis_manifest import (
    active_pinned_analysis_manifest,
    active_pinned_derived_artifact_manifest,
    load_latest_analysis_manifest,
    missing_module_summary_items,
    read_verified_derived_artifact_bytes,
)
from missing_summary import write_missing_summary
from report_build_manifest import write_report_build_manifest
from report_context import ReportBuildContext
from docx_header_fields import (
    HONGTANG_FRONT_MATTER_PAGES,
    audit_header_pagination_fields,
    ensure_header_pagination_fields,
)
from report_qc import check_report, write_report_qc_report


LOWFREQ_MODULES = {
    "strain": "结构应变监测",
    "tilt": "主塔倾斜监测",
    "bearing_displacement": "支座变位监测",
}

HIGHFREQ_MODULES = {
    "cable_accel": "吊索索力监测",
    "acceleration": "主梁、主塔振动监测",
    "wind": "风向风速监测",
    "eq": "地震动监测",
}

STATIC_CAPTION_RE = re.compile(r"^(?P<prefix>图|表|续表)\s*\d+(?:[-－]\d+){1,2}\s*.+")


def _paragraph_has_auto_caption_field(para) -> bool:
    for node in para._p.iter():
        if node.tag.endswith("}instrText") and node.text and "SEQ " in node.text:
            return True
    return False


def _static_caption_kind(text: str) -> str | None:
    stripped = re.sub(r"\s+", " ", str(text or "").strip())
    match = STATIC_CAPTION_RE.match(stripped)
    if not match:
        return None
    prefix = match.group("prefix")
    if prefix == "图":
        return "figure"
    if prefix == "续表":
        return "table_continued"
    return "table"


def convert_static_captions_to_auto_number(doc: Document) -> int:
    converted = 0
    for para in doc.paragraphs:
        if _paragraph_has_auto_caption_field(para):
            continue
        kind = _static_caption_kind(para.text)
        if kind is None:
            continue
        set_auto_caption_paragraph(para, para.text, kind)
        converted += 1
    return converted


def default_period_template(repo_root: Path) -> Path:
    reports_dir = repo_root / "reports"
    candidates = [
        reports_dir / "洪塘大桥健康监测周期报模板-自动报告.docx",
        reports_dir / "洪塘大桥健康监测2026年第一季季报-改4.docx",
        reports_dir / "洪塘大桥健康监测周期报模板0318.docx",
        reports_dir / "洪塘大桥健康监测周期报模板.docx",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def parse_args() -> argparse.Namespace:
    repo_root = Path(__file__).resolve().parents[1]
    default_template = default_period_template(repo_root)
    parser = argparse.ArgumentParser(description="Build full period monitoring report, including WIM.")
    parser.add_argument("--template", type=Path, default=default_template)
    parser.add_argument("--config", type=Path, default=repo_root / "config" / "hongtang_config.json")
    parser.add_argument("--result-root", type=Path, required=True)
    parser.add_argument("--analysis-root", type=Path, default=repo_root)
    parser.add_argument("--image-root", type=Path, default=None)
    parser.add_argument("--wim-root", type=Path, default=None, help="Processed monthly WIM result root, e.g. <result-root>/WIM/results/hongtang")
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--period-label", default="2026年1-3月")
    parser.add_argument("--monitoring-range", default="2026年01月01日~2026年03月31日")
    parser.add_argument("--report-number", default=None, help="Report number, e.g. BG02FQJC2600002-J2. Defaults to the quarter suffix.")
    parser.add_argument("--report-date", default=datetime.now().strftime("%Y年%m月%d日"))
    parser.add_argument("--start-date", default="2026-01-01")
    parser.add_argument("--end-date", default="2026-03-31")
    parser.add_argument(
        "--debug-section",
        default=None,
        help="Print a generated manifest section for debugging, e.g. cable_force, vibration, wim, health_status.",
    )
    parser.add_argument("--skip-template-precheck", action="store_true", help="Skip DOCX template anchor precheck.")
    return parser.parse_args()


def parse_date_str(text: str) -> date:
    return datetime.strptime(text, "%Y-%m-%d").date()


def extract_dates_from_range(text: str) -> tuple[date, date] | None:
    import re

    pattern = re.compile(r"(\d{4})[年.-](\d{1,2})[月.-](\d{1,2})日?.*?(\d{4})[年.-](\d{1,2})[月.-](\d{1,2})日?")
    match = pattern.search(text)
    if not match:
        return None
    y1, m1, d1, y2, m2, d2 = map(int, match.groups())
    return date(y1, m1, d1), date(y2, m2, d2)


def months_between(start_date: date, end_date: date) -> list[str]:
    if end_date < start_date:
        raise ValueError("end-date must not be earlier than start-date")
    months: list[str] = []
    year = start_date.year
    month = start_date.month
    while (year, month) <= (end_date.year, end_date.month):
        months.append(f"{year:04d}{month:02d}")
        if month == 12:
            year += 1
            month = 1
        else:
            month += 1
    return months


def resolve_wim_root(result_root: Path, analysis_root: Path, explicit_root: Path | None) -> Path:
    candidates: list[Path] = []
    if explicit_root is not None:
        candidates.append(explicit_root)
    candidates.extend(
        [
            result_root / "WIM" / "results" / "hongtang",
            result_root / "WIM" / "results",
            result_root / "WIM_results",
            result_root / "wim_results",
            analysis_root / "outputs" / "wim_quarter_sql" / "hongtang",
        ]
    )
    for candidate in candidates:
        if candidate.exists() and candidate.is_dir():
            return candidate
    checked = ", ".join(str(p) for p in candidates)
    raise FileNotFoundError(f"Processed WIM result root not found. Checked: {checked}")


def build_wim_period_section(wim_root: Path, months: list[str], cfg: dict | None = None) -> dict:
    thresholds = resolve_wim_thresholds(cfg)
    summaries = []
    warnings: list[str] = []
    for yyyymm in months:
        month_dir = wim_root / yyyymm
        if not month_dir.exists():
            warnings.append(f"Missing WIM month directory: {month_dir}")
            continue
        summaries.append(parse_month_summary(wim_root, yyyymm, thresholds))
    return {
        "enabled": bool(summaries),
        "wim_root": str(wim_root),
        "months": months,
        "warnings": warnings,
        "summary": make_quarter_summary(summaries) if summaries else "",
        "month_summaries": summaries,
    }


def _next_nonempty_paragraph_after(doc: Document, heading_text: str) -> object:
    paragraphs = doc.paragraphs
    for idx, para in enumerate(paragraphs):
        if para.text.strip() == heading_text:
            for nxt in paragraphs[idx + 1 :]:
                if nxt.text.strip():
                    return nxt
            break
    raise ValueError(f'Paragraph after "{heading_text}" not found')


def insert_table_after(paragraph, rows: int, cols: int):
    body = paragraph._parent
    table = body.add_table(rows=rows, cols=cols, width=Mm(160))
    paragraph._p.addnext(table._tbl)
    return table


def _normalize_missing_value(value: object) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        stripped = value.strip()
        return stripped == "" or stripped == "--"
    return False


def _iter_days(start_date: date, end_date: date) -> Iterable[date]:
    current = start_date
    while current <= end_date:
        yield current
        current += timedelta(days=1)


def _group_datetime_ranges(timestamps: list[datetime], missing_flags: list[bool]) -> list[tuple[datetime, datetime]]:
    ranges: list[tuple[datetime, datetime]] = []
    start_ts: datetime | None = None
    prev_ts: datetime | None = None
    for ts, missing in zip(timestamps, missing_flags):
        if missing:
            if start_ts is None:
                start_ts = ts
            prev_ts = ts
            continue
        if start_ts is not None and prev_ts is not None:
            ranges.append((start_ts, prev_ts))
            start_ts = None
            prev_ts = None
    if start_ts is not None and prev_ts is not None:
        ranges.append((start_ts, prev_ts))
    return ranges


def _group_date_ranges(days: list[date]) -> list[tuple[date, date]]:
    if not days:
        return []
    sorted_days = sorted(set(days))
    ranges: list[tuple[date, date]] = []
    start_day = sorted_days[0]
    prev_day = sorted_days[0]
    for current in sorted_days[1:]:
        if current == prev_day + timedelta(days=1):
            prev_day = current
            continue
        ranges.append((start_day, prev_day))
        start_day = current
        prev_day = current
    ranges.append((start_day, prev_day))
    return ranges


def _format_dt_range(start_dt: datetime, end_dt: datetime) -> str:
    if start_dt == end_dt:
        return start_dt.strftime("%Y-%m-%d %H:%M:%S")
    return f"{start_dt:%Y-%m-%d %H:%M:%S}~{end_dt:%Y-%m-%d %H:%M:%S}"


def _format_day_range(start_day: date, end_day: date) -> str:
    if start_day == end_day:
        return start_day.strftime("%Y-%m-%d")
    return f"{start_day:%Y-%m-%d}~{end_day:%Y-%m-%d}"


def _parse_range_bounds(range_text: str) -> tuple[date, date]:
    parts = range_text.split("~", 1)
    start_part = parts[0][:10]
    end_part = (parts[1][:10] if len(parts) == 2 else parts[0][:10])
    return (
        datetime.strptime(start_part, "%Y-%m-%d").date(),
        datetime.strptime(end_part, "%Y-%m-%d").date(),
    )


def _format_range_short(range_text: str) -> str:
    start_day, end_day = _parse_range_bounds(range_text)
    if start_day == end_day:
        return start_day.strftime("%Y-%m-%d")
    if start_day.year == end_day.year and start_day.month == end_day.month:
        return f"{start_day:%Y-%m-%d}~{end_day:%m-%d}"
    return f"{start_day:%Y-%m-%d}~{end_day:%Y-%m-%d}"


def _point_sort_key(point: str) -> tuple[int, str]:
    import re

    tokens = re.split(r"(\d+)", point)
    key: list[object] = []
    for token in tokens:
        if not token:
            continue
        key.append(int(token) if token.isdigit() else token)
    return (0, str(key))


def _join_points(points: Iterable[str], max_names: int = 4) -> str:
    ordered = sorted(dict.fromkeys(points), key=_point_sort_key)
    if len(ordered) <= max_names:
        return "、".join(ordered)
    return f"{'、'.join(ordered[:max_names])}等{len(ordered)}个测点"


def _join_module_names(names: Iterable[str]) -> str:
    ordered = []
    for module in list(LOWFREQ_MODULES.values()) + list(HIGHFREQ_MODULES.values()):
        if module in names and module not in ordered:
            ordered.append(module)
    return "、".join(ordered)


def _summarize_ranges(range_texts: Iterable[str], max_ranges: int = 3) -> str:
    ordered = sorted(dict.fromkeys(range_texts), key=_parse_range_bounds)
    formatted = [_format_range_short(item) for item in ordered]
    if not formatted:
        return ""
    if len(formatted) <= max_ranges:
        return "、".join(formatted)
    return f"{'、'.join(formatted[:max_ranges])}等{len(formatted)}段时段"


def _full_period_text(start_date: date, end_date: date) -> str:
    return _format_dt_range(datetime.combine(start_date, datetime.min.time()), datetime.combine(end_date, datetime.max.time()))


def period_report_title(start_date: date, end_date: date) -> str:
    if start_date.year == end_date.year:
        quarter_by_month = {
            (1, 3): "第一季度报告",
            (4, 6): "第二季度报告",
            (7, 9): "第三季度报告",
            (10, 12): "第四季度报告",
        }
        title = quarter_by_month.get((start_date.month, end_date.month))
        if title and start_date.day == 1:
            return title
    return "周期报告"


def period_report_number(start_date: date, end_date: date, prefix: str = "BG02FQJC2600002") -> str:
    quarter = ((start_date.month - 1) // 3) + 1
    return f"{prefix}-J{quarter}"


def safe_filename_text(text: str) -> str:
    text = re.sub(r'[<>:"/\\|?*\r\n\t]+', "_", str(text)).strip(" ._")
    return text or "周期"


def period_report_filename(period_label: str, timestamp: str) -> str:
    label = safe_filename_text(period_label)
    return f"洪塘大桥健康监测{label}周期报_{timestamp}.docx"


def apply_period_cover_title(doc: Document, start_date: date, end_date: date) -> None:
    title = period_report_title(start_date, end_date)
    pattern = re.compile(r"^第[一二三四]季度报告$")
    for para in doc.paragraphs:
        if pattern.match(para.text.strip()):
            replace_paragraph_text(para, title)


def apply_period_toc_months(doc: Document, start_date: date, end_date: date) -> None:
    labels = []
    for item in months_between(start_date, end_date):
        year = int(item[:4])
        month = int(item[4:])
        labels.append(f"{year}年{month}月交通状况监测")
    if not labels:
        return

    toc_pattern = re.compile(r"^(4\.1\.\d+\s+)\d{4}年\d{1,2}月交通状况监测(.*)$")
    label_idx = 0
    for para in doc.paragraphs:
        match = toc_pattern.match(para.text.strip())
        if not match:
            continue
        if label_idx >= len(labels):
            break
        replace_paragraph_text(para, f"{match.group(1)}{labels[label_idx]}{match.group(2)}")
        label_idx += 1


def apply_period_maintenance_text(doc: Document) -> None:
    stale_fragments = (
        "2026年2月4日对断电恢复后的称重系统进行调试",
        "2026年3月9日，对监测传感器、设备进行检查与维护",
    )
    replacement = (
        "本周期硬件维护情况以现场运维记录为准；本报告自动生成结果主要反映监测数据完整性和统计分析结果。"
        "对健康监测系统运行状况中列出的缺失或异常测点，建议结合现场巡检和设备台账复核。"
    )
    for para in doc.paragraphs:
        text = para.text.strip()
        if any(fragment in text for fragment in stale_fragments):
            replace_paragraph_text(para, replacement)


HONGTANG_PERIOD_MAINTENANCE_LOG = [
    ("2026-04-06", "软件系统日常检查"),
    ("2026-04-09", "基康采集设备（南侧）离线，现场维护，重启后设备恢复"),
    ("2026-04-13", "软件系统日常检查"),
    ("2026-04-20", "软件系统日常检查"),
    ("2026-04-27", "软件系统日常检查及定期维护"),
    ("2026-05-04", "软件系统日常检查"),
    ("2026-05-07", "部分测点离线，现场维护，线头老化已维保"),
    ("2026-05-11", "软件系统日常检查"),
    ("2026-05-18", "软件系统日常检查"),
    ("2026-05-25", "软件系统日常检查及定期维护"),
    ("2026-06-01", "软件系统日常检查"),
    ("2026-06-08", "软件系统日常检查"),
    ("2026-06-15", "软件系统日常检查，现场对基康采集仪设备进行维保（已修复）"),
    ("2026-06-22", "软件系统日常检查"),
    ("2026-06-29", "软件系统日常检查及定期维护"),
]


def period_maintenance_log_rows(start_date: date, end_date: date) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    for date_text, maintenance_type in HONGTANG_PERIOD_MAINTENANCE_LOG:
        current = parse_date_str(date_text)
        if start_date <= current <= end_date:
            rows.append((date_text, maintenance_type))
    return rows


def _next_table_after_paragraph(paragraph) -> Table | None:
    current = paragraph._p.getnext()
    while current is not None:
        if current.tag == qn("w:tbl"):
            return Table(current, paragraph._parent)
        current = current.getnext()
    return None


def _set_cell_text_preserve_style(cell, text: object) -> None:
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.text = str(text)
        return
    first = paragraphs[0]
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    if first.runs:
        first.runs[0].text = str(text)
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(str(text))


def _remove_table_row(row) -> None:
    row._tr.getparent().remove(row._tr)


def apply_period_maintenance_log(doc: Document, start_date: date, end_date: date) -> None:
    rows = period_maintenance_log_rows(start_date, end_date)
    if not rows:
        return

    target_table: Table | None = None
    for para in doc.paragraphs:
        if "健康监测系统维护日志表" in para.text:
            target_table = _next_table_after_paragraph(para)
            break
    if target_table is None:
        raise ValueError("Maintenance log table after 表 1-2 not found")
    if len(target_table.columns) < 3:
        raise ValueError("Maintenance log table must have at least three columns")

    required_rows = len(rows) + 1
    while len(target_table.rows) < required_rows:
        target_table.add_row()
    while len(target_table.rows) > required_rows:
        _remove_table_row(target_table.rows[-1])

    headers = ("序号", "维护日期", "维护类型")
    for idx, value in enumerate(headers):
        _set_cell_text_preserve_style(target_table.cell(0, idx), value)
    for ridx, (date_text, maintenance_type) in enumerate(rows, start=1):
        values = (str(ridx), date_text, maintenance_type)
        for cidx, value in enumerate(values):
            _set_cell_text_preserve_style(target_table.cell(ridx, cidx), value)


def _summarize_lowfreq_module(module: str, module_events: list[dict], start_date: date, end_date: date) -> str:
    full_period = _full_period_text(start_date, end_date)
    reason_priority = {"原始记录缺失": 0, "lowfreq/data.xlsx 缺失": 1, "lowfreq 数据期内无原始记录": 2}
    grouped_by_range: dict[tuple[str, str], set[str]] = {}
    for event in module_events:
        grouped_by_range.setdefault((event["range"], event["reason"]), set()).update(event["points"])

    grouped: dict[tuple[tuple[str, ...], str], list[str]] = {}
    for (range_text, reason), points in grouped_by_range.items():
        point_key = tuple(sorted(points, key=_point_sort_key))
        grouped.setdefault((point_key, reason), []).append(range_text)

    persistent_parts: list[str] = []
    intermittent_parts: list[str] = []
    for (points, reason), ranges in sorted(
        grouped.items(),
        key=lambda item: (reason_priority.get(item[0][1], 99), -len(item[0][0]), item[0][0]),
    ):
        unique_ranges = sorted(dict.fromkeys(ranges), key=_parse_range_bounds)
        point_text = _join_points(points)
        if len(unique_ranges) == 1 and unique_ranges[0] == full_period:
            if points == ("全测点",):
                persistent_parts.append(reason)
            else:
                persistent_parts.append(f"{point_text}全周期{reason}")
            continue
        range_text = _summarize_ranges(unique_ranges)
        if points == ("全测点",):
            intermittent_parts.append(f"{range_text}{reason}")
        else:
            intermittent_parts.append(f"{point_text}在{range_text}{reason}")

    parts: list[str] = []
    if persistent_parts:
        parts.append("；".join(persistent_parts))
    if intermittent_parts:
        parts.append("；".join(intermittent_parts))
    if not parts:
        return ""
    return f"{module}中，" + "；".join(parts)


def _summarize_highfreq_events(module_events: list[dict]) -> str:
    if not module_events:
        return ""

    range_reason_modules: dict[tuple[str, str], dict[str, set[str]]] = {}
    for event in module_events:
        key = (event["range"], event["reason"])
        bucket = range_reason_modules.setdefault(key, {})
        bucket.setdefault(event["module"], set()).update(event["points"])

    common_parts: list[str] = []
    common_keys: set[tuple[str, str]] = set()
    for (range_text, reason), modules in sorted(range_reason_modules.items(), key=lambda item: _parse_range_bounds(item[0][0])):
        if len(modules) < 2:
            continue
        common_keys.add((range_text, reason))
        common_parts.append(f"{_format_range_short(range_text)}出现{reason}，影响{_join_module_names(modules.keys())}")

    residual_by_module: dict[str, list[tuple[str, str, set[str]]]] = {}
    for (range_text, reason), modules in range_reason_modules.items():
        if (range_text, reason) in common_keys:
            continue
        for module, points in modules.items():
            residual_by_module.setdefault(module, []).append((range_text, reason, points))

    residual_parts: list[str] = []
    for module in HIGHFREQ_MODULES.values():
        entries = residual_by_module.get(module, [])
        if not entries:
            continue
        grouped: dict[tuple[tuple[str, ...], str], list[str]] = {}
        for range_text, reason, points in entries:
            grouped.setdefault((tuple(sorted(points, key=_point_sort_key)), reason), []).append(range_text)
        module_parts = []
        for (points, reason), ranges in sorted(grouped.items(), key=lambda item: _parse_range_bounds(item[1][0])):
            point_text = _join_points(points)
            range_text = _summarize_ranges(ranges)
            module_parts.append(f"{point_text}在{range_text}{reason}")
        residual_parts.append(f"{module}中，" + "；".join(module_parts))

    parts: list[str] = []
    if common_parts:
        parts.append("高频监测系统在" + "；".join(common_parts))
    if residual_parts:
        parts.append("此外，" + "；".join(residual_parts))
    return "。".join(parts)


def _build_lowfreq_health_rows(module: str, module_events: list[dict], start_date: date, end_date: date) -> list[dict[str, str]]:
    full_period = _full_period_text(start_date, end_date)
    grouped_by_range: dict[tuple[str, str], set[str]] = {}
    for event in module_events:
        grouped_by_range.setdefault((event["range"], event["reason"]), set()).update(event["points"])

    grouped: dict[tuple[tuple[str, ...], str], list[str]] = {}
    for (range_text, reason), points in grouped_by_range.items():
        grouped.setdefault((tuple(sorted(points, key=_point_sort_key)), reason), []).append(range_text)

    rows: list[dict[str, str]] = []
    for (points, reason), ranges in sorted(grouped.items(), key=lambda item: (_parse_range_bounds(item[1][0])[0], item[0][0])):
        unique_ranges = sorted(dict.fromkeys(ranges), key=_parse_range_bounds)
        rows.append(
            {
                "module": module,
                "points": "全测点" if points == ("全测点",) else _join_points(points, max_names=6),
                "range": "全周期" if len(unique_ranges) == 1 and unique_ranges[0] == full_period else _summarize_ranges(unique_ranges, max_ranges=4),
                "reason": reason,
            }
        )
    return rows


def _build_highfreq_health_rows(module_events: list[dict]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    range_reason_modules: dict[tuple[str, str], dict[str, set[str]]] = {}
    for event in module_events:
        key = (event["range"], event["reason"])
        range_reason_modules.setdefault(key, {}).setdefault(event["module"], set()).update(event["points"])

    common_keys: set[tuple[str, str]] = set()
    for (range_text, reason), modules in sorted(range_reason_modules.items(), key=lambda item: _parse_range_bounds(item[0][0])):
        if len(modules) < 2:
            continue
        common_keys.add((range_text, reason))
        rows.append(
            {
                "module": _join_module_names(modules.keys()),
                "points": "相关测点",
                "range": _format_range_short(range_text),
                "reason": reason,
            }
        )

    residual_by_module: dict[str, list[tuple[str, str, set[str]]]] = {}
    for (range_text, reason), modules in range_reason_modules.items():
        if (range_text, reason) in common_keys:
            continue
        for module, points in modules.items():
            residual_by_module.setdefault(module, []).append((range_text, reason, points))

    for module in HIGHFREQ_MODULES.values():
        entries = residual_by_module.get(module, [])
        if not entries:
            continue
        grouped: dict[tuple[tuple[str, ...], str], list[str]] = {}
        for range_text, reason, points in entries:
            grouped.setdefault((tuple(sorted(points, key=_point_sort_key)), reason), []).append(range_text)
        for (points, reason), ranges in sorted(grouped.items(), key=lambda item: _parse_range_bounds(item[1][0])):
            rows.append(
                {
                    "module": module,
                    "points": _join_points(points, max_names=6),
                    "range": _summarize_ranges(ranges, max_ranges=4),
                    "reason": reason,
                }
            )
    return rows


def build_health_status_rows(
    cfg: dict,
    result_root: Path,
    start_date: date,
    end_date: date,
    lowfreq_events: list[dict] | None = None,
    highfreq_events: list[dict] | None = None,
) -> list[dict[str, str]]:
    if lowfreq_events is None:
        lowfreq_events = collect_lowfreq_missing_events(cfg, result_root, start_date, end_date)
    if highfreq_events is None:
        highfreq_events = collect_highfreq_missing_events(cfg, result_root, start_date, end_date)

    rows: list[dict[str, str]] = []
    for module in LOWFREQ_MODULES.values():
        module_events = [event for event in lowfreq_events if event["module"] == module]
        rows.extend(_build_lowfreq_health_rows(module, module_events, start_date, end_date))
    rows.extend(_build_highfreq_health_rows(highfreq_events))
    return rows


def build_report_missing_rows(manifest: dict, wim_section: dict) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for key, title in SECTION_TITLES.items():
        section = manifest.get("sections", {}).get(key, {})
        if not section.get("enabled", True):
            continue
        if section.get("available", True):
            continue
        rows.append(
            {
                "module": title,
                "points": "-",
                "range": "本周期",
                "reason": "本周期未获取到有效数据",
            }
        )

    if not wim_section.get("enabled"):
        rows.append(
            {
                "module": "交通状况监测",
                "points": "-",
                "range": "本周期",
                "reason": "本周期未获取到有效数据",
            }
        )
        return rows

    available_months = {item.yyyymm for item in wim_section.get("month_summaries", [])}
    missing_months = [month for month in wim_section.get("months", []) if month not in available_months]
    if missing_months:
        rows.append(
            {
                "module": "交通状况监测",
                "points": "-",
                "range": "、".join(missing_months),
                "reason": "部分月份WIM结果缺失",
            }
        )
    return rows


def merge_health_status_summary(summary_text: str, missing_rows: list[dict[str, str]]) -> str:
    if not missing_rows:
        return summary_text
    extra = "此外，部分监测内容缺失，详见下表。"
    if not summary_text or "未发现原始数据缺失" in summary_text:
        return extra
    return summary_text.rstrip("。") + "。" + extra


def _pattern_for_point(cfg: dict, module: str, point_id: str, file_id: str | None = None) -> str:
    return _patterns_for_point(cfg, module, point_id, file_id=file_id)[0]


def _patterns_for_point(cfg: dict, module: str, point_id: str, file_id: str | None = None) -> list[str]:
    patterns = cfg.get("file_patterns", {}).get(module, {})
    per_point = patterns.get("per_point", {})
    pattern = per_point.get(point_id)
    if pattern is None:
        default_patterns = patterns.get("default") or []
        if not default_patterns:
            raise KeyError(f"No file pattern configured for module={module}, point={point_id}")
        values = default_patterns if isinstance(default_patterns, list) else [default_patterns]
    elif isinstance(pattern, list):
        if not pattern:
            raise KeyError(f"Empty file pattern configured for module={module}, point={point_id}")
        values = pattern
    else:
        values = [pattern]
    return [str(item).format(point=point_id, file_id=file_id or "") for item in values]


def _csv_has_records(path: Path) -> bool:
    if not path.exists() or path.stat().st_size <= 0:
        return False
    with path.open("r", encoding="utf-8", errors="ignore") as fh:
        for line in fh:
            if line.strip():
                return True
    return False


def _cache_patterns(csv_pattern: str) -> list[str]:
    """Return MAT-cache glob candidates for one configured CSV pattern."""
    name = Path(str(csv_pattern).replace("\\", "/")).name
    stem = Path(name).stem
    patterns = [stem + ".mat"]
    # Cache files are keyed by the logical point name.  A source pattern such
    # as A1_174.csv therefore becomes A1_*.mat; the same rule also covers
    # X_144.csv, 风速_162.csv and their timestamped cache variants.
    logical_stem = re.sub(r"_\d+$", "_*", stem)
    if logical_stem != stem:
        patterns.append(logical_stem + ".mat")
    return list(dict.fromkeys(patterns))


def _raw_or_cache_has_records(raw_dir: Path, csv_patterns: Iterable[str]) -> bool:
    patterns = [str(item) for item in csv_patterns]
    for pattern in patterns:
        if any(_csv_has_records(path) for path in raw_dir.glob(pattern)):
            return True

    cache_dir = raw_dir / "cache"
    if not cache_dir.is_dir():
        return False
    for pattern in patterns:
        for cache_pattern in _cache_patterns(pattern):
            if any(path.is_file() and path.stat().st_size > 0 for path in cache_dir.glob(cache_pattern)):
                return True
    return False


def _normalize_plot_series(value) -> list[dict]:
    if isinstance(value, list):
        return [item for item in value if isinstance(item, dict)]
    if isinstance(value, dict):
        if "source" in value or "sampling_mode" in value:
            return [value]
        return [item for item in value.values() if isinstance(item, dict)]
    return []


def _provenance_point_id(series: dict, sidecar: Path) -> str:
    point_id = str(series.get("point_id") or "").split(":", 1)[0].strip()
    if point_id:
        return point_id
    stem = sidecar.name.removesuffix(".plot.json")
    stem = re.sub(r"_20\d{2}-\d{2}-\d{2}_20\d{2}-\d{2}-\d{2}$", "", stem)
    stem = re.sub(r"_20\d{6}_20\d{6}$", "", stem)
    return stem.replace("EQ_", "EQ-")


def collect_highfreq_provenance_events(
    result_root: Path, start_date: date, end_date: date
) -> list[dict]:
    """Build natural-day completeness events from the pinned analysis artifacts.

    Date-folder presence alone cannot prove a natural day is complete for the
    rolling D+D+1 exports used by Hongtang.  The plot sidecars record that
    reconstruction explicitly, including boundary days affected by a missing
    adjacent export.
    """
    _manifest_path, manifest = load_latest_analysis_manifest(result_root)
    if not isinstance(manifest, dict):
        return []

    module_labels = {
        "acceleration": HIGHFREQ_MODULES["acceleration"],
        "cable_accel": HIGHFREQ_MODULES["cable_accel"],
        "wind": HIGHFREQ_MODULES["wind"],
        "earthquake": HIGHFREQ_MODULES["eq"],
    }
    records = manifest.get("module_results") or manifest.get("module_logs") or []
    dedup: set[tuple[str, str, date, str]] = set()
    events: list[dict] = []
    root_resolved = result_root.resolve()
    coverage_audit = load_data_coverage_audit(result_root)
    boundary = coverage_audit.get("boundary_disclosure") if isinstance(coverage_audit, dict) else None
    explicit_boundary_day = ""
    if isinstance(boundary, dict) and str(boundary.get("statement") or "").strip():
        explicit_boundary_day = str(boundary.get("affected_day") or "")[:10]

    for record in records if isinstance(records, list) else []:
        if not isinstance(record, dict):
            continue
        module_key = str(record.get("key") or record.get("module") or "")
        module_label = module_labels.get(module_key)
        if module_label is None:
            continue
        for artifact in record.get("artifacts") or []:
            raw_path = artifact.get("path") if isinstance(artifact, dict) else artifact
            if not raw_path or not str(raw_path).lower().endswith(".plot.json"):
                continue
            sidecar = Path(str(raw_path))
            try:
                sidecar.resolve().relative_to(root_resolved)
            except (OSError, ValueError):
                continue
            if not sidecar.is_file():
                continue
            try:
                payload = json.loads(sidecar.read_text(encoding="utf-8-sig"))
            except (OSError, ValueError):
                continue
            for series in _normalize_plot_series(payload.get("series")):
                source = series.get("source") if isinstance(series.get("source"), dict) else {}
                missing_sources = [str(item) for item in source.get("missing_required_sources") or []]
                reason = "自然日数据不完整（缺少相邻滚动导出）" if missing_sources else "自然日数据不完整"
                point_id = _provenance_point_id(series, sidecar)
                for value in source.get("incomplete_days") or []:
                    try:
                        day = date.fromisoformat(str(value)[:10])
                    except ValueError:
                        continue
                    if not (start_date <= day <= end_date):
                        continue
                    if day.isoformat() == explicit_boundary_day:
                        # The independently audited boundary statement below is
                        # more precise than the generic sidecar label.
                        continue
                    marker = (module_label, point_id, day, reason)
                    if marker in dedup:
                        continue
                    dedup.add(marker)
                    events.append(
                        {
                            "module": module_label,
                            "points": [point_id],
                            "range": _format_day_range(day, day),
                            "reason": reason,
                        }
                    )
    return events


def load_data_coverage_audit(result_root: Path) -> dict:
    conventional_path = (result_root / "run_logs" / "data_coverage_audit.json").resolve()
    if active_pinned_analysis_manifest() is not None:
        if active_pinned_derived_artifact_manifest() is None:
            if conventional_path.is_file():
                raise ValueError(
                    "Strict source provenance forbids an unlisted data coverage audit: "
                    f"{conventional_path}"
                )
            return {}
        verified = read_verified_derived_artifact_bytes(
            kind="coverage_audit", role="data_coverage_audit", module="run_audit"
        )
        if verified is None:
            if conventional_path.is_file():
                raise ValueError(
                    "Strict source provenance requires data_coverage_audit.json to be "
                    "listed in the active derived-artifact manifest."
                )
            return {}
        path, raw, record = verified
    else:
        path = conventional_path
        if not path.is_file():
            return {}
        raw = path.read_bytes()
        record = None
    try:
        payload = json.loads(raw.decode("utf-8-sig"))
    except (UnicodeDecodeError, ValueError) as exc:
        raise ValueError(f"Invalid data coverage audit: {path}: {exc}") from exc
    if not isinstance(payload, dict):
        raise ValueError(f"Data coverage audit root must be an object: {path}")
    return {
        **payload,
        "path": str(path),
        "bytes": len(raw),
        "sha256": hashlib.sha256(raw).hexdigest().upper(),
        "source": "derived_artifact_manifest" if record is not None else "filesystem",
    }


def collect_declared_data_coverage_events(
    result_root: Path, start_date: date, end_date: date
) -> list[dict]:
    """Read independently audited timestamp gaps for report disclosure."""
    audit = load_data_coverage_audit(result_root)
    raw_events = audit.get("events") if isinstance(audit, dict) else []
    if not isinstance(raw_events, list):
        return []

    period_start = datetime.combine(start_date, datetime.min.time())
    period_end = datetime.combine(end_date, datetime.max.time())
    result: list[dict] = []
    for raw in raw_events:
        if not isinstance(raw, dict):
            continue
        try:
            event_start = datetime.fromisoformat(str(raw.get("start") or ""))
            event_end = datetime.fromisoformat(str(raw.get("end") or ""))
        except ValueError:
            continue
        if event_end < event_start or event_end < period_start or event_start > period_end:
            continue
        module_key = str(raw.get("module") or "").strip()
        module_label = HIGHFREQ_MODULES.get(module_key, module_key)
        points = raw.get("points")
        if isinstance(points, str):
            points = [points]
        if not module_label or not isinstance(points, list) or not points:
            continue
        reason = str(raw.get("reason") or "经时间戳审计确认的数据缺口").strip()
        result.append(
            {
                "module": module_label,
                "points": [str(item) for item in points if str(item).strip()],
                "range": _format_dt_range(max(event_start, period_start), min(event_end, period_end)),
                "reason": reason,
            }
        )

    boundary = audit.get("boundary_disclosure") if isinstance(audit, dict) else None
    if isinstance(boundary, dict):
        try:
            affected_day = date.fromisoformat(str(boundary.get("affected_day") or "")[:10])
        except ValueError:
            affected_day = None
        statement = str(boundary.get("statement") or "").strip()
        if affected_day is not None and start_date <= affected_day <= end_date and statement:
            for module_key in ("acceleration", "cable_accel", "wind", "eq"):
                result.append(
                    {
                        "module": HIGHFREQ_MODULES[module_key],
                        "points": ["相关测点"],
                        "range": _format_day_range(affected_day, affected_day),
                        "reason": statement,
                    }
                )
    return result


def collect_lowfreq_missing_events(cfg: dict, result_root: Path, start_date: date, end_date: date) -> list[dict]:
    workbook = result_root / "lowfreq" / "data.xlsx"
    points_by_module = {key: cfg.get("points", {}).get(key, []) for key in LOWFREQ_MODULES}
    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt = datetime.combine(end_date, datetime.max.time())
    events: list[dict] = []
    if not workbook.exists():
        for module, label in LOWFREQ_MODULES.items():
            if points_by_module.get(module):
                events.append(
                    {
                        "module": label,
                        "points": ["全测点"],
                        "range": _format_dt_range(start_dt, end_dt),
                        "reason": "lowfreq/data.xlsx 缺失",
                    }
                )
        return events

    wb = load_workbook(workbook, read_only=True, data_only=True)
    ws = wb.active
    rows = ws.iter_rows(values_only=True)
    headers = list(next(rows))
    header_index = {str(name).strip(): idx for idx, name in enumerate(headers) if name is not None}
    timestamps: list[datetime] = []
    missing_days_by_point: dict[str, set[date]] = {}
    for module_points in points_by_module.values():
        for point in module_points:
            missing_days_by_point[point] = set()

    for row in rows:
        ts = row[0]
        if ts is None:
            continue
        if isinstance(ts, datetime):
            dt = ts
        else:
            try:
                dt = datetime.fromisoformat(str(ts))
            except ValueError:
                continue
        if dt < start_dt or dt > end_dt:
            continue
        timestamps.append(dt)
        for point in missing_days_by_point:
            col_idx = header_index.get(point)
            is_missing = col_idx is None or col_idx >= len(row) or _normalize_missing_value(row[col_idx])
            if is_missing:
                missing_days_by_point[point].add(dt.date())
    wb.close()

    if not timestamps:
        for module, label in LOWFREQ_MODULES.items():
            if points_by_module.get(module):
                events.append(
                    {
                        "module": label,
                        "points": ["全测点"],
                        "range": _format_dt_range(start_dt, end_dt),
                        "reason": "lowfreq 数据期内无原始记录",
                    }
                )
        return events

    for module, label in LOWFREQ_MODULES.items():
        for point in points_by_module.get(module, []):
            ranges = _group_date_ranges(sorted(missing_days_by_point.get(point, set())))
            for start_day, end_day in ranges:
                events.append(
                    {
                        "module": label,
                        "points": [point],
                        "range": _format_day_range(start_day, end_day),
                        "reason": "原始记录缺失",
                    }
                )
    return events


def collect_highfreq_missing_events(cfg: dict, result_root: Path, start_date: date, end_date: date) -> list[dict]:
    events: list[dict] = []
    subfolders = cfg.get("subfolders", {})
    points_cfg = cfg.get("points", {})
    file_patterns = cfg.get("file_patterns", {})

    def add_events(module_label: str, point_label: str, days: list[date], reason: str) -> None:
        for start_day, end_day in _group_date_ranges(days):
            events.append(
                {
                    "module": module_label,
                    "points": [point_label],
                    "range": _format_day_range(start_day, end_day),
                    "reason": reason,
                }
            )

    # acceleration and cable acceleration
    for module in ("acceleration", "cable_accel"):
        module_label = HIGHFREQ_MODULES[module]
        day_missing: dict[tuple[str, str], list[date]] = {}
        folder_name = subfolders.get(module)
        if not folder_name:
            continue
        for current_day in _iter_days(start_date, end_date):
            raw_dir = result_root / current_day.strftime("%Y-%m-%d") / folder_name
            for point in points_cfg.get(module, []):
                key = (point, "原始文件缺失")
                if not raw_dir.exists():
                    day_missing.setdefault(key, []).append(current_day)
                    continue
                patterns = _patterns_for_point(cfg, module, point)
                if not _raw_or_cache_has_records(raw_dir, patterns):
                    day_missing.setdefault((point, "无原始记录"), []).append(current_day)
        for (point, reason), days in day_missing.items():
            add_events(module_label, point, days, reason)

    # wind speed / direction
    wind_folder = subfolders.get("wind_raw") or subfolders.get("wind")
    if wind_folder:
        day_missing: dict[tuple[str, str], list[date]] = {}
        for current_day in _iter_days(start_date, end_date):
            raw_dir = result_root / current_day.strftime("%Y-%m-%d") / wind_folder
            for point, point_cfg in cfg.get("per_point", {}).get("wind", {}).items():
                for suffix, field, module_key in (("风速", "speed_point_id", "wind_speed"), ("风向", "dir_point_id", "wind_direction")):
                    label = f"{point}-{suffix}"
                    file_id = point_cfg.get(field)
                    if not file_id:
                        continue
                    key = (label, "原始文件缺失")
                    if not raw_dir.exists():
                        day_missing.setdefault(key, []).append(current_day)
                        continue
                    patterns = _patterns_for_point(cfg, module_key, point, file_id=file_id)
                    if not _raw_or_cache_has_records(raw_dir, patterns):
                        day_missing.setdefault((label, "无原始记录"), []).append(current_day)
        for (point, reason), days in day_missing.items():
            add_events(HIGHFREQ_MODULES["wind"], point, days, reason)

    # earthquake
    eq_folder = subfolders.get("eq_raw") or subfolders.get("eq")
    if eq_folder:
        day_missing: dict[tuple[str, str], list[date]] = {}
        for current_day in _iter_days(start_date, end_date):
            raw_dir = result_root / current_day.strftime("%Y-%m-%d") / eq_folder
            for point, point_cfg in cfg.get("per_point", {}).get("eq", {}).items():
                file_id = point_cfg.get("file_id")
                if not file_id:
                    continue
                key = (point, "原始文件缺失")
                if not raw_dir.exists():
                    day_missing.setdefault(key, []).append(current_day)
                    continue
                module_key = point.lower().replace("-", "_")
                patterns = _patterns_for_point(cfg, module_key, point, file_id=file_id)
                if not _raw_or_cache_has_records(raw_dir, patterns):
                    day_missing.setdefault((point, "无原始记录"), []).append(current_day)
        for (point, reason), days in day_missing.items():
            add_events(HIGHFREQ_MODULES["eq"], point, days, reason)

    events.extend(collect_highfreq_provenance_events(result_root, start_date, end_date))
    events.extend(collect_declared_data_coverage_events(result_root, start_date, end_date))
    return events


def build_health_status_summary(
    cfg: dict,
    result_root: Path,
    start_date: date,
    end_date: date,
    lowfreq_events: list[dict] | None = None,
    highfreq_events: list[dict] | None = None,
) -> str:
    if lowfreq_events is None:
        lowfreq_events = collect_lowfreq_missing_events(cfg, result_root, start_date, end_date)
    if highfreq_events is None:
        highfreq_events = collect_highfreq_missing_events(cfg, result_root, start_date, end_date)
    if not lowfreq_events and not highfreq_events:
        return "监测周期内未发现原始数据缺失、无文件或无记录情况。"

    parts: list[str] = []
    if lowfreq_events:
        lowfreq_parts = []
        for module in LOWFREQ_MODULES.values():
            module_events = [event for event in lowfreq_events if event["module"] == module]
            text = _summarize_lowfreq_module(module, module_events, start_date, end_date)
            if text:
                lowfreq_parts.append(text)
        if lowfreq_parts:
            parts.append("低频监测系统存在持续性和阶段性原始记录缺失。" + "；".join(lowfreq_parts) + "。")

    highfreq_text = _summarize_highfreq_events(highfreq_events)
    if highfreq_text:
        parts.append(highfreq_text + "。")

    return "监测周期内原始数据缺失、无文件或无记录情况见下表。"


def apply_health_status_to_doc(doc: Document, summary_text: str, rows: list[dict[str, str]]) -> None:
    paragraph = _next_nonempty_paragraph_after(doc, "健康监测系统运行状况")
    replace_paragraph_text(paragraph, summary_text)
    if not rows:
        return
    table = insert_table_after(paragraph, rows=len(rows) + 1, cols=4)
    headers = ["监测项目", "异常测点/测点组", "时间段", "异常类型"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for ridx, row in enumerate(rows, start=1):
        values = [row["module"], row["points"], row["range"], row["reason"]]
        for cidx, value in enumerate(values):
            table.cell(ridx, cidx).text = value
    style_table(table, left=True)
    set_header_bold(table)
    set_table_column_widths(table, [28, 66, 42, 24])


def _patch_report_number_xml(xml_text: str, report_number: str) -> tuple[str, int]:
    report_no_re = re.compile(r"BG02FQJC2600002-J\d+")
    return report_no_re.subn(report_number, xml_text)


def _patch_report_number_in_docx(docx_path: Path, report_number: str | None) -> int:
    if not report_number:
        return 0
    if not docx_path.exists():
        return 0

    patched_count = 0
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    try:
                        text = data.decode("utf-8")
                    except UnicodeDecodeError:
                        text = ""
                    if text:
                        text, count = _patch_report_number_xml(text, report_number)
                        if count:
                            patched_count += count
                            data = text.encode("utf-8")
                zout.writestr(info, data)
        tmp_path.replace(docx_path)
    except zipfile.BadZipFile:
        return 0
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass
    return patched_count


def _docx_contains_broken_reference_text(docx_path: Path) -> bool:
    phrases = ("引用源未找到", "Error! Reference source not found")
    try:
        with zipfile.ZipFile(docx_path, "r") as archive:
            for name in archive.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                try:
                    root = ET.fromstring(archive.read(name))
                except ET.ParseError:
                    continue
                text = "".join(
                    node.text or ""
                    for node in root.iter()
                    if node.tag.endswith("}t")
                )
                if any(phrase in text for phrase in phrases):
                    return True
    except (OSError, zipfile.BadZipFile):
        return False
    return False


def _run_python_word_field_update(docx_path: Path) -> tuple[bool, str]:
    script = f"""
import pythoncom
import win32com.client
p = r'''{str(docx_path)}'''
pythoncom.CoInitialize()
word = None
doc = None
try:
    com_errors = []
    for prog_id in ("Word.Application", "KWPS.Application"):
        try:
            word = win32com.client.DispatchEx(prog_id)
            break
        except Exception as exc:
            com_errors.append(f"{{prog_id}}: {{exc}}")
    if word is None:
        raise RuntimeError("; ".join(com_errors))
    word.Visible = False
    word.DisplayAlerts = 0
    doc = word.Documents.Open(p)
    def update_all_fields():
        def update_shape_fields(container):
            try:
                shapes = container.Shapes
                count = shapes.Count
            except Exception:
                return
            for idx in range(1, count + 1):
                try:
                    shape = shapes.Item(idx)
                    if shape.TextFrame.HasText:
                        shape.TextFrame.TextRange.Fields.Update()
                except Exception:
                    pass

        try:
            doc.Repaginate()
        except Exception:
            pass
        for section in doc.Sections:
            for header in section.Headers:
                try:
                    if header.Exists:
                        header.Range.Fields.Update()
                except Exception:
                    pass
                update_shape_fields(header)
            for footer in section.Footers:
                try:
                    if footer.Exists:
                        footer.Range.Fields.Update()
                except Exception:
                    pass
                update_shape_fields(footer)
        for story_start in doc.StoryRanges:
            story = story_start
            while story is not None:
                try:
                    story.Fields.Update()
                except Exception:
                    pass
                try:
                    story = story.NextStoryRange
                except Exception:
                    story = None
        try:
            doc.TablesOfContents(1).Update()
        except Exception:
            pass
        try:
            doc.Fields.Update()
        except Exception:
            pass

    update_all_fields()
    update_all_fields()
    update_all_fields()
    doc.Save()
finally:
    if doc is not None:
        doc.Close(SaveChanges=True)
    if word is not None:
        word.Quit()
    pythoncom.CoUninitialize()
"""
    candidates = [Path(__file__).resolve().parent / ".venv" / "Scripts" / "python.exe"]
    if os.environ.get("PYTHON"):
        candidates.append(Path(os.environ["PYTHON"]))
    candidates.append(Path("python"))
    errors: list[str] = []
    original_bytes = docx_path.read_bytes()
    for py in candidates:
        if py != Path("python") and not py.exists():
            continue
        docx_path.write_bytes(original_bytes)
        try:
            result = subprocess.run([str(py), "-c", script], check=False, timeout=180, capture_output=True, text=True)
            if result.returncode == 0:
                if _docx_contains_broken_reference_text(docx_path):
                    errors.append(f"{py}: field update produced broken reference results")
                    docx_path.write_bytes(original_bytes)
                    continue
                return True, ""
            detail = (result.stderr or result.stdout or f"exit={result.returncode}").strip()
            errors.append(f"{py}: {detail[:300]}")
        except (OSError, subprocess.SubprocessError) as exc:
            errors.append(f"{py}: {exc}")
    docx_path.write_bytes(original_bytes)
    return False, "; ".join(errors[-3:])


def _run_powershell_word_field_update(docx_path: Path) -> tuple[bool, str]:
    ps_script = r"""
$ErrorActionPreference = 'Stop'
$p = $env:BMS_DOCX_PATH
$word = $null
$doc = $null
$saved = $false
try {
    $comErrors = @()
    foreach ($progId in @('Word.Application', 'KWPS.Application')) {
        try {
            $word = New-Object -ComObject $progId
            break
        } catch {
            $comErrors += "${progId}: $($_.Exception.Message)"
        }
    }
    if ($null -eq $word) { throw ($comErrors -join '; ') }
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $doc = $word.Documents.Open($p)
    function Update-ShapeFields($shapes) {
        try { $count = $shapes.Count } catch { return }
        for ($idx = 1; $idx -le $count; $idx++) {
            try {
                $shape = $shapes.Item($idx)
                if ($shape.TextFrame.HasText) {
                    $shape.TextFrame.TextRange.Fields.Update() | Out-Null
                }
            } catch {}
        }
    }
    function Update-WordFields($doc) {
        try { $doc.Repaginate() | Out-Null } catch {}
        foreach ($section in @($doc.Sections)) {
            foreach ($header in @($section.Headers)) {
                try {
                    if ($header.Exists) { $header.Range.Fields.Update() | Out-Null }
                } catch {}
                try { Update-ShapeFields $header.Shapes } catch {}
            }
            foreach ($footer in @($section.Footers)) {
                try {
                    if ($footer.Exists) { $footer.Range.Fields.Update() | Out-Null }
                } catch {}
                try { Update-ShapeFields $footer.Shapes } catch {}
            }
        }
        foreach ($storyStart in @($doc.StoryRanges)) {
            $story = $storyStart
            while ($null -ne $story) {
                try { $story.Fields.Update() | Out-Null } catch {}
                $story = $story.NextStoryRange
            }
        }
        try { $doc.TablesOfContents.Item(1).Update() | Out-Null } catch {}
        try { $doc.Fields.Update() | Out-Null } catch {}
    }
    Update-WordFields $doc
    Update-WordFields $doc
    Update-WordFields $doc
    $doc.Save()
    $saved = $true
} finally {
    if ($null -ne $doc) { try { $doc.Close($true) | Out-Null } catch {} }
    if ($null -ne $word) { try { $word.Quit() | Out-Null } catch {} }
}
if (-not $saved) { exit 1 }
"""
    env = os.environ.copy()
    env["BMS_DOCX_PATH"] = str(docx_path)
    errors: list[str] = []
    script_path: Path | None = None
    original_bytes = docx_path.read_bytes()
    try:
        with tempfile.NamedTemporaryFile("w", encoding="utf-8", suffix=".ps1", delete=False) as tmp:
            tmp.write(ps_script)
            script_path = Path(tmp.name)
        for exe in ("powershell.exe", "powershell"):
            docx_path.write_bytes(original_bytes)
            try:
                result = subprocess.run(
                    [exe, "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(script_path)],
                    check=False,
                    timeout=180,
                    capture_output=True,
                    text=True,
                    env=env,
                )
                if result.returncode == 0:
                    if _docx_contains_broken_reference_text(docx_path):
                        errors.append(f"{exe}: field update produced broken reference results")
                        docx_path.write_bytes(original_bytes)
                        continue
                    return True, ""
                detail = (result.stderr or result.stdout or f"exit={result.returncode}").strip()
                errors.append(f"{exe}: {detail[:300]}")
            except (OSError, subprocess.SubprocessError) as exc:
                errors.append(f"{exe}: {exc}")
    finally:
        if script_path is not None:
            try:
                script_path.unlink()
            except OSError:
                pass
    docx_path.write_bytes(original_bytes)
    return False, "; ".join(errors[-2:])


def update_fields_with_word(docx_path: Path) -> list[str]:
    if os.environ.get("BMS_NO_WORD") == "1":
        return ["word_field_update_skipped:BMS_NO_WORD=1"]

    # Word/KWPS COM resolves relative paths against the application process,
    # not the caller's working directory. Always hand the field engine a fully
    # qualified path so CLI output directories work exactly like GUI paths.
    docx_path = Path(docx_path).resolve()
    ok, python_error = _run_python_word_field_update(docx_path)
    if ok:
        return []
    ok, powershell_error = _run_powershell_word_field_update(docx_path)
    if ok:
        return []
    detail = "; ".join(part for part in (python_error, powershell_error) if part)
    return [f"word_field_update_failed:{detail[:1000]}"]


def find_last_paragraph_contains(doc: Document, fragment: str):
    matches = [para for para in doc.paragraphs if fragment in para.text.strip()]
    if not matches:
        raise ValueError(f"Paragraph containing '{fragment}' not found")
    return matches[-1]


def find_last_paragraph_matching(doc: Document, pattern: str):
    regex = re.compile(pattern)
    matches = [para for para in doc.paragraphs if regex.search(para.text.strip())]
    if not matches:
        raise ValueError(f"Paragraph matching '{pattern}' not found")
    return matches[-1]


def apply_wim_period_to_doc(doc: Document, section: dict) -> None:
    if not section.get("enabled"):
        return

    summaries = section["month_summaries"]
    summary_text = section["summary"]

    set_summary_table(doc, summary_text)

    wim_heading = find_last_paragraph(doc, T_WIM)
    next_heading = find_last_paragraph(doc, T_NEXT)
    table_templates = capture_wim_table_templates(wim_heading, next_heading)
    heading_tpl = capture_paragraph_template(find_last_paragraph_matching(doc, r"\d{4}年\d{1,2}月交通状况监测"))
    body_tpl = capture_paragraph_template(find_last_paragraph_contains(doc, "桥梁共通过车辆"))
    caption_tpl = capture_paragraph_template(find_last_paragraph_contains(doc, "季度交通状况分月统计表"))
    subcap_tpl = capture_paragraph_template(find_last_paragraph_matching(doc, r"续表\s*4-\d+"))
    fig_tpl = capture_paragraph_template(find_last_paragraph_contains(doc, "桥梁交通流参数分析"))

    clear_section_between(wim_heading, next_heading)
    add_text_paragraph_before(next_heading, summary_text, body_tpl)
    add_quarter_overview(next_heading, summaries, caption_tpl, table_templates)
    for idx, item in enumerate(summaries, start=1):
        add_month_block(
            next_heading,
            item,
            heading_tpl,
            body_tpl,
            caption_tpl,
            fig_tpl,
            subcap_tpl,
            section_index=idx,
            base_table_no=2 + (idx - 1) * 3,
            figure_no=idx,
            table_templates=table_templates,
        )


def summarize_missing_wim_images(section: dict) -> list[str]:
    missing: list[str] = []
    if not section.get("enabled"):
        return missing
    expected_labels = {
        "(a) 不同车道车辆数",
        "(b) 不同车速车辆数",
        "(c) 不同重量车辆数",
        "(d) 不同时间段车辆总数",
        "(e) 不同时间段平均车速",
        "(f) 大于50t车辆时间分布",
    }
    for item in section.get("month_summaries", []):
        found_labels = {label for label, _ in item.plot_paths}
        for label in sorted(expected_labels - found_labels):
            missing.append(f"wim:{item.yyyymm}:{label}")
    return missing


def build_period_report(
    template: Path,
    config_path: Path,
    result_root: Path,
    analysis_root: Path | None = None,
    image_root: Path | None = None,
    wim_root: Path | None = None,
    output_dir: Path | None = None,
    period_label: str = "2026年1-3月",
    monitoring_range: str = "2026年01月01日~2026年03月31日",
    report_number: str | None = None,
    report_date: str | None = None,
    start_date: str | None = None,
    end_date: str | None = None,
    precheck_template: bool = True,
) -> tuple[Path, Path, list[str]]:
    if report_date is None:
        report_date = datetime.now().strftime("%Y年%m月%d日")

    if start_date and end_date:
        start_dt = parse_date_str(start_date)
        end_dt = parse_date_str(end_date)
    else:
        extracted = extract_dates_from_range(monitoring_range)
        if extracted is None:
            raise ValueError("Unable to derive start/end dates. Provide --start-date and --end-date.")
        start_dt, end_dt = extracted
    if report_number is None:
        report_number = period_report_number(start_dt, end_dt)

    ctx = ReportBuildContext.from_inputs(
        template=template,
        config_path=config_path,
        result_root=result_root,
        analysis_root=analysis_root,
        image_root=image_root,
        output_dir=output_dir,
        wim_root=wim_root,
        assets_subdir="generated_assets",
    )

    cfg = load_json(config_path)
    reporting_cfg = cfg.get("reporting", {}) if isinstance(cfg.get("reporting", {}), dict) else {}
    front_matter_pages = int(reporting_cfg.get("front_matter_pages", HONGTANG_FRONT_MATTER_PAGES))
    if front_matter_pages < 0:
        raise ValueError("reporting.front_matter_pages must be non-negative")
    manifest = build_manifest(cfg, ctx.stats_root, ctx.fallback_stats_root, ctx.image_root, template, ctx.assets_dir, period_label, monitoring_range, report_date)
    manifest["report_number"] = report_number
    manifest["front_matter_pages"] = front_matter_pages
    manifest["analysis_run_manifest"] = ctx.analysis_context()
    wim_months = months_between(start_dt, end_dt)
    try:
        resolved_wim_root = resolve_wim_root(result_root, ctx.analysis_root, wim_root)
        manifest["wim"] = build_wim_period_section(resolved_wim_root, wim_months, cfg)
    except FileNotFoundError as exc:
        fallback_wim_root = wim_root if wim_root is not None else (result_root / "WIM" / "results" / "hongtang")
        manifest["wim"] = {
            "enabled": False,
            "wim_root": str(fallback_wim_root),
            "months": wim_months,
            "warnings": [str(exc)],
            "summary": "",
            "month_summaries": [],
        }
    manifest["data_coverage_audit"] = load_data_coverage_audit(result_root)
    lowfreq_missing_events = collect_lowfreq_missing_events(cfg, result_root, start_dt, end_dt)
    highfreq_missing_events = collect_highfreq_missing_events(cfg, result_root, start_dt, end_dt)
    raw_health_summary = build_health_status_summary(
        cfg,
        result_root,
        start_dt,
        end_dt,
        lowfreq_events=lowfreq_missing_events,
        highfreq_events=highfreq_missing_events,
    )
    raw_health_rows = build_health_status_rows(
        cfg,
        result_root,
        start_dt,
        end_dt,
        lowfreq_events=lowfreq_missing_events,
        highfreq_events=highfreq_missing_events,
    )
    missing_rows = build_report_missing_rows(manifest, manifest["wim"])
    manifest["health_status_summary"] = merge_health_status_summary(raw_health_summary, missing_rows)
    manifest["health_status_rows"] = raw_health_rows + missing_rows

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if precheck_template:
        raise_for_template("hongtang_period", template, manifest)

    doc = Document(str(template))
    apply_manifest_to_doc(doc, manifest)
    apply_period_cover_title(doc, start_dt, end_dt)
    apply_period_maintenance_text(doc)
    apply_period_maintenance_log(doc, start_dt, end_dt)
    apply_health_status_to_doc(doc, manifest["health_status_summary"], manifest["health_status_rows"])
    apply_wim_period_to_doc(doc, manifest["wim"])
    apply_period_toc_months(doc, start_dt, end_dt)
    convert_static_captions_to_auto_number(doc)
    if ensure_header_pagination_fields(doc, front_matter_pages=front_matter_pages) != 1:
        raise RuntimeError("Hongtang period template must contain exactly one report header pagination cell")

    output_docx = ctx.output_dir / period_report_filename(period_label, timestamp)
    doc.save(str(output_docx))
    _patch_report_number_in_docx(output_docx, report_number)
    field_update_warnings = update_fields_with_word(output_docx)
    header_audit = audit_header_pagination_fields(output_docx, front_matter_pages=front_matter_pages)
    if not header_audit.valid:
        detail = list(header_audit.details) + list(header_audit.formatting_errors)
        raise RuntimeError("Invalid Hongtang PAGE/adjusted-NUMPAGES header fields: " + "; ".join(detail))

    missing = summarize_missing_images(manifest) + summarize_missing_wim_images(manifest["wim"])
    missing.extend(missing_module_summary_items(manifest.get("analysis_run_manifest")))
    warnings = list(manifest["wim"].get("warnings", [])) + field_update_warnings
    qc_paths: dict[str, str] = {}
    try:
        qc_result = check_report("hongtang_period", output_docx)
        qc_txt, qc_json = write_report_qc_report(qc_result, ctx.output_dir, timestamp=timestamp)
        qc_paths = {"report_qc_txt": str(qc_txt), "report_qc_json": str(qc_json), "report_qc_status": qc_result.status}
        warnings = list(warnings) + [
            f"{issue.code}: {issue.message}"
            for issue in qc_result.issues
            if issue.severity == "warning"
        ]
    except Exception as exc:
        warnings = list(warnings) + [f"report_qc_failed: {exc}"]
    missing.extend(f"warning:{msg}" for msg in warnings)
    manifest_path = write_report_build_manifest(
        context=ctx,
        report_type="hongtang_period",
        output_docx=output_docx,
        timestamp=timestamp,
        legacy_manifest=manifest,
        missing=missing,
        warnings=warnings,
        extra={"report_number": report_number, **qc_paths},
        filename_prefix="period_report_manifest",
    )
    write_missing_summary(
        "洪塘周期报",
        output_docx,
        missing,
        context={"manifest": str(manifest_path), "result_root": str(result_root), "wim_root": str(wim_root or "")},
    )
    return manifest_path, output_docx, missing


def _json_default(value):
    if is_dataclass(value):
        return asdict(value)
    if isinstance(value, Path):
        return str(value)
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    raise TypeError(f"Object of type {value.__class__.__name__} is not JSON serializable")


def main() -> None:
    args = parse_args()
    if args.template is None or not args.template.exists():
        raise SystemExit("Template docx not found.")
    if not args.config.exists():
        raise SystemExit("Config file not found.")
    if not args.result_root.exists():
        raise SystemExit("Result root not found.")

    manifest_path, report_path, missing = build_period_report(
        template=args.template,
        config_path=args.config,
        result_root=args.result_root,
        analysis_root=args.analysis_root,
        image_root=args.image_root,
        wim_root=args.wim_root,
        output_dir=args.output_dir,
        period_label=args.period_label,
        monitoring_range=args.monitoring_range,
        report_number=args.report_number,
        report_date=args.report_date,
        start_date=args.start_date,
        end_date=args.end_date,
        precheck_template=not args.skip_template_precheck,
    )
    print(f"Manifest written to: {manifest_path}")
    print(f"Report written to:   {report_path}")
    if args.debug_section:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8-sig"))
        key = args.debug_section
        if key == "wim":
            payload = manifest.get("wim", {})
        elif key == "health_status":
            payload = {
                "summary": manifest.get("health_status_summary", ""),
                "rows": manifest.get("health_status_rows", []),
            }
        else:
            payload = manifest.get("sections", {}).get(key, {})
        print(f"Debug section [{key}]:")
        print(json.dumps(payload, ensure_ascii=False, indent=2))
    if missing:
        print("Warnings / missing assets:")
        for item in missing:
            print(f"  - {item}")


if __name__ == "__main__":
    main()
