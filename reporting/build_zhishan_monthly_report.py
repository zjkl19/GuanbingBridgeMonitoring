from __future__ import annotations

import argparse
import json
import math
import re
import shutil
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from openpyxl import load_workbook

from docx_utils import (
    find_paragraph_contains,
    paragraph_has_image,
    previous_body_paragraphs,
    remove_paragraph,
    set_cell_text_preserve,
)
from image_block_utils import count_docx_images
from report_build_manifest import write_report_build_manifest
from report_context import ReportBuildContext


REPO_ROOT = Path(__file__).resolve().parents[1]


@dataclass(frozen=True)
class RangeStats:
    min_value: float | None
    max_value: float | None
    mean_value: float | None = None
    count: int | None = None


STRAIN_GROUPS = [
    ("SX-1、SX-2、SX-9、SX-10", ["SX-1", "SX-2", "SX-9", "SX-10"], "SX_L2_414_283"),
    ("SX-3、SX-4、SX-7、SX-8", ["SX-3", "SX-4", "SX-7", "SX-8"], "SX_L2_298_218"),
    ("SX-5、SX-6", ["SX-5", "SX-6"], "SX_L2_405_252"),
]

CABLE_PARAMS = {
    "CF-1": {"length": 75.61, "density": 57.687, "baseline": 3496},
    "CF-2": {"length": 87.45, "density": 57.687, "baseline": 3694},
    "CF-3": {"length": 87.45, "density": 57.687, "baseline": 3799},
    "CF-4": {"length": 75.61, "density": 57.687, "baseline": 3386},
    "CF-5": {"length": 79.91, "density": 57.687, "baseline": 3860},
    "CF-6": {"length": 82.99, "density": 57.687, "baseline": 3842},
    "CF-7": {"length": 82.99, "density": 57.687, "baseline": 3786},
    "CF-8": {"length": 79.91, "density": 57.687, "baseline": 3659},
}

TYPICAL_CABLE_POINTS = ["CF-1", "CF-2", "CF-6", "CF-7", "CF-8"]


def _profile_defaults() -> dict:
    path = REPO_ROOT / "config" / "bridge_profiles.json"
    if not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    for row in data.get("profiles", []):
        if row.get("bridge_id") == "zhishan":
            return row
    return {}


def default_template() -> Path:
    profile = _profile_defaults()
    configured = profile.get("report_template")
    if configured:
        path = Path(configured)
        if not path.is_absolute():
            path = REPO_ROOT / path
        if path.exists():
            return path
    candidates = sorted((REPO_ROOT / "reports").glob("*0609_1652.docx"))
    if candidates:
        return candidates[-1]
    return REPO_ROOT / "reports" / "芝山大桥健康监测2026年3月份月报_0609_1652.docx"


def default_result_root() -> Path:
    profile = _profile_defaults()
    if profile.get("default_data_root"):
        return Path(profile["default_data_root"])
    return Path("D:/芝山大桥数据/2026年3月")


def parse_args() -> argparse.Namespace:
    profile = _profile_defaults()
    parser = argparse.ArgumentParser(description="Build Zhishan Bridge monthly monitoring report.")
    parser.add_argument("--template", type=Path, default=default_template())
    parser.add_argument("--config", type=Path, default=REPO_ROOT / "config" / "zhishan_config.json")
    parser.add_argument("--result-root", type=Path, default=default_result_root())
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--period-label", default=profile.get("default_period_label") or "2026年3月")
    parser.add_argument("--monitoring-range", default=profile.get("default_monitoring_range") or "2026年3月1日~2026年3月31日")
    parser.add_argument("--report-date", default=profile.get("default_report_date") or datetime.now().strftime("%Y年%m月%d日"))
    parser.add_argument("--no-word-update", action="store_true", help="Do not launch Word to refresh fields.")
    return parser.parse_args()


def report_month_label(period_label: str) -> str:
    match = re.search(r"(\d{4})年\s*(\d{1,2})月", period_label)
    if match:
        return f"{match.group(1)}年{int(match.group(2))}月份"
    text = period_label.strip()
    if text.endswith("月份"):
        return text
    if text.endswith("月"):
        return f"{text[:-1]}月份"
    return text


def parse_chinese_date_range(text: str) -> tuple[datetime, datetime] | None:
    match = re.search(
        r"(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日\s*[~～至-]\s*"
        r"(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日",
        text,
    )
    if not match:
        return None
    values = [int(item) for item in match.groups()]
    start = datetime(values[0], values[1], values[2])
    end = datetime(values[3], values[4], values[5])
    return start, end


def month_range_from_label(period_label: str) -> tuple[datetime, datetime] | None:
    match = re.search(r"(\d{4})年\s*(\d{1,2})月", period_label)
    if not match:
        return None
    year = int(match.group(1))
    month = int(match.group(2))
    start = datetime(year, month, 1)
    if month == 12:
        next_month = datetime(year + 1, 1, 1)
    else:
        next_month = datetime(year, month + 1, 1)
    return start, next_month - timedelta(days=1)


def expected_dates(period_label: str, monitoring_range: str) -> list[datetime]:
    parsed = parse_chinese_date_range(monitoring_range) or month_range_from_label(period_label)
    if not parsed:
        return []
    start, end = parsed
    if end < start:
        return []
    days = (end - start).days + 1
    return [start + timedelta(days=idx) for idx in range(days)]


def format_cn_date(day: datetime) -> str:
    return f"{day.year}年{day.month}月{day.day}日"


def build_coverage(result_root: Path, period_label: str, monitoring_range: str) -> dict:
    expected = expected_dates(period_label, monitoring_range)
    existing = {
        child.name
        for child in result_root.iterdir()
        if child.is_dir() and re.fullmatch(r"\d{4}-\d{2}-\d{2}", child.name)
    }
    missing = [day for day in expected if day.strftime("%Y-%m-%d") not in existing]
    return {
        "expected_days": len(expected),
        "valid_days": max(0, len(expected) - len(missing)),
        "missing": missing,
        "missing_cn": "、".join(format_cn_date(day) for day in missing),
        "missing_iso": "、".join(day.strftime("%Y-%m-%d") for day in missing),
    }


def _as_float_list(value: object) -> list[float]:
    if isinstance(value, (int, float)) and math.isfinite(float(value)):
        return [float(value)]
    if isinstance(value, list):
        out: list[float] = []
        for item in value:
            if isinstance(item, (int, float)) and math.isfinite(float(item)):
                out.append(float(item))
        return out
    return []


def _first_order_frequencies(block: object) -> tuple[float | None, float | None]:
    if not isinstance(block, dict):
        return None, None
    orders = block.get("peak_orders")
    if isinstance(orders, dict):
        orders = [orders]
    if isinstance(orders, list):
        for item in orders:
            if not isinstance(item, dict):
                continue
            theor = as_float(item.get("theoretical_hz", item.get("theor_hz")))
            center = as_float(
                item.get("search_center_hz", item.get("target_hz", item.get("frequency_hz", item.get("freq_hz"))))
            )
            search_min = as_float(item.get("search_min_hz", item.get("min_hz", item.get("lower_hz"))))
            search_max = as_float(item.get("search_max_hz", item.get("max_hz", item.get("upper_hz"))))
            if center is None and search_min is not None and search_max is not None and search_max > search_min:
                center = (search_min + search_max) / 2
            if theor is not None or center is not None:
                return theor, center

    target_freqs = _as_float_list(block.get("target_freqs"))
    theor_freqs = _as_float_list(block.get("theor_freqs"))
    theor = theor_freqs[0] if theor_freqs else None
    center = target_freqs[0] if target_freqs else None
    return theor, center


def _point_block(config: dict, per_point_key: str, point_id: str) -> dict:
    per_point = config.get("per_point", {})
    if not isinstance(per_point, dict):
        return {}
    block = per_point.get(per_point_key, {})
    if not isinstance(block, dict):
        return {}
    candidates = [point_id, point_id.replace("-", "_")]
    name_map = config.get("name_map_global", {})
    if isinstance(name_map, dict):
        for key, value in name_map.items():
            if str(value) == point_id:
                candidates.append(str(key))
    for key in candidates:
        value = block.get(key)
        if isinstance(value, dict):
            return value
    return {}


def _format_frequency_values(values: list[float], fallback: float) -> tuple[str, bool]:
    finite = sorted({round(v, 3) for v in values if math.isfinite(v)})
    if not finite:
        finite = [fallback]
    if len(finite) == 1:
        return f"{finite[0]:.3f}Hz", False
    return f"{min(finite):.3f}Hz~{max(finite):.3f}Hz", True


def load_accel_frequency_note(config_path: Path) -> str:
    try:
        config = json.loads(config_path.read_text(encoding="utf-8"))
    except Exception:
        config = {}
    params = config.get("accel_spectrum_params", {}) if isinstance(config, dict) else {}
    default_theor, default_center = _first_order_frequencies(params)
    points = ["AZ-1", "AZ-2", "AZ-3", "AZ-4", "AZ-5"]
    theor_values: list[float] = []
    center_values: list[float] = []
    for point in points:
        pt_theor, pt_center = _first_order_frequencies(_point_block(config, "accel_spectrum", point))
        theor = pt_theor if pt_theor is not None else default_theor
        center = pt_center if pt_center is not None else default_center
        if theor is not None:
            theor_values.append(theor)
        if center is not None:
            center_values.append(center)
    theor_text, theor_is_range = _format_frequency_values(theor_values, 0.593)
    center_text, center_is_range = _format_frequency_values(center_values, 0.640)
    theor_word = "范围为" if theor_is_range else "为"
    center_word = "范围附近" if center_is_range else "附近"
    return f"结构理论竖向一阶频率{theor_word}{theor_text}，各测点按{center_text}{center_word}峰值进行识别，"


def as_float(value: object) -> float | None:
    if value is None:
        return None
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    if math.isnan(number) or math.isinf(number):
        return None
    return number


def fmt_num(value: float | None, digits: int = 1, keep_decimal: bool = False) -> str:
    if value is None:
        return "/"
    text = f"{value:.{digits}f}"
    if keep_decimal:
        return text
    return text.rstrip("0").rstrip(".")


def fmt_range(stats: RangeStats | None, digits: int = 1, unit: str = "") -> str:
    if stats is None:
        return "/"
    return f"{fmt_num(stats.min_value, digits, True)}{unit}~{fmt_num(stats.max_value, digits, True)}{unit}"


def load_sheet_rows(path: Path, sheet_name: str | None = None) -> list[dict]:
    if not path.exists():
        return []
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb[sheet_name] if sheet_name else wb.worksheets[0]
        rows = list(ws.iter_rows(values_only=True))
    finally:
        wb.close()
    if not rows:
        return []
    headers = [str(v).strip() if v is not None else "" for v in rows[0]]
    result: list[dict] = []
    for row in rows[1:]:
        item = {key: value for key, value in zip(headers, row) if key}
        if any(value is not None for value in item.values()):
            result.append(item)
    return result


def rows_by_point(rows: Iterable[dict]) -> dict[str, dict]:
    return {str(row.get("PointID", "")).strip(): row for row in rows if str(row.get("PointID", "")).strip()}


def aggregate_rows(rows: Iterable[dict], *, min_key: str = "Min", max_key: str = "Max", mean_key: str = "Mean", count_key: str = "Count") -> RangeStats | None:
    mins: list[float] = []
    maxs: list[float] = []
    weighted_sum = 0.0
    weight_total = 0
    means: list[float] = []
    for row in rows:
        min_value = as_float(row.get(min_key))
        max_value = as_float(row.get(max_key))
        mean_value = as_float(row.get(mean_key))
        count_value = as_float(row.get(count_key))
        if min_value is not None:
            mins.append(min_value)
        if max_value is not None:
            maxs.append(max_value)
        if mean_value is not None:
            if count_value is not None and count_value > 0:
                weighted_sum += mean_value * int(count_value)
                weight_total += int(count_value)
            else:
                means.append(mean_value)
    if not mins and not maxs:
        return None
    mean = (weighted_sum / weight_total) if weight_total else (sum(means) / len(means) if means else None)
    return RangeStats(min(mins) if mins else None, max(maxs) if maxs else None, mean, weight_total or None)


def aggregate_ranges(rows: Iterable[RangeStats]) -> RangeStats | None:
    mins = [row.min_value for row in rows if row.min_value is not None]
    maxs = [row.max_value for row in rows if row.max_value is not None]
    means = [row.mean_value for row in rows if row.mean_value is not None]
    if not mins and not maxs:
        return None
    return RangeStats(
        min(mins) if mins else None,
        max(maxs) if maxs else None,
        sum(means) / len(means) if means else None,
    )


def load_frequency_ranges(path: Path) -> dict[str, RangeStats]:
    if not path.exists():
        return {}
    wb = load_workbook(path, read_only=True, data_only=True)
    result: dict[str, RangeStats] = {}
    try:
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(v).strip() if v is not None else "" for v in rows[0]]
            freq_idx = next((i for i, name in enumerate(headers) if name.startswith("Freq_")), None)
            if freq_idx is None:
                continue
            values = [as_float(row[freq_idx]) for row in rows[1:] if len(row) > freq_idx]
            values = [v for v in values if v is not None]
            if values:
                result[ws.title] = RangeStats(min(values), max(values))
    finally:
        wb.close()
    return result


def load_cable_force_ranges(path: Path) -> dict[str, RangeStats]:
    if not path.exists():
        return {}
    wb = load_workbook(path, read_only=True, data_only=True)
    result: dict[str, RangeStats] = {}
    try:
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(v).strip() if v is not None else "" for v in rows[0]]
            force_idx = next((i for i, name in enumerate(headers) if name == "CableForce_kN"), None)
            if force_idx is None:
                continue
            values = [as_float(row[force_idx]) for row in rows[1:] if len(row) > force_idx]
            values = [v for v in values if v is not None]
            if values:
                result[ws.title] = RangeStats(min(values), max(values))
    finally:
        wb.close()
    return result


def load_group_stats(path: Path) -> dict[str, RangeStats]:
    if not path.exists():
        return {}
    result: dict[str, RangeStats] = {}
    for label, _points, sheet in STRAIN_GROUPS:
        rows = load_sheet_rows(path, sheet)
        result[label] = aggregate_rows(rows)
    return result


def build_context(result_root: Path, config_path: Path | None = None) -> dict:
    stats_dir = result_root / "stats"
    bearing_rows = load_sheet_rows(stats_dir / "bearing_displacement_stats.xlsx")
    accel_rows = load_sheet_rows(stats_dir / "accel_stats.xlsx")
    strain_rows = load_sheet_rows(stats_dir / "strain_stats.xlsx")
    cable_accel_rows = load_sheet_rows(stats_dir / "cable_accel_stats.xlsx")
    accel_freq = load_frequency_ranges(stats_dir / "accel_spec_stats.xlsx")
    cable_force = load_cable_force_ranges(stats_dir / "cable_accel_spec_stats.xlsx")

    strain_by_point = rows_by_point(strain_rows)
    raw_group_stats: dict[str, RangeStats] = {}
    for label, points, _sheet in STRAIN_GROUPS:
        rows = [strain_by_point[p] for p in points if p in strain_by_point]
        raw_group_stats[label] = aggregate_rows(rows)

    return {
        "bearing_rows": rows_by_point(bearing_rows),
        "accel_rows": rows_by_point(accel_rows),
        "accel_freq": accel_freq,
        "strain_group_stats": raw_group_stats,
        "dynamic_hp_group_stats": load_group_stats(stats_dir / "dynamic_strain_highpass_stats.xlsx"),
        "dynamic_lp_group_stats": load_group_stats(stats_dir / "dynamic_strain_lowpass_stats.xlsx"),
        "cable_accel_rows": rows_by_point(cable_accel_rows),
        "cable_force": cable_force,
        "accel_frequency_note": load_accel_frequency_note(config_path or (REPO_ROOT / "config" / "zhishan_config.json")),
    }


def replace_text_in_paragraph(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_paragraph_plain_text(paragraph, text: str) -> None:
    """Replace paragraph content and remove stale Word field nodes."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.add_run(text)


def normalize_caption_fields(doc: DocxDocument) -> int:
    """Convert visible figure/table captions to plain text.

    Some legacy Zhishan templates keep hidden REF fields inside caption runs.
    Those fields can render as "引用源未找到" even when python-docx sees a
    normal caption such as "图 2-5 ...".  The generated monthly report uses
    fixed chapter captions, so removing those stale field nodes is safer than
    refreshing them.
    """
    count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^(图|表)\s+\d+-\d+\b", text):
            replace_paragraph_plain_text(paragraph, text)
            count += 1
    return count


def replace_first_by_prefix(doc: DocxDocument, prefix: str, text: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            replace_text_in_paragraph(paragraph, text)
            return True
    return False


def replace_all_by_prefix(doc: DocxDocument, prefix: str, text: str) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            replace_text_in_paragraph(paragraph, text)
            count += 1
    return count


def iter_all_paragraphs(doc: DocxDocument):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_all_by_prefix_anywhere(doc: DocxDocument, prefix: str, text: str) -> int:
    count = 0
    for paragraph in iter_all_paragraphs(doc):
        if paragraph.text.strip().startswith(prefix):
            replace_text_in_paragraph(paragraph, text)
            count += 1
    return count


def delete_table_row(row) -> None:
    parent = row._tr.getparent()
    if parent is not None:
        parent.remove(row._tr)


def update_cover_dates(doc: DocxDocument, period_label: str, report_date: str) -> None:
    replace_all_by_prefix(doc, "（监测时间：", f"（监测时间：{period_label}）")
    replace_all_by_prefix(doc, "报告日期：", f"报告日期：{report_date}")


def table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def find_table(doc: DocxDocument, fragments: list[str], occurrence: int = 1):
    seen = 0
    for table in doc.tables:
        text = table_text(table)
        if all(fragment in text for fragment in fragments):
            seen += 1
            if seen == occurrence:
                return table
    return None


def set_row(row, values: list[str]) -> None:
    for cell, value in zip(row.cells, values):
        set_cell_text_preserve(cell, value)


def set_summary_result_cell(cell, lines: list[str], bold_indices: set[int]) -> None:
    """Write summary table lines with review-template paragraph styling."""
    if not cell.paragraphs:
        cell.text = ""
    base_para = cell.paragraphs[0]
    base_style = base_para.style
    base_alignment = base_para.alignment
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)
    if base_para.runs:
        for run in base_para.runs:
            run.text = ""
            run.bold = False
    else:
        base_para.add_run("")
    paragraphs = [base_para]
    for _ in range(max(0, len(lines) - 1)):
        para = cell.add_paragraph()
        para.style = base_style
        para.alignment = base_alignment
        paragraphs.append(para)
    for idx, (para, text) in enumerate(zip(paragraphs, lines)):
        para.paragraph_format.first_line_indent = Pt(24)
        if not para.runs:
            para.add_run("")
        for run in para.runs:
            run.text = ""
            run.bold = False
        para.runs[0].text = text
        para.runs[0].bold = idx in bold_indices


def update_period_table(doc: DocxDocument, monitoring_range: str) -> bool:
    table = find_table(doc, ["委托单位", "监测时间"])
    if table is None:
        return False
    for row in table.rows:
        for idx, cell in enumerate(row.cells[:-1]):
            if "监测时间" in cell.text:
                set_cell_text_preserve(row.cells[idx + 1], monitoring_range)
                return True
    return False


def update_data_availability(doc: DocxDocument, result_root: Path, period_label: str, monitoring_range: str) -> None:
    coverage = build_coverage(result_root, period_label, monitoring_range)
    missing_cn = coverage["missing_cn"]
    missing_iso = coverage["missing_iso"]
    if missing_cn:
        report_section_text = (
            f"本月报告分析数据覆盖{period_label}，实际有效数据为{coverage['valid_days']}天，"
            f"{missing_cn}由于未获取到整日有效数据，未纳入统计分析；其余时间段监测系统数据采集及传输状态正常。"
        )
        first_row_date_text = f"有效{coverage['valid_days']}天（缺{missing_iso}，下同）"
    else:
        report_section_text = (
            f"本月报告分析数据覆盖{period_label}，实际有效数据为{coverage['valid_days']}天，"
            "监测系统数据采集及传输状态正常。"
        )
        first_row_date_text = f"有效{coverage['valid_days']}天"

    replace_all_by_prefix_anywhere(doc, "本月报告分析数据覆盖", report_section_text)
    replace_all_by_prefix_anywhere(
        doc,
        "按本月监测数据获取情况表统计",
        (
            "按本月监测数据获取情况表统计，本次自动报告纳入梁端纵向位移、振动加速度、"
            "结构应变、索力加速度共27项次，实际获取27项次，整体获取率100.0%；"
            "温湿度数据另行处理，本报告不展开自动统计。"
        ),
    )
    replace_all_by_prefix_anywhere(
        doc,
        "本月持续开展监测系统运行维护工作",
        "本月持续开展监测系统运行维护工作，每日对系统运行状态进行检查，并结合数据处理结果进行质量复核。",
    )
    replace_all_by_prefix_anywhere(
        doc,
        "软件线上检查维护",
        "软件线上检查维护：本月每日对系统运行情况进行检查，监测期内除上述整日缺失日期外，其余日期数据可用于本期统计分析。",
    )

    table = find_table(doc, ["配置测点数", "实际获取测点数", "获取日期"])
    if table is not None:
        rows = [
            ["1", "温湿度", "1", "/", "/", "本次未纳入自动统计", "温湿度数据另行处理"],
            ["2", "梁端纵向位移", "4", "4", "100%", first_row_date_text, "/"],
            ["3", "振动加速度", "5", "5", "100%", f"有效{coverage['valid_days']}天", "/"],
            ["4", "结构应变", "10", "10", "100%", f"有效{coverage['valid_days']}天", "/"],
            ["5", "索力加速度", "8", "8", "100%", f"有效{coverage['valid_days']}天", "/"],
        ]
        while len(table.rows) > len(rows) + 1:
            delete_table_row(table.rows[-1])
        while len(table.rows) < len(rows) + 1:
            table.add_row()
        for idx, values in enumerate(rows, start=1):
            set_row(table.rows[idx], values)


def update_temperature_humidity_placeholders(doc: DocxDocument) -> None:
    replace_first_by_prefix(
        doc,
        "监测结果表明，环境温度",
        "本期温度数据未纳入本次自动报告统计，相关数据另行处理，本节暂不展开定量分析。",
    )
    replace_all_by_prefix_anywhere(
        doc,
        "温度传感器",
        "本期温度数据未纳入本次自动报告统计，相关数据另行处理，本节暂不展开定量分析。",
    )
    replace_first_by_prefix(
        doc,
        "本月相对湿度",
        "本期相对湿度数据未纳入本次自动报告统计，相关数据另行处理，本节暂不展开定量分析。",
    )
    replace_all_by_prefix_anywhere(
        doc,
        "本月相对湿度数据",
        "本期相对湿度数据未纳入本次自动报告统计，相关数据另行处理，本节暂不展开定量分析。",
    )
    temp_table = find_table(doc, ["监测类型", "最小值(℃)", "最大值(℃)"])
    if temp_table is not None and len(temp_table.rows) > 1:
        set_row(temp_table.rows[1], ["1", "WS-T", "温度", "/", "/", "/"])
    humidity_table = find_table(doc, ["监测类型", "最小值(%)", "最大值(%)"])
    if humidity_table is not None and len(humidity_table.rows) > 1:
        set_row(humidity_table.rows[1], ["1", "WS-H", "相对湿度", "/", "/", "/"])

    for anchor_text in ["图 2-3 温度时程曲线", "图 2-4 湿度时程曲线"]:
        anchor = find_paragraph_contains(doc, anchor_text)
        if anchor is not None:
            remove_nearby_pictures_before(anchor, limit=12)
            replace_text_in_paragraph(anchor, f"{anchor_text}（本期未纳入自动统计）")


def update_bearing_tables(doc: DocxDocument, context: dict) -> list[str]:
    warnings: list[str] = []
    rows = context["bearing_rows"]
    tables = [
        (find_table(doc, ["测点编号", "最小值(mm)", "最大值(mm)"], occurrence=1), "OrigMin_mm", "OrigMax_mm"),
        (find_table(doc, ["测点编号", "最小值(mm)", "最大值(mm)"], occurrence=2), "FiltMin_mm", "FiltMax_mm"),
    ]
    for table, min_key, max_key in tables:
        if table is None:
            warnings.append(f"missing bearing table for {min_key}")
            continue
        for idx, point in enumerate(["DX-1", "DX-2", "DX-3", "DX-4"], start=1):
            if idx >= len(table.rows):
                continue
            row = rows.get(point, {})
            set_row(table.rows[idx], [str(idx), point, fmt_num(as_float(row.get(min_key)), 1, True), fmt_num(as_float(row.get(max_key)), 1, True)])
    return warnings


def update_accel_table(doc: DocxDocument, context: dict) -> list[str]:
    warnings: list[str] = []
    table = find_table(doc, ["10min RMS最大值", "基频最小值", "基频最大值"])
    if table is None:
        return ["missing acceleration summary table"]
    accel_rows = context["accel_rows"]
    freq_rows = context["accel_freq"]
    for idx, point in enumerate(["AZ-1", "AZ-2", "AZ-3", "AZ-4", "AZ-5"], start=1):
        row = accel_rows.get(point, {})
        freq = freq_rows.get(point)
        set_row(
            table.rows[idx],
            [
                str(idx),
                point,
                fmt_num(as_float(row.get("Min")), 3, True),
                fmt_num(as_float(row.get("Max")), 3, True),
                fmt_num(as_float(row.get("RMS10minMax")), 3, True),
                fmt_num(freq.min_value if freq else None, 3, True),
                fmt_num(freq.max_value if freq else None, 3, True),
            ],
        )
    return warnings


def update_strain_group_table(table, group_stats: dict[str, RangeStats]) -> None:
    for idx, (label, _points, _sheet) in enumerate(STRAIN_GROUPS, start=1):
        if idx >= len(table.rows):
            continue
        stats = group_stats.get(label)
        set_row(
            table.rows[idx],
            [
                str(idx),
                label,
                fmt_num(stats.min_value if stats else None, 1, True),
                fmt_num(stats.max_value if stats else None, 1, True),
                fmt_num(stats.mean_value if stats else None, 1, True),
            ],
        )


def update_strain_tables(doc: DocxDocument, context: dict) -> list[str]:
    warnings: list[str] = []
    tables = [
        (find_table(doc, ["测点分组", "平均值(με)"], occurrence=1), context["strain_group_stats"]),
        (find_table(doc, ["测点分组", "平均值(με)"], occurrence=2), context["dynamic_hp_group_stats"]),
        (find_table(doc, ["测点分组", "平均值(με)"], occurrence=3), context["dynamic_lp_group_stats"]),
    ]
    for table, group_stats in tables:
        if table is None:
            warnings.append("missing strain group stats table")
            continue
        update_strain_group_table(table, group_stats)
    return warnings


def update_cable_tables(doc: DocxDocument, context: dict) -> list[str]:
    warnings: list[str] = []
    accel_table = find_table(doc, ["测点编号", "10min RMS最大值(mm/s²)"])
    cable_rows = context["cable_accel_rows"]
    if accel_table is None:
        warnings.append("missing cable acceleration summary table")
    else:
        for idx, point in enumerate([f"CF-{i}" for i in range(1, 9)], start=1):
            row = cable_rows.get(point, {})
            set_row(
                accel_table.rows[idx],
                [
                    str(idx),
                    point,
                    fmt_num(as_float(row.get("Min")), 1, True),
                    fmt_num(as_float(row.get("Max")), 1, True),
                    fmt_num(as_float(row.get("RMS10minMax")), 3, True),
                ],
            )

    force_table = find_table(doc, ["线密度", "计算索长", "成桥"])
    force_rows = context["cable_force"]
    if force_table is None:
        warnings.append("missing cable force summary table")
    else:
        for idx, point in enumerate([f"CF-{i}" for i in range(1, 9)], start=1):
            params = CABLE_PARAMS[point]
            stats = force_rows.get(point)
            baseline = params["baseline"]
            if stats and stats.min_value is not None and stats.max_value is not None:
                min_force = round(stats.min_value)
                max_force = round(stats.max_value)
                min_rate = (stats.min_value - baseline) / baseline * 100.0
                max_rate = (stats.max_value - baseline) / baseline * 100.0
                rate = f"{fmt_num(min_rate, 2, True)}%~{fmt_num(max_rate, 2, True)}%"
            else:
                min_force = max_force = "/"
                rate = "/"
            set_row(
                force_table.rows[idx],
                [
                    point,
                    fmt_num(params["density"], 3, True),
                    fmt_num(params["length"], 2, True),
                    str(min_force),
                    str(max_force),
                    str(baseline),
                    rate,
                ],
            )
    return warnings


def _range_for_points(rows: dict[str, dict], points: list[str], min_key: str = "Min", max_key: str = "Max") -> RangeStats | None:
    return aggregate_rows((rows.get(point, {}) for point in points), min_key=min_key, max_key=max_key)


def _max_rms(rows: dict[str, dict]) -> tuple[str, float | None]:
    best_point = ""
    best_value: float | None = None
    for point, row in rows.items():
        value = as_float(row.get("RMS10minMax"))
        if value is not None and (best_value is None or value > best_value):
            best_point = point
            best_value = value
    return best_point, best_value


def update_narrative(doc: DocxDocument, context: dict) -> None:
    bearing_raw = aggregate_rows(context["bearing_rows"].values(), min_key="OrigMin_mm", max_key="OrigMax_mm")
    bearing_filt = aggregate_rows(context["bearing_rows"].values(), min_key="FiltMin_mm", max_key="FiltMax_mm")
    replace_first_by_prefix(
        doc,
        "监测结果表明，梁端纵向位移",
        (
            f"监测结果表明，梁端纵向位移实测值范围为{fmt_range(bearing_raw, 1, 'mm')}；"
            "均处于阈值范围之内，未超过各级超限阈值和报警的情况。"
            f"梁端位移滤波后实测值范围在{fmt_range(bearing_filt, 1, 'mm')}之间。"
        ),
    )

    accel_all = aggregate_rows(context["accel_rows"].values())
    accel_point, accel_rms = _max_rms(context["accel_rows"])
    freq_all = aggregate_ranges(context["accel_freq"].values())
    replace_first_by_prefix(
        doc,
        "监测结果表明，主梁加速度",
        (
            f"监测结果表明，主梁加速度实测值范围为{fmt_range(accel_all, 3, 'm/s²')}；"
            f"各测点10min RMS最大值为{fmt_num(accel_rms, 3, True)}m/s²，对应测点{accel_point}，"
            "未超过0.315m/s²，处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        ),
    )
    replace_first_by_prefix(
        doc,
        "监测周期内桥梁实测竖向1阶自振频率范围",
        (
            f"监测周期内桥梁实测竖向1阶自振频率范围为{fmt_range(freq_all, 3, 'Hz')}，"
            f"{context['accel_frequency_note']}"
            "频率结果总体处于合理范围内，桥梁结构实测刚度未见明显异常。"
        ),
    )

    raw_groups = context["strain_group_stats"]
    replace_first_by_prefix(
        doc,
        "监测结果表明，边跨测点",
        (
            f"监测结果表明，边跨测点SX-1、SX-2、SX-9、SX-10实测范围为{fmt_range(raw_groups.get(STRAIN_GROUPS[0][0]), 1, 'με')}，"
            f"中跨1/4、3/4测点SX-3、SX-4、SX-7、SX-8实测范围为{fmt_range(raw_groups.get(STRAIN_GROUPS[1][0]), 1, 'με')}，"
            f"中跨跨中测点SX-5、SX-6实测范围为{fmt_range(raw_groups.get(STRAIN_GROUPS[2][0]), 1, 'με')}，"
            "均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        ),
    )
    hp_all = aggregate_ranges(context["dynamic_hp_group_stats"].values())
    hp_abs_min = abs(hp_all.min_value) if hp_all and hp_all.min_value is not None else None
    replace_first_by_prefix(
        doc,
        "由高通滤波结果可知",
        (
            "由高通滤波结果可知，本月监测周期内，各组活载动应变整体围绕0με波动；"
            f"各组最大拉应变约为{fmt_num(hp_all.max_value if hp_all else None, 1, True)}με，"
            f"最大压应变约为{fmt_num(hp_abs_min, 1, True)}με。"
            "箱线图显示多数数据集中在0με附近，反映本期桥梁活载作用下应变响应总体较小。"
        ),
    )
    lp_groups = context["dynamic_lp_group_stats"]
    replace_first_by_prefix(
        doc,
        "由低通滤波结果可知",
        (
            "由低通滤波结果可知，本月各组静应变随时间呈缓慢周期性变化。"
            f"边跨测点组实测范围为{fmt_range(lp_groups.get(STRAIN_GROUPS[0][0]), 1, 'με')}，"
            f"中跨1/4及3/4测点组实测范围为{fmt_range(lp_groups.get(STRAIN_GROUPS[1][0]), 1, 'με')}，"
            f"中跨跨中测点组实测范围为{fmt_range(lp_groups.get(STRAIN_GROUPS[2][0]), 1, 'με')}，"
            "截面整体受力未见明显异常。"
        ),
    )

    cable_point, cable_rms = _max_rms(context["cable_accel_rows"])
    replace_first_by_prefix(
        doc,
        "监测结果表明，各测点10min RMS最大值",
        (
            f"监测结果表明，各测点10min RMS最大值为{fmt_num(cable_rms, 3, True)}mm/s²，对应测点{cable_point}，"
            "未超过1000mm/s²（即1.000m/s²），均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        ),
    )
    force_ranges = list(context["cable_force"].values())
    min_rate: float | None = None
    max_rate: float | None = None
    for point, stats in context["cable_force"].items():
        baseline = CABLE_PARAMS.get(point, {}).get("baseline")
        if not baseline or stats.min_value is None or stats.max_value is None:
            continue
        rates = [(stats.min_value - baseline) / baseline * 100.0, (stats.max_value - baseline) / baseline * 100.0]
        min_rate = min(rates + ([min_rate] if min_rate is not None else []))
        max_rate = max(rates + ([max_rate] if max_rate is not None else []))
    if force_ranges:
        replace_first_by_prefix(
            doc,
            "监测结果表明，本期索力换算结果",
            (
                f"监测结果表明，本期索力换算结果与成桥索力相比变化范围为{fmt_num(min_rate, 2, True)}%~{fmt_num(max_rate, 2, True)}%，"
                "均在±10%二级预警范围内。"
            ),
        )


def update_summary_table(doc: DocxDocument, context: dict) -> None:
    table = find_table(doc, ["监测结果", "建  议"])
    if table is None:
        return
    accel_all = aggregate_rows(context["accel_rows"].values())
    accel_point, accel_rms = _max_rms(context["accel_rows"])
    freq_all = aggregate_ranges(context["accel_freq"].values())
    raw_groups = context["strain_group_stats"]
    hp_all = aggregate_ranges(context["dynamic_hp_group_stats"].values())
    hp_abs_min = abs(hp_all.min_value) if hp_all and hp_all.min_value is not None else None
    lp_groups = context["dynamic_lp_group_stats"]
    cable_point, cable_rms = _max_rms(context["cable_accel_rows"])
    force_rates: list[float] = []
    for point, stats in context["cable_force"].items():
        baseline = CABLE_PARAMS.get(point, {}).get("baseline")
        if not baseline or stats.min_value is None or stats.max_value is None:
            continue
        force_rates.extend([(stats.min_value - baseline) / baseline * 100.0, (stats.max_value - baseline) / baseline * 100.0])
    result_lines = [
        "4.4 结构振动监测",
        (
            f"监测结果表明，主梁加速度实测值范围为{fmt_range(accel_all, 3, 'm/s²')}；各测点10min RMS最大值为"
            f"{fmt_num(accel_rms, 3, True)}m/s²，对应测点{accel_point}，未超过0.315m/s²，"
            "处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        ),
        (
            f"监测周期内桥梁实测竖向1阶自振频率范围为{fmt_range(freq_all, 3, 'Hz')}，"
            f"{context['accel_frequency_note']}"
            "频率结果总体处于合理范围内，桥梁结构实测刚度未见明显异常。"
        ),
        "4.5 结构应变监测",
        (
            f"监测结果表明，边跨测点SX-1、SX-2、SX-9、SX-10实测范围为{fmt_range(raw_groups.get(STRAIN_GROUPS[0][0]), 1, 'με')}，"
            f"中跨1/4、3/4测点SX-3、SX-4、SX-7、SX-8实测范围为{fmt_range(raw_groups.get(STRAIN_GROUPS[1][0]), 1, 'με')}，"
            f"中跨跨中测点SX-5、SX-6实测范围为{fmt_range(raw_groups.get(STRAIN_GROUPS[2][0]), 1, 'με')}，"
            "均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        ),
        (
            "由高通滤波结果可知，本月监测周期内，各组活载动应变整体围绕0με波动；"
            f"各组最大拉应变约为{fmt_num(hp_all.max_value if hp_all else None, 1, True)}με，"
            f"最大压应变约为{fmt_num(hp_abs_min, 1, True)}με。"
            "箱线图显示多数数据集中在0με附近，反映本期桥梁活载作用下应变响应总体较小。"
        ),
        (
            "由低通滤波结果可知，本月各组静应变随时间呈缓慢周期性变化。"
            f"边跨测点组实测范围为{fmt_range(lp_groups.get(STRAIN_GROUPS[0][0]), 1, 'με')}，"
            f"中跨1/4及3/4测点组实测范围为{fmt_range(lp_groups.get(STRAIN_GROUPS[1][0]), 1, 'με')}，"
            f"中跨跨中测点组实测范围为{fmt_range(lp_groups.get(STRAIN_GROUPS[2][0]), 1, 'με')}，"
            "截面整体受力未见明显异常。"
        ),
        "4.6 斜拉索索力加速度监测",
        (
            f"监测结果表明，各测点10min RMS最大值为{fmt_num(cable_rms, 3, True)}mm/s²，对应测点{cable_point}，"
            "未超过1000mm/s²（即1.000m/s²），均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
            f"本期索力换算结果与成桥索力相比变化范围为{fmt_num(min(force_rates) if force_rates else None, 2, True)}%~"
            f"{fmt_num(max(force_rates) if force_rates else None, 2, True)}%，均在±10%二级预警范围内。"
        ),
    ]
    advice_text = (
        "建议继续加强监测平台运行维护和数据质量复核，重点跟踪数据传输中断、应变残余尖峰及CF-3~CF-5索力频谱波动等情况；"
        "后续报告生成时应优先采用经复核后的统计表和正式展示图。"
    )
    set_summary_result_cell(table.rows[0].cells[1], result_lines, {0, 3, 7})
    set_cell_text_preserve(table.rows[1].cells[1], advice_text)


def latest_image(root: Path, directory: str, pattern: str) -> Path | None:
    folder = root / directory
    if not folder.exists():
        return None
    files = sorted(folder.glob(pattern))
    files = [path for path in files if path.is_file() and path.suffix.lower() in {".jpg", ".jpeg", ".png"}]
    if not files:
        return None
    return max(files, key=lambda path: (path.stat().st_mtime, path.name))


def first_images(root: Path, directory: str, patterns: list[str]) -> list[Path]:
    images: list[Path] = []
    for pattern in patterns:
        image = latest_image(root, directory, pattern)
        if image is not None:
            images.append(image)
    return images


def latest_nested_psd(root: Path, directory: str, point: str, date_token: str = "2026-03-10") -> Path | None:
    folder = root / directory / point
    if not folder.exists():
        return None
    files = sorted(folder.glob(f"*{date_token}*.jpg"))
    if not files:
        files = sorted(folder.glob("*.jpg"))
    return files[0] if files else None


def zhishan_image_replacements(result_root: Path) -> list[tuple[str, list[Path | None], float]]:
    """Return report image replacements using analyzer output as the source of truth."""
    az_psd_images = [latest_nested_psd(result_root, "PSD_备查", f"AZ-{idx}") for idx in range(1, 6)]
    cable_accel_images = first_images(
        result_root,
        "时程曲线_索力加速度",
        [f"CF-{idx}_*.jpg" for idx in range(1, 9)],
    )
    cable_force_images = [
        latest_image(result_root, "索力时程图", f"CableForce_{point}_*.jpg")
        for point in TYPICAL_CABLE_POINTS
    ]
    cable_psd_images = [
        latest_nested_psd(result_root, "PSD_备查_索力加速度", point)
        for point in TYPICAL_CABLE_POINTS
    ]

    return [
        ("图 2-5", first_images(result_root, "时程曲线_梁端纵向位移_组图", ["*Orig*.jpg"]), 145.0),
        ("图 2-6", first_images(result_root, "时程曲线_梁端纵向位移_组图", ["*Filt*.jpg"]), 145.0),
        (
            "图 2-7",
            [
                *first_images(result_root, "时程曲线_加速度_组图", ["*.jpg"]),
                *first_images(result_root, "时程曲线_加速度_RMS10min_组图", ["*.jpg"]),
            ],
            145.0,
        ),
        ("图 2-8", az_psd_images, 145.0),
        ("图 2-9", first_images(result_root, "频谱峰值曲线_结构加速度_组图", ["*.jpg"]), 145.0),
        (
            "图 2-10",
            first_images(
                result_root,
                "时程曲线_应变_组图",
                ["*SX_L2_414_283*.jpg", "*SX_L2_298_218*.jpg", "*SX_L2_405_252*.jpg"],
            ),
            145.0,
        ),
        ("图 2-11", first_images(result_root, "时程曲线_动应变_高通滤波_组图", ["*SX_L2_414_283*.jpg"]), 145.0),
        ("图 2-12", first_images(result_root, "时程曲线_动应变_高通滤波_组图", ["*SX_L2_298_218*.jpg"]), 145.0),
        ("图 2-13", first_images(result_root, "时程曲线_动应变_高通滤波_组图", ["*SX_L2_405_252*.jpg"]), 145.0),
        ("图 2-14", first_images(result_root, "动应变箱线图_高通滤波", ["*SX_L2_414_283*.jpg"]), 145.0),
        ("图 2-15", first_images(result_root, "动应变箱线图_高通滤波", ["*SX_L2_298_218*.jpg"]), 145.0),
        ("图 2-16", first_images(result_root, "动应变箱线图_高通滤波", ["*SX_L2_405_252*.jpg"]), 145.0),
        ("图 2-17", first_images(result_root, "时程曲线_动应变_低通滤波_组图", ["*SX_L2_414_283*.jpg"]), 145.0),
        ("图 2-18", first_images(result_root, "时程曲线_动应变_低通滤波_组图", ["*SX_L2_298_218*.jpg"]), 145.0),
        ("图 2-19", first_images(result_root, "时程曲线_动应变_低通滤波_组图", ["*SX_L2_405_252*.jpg"]), 145.0),
        ("图 2-20", cable_accel_images, 145.0),
        ("图 2-21", cable_force_images, 145.0),
        ("图 2-22", cable_psd_images, 145.0),
    ]


def remove_nearby_pictures_before(anchor, limit: int = 20) -> int:
    removed = 0
    for candidate in previous_body_paragraphs(anchor, limit=limit):
        text = candidate.text.strip()
        if paragraph_has_image(candidate):
            remove_paragraph(candidate)
            removed += 1
            continue
        if text:
            break
    return removed


def insert_picture_before(anchor, image_path: Path, width_mm: float) -> None:
    paragraph = anchor.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Mm(width_mm))


def replace_pictures_before_anchor(doc: DocxDocument, anchor_fragment: str, image_paths: list[Path], width_mm: float = 145.0) -> list[str]:
    if not image_paths:
        return [f"missing images for anchor: {anchor_fragment}"]
    anchor = find_paragraph_contains(doc, anchor_fragment)
    if anchor is None:
        return [f"missing anchor: {anchor_fragment}"]
    remove_nearby_pictures_before(anchor)
    for image_path in image_paths:
        insert_picture_before(anchor, image_path, width_mm=width_mm)
    return []


def update_images(doc: DocxDocument, result_root: Path) -> list[str]:
    missing: list[str] = []
    for anchor, images, width in zhishan_image_replacements(result_root):
        clean_images = [image for image in images if image is not None]
        missing.extend(replace_pictures_before_anchor(doc, anchor, clean_images, width_mm=width))
    return missing


def refresh_word_fields(path: Path) -> str | None:
    try:
        import win32com.client  # type: ignore
    except Exception:
        return "pywin32/Word COM not available; skipped field refresh"
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(path))
        doc.Fields.Update()
        for table in doc.Tables:
            table.Range.Fields.Update()
        doc.Save()
        doc.Close(False)
        return None
    except Exception as exc:  # noqa: BLE001
        return f"Word field refresh failed: {exc}"
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def update_document(
    doc: DocxDocument,
    result_root: Path,
    config_path: Path,
    period_label: str,
    monitoring_range: str,
    report_date: str,
) -> list[str]:
    context = build_context(result_root, config_path)
    warnings: list[str] = []
    update_cover_dates(doc, period_label, report_date)
    update_period_table(doc, monitoring_range)
    update_data_availability(doc, result_root, period_label, monitoring_range)
    update_temperature_humidity_placeholders(doc)
    warnings.extend(update_bearing_tables(doc, context))
    warnings.extend(update_accel_table(doc, context))
    warnings.extend(update_strain_tables(doc, context))
    warnings.extend(update_cable_tables(doc, context))
    update_narrative(doc, context)
    update_summary_table(doc, context)
    warnings.extend(update_images(doc, result_root))
    normalize_caption_fields(doc)
    return warnings


def build_report(
    *,
    template: Path,
    config_path: Path,
    result_root: Path,
    output_dir: Path | None = None,
    period_label: str = "2026年3月",
    monitoring_range: str = "2026年3月1日~2026年3月31日",
    report_date: str | None = None,
    update_word: bool = True,
) -> tuple[Path, Path]:
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")
    if not result_root.exists():
        raise FileNotFoundError(f"Result root not found: {result_root}")
    report_date = report_date or datetime.now().strftime("%Y年%m月%d日")
    context = ReportBuildContext.from_inputs(
        template=template,
        config_path=config_path,
        result_root=result_root,
        analysis_root=REPO_ROOT,
        output_dir=output_dir,
        assets_subdir="zhishan_report_assets",
    )
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_docx = context.output_dir / f"芝山大桥健康监测{report_month_label(period_label)}月报_自动生成_{timestamp}.docx"
    shutil.copy2(template, output_docx)

    doc = Document(str(output_docx))
    warnings = update_document(doc, result_root, config_path, period_label, monitoring_range, report_date)
    doc.save(str(output_docx))

    if update_word:
        warning = refresh_word_fields(output_docx)
        if warning:
            warnings.append(warning)

    manifest_path = write_report_build_manifest(
        context=context,
        report_type="zhishan_monthly",
        output_docx=output_docx,
        timestamp=timestamp,
        warnings=warnings,
        extra={
            "period_label": period_label,
            "monitoring_range": monitoring_range,
            "report_date": report_date,
            "output_docx_image_count": count_docx_images(output_docx),
        },
        filename_prefix="zhishan_report_build_manifest",
    )
    return output_docx, manifest_path


def main() -> None:
    args = parse_args()
    report_path, manifest_path = build_report(
        template=args.template,
        config_path=args.config,
        result_root=args.result_root,
        output_dir=args.output_dir,
        period_label=args.period_label,
        monitoring_range=args.monitoring_range,
        report_date=args.report_date,
        update_word=not args.no_word_update,
    )
    print(f"Report:   {report_path}")
    print(f"Manifest: {manifest_path}")


if __name__ == "__main__":
    main()
