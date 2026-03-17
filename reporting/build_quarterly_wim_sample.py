from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

T_REPORT_MONTH = "\u5065\u5eb7\u76d1\u6d4b\u6708\u62a5"
T_REPORT_QUARTER = "\u5065\u5eb7\u76d1\u6d4b\u5b63\u62a5\uff08WIM\u6d4b\u8bd5\u7a3f\uff09"
T_RANGE_PREFIX = "\uff08\u76d1\u6d4b\u65f6\u95f4\uff1a"
T_RANGE_SUFFIX = "\uff09"
T_REPORT_DATE_PREFIX = "\u62a5\u544a\u65e5\u671f\uff1a"
T_SUMMARY = "\u76d1\u6d4b\u7ed3\u679c"
T_WIM = "\u4ea4\u901a\u72b6\u51b5\u76d1\u6d4b"
T_NEXT = "\u7ed3\u6784\u5e94\u53d8\u76d1\u6d4b"
T_TABLE4 = "\u8868 4-4 \u8f66\u6d41\u91cf\u7edf\u8ba1\u8868"
T_TABLE5 = "\u8868 4-5 \u524d10\u603b\u91cd\u6700\u91cd\u8f66\u8f86\u53c2\u6570"
T_CONT5 = "\u7eed\u8868 4-5"
T_TABLE6 = "\u8868 4-6 \u524d10\u6700\u5927\u8f74\u91cd\u8f66\u8f86\u53c2\u6570"
T_FIG = "\u56fe 4-3 \u6865\u6881\u4ea4\u901a\u6d41\u53c2\u6570\u5206\u6790"
OUT_SUFFIX = "\u5b63\u62a5\u6837\u7a3f_WIM"

PLOT_LABELS = [
    ("(a) \u4e0d\u540c\u8f66\u9053\u8f66\u8f86\u6570", "WIM_\u4e0d\u540c\u8f66\u9053\u8f66\u8f86\u6570_hongtang_{yyyymm}.png"),
    ("(b) \u4e0d\u540c\u8f66\u901f\u8f66\u8f86\u6570", "WIM_\u4e0d\u540c\u8f66\u901f\u533a\u95f4\u8f66\u8f86\u6570_hongtang_{yyyymm}.png"),
    ("(c) \u4e0d\u540c\u91cd\u91cf\u8f66\u8f86\u6570", "WIM_\u4e0d\u540c\u8f66\u91cd\u533a\u95f4\u8f66\u8f86\u6570_hongtang_{yyyymm}.png"),
    ("(d) \u4e0d\u540c\u65f6\u95f4\u6bb5\u8f66\u8f86\u603b\u6570", "WIM_\u4e0d\u540c\u5c0f\u65f6\u533a\u95f4\u8f66\u8f86\u6570_hongtang_{yyyymm}.png"),
    ("(e) \u4e0d\u540c\u65f6\u95f4\u6bb5\u5e73\u5747\u8f66\u901f", "WIM_\u4e0d\u540c\u5c0f\u65f6\u533a\u95f4\u5e73\u5747\u8f66\u901f_hongtang_{yyyymm}.png"),
    ("(f) \u5927\u4e8e50t\u8f66\u8f86\u65f6\u95f4\u5206\u5e03", "WIM_\u5927\u4e8e50t\u8f66\u8f86\u65f6\u95f4\u5206\u5e03_hongtang_{yyyymm}.png"),
]


@dataclass
class MonthWimSummary:
    yyyymm: str
    total_count: int
    up_count: int
    down_count: int
    days_in_month: int
    max_gross_t: float
    max_axle_t: float
    total_over_count: int
    axle_over_count: int
    daily_rows: list[dict]
    top_gross_rows: list[dict]
    top_axle_rows: list[dict]
    plot_paths: list[tuple[str, Path]]


@dataclass
class ParagraphTemplate:
    style_name: str | None
    alignment: WD_ALIGN_PARAGRAPH | None
    left_indent: object
    right_indent: object
    first_line_indent: object
    space_before: object
    space_after: object
    line_spacing: object
    line_spacing_rule: object
    keep_together: object
    keep_with_next: object
    page_break_before: object
    widow_control: object
    run_bold: bool | None
    run_italic: bool | None
    run_underline: object
    run_font_name: str | None
    run_font_size: object


def parse_args() -> argparse.Namespace:
    repo_root = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser(description="Build quarterly WIM report section from monthly template.")
    parser.add_argument("--template", type=Path, default=repo_root / "reports" / "\u6d2a\u5858\u5927\u6865\u5065\u5eb7\u76d1\u6d4b2025\u5e7412\u6708\u4efd\u6708\u62a5 - \u65b0\u6a21\u677f2.docx")
    parser.add_argument("--wim-root", type=Path, default=None)
    parser.add_argument("--months", nargs="+", default=["202601", "202602", "202603"])
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--period-label", default="2026\u5e74\u7b2c\u4e00\u5b63\u5ea6")
    parser.add_argument("--report-date", default=datetime.now().strftime("%Y\u5e74%m\u6708%d\u65e5"))
    return parser.parse_args()


def load_sheet_rows(path: Path, sheet: str) -> list[dict]:
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb[sheet]
    rows = list(ws.iter_rows(values_only=True))
    wb.close()
    if not rows:
        return []
    header = [str(v) if v is not None else "" for v in rows[0]]
    out: list[dict] = []
    for row in rows[1:]:
        item = {}
        for k, v in zip(header, row):
            item[k] = v
        out.append(item)
    return out


def find_plot_paths(month_dir: Path, yyyymm: str) -> list[tuple[str, Path]]:
    plot_dir = month_dir / "plots"
    items: list[tuple[str, Path]] = []
    for label, pattern in PLOT_LABELS:
        path = plot_dir / pattern.format(yyyymm=yyyymm)
        if path.exists():
            items.append((label, path))
    return items


def parse_month_summary(root: Path, yyyymm: str) -> MonthWimSummary:
    month_dir = root / yyyymm
    excel = month_dir / f"WIM_Report_hongtang_{yyyymm}.xlsx"
    if not excel.exists():
        raise FileNotFoundError(f"WIM workbook not found: {excel}")

    daily = load_sheet_rows(excel, "DailyTraffic")
    top_gross = load_sheet_rows(excel, "TopN")
    top_axle = load_sheet_rows(excel, "TopN_MaxAxle")
    overload = load_sheet_rows(excel, "Overload_Summary")

    up_count = int(sum(int(row.get("up_cnt") or 0) for row in daily))
    down_count = int(sum(int(row.get("down_cnt") or 0) for row in daily))
    total_count = int(sum(int(row.get("total") or 0) for row in daily))

    max_gross_t = 0.0
    for row in top_gross:
        gross = float(row.get("gross_kg") or 0)
        max_gross_t = max(max_gross_t, gross / 1000.0)

    max_axle_t = 0.0
    axle_cols = [f"axle{i}" for i in range(1, 9)]
    for row in top_axle:
        for col in axle_cols:
            axle = float(row.get(col) or 0)
            max_axle_t = max(max_axle_t, axle / 1000.0)

    total_over_count = 0
    axle_over_count = 0
    total_thresholds = []
    axle_thresholds = []
    for row in overload:
        row_type = str(row.get("type") or "")
        count = int(row.get("count") or 0)
        threshold = float(row.get("threshold_kg") or 0)
        if row_type == "total":
            total_thresholds.append((threshold, count))
        elif row_type == "axle":
            axle_thresholds.append((threshold, count))
    if total_thresholds:
        total_over_count = max(total_thresholds, key=lambda x: x[0])[1]
    if axle_thresholds:
        axle_over_count = max(axle_thresholds, key=lambda x: x[0])[1]

    return MonthWimSummary(
        yyyymm=yyyymm,
        total_count=total_count,
        up_count=up_count,
        down_count=down_count,
        days_in_month=len(daily),
        max_gross_t=max_gross_t,
        max_axle_t=max_axle_t,
        total_over_count=total_over_count,
        axle_over_count=axle_over_count,
        daily_rows=daily,
        top_gross_rows=top_gross[:10],
        top_axle_rows=top_axle[:10],
        plot_paths=find_plot_paths(month_dir, yyyymm),
    )


def find_last_paragraph(doc: Document, text: str) -> Paragraph:
    matches = [para for para in doc.paragraphs if para.text.strip() == text]
    if not matches:
        raise ValueError(f"Paragraph not found: {text}")
    return matches[-1]


def find_exact_paragraphs(doc: Document, text: str) -> list[Paragraph]:
    return [para for para in doc.paragraphs if para.text.strip() == text]


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_table_before(paragraph: Paragraph, rows: int, cols: int) -> Table:
    body = paragraph._parent
    table = body.add_table(rows=rows, cols=cols, width=Mm(160))
    paragraph._p.addprevious(table._tbl)
    return table


def clear_section_between(start_paragraph: Paragraph, end_paragraph: Paragraph) -> None:
    parent = start_paragraph._p.getparent()
    current = start_paragraph._p.getnext()
    end = end_paragraph._p
    while current is not None and current is not end:
        nxt = current.getnext()
        parent.remove(current)
        current = nxt


def capture_paragraph_template(paragraph: Paragraph) -> ParagraphTemplate:
    pf = paragraph.paragraph_format
    run0 = paragraph.runs[0] if paragraph.runs else None
    return ParagraphTemplate(
        style_name=paragraph.style.name if paragraph.style is not None else None,
        alignment=paragraph.alignment,
        left_indent=pf.left_indent,
        right_indent=pf.right_indent,
        first_line_indent=pf.first_line_indent,
        space_before=pf.space_before,
        space_after=pf.space_after,
        line_spacing=pf.line_spacing,
        line_spacing_rule=pf.line_spacing_rule,
        keep_together=pf.keep_together,
        keep_with_next=pf.keep_with_next,
        page_break_before=pf.page_break_before,
        widow_control=pf.widow_control,
        run_bold=run0.bold if run0 else None,
        run_italic=run0.italic if run0 else None,
        run_underline=run0.underline if run0 else None,
        run_font_name=run0.font.name if run0 else None,
        run_font_size=run0.font.size if run0 else None,
    )


def apply_paragraph_template(paragraph: Paragraph, template: ParagraphTemplate) -> None:
    if template.style_name:
        paragraph.style = template.style_name
    paragraph.alignment = template.alignment
    pf = paragraph.paragraph_format
    pf.left_indent = template.left_indent
    pf.right_indent = template.right_indent
    pf.first_line_indent = template.first_line_indent
    pf.space_before = template.space_before
    pf.space_after = template.space_after
    pf.line_spacing = template.line_spacing
    pf.line_spacing_rule = template.line_spacing_rule
    pf.keep_together = template.keep_together
    pf.keep_with_next = template.keep_with_next
    pf.page_break_before = template.page_break_before
    pf.widow_control = template.widow_control
    for run in paragraph.runs:
        if template.run_bold is not None:
            run.bold = template.run_bold
        if template.run_italic is not None:
            run.italic = template.run_italic
        if template.run_underline is not None:
            run.underline = template.run_underline
        if template.run_font_name:
            run.font.name = template.run_font_name
        if template.run_font_size:
            run.font.size = template.run_font_size


def add_text_paragraph_before(anchor: Paragraph, text: str, template: ParagraphTemplate) -> Paragraph:
    para = insert_paragraph_before(anchor)
    para.add_run(text)
    apply_paragraph_template(para, template)
    return para


def style_table(table: Table, left: bool = False) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.CENTER


def remove_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


def month_label(yyyymm: str) -> str:
    return f"{yyyymm[:4]}\u5e74{int(yyyymm[4:]):d}\u6708"


def max_or_zero(values: Iterable[float]) -> float:
    vals = list(values)
    return max(vals) if vals else 0.0


def make_month_narrative(item: MonthWimSummary) -> str:
    start_text = f"{item.yyyymm[:4]}\u5e74{int(item.yyyymm[4:]):d}\u67081\u65e5"
    end_text = f"{item.yyyymm[:4]}\u5e74{int(item.yyyymm[4:]):d}\u6708{item.days_in_month}\u65e5"
    daily_avg = round(item.total_count / item.days_in_month) if item.days_in_month else 0
    total_threshold = 110.0
    axle_threshold = 42.0
    gross_text = f"\u8d85\u8fc72.0\u500d\u8bbe\u8ba1\u8f66\u8f86\u8377\u8f7d{total_threshold:.0f}t" if item.max_gross_t >= total_threshold else f"\u672a\u8fbe\u52302.0\u500d\u8bbe\u8ba1\u8f66\u8f86\u8377\u8f7d{total_threshold:.0f}t"
    axle_text = f"\u8fbe\u52301.5\u500d\u8bbe\u8ba1\u8f66\u8f86\u8377\u8f7d{axle_threshold:.0f}t" if item.max_axle_t >= axle_threshold else f"\u672a\u8fbe\u52301.5\u500d\u8bbe\u8ba1\u8f66\u8f86\u8377\u8f7d{axle_threshold:.0f}t"
    return (
        f"{start_text}\u81f3{end_text}\uff0c\u6865\u6881\u5171\u901a\u8fc7\u8f66\u8f86{item.total_count}\u8f86\uff0c\u65e5\u5747{daily_avg}\u8f86\u3002"
        f"\u5176\u4e2d\u4e0a\u884c\u65b9\u5411\uff08\u95fd\u4faf-\u519c\u5927\uff0c\u8f66\u90531\uff5e\u8f66\u90534\uff09\u6240\u901a\u8fc7\u8f66\u8f86\u4e3a{item.up_count}\u8f86\uff0c"
        f"\u4e0b\u884c\u65b9\u5411\uff08\u519c\u5927-\u95fd\u4faf\uff0c\u8f66\u90535\uff5e\u8f66\u90538\uff09\u6240\u901a\u8fc7\u8f66\u8f86\u4e3a{item.down_count}\u8f86\u3002"
        f"\u671f\u95f4\u7cfb\u7edf\u8bb0\u5f55\u5230\u7684\u6700\u5927\u8f66\u91cd\u4e3a{item.max_gross_t:.2f}t\uff0c{gross_text}\u3002"
        f"\u6700\u5927\u8f74\u91cd{item.max_axle_t:.2f}t\uff0c{axle_text}\u3002"
    )


def make_quarter_summary(summaries: list[MonthWimSummary]) -> str:
    total = sum(x.total_count for x in summaries)
    up = sum(x.up_count for x in summaries)
    down = sum(x.down_count for x in summaries)
    total_days = sum(x.days_in_month for x in summaries)
    daily_avg = round(total / total_days) if total_days else 0
    max_gross = max_or_zero(x.max_gross_t for x in summaries)
    max_axle = max_or_zero(x.max_axle_t for x in summaries)
    total_threshold = 110.0
    axle_threshold = 42.0
    gross_text = f"\u8d85\u8fc72.0\u500d\u8bbe\u8ba1\u8f66\u8f86\u8377\u8f7d{total_threshold:.0f}t" if max_gross >= total_threshold else f"\u672a\u8fbe\u52302.0\u500d\u8bbe\u8ba1\u8f66\u8f86\u8377\u8f7d{total_threshold:.0f}t"
    axle_text = f"\u8fbe\u52301.5\u500d\u8bbe\u8ba1\u8f66\u8f86\u8377\u8f7d{axle_threshold:.0f}t" if max_axle >= axle_threshold else f"\u672a\u8fbe\u52301.5\u500d\u8bbe\u8ba1\u8f66\u8f86\u8377\u8f7d{axle_threshold:.0f}t"
    return (
        f"\u76d1\u6d4b\u7ed3\u679c\u8868\u660e\uff0c\u6865\u6881\u5171\u901a\u8fc7\u8f66\u8f86{total}\u8f86\uff0c\u65e5\u5747{daily_avg}\u8f86\u3002"
        f"\u5176\u4e2d\u4e0a\u884c\u65b9\u5411\uff08\u95fd\u4faf-\u519c\u5927\uff0c\u8f66\u90531\uff5e\u8f66\u90534\uff09\u6240\u901a\u8fc7\u8f66\u8f86\u4e3a{up}\u8f86\uff0c"
        f"\u4e0b\u884c\u65b9\u5411\uff08\u519c\u5927-\u95fd\u4faf\uff0c\u8f66\u90535\uff5e\u8f66\u90538\uff09\u6240\u901a\u8fc7\u8f66\u8f86\u4e3a{down}\u8f86\u3002"
        f"\u671f\u95f4\u7cfb\u7edf\u8bb0\u5f55\u5230\u7684\u6700\u5927\u8f66\u91cd\u4e3a{max_gross:.2f}t\uff0c{gross_text}\u3002"
        f"\u6700\u5927\u8f74\u91cd{max_axle:.2f}t\uff0c{axle_text}\u3002"
    )


def set_summary_table(doc: Document, summary_text: str) -> None:
    for table in doc.tables:
        if len(table.rows) >= 4 and table.rows[3].cells and table.rows[3].cells[0].text.strip() == T_SUMMARY:
            cell = table.rows[3].cells[min(2, len(table.rows[3].cells) - 1)]
            if len(cell.paragraphs) >= 9:
                replace_paragraph_text(cell.paragraphs[7], "4\u3001\u4ea4\u901a\u72b6\u51b5\u76d1\u6d4b")
                replace_paragraph_text(cell.paragraphs[8], summary_text)
                return
    raise ValueError("Summary result cell not found")


def add_quarter_overview(anchor: Paragraph, summaries: list[MonthWimSummary], caption_tpl: ParagraphTemplate) -> None:
    add_text_paragraph_before(anchor, "\u8868 4-4 \u5b63\u5ea6\u4ea4\u901a\u72b6\u51b5\u5206\u6708\u7edf\u8ba1\u8868", caption_tpl)
    table = insert_table_before(anchor, rows=len(summaries) + 2, cols=8)
    style_table(table)
    headers = ["\u6708\u4efd", "\u603b\u8f66\u6d41\u91cf", "\u4e0a\u884c\u8f66\u6570", "\u4e0b\u884c\u8f66\u6570", "\u65e5\u5747\u8f66\u6d41\u91cf", "\u6700\u5927\u8f66\u91cd(t)", "\u6700\u5927\u8f74\u91cd(t)", "\u603b\u91cd/\u8f74\u91cd\u8d85\u9650\u8f66\u6b21"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for r, item in enumerate(summaries, start=1):
        daily_avg = round(item.total_count / item.days_in_month) if item.days_in_month else 0
        values = [
            month_label(item.yyyymm),
            str(item.total_count),
            str(item.up_count),
            str(item.down_count),
            str(daily_avg),
            f"{item.max_gross_t:.2f}",
            f"{item.max_axle_t:.2f}",
            f"{item.total_over_count}/{item.axle_over_count}",
        ]
        for c, value in enumerate(values):
            table.cell(r, c).text = value
    totals = [
        "\u5408\u8ba1",
        str(sum(item.total_count for item in summaries)),
        str(sum(item.up_count for item in summaries)),
        str(sum(item.down_count for item in summaries)),
        str(round(sum(item.total_count for item in summaries) / max(1, sum(item.days_in_month for item in summaries)))),
        f"{max_or_zero(item.max_gross_t for item in summaries):.2f}",
        f"{max_or_zero(item.max_axle_t for item in summaries):.2f}",
        f"{sum(item.total_over_count for item in summaries)}/{sum(item.axle_over_count for item in summaries)}",
    ]
    for c, value in enumerate(totals):
        table.cell(len(summaries) + 1, c).text = value


def add_daily_traffic_table(anchor: Paragraph, item: MonthWimSummary, caption_tpl: ParagraphTemplate) -> None:
    add_text_paragraph_before(anchor, f"\u8868 4-4-{int(item.yyyymm[4:])} {month_label(item.yyyymm)}\u8f66\u6d41\u91cf\u7edf\u8ba1\u8868", caption_tpl)
    table = insert_table_before(anchor, rows=len(item.daily_rows) + 1, cols=4)
    style_table(table)
    headers = ["\u65e5\u671f", "\u4e0a\u884c\u8f66\u6570", "\u4e0b\u884c\u8f66\u6570", "\u603b\u8f66\u6570"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for r, row in enumerate(item.daily_rows, start=1):
        date_val = row.get("date")
        date_text = date_val.strftime("%Y-%m-%d") if hasattr(date_val, "strftime") else str(date_val or "")
        values = [date_text, str(int(row.get("up_cnt") or 0)), str(int(row.get("down_cnt") or 0)), str(int(row.get("total") or 0))]
        for c, value in enumerate(values):
            table.cell(r, c).text = value


def add_topn_main_table(anchor: Paragraph, title: str, rows: list[dict], caption_tpl: ParagraphTemplate) -> None:
    add_text_paragraph_before(anchor, title, caption_tpl)
    table = insert_table_before(anchor, rows=len(rows) + 1, cols=6)
    style_table(table)
    headers = ["\u5e8f\u53f7", "\u8f66\u9053", "\u65f6\u95f4", "\u8f74\u6570", "\u603b\u91cd\uff08kg\uff09", "\u901f\u5ea6\uff08km/h\uff09"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for r, row in enumerate(rows, start=1):
        values = [
            str(row.get("rank") or r),
            str(row.get("lane") or ""),
            str(row.get("time") or ""),
            str(row.get("axle_num") or ""),
            str(int(row.get("gross_kg") or 0)),
            str(int(row.get("speed_kmh") or 0)),
        ]
        for c, value in enumerate(values):
            table.cell(r, c).text = value


def add_topn_cont_table(anchor: Paragraph, title: str, rows: list[dict], caption_tpl: ParagraphTemplate) -> None:
    add_text_paragraph_before(anchor, title, caption_tpl)
    table = insert_table_before(anchor, rows=len(rows) + 1, cols=12)
    style_table(table)
    headers = [
        "\u5e8f\u53f7", "\u8f741\u91cd\uff08kg\uff09", "\u8f742\u91cd\uff08kg\uff09", "\u8f743\u91cd\uff08kg\uff09", "\u8f744\u91cd\uff08kg\uff09", "\u8f745\u91cd\uff08kg\uff09", "\u8f746\u91cd\uff08kg\uff09",
        "\u8f74\u8ddd1\uff08m\uff09", "\u8f74\u8ddd2\uff08m\uff09", "\u8f74\u8ddd3\uff08m\uff09", "\u8f74\u8ddd4\uff08m\uff09", "\u8f74\u8ddd5\uff08m\uff09",
    ]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for r, row in enumerate(rows, start=1):
        values = [str(row.get("rank") or r)]
        for idx in range(1, 7):
            values.append(str(int(row.get(f"axle{idx}") or 0)))
        for idx in range(1, 6):
            dist = float(row.get(f"axledis{idx}") or 0) / 1000.0
            values.append(f"{dist:.3f}" if dist else "0")
        for c, value in enumerate(values):
            table.cell(r, c).text = value


def add_plot_grid(anchor: Paragraph, item: MonthWimSummary, fig_tpl: ParagraphTemplate, subcap_tpl: ParagraphTemplate) -> None:
    if not item.plot_paths:
        return
    add_text_paragraph_before(anchor, f"\u56fe 4-3-{int(item.yyyymm[4:])} {month_label(item.yyyymm)}\u6865\u6881\u4ea4\u901a\u6d41\u53c2\u6570\u5206\u6790", fig_tpl)
    table = insert_table_before(anchor, rows=(len(item.plot_paths) + 1) // 2, cols=2)
    remove_table_borders(table)
    for r in range(len(table.rows)):
        for c in range(2):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.width = Mm(80)
    for idx, (label, path) in enumerate(item.plot_paths):
        cell = table.cell(idx // 2, idx % 2)
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(path), width=Mm(78))
        p_cap = cell.add_paragraph()
        p_cap.add_run(label)
        apply_paragraph_template(p_cap, subcap_tpl)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_month_block(anchor: Paragraph, item: MonthWimSummary, heading_tpl: ParagraphTemplate, body_tpl: ParagraphTemplate, caption_tpl: ParagraphTemplate, fig_tpl: ParagraphTemplate, subcap_tpl: ParagraphTemplate) -> None:
    add_text_paragraph_before(anchor, f"4.4.{int(item.yyyymm[4:]):d} {month_label(item.yyyymm)}\u4ea4\u901a\u72b6\u51b5\u76d1\u6d4b", heading_tpl)
    add_text_paragraph_before(anchor, make_month_narrative(item), body_tpl)
    add_daily_traffic_table(anchor, item, caption_tpl)
    add_topn_main_table(anchor, f"\u8868 4-5-{int(item.yyyymm[4:])} {month_label(item.yyyymm)}\u524d10\u603b\u91cd\u6700\u91cd\u8f66\u8f86\u53c2\u6570", item.top_gross_rows, caption_tpl)
    add_topn_cont_table(anchor, f"\u7eed\u8868 4-5-{int(item.yyyymm[4:])}", item.top_gross_rows, caption_tpl)
    add_topn_main_table(anchor, f"\u8868 4-6-{int(item.yyyymm[4:])} {month_label(item.yyyymm)}\u524d10\u6700\u5927\u8f74\u91cd\u8f66\u8f86\u53c2\u6570", item.top_axle_rows, caption_tpl)
    add_topn_cont_table(anchor, f"\u7eed\u8868 4-6-{int(item.yyyymm[4:])}", item.top_axle_rows, caption_tpl)
    add_plot_grid(anchor, item, fig_tpl, subcap_tpl)


def update_cover_and_metadata(doc: Document, period_label: str, report_date: str, months: list[str]) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == T_REPORT_MONTH:
            replace_paragraph_text(para, T_REPORT_QUARTER)
        elif text.startswith(T_RANGE_PREFIX) and text.endswith(T_RANGE_SUFFIX):
            replace_paragraph_text(para, f"{T_RANGE_PREFIX}{period_label}{T_RANGE_SUFFIX}")
        elif text.startswith(T_REPORT_DATE_PREFIX):
            replace_paragraph_text(para, f"{T_REPORT_DATE_PREFIX}{report_date}")
    if len(doc.tables) >= 2 and len(doc.tables[1].rows) >= 2 and len(doc.tables[1].rows[1].cells) >= 5:
        start_label = f"{months[0][:4]}.{months[0][4:]}.01"
        end_month = months[-1]
        days = month_days(end_month)
        end_label = f"{end_month[:4]}.{end_month[4:]}.{days:02d}"
        doc.tables[1].rows[1].cells[4].text = f"{start_label}\uff5e{end_label}"


def month_days(yyyymm: str) -> int:
    year = int(yyyymm[:4])
    month = int(yyyymm[4:])
    if month == 12:
        next_month = datetime(year + 1, 1, 1)
    else:
        next_month = datetime(year, month + 1, 1)
    return (next_month - datetime(year, month, 1)).days


def build_quarterly_wim_sample(template: Path, wim_root: Path, months: list[str], output_dir: Path, period_label: str, report_date: str) -> Path:
    summaries = [parse_month_summary(wim_root, yyyymm) for yyyymm in months]
    doc = Document(str(template))
    update_cover_and_metadata(doc, period_label, report_date, months)
    set_summary_table(doc, make_quarter_summary(summaries))

    wim_heading = find_last_paragraph(doc, T_WIM)
    next_heading = find_last_paragraph(doc, T_NEXT)
    table4_caption = find_last_paragraph(doc, T_TABLE4)
    cont5_caption = find_last_paragraph(doc, T_CONT5)
    fig_caption = find_last_paragraph(doc, T_FIG)

    heading_tpl = capture_paragraph_template(wim_heading)
    body_tpl = capture_paragraph_template(doc.tables[1].rows[3].cells[2].paragraphs[8])
    caption_tpl = capture_paragraph_template(table4_caption)
    subcap_tpl = capture_paragraph_template(cont5_caption)
    fig_tpl = capture_paragraph_template(fig_caption)

    clear_section_between(wim_heading, next_heading)
    add_text_paragraph_before(next_heading, make_quarter_summary(summaries), body_tpl)
    add_quarter_overview(next_heading, summaries, caption_tpl)
    for item in summaries:
        add_month_block(next_heading, item, heading_tpl, body_tpl, caption_tpl, fig_tpl, subcap_tpl)

    output_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"{template.stem}_{OUT_SUFFIX}_{ts}.docx"
    doc.save(output_path)
    return output_path


def main() -> None:
    args = parse_args()
    wim_root = args.wim_root or (Path.cwd() / "WIM" / "results" / "hongtang")
    output_dir = args.output_dir or (wim_root.parents[2] / "自动报告" if len(wim_root.parents) >= 3 else (Path.cwd() / "自动报告"))
    output = build_quarterly_wim_sample(
        template=args.template,
        wim_root=wim_root,
        months=args.months,
        output_dir=output_dir,
        period_label=args.period_label,
        report_date=args.report_date,
    )
    print(f"Quarterly WIM sample report generated: {output}")


if __name__ == "__main__":
    main()
