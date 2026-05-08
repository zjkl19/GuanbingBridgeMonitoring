from __future__ import annotations

from collections.abc import Iterable, Sequence

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import Table


DXA_PER_MM = 56.6929


def iter_unique_cells(table: Table):
    """Yield each physical cell once, even when Word exposes merged cells repeatedly."""
    seen: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            yield cell


def style_table(
    table: Table,
    *,
    left: bool = False,
    autofit: bool = True,
    align_center: bool = True,
    style: str = "Table Grid",
) -> None:
    table.style = style
    if align_center:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_autofit(table, autofit)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.CENTER


def set_table_autofit(table: Table, enabled: bool = True) -> None:
    table.autofit = enabled
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit" if enabled else "fixed")


def set_table_width(table: Table, width_mm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(round(width_mm * DXA_PER_MM)))


def set_table_auto_width(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "auto")
    tbl_w.set(qn("w:w"), "0")


def set_table_column_widths(table: Table, widths_mm: Sequence[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_mm):
            if idx < len(row.cells):
                row.cells[idx].width = Mm(width)


def set_header_bold(table: Table, header_rows: int = 1) -> None:
    for row in table.rows[:header_rows]:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True


def set_table_borders(
    table: Table,
    *,
    outer_size_eighth_pt: int = 12,
    inside_size_eighth_pt: int | None = None,
    color: str = "000000",
) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    inside_size = outer_size_eighth_pt if inside_size_eighth_pt is None else inside_size_eighth_pt
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        size = outer_size_eighth_pt if edge in {"top", "left", "bottom", "right"} else inside_size
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_outer_border(table: Table, size_eighth_pt: int = 12) -> None:
    set_table_borders(table, outer_size_eighth_pt=size_eighth_pt, inside_size_eighth_pt=4)


def remove_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


def set_table_font_size(table: Table, size_pt: int | float) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(size_pt)


def set_cell_line_spacing(
    cell,
    *,
    line_spacing: float = 1.5,
    rule=WD_LINE_SPACING.MULTIPLE,
) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.line_spacing_rule = rule


def set_table_line_spacing(
    table: Table,
    *,
    line_spacing: float = 1.5,
    rule=WD_LINE_SPACING.MULTIPLE,
) -> None:
    for cell in iter_unique_cells(table):
        set_cell_line_spacing(cell, line_spacing=line_spacing, rule=rule)


def set_cell_alignment(cell, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment


def apply_report_table_format(
    table: Table,
    *,
    line_spacing: float | None = None,
    align_center: bool = True,
    vertical_center: bool = True,
) -> None:
    """Apply common Word table formatting without changing table contents."""
    for cell in iter_unique_cells(table):
        if vertical_center:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if align_center:
            set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.CENTER)
        if line_spacing is not None:
            set_cell_line_spacing(cell, line_spacing=line_spacing)


def fill_table(
    table: Table,
    rows: Iterable[Sequence[object]],
    *,
    start_row: int = 0,
    start_col: int = 0,
    blank: str = "/",
) -> None:
    for ridx, row_values in enumerate(rows, start=start_row):
        if ridx >= len(table.rows):
            break
        row = table.rows[ridx]
        for cidx, value in enumerate(row_values, start=start_col):
            if cidx >= len(row.cells):
                break
            row.cells[cidx].text = blank if value is None else str(value)
