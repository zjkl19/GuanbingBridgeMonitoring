from __future__ import annotations

import argparse
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

from build_jlj_monthly_report import build_report as build_jlj_monthly_report
from build_period_report import build_period_report as build_hongtang_period_report
from template_precheck import TemplatePrecheckError, check_template, write_precheck_report


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def _default_hongtang_template(repo_root: Path) -> Path:
    candidates = [
        repo_root / "reports" / "洪塘大桥健康监测2026年第一季季报-改4.docx",
        repo_root / "reports" / "洪塘大桥健康监测周期报模板-自动报告.docx",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def _default_jlj_template(repo_root: Path) -> Path:
    return repo_root / "reports" / "九龙江大桥健康监测2026年3月份月报_修订5.docx"


def _docx_contains(path: Path, fragment: str) -> bool:
    doc = Document(str(path))
    for para in doc.paragraphs:
        if fragment in para.text:
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if fragment in cell.text:
                    return True
    return False


def _assert_generated_docx(path: Path, fragments: list[str]) -> None:
    if not path.exists() or path.stat().st_size <= 0:
        raise AssertionError(f"Generated report is missing or empty: {path}")
    for fragment in fragments:
        if not _docx_contains(path, fragment):
            raise AssertionError(f'Generated report does not contain required text "{fragment}": {path}')


def smoke_hongtang(args: argparse.Namespace, output_root: Path) -> None:
    print(f"[hongtang] template: {args.hongtang_template}")
    issues = check_template("hongtang_period", args.hongtang_template)
    txt_path, json_path = write_precheck_report(
        "hongtang_period",
        args.hongtang_template,
        issues,
        output_root / "precheck",
        context={"smoke_generate": args.generate},
    )
    print(f"[hongtang] precheck report: {txt_path}")
    if issues:
        raise TemplatePrecheckError(args.hongtang_template, issues)
    print("[hongtang] template precheck OK")
    if not args.generate:
        return
    if not args.hongtang_result_root.exists():
        raise FileNotFoundError(f"Hongtang result root not found: {args.hongtang_result_root}")
    out_dir = output_root / "hongtang"
    out_dir.mkdir(parents=True, exist_ok=True)
    _, report_path, missing = build_hongtang_period_report(
        template=args.hongtang_template,
        config_path=args.hongtang_config,
        result_root=args.hongtang_result_root,
        wim_root=args.hongtang_wim_root,
        output_dir=out_dir,
        period_label=args.hongtang_period_label,
        monitoring_range=args.hongtang_monitoring_range,
        report_date=args.report_date,
        start_date=args.hongtang_start_date,
        end_date=args.hongtang_end_date,
    )
    _assert_generated_docx(report_path, ["健康监测系统运行状况", "交通状况监测"])
    print(f"[hongtang] generated OK: {report_path}")
    if missing:
        print("[hongtang] warnings/missing:")
        for item in missing:
            print(f"  - {item}")


def smoke_jlj(args: argparse.Namespace, output_root: Path) -> None:
    print(f"[jlj] template: {args.jlj_template}")
    issues = check_template("jlj_monthly", args.jlj_template)
    txt_path, json_path = write_precheck_report(
        "jlj_monthly",
        args.jlj_template,
        issues,
        output_root / "precheck",
        context={"smoke_generate": args.generate},
    )
    print(f"[jlj] precheck report: {txt_path}")
    if issues:
        raise TemplatePrecheckError(args.jlj_template, issues)
    print("[jlj] template precheck OK")
    if not args.generate:
        return
    if not args.jlj_result_root.exists():
        raise FileNotFoundError(f"Jiulongjiang result root not found: {args.jlj_result_root}")
    out_dir = output_root / "jlj"
    out_dir.mkdir(parents=True, exist_ok=True)
    report_path = build_jlj_monthly_report(
        template=args.jlj_template,
        config_path=args.jlj_config,
        result_root=args.jlj_result_root,
        output_dir=out_dir,
        period_label=args.jlj_period_label,
        monitoring_range=args.jlj_monitoring_range,
        report_date=args.report_date,
    )
    _assert_generated_docx(report_path, ["健康监测系统运行状况", "主桥环境与作用监测", "（以下无正文）"])
    print(f"[jlj] generated OK: {report_path}")


def parse_args() -> argparse.Namespace:
    repo_root = _repo_root()
    parser = argparse.ArgumentParser(description="Smoke test Hongtang/Jiulongjiang report templates and generation.")
    parser.add_argument("--kind", choices=["all", "hongtang", "jlj"], default="all")
    parser.add_argument("--generate", action="store_true", help="Also generate reports from existing stats/images.")
    parser.add_argument("--keep-output", action="store_true", help="Keep tmp smoke-test outputs.")
    parser.add_argument("--output-root", type=Path, default=repo_root / "tmp" / "report_smoke")
    parser.add_argument("--report-date", default=datetime.now().strftime("%Y年%m月%d日"))

    parser.add_argument("--hongtang-template", type=Path, default=_default_hongtang_template(repo_root))
    parser.add_argument("--hongtang-config", type=Path, default=repo_root / "config" / "hongtang_config.json")
    parser.add_argument("--hongtang-result-root", type=Path, default=Path(r"E:\洪塘大桥数据\2026年1-3月"))
    parser.add_argument("--hongtang-wim-root", type=Path, default=None)
    parser.add_argument("--hongtang-period-label", default="2026年1-3月")
    parser.add_argument("--hongtang-monitoring-range", default="2026年01月01日~2026年03月31日")
    parser.add_argument("--hongtang-start-date", default="2026-01-01")
    parser.add_argument("--hongtang-end-date", default="2026-03-31")

    parser.add_argument("--jlj-template", type=Path, default=_default_jlj_template(repo_root))
    parser.add_argument("--jlj-config", type=Path, default=repo_root / "config" / "jiulongjiang_config.json")
    parser.add_argument("--jlj-result-root", type=Path, default=Path(r"E:\九龙江数据\2026年3月"))
    parser.add_argument("--jlj-period-label", default="2026年3月份")
    parser.add_argument("--jlj-monitoring-range", default="2026.03.23~2026.03.31")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_root = args.output_root
    if output_root.exists() and not args.keep_output:
        shutil.rmtree(output_root)
    output_root.mkdir(parents=True, exist_ok=True)

    if args.kind in ("all", "hongtang"):
        smoke_hongtang(args, output_root)
    if args.kind in ("all", "jlj"):
        smoke_jlj(args, output_root)

    if not args.generate and not args.keep_output and output_root.exists():
        shutil.rmtree(output_root)
    print("Smoke test OK.")


if __name__ == "__main__":
    main()
