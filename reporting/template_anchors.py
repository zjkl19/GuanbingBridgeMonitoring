from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Iterable

from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class TemplateAnchor:
    key: str
    text: str
    description: str
    required: bool = False


JLJ_SUMMARY_TABLE = TemplateAnchor(
    key="jlj.summary_table",
    text="{{JLJ_SUMMARY_TABLE}}",
    description="九龙江结论页监测结果表起始位置",
)

JLJ_DATA_STATUS_SECTION = TemplateAnchor(
    key="jlj.data_status_section",
    text="{{JLJ_DATA_STATUS_SECTION}}",
    description="九龙江本月监测数据情况章节",
)

JLJ_PATROL_SECTION = TemplateAnchor(
    key="jlj.patrol_section",
    text="{{JLJ_PATROL_SECTION}}",
    description="九龙江人工巡查报告插入位置",
)


JLJ_OPTIONAL_ANCHORS = (
    JLJ_SUMMARY_TABLE,
    JLJ_DATA_STATUS_SECTION,
    JLJ_PATROL_SECTION,
)


def normalize_anchor_text(text: object) -> str:
    return re.sub(r"\s+", "", "" if text is None else str(text))


def text_contains_anchor(text: object, anchor: TemplateAnchor | str) -> bool:
    needle = normalize_anchor_text(anchor.text if isinstance(anchor, TemplateAnchor) else anchor)
    return bool(needle) and needle in normalize_anchor_text(text)


def paragraph_contains_anchor(paragraph: Paragraph, anchor: TemplateAnchor | str) -> bool:
    return text_contains_anchor(paragraph.text, anchor)


def table_contains_anchor(table: Table, anchor: TemplateAnchor | str) -> bool:
    for row in table.rows:
        for cell in row.cells:
            if text_contains_anchor(cell.text, anchor):
                return True
    return False


def iter_body_blocks(doc: DocxDocument):
    """Yield document body paragraphs and tables in Word order."""
    body = doc.element.body
    parent = doc._body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, parent)


def find_paragraph_containing_anchor(doc: DocxDocument, anchor: TemplateAnchor | str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph_contains_anchor(paragraph, anchor):
            return paragraph
    return None


def find_table_containing_anchor(doc: DocxDocument, anchor: TemplateAnchor | str) -> Table | None:
    for table in doc.tables:
        if table_contains_anchor(table, anchor):
            return table
    return None


def find_table_after_anchor(doc: DocxDocument, anchor: TemplateAnchor | str) -> Table | None:
    """Return the first table after an anchor paragraph/table, if the template has one."""
    found = False
    for kind, block in iter_body_blocks(doc):
        if kind == "paragraph" and paragraph_contains_anchor(block, anchor):
            found = True
            continue
        if kind == "table":
            table = block
            if table_contains_anchor(table, anchor):
                return table
            if found:
                return table
    return None


def remove_anchor_paragraphs(doc: DocxDocument, anchors: Iterable[TemplateAnchor | str]) -> int:
    removed = 0
    for paragraph in list(doc.paragraphs):
        if any(paragraph_contains_anchor(paragraph, anchor) for anchor in anchors):
            parent = paragraph._p.getparent()
            if parent is not None:
                parent.remove(paragraph._p)
                removed += 1
    return removed
