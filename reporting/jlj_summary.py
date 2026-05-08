from __future__ import annotations

import re
from copy import deepcopy
from typing import Any, Mapping

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from docx_table_utils import find_first_row_index_by_first_cell, find_summary_table, table_has_first_cell
from docx_utils import set_cell_paragraphs, set_cell_text_preserve


def is_summary_table(table: Table) -> bool:
    return table_has_first_cell(table, "监测结果")


def clear_repeat_table_headers(table: Table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        for child in list(tr_pr):
            if child.tag == qn("w:tblHeader"):
                tr_pr.remove(child)


def allow_table_rows_to_expand_and_split(table: Table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        for child in list(tr_pr):
            if child.tag in {qn("w:trHeight"), qn("w:cantSplit")}:
                tr_pr.remove(child)


def append_table_row_from_template(table: Table, template_tr) -> object:
    new_tr = deepcopy(template_tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def clear_table_rows(table: Table) -> None:
    for tr in list(table._tbl.tr_lst):
        table._tbl.remove(tr)


def remove_table(table: Table) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_stale_summary_tables(doc: Document, keep_table: Table) -> None:
    """Remove old one-row summary fragments left by manually edited templates."""
    keep_element = keep_table._element
    for table in list(doc.tables):
        if table._element is keep_element:
            continue
        if is_summary_table(table) and not table_has_first_cell(table, "建议"):
            remove_table(table)


def summary_line_is_heading(line: str) -> bool:
    text = str(line or "").strip()
    return bool(
        re.match(r"^[一二三四五六七八九十]+、", text)
        or re.match(r"^\d+、", text)
        or re.match(r"^\d+(?:\.\d+)+\s*", text)
        or re.match(r"^（\d+）", text)
    )


def is_continue_marker(line: str) -> bool:
    return str(line or "").strip() in {"（续上页）", "（转下页）"}


def build_summary_result_rows(
    result_lines: list[str],
    *,
    include_continue_marker_in_cell: bool = True,
) -> list[tuple[list[str], set[int]]]:
    lines = [str(line or "") for line in result_lines]
    if not lines:
        return []

    # Word/WPS pagination is renderer-dependent. Split the long summary at
    # stable section boundaries so the generated report can carry the same
    # continuation cues as the manually adjusted template.
    boundary_titles = ["1.4 风向风速监测", "2.5 结构应变监测"]
    boundaries: list[int] = []
    for title in boundary_titles:
        try:
            idx = next(i for i, item in enumerate(lines) if item.strip() == title)
        except StopIteration:
            continue
        if 0 < idx < len(lines):
            boundaries.append(idx)
    boundaries = sorted(set(boundaries))

    starts = [0] + boundaries
    ends = boundaries + [len(lines)]
    rows: list[tuple[list[str], set[int]]] = []
    for chunk_idx, (start, end) in enumerate(zip(starts, ends)):
        chunk = list(lines[start:end])
        if not chunk:
            continue
        if chunk_idx > 0 and include_continue_marker_in_cell:
            chunk.insert(0, "（续上页）")
        if chunk_idx < len(starts) - 1:
            chunk.append("（转下页）")
        bold_indices = {
            idx
            for idx, item in enumerate(chunk)
            if summary_line_is_heading(item) and not is_continue_marker(item)
        }
        rows.append((chunk, bold_indices))
    return rows


def clear_paragraph_numbering(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def set_page_break_before_paragraph(paragraph: Paragraph | None) -> None:
    if paragraph is None:
        return
    paragraph.paragraph_format.page_break_before = True


def previous_paragraph_before_table(table: Table) -> Paragraph | None:
    prev = table._tbl.getprevious()
    while prev is not None:
        if prev.tag == qn("w:p"):
            return Paragraph(prev, table._parent)
        if prev.tag == qn("w:tbl"):
            return None
        prev = prev.getprevious()
    return None


def row_label_content_cells(row) -> tuple[object, object]:
    label_cell = row.cells[0]
    for cell in row.cells[1:]:
        if cell._tc is not label_cell._tc:
            return label_cell, cell
    return label_cell, row.cells[-1]


def clear_cell_numbering(cell) -> None:
    for paragraph in cell.paragraphs:
        clear_paragraph_numbering(paragraph)


def apply_summary_paragraph_format(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        text = paragraph.text.strip()
        if text == "（转下页）":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = False
        elif is_continue_marker(text):
            for run in paragraph.runs:
                run.bold = False


def apply_summary_table_line_spacing(table: Table) -> None:
    seen: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            apply_summary_paragraph_format(cell)


def rebuild_summary_table_rows(
    table: Table,
    result_lines: list[str],
    advice_left: str,
    advice_lines: list[str],
) -> None:
    if not table.rows:
        return
    advice_template_idx = find_first_row_index_by_first_cell(table, "建议")
    result_template_tr = deepcopy(table.rows[0]._tr)
    advice_template_tr = deepcopy(table.rows[advice_template_idx]._tr if advice_template_idx is not None else table.rows[-1]._tr)
    clear_table_rows(table)

    result_rows = build_summary_result_rows(result_lines)
    for idx, (lines, bold_indices) in enumerate(result_rows):
        row = append_table_row_from_template(table, result_template_tr)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_text_preserve(row.cells[0], "监测结果")
        set_cell_paragraphs(row.cells[1], lines, bold_indices=bold_indices)
        apply_summary_paragraph_format(row.cells[0])
        apply_summary_paragraph_format(row.cells[1])
        if idx > 0 and row.cells[1].paragraphs:
            row.cells[1].paragraphs[0].paragraph_format.page_break_before = True
        clear_cell_numbering(row.cells[0])
        clear_cell_numbering(row.cells[1])

    row = append_table_row_from_template(table, advice_template_tr)
    row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_preserve(row.cells[0], advice_left or "建  议")
    set_cell_paragraphs(row.cells[1], advice_lines or ["建议结合监测数据变化情况进行持续跟踪。"])
    apply_summary_paragraph_format(row.cells[0])
    apply_summary_paragraph_format(row.cells[1])
    clear_cell_numbering(row.cells[0])
    clear_cell_numbering(row.cells[1])
    allow_table_rows_to_expand_and_split(table)
    clear_repeat_table_headers(table)


def find_cover_summary_table(doc: Document) -> Table | None:
    for table in doc.tables:
        if len(table.columns) < 2:
            continue
        if not find_first_row_index_by_first_cell(table, "委托单位") == 0:
            continue
        if find_first_row_index_by_first_cell(table, "监测结果") is not None:
            return table
    return None


def following_front_summary_tables(doc: Document, cover_table: Table) -> list[Table]:
    out: list[Table] = []
    seen_cover = False
    for table in doc.tables:
        if table._element is cover_table._element:
            seen_cover = True
            continue
        if not seen_cover:
            continue
        if table_has_first_cell(table, "监测结果"):
            out.append(table)
            continue
        break
    return out


def update_cover_summary_tables(
    doc: Document,
    cover_table: Table,
    result_lines: list[str],
    advice_left: str,
    advice_lines: list[str],
) -> bool:
    summary_row_idx = find_first_row_index_by_first_cell(cover_table, "监测结果")
    if summary_row_idx is None:
        return False

    continuation_tables = following_front_summary_tables(doc, cover_table)
    slots: list[tuple[Table, object]] = [(cover_table, cover_table.rows[summary_row_idx])]
    for table in continuation_tables:
        row_idx = find_first_row_index_by_first_cell(table, "监测结果")
        if row_idx is not None:
            slots.append((table, table.rows[row_idx]))

    if not slots:
        return False

    result_rows = build_summary_result_rows(result_lines, include_continue_marker_in_cell=False)
    if len(result_rows) > len(slots):
        overflow_lines: list[str] = []
        overflow_bold: set[int] = set()
        for lines, bold_indices in result_rows[len(slots) - 1 :]:
            offset = len(overflow_lines)
            overflow_lines.extend(lines)
            overflow_bold.update(offset + idx for idx in bold_indices)
        result_rows = result_rows[: len(slots) - 1] + [(overflow_lines, overflow_bold)]

    for idx, (table, row) in enumerate(slots):
        if idx >= len(result_rows):
            if table is not cover_table:
                remove_table(table)
            continue
        lines, bold_indices = result_rows[idx]
        label_cell, content_cell = row_label_content_cells(row)
        set_cell_text_preserve(label_cell, "监测结果")
        set_cell_paragraphs(content_cell, lines, bold_indices=bold_indices)
        apply_summary_paragraph_format(label_cell)
        apply_summary_paragraph_format(content_cell)
        clear_cell_numbering(label_cell)
        clear_cell_numbering(content_cell)
        if idx > 0:
            set_page_break_before_paragraph(previous_paragraph_before_table(table))

    advice_table = next((table for table in reversed(continuation_tables) if find_first_row_index_by_first_cell(table, "建议") is not None), None)
    if advice_table is not None:
        advice_row_idx = find_first_row_index_by_first_cell(advice_table, "建议")
        if advice_row_idx is not None:
            label_cell, content_cell = row_label_content_cells(advice_table.rows[advice_row_idx])
            set_cell_text_preserve(label_cell, advice_left or "建  议")
            set_cell_paragraphs(content_cell, advice_lines or ["建议结合监测数据变化情况进行持续跟踪。"])
            apply_summary_paragraph_format(label_cell)
            apply_summary_paragraph_format(content_cell)
            clear_cell_numbering(label_cell)
            clear_cell_numbering(content_cell)

    for table in [cover_table, *continuation_tables]:
        apply_summary_table_line_spacing(table)
        allow_table_rows_to_expand_and_split(table)
        clear_repeat_table_headers(table)
    return True


def _summary_text(section_map: Mapping[str, Any], key: str) -> str:
    content = section_map[key]
    return str(getattr(content, "summary_sentence", "") or "")


def build_summary_result_lines(section_map: Mapping[str, Any], data_acquisition_summary: str | None = None) -> list[str]:
    return [
        "一、监测系统运行情况",
        "",
        "二、本月监测数据情况",
        data_acquisition_summary or "本月监测数据获取情况详见正文。",
        "三、监测数据分析结果",
        "1、主桥环境与作用监测",
        "1.1 温度监测",
        _summary_text(section_map, "main_env"),
        "1.2 湿度监测",
        _summary_text(section_map, "main_humidity"),
        "1.3 雨量监测",
        _summary_text(section_map, "main_rainfall"),
        "1.4 风向风速监测",
        _summary_text(section_map, "main_wind"),
        "1.5 地震动监测",
        _summary_text(section_map, "main_eq"),
        "1.6 车辆荷载监测",
        _summary_text(section_map, "main_traffic"),
        "2、主桥结构响应与结构变化监测",
        "2.1 主梁挠度监测",
        _summary_text(section_map, "main_deflection"),
        "2.2 支座、梁段纵向位移监测",
        _summary_text(section_map, "main_bearing"),
        "2.3 拱顶、拱脚位移监测（GNSS）",
        _summary_text(section_map, "main_gnss"),
        "2.4 结构振动监测",
        _summary_text(section_map, "main_vibration"),
        "2.5 结构应变监测",
        _summary_text(section_map, "main_strain"),
        "2.6 裂缝监测",
        _summary_text(section_map, "main_crack"),
        "2.7 吊杆索力监测",
        _summary_text(section_map, "main_cable"),
        "3、北江滨匝道桥监测",
        "3.1 结构应变监测",
        _summary_text(section_map, "north_strain"),
        "3.2 支座位移监测",
        _summary_text(section_map, "north_bearing"),
        "3.3 墩柱倾斜监测",
        _summary_text(section_map, "north_tilt"),
        "4、南江滨匝道桥监测",
        "4.1 结构应变监测",
        _summary_text(section_map, "south_strain"),
        "4.2 支座位移监测",
        _summary_text(section_map, "south_bearing"),
        "4.3 墩柱倾斜监测",
        _summary_text(section_map, "south_tilt"),
    ]


def update_summary_table(doc: Document, section_map: Mapping[str, Any], data_acquisition_summary: str | None = None) -> None:
    result_lines = build_summary_result_lines(section_map, data_acquisition_summary)
    cover_table = find_cover_summary_table(doc)
    if cover_table is not None:
        advice_left = "建  议"
        advice_lines = ["建议结合监测数据变化情况进行持续跟踪。"]
        for table in following_front_summary_tables(doc, cover_table):
            advice_row_idx = find_first_row_index_by_first_cell(table, "建议")
            if advice_row_idx is None:
                continue
            advice_left = table.cell(advice_row_idx, 0).text.strip() or advice_left
            advice_lines = [
                para.text.strip()
                for para in row_label_content_cells(table.rows[advice_row_idx])[1].paragraphs
                if para.text.strip()
            ] or advice_lines
            break
        if update_cover_summary_tables(doc, cover_table, result_lines, advice_left, advice_lines):
            return

    summary_table = find_summary_table(doc)
    if summary_table is None:
        raise ValueError("Jiulongjiang summary table not found: expected left-column labels '监测结果' and '建  议'.")
    remove_stale_summary_tables(doc, summary_table)
    advice_row_idx = find_first_row_index_by_first_cell(summary_table, "建议")
    advice_left = summary_table.cell(advice_row_idx, 0).text.strip() if advice_row_idx is not None else "建  议"
    advice_lines = []
    if advice_row_idx is not None:
        advice_lines = [
            para.text.strip()
            for para in summary_table.cell(advice_row_idx, 1).paragraphs
            if para.text.strip()
        ]
    rebuild_summary_table_rows(
        summary_table,
        result_lines,
        advice_left,
        advice_lines,
    )
