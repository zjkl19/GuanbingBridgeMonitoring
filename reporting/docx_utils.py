from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm
from docx.text.paragraph import Paragraph


def find_paragraph_contains(doc: Document, fragment: str, occurrence: int = 1) -> Paragraph | None:
    seen = 0
    for paragraph in doc.paragraphs:
        if fragment in paragraph.text:
            seen += 1
            if seen == occurrence:
                return paragraph
    return None


def paragraph_has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing") or paragraph._p.xpath(".//w:pict"))


def prune_unused_document_image_relationships(doc: Document) -> list[str]:
    """Drop image relationships no longer referenced by the document body.

    Report templates are frequently reused after their old result paragraphs
    have been removed.  ``python-docx`` does not automatically delete those
    relationships, so the old result media otherwise remains packaged inside
    the new DOCX even though it is no longer visible.
    """
    used_rel_ids = set(doc.element.body.xpath(".//a:blip/@r:embed"))
    used_rel_ids.update(doc.element.body.xpath(".//a:blip/@r:link"))
    relationship_id_attr = (
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
    )
    for element in doc.element.body.iter():
        if str(element.tag).endswith("}imagedata"):
            rel_id = element.get(relationship_id_attr)
            if rel_id:
                used_rel_ids.add(rel_id)

    dropped: list[str] = []
    for rel_id, relationship in list(doc.part.rels.items()):
        if relationship.reltype == RT.IMAGE and rel_id not in used_rel_ids:
            doc.part.drop_rel(rel_id)
            dropped.append(rel_id)
    return dropped


def paragraph_from_element(element, parent) -> Paragraph:
    return Paragraph(element, parent)


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def previous_body_paragraphs(paragraph: Paragraph, limit: int = 8) -> list[Paragraph]:
    out: list[Paragraph] = []
    element = paragraph._p.getprevious()
    while element is not None and len(out) < limit:
        if element.tag == qn("w:p"):
            out.append(paragraph_from_element(element, paragraph._parent))
        element = element.getprevious()
    return out


def remove_nearby_picture_before(anchor: Paragraph, limit: int = 8) -> int:
    removed = 0
    for candidate in previous_body_paragraphs(anchor, limit=limit):
        text = candidate.text.strip()
        if paragraph_has_image(candidate):
            remove_paragraph(candidate)
            removed += 1
            continue
        if text:
            break
    return removed


def _is_short_picture_label(text: str) -> bool:
    value = text.strip()
    if not value or len(value) > 80:
        return False
    if re.match(r"^(图|表|续表|第[一二三四五六七八九十0-9]+[章节条]|[0-9]+(?:\.[0-9]+)+)", value):
        return False
    return True


def remove_nearby_picture_block_before(anchor: Paragraph, limit: int = 120) -> int:
    """Remove a generated picture block immediately before a caption anchor.

    Existing report files are often reused as templates. In that case the old
    report pictures sit directly before their captions; remove that contiguous
    image/short-label block before inserting fresh images.
    """
    candidates = previous_body_paragraphs(anchor, limit=limit)
    removed: list[Paragraph] = []
    in_block = False
    for idx, candidate in enumerate(candidates):
        text = candidate.text.strip()
        has_image = paragraph_has_image(candidate)
        previous_is_image = idx + 1 < len(candidates) and paragraph_has_image(candidates[idx + 1])
        if has_image:
            removed.append(candidate)
            in_block = True
            continue
        if not text and in_block:
            removed.append(candidate)
            continue
        if _is_short_picture_label(text) and previous_is_image:
            removed.append(candidate)
            in_block = True
            continue
        break
    for paragraph in removed:
        remove_paragraph(paragraph)
    return len(removed)


def insert_picture_before(anchor: Paragraph, image_path: Path, width_mm: float = 145.0) -> None:
    paragraph = anchor.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Mm(width_mm))


def replace_picture_before_anchor(
    doc: Document,
    anchor_fragment: str,
    image_path: Path | None,
    occurrence: int = 1,
    width_mm: float = 145.0,
) -> tuple[bool, str]:
    if image_path is None or not image_path.exists():
        return False, f"missing image for anchor: {anchor_fragment}"
    anchor = find_paragraph_contains(doc, anchor_fragment, occurrence=occurrence)
    if anchor is None:
        return False, f"missing anchor: {anchor_fragment}"
    remove_nearby_picture_before(anchor)
    insert_picture_before(anchor, image_path, width_mm=width_mm)
    return True, str(image_path)


def set_cell_text_preserve(cell, text: str) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text
        return
    first = paragraphs[0]
    if first.runs:
        for run in first.runs:
            run.text = ""
        first.runs[0].text = text
    else:
        first.add_run(text)
    for para in paragraphs[1:]:
        for run in para.runs:
            run.text = ""


def set_cell_paragraphs(cell, lines: list[str], bold_indices: set[int] | None = None) -> None:
    bold_indices = bold_indices or set()
    if not cell.paragraphs:
        cell.text = ""
    base_para = cell.paragraphs[0]
    base_style = base_para.style
    base_alignment = base_para.alignment
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)
    if base_para.runs:
        for run in base_para.runs:
            run.text = ""
    else:
        base_para.add_run("")
    paragraphs = [base_para]
    for _ in range(max(0, len(lines) - 1)):
        para = cell.add_paragraph()
        para.style = base_style
        para.alignment = base_alignment
        paragraphs.append(para)
    for idx, (para, text) in enumerate(zip(paragraphs, lines)):
        if not para.runs:
            para.add_run("")
        for run in para.runs:
            run.text = ""
            run.bold = False
        para.runs[0].text = text
        para.runs[0].bold = idx in bold_indices
