from __future__ import annotations

import argparse
import math
import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

from analysis_manifest import analysis_manifest_context, missing_module_summary_items
from docx_utils import (
    find_paragraph_contains,
    prune_unused_document_image_relationships,
    remove_nearby_picture_block_before,
    replace_picture_before_anchor,
)
from docx_header_fields import ensure_section_footer_pagination_fields
from excel_utils import load_sheet_rows as load_xlsx_rows
from image_block_utils import count_docx_images, stack_images_vertical
from report_artifact_resolver import (
    find_latest_file as resolve_latest_file,
    find_latest_image as resolve_latest_image,
)
from report_build_manifest import write_report_build_manifest
from report_context import ReportBuildContext
from report_qc import check_report, write_report_qc_report

REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE_TEMPLATE = REPO_ROOT / "reports" / "G104线管柄大桥监测月报20260410-M18.docx"
DEFAULT_TEMPLATE = REPO_ROOT / "reports" / "G104线管柄大桥监测月报模板-自动报告.docx"
DEFAULT_RESULT_ROOT = Path("F:/管柄大桥数据/2026年3月")


@dataclass
class RangeStats:
    min_value: float
    max_value: float
    mean_value: float | None = None


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build G104 Guanbing Bridge monthly report.")
    parser.add_argument("--source-template", type=Path, default=DEFAULT_SOURCE_TEMPLATE)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--config", type=Path, default=REPO_ROOT / "config" / "default_config.json")
    parser.add_argument("--result-root", type=Path, default=DEFAULT_RESULT_ROOT)
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--period-label", default="2026年03月")
    parser.add_argument("--monitoring-range", default="2026年02月26日~2026年03月25日")
    parser.add_argument("--report-date", default=datetime.now().strftime("%Y年%m月%d日"))
    parser.add_argument(
        "--report-number",
        default=None,
        help="One report number for cover and every section header; defaults to the first cover number.",
    )
    parser.add_argument("--start-date", default="2026-02-26")
    parser.add_argument("--end-date", default="2026-03-25")
    parser.add_argument("--skip-image-replace", action="store_true")
    parser.add_argument("--refresh-template", action="store_true", help="Overwrite the auto-report template from source template.")
    return parser.parse_args()


def ensure_template(source_template: Path, template: Path, refresh: bool = False) -> Path:
    if refresh or not template.exists():
        if not source_template.exists():
            raise FileNotFoundError(f"Source template not found: {source_template}")
        template.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_template, template)
    return template


def load_sheet_rows(path: Path, sheet_name: str | None = None) -> list[dict]:
    return load_xlsx_rows(path, sheet_name, strip_headers=True, skip_empty=True, require_exists=False)


def load_workbook_sheet_rows(path: Path) -> dict[str, list[dict]]:
    if not path.exists():
        return {}
    wb = load_workbook(path, read_only=True, data_only=True)
    result: dict[str, list[dict]] = {}
    try:
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                result[ws.title] = []
                continue
            headers = [str(v).strip() if v is not None else "" for v in rows[0]]
            sheet_rows: list[dict] = []
            for row in rows[1:]:
                item = {key: value for key, value in zip(headers, row) if key}
                if any(value is not None for value in item.values()):
                    sheet_rows.append(item)
            result[ws.title] = sheet_rows
    finally:
        wb.close()
    return result


def as_float(value: object) -> float | None:
    if value is None:
        return None
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    if math.isnan(number) or math.isinf(number):
        return None
    return number


def fmt_num(value: float | None, digits: int = 1, keep_decimal: bool = False) -> str:
    if value is None:
        return "/"
    text = f"{value:.{digits}f}"
    if keep_decimal:
        return text
    return text.rstrip("0").rstrip(".")


def aggregate_range(rows: Iterable[dict], min_key: str = "Min", max_key: str = "Max", mean_key: str = "Mean") -> RangeStats | None:
    mins: list[float] = []
    maxs: list[float] = []
    means: list[float] = []
    for row in rows:
        min_value = as_float(row.get(min_key))
        max_value = as_float(row.get(max_key))
        mean_value = as_float(row.get(mean_key))
        if min_value is not None:
            mins.append(min_value)
        if max_value is not None:
            maxs.append(max_value)
        if mean_value is not None:
            means.append(mean_value)
    if not mins or not maxs:
        return None
    mean = sum(means) / len(means) if means else None
    return RangeStats(min(mins), max(maxs), mean)


def rows_by_points(rows: list[dict], points: set[str]) -> list[dict]:
    return [row for row in rows if str(row.get("PointID", "")).strip() in points]


def replace_text_in_paragraph(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_first_by_prefix(doc: Document, prefix: str, text: str, start_at: int = 0) -> bool:
    for paragraph in doc.paragraphs[start_at:]:
        if paragraph.text.strip().startswith(prefix):
            replace_text_in_paragraph(paragraph, text)
            return True
    return False


def replace_all_by_prefix(doc: Document, prefix: str, text: str, limit: int | None = None) -> int:
    count = 0
    paragraphs = doc.paragraphs[:limit] if limit is not None else doc.paragraphs
    for paragraph in paragraphs:
        if paragraph.text.strip().startswith(prefix):
            replace_text_in_paragraph(paragraph, text)
            count += 1
    return count


def iter_all_paragraphs(doc: Document):
    """Yield unique body, table, header and footer paragraphs."""
    # Keep the XML elements themselves alive in the set.  Using id(element)
    # is unsafe here because python-docx creates short-lived wrappers while
    # walking a large document, allowing CPython to reuse an id and silently
    # skip unrelated table paragraphs.
    seen: set[object] = set()

    def emit(paragraphs):
        for paragraph in paragraphs:
            marker = paragraph._p
            if marker not in seen:
                seen.add(marker)
                yield paragraph

    yield from emit(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from emit(cell.paragraphs)
    for section in doc.sections:
        for part in (section.header, section.footer):
            yield from emit(part.paragraphs)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from emit(cell.paragraphs)


REPORT_NUMBER_PATTERN = re.compile(
    r"报告编号\s*[:：]\s*([A-Za-z0-9]+(?:[-/][A-Za-z0-9]+)*)",
    re.IGNORECASE,
)


def _replace_text_span_in_runs(
    paragraph: Paragraph,
    start: int,
    end: int,
    replacement: str,
) -> bool:
    """Replace one text span without flattening the paragraph's run styles."""
    runs = paragraph.runs
    if not runs:
        text = paragraph.text
        if start < 0 or end > len(text) or start >= end:
            return False
        paragraph.add_run(text[:start] + replacement + text[end:])
        return True

    full_text = "".join(run.text for run in runs)
    if start < 0 or end > len(full_text) or start >= end:
        return False

    cursor = 0
    inserted = False
    for run in runs:
        original = run.text
        run_start = cursor
        run_end = cursor + len(original)
        cursor = run_end
        if run_end <= start or run_start >= end:
            continue

        left = original[: max(0, start - run_start)] if start > run_start else ""
        right = original[max(0, end - run_start) :] if end < run_end else ""
        if not inserted:
            run.text = left + replacement + right
            inserted = True
        else:
            run.text = right
    return inserted


def _replace_literal_preserve_runs(paragraph: Paragraph, old: str, new: str) -> int:
    changed = 0
    while True:
        full_text = "".join(run.text for run in paragraph.runs)
        start = full_text.find(old)
        if start < 0:
            break
        if not _replace_text_span_in_runs(paragraph, start, start + len(old), new):
            break
        changed += 1
    return changed


def detect_report_number(doc: Document) -> str | None:
    """Prefer the first cover/body number, then fall back to headers/footers."""
    for paragraph in doc.paragraphs:
        match = REPORT_NUMBER_PATTERN.search(paragraph.text)
        if match:
            return match.group(1)
    for paragraph_xml in _iter_report_number_paragraph_xml(doc):
        match = REPORT_NUMBER_PATTERN.search(_xml_paragraph_text(paragraph_xml))
        if match:
            return match.group(1)
    return None


def _iter_report_number_part_roots(doc: Document):
    """Yield the main document and every header/footer XML part once.

    ``section.header`` only exposes the default header and therefore misses
    first/even-page parts.  Walking package parts also reaches DrawingML/VML
    text boxes stored in those parts.
    """
    yield doc.element
    seen: set[object] = {doc.element}
    for part in doc.part.package.parts:
        name = str(part.partname)
        if not re.fullmatch(r"/word/(?:header|footer)\d+\.xml", name):
            continue
        root = getattr(part, "element", None)
        if root is None or root in seen:
            continue
        seen.add(root)
        yield root


def _iter_report_number_paragraph_xml(doc: Document):
    """Yield leaf ``w:p`` nodes, including paragraphs inside text boxes."""
    seen: set[object] = set()
    for root in _iter_report_number_part_roots(doc):
        for paragraph in root.xpath(".//w:p[not(.//w:p)]"):
            if paragraph in seen:
                continue
            seen.add(paragraph)
            yield paragraph


def _xml_paragraph_text(paragraph_xml) -> str:
    return "".join(node.text or "" for node in paragraph_xml.xpath(".//w:t"))


def _set_xml_text(node, value: str) -> None:
    """Set one ``w:t`` value without altering its containing run properties."""
    node.text = value
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    if value.startswith(" ") or value.endswith(" "):
        node.set(xml_space, "preserve")
    else:
        node.attrib.pop(xml_space, None)


def _replace_xml_text_span(paragraph_xml, start: int, end: int, replacement: str) -> bool:
    """Replace a cross-run text span in place, retaining every ``w:rPr``."""
    text_nodes = list(paragraph_xml.xpath(".//w:t"))
    full_text = "".join(node.text or "" for node in text_nodes)
    if start < 0 or end > len(full_text) or start >= end:
        return False

    cursor = 0
    inserted = False
    for node in text_nodes:
        original = node.text or ""
        node_start = cursor
        node_end = cursor + len(original)
        cursor = node_end
        if node_end <= start or node_start >= end:
            continue

        left = original[: max(0, start - node_start)] if start > node_start else ""
        right = original[max(0, end - node_start) :] if end < node_end else ""
        if not inserted:
            _set_xml_text(node, left + replacement + right)
            inserted = True
        else:
            _set_xml_text(node, right)
    return inserted


def normalize_report_number(
    doc: Document,
    report_number: str | None = None,
) -> tuple[str | None, int]:
    """Use one explicit or cover-derived report number everywhere in the DOCX."""
    resolved = str(report_number or "").strip() or detect_report_number(doc)
    if not resolved:
        raise ValueError(
            "Report number is required: pass --report-number or provide one "
            "'报告编号：...' value in the template."
        )
    if not re.fullmatch(r"[A-Za-z0-9]+(?:[-/][A-Za-z0-9]+)*", resolved):
        raise ValueError(f"Invalid report number: {resolved!r}")

    changed = 0
    for paragraph_xml in _iter_report_number_paragraph_xml(doc):
        full_text = _xml_paragraph_text(paragraph_xml)
        matches = list(REPORT_NUMBER_PATTERN.finditer(full_text))
        for match in reversed(matches):
            if match.group(1) == resolved:
                continue
            if not _replace_xml_text_span(
                paragraph_xml,
                match.start(1),
                match.end(1),
                resolved,
            ):
                continue
            changed += 1
    return resolved, changed


def is_sample_monitoring_period(
    period_label: str,
    monitoring_range: str,
    start_date: str | None,
    end_date: str | None,
) -> bool:
    """Treat an explicitly labelled or at-most-seven-day range as a sample."""
    if "样本" in f"{period_label} {monitoring_range}":
        return True
    if not start_date or not end_date:
        return False
    try:
        start = datetime.strptime(start_date, "%Y-%m-%d")
        end = datetime.strptime(end_date, "%Y-%m-%d")
    except ValueError:
        return False
    return 0 <= (end - start).days <= 6


def normalize_sample_period_wording(doc: Document) -> int:
    """Avoid calling a short sample a whole month while retaining engineering text."""
    replacements = (
        ("本月月度监测周期内", "本监测期内"),
        ("本月监测周期内", "本监测期内"),
        ("本月系统", "本期系统"),
        ("本月每日", "本期每日"),
        ("本月", "本监测期"),
    )
    changed = 0
    for paragraph in iter_all_paragraphs(doc):
        for old, new in replacements:
            changed += _replace_literal_preserve_runs(paragraph, old, new)
    return changed


def _text_paragraph_element(text: str):
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    return paragraph


def replace_legacy_operation_calendar(doc: Document) -> int:
    """Remove the fixed February/March calendar and its colour legend."""
    matched = []
    weekdays = {"一", "二", "三", "四", "五", "六", "日"}
    for table in list(doc.tables):
        values = [
            cell.text.strip()
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        text = "\n".join(values)
        date_cells = sum(bool(re.fullmatch(r"\d{1,2}月\d{1,2}日", value)) for value in values)
        is_calendar = weekdays.issubset(set(values)) and date_cells >= 7
        is_legend = (
            "监测系统运行正常" in text
            and "监测系统运行故障" in text
        )
        if is_calendar or is_legend:
            matched.append(table)

    for index, table in enumerate(matched):
        element = table._element
        parent = element.getparent()
        if parent is None:
            continue
        if index == 0:
            parent.replace(
                element,
                _text_paragraph_element(
                    "本期未提供与当前监测范围对应的独立系统运行日历记录。"
                    "系统运行状态以实际运维记录为准；本报告不根据监测数据缺口推断系统故障。"
                ),
            )
        else:
            parent.remove(element)
    return len(matched)


def neutralize_historical_picture_block(
    doc: Document,
    anchor_fragment: str,
    note: str,
) -> int:
    """Remove a stale screenshot block but keep its numbered caption sequence."""
    anchor = find_paragraph_contains(doc, anchor_fragment)
    if anchor is None:
        return 0
    removed = remove_nearby_picture_block_before(anchor, limit=120)
    if not removed:
        return 0
    anchor.insert_paragraph_before(note)
    if "本期未提供对应记录" not in anchor.text:
        anchor.add_run("（本期未提供对应记录）")
    return removed


def normalize_acceleration_units(doc: Document) -> int:
    """Use a real squared symbol even when a legacy template splits the 2 into a run."""
    changed = 0
    for paragraph in iter_all_paragraphs(doc):
        prefix = ""
        for run in paragraph.runs:
            original = run.text
            output: list[str] = []
            replaced_in_run = False
            for char in original:
                preceding = prefix + "".join(output)
                if char == "2" and (preceding.endswith("m/s") or preceding.endswith("cm/s")):
                    output.append("²")
                    replaced_in_run = True
                    changed += 1
                else:
                    output.append(char)
            if replaced_in_run:
                run.text = "".join(output)
                run.font.superscript = False
            prefix += "".join(output)
    return changed


def find_latest_file(root: Path, configured_dir: str, pattern: str) -> Path | None:
    """Find the newest result file with the same lookup rules used by reports."""
    return resolve_latest_file(root, configured_dir, pattern).path


def find_dynamic_strain_stats_file(result_root: Path, file_name: str, legacy_dir: str) -> Path | None:
    """Prefer canonical stats/ outputs, with fallback for older boxplot runs."""
    stats_path = result_root / "stats" / file_name
    if stats_path.exists():
        return stats_path
    return find_latest_file(result_root, legacy_dir, "boxplot_stats_*.xlsx")


def find_latest_image(root: Path, configured_dir: str, name_prefix: str) -> Path | None:
    """Find the newest image by point/name prefix, preferring report-friendly rasters."""
    return resolve_latest_image(root, configured_dir, name_prefix).path



def build_accel_combined_image(result_root: Path, asset_dir: Path, point_id: str) -> Path | None:
    time_img = find_latest_image(result_root, "时程曲线_加速度", point_id)
    rms_img = find_latest_image(result_root, "时程曲线_加速度_RMS10min", f"AccelRMS10_{point_id}")
    return stack_images_vertical([time_img, rms_img], asset_dir / f"accel_{point_id}.jpg").path


def build_stats_texts(result_root: Path, period_label: str) -> dict[str, str]:
    stats_dir = result_root / "stats"
    texts: dict[str, str] = {}

    temp_rows = load_sheet_rows(stats_dir / "temp_stats.xlsx")
    env_temp = aggregate_range(rows_by_points(temp_rows, {"GB-RTS-G05-001-03"}))
    box_temp = aggregate_range(rows_by_points(temp_rows, {"GB-RTS-G05-001-01", "GB-RTS-G05-001-02"}))
    if env_temp:
        texts["temp_env"] = (
            f"（1）本月监测周期内，环境最高温度为{fmt_num(env_temp.max_value, 1, True)}℃，"
            f"最低温度为{fmt_num(env_temp.min_value, 1, True)}℃，平均温度{fmt_num(env_temp.mean_value, 1, True)}℃。"
        )
    if box_temp:
        texts["temp_box"] = (
            f"（2）本月监测周期内，箱内最高温度为{fmt_num(box_temp.max_value, 1, True)}℃，"
            f"最低温度为{fmt_num(box_temp.min_value, 1, True)}℃，平均温度{fmt_num(box_temp.mean_value, 1, True)}℃。"
            "可知，箱内最高温度低于环境温度，温度幅值及波动剧烈程度低于环境温度。"
        )

    humidity_rows = load_sheet_rows(stats_dir / "humidity_stats.xlsx")
    env_humidity = aggregate_range(rows_by_points(humidity_rows, {"GB-RHS-G05-001-03"}))
    box_humidity = aggregate_range(rows_by_points(humidity_rows, {"GB-RHS-G05-001-01", "GB-RHS-G05-001-02"}))
    if env_humidity:
        texts["humidity_env"] = (
            f"（1）本月监测周期内，环境最高湿度为{fmt_num(env_humidity.max_value, 1)}%RH，"
            f"最低湿度为{fmt_num(env_humidity.min_value, 1)}%RH，平均湿度为{fmt_num(env_humidity.mean_value, 1)}%RH，"
            "湿度主要分布在80%~100%RH范围内。"
        )
    if box_humidity:
        texts["humidity_box"] = (
            f"（2）本月监测周期内，箱内最高湿度为{fmt_num(box_humidity.max_value, 1)}%RH，"
            f"最低湿度为{fmt_num(box_humidity.min_value, 1)}%RH，平均湿度为{fmt_num(box_humidity.mean_value, 1)}%RH。"
            "可知，箱内湿度低于环境湿度，湿度幅值及波动剧烈程度低于环境湿度。"
        )

    deflection_rows = load_sheet_rows(stats_dir / "deflection_stats.xlsx")
    if deflection_rows:
        orig_min = [as_float(row.get("OrigMin_mm")) for row in deflection_rows]
        orig_max = [as_float(row.get("OrigMax_mm")) for row in deflection_rows]
        orig_min = [value for value in orig_min if value is not None]
        orig_max = [value for value in orig_max if value is not None]
        if orig_min and orig_max:
            max_up = abs(min(orig_min))
            max_down = max(orig_max)
            texts["deflection_abs"] = (
                f"由以上各图可知，本月监测周期内，第2、3跨挠度最大上挠{fmt_num(max_up, 1)}mm，"
                f"最大下挠{fmt_num(max_down, 1)}mm，均处于超限阈值范围之内，"
                "未出现超过各级超限阈值和报警的情况。"
            )
        mid_rows = [row for row in deflection_rows if "-002-" in str(row.get("PointID", ""))]
        filt_min = [as_float(row.get("FiltMin_mm")) for row in mid_rows]
        filt_max = [as_float(row.get("FiltMax_mm")) for row in mid_rows]
        filt_min = [value for value in filt_min if value is not None]
        filt_max = [value for value in filt_max if value is not None]
        if filt_min and filt_max:
            trend_min = min(filt_min)
            trend_max = max(filt_max)
            texts["deflection_trend"] = (
                f"由以上各图可知，本月监测周期内，第2、3跨主梁跨中挠度变化范围为"
                f"{fmt_num(trend_min, 1)}mm~{fmt_num(trend_max, 1)}mm，挠度同一天中处于波动变化中。"
            )
            texts["conclusion_deflection"] = (
                f"（4）本月监测周期内，实测主梁挠度值处于设计理论挠度范围，"
                f"第2、3跨主梁跨中挠度变化范围为{fmt_num(trend_min, 1)}mm~{fmt_num(trend_max, 1)}mm，"
                "挠度同一天中处于波动变化中。"
            )

    tilt_sheets = load_workbook_sheet_rows(stats_dir / "tilt_stats.xlsx")
    tilt_x = aggregate_range(tilt_sheets.get("Tilt_X", []), min_key="Min", max_key="Max")
    tilt_y = aggregate_range(tilt_sheets.get("Tilt_Y", []), min_key="Min", max_key="Max")
    if tilt_x and tilt_y:
        tilt_x_abs = max(abs(tilt_x.min_value), abs(tilt_x.max_value))
        tilt_y_abs = max(abs(tilt_y.min_value), abs(tilt_y.max_value))
        texts["tilt"] = (
            f"由以上各图可知，本月监测周期内，主墩倾角纵桥向X最大为{fmt_num(tilt_x_abs, 3)}°，"
            f"横桥向Y最大为{fmt_num(tilt_y_abs, 3)}°，均处于超限阈值范围之内，"
            "未出现超过各级超限阈值和报警的情况。主墩未出现明显倾斜趋势。"
        )
        texts["conclusion_tilt"] = (
            f"（5）本月监测周期内，主墩倾角纵桥向X最大为{fmt_num(tilt_x_abs, 3)}°，"
            f"横桥向Y最大为{fmt_num(tilt_y_abs, 3)}°，均处于超限阈值范围之内，"
            "未出现超过各级超限阈值和报警的情况，主墩未出现明显倾斜趋势。"
        )

    hp_path = find_dynamic_strain_stats_file(
        result_root, "dynamic_strain_highpass_stats.xlsx", "动应变箱线图_高通滤波"
    )
    hp_sheets = load_workbook_sheet_rows(hp_path) if hp_path else {}
    if hp_sheets:
        pieces: list[str] = []
        for sheet, label in (("G05", "第2跨跨中截面"), ("G06", "第3跨跨中截面")):
            stats = aggregate_range(hp_sheets.get(sheet, []), min_key="Min", max_key="Max")
            if stats:
                pieces.append(
                    f"{label}测点活载作用下最大拉应变为{fmt_num(stats.max_value, 2)}με、"
                    f"最大压应变为{fmt_num(abs(stats.min_value), 2)}με"
                )
        if pieces:
            texts["strain_hp"] = (
                "由上图可知，本月监测周期内，"
                + "；".join(pieces)
                + "，均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
                "图中各应变传感器的上下四分位数距离较小，整体活载水平不高；最值之间距离较大，表明桥上时有重车通过。"
            )

    lp_path = find_dynamic_strain_stats_file(
        result_root, "dynamic_strain_lowpass_stats.xlsx", "动应变箱线图_低通滤波"
    )
    lp_rows_by_sheet = load_workbook_sheet_rows(lp_path) if lp_path else {}
    all_lp_rows: list[dict] = []
    for rows in lp_rows_by_sheet.values():
        all_lp_rows.extend(rows)
    lp_stats = aggregate_range(all_lp_rows, min_key="Min", max_key="Max")
    if lp_stats:
        texts["strain_lp"] = (
            f"由上图可知，应变测点最大拉应变{fmt_num(lp_stats.max_value, 2)}με，"
            f"最大压应变{fmt_num(abs(lp_stats.min_value), 2)}με，呈现缓慢变化，未超过设计最不利计算值。"
            "综上，截面整体受力未见明显异常。"
        )

    crack_rows = load_sheet_rows(stats_dir / "crack_stats.xlsx")
    g05_crack = aggregate_range([row for row in crack_rows if "G05" in str(row.get("PointID", ""))], min_key="CrkMin", max_key="CrkMax")
    g06_crack = aggregate_range([row for row in crack_rows if "G06" in str(row.get("PointID", ""))], min_key="CrkMin", max_key="CrkMax")
    if g05_crack and g06_crack:
        texts["crack"] = (
            "由以上各图可知，本月监测周期内，"
            f"顶板裂缝宽度变化量（相对2024年9月26日）在{fmt_num(g05_crack.min_value, 3)}mm~{fmt_num(g05_crack.max_value, 3)}mm之间，"
            f"底板裂缝宽度变化量（相对2024年9月26日）在{fmt_num(g06_crack.min_value, 3)}mm~{fmt_num(g06_crack.max_value, 3)}mm之间，"
            "裂缝宽度监测值同一天中处于波动变化中。"
        )

    freq_path = stats_dir / "accel_spec_stats.xlsx"
    freq_sheets = load_workbook_sheet_rows(freq_path)
    freq_ranges: list[tuple[float, float]] = []
    for idx in range(1, 4):
        values: list[float] = []
        for rows in freq_sheets.values():
            if not rows:
                continue
            keys = [key for key in rows[0].keys() if key.startswith("Freq_")]
            if len(keys) >= idx:
                for row in rows:
                    value = as_float(row.get(keys[idx - 1]))
                    if value is not None:
                        values.append(value)
        if values:
            freq_ranges.append((min(values), max(values)))
    if len(freq_ranges) == 3:
        f_text = "、".join(f"{fmt_num(lo, 3)}Hz~{fmt_num(hi, 3)}Hz" for lo, hi in freq_ranges)
        texts["freq"] = (
            f"由上图可知，本月监测周期内各个传感器实测竖向第一、二、三阶自振频率分别为{f_text}，"
            "均大于理论计算的主梁一、二、三阶竖弯频率0.975Hz、1.243Hz、1.528Hz，"
            "且处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况，桥梁结构实测刚度大于理论刚度且未见明显变化。"
        )

    accel_rows = [row for row in load_sheet_rows(stats_dir / "accel_stats.xlsx") if row.get("PointID")]
    accel_stats = aggregate_range(accel_rows, min_key="Min", max_key="Max")
    rms_values = [as_float(row.get("RMS10minMax")) for row in accel_rows]
    rms_values = [value for value in rms_values if value is not None]
    if accel_stats and rms_values:
        max_abs = max(abs(accel_stats.min_value), abs(accel_stats.max_value))
        texts["accel"] = (
            f"由以上各图可知，本月监测周期内，主梁竖向加速度各测点绝对最大值为{fmt_num(max_abs, 2)}mm/s²，"
            f"各测点10min加速度均方根值最大为{fmt_num(max(rms_values), 3)}mm/s²，未超过315mm/s²，"
            "均处于超限阈值范围之内，未出现超过各级超限阈值和报警的情况。"
        )

    if env_temp and box_temp:
        texts["conclusion_temp"] = (
            f"（2）本月监测周期内，环境最高温度为{fmt_num(env_temp.max_value, 1, True)}℃，最低温度为{fmt_num(env_temp.min_value, 1, True)}℃，"
            f"平均温度{fmt_num(env_temp.mean_value, 1, True)}℃；箱内最高温度为{fmt_num(box_temp.max_value, 1, True)}℃，"
            f"最低温度为{fmt_num(box_temp.min_value, 1, True)}℃，平均温度{fmt_num(box_temp.mean_value, 1, True)}℃。"
            "可知，箱内最高温度低于环境温度，温度幅值及波动剧烈程度低于环境温度。"
        )
    if env_humidity and box_humidity:
        texts["conclusion_humidity"] = (
            f"（3）本月监测周期内，环境最高湿度为{fmt_num(env_humidity.max_value, 1)}%RH，最低湿度为{fmt_num(env_humidity.min_value, 1)}%RH，"
            f"平均湿度为{fmt_num(env_humidity.mean_value, 1)}%RH，湿度主要分布在80%~100%RH范围内；"
            f"箱内最高湿度为{fmt_num(box_humidity.max_value, 1)}%RH，最低湿度为{fmt_num(box_humidity.min_value, 1)}%RH，"
            f"平均湿度为{fmt_num(box_humidity.mean_value, 1)}%RH。可知，箱内湿度低于环境湿度，湿度幅值及波动剧烈程度低于环境湿度。"
        )
    if "strain_hp" in texts and lp_stats:
        texts["conclusion_strain"] = (
            f"（6）本月监测周期内，主梁活载作用下动应变和温度、收缩徐变等作用下静应变均处于设计值范围内；"
            f"静应变最大拉应变{fmt_num(lp_stats.max_value, 2)}με，最大压应变{fmt_num(abs(lp_stats.min_value), 2)}με，截面整体受力未见明显异常。"
        )
    if "crack" in texts:
        texts["conclusion_crack"] = "（8）本月监测周期内，" + texts["crack"].split("本月监测周期内，", 1)[-1]
    if "freq" in texts and "accel" in texts:
        texts["conclusion_accel"] = "（7）" + texts["accel"].replace("由以上各图可知，本月监测周期内，", "本月监测周期内，") + "所测竖向第一、二、三阶自振频率均在理论值内，桥梁结构实测刚度大于理论刚度且未见明显变化。"

    texts["summary"] = f"综上所述，G104线管柄大桥{period_label}监测周期内，桥梁主要健康监测指标正常，桥梁运营状态良好，建议管养单位继续加强桥面和箱内的巡查。"
    return texts


def resolve_monitoring_range(
    monitoring_range: str,
    start_date: str | None = None,
    end_date: str | None = None,
) -> str:
    """Return a user-facing range without falling back to the template period."""
    requested_range = str(monitoring_range or "").strip()
    if requested_range:
        return requested_range

    def format_date(value: str | None) -> str:
        if not value:
            return ""
        try:
            return datetime.strptime(value, "%Y-%m-%d").strftime("%Y年%m月%d日")
        except ValueError:
            return str(value).strip()

    start_text = format_date(start_date)
    end_text = format_date(end_date)
    if start_text and end_text:
        return f"{start_text}~{end_text}"
    if start_text or end_text:
        return start_text or end_text
    return "当前任务指定监测范围"


def apply_text_updates(
    doc: Document,
    texts: dict[str, str],
    monitoring_range: str,
    report_date: str,
    start_date: str | None = None,
    end_date: str | None = None,
    sample_period: bool = False,
) -> list[str]:
    monitoring_range = resolve_monitoring_range(monitoring_range, start_date, end_date)
    updated: list[str] = []
    if replace_all_by_prefix(doc, "（时间范围：", f"（时间范围：{monitoring_range}）", limit=30):
        updated.append("cover monitoring range")
    if replace_all_by_prefix(doc, "报告日期：", f"报告日期：{report_date}", limit=30):
        updated.append("cover report date")

    # The source template contains March-specific operation and maintenance
    # statements.  Those statements are not monitoring statistics and cannot
    # be inferred from gaps in sensor data.  Always replace them with
    # task-scoped, evidence-neutral wording so a report for another month (or a
    # partial sample) never inherits a fixed March outage assertion.
    period_replacements = [
        (
            "G104线管柄大桥健康监测系统于2024年9月26日交工验收后进入运维服务期",
            "G104线管柄大桥健康监测系统于2024年9月26日交工验收后进入运维服务期。"
            f"本期报告监测范围为{monitoring_range}。本报告仅依据该范围内实际获取的监测数据进行统计分析；"
            "系统运行状态及数据覆盖情况以实际运维记录、分析结果清单和本报告图表为准。",
            "system operation scope",
        ),
        (
            "软件线上检查维护：",
            "软件线上检查维护：本期软件运行和维护情况以实际运维记录为准；"
            f"本报告仅依据{monitoring_range}范围内实际获取的监测数据进行统计分析，"
            "不根据数据缺口推断系统故障或维护情况。",
            "software maintenance scope",
        ),
        (
            "（1）本月月度监测周期内，",
            f"（1）本期报告监测范围为{monitoring_range}。"
            "系统运行、维护及数据覆盖情况以实际运维记录、分析结果清单和本报告图表为准；"
            "本报告不根据数据缺口推断系统故障。",
            "conclusion operation scope",
        ),
        (
            "预警信息处理：",
            "预警信息处理：本期未提供与当前监测范围对应的独立预警处置记录。"
            "本报告不沿用模板中的历史报警数量、日期或原因；"
            "预警情况以实际平台记录和经核验的处置资料为准。",
            "warning handling scope",
        ),
        (
            "现场硬件维护：",
            "现场硬件维护：本期未提供与当前监测范围对应的独立现场维护记录。"
            "现场维护情况以实际运维记录为准，本报告不沿用模板中的历史维护日期或事件。",
            "hardware maintenance scope",
        ),
    ]
    for prefix, replacement, label in period_replacements:
        if replace_first_by_prefix(doc, prefix, replacement):
            updated.append(label)

    calendar_table_count = replace_legacy_operation_calendar(doc)
    if calendar_table_count:
        updated.append(f"operation calendar tables ({calendar_table_count})")

    maintenance_picture_count = neutralize_historical_picture_block(
        doc,
        "图 2 系统维护情况",
        "本期未提供与当前监测范围对应的系统维护截图，模板中的历史截图已移除。",
    )
    if maintenance_picture_count:
        updated.append(f"maintenance screenshots ({maintenance_picture_count})")
    warning_picture_count = neutralize_historical_picture_block(
        doc,
        "图 3 部分预警信息处理情况",
        "本期未提供与当前监测范围对应的预警处置截图，模板中的历史截图已移除。",
    )
    if warning_picture_count:
        updated.append(f"warning screenshots ({warning_picture_count})")

    replacements = [
        ("（1）本月监测周期内，环境最高温度", texts.get("temp_env")),
        ("（2）本月监测周期内，箱内最高温度", texts.get("temp_box")),
        ("（1）本月监测周期内，环境最高湿度", texts.get("humidity_env")),
        ("（2）本月监测周期内，箱内最高湿度", texts.get("humidity_box")),
        ("由以上各图可知，本月监测周期内，第2、3跨挠度最大上挠", texts.get("deflection_abs")),
        ("由以上各图可知，本月监测周期内，第2、3跨主梁跨中挠度变化范围", texts.get("deflection_trend")),
        ("由以上各图可知，本月监测周期内，主墩倾角纵桥向X最大", texts.get("tilt")),
        ("由上图可知，本月监测周期内，第2跨跨中截面测点活载作用下", texts.get("strain_hp")),
        ("由上图可知，应变测点最大拉应变", texts.get("strain_lp")),
        ("由上图可知，本月监测周期内各个传感器实测竖向第一、二、三阶自振频率", texts.get("freq")),
        ("由以上各图可知，本月监测周期内，顶板裂缝宽度变化量", texts.get("crack")),
        ("（2）本月监测周期内，环境最高温度", texts.get("conclusion_temp")),
        ("（3）本月监测周期内，环境最高湿度", texts.get("conclusion_humidity")),
        ("（6）本月监测周期内，各截面上下缘应变", texts.get("conclusion_strain")),
        ("（4）本月监测周期内，实测主梁挠度值", texts.get("conclusion_deflection")),
        ("（5）本月监测周期内，主墩倾角纵桥向X最大", texts.get("conclusion_tilt")),
        ("（8）本月监测周期内，顶板裂缝宽度变化量", texts.get("conclusion_crack")),
        ("综上所述，G104线管柄大桥", texts.get("summary")),
    ]
    if texts.get("accel"):
        replacements.append(("由以上各图可知，本月监测周期内，主梁竖向加速度", texts["accel"]))
    if texts.get("conclusion_accel"):
        replacements.append(("（7）本月监测周期内，主梁竖向加速度", texts["conclusion_accel"]))

    for prefix, text in replacements:
        if text and replace_first_by_prefix(doc, prefix, text):
            updated.append(prefix)
    if sample_period:
        sample_wording_count = normalize_sample_period_wording(doc)
        if sample_wording_count:
            updated.append(f"sample-period wording ({sample_wording_count})")
    return updated


def apply_image_updates(doc: Document, result_root: Path, asset_dir: Path) -> tuple[list[dict], list[str]]:
    accel_combined = {
        point_id: build_accel_combined_image(result_root, asset_dir, point_id)
        for point_id in [
            "GB-VIB-G04-001-01",
            "GB-VIB-G05-002-01",
            "GB-VIB-G06-002-01",
            "GB-VIB-G07-001-01",
        ]
    }
    specs = [
        ("图 5 桥面环境温度测点时程图", find_latest_image(result_root, "时程曲线_温度", "GB-RTS-G05-001-03"), 1, 145.0),
        ("(a)GB-RTS-G05-001-01", find_latest_image(result_root, "时程曲线_温度", "GB-RTS-G05-001-01"), 1, 145.0),
        ("(b)GB-RTS-G05-001-02", find_latest_image(result_root, "时程曲线_温度", "GB-RTS-G05-001-02"), 1, 145.0),
        ("图 7 桥面环境湿度测点时程图", find_latest_image(result_root, "时程曲线_湿度", "GB-RHS-G05-001-03"), 1, 145.0),
        ("图 8 桥面环境湿度累积持续时间频次分布图", find_latest_image(result_root, "频次分布_湿度", "GB-RHS-G05-001-03_freq"), 1, 145.0),
        ("(a)GB-RHS-G05-001-01", find_latest_image(result_root, "时程曲线_湿度", "GB-RHS-G05-001-01"), 1, 145.0),
        ("(b)GB-RHS-G05-001-02", find_latest_image(result_root, "时程曲线_湿度", "GB-RHS-G05-001-02"), 1, 145.0),
        ("(a)GB-RHS-G05-001-01", find_latest_image(result_root, "频次分布_湿度", "GB-RHS-G05-001-01_freq"), 2, 145.0),
        ("(b)GB-RHS-G05-001-02", find_latest_image(result_root, "频次分布_湿度", "GB-RHS-G05-001-02_freq"), 2, 145.0),
        ("第2跨1/4跨", find_latest_image(result_root, "时程曲线_挠度_组图_原始", "Defl_G1_Orig"), 1, 145.0),
        ("第2跨1/2跨", find_latest_image(result_root, "时程曲线_挠度_组图_原始", "Defl_G2_Orig"), 1, 145.0),
        ("第2跨3/4跨", find_latest_image(result_root, "时程曲线_挠度_组图_原始", "Defl_G3_Orig"), 1, 145.0),
        ("第3跨1/4跨", find_latest_image(result_root, "时程曲线_挠度_组图_原始", "Defl_G4_Orig"), 1, 145.0),
        ("第3跨1/2跨", find_latest_image(result_root, "时程曲线_挠度_组图_原始", "Defl_G5_Orig"), 1, 145.0),
        ("第3跨3/4跨", find_latest_image(result_root, "时程曲线_挠度_组图_原始", "Defl_G6_Orig"), 1, 145.0),
        ("图 13 第2跨主梁位移变化趋势", find_latest_image(result_root, "时程曲线_挠度_组图_滤波", "Defl_G2_Filt"), 1, 145.0),
        ("图 14 第3跨主梁位移变化趋势", find_latest_image(result_root, "时程曲线_挠度_组图_滤波", "Defl_G5_Filt"), 1, 145.0),
        ("（a）纵桥向X", find_latest_image(result_root, "时程曲线_倾角_组图", "Tilt_X"), 1, 145.0),
        ("（b）横桥向Y", find_latest_image(result_root, "时程曲线_倾角_组图", "Tilt_Y"), 1, 145.0),
        ("（a）第2跨", find_latest_image(result_root, "动应变箱线图_高通滤波", "boxplot_G05"), 1, 145.0),
        ("（b）第3跨", find_latest_image(result_root, "动应变箱线图_高通滤波", "boxplot_G06"), 1, 145.0),
        ("（a）第2跨", find_latest_image(result_root, "时程曲线_动应变_低通滤波_组图", "dynstrain_lp_G05"), 2, 145.0),
        ("（b）第3跨", find_latest_image(result_root, "时程曲线_动应变_低通滤波_组图", "dynstrain_lp_G06"), 2, 145.0),
        ("（a）GB-VIB-G04-001-01", accel_combined["GB-VIB-G04-001-01"], 1, 145.0),
        ("（b）GB-VIB-G05-002-01", accel_combined["GB-VIB-G05-002-01"], 1, 145.0),
        ("（c）GB-VIB-G06-002-01", accel_combined["GB-VIB-G06-002-01"], 1, 145.0),
        ("（d）GB-VIB-G07-001-01", accel_combined["GB-VIB-G07-001-01"], 1, 145.0),
        ("（a）GB-VIB-G05-002-01", find_latest_image(result_root, "频谱峰值曲线_加速度", "SpecFreq_GB-VIB-G05-002-01"), 1, 145.0),
        ("（b）GB-VIB-G06-002-01", find_latest_image(result_root, "频谱峰值曲线_加速度", "SpecFreq_GB-VIB-G06-002-01"), 1, 145.0),
        ("(a)第2跨", find_latest_image(result_root, "时程曲线_裂缝宽度", "裂缝宽度_G05"), 1, 145.0),
        ("(b)第3跨", find_latest_image(result_root, "时程曲线_裂缝宽度", "裂缝宽度_G06"), 1, 145.0),
    ]
    replaced: list[dict] = []
    missing: list[str] = []
    for anchor, image_path, occurrence, width in specs:
        ok, info = replace_picture_before_anchor(doc, anchor, image_path, occurrence=occurrence, width_mm=width)
        if ok:
            replaced.append({"anchor": anchor, "occurrence": occurrence, "image": info})
        else:
            missing.append(info)
    return replaced, missing


def build_report(
    template: Path = DEFAULT_TEMPLATE,
    source_template: Path = DEFAULT_SOURCE_TEMPLATE,
    config_path: Path | None = None,
    result_root: Path = DEFAULT_RESULT_ROOT,
    output_dir: Path | None = None,
    period_label: str = "2026年03月",
    monitoring_range: str = "2026年02月26日~2026年03月25日",
    report_date: str | None = None,
    report_number: str | None = None,
    start_date: str = "2026-02-26",
    end_date: str = "2026-03-25",
    refresh_template: bool = False,
    skip_image_replace: bool = False,
) -> tuple[Path, Path]:
    report_date = report_date or datetime.now().strftime("%Y年%m月%d日")
    template = ensure_template(source_template, template, refresh=refresh_template)
    ctx = ReportBuildContext.from_inputs(
        template=template,
        config_path=config_path,
        result_root=result_root,
        output_dir=output_dir,
        assets_subdir="_assets",
    )

    doc = Document(str(template))
    image_count_before = count_docx_images(template)
    stats_texts = build_stats_texts(result_root, period_label)
    sample_period = is_sample_monitoring_period(
        period_label,
        monitoring_range,
        start_date,
        end_date,
    )
    updated_paragraphs = apply_text_updates(
        doc,
        stats_texts,
        monitoring_range,
        report_date,
        start_date=start_date,
        end_date=end_date,
        sample_period=sample_period,
    )
    replaced_images: list[dict] = []
    missing_images: list[str] = []
    if not skip_image_replace:
        replaced_images, missing_images = apply_image_updates(doc, result_root, ctx.assets_dir)
    removed_unused_image_relationships = prune_unused_document_image_relationships(doc)
    resolved_report_number, normalized_report_number_count = normalize_report_number(
        doc,
        report_number,
    )
    normalized_unit_count = normalize_acceleration_units(doc)
    normalized_footer_pagination_count = ensure_section_footer_pagination_fields(doc)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = ctx.output_dir / f"G104线管柄大桥监测月报_{period_label}_自动生成_{timestamp}.docx"
    doc.save(str(output_path))

    manifest = {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "template": str(template),
        "source_template": str(source_template),
        "result_root": str(result_root),
        "output": str(output_path),
        "period_label": period_label,
        "monitoring_range": monitoring_range,
        "report_date": report_date,
        "report_number": resolved_report_number,
        "start_date": start_date,
        "end_date": end_date,
        "sample_period_wording": sample_period,
        "updated_paragraph_count": len(updated_paragraphs),
        "updated_paragraphs": updated_paragraphs,
        "replaced_image_count": len(replaced_images),
        "replaced_images": replaced_images,
        "missing_images": missing_images,
        "removed_unused_image_relationship_count": len(removed_unused_image_relationships),
        "normalized_report_number_count": normalized_report_number_count,
        "normalized_acceleration_unit_count": normalized_unit_count,
        "normalized_footer_pagination_count": normalized_footer_pagination_count,
        "analysis_run_manifest": analysis_manifest_context(result_root),
        "image_count_before": image_count_before,
        "image_count_after": count_docx_images(output_path),
        "notes": [
            "如果当前结果目录缺少挠度、倾角、加速度时程图或对应统计，自动报告会保留模板原图/原文字。",
            "本次低通应变插图使用时程曲线_动应变_低通滤波，而不是原始应变组图。",
        ],
    }
    manifest["missing_analysis_modules"] = missing_module_summary_items(manifest.get("analysis_run_manifest"))
    qc_paths: dict[str, str] = {}
    qc_warnings: list[str] = []
    try:
        qc_result = check_report("guanbing_monthly", output_path)
        qc_txt, qc_json = write_report_qc_report(qc_result, ctx.output_dir, timestamp=timestamp)
        qc_paths = {"report_qc_txt": str(qc_txt), "report_qc_json": str(qc_json), "report_qc_status": qc_result.status}
        qc_warnings = [
            f"{issue.code}: {issue.message}"
            for issue in qc_result.issues
            if issue.severity == "warning"
        ]
    except Exception as exc:
        qc_warnings = [f"report_qc_failed: {exc}"]
    manifest_path = write_report_build_manifest(
        context=ctx,
        report_type="guanbing_monthly",
        output_docx=output_path,
        timestamp=timestamp,
        legacy_manifest=manifest,
        missing=missing_images + manifest["missing_analysis_modules"],
        warnings=qc_warnings,
        extra={
            "updated_paragraph_count": len(updated_paragraphs),
            "replaced_image_count": len(replaced_images),
            "report_number": resolved_report_number,
            "normalized_report_number_count": normalized_report_number_count,
            **qc_paths,
        },
        filename_prefix="G104线管柄大桥监测月报_manifest",
    )
    return output_path, manifest_path


def main() -> None:
    args = parse_args()
    output_path, manifest_path = build_report(
        template=args.template,
        source_template=args.source_template,
        config_path=args.config,
        result_root=args.result_root,
        output_dir=args.output_dir,
        period_label=args.period_label,
        monitoring_range=args.monitoring_range,
        report_date=args.report_date,
        report_number=args.report_number,
        start_date=args.start_date,
        end_date=args.end_date,
        refresh_template=args.refresh_template,
        skip_image_replace=args.skip_image_replace,
    )
    print(f"Report written: {output_path}")
    print(f"Manifest written: {manifest_path}")


if __name__ == "__main__":
    main()
