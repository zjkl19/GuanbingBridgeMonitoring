from __future__ import annotations

import argparse
import json
import posixpath
import shutil
import subprocess
from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageChops


BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
AMBER = "FFF2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = "000000") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for name, size, color in (("Title", 24, BLUE), ("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE)):
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.add_run(text)


def extract_layout_image(report: Path, output_dir: Path) -> Path:
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    with ZipFile(report, "r") as archive:
        document = ET.fromstring(archive.read("word/document.xml"))
        rels = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
        targets = {
            rel.attrib["Id"]: rel.attrib["Target"]
            for rel in rels.findall("pr:Relationship", ns)
            if rel.attrib.get("Type", "").endswith("/image")
        }
        paragraphs = document.findall(".//w:p", ns)
        member = ""
        for index, paragraph in enumerate(paragraphs):
            text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns))
            if "风向风速仪布置示意图" not in text:
                continue
            for previous in reversed(paragraphs[max(0, index - 3):index]):
                ids = [
                    node.attrib.get(f"{{{ns['r']}}}embed")
                    for node in previous.findall(".//a:blip", ns)
                ]
                ids = [value for value in ids if value in targets]
                if len(ids) == 1:
                    member = posixpath.normpath(posixpath.join("word", targets[ids[0]]))
                    break
            if member:
                break
        if not member:
            raise RuntimeError("wind-sensor layout image was not found in the source report")
        suffix = Path(member).suffix or ".png"
        output = output_dir / f"hongtang_wind_sensor_layout{suffix}"
        output.write_bytes(archive.read(member))
        if suffix.lower() in {".emf", ".wmf"}:
            png = output.with_suffix(".png")
            if not png.exists():
                soffice = shutil.which("soffice.com") or shutil.which("soffice")
                if not soffice:
                    raise RuntimeError("LibreOffice is required to convert the layout EMF to PNG")
                proc = subprocess.run(
                    [soffice, "--headless", "--convert-to", "png", "--outdir", str(output_dir), str(output)],
                    text=True,
                    capture_output=True,
                )
                if proc.returncode or not png.exists():
                    raise RuntimeError(f"layout image conversion failed: {proc.stderr or proc.stdout}")
            cropped = png.with_name(f"{png.stem}_cropped.png")
            image = Image.open(png).convert("RGB")
            background = Image.new("RGB", image.size, "white")
            difference = ImageChops.difference(image, background).convert("L")
            mask = difference.point(lambda value: 255 if value > 12 else 0)
            bbox = mask.getbbox()
            if bbox:
                pad = 18
                left = max(0, bbox[0] - pad)
                top = max(0, bbox[1] - pad)
                right = min(image.width, bbox[2] + pad)
                bottom = min(image.height, bbox[3] + pad)
                image.crop((left, top, right, bottom)).save(cropped)
                return cropped
            return png
        return output


def build(args: argparse.Namespace) -> Path:
    output_dir = args.output.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    diagnostic = json.loads(args.diagnostic.read_text(encoding="utf-8-sig"))
    layout = extract_layout_image(args.source_report, output_dir)

    doc = Document()
    configure_document(doc)
    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("洪塘大桥 W1/W2 风速差异专项排查简报")
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("2026年第二季度｜专项报告工具 v1.7.37｜风数据重算 v1.7.36").font.size = Pt(13)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(15.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(
        "结论先行：领导对“桥面 W1 与塔顶 W2 的结果关系值得复核”的判断是合理的。"
        "两通道不是重复数据，程序重算口径已纠正；但 W2 长期系统性低于 W1，不能只用高度增风规律解释。"
        "现有证据更符合测点暴露、桥塔/构件绕流与可能的安装遮挡或标定差异共同作用，建议对 W2 做现场复核。"
    )
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_heading("1. 核心判断", level=1)
    add_bullet(doc, "不是“领导看错”：报告中两测点瞬时/10分钟峰值确实接近，容易形成两者风速差不多的直观印象。")
    add_bullet(doc, "不是“程序把 W1/W2 用成同一通道”：配置与源文件标识不同，配对序列相关但不相同，且两测点风向分布差异显著。")
    add_bullet(doc, "程序曾存在滚动导出自然日拼接与风玫瑰方位显示问题，现已修复并按完整自然日重算；修复后，W2 长期偏低的主结论仍存在。")
    add_bullet(doc, "尚不能直接判定仪器故障：数据无持续饱和、复制或大面积失真证据，但 W2 的安装遮挡、朝向基准、标定/换算系数仍需现场验证。")

    doc.add_page_break()
    doc.add_heading("2. 工程与测点关系", level=1)
    doc.add_paragraph(
        "洪塘大桥主桥跨越乌龙江，主桥为 50 m + 150 m + 150 m + 50 m 的独塔自锚式悬索桥。"
        "原报告工程概况及布置说明明确：W1 位于右幅桥面 12 号墩附近散索鞍保护罩，W2 位于 13 号主塔塔顶。"
        "两者既不在同一平面位置，也不是同一竖向测风剖面，因此不能直接套用“高度越高风速越大”的单一边界层关系。"
    )
    doc.add_picture(str(layout), width=Cm(15.6))
    add_caption(doc, "图 1  原季度报告中的风向风速仪布置示意图")
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("测点", "位置", "主要局部影响", "判断注意事项")
    for idx, text in enumerate(headers):
        set_cell_shading(table.cell(0, idx), BLUE)
        set_cell_text(table.cell(0, idx), text, bold=True, color="FFFFFF")
    rows = (
        ("W1", "右幅桥面12号墩附近散索鞍保护罩", "江面来流、桥面/箱梁边缘及保护罩局部加速", "代表桥面局部暴露风场"),
        ("W2", "13号主塔塔顶", "塔柱、横梁、检修构件及探头安装方位的尾流/遮挡", "需核实探头距构件距离、方向零位与标定"),
    )
    for r_idx, row in enumerate(rows, 1):
        for c_idx, text in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), text)

    doc.add_page_break()
    doc.add_heading("3. 数据与程序复核结果", level=1)
    table = doc.add_table(rows=8, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, text in enumerate(("指标", "W1 桥面", "W2 塔顶", "解释")):
        set_cell_shading(table.cell(0, idx), BLUE)
        set_cell_text(table.cell(0, idx), text, bold=True, color="FFFFFF")
    rows = (
        ("全期平均风速", "2.74 m/s", "1.36 m/s", "W2/W1 = 49.7%"),
        ("全期原始最大风速", "12.25 m/s", "12.68 m/s", "峰值接近，解释了直观上的“差不多”"),
        ("最大10分钟平均风速", "6.89 m/s", "6.51 m/s", "峰值接近，但发生时间不同"),
        ("配对10分钟窗", "12,051", "12,051", "同时间窗严格配对"),
        ("配对相关系数", "—", "0.682", "共同受区域天气驱动，但不是复制通道"),
        ("W2 高于 W1 的时段占比", "—", "1.77%", "W2 偏低具有持续性"),
        ("两测点风向差中位数", "—", "91.3°", "局部流场/朝向基准存在显著差异"),
    )
    for r_idx, row in enumerate(rows, 1):
        for c_idx, text in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), text)
        if r_idx in (1, 6, 7):
            for c_idx in range(4):
                set_cell_shading(table.cell(r_idx, c_idx), AMBER)

    doc.add_heading("程序与数据质量检查", level=2)
    add_bullet(doc, "已采用滚动导出 D 日 + D+1 日文件重建自然日，避免每日只保留约 09:00 前数据。")
    add_bullet(doc, "W1/W2 风速、风向源文件与配置点号分别独立；全季源数据/输入点数闭环，正式图均带 provenance。")
    add_bullet(doc, "W1 发现 1 个 -44.39 m/s 物理无效负值并按规则剔除；W2 未见同类负值或上限持续饱和。")
    add_bullet(doc, "第二季度存在 9 个自然日源覆盖不完整，已在 provenance 中披露；配对结论使用双方同时有效的 10 分钟窗。")

    doc.add_page_break()
    doc.add_heading("4. 配对数据证据", level=1)
    doc.add_picture(str(args.scatter), width=Cm(14.8))
    add_caption(doc, "图 2  W1/W2 配对10分钟平均风速（虚线为两者相等）")
    doc.add_paragraph(
        "大部分点位于 W2=W1 虚线下方，说明“塔顶低于桥面”不是少数异常峰造成，而是贯穿全季的系统性关系。"
        "散点颜色所示方向差与速度比存在分组，表明结果受来流方向影响。"
    )
    doc.add_picture(str(args.sector_ratio), width=Cm(15.5))
    add_caption(doc, "图 3  按 W1 来流方向分区的 W2/W1 平均风速比")
    doc.add_paragraph(
        "各方向扇区的塔顶/桥面风速比约为 0.32～0.65，方向依赖明显。"
        "如果只是简单的统一比例换算错误，风速比通常不会呈现如此清晰的方向变化；该特征更支持局部流场和安装暴露影响。"
    )

    doc.add_page_break()
    doc.add_heading("5. 原因评估与建议", level=1)
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, text in enumerate(("候选原因", "现有证据", "当前判断")):
        set_cell_shading(table.cell(0, idx), BLUE)
        set_cell_text(table.cell(0, idx), text, bold=True, color="FFFFFF")
    rows = (
        ("领导误读", "峰值确实相近，但均值相差约一倍", "排除“纯误读”，疑问成立"),
        ("程序串点/复制", "源点号、文件、风向和配对序列均不同", "基本排除"),
        ("重算算法错误", "自然日与风玫瑰问题已修复，重算后关系仍存在", "不是当前主因"),
        ("桥位局部风场", "速度比随来流方向显著变化；两测点非同一剖面", "较强支持"),
        ("W2 安装/标定问题", "长期系统性偏低且风向差大，但无复制/饱和证据", "不能排除，需现场验证"),
    )
    for r_idx, row in enumerate(rows, 1):
        for c_idx, text in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), text)

    doc.add_heading("建议的现场闭环", level=2)
    add_bullet(doc, "拍照记录 W2 探头与塔柱、横梁、栏杆、避雷针等构件的相对位置，核对是否位于尾流区，测量探头距最近构件距离。")
    add_bullet(doc, "核对 W1/W2 风向零位、坐标定义及安装朝向；用便携式经标定风速仪在 W2 邻近位置做不少于 30～60 分钟并行比测。")
    add_bullet(doc, "核对两套仪器型号、量程、输出协议、换算系数、最近标定证书和维护记录；必要时交换探头或采集通道进行交叉验证。")
    add_bullet(doc, "在台风过程中特别观察 W1/W2 比值是否随风向按本次扇区规律变化；若 W2 在强风下仍固定偏低或出现平台，应优先检修。")

    doc.add_heading("可直接向领导汇报的表述", level=2)
    table = doc.add_table(rows=1, cols=1)
    set_cell_shading(table.cell(0, 0), LIGHT_GRAY)
    paragraph = table.cell(0, 0).paragraphs[0]
    run = paragraph.add_run(
        "“复核后确认，W1 为桥面散索鞍附近测点，W2 为主塔塔顶测点，两者不在同一竖向剖面。"
        "第二季度原始峰值接近，但平均风速分别为 2.74 m/s 和 1.36 m/s，并非真正相同。"
        "两通道数据独立，程序拼接问题已修复；速度比随风向明显变化，说明桥位局部绕流和安装暴露条件影响较大。"
        "但塔顶长期偏低与通常高度效应不一致，暂不能只作气象解释，我们建议立即核查 W2 的遮挡、朝向和标定，并在台风过程中开展并行比测。”"
    )
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)

    doc.add_paragraph()
    source = doc.add_paragraph(
        "依据：洪塘大桥健康监测2026年4—6月周期报工程概况与测点布置；v1.7.36 全自然日重算统计；"
        "W1/W2 配对10分钟诊断结果。结论仅针对现有监测数据，不替代现场仪器检定。"
    )
    source.runs[0].italic = True
    source.runs[0].font.size = Pt(8.5)

    doc.save(args.output)
    return args.output


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-report", type=Path, required=True)
    parser.add_argument("--diagnostic", type=Path, required=True)
    parser.add_argument("--scatter", type=Path, required=True)
    parser.add_argument("--sector-ratio", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    print(build(args))


if __name__ == "__main__":
    main()
