from pathlib import Path
import sys
import tempfile
import unittest

from docx import Document
from docx.oxml import OxmlElement


REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "reporting"))

from build_zhishan_monthly_report import normalize_caption_fields, zhishan_image_replacements  # noqa: E402


def _touch(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(b"placeholder")
    return path


def test_zhishan_report_uses_program_output_images(tmp_path):
    root = tmp_path

    for idx in range(1, 6):
        _touch(root / "PSD_备查" / f"AZ-{idx}" / f"PSD_AZ-{idx}_2026-03-10.jpg")
    for idx in range(1, 9):
        _touch(root / "时程曲线_索力加速度" / f"CF-{idx}_20260302_20260330.jpg")
        _touch(root / "时程曲线_索力加速度_严格最终推荐展示" / f"Strict_CF-{idx}_20260301_20260331.jpg")
    _touch(root / "时程曲线_梁端纵向位移_组图_原始" / "BearingDisp_G1_Orig_20260401-20260430.jpg")
    _touch(root / "时程曲线_梁端纵向位移_组图_滤波" / "BearingDisp_G1_Filt_20260401-20260430.jpg")
    for point in ["CF-1", "CF-2", "CF-6", "CF-7", "CF-8"]:
        _touch(root / "索力时程图" / f"CableForce_{point}_20260301_20260331.jpg")
        _touch(root / "PSD_备查_索力加速度" / point / f"PSD_{point}_2026-03-10.jpg")

    _touch(root / "report_assets_review" / "AZ_PSD_contact_20260609.jpg")
    _touch(root / "report_assets_review" / "Cable_PSD_contact_20260609.jpg")
    _touch(root / "索力时程图_组图" / "CableForce_CF-1-CF-2-CF-3-CF-4.jpg")

    replacements = {
        anchor: [path for path in images if path is not None]
        for anchor, images, _ in zhishan_image_replacements(
            root,
            "2026年3月",
            "2026年3月1日~2026年3月31日",
        )
    }

    assert len(replacements["图 2-8"]) == 5
    assert all(path.parent.parent.name == "PSD_备查" for path in replacements["图 2-8"])

    assert len(replacements["图 2-5"]) == 1
    assert replacements["图 2-5"][0].parent.name == "时程曲线_梁端纵向位移_组图_原始"
    assert len(replacements["图 2-6"]) == 1
    assert replacements["图 2-6"][0].parent.name == "时程曲线_梁端纵向位移_组图_滤波"

    assert len(replacements["图 2-20"]) == 8
    assert all(path.parent.name == "时程曲线_索力加速度" for path in replacements["图 2-20"])

    assert [path.stem.split("_")[1] for path in replacements["图 2-21"]] == ["CF-1", "CF-2", "CF-6", "CF-7", "CF-8"]
    assert all(path.parent.name == "索力时程图" for path in replacements["图 2-21"])

    assert len(replacements["图 2-22"]) == 5
    assert all(path.parent.parent.name == "PSD_备查_索力加速度" for path in replacements["图 2-22"])


def test_normalize_caption_fields_removes_stale_ref_field():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("\u56fe 2-1 Caption")
    instr = OxmlElement("w:instrText")
    instr.text = " REF _RefMissing "
    run._r.append(instr)

    assert paragraph._p.xpath(".//w:instrText")
    assert normalize_caption_fields(doc) == 1
    assert paragraph.text == "\u56fe 2-1 Caption"
    assert not paragraph._p.xpath(".//w:instrText")


def _test_assets_with_temporary_path():
    with tempfile.TemporaryDirectory() as tmp:
        test_zhishan_report_uses_program_output_images(Path(tmp))


def load_tests(loader, tests, pattern):
    del loader, tests, pattern
    return unittest.TestSuite(
        [
            unittest.FunctionTestCase(_test_assets_with_temporary_path),
            unittest.FunctionTestCase(test_normalize_caption_fields_removes_stale_ref_field),
        ]
    )


if __name__ == "__main__":

    with tempfile.TemporaryDirectory() as tmp:
        test_zhishan_report_uses_program_output_images(Path(tmp))
    test_normalize_caption_fields_removes_stale_ref_field()
    print("zhishan report asset smoke ok")
