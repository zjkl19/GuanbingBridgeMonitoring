"""Freeze audited caption text and the final page total in a Hongtang DOCX."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


CAPTION_PATTERN = re.compile(r"^[图表]\s+\d+-\d+\s+")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        keeper = paragraph.runs[0]
        keeper.text = text
        for run in list(paragraph.runs[1:]):
            paragraph._p.remove(run._r)
    else:
        paragraph.add_run(text)


def iter_unique_headers(document: Document):
    seen_parts = set()
    for section in document.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            key = header.part.partname
            if key in seen_parts:
                continue
            seen_parts.add(key)
            yield header


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--total-pages", type=int, required=True)
    parser.add_argument("--audit-output", type=Path, required=True)
    args = parser.parse_args()

    document = Document(args.input)
    staticized = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if CAPTION_PATTERN.match(text):
            replace_paragraph_text(paragraph, text)
            staticized.append({"paragraph_index": index, "text": text})

    page_total_replacements = 0
    for header in iter_unique_headers(document):
        for text_node in header.part.element.xpath(".//w:t"):
            if (text_node.text or "").strip() == "82":
                text_node.text = (text_node.text or "").replace("82", str(args.total_pages))
                page_total_replacements += 1

    if page_total_replacements == 0:
        raise RuntimeError("Expected at least one static header total of 82 pages")
    if not any(item["text"].startswith("图 4-10 ") for item in staticized):
        raise RuntimeError("Expected the audited static caption '图 4-10 地震动时程图'")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    audit = {
        "status": "ok",
        "input_docx": str(args.input.resolve()),
        "output_docx": str(args.output.resolve()),
        "total_pages": args.total_pages,
        "page_total_replacements": page_total_replacements,
        "staticized_caption_count": len(staticized),
        "staticized_captions": staticized,
    }
    args.audit_output.parent.mkdir(parents=True, exist_ok=True)
    args.audit_output.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(args.output)
    print(args.audit_output)


if __name__ == "__main__":
    main()
