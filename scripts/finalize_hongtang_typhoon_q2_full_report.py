"""Finalize the refreshed Hongtang Q2 typhoon report with audited RMS details."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path

from docx import Document


ACCEL_SHA256 = "29E2FFBF4B1BAB2CF528AA2318C4A46B5226E5B49FB47C9B4AF323BC672E1444"
CABLE_SHA256 = "533F598481E07F0BF550626439A927A413F061E9241557DC7ADB8FD685633B5C"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        keeper = paragraph.runs[0]
        keeper.text = text
        for run in list(paragraph.runs[1:]):
            paragraph._p.remove(run._r)
    else:
        paragraph.add_run(text)


def find_body_paragraph(document: Document, prefix: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one body paragraph starting with {prefix!r}, found {len(matches)}")
    return matches[0]


def find_table_paragraphs(document: Document, prefix: str):
    found = []
    seen = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    key = paragraph._p
                    if key in seen:
                        continue
                    seen.add(key)
                    if paragraph.text.strip().startswith(prefix):
                        found.append(paragraph)
    if not found:
        raise RuntimeError(f"No table paragraph starts with {prefix!r}")
    return found


def all_visible_text(document: Document) -> str:
    body = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(body + cells)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--audit-output", type=Path, required=True)
    parser.add_argument("--accel-stats", type=Path, required=True)
    parser.add_argument("--cable-stats", type=Path, required=True)
    args = parser.parse_args()

    if sha256(args.accel_stats) != ACCEL_SHA256:
        raise RuntimeError("accel_stats.xlsx SHA256 mismatch")
    if sha256(args.cable_stats) != CABLE_SHA256:
        raise RuntimeError("cable_accel_stats.xlsx SHA256 mismatch")

    document = Document(args.input)

    method = find_body_paragraph(document, "说明：第4章保留季报既有章节体系")
    set_paragraph_text(
        method,
        "说明：第4章保留季报既有章节体系。因本次台风窗口尚未取得交通称重、结构应变、主塔倾斜和支座变位数据，"
        "4.1至4.4仅保留章节标题，正文、表格和图片留空，不沿用第二季度旧值。风向风速采用原始波形按10min统计；"
        "吊索、主梁及主塔振动同时采用完整20Hz原始波形按自然钟10min计算加速度均方根进行阈值核验，"
        "并采用特征值包络峰值按10min汇总开展登陆前后趋势比较。",
    )

    cable_summary = (
        "选取典型监测数据进行分析，监测结果表明，吊索加速度各测点10min均方根总体最大值为"
        "453mm/s²（0.453m/s²），出现在南侧CS11测点，对应时段为2026年7月11日21:50-22:00；"
        "北侧最大值为424mm/s²（0.424m/s²），出现在CX5测点，时段相同。两者分别为一级阈值"
        "1000mm/s²（1.000m/s²）的45.3%和42.4%，均低于一级超限阈值，未触发相应报警。"
    )
    set_paragraph_text(
        find_body_paragraph(document, "选取典型监测数据进行分析，监测结果表明，吊索加速度"),
        cable_summary,
    )

    cable_trend = find_body_paragraph(document, "台风窗口内，南侧、北侧吊索10min包络峰值")
    set_paragraph_text(
        cable_trend,
        cable_trend.text
        + " 该包络峰值用于趋势筛查，不等同于报警口径的10min加速度均方根。按完整20Hz原始波形核验，"
        "南侧CS11和北侧CX5的最大均方根均出现在登陆前约1小时25分钟，登陆后未形成更高的10min均方根峰值。",
    )

    vibration_summary = (
        "选取典型监测数据进行分析，监测结果表明，主梁和主塔加速度各测点10min均方根最大值分别为"
        "27mm/s²（0.027m/s²）和8mm/s²（0.008m/s²），分别出现在A4和A10-X测点，对应时段均为"
        "2026年7月10日23:20-23:30。按本报告既有主梁/主塔振动评价口径，两者分别为一级阈值"
        "315mm/s²（0.315m/s²）的8.6%和2.5%，均低于一级超限阈值，未触发相应报警。"
    )
    set_paragraph_text(
        find_body_paragraph(document, "选取典型监测数据进行分析，监测结果表明，主梁及主塔加速度"),
        vibration_summary,
    )

    vibration_trend = find_body_paragraph(document, "台风窗口内，主梁、主塔10min包络峰值")
    set_paragraph_text(
        vibration_trend,
        vibration_trend.text
        + " 该包络峰值用于登陆前后相对趋势筛查，不替代报警口径的10min加速度均方根。"
        "按完整20Hz原始波形核验，主梁A4和主塔A10-X最大均方根均出现在监测窗口首个完整时段，"
        "登陆后未形成更高的10min均方根峰值。",
    )

    # Front result panels: keep the Q2 layout while closing the summary/body loop.
    for paragraph in find_table_paragraphs(document, "选取典型监测数据进行分析。监测结果表明，吊索加速度"):
        original = paragraph.text
        force_tail = ""
        marker = "与成桥索力相比"
        if marker in original:
            force_tail = marker + original.split(marker, 1)[1]
        replacement = cable_summary
        if force_tail:
            replacement += force_tail
        set_paragraph_text(paragraph, replacement)

    front_vibration = (
        "台风窗口增补：主梁、主塔10min包络峰值的登陆后/登陆前中位值比分别为1.03和1.04，"
        "未见持续、多测点同步放大。另按完整20Hz原始波形计算自然钟10min均方根，主梁A4最大值为"
        "27mm/s²，主塔A10-X最大值为8mm/s²，对应时段均为07-10 23:20-23:30，"
        "分别占315mm/s²一级阈值的8.6%和2.5%，均未达到一级阈值。"
    )
    for paragraph in find_table_paragraphs(document, "台风窗口增补：台风登陆前后比较中"):
        set_paragraph_text(paragraph, front_vibration)

    conclusion = find_body_paragraph(document, "综合判断：台风窗口内W1未达一级")
    set_paragraph_text(
        conclusion,
        "综合判断：台风窗口内W1桥面和W2塔顶最大10min平均风速分别为5.21m/s和7.60m/s，均未达到一级阈值；"
        "原始风速最大值分别为14.74m/s和17.42m/s，作为短时阵风指标单独报告。主梁、主塔10min加速度均方根"
        "最大值分别为27mm/s²和8mm/s²；南、北侧吊索最大值分别为453mm/s²和424mm/s²，均低于相应一级阈值。"
        "结构响应包络峰值登陆后/登陆前中位值比最大的类别为主塔加速度，比值为1.04；登陆后未见持续、多测点同步放大，"
        "现有证据不支持判定出现台风诱发结构异常。",
    )

    visible = all_visible_text(document)
    forbidden_patterns = {
        "standalone_315m/s²": r"(?<![\d.])315m/s²",
        "standalone_1000m/s²": r"(?<![\d.])1000m/s²",
        "起至至": "起至至",
        "1.05m/s。。": r"1\.05m/s。。",
    }
    hits = {name: len(re.findall(pattern, visible)) for name, pattern in forbidden_patterns.items()}
    if any(hits.values()):
        raise RuntimeError(f"Forbidden text remains: {hits}")

    required = [
        "27mm/s²（0.027m/s²）",
        "8mm/s²（0.008m/s²）",
        "453mm/s²（0.453m/s²）",
        "424mm/s²（0.424m/s²）",
        "2026年7月11日21:50-22:00",
    ]
    missing = [item for item in required if item not in visible]
    # Recompute after edits before deciding whether required strings are present.
    visible = all_visible_text(document)
    missing = [item for item in required if item not in visible]
    if missing:
        raise RuntimeError(f"Required audited report text missing: {missing}")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)

    audit = {
        "status": "ok",
        "input_docx": str(args.input.resolve()),
        "input_sha256": sha256(args.input),
        "output_docx": str(args.output.resolve()),
        "output_sha256": sha256(args.output),
        "source_workbooks": {
            "acceleration": {"path": str(args.accel_stats.resolve()), "sha256": ACCEL_SHA256},
            "cable_acceleration": {"path": str(args.cable_stats.resolve()), "sha256": CABLE_SHA256},
        },
        "method": "20 Hz raw waveform; natural-clock 10-minute RMS; minimum bin coverage 70%",
        "results": {
            "main_girder": {"point": "A4", "value_mm_s2": 27, "interval": "2026-07-10 23:20/23:30", "threshold_ratio_pct": 8.6},
            "main_tower": {"point": "A10-X", "value_mm_s2": 8, "interval": "2026-07-10 23:20/23:30", "threshold_ratio_pct": 2.5},
            "south_cable": {"point": "CS11", "value_mm_s2": 453, "interval": "2026-07-11 21:50/22:00", "threshold_ratio_pct": 45.3},
            "north_cable": {"point": "CX5", "value_mm_s2": 424, "interval": "2026-07-11 21:50/22:00", "threshold_ratio_pct": 42.4},
        },
        "forbidden_text_hits": hits,
    }
    args.audit_output.parent.mkdir(parents=True, exist_ok=True)
    args.audit_output.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(args.output)
    print(args.audit_output)


if __name__ == "__main__":
    main()
