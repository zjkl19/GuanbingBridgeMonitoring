"""Add audited 10-minute acceleration RMS results to the Hongtang typhoon report.

The script deliberately edits a copy of the user-provided report.  It keeps the
existing report structure, adds one compact results table, and writes a small
JSON audit record beside the project copy of the output.
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ACCEL_SHA256 = "29E2FFBF4B1BAB2CF528AA2318C4A46B5226E5B49FB47C9B4AF323BC672E1444"
CABLE_SHA256 = "533F598481E07F0BF550626439A927A413F061E9241557DC7ADB8FD685633B5C"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def set_paragraph_text(paragraph, text: str) -> None:
    """Replace paragraph text while retaining its first run's formatting."""

    if paragraph.runs:
        keeper = paragraph.runs[0]
        keeper.text = text
        for run in list(paragraph.runs[1:]):
            paragraph._p.remove(run._r)
    else:
        paragraph.add_run(text)


def find_paragraph(document: Document, startswith: str):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(startswith)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {startswith!r}, found {len(matches)}")
    return matches[0]


def set_cell_text_like(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)


def replace_text_preserving_runs(paragraph, old: str, new: str) -> None:
    full = paragraph.text
    if old not in full:
        raise RuntimeError(f"Could not find {old!r} in paragraph {full!r}")
    set_paragraph_text(paragraph, full.replace(old, new))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--project-output", type=Path, required=True)
    parser.add_argument("--accel-stats", type=Path, required=True)
    parser.add_argument("--cable-stats", type=Path, required=True)
    args = parser.parse_args()

    if sha256(args.accel_stats) != ACCEL_SHA256:
        raise RuntimeError("accel_stats.xlsx SHA256 does not match the audited source")
    if sha256(args.cable_stats) != CABLE_SHA256:
        raise RuntimeError("cable_accel_stats.xlsx SHA256 does not match the audited source")

    document = Document(args.input)

    # Executive result panel (physical pages 6-7).
    result_cell = document.tables[0].cell(3, 2)
    result_text = result_cell.paragraphs[0].text
    old_vibration = (
        "2、主要结构振动监测\n"
        "登陆后相对登陆前中位水平变化最大的类别为主塔加速度，后/前比为1.04。"
        "各类结构响应按持续性和多测点同步性进行筛查，单个短时尖峰不直接判定为结构异常。"
    )
    new_vibration = (
        "2、主要结构振动监测\n"
        "登陆后相对登陆前中位水平变化最大的类别为主塔加速度，后/前比为1.04。"
        "按原始波形10min均方根口径，主梁、主塔最大值分别为27mm/s²（A4）和8mm/s²（A10-X），"
        "分别占315mm/s²一级阈值的8.6%和2.5%；吊索总体最大值为453mm/s²（CS11），"
        "占1000mm/s²一级阈值的45.3%，均未达到一级阈值。"
    )
    if old_vibration not in result_text:
        raise RuntimeError("Executive vibration summary anchor not found")
    set_paragraph_text(result_cell.paragraphs[0], result_text.replace(old_vibration, new_vibration))

    judgement = document.tables[1].cell(0, 1).paragraphs[0]
    judgement_text = judgement.text
    judgement_addition = (
        "另按20Hz原始波形计算自然钟10min均方根，主梁、主塔和吊索实际最大值分别为"
        "27mm/s²、8mm/s²和453mm/s²，均未达到相应一级阈值。"
    )
    anchor = "主要结构振动响应采用特征包络峰值进行趋势筛查；"
    if anchor not in judgement_text:
        raise RuntimeError("Comprehensive judgement anchor not found")
    set_paragraph_text(judgement, judgement_text.replace(anchor, judgement_addition + anchor))

    # Directly related copy-edit fixes found during the visual/DOM audit.
    event_paragraph = find_paragraph(document, "今年第9号台风")
    replace_text_preserving_runs(event_paragraph, "起至至2026年", "起至2026年")
    wind_paragraph = find_paragraph(document, "W1登陆前后平均风速")
    replace_text_preserving_runs(wind_paragraph, "1.05m/s。。", "1.05m/s。")

    # Add the RMS method/interpretation immediately after the existing Table 4-4.
    blank_analysis = next(
        p
        for p in document.paragraphs
        if not p.text.strip()
        and p.style is not None
        and p.style.name == "洪塘大桥月报正文"
        and p._p.getnext() is find_paragraph(document, "结论")._p
    )
    analysis_text = (
        "表4-4反映各类别10min包络峰值相对登陆前中位水平的变化，用于趋势筛查，"
        "不等同于报警评价采用的10min加速度均方根。为补充报警口径分析，本次另对20Hz原始波形"
        "按自然钟10min分箱计算均方根（单箱有效覆盖率不低于70%），结果见表4‑5。"
        "主梁、主塔及南北侧吊索最大均方根均低于一级阈值；吊索最大窗口发生在台风登陆前约1小时25分钟，"
        "登陆后未形成更高的10min均方根峰值。主塔按本报告既有主梁/主塔振动评价口径进行对比。"
    )
    set_paragraph_text(blank_analysis, analysis_text)

    existing_caption = find_paragraph(document, "表 4-4 主要结构振动响应登陆前后比较")
    caption = copy.deepcopy(existing_caption._p)
    caption_para = existing_caption._parent.add_paragraph()._p
    caption_para.getparent().remove(caption_para)
    # Replace the text in the cloned paragraph without disturbing paragraph properties.
    for child in list(caption):
        if child.tag == qn("w:r"):
            caption.remove(child)
    run = copy.deepcopy(existing_caption.runs[0]._r)
    for text_node in run.findall(".//" + qn("w:t")):
        text_node.text = ""
    first_text = run.find(".//" + qn("w:t"))
    if first_text is None:
        first_text = run._new_t()
    first_text.text = "表 4-5 10min加速度均方根最大值及阈值对比"
    caption.append(run)
    # Keep the compact RMS table together on a fresh page instead of leaving a
    # header and two rows stranded at the foot of the preceding page.
    caption_ppr = caption.get_or_add_pPr()
    page_break_before = OxmlElement("w:pageBreakBefore")
    caption_ppr.append(page_break_before)
    keep_with_next = OxmlElement("w:keepNext")
    caption_ppr.append(keep_with_next)
    blank_analysis._p.addnext(caption)

    # Clone Table 4-4 to inherit widths, borders, shading, fonts and row heights.
    source_table = document.tables[-1]
    rms_table = copy.deepcopy(source_table._tbl)
    caption.addnext(rms_table)
    inserted_table = document.tables[-1]
    table_rows = [
        ["部位", "最大10min RMS\n(mm/s²)", "测点", "统计时段", "一级阈值\n(mm/s²)", "占比/判断"],
        ["主梁加速度", "27", "A4", "07-10 23:20–23:30", "315", "8.6% / 未达一级"],
        ["主塔加速度", "8", "A10-X", "07-10 23:20–23:30", "315", "2.5% / 未达一级"],
        ["南侧吊索", "453", "CS11", "07-11 21:50–22:00", "1000", "45.3% / 未达一级"],
        ["北侧吊索", "424", "CX5", "07-11 21:50–22:00", "1000", "42.4% / 未达一级"],
    ]
    for row, values in zip(inserted_table.rows, table_rows, strict=True):
        for cell, value in zip(row.cells, values, strict=True):
            set_cell_text_like(cell, value)

    # Replace the three conclusion statements with measured values and unambiguous units.
    main_conclusion = find_paragraph(document, "主梁及主塔加速度各测点")
    set_paragraph_text(
        main_conclusion,
        "主梁和主塔加速度各测点10min加速度均方根最大值分别为27mm/s²（0.027m/s²）和"
        "8mm/s²（0.008m/s²），分别出现在A4和A10-X测点，对应时段均为2026年7月10日23:20—23:30。"
        "按本报告既有主梁/主塔振动评价口径，两者分别为315mm/s²一级阈值的8.6%和2.5%，"
        "均低于一级超限阈值，未触发相应报警。",
    )
    cable_conclusion = find_paragraph(document, "吊索加速度各测点")
    set_paragraph_text(
        cable_conclusion,
        "吊索加速度各测点10min加速度均方根总体最大值为453mm/s²（0.453m/s²），"
        "出现在南侧CS11测点，对应时段为2026年7月11日21:50—22:00；北侧最大值为"
        "424mm/s²（0.424m/s²），出现在CX5测点，时段相同。两者分别为一级阈值"
        "1000mm/s²（1.000m/s²）的45.3%和42.4%，均低于一级超限阈值，未触发相应报警。",
    )
    wind_conclusion = find_paragraph(document, "桥面 10min 平均风速最大值")
    set_paragraph_text(
        wind_conclusion,
        "桥面W1测点10min平均风速最大值为5.21m/s，为25m/s一级阈值的20.8%，"
        "低于一级超限阈值，未触发相应报警。",
    )

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.project_output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    document.save(args.project_output)

    audit = {
        "report_input": str(args.input.resolve()),
        "report_output": str(args.output.resolve()),
        "metric_definition": {
            "source": "20 Hz raw waveform",
            "aggregation": "natural-clock 10-minute RMS sqrt(mean(x^2))",
            "minimum_bin_coverage": 0.70,
            "workbook_time_semantics": "bin center; displayed times are interval start-end",
        },
        "source_workbooks": {
            "acceleration": {"path": str(args.accel_stats.resolve()), "sha256": ACCEL_SHA256},
            "cable_acceleration": {"path": str(args.cable_stats.resolve()), "sha256": CABLE_SHA256},
        },
        "results": [
            {"category": "main_girder", "point": "A4", "rms_m_s2": 0.027, "rms_mm_s2": 27, "interval": "2026-07-10 23:20/23:30", "threshold_mm_s2": 315, "ratio_pct": 8.6},
            {"category": "main_tower", "point": "A10-X", "rms_m_s2": 0.008, "rms_mm_s2": 8, "interval": "2026-07-10 23:20/23:30", "threshold_mm_s2": 315, "ratio_pct": 2.5, "threshold_note": "existing report main-girder/tower evaluation convention"},
            {"category": "south_cable", "point": "CS11", "rms_m_s2": 0.453, "rms_mm_s2": 453, "interval": "2026-07-11 21:50/22:00", "threshold_mm_s2": 1000, "ratio_pct": 45.3},
            {"category": "north_cable", "point": "CX5", "rms_m_s2": 0.424, "rms_mm_s2": 424, "interval": "2026-07-11 21:50/22:00", "threshold_mm_s2": 1000, "ratio_pct": 42.4},
        ],
        "unit_correction": "1000 m/s^2 corrected to 1000 mm/s^2 (1.000 m/s^2)",
    }
    audit_path = args.project_output.with_suffix(".audit.json")
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(args.output)
    print(args.project_output)
    print(audit_path)


if __name__ == "__main__":
    main()
