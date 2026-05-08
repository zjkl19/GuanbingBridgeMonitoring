import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from docx import Document  # noqa: E402
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from table_utils import (  # noqa: E402
    apply_report_table_format,
    fill_table,
    iter_unique_cells,
    remove_table_borders,
    set_cell_line_spacing,
    set_header_bold,
    set_table_autofit,
    set_table_borders,
    set_table_column_widths,
    set_table_width,
    style_table,
)


class TestTableUtils(unittest.TestCase):
    def test_style_table_centers_cells_and_sets_autofit(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        style_table(table, autofit=True)
        self.assertTrue(table.autofit)
        self.assertEqual(table.cell(0, 0).vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
        self.assertEqual(table.cell(0, 0).paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_table_layout_and_widths_are_written(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        set_table_autofit(table, False)
        set_table_width(table, 100)
        set_table_column_widths(table, [20, 30])
        tbl_pr = table._tbl.tblPr
        self.assertEqual(tbl_pr.first_child_found_in("w:tblLayout").get(qn("w:type")), "fixed")
        self.assertEqual(tbl_pr.first_child_found_in("w:tblW").get(qn("w:type")), "dxa")

    def test_header_bold_and_fill_table(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        fill_table(table, [["h1", "h2"], ["v1", None]])
        set_header_bold(table)
        self.assertEqual(table.cell(1, 1).text, "/")
        self.assertTrue(table.cell(0, 0).paragraphs[0].runs[0].bold)

    def test_borders_can_be_set_and_removed(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        set_table_borders(table, outer_size_eighth_pt=12, inside_size_eighth_pt=4)
        borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
        self.assertEqual(borders.find(qn("w:top")).get(qn("w:sz")), "12")
        remove_table_borders(table)
        self.assertEqual(borders.find(qn("w:top")).get(qn("w:val")), "nil")

    def test_line_spacing_helpers_handle_merged_cells_once(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).merge(table.cell(1, 1)).text = "merged"

        cells = list(iter_unique_cells(table))
        self.assertEqual(len(cells), 3)
        set_cell_line_spacing(table.cell(0, 0), line_spacing=1.5)
        self.assertEqual(table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing, 1.5)

        apply_report_table_format(table, line_spacing=1.5)
        self.assertEqual(table.cell(1, 0).paragraphs[0].paragraph_format.line_spacing, 1.5)
        self.assertEqual(table.cell(1, 0).vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)


if __name__ == "__main__":
    unittest.main()
