import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

from docx_table_utils import (  # noqa: E402
    find_first_row_index_by_first_cell,
    find_summary_table,
    normalize_cell_text,
)
from build_jlj_monthly_report import SectionContent  # noqa: E402
from jlj_summary import update_summary_table  # noqa: E402


class TestDocxTableUtils(unittest.TestCase):
    def test_find_summary_table_uses_left_column_labels(self):
        doc = Document()
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "普通表格"
        summary = doc.add_table(rows=2, cols=2)
        summary.cell(0, 0).text = "监测结果"
        summary.cell(0, 1).text = "旧内容"
        summary.cell(1, 0).text = "建  议"
        summary.cell(1, 1).text = "旧建议"

        found = find_summary_table(doc)
        self.assertIsNotNone(found)
        self.assertEqual(found.cell(0, 0).text, summary.cell(0, 0).text)
        self.assertEqual(find_first_row_index_by_first_cell(summary, "建议"), 1)
        self.assertEqual(normalize_cell_text("建  议"), "建议")

    def test_jlj_summary_update_does_not_depend_on_table_index(self):
        doc = Document()
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "封面表"
        stale = doc.add_table(rows=1, cols=2)
        stale.cell(0, 0).text = "监测结果"
        stale.cell(0, 1).text = "旧拆分页残留"
        summary = doc.add_table(rows=2, cols=2)
        summary.cell(0, 0).text = "监测结果"
        summary.cell(0, 1).text = "旧监测结果"
        summary.cell(1, 0).text = "建  议"
        summary.cell(1, 1).text = "旧建议"
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "后续表"

        keys = [
            "main_env",
            "main_humidity",
            "main_rainfall",
            "main_wind",
            "main_eq",
            "main_traffic",
            "main_deflection",
            "main_bearing",
            "main_gnss",
            "main_vibration",
            "main_strain",
            "main_crack",
            "main_cable",
            "north_strain",
            "north_bearing",
            "north_tilt",
            "south_strain",
            "south_bearing",
            "south_tilt",
        ]
        section_map = {key: SectionContent(narrative="", summary_sentence=f"{key} summary") for key in keys}

        update_summary_table(doc, section_map, "data summary")

        self.assertEqual(len(summary.rows), 4)
        self.assertEqual(summary.cell(0, 0).text.strip(), "监测结果")
        self.assertEqual(summary.cell(1, 0).text.strip(), "监测结果")
        self.assertEqual(summary.cell(2, 0).text.strip(), "监测结果")
        self.assertEqual(summary.cell(3, 0).text.strip(), "建  议")
        self.assertIn("一、监测系统运行情况", summary.cell(0, 1).text)
        self.assertIn("二、本月监测数据情况", summary.cell(0, 1).text)
        self.assertIn("（转下页）", summary.cell(0, 1).text)
        self.assertIn("（续上页）", summary.cell(1, 1).text)
        self.assertIn("三、监测数据分析结果", summary.cell(0, 1).text)
        self.assertIn("main_env summary", summary.cell(0, 1).text)
        self.assertEqual(doc.tables[0].cell(0, 0).text, "封面表")
        self.assertEqual(doc.tables[1].cell(0, 0).text, "监测结果")
        self.assertEqual(doc.tables[2].cell(0, 0).text, "后续表")

    def test_jlj_summary_updates_cover_table_after_client_info(self):
        doc = Document()
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "封面表"
        cover = doc.add_table(rows=4, cols=5)
        cover.cell(0, 0).text = "委托单位"
        cover.cell(0, 1).text = "名称"
        cover.cell(1, 3).text = "监测时间"
        cover.cell(3, 0).text = "监测结果"
        cover.cell(3, 1).text = "旧首页结果"
        doc.add_paragraph("（续上页）")
        cont1 = doc.add_table(rows=1, cols=2)
        cont1.cell(0, 0).text = "监测结果"
        cont1.cell(0, 1).text = "旧续页1"
        doc.add_paragraph("（续上页）")
        cont2 = doc.add_table(rows=2, cols=2)
        cont2.cell(0, 0).text = "监测结果"
        cont2.cell(0, 1).text = "旧续页2"
        cont2.cell(1, 0).text = "建  议"
        cont2.cell(1, 1).text = "旧建议"
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "后续表"

        keys = [
            "main_env",
            "main_humidity",
            "main_rainfall",
            "main_wind",
            "main_eq",
            "main_traffic",
            "main_deflection",
            "main_bearing",
            "main_gnss",
            "main_vibration",
            "main_strain",
            "main_crack",
            "main_cable",
            "north_strain",
            "north_bearing",
            "north_tilt",
            "south_strain",
            "south_bearing",
            "south_tilt",
        ]
        section_map = {key: SectionContent(narrative="", summary_sentence=f"{key} summary") for key in keys}

        update_summary_table(doc, section_map, "data summary")

        self.assertEqual(doc.tables[1].cell(0, 0).text, "委托单位")
        self.assertIn("一、监测系统运行情况", doc.tables[1].cell(3, 1).text)
        self.assertIn("（转下页）", doc.tables[1].cell(3, 1).text)
        transfer_para = next(
            para for para in doc.tables[1].cell(3, 1).paragraphs
            if para.text.strip() == "（转下页）"
        )
        self.assertEqual(transfer_para.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(transfer_para.paragraph_format.line_spacing, 1.5)
        self.assertFalse(any(run.bold for run in transfer_para.runs))
        self.assertEqual(doc.tables[1].cell(0, 0).paragraphs[0].paragraph_format.line_spacing, 1.5)
        self.assertIn("1.4 风向风速监测", doc.tables[2].cell(0, 1).text)
        self.assertNotIn("（续上页）", doc.tables[2].cell(0, 1).text)
        self.assertIn("2.5 结构应变监测", doc.tables[3].cell(0, 1).text)
        self.assertEqual(doc.tables[4].cell(0, 0).text, "后续表")


if __name__ == "__main__":
    unittest.main()
