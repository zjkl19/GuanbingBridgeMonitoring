import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORTING = ROOT / "reporting"
if str(REPORTING) not in sys.path:
    sys.path.insert(0, str(REPORTING))

from build_zhishan_monthly_report import (  # noqa: E402
    RangeStats,
    latest_nested_psd,
    lowpass_alarm_note,
    _psd_period_tokens,
    update_data_availability,
)


class ZhishanReportGeneratorTests(unittest.TestCase):
    def test_psd_media_selection_uses_report_month_and_never_stale_march(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            folder = root / "PSD_备查" / "AZ-1"
            folder.mkdir(parents=True)
            for name in (
                "PSD_AZ-1_2026-03-10.jpg",
                "PSD_AZ-1_2026-04-01.jpg",
                "PSD_AZ-1_2026-04-10.jpg",
            ):
                (folder / name).write_bytes(b"fixture")

            preferred, month = _psd_period_tokens(
                "2026年4月",
                "2026年4月1日~2026年4月30日",
            )
            selected = latest_nested_psd(
                root,
                "PSD_备查",
                "AZ-1",
                preferred_date_token=preferred,
                month_token=month,
            )

            self.assertIsNotNone(selected)
            self.assertEqual(selected.name, "PSD_AZ-1_2026-04-10.jpg")

    def test_psd_media_selection_falls_back_only_within_report_month(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            folder = root / "PSD_备查" / "AZ-1"
            folder.mkdir(parents=True)
            (folder / "PSD_AZ-1_2026-03-10.jpg").write_bytes(b"stale")
            (folder / "PSD_AZ-1_2026-04-03.jpg").write_bytes(b"current")

            selected = latest_nested_psd(
                root,
                "PSD_备查",
                "AZ-1",
                preferred_date_token="2026-04-10",
                month_token="2026-04",
            )
            self.assertIsNotNone(selected)
            self.assertEqual(selected.name, "PSD_AZ-1_2026-04-03.jpg")

            (folder / "PSD_AZ-1_2026-04-03.jpg").unlink()
            self.assertIsNone(
                latest_nested_psd(
                    root,
                    "PSD_备查",
                    "AZ-1",
                    preferred_date_token="2026-04-10",
                    month_token="2026-04",
                )
            )

    def test_lowpass_alarm_note_uses_configured_point_bound(self) -> None:
        context = {
            "dynamic_lp_point_stats": {
                "SX-5": RangeStats(min_value=317.1, max_value=999.9),
                "SX-6": RangeStats(min_value=-15.4, max_value=54.9),
            },
            "strain_alarm_bounds": {
                "SX-5": (-252.0, 405.0),
                "SX-6": (-252.0, 405.0),
            },
        }

        note = lowpass_alarm_note(context)

        self.assertIn("SX-5", note)
        self.assertIn("999.9με", note)
        self.assertIn("+405.0με", note)
        self.assertIn("不宜仅据此直接判定为结构异常", note)
        self.assertNotIn("SX-6", note)

    def test_source_quality_note_is_written_to_coverage_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            result_root = Path(tmp)
            (result_root / "2026-06-01").mkdir()
            doc = Document()
            doc.add_paragraph("本月报告分析数据覆盖旧内容")
            doc.add_paragraph("按本月监测数据获取情况表统计旧内容")
            doc.add_paragraph("本月持续开展监测系统运行维护工作旧内容")
            doc.add_paragraph("软件线上检查维护旧内容")
            note = (
                "源数据完整性说明：缺少2026-07-01相邻滚动文件，"
                "6月30日尾段不完整；仅使用实际获取数据，不作伪造补齐。"
            )

            update_data_availability(
                doc,
                result_root,
                "2026年6月",
                "2026年6月1日~2026年6月30日",
                note,
            )

            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertIn(note, text)
            self.assertIn("实际有效数据为1天", text)


if __name__ == "__main__":
    unittest.main()
