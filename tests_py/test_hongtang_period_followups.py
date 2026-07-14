import json
import sys
import hashlib
import tempfile
import unittest
from datetime import date
from pathlib import Path

from docx import Document
from openpyxl import Workbook

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from build_monthly_report import (  # noqa: E402
    build_eq_section,
    build_overview_items,
    build_wind_section,
    format_rms_threshold,
    format_rms_value,
    rms_threshold_assessment,
    normalize_eq_peak_key,
    parse_wind_summary,
    replace_wind_speed_caption,
    update_wind_table,
)
from build_period_report import (  # noqa: E402
    HIGHFREQ_MODULES,
    apply_period_maintenance_log,
    collect_highfreq_missing_events,
    collect_highfreq_provenance_events,
    collect_declared_data_coverage_events,
    load_data_coverage_audit,
)
from analysis_manifest import pinned_analysis_manifest_scope  # noqa: E402
from build_quarterly_wim_sample import format_load_limit_text  # noqa: E402


class TestHongtangPeriodFollowups(unittest.TestCase):
    def test_wind_summary_parser_reads_all_report_table_fields(self):
        with tempfile.TemporaryDirectory() as td:
            summary = Path(td) / "W1_windrose_summary.txt"
            summary.write_text(
                "风玫瑰简要结论（W1）\n"
                "平均风向: 198.3°\n"
                "主导风向: 180.0°-202.5°，占比 10.2%\n"
                "平均风速: 2.75 m/s\n"
                "最大风速: 12.25 m/s\n"
                "主要风速等级: 2-4 m/s（依据：全样本风速分级占比最高）\n",
                encoding="utf-8",
            )

            self.assertEqual(
                parse_wind_summary(summary),
                {
                    "mean_dir": "198.3°",
                    "dominant_dir": "180.0°-202.5°，占比 10.2%",
                    "mean_speed": "2.75",
                    "max_speed": "12.25",
                    "main_grade": "2-4 m/s",
                },
            )

    def test_strict_mode_rejects_unlisted_data_coverage_audit(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            run_logs = root / "run_logs"
            run_logs.mkdir()
            audit = run_logs / "data_coverage_audit.json"
            audit.write_text('{"schema_version":1,"events":[]}', encoding="utf-8")
            manifest = run_logs / "analysis_manifest.json"
            manifest.write_text('{"status":"ok"}', encoding="utf-8")
            manifest_hash = hashlib.sha256(manifest.read_bytes()).hexdigest().upper()

            with pinned_analysis_manifest_scope(
                manifest,
                manifest_hash,
                require_source_provenance=True,
                result_root=root,
            ):
                with self.assertRaisesRegex(ValueError, "unlisted data coverage audit"):
                    load_data_coverage_audit(root)

    def test_highfreq_health_check_accepts_mat_only_cache(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            raw_dir = root / "2026-06-30" / "waveform"
            cache_dir = raw_dir / "cache"
            cache_dir.mkdir(parents=True)
            (cache_dir / "A1_20260705224345190.mat").write_bytes(b"MAT-cache")
            cfg = {
                "subfolders": {"acceleration": "waveform"},
                "points": {"acceleration": ["A1"]},
                "file_patterns": {
                    "acceleration": {
                        "default": ["{point}_*.csv"],
                        "per_point": {"A1": "{point}_174.csv"},
                    }
                },
            }

            events = collect_highfreq_missing_events(
                cfg, root, date(2026, 6, 30), date(2026, 6, 30)
            )

            self.assertEqual(events, [])

    def test_highfreq_health_check_uses_manifest_natural_day_completeness(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            raw_dir = root / "2026-06-30" / "waveform"
            cache_dir = raw_dir / "cache"
            cache_dir.mkdir(parents=True)
            (cache_dir / "A1_20260705224345190.mat").write_bytes(b"MAT-cache")
            sidecar = root / "plots" / "A1_20260401_20260630.plot.json"
            sidecar.parent.mkdir()
            sidecar.write_text(
                json.dumps(
                    {
                        "series": {
                            "point_id": "A1",
                            "source": {
                                "incomplete_days": ["2026-06-30"],
                                "missing_required_sources": ["2026-07-01"],
                            },
                        }
                    }
                ),
                encoding="utf-8",
            )
            run_logs = root / "run_logs"
            run_logs.mkdir()
            (run_logs / "analysis_manifest_20260713_230000.json").write_text(
                json.dumps(
                    {
                        "status": "ok",
                        "module_results": [
                            {
                                "key": "acceleration",
                                "status": "ok",
                                "artifacts": [{"path": str(sidecar)}],
                            }
                        ],
                    }
                ),
                encoding="utf-8",
            )
            cfg = {
                "subfolders": {"acceleration": "waveform"},
                "points": {"acceleration": ["A1"]},
                "file_patterns": {
                    "acceleration": {
                        "default": ["{point}_*.csv"],
                        "per_point": {"A1": "{point}_174.csv"},
                    }
                },
            }

            provenance_events = collect_highfreq_provenance_events(
                root, date(2026, 6, 30), date(2026, 6, 30)
            )
            all_events = collect_highfreq_missing_events(
                cfg, root, date(2026, 6, 30), date(2026, 6, 30)
            )

            self.assertEqual(provenance_events, all_events)
            self.assertEqual(len(all_events), 1)
            self.assertEqual(all_events[0]["module"], HIGHFREQ_MODULES["acceleration"])
            self.assertEqual(all_events[0]["points"], ["A1"])
            self.assertIn("2026-06-30", all_events[0]["range"])
            self.assertTrue(all_events[0]["reason"])

    def test_highfreq_health_check_includes_independent_timestamp_gap_audit(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            run_logs = root / "run_logs"
            run_logs.mkdir()
            (run_logs / "data_coverage_audit.json").write_text(
                json.dumps(
                    {
                        "schema_version": 1,
                        "events": [
                            {
                                "module": "wind",
                                "points": ["W1", "W2"],
                                "start": "2026-06-29 08:46:26.532",
                                "end": "2026-06-29 09:48:30.684",
                                "reason": "timestamp gap",
                            }
                        ],
                    }
                ),
                encoding="utf-8",
            )

            events = collect_declared_data_coverage_events(
                root, date(2026, 4, 1), date(2026, 6, 30)
            )

            self.assertEqual(len(events), 1)
            self.assertEqual(events[0]["module"], HIGHFREQ_MODULES["wind"])
            self.assertEqual(events[0]["points"], ["W1", "W2"])
            self.assertIn("2026-06-29 08:46:26", events[0]["range"])
            self.assertIn("2026-06-29 09:48:30", events[0]["range"])
            self.assertEqual(events[0]["reason"], "timestamp gap")

    def test_invalid_data_coverage_audit_fails_closed(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            run_logs = root / "run_logs"
            run_logs.mkdir()
            (run_logs / "data_coverage_audit.json").write_text("{broken", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "Invalid data coverage audit"):
                load_data_coverage_audit(root)

    def test_data_coverage_audit_records_exact_file_hash(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            path = root / "run_logs" / "data_coverage_audit.json"
            path.parent.mkdir()
            path.write_text('{"schema_version": 1, "events": []}', encoding="utf-8")

            audit = load_data_coverage_audit(root)

            self.assertEqual(audit["path"], str(path))
            self.assertEqual(audit["bytes"], path.stat().st_size)
            self.assertEqual(
                audit["sha256"],
                hashlib.sha256(path.read_bytes()).hexdigest().upper(),
            )

    def test_boundary_disclosure_is_renderable_for_all_highfreq_modules(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            run_logs = root / "run_logs"
            run_logs.mkdir()
            statement = "2026-06-30 about 09:00 tail is incomplete"
            (run_logs / "data_coverage_audit.json").write_text(
                json.dumps(
                    {
                        "schema_version": 1,
                        "events": [],
                        "boundary_disclosure": {
                            "affected_day": "2026-06-30",
                            "statement": statement,
                        },
                    }
                ),
                encoding="utf-8",
            )

            events = collect_declared_data_coverage_events(
                root, date(2026, 4, 1), date(2026, 6, 30)
            )

            self.assertEqual(len(events), 4)
            self.assertEqual({event["module"] for event in events}, set(HIGHFREQ_MODULES.values()))
            self.assertTrue(all(event["range"] == "2026-06-30" for event in events))
            self.assertTrue(all(event["reason"] == statement for event in events))

    def test_wim_load_limit_wording_distinguishes_equal_from_exceeded(self):
        self.assertEqual(
            format_load_limit_text(30.0, 30.0, 40.0),
            "达到1.5倍设计车辆荷载30t，未达到2.0倍设计车辆荷载40t",
        )
        self.assertEqual(
            format_load_limit_text(30.01, 30.0, 40.0),
            "超过1.5倍设计车辆荷载30t，未达到2.0倍设计车辆荷载40t",
        )
        self.assertEqual(format_load_limit_text(40.0, 30.0, 40.0), "达到2.0倍设计车辆荷载40t")

    def test_eq_section_maps_base_point_id_and_component_to_direction(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            stats = root / "stats"
            assets = root / "assets"
            stats.mkdir()
            assets.mkdir()

            wb = Workbook()
            ws = wb.active
            ws.append(["PointID", "Component", "Peak", "PeakTime"])
            ws.append(["EQ", "X", 0.005, "2026-04-20 06:42:09"])
            ws.append(["EQ", "Y", 0.018, "2026-06-17 03:22:18"])
            ws.append(["EQ", "Z", 0.019, "2026-06-17 01:02:38"])
            wb.save(stats / "eq_stats.xlsx")

            cfg = {
                "points": {"eq": ["EQ-X", "EQ-Y", "EQ-Z"]},
                "per_point": {
                    "eq": {
                        "EQ-X": {"alarm_levels": [0.1]},
                        "EQ-Y": {"alarm_levels": [0.1]},
                        "EQ-Z": {"alarm_levels": [0.1]},
                    }
                },
            }
            section = build_eq_section(cfg, stats, None, root, assets)

            self.assertEqual(normalize_eq_peak_key("EQ", "X"), "EQ-X")
            self.assertAlmostEqual(section["horizontal_peak"], 0.018)
            self.assertAlmostEqual(section["vertical_peak"], 0.019)
            self.assertIn("0.018m/s²", section["summary"])
            self.assertIn("0.019m/s²", section["summary"])

    def test_period_maintenance_log_replaces_table_with_q2_rows(self):
        doc = Document()
        doc.add_paragraph("表 1-2 健康监测系统维护日志表")
        table = doc.add_table(rows=19, cols=3)
        table.cell(0, 0).text = "序号"
        table.cell(0, 1).text = "维护日期"
        table.cell(0, 2).text = "维护类型"
        for idx in range(1, 19):
            table.cell(idx, 0).text = str(idx)
            table.cell(idx, 1).text = f"2026-03-{idx:02d}"
            table.cell(idx, 2).text = "旧记录"

        apply_period_maintenance_log(doc, date(2026, 4, 1), date(2026, 6, 30))

        self.assertEqual(len(table.rows), 16)
        self.assertEqual([cell.text for cell in table.rows[0].cells], ["序号", "维护日期", "维护类型"])
        self.assertEqual(table.cell(1, 1).text, "2026-04-06")
        self.assertIn("基康采集设备", table.cell(2, 2).text)
        self.assertEqual(table.cell(15, 1).text, "2026-06-29")
        self.assertNotIn("旧记录", "\n".join(cell.text for row in table.rows for cell in row.cells))

    def test_overview_uses_accepted_q2_wording(self):
        manifest = {
            "sections": {
                "wim": {"enabled": False},
                "traffic": {},
                "strain": {"enabled": False, "available": False},
                "tilt": {"enabled": False, "available": False},
                "bearing_displacement": {"enabled": False, "available": False},
                "cable_force": {
                    "enabled": True,
                    "available": True,
                    "accel_available": True,
                    "force_available": True,
                    "max_abs": 2.44,
                    "max_rms": 0.37,
                    "min_change": -7.42,
                    "max_change": 1.77,
                },
                "vibration": {
                    "enabled": True,
                    "available": True,
                    "max_abs": 0.35,
                    "max_rms": 0.04,
                },
                "wind": {"enabled": True, "available": True, "max_10min": 5.46},
                "eq": {"enabled": False, "available": False},
            }
        }

        items = build_overview_items(manifest)
        cable_text = items["吊索索力监测"][0]
        wind_text = items["风向风速监测"][0]
        vibration_text = items["主梁、主塔振动监测"][0]

        self.assertIn("监测结果表明，吊索加速度", cable_text)
        self.assertIn("370mm/s²（0.370m/s²）", cable_text)
        self.assertIn("1000mm/s²（1.000m/s²）", cable_text)
        self.assertNotIn("1000m/s²", cable_text)
        self.assertIn("变化幅度均在10%以内", cable_text)
        self.assertNotIn("监测结果表明吊索加速度", cable_text)
        self.assertNotIn("与成桥索力相比变化范围在10%以内", cable_text)
        self.assertIn("桥面测点W1的10min平均风速", wind_text)
        self.assertIn("40mm/s²（0.040m/s²）", vibration_text)
        self.assertIn("315mm/s²（0.315m/s²）", vibration_text)

    def test_rms_format_preserves_mm_per_second_squared_precision(self):
        self.assertEqual(format_rms_value(0.027), "27mm/s²（0.027m/s²）")
        self.assertEqual(format_rms_value(0.453), "453mm/s²（0.453m/s²）")
        self.assertEqual(format_rms_threshold(315), "315mm/s²（0.315m/s²）")
        self.assertEqual(format_rms_threshold(1000), "1000mm/s²（1.000m/s²）")

    def test_rms_threshold_assessment_never_calls_an_exceedance_below_threshold(self):
        self.assertIn("低于一级超限阈值", rms_threshold_assessment(0.37, 1000))
        self.assertIn("未达到一级超限条件", rms_threshold_assessment(0.37, 1000))
        exceeded = rms_threshold_assessment(2.468, 1000)
        self.assertIn("超过一级超限阈值", exceeded)
        self.assertIn("已达到一级超限条件", exceeded)
        self.assertNotIn("低于", exceeded)

    def test_wind_section_uses_accepted_spacing_and_caption(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            stats = root / "stats"
            stats.mkdir()

            wb = Workbook()
            ws = wb.active
            ws.append(["PointID", "Mean10minMax", "MeanSpeed", "MaxSpeed"])
            ws.append(["W1", 5.46, 2.28, 9.25])
            wb.save(stats / "wind_stats.xlsx")

            cfg = {"points": {"wind": ["W1"]}}
            section = build_wind_section(cfg, stats, None, root, root / "assets")

            self.assertIn("桥面测点W1的10min平均风速最大值为5.46m/s", section["summary"])
            self.assertIn("W1瞬时最大风速为9.25m/s", section["summary"])
            self.assertEqual(section["speed_caption"], "W1桥面与W2塔顶10min平均风速时程图")

    def test_wind_speed_caption_prefers_v1738_and_falls_back_to_legacy(self):
        accepted = "W1桥面与W2塔顶10min平均风速时程图"

        current_doc = Document()
        current = current_doc.add_paragraph(accepted)
        legacy = current_doc.add_paragraph("图 4-15 桥面 10min 平均风速时程图")
        replace_wind_speed_caption(current_doc, accepted)
        self.assertEqual(current.text, accepted)
        self.assertEqual(legacy.text, "图 4-15 桥面 10min 平均风速时程图")

        legacy_doc = Document()
        legacy_only = legacy_doc.add_paragraph("图 4-15 桥面 10min 平均风速时程图")
        replace_wind_speed_caption(legacy_doc, accepted)
        self.assertEqual(legacy_only.text, accepted)

    def test_wind_section_distinguishes_locations_and_explains_comparison_limit(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            stats = root / "stats"
            stats.mkdir()

            wb = Workbook()
            ws = wb.active
            ws.append(["PointID", "Mean10minMax", "MeanSpeed", "MaxSpeed"])
            ws.append(["W1", 6.4, 2.74, 10.1])
            ws.append(["W2", 4.2, 1.37, 12.0])
            wb.save(stats / "wind_stats.xlsx")

            cfg = {
                "points": {"wind": ["W1", "W2"]},
                "per_point": {
                    "wind": {
                        "W1": {"location": "右幅桥面12号墩附近散索鞍保护罩"},
                        "W2": {"location": "13号主塔塔顶"},
                    }
                },
            }
            section = build_wind_section(cfg, stats, None, root, root / "assets")

            self.assertIn("桥面测点W1的10min平均风速最大值为6.40m/s", section["summary"])
            self.assertIn("塔顶测点W2的10min平均风速最大值为4.20m/s", section["summary"])
            self.assertIn("两者并非同一竖向测风剖面", section["summary"])
            self.assertIn("塔顶/桥面平均风速比约为50.0%", section["summary"])
            self.assertIn("不能按一般大气边界层的高度增风规律作简单对比", section["summary"])
            self.assertIn("不支持直接判定仪器故障或结构异常", section["summary"])
            self.assertEqual(section["deck_max_10min"], 6.4)
            self.assertEqual(section["tower_max_10min"], 4.2)

    def test_wind_table_clears_stale_template_rows_when_point_missing(self):
        doc = Document()
        table = doc.add_table(rows=3, cols=6)
        headers = ["测点", "平均风向", "主导风向", "平均风速", "最大风速", "主要风速等级"]
        for idx, value in enumerate(headers):
            table.cell(0, idx).text = value
        for row_idx, point_id in ((1, "W1"), (2, "W2")):
            table.cell(row_idx, 0).text = point_id
            for col_idx in range(1, 6):
                table.cell(row_idx, col_idx).text = "旧值"

        update_wind_table(
            doc,
            [{
                "PointID": "W1",
                "mean_dir": "270°",
                "dominant_dir": "270°-292.5°",
                "mean_speed": 2.74,
                "max_speed": 10.1,
                "mean10min_max": 6.4,
                "main_grade": "2-4 m/s",
            }],
        )

        self.assertEqual(len(table.columns), 7)
        self.assertEqual(table.cell(1, 3).text, "2.74")
        self.assertEqual(table.cell(0, 4).text, "瞬时最大风速（m/s）")
        self.assertEqual(table.cell(0, 5).text, "10min平均风速最大值（m/s）")
        self.assertEqual(table.cell(1, 5).text, "6.40")
        self.assertEqual(table.cell(1, 6).text, "2-4 m/s")
        self.assertEqual([table.cell(2, idx).text for idx in range(1, 7)], ["", "", "", "", "", ""])

    def test_wind_table_schema_upgrade_is_idempotent(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=6)
        headers = ["测点", "平均风向", "主导风向", "平均风速", "最大风速", "主要风速等级"]
        for idx, value in enumerate(headers):
            table.cell(0, idx).text = value
        table.cell(1, 0).text = "W1"
        table.cell(1, 5).text = "2-4 m/s"

        update_wind_table(
            doc,
            [{
                "PointID": "W1",
                "mean_dir": "270°",
                "dominant_dir": "270°-292.5°",
                "mean_speed": 2.74,
                "max_speed": 10.1,
                "mean10min_max": 6.4,
                "main_grade": "2-4 m/s",
            }],
        )
        update_wind_table(
            doc,
            [{
                "PointID": "W1",
                "mean_dir": "270°",
                "dominant_dir": "270°-292.5°",
                "mean_speed": 2.74,
                "max_speed": 10.1,
                "mean10min_max": 6.4,
                "main_grade": "2-4 m/s",
            }],
        )

        self.assertEqual(len(table.columns), 7)
        self.assertEqual(table.cell(1, 5).text, "6.40")


if __name__ == "__main__":
    unittest.main()
