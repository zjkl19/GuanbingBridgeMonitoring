import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from docx import Document  # noqa: E402
from jlj_patrol import (  # noqa: E402
    insert_docx_body_after_heading,
    patrol_report_matches_period,
    patrol_source_availability_record,
    period_label_month,
    replace_patrol_section_with_note,
    replace_patrol_report_dates,
    resolve_patrol_report_source,
)


class TestJljPatrol(unittest.TestCase):
    def test_replace_patrol_report_dates_uses_target_month(self):
        text = "2026年2月09日，02月巡查，2026-02-09，2026.02.09"
        out = replace_patrol_report_dates(text, target_month=3)
        self.assertIn("2026年3月09日", out)
        self.assertIn("3月巡查", out)
        self.assertIn("2026-03-09", out)
        self.assertIn("2026.03.09", out)

    def test_period_label_month(self):
        self.assertEqual(period_label_month("2026年4月份"), 4)
        self.assertEqual(period_label_month("无月份", default_month=3), 3)

    def test_insert_docx_body_after_heading_clears_old_body_without_falsifying_dates(self):
        target = Document()
        target.add_heading("桥梁人工巡查结果", level=1)
        target.add_paragraph("旧巡查内容")
        target.add_heading("后续章节", level=1)

        source = Document()
        source.add_paragraph("2026年2月09日巡查")
        source.add_paragraph("附件：现场照片")

        with self.subTest("insert"):
            tmp = Path("tmp") / "test_jlj_patrol_source.docx"
            tmp.parent.mkdir(exist_ok=True)
            source.save(tmp)
            self.assertTrue(insert_docx_body_after_heading(target, "桥梁人工巡查结果", tmp, target_month=3))
            text = "\n".join(p.text for p in target.paragraphs)
            self.assertIn("2026年2月09日巡查", text)
            self.assertNotIn("2026年3月09日巡查", text)
            self.assertIn("附件：现场照片", text)
            self.assertNotIn("旧巡查内容", text)
            tmp.unlink(missing_ok=True)

    def test_patrol_period_matching_and_missing_note(self):
        target = Document()
        target.add_heading("桥梁人工巡查结果", level=1)
        target.add_paragraph("2026年3月旧巡查")
        target.add_heading("后续章节", level=1)

        source = Document()
        source.add_paragraph("2026年05月09日上午巡查")
        tmp = Path("tmp") / "test_jlj_patrol_period.docx"
        tmp.parent.mkdir(exist_ok=True)
        source.save(tmp)
        try:
            self.assertTrue(patrol_report_matches_period(tmp, 2026, 5))
            self.assertFalse(patrol_report_matches_period(tmp, 2026, 3))
            replace_patrol_section_with_note(target, "桥梁人工巡查结果")
            text = "\n".join(p.text for p in target.paragraphs)
            self.assertIn("本期巡查资料未提供", text)
            self.assertNotIn("2026年3月旧巡查", text)
        finally:
            tmp.unlink(missing_ok=True)

    def test_patrol_source_availability_records_verified_source_hash_and_period(self):
        source = Document()
        source.add_paragraph("2026年05月09日上午巡查")
        tmp = Path("tmp") / "test_jlj_patrol_availability.docx"
        tmp.parent.mkdir(exist_ok=True)
        source.save(tmp)
        try:
            record = patrol_source_availability_record(
                tmp,
                required=False,
                target_year=2026,
                target_month=5,
                action="verified_source_inserted",
            )
            self.assertEqual(record["status"], "available")
            self.assertEqual(record["target_period"], "2026-05")
            self.assertEqual(record["source_period"], "2026-05")
            self.assertEqual(record["source"], str(tmp.resolve()))
            self.assertEqual(len(record["source_sha256"]), 64)
            self.assertEqual(record["action"], "verified_source_inserted")
        finally:
            tmp.unlink(missing_ok=True)

    def test_patrol_source_availability_records_nonblocking_absence(self):
        record = patrol_source_availability_record(
            None,
            required=False,
            target_year=2026,
            target_month=5,
            action="template_content_cleared_and_note_inserted",
        )
        self.assertEqual(
            record,
            {
                "required": False,
                "status": "not_available",
                "target_period": "2026-05",
                "source": "",
                "source_sha256": "",
                "source_period": "",
                "action": "template_content_cleared_and_note_inserted",
            },
        )

    def test_explicit_patrol_source_from_another_month_fails_closed(self):
        source = Document()
        source.add_paragraph("2026年03月09日巡查")
        tmp = Path("tmp") / "test_jlj_patrol_wrong_period.docx"
        tmp.parent.mkdir(exist_ok=True)
        source.save(tmp)
        try:
            with self.assertRaisesRegex(ValueError, "period does not match"):
                resolve_patrol_report_source(
                    Path("template.docx"),
                    tmp,
                    target_year=2026,
                    target_month=5,
                )
        finally:
            tmp.unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
