from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parents[1]
REPORTING_DIR = REPO_ROOT / "reporting"
if str(REPORTING_DIR) not in sys.path:
    sys.path.insert(0, str(REPORTING_DIR))

from build_quarterly_wim_sample import (  # noqa: E402
    add_caption_paragraph_before,
    capture_paragraph_template,
    overload_counts_text,
)
from build_period_report import _pattern_for_point, convert_static_captions_to_auto_number  # noqa: E402


def paragraph_fields(paragraph) -> str:
    return "|".join(node.text or "" for node in paragraph._p.iter() if node.tag.endswith("}instrText"))


def add_bookmark(paragraph, bookmark_id: str, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


class WimAutoCaptionTests(unittest.TestCase):
    def test_wim_table_and_figure_captions_use_word_seq_fields(self) -> None:
        doc = Document()
        template_para = doc.add_paragraph("表 4-1 模板")
        anchor = doc.add_paragraph("anchor")
        template = capture_paragraph_template(template_para)

        table_caption = add_caption_paragraph_before(anchor, "表 4-1 2026年第二季度交通状况分月统计表", template, "table")
        figure_caption = add_caption_paragraph_before(anchor, "图 4-1 2026年4月桥梁交通流参数分析", template, "figure")
        continued_caption = add_caption_paragraph_before(anchor, "续表 4-3（轴重单位：kg）", template, "table_continued")

        self.assertIn("STYLEREF 1 \\s", paragraph_fields(table_caption))
        self.assertIn("SEQ 表 \\* ARABIC \\s 1", paragraph_fields(table_caption))
        self.assertIn("SEQ 图 \\* ARABIC \\s 1", paragraph_fields(figure_caption))
        self.assertIn("SEQ 表 \\c", paragraph_fields(continued_caption))
        self.assertIn("2026年第二季度交通状况分月统计表", table_caption.text)
        self.assertIn("2026年4月桥梁交通流参数分析", figure_caption.text)

        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "captions.docx"
            doc.save(out)
            reopened = Document(str(out))
            fields = "\n".join(paragraph_fields(p) for p in reopened.paragraphs)
            self.assertIn("SEQ 表 \\* ARABIC \\s 1", fields)
            self.assertIn("SEQ 图 \\* ARABIC \\s 1", fields)
            self.assertIn("SEQ 表 \\c", fields)

    def test_overload_counts_show_two_threshold_levels_explicitly(self) -> None:
        text = overload_counts_text(34, 0, 1, 0)

        self.assertIn("总重1.5/2.0倍：34/0", text)
        self.assertIn("轴重1.5/2.0倍：1/0", text)

    def test_period_report_file_pattern_accepts_per_point_lists(self) -> None:
        cfg = {
            "file_patterns": {
                "wind_speed": {
                    "per_point": {
                        "W1": ["*{file_id}*.csv", "*fallback*.csv"],
                    }
                }
            }
        }

        self.assertEqual(_pattern_for_point(cfg, "wind_speed", "W1", file_id="20260630"), "*20260630*.csv")

    def test_static_period_captions_are_converted_to_word_seq_fields(self) -> None:
        doc = Document()
        figure_caption = doc.add_paragraph("图 4-4 主梁各截面位置应变时程曲线图")
        table_caption = doc.add_paragraph("表 4-1 2026年第二季度交通状况分月统计表")
        continued_caption = doc.add_paragraph("续表 4-3（轴重单位：kg）")
        body = doc.add_paragraph("监测结果如图 4-4 所示。")

        count = convert_static_captions_to_auto_number(doc)

        self.assertEqual(count, 3)
        self.assertIn("SEQ 图 \\* ARABIC \\s 1", paragraph_fields(figure_caption))
        self.assertIn("SEQ 表 \\* ARABIC \\s 1", paragraph_fields(table_caption))
        self.assertIn("SEQ 表 \\c", paragraph_fields(continued_caption))
        self.assertEqual(paragraph_fields(body), "")


    def test_static_caption_conversion_preserves_cross_reference_bookmarks(self) -> None:
        doc = Document()
        caption = doc.add_paragraph("\u56fe 4-6 tower strain time history")
        add_bookmark(caption, "42", "_Ref4508")

        count = convert_static_captions_to_auto_number(doc)

        self.assertEqual(count, 1)
        xml = caption._p.xml
        self.assertIn("_Ref4508", xml)
        self.assertLess(xml.find("bookmarkStart"), xml.find("SEQ"))
        self.assertLess(xml.find("SEQ"), xml.find("bookmarkEnd"))
        self.assertLess(xml.find("bookmarkEnd"), xml.find("tower strain"))


if __name__ == "__main__":
    unittest.main()
