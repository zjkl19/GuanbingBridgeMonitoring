import sys
import tempfile
import unittest
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REPORTING = Path(__file__).resolve().parents[1] / "reporting"
if str(REPORTING) not in sys.path:
    sys.path.insert(0, str(REPORTING))

from build_hongtang_typhoon_brief import (  # noqa: E402
    aggregate_entry,
    parse_data_line,
    point_channel,
    ten_minute_bin,
)
from build_hongtang_typhoon_template_report import (  # noqa: E402
    add_numbered_heading,
    remove_empty_section_break_before,
    wind_alarm_status,
)


class HongtangTyphoonBriefTests(unittest.TestCase):
    def test_time_bin_and_channel_mapping(self):
        value = datetime(2026, 7, 11, 8, 59, 59, 123000)
        self.assertEqual(ten_minute_bin(value), datetime(2026, 7, 11, 8, 50))
        self.assertEqual(point_channel("A10-X"), 156)
        self.assertEqual(point_channel("CS12"), 165)
        self.assertEqual(point_channel("CX3"), 120)

    def test_parse_data_line(self):
        parsed = parse_data_line("2026-07-11 08:00:00.125,1.25\n")
        self.assertEqual(parsed, (datetime(2026, 7, 11, 8, 0, 0, 125000), 1.25))
        self.assertIsNone(parse_data_line("开始时间: 2026-07-11 08:00:00"))
        self.assertIsNone(parse_data_line("bad,row"))

    def test_utf16_zip_entry_is_aggregated(self):
        content = (
            "开始时间: 2026-07-11 08:00:00\n"
            "序列号: 1\n"
            "通道号: 2\n"
            "2026-07-11 08:00:00.000,1.0\n"
            "2026-07-11 08:09:59.000,3.0\n"
            "2026-07-11 08:10:00.000,-1.0\n"
        ).encode("utf-16")
        with tempfile.TemporaryDirectory() as temp_dir:
            zip_path = Path(temp_dir) / "sample.zip"
            with ZipFile(zip_path, "w") as archive:
                archive.writestr("风速_162.csv", content)
            with ZipFile(zip_path) as archive:
                bins, audit = aggregate_entry(
                    archive,
                    "风速_162.csv",
                    key="W1_speed",
                    zip_path=zip_path,
                    validator=lambda value: value >= 0,
                )
        self.assertEqual(audit.rows, 2)
        self.assertEqual(audit.rejected_rows, 1)
        self.assertEqual(bins[datetime(2026, 7, 11, 8, 0)].mean, 2.0)
        self.assertEqual(
            bins[datetime(2026, 7, 11, 8, 0)].maximum_time,
            datetime(2026, 7, 11, 8, 9, 59),
        )

    def test_wind_alarm_status_uses_ten_minute_thresholds(self):
        self.assertEqual(wind_alarm_status(24.99), "未达一级")
        self.assertEqual(wind_alarm_status(25.0), "达到一级")
        self.assertEqual(wind_alarm_status(29.92), "达到二级")
        self.assertEqual(wind_alarm_status(37.4), "达到三级")

    def test_template_heading_inherits_direct_multilevel_numbering(self):
        doc = Document()
        source = doc.add_heading("既有二级标题", level=2)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "1")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "2")
        num_pr.extend([ilvl, num_id])
        source._p.get_or_add_pPr().append(num_pr)

        added = add_numbered_heading(doc, "新增二级标题", level=2)

        self.assertEqual(added._p.pPr.numPr.numId.get(qn("w:val")), "2")
        self.assertEqual(added._p.pPr.numPr.ilvl.get(qn("w:val")), "1")

    def test_redundant_empty_section_break_is_removed(self):
        doc = Document()
        doc.add_paragraph("报警阈值表结束")
        empty = doc.add_paragraph()
        empty._p.get_or_add_pPr().append(OxmlElement("w:sectPr"))
        heading = doc.add_heading("监测结果", level=1)
        body = doc._element.body
        heading_index = list(body.iterchildren()).index(heading._p)

        remove_empty_section_break_before(doc, heading_index)

        self.assertNotIn(empty._p, list(body.iterchildren()))


if __name__ == "__main__":
    unittest.main()
