from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Mm
from PIL import Image

REPO_ROOT = Path(__file__).resolve().parents[1]
REPORTING_DIR = REPO_ROOT / "reporting"
if str(REPORTING_DIR) not in sys.path:
    sys.path.insert(0, str(REPORTING_DIR))

from docx_utils import paragraph_has_image, remove_nearby_picture_block_before  # noqa: E402


class DocxImageBlockTests(unittest.TestCase):
    def test_remove_nearby_picture_block_removes_labels_and_pictures_only(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            image_path = Path(tmp) / "point.png"
            Image.new("RGB", (40, 30), "white").save(image_path)

            doc = Document()
            body = doc.add_paragraph("body paragraph")
            pic = doc.add_paragraph()
            pic.add_run().add_picture(str(image_path), width=Mm(20))
            doc.add_paragraph("CS1")
            caption = doc.add_paragraph("图 4-10 caption")

            removed = remove_nearby_picture_block_before(caption)

            self.assertEqual(removed, 2)
            self.assertEqual(body.text, "body paragraph")
            self.assertFalse(any(paragraph_has_image(p) for p in doc.paragraphs))
            self.assertIn("图 4-10 caption", [p.text for p in doc.paragraphs])


if __name__ == "__main__":
    unittest.main()
