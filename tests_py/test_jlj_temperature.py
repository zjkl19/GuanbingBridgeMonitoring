import shutil
import sys
import unittest
from pathlib import Path

from openpyxl import Workbook
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from build_jlj_monthly_report import build_temperature_section, clean_jlj_report_xml_text  # noqa: E402


class TestJljTemperatureSection(unittest.TestCase):
    def setUp(self):
        self.tmp = Path("tmp") / "test_jlj_temperature"
        shutil.rmtree(self.tmp, ignore_errors=True)
        (self.tmp / "stats").mkdir(parents=True)
        wb = Workbook()
        ws = wb.active
        ws.append(["PointID", "Min", "Max", "Mean"])
        ws.append(["JGWD-01-K16-ZFQM-G20", 21.7, 69.8, 40.0])
        ws.append(["WDCGQ-01-K16-A20", 42.0, 43.4, 42.7])
        ws.append(["WDCGQ-02-K16-X1-G20", 19.8, 28.4, 24.1])
        ws.append(["WSDJ-01-K16-X1-G20", 17.8, 57.2, 31.2])
        ws.append(["WSDJ-04-K16-QM-G20", 15.7, 90.7, 48.0])
        wb.save(self.tmp / "stats" / "temp_stats.xlsx")

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_temperature_warning_uses_sensor_category_for_maximum(self):
        section = build_temperature_section({}, self.tmp, self.tmp / "stats", None, self.tmp)

        self.assertIn("\u6865\u9762\u6e29\u5ea6\u6700\u5927\u503c\u4e3a69.8\u2103", section.narrative)
        self.assertIn("\u6865\u5740\u533a\u73af\u5883\u6e29\u5ea6\u6700\u5927\u503c\u4e3a90.7\u2103", section.narrative)
        self.assertNotIn("\u6865\u9762\u6e29\u5ea6\u6700\u5927\u503c\u4e3a90.7\u2103", section.narrative)
        self.assertIn("\u4e3b\u6881\u5185\u6e29\u5ea6\u6700\u5927\u503c\u4e3a57.2\u2103", section.narrative)

    def test_report_cleanup_removes_review_phrases(self):
        path = self.tmp / "cleanup.docx"
        doc = Document()
        doc.add_paragraph(
            "\u6865\u9762\u6e29\u5ea6\u6700\u5927\u503c\u4e3a69.8\u2103\uff0c"
            "\u5efa\u8bae\u7ed3\u5408\u539f\u59cb\u6570\u636e\u548c\u4f20\u611f\u5668\u72b6\u6001\u8fdb\u4e00\u6b65\u590d\u6838\u3002"
        )
        doc.add_paragraph(
            "\u5f53\u524d\u540a\u6746\u53c2\u6570\u914d\u7f6e\u5c1a\u672a\u5b8c\u6574\u6821\u6838\uff0c"
            "\u7d22\u529b\u6362\u7b97\u7ed3\u679c\u6682\u4ec5\u7528\u4e8e\u65f6\u7a0b\u5c55\u793a\u3002"
        )
        doc.save(path)

        clean_jlj_report_xml_text(path)

        cleaned = "\n".join(p.text for p in Document(path).paragraphs)
        self.assertIn("\u6865\u9762\u6e29\u5ea6\u6700\u5927\u503c\u4e3a69.8\u2103\u3002", cleaned)
        self.assertNotIn("\u5efa\u8bae\u7ed3\u5408\u539f\u59cb\u6570\u636e", cleaned)
        self.assertNotIn("\u7d22\u529b\u6362\u7b97\u7ed3\u679c\u6682\u4ec5\u7528\u4e8e\u65f6\u7a0b\u5c55\u793a", cleaned)


if __name__ == "__main__":
    unittest.main()
