from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
import sys

from docx import Document
from openpyxl import Workbook

REPO_ROOT = Path(__file__).resolve().parents[1]
REPORTING_ROOT = REPO_ROOT / "reporting"
for candidate in (REPO_ROOT, REPORTING_ROOT):
    if str(candidate) not in sys.path:
        sys.path.insert(0, str(candidate))

from reporting.build_jlj_monthly_report import (
    ImageItem,
    SectionContent,
    build_traffic_section,
    collect_missing_items,
    jlj_patrol_required,
    jlj_traffic_required,
)
from reporting.report_job import build_qc


class JiulongjiangMissingItemsTests(unittest.TestCase):
    def test_unavailable_section_is_reported(self) -> None:
        section_map = {
            "main_eq": SectionContent(
                narrative="本月未获取到主桥地震动有效数据。",
                summary_sentence="本月未获取到主桥地震动有效数据。",
                available=False,
            )
        }

        rows = collect_missing_items(section_map)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["category"], "章节内容缺失")
        self.assertIn("地震动监测", rows[0]["section"])
        self.assertIn("地震动", rows[0]["detail"])

    def test_missing_image_is_reported_but_existing_image_is_ignored(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            existing = Path(tmp) / "exists.jpg"
            existing.write_bytes(b"fake")
            missing = Path(tmp) / "missing.jpg"
            section_map = {
                "main_wind": SectionContent(
                    narrative="",
                    summary_sentence="",
                    image_items=[
                        ImageItem("已存在图片", existing),
                        ImageItem("缺失图片", missing),
                    ],
                )
            }

            rows = collect_missing_items(section_map)

            self.assertEqual(len(rows), 1)
            self.assertEqual(rows[0]["category"], "图表/资源缺失")
            self.assertEqual(rows[0]["item"], "缺失图片")
            self.assertIn("风向风速监测", rows[0]["section"])

    def test_optional_traffic_without_wim_is_not_applicable(self) -> None:
        content = build_traffic_section(None, required=False)

        self.assertFalse(content.available)
        self.assertTrue(content.not_applicable)
        self.assertEqual(collect_missing_items({"main_traffic": content}), [])

    def test_provided_wim_is_not_falsely_reported_as_available(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            wim_root = Path(tmp) / "wim"
            wim_root.mkdir()
            Workbook().save(wim_root / "WIM_Report_jiulongjiang_202605.xlsx")

            content = build_traffic_section(wim_root, required=False)
            rows = collect_missing_items({"main_traffic": content})

        self.assertFalse(content.available)
        self.assertFalse(content.not_applicable)
        self.assertEqual(len(rows), 1)
        self.assertIn("待接入", rows[0]["detail"])

    def test_traffic_requirement_is_config_driven(self) -> None:
        self.assertFalse(jlj_traffic_required({}))
        self.assertFalse(jlj_traffic_required({"reporting": {"traffic": {"required": False}}}))
        self.assertTrue(jlj_traffic_required({"reporting": {"traffic": {"required": True}}}))

    def test_patrol_requirement_defaults_to_required_and_can_be_optional(self) -> None:
        self.assertTrue(jlj_patrol_required({}))
        self.assertTrue(jlj_patrol_required({"reporting": {"patrol": {"required": True}}}))
        self.assertFalse(jlj_patrol_required({"reporting": {"patrol": {"required": False}}}))
        self.assertFalse(jlj_patrol_required({"reporting": {"patrol": {"required": "off"}}}))

    def test_strict_qc_accepts_not_applicable_traffic(self) -> None:
        content = build_traffic_section(None, required=False)
        missing = collect_missing_items({"main_traffic": content})
        self.assertEqual(missing, [])

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            report = root / "report.docx"
            Document().save(report)
            manifest = root / "jlj_report_build_manifest.json"
            manifest.write_text(
                json.dumps({
                    "status": "ok" if not missing else "warning",
                    "missing_count": len(missing),
                    "missing_items": missing,
                    "warnings": [],
                }),
                encoding="utf-8",
            )

            qc = build_qc(
                report,
                manifest,
                None,
                {"status": "passed", "page_count": 1, "pages": []},
                require_source_provenance=True,
            )

        self.assertEqual(qc["status"], "passed")

    def test_strict_qc_accepts_optional_patrol_marked_not_available(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            report = root / "report.docx"
            Document().save(report)
            manifest = root / "jlj_report_build_manifest.json"
            manifest.write_text(
                json.dumps({
                    "status": "ok",
                    "missing_count": 0,
                    "missing_items": [],
                    "warnings": [],
                    "source_availability": {
                        "patrol": {
                            "required": False,
                            "status": "not_available",
                            "target_period": "2026-05",
                            "source": "",
                            "source_sha256": "",
                            "source_period": "",
                            "action": "template_content_cleared_and_note_inserted",
                        }
                    },
                }),
                encoding="utf-8",
            )
            qc = build_qc(
                report,
                manifest,
                None,
                {"status": "passed", "page_count": 1, "pages": []},
                require_source_provenance=True,
            )
        self.assertEqual(qc["status"], "passed")

    def test_strict_qc_rejects_required_missing_patrol(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            report = root / "report.docx"
            Document().save(report)
            manifest = root / "jlj_report_build_manifest.json"
            manifest.write_text(
                json.dumps({
                    "status": "warning",
                    "missing_count": 1,
                    "missing_items": [{"category": "巡查资料缺失", "label": "2026-05"}],
                    "warnings": [],
                }),
                encoding="utf-8",
            )
            qc = build_qc(
                report,
                manifest,
                None,
                {"status": "passed", "page_count": 1, "pages": []},
                require_source_provenance=True,
            )
        self.assertEqual(qc["status"], "failed")


if __name__ == "__main__":
    unittest.main()
