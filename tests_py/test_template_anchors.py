import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from docx import Document  # noqa: E402

from template_anchors import (  # noqa: E402
    JLJ_SUMMARY_TABLE,
    find_paragraph_containing_anchor,
    find_table_after_anchor,
    remove_anchor_paragraphs,
)


class TestTemplateAnchors(unittest.TestCase):
    def test_find_table_after_anchor(self):
        doc = Document()
        doc.add_paragraph("intro")
        doc.add_paragraph(JLJ_SUMMARY_TABLE.text)
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "target"

        self.assertIsNotNone(find_paragraph_containing_anchor(doc, JLJ_SUMMARY_TABLE))
        found = find_table_after_anchor(doc, JLJ_SUMMARY_TABLE)
        self.assertIsNotNone(found)
        self.assertEqual(found.cell(0, 0).text, "target")

    def test_remove_anchor_paragraphs(self):
        doc = Document()
        doc.add_paragraph("before")
        doc.add_paragraph(" " + JLJ_SUMMARY_TABLE.text + " ")
        doc.add_paragraph("after")

        removed = remove_anchor_paragraphs(doc, [JLJ_SUMMARY_TABLE])

        self.assertEqual(removed, 1)
        self.assertEqual([para.text for para in doc.paragraphs], ["before", "after"])


if __name__ == "__main__":
    unittest.main()
