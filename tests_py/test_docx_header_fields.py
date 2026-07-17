import sys
import tempfile
import unittest
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from calibrate_hongtang_period_template import calibrate_template  # noqa: E402
from docx_header_fields import (  # noqa: E402
    HEADER_CJK_FONT,
    HEADER_FONT_SIZE_PT,
    HEADER_LATIN_FONT,
    HONGTANG_FRONT_MATTER_PAGES,
    HONGTANG_HEADER_WIDTHS_TWIPS,
    audit_header_pagination_fields,
    audit_section_footer_pagination_fields,
    ensure_header_pagination_fields,
    ensure_section_footer_pagination_fields,
)


def add_hongtang_header(doc: Document) -> None:
    table = doc.sections[0].header.add_table(rows=1, cols=3, width=1)
    table.cell(0, 0).text = "报告编号：BG02FQJC2600002-J2"
    table.cell(0, 1).text = "福建省建筑工程质量检测中心有限公司"
    table.cell(0, 2).text = "第 1 页 共 76 页第 1 页 共 76 页"


def add_legacy_wind_table(doc: Document) -> None:
    table = doc.add_table(rows=3, cols=6)
    headers = [
        "测点",
        "平均风向（°）",
        "主导风向（°）",
        "平均风速（m/s）",
        "最大风速（m/s）",
        "主要风速等级（m/s）",
    ]
    for column_index, value in enumerate(headers):
        table.cell(0, column_index).text = value
    table.cell(1, 0).text = "W1"
    table.cell(1, 5).text = "2-4 m/s"
    table.cell(2, 0).text = "W2"
    table.cell(2, 5).text = "0-2 m/s"


def write_word_normalized_header_fixture(path: Path, *, east_asia_font: str = HEADER_CJK_FONT) -> None:
    """Write the field/run shape produced by desktop Word after Fields.Update."""

    header = f'''<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:p><w:r><w:t>报告编号：BG02FQJC2600002-J2</w:t></w:r></w:p>
<w:p><w:pPr><w:rPr><w:sz w:val="18"/><w:szCs w:val="18"/></w:rPr></w:pPr>
<w:r><w:t>第</w:t></w:r><w:r><w:t xml:space="preserve"> </w:t></w:r>
<w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>
<w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>1</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r>
<w:r><w:t xml:space="preserve"> 页 共 </w:t></w:r>
<w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> = </w:instrText></w:r>
<w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>
<w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:instrText>79</w:instrText></w:r>
<w:r><w:fldChar w:fldCharType="end"/></w:r><w:r><w:instrText xml:space="preserve"> - {HONGTANG_FRONT_MATTER_PAGES} </w:instrText></w:r>
<w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>{79 - HONGTANG_FRONT_MATTER_PAGES}</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r>
<w:r><w:t xml:space="preserve"> 页</w:t></w:r></w:p></w:hdr>'''
    styles = f'''<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:ascii="{HEADER_LATIN_FONT}" w:hAnsi="{HEADER_LATIN_FONT}" w:cs="{HEADER_LATIN_FONT}" w:eastAsia="宋体"/></w:rPr></w:rPrDefault></w:docDefaults>
<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:rPr><w:rFonts w:eastAsia="{east_asia_font}"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr></w:style>
</w:styles>'''
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("word/header1.xml", header)
        archive.writestr("word/styles.xml", styles)


def write_footer_audit_fixture(
    path: Path,
    *,
    inherited_absolute_footer: bool,
) -> None:
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rel_namespace = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    package_rel_namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    if inherited_absolute_footer:
        sections = f'''<w:p><w:pPr><w:sectPr>
<w:footerReference w:type="default" r:id="rId1"/>
</w:sectPr></w:pPr></w:p><w:sectPr><w:pgNumType w:start="1"/></w:sectPr>'''
        target = "/word/footer1.xml"
    else:
        sections = '<w:sectPr><w:pgNumType w:start="1"/></w:sectPr>'
        target = "footer1.xml"
    document = f'''<w:document xmlns:w="{namespace}" xmlns:r="{rel_namespace}">
<w:body>{sections}</w:body></w:document>'''
    relationships = f'''<Relationships xmlns="{package_rel_namespace}">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" Target="{target}"/>
</Relationships>'''
    visible_properties = (
        f'<w:rPr><w:rFonts w:ascii="{HEADER_LATIN_FONT}" '
        f'w:hAnsi="{HEADER_LATIN_FONT}" w:cs="{HEADER_LATIN_FONT}" '
        f'w:eastAsia="{HEADER_CJK_FONT}"/><w:sz w:val="18"/>'
        f'<w:szCs w:val="18"/></w:rPr>'
    )
    footer = f'''<w:ftr xmlns:w="{namespace}"><w:p>
<w:r>{visible_properties}<w:t xml:space="preserve">第 </w:t></w:r>
<w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> PAGE </w:instrText></w:r>
<w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r>{visible_properties}<w:t>1</w:t></w:r>
<w:r><w:fldChar w:fldCharType="end"/></w:r>
<w:r>{visible_properties}<w:t xml:space="preserve"> 页 共 </w:t></w:r>
<w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> SECTIONPAGES </w:instrText></w:r>
<w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r>{visible_properties}<w:t>1</w:t></w:r>
<w:r><w:fldChar w:fldCharType="end"/></w:r>
<w:r>{visible_properties}<w:t xml:space="preserve"> 页</w:t></w:r>
</w:p></w:ftr>'''
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("word/document.xml", document)
        archive.writestr("word/_rels/document.xml.rels", relationships)
        archive.writestr("word/footer1.xml", footer)


def _append_complex_field(paragraph, instruction: str, result: str = "1") -> None:
    def append_field_char(field_type: str) -> None:
        node = OxmlElement("w:fldChar")
        node.set(qn("w:fldCharType"), field_type)
        paragraph.add_run()._r.append(node)

    append_field_char("begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = f" {instruction} "
    paragraph.add_run()._r.append(instruction_node)
    append_field_char("separate")
    paragraph.add_run(result)
    append_field_char("end")


def write_legacy_guanbing_footer_fixture(path: Path) -> None:
    """Create the legacy static-total footer contract without a private template."""

    document = Document()
    document.add_paragraph(
        "G104 \u7ba1\u67c4\u5927\u6865 "
        "\u62a5\u544a\u7f16\u53f7\uff1aBG02FQJC2400001-M18"
    )
    section = document.sections[0]
    page_numbering = OxmlElement("w:pgNumType")
    page_numbering.set(qn("w:start"), "1")
    section._sectPr.append(page_numbering)
    paragraph = section.footer.paragraphs[0]
    paragraph.add_run("\u7b2c ")
    _append_complex_field(paragraph, "PAGE")
    paragraph.add_run(" \u9875 \u5171 76 \u9875")
    document.save(path)


class TestDocxHeaderFields(unittest.TestCase):
    def test_guanbing_body_footer_uses_dynamic_section_total_and_required_fonts(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "legacy-static-pagination.docx"
            output = Path(temp_dir) / "guanbing-pagination.docx"
            write_legacy_guanbing_footer_fixture(template)
            document = Document(template)
            self.assertGreaterEqual(ensure_section_footer_pagination_fields(document), 1)
            document.save(output)

            audit = audit_section_footer_pagination_fields(output)

        self.assertTrue(audit.valid, (audit.details, audit.formatting_errors))
        self.assertEqual(audit.footer_parts, 1)
        self.assertGreaterEqual(audit.pagination_paragraphs, 1)
        self.assertEqual(audit.page_fields, audit.pagination_paragraphs)
        self.assertEqual(audit.sectionpages_fields, audit.pagination_paragraphs)
        self.assertEqual(audit.static_total_paragraphs, 0)
        self.assertEqual(audit.formatting_errors, ())

    def test_footer_audit_resolves_absolute_target_inherited_by_restart_section(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "inherited-absolute-footer.docx"
            write_footer_audit_fixture(path, inherited_absolute_footer=True)

            audit = audit_section_footer_pagination_fields(path)

        self.assertTrue(audit.valid, (audit.details, audit.formatting_errors))
        self.assertEqual(audit.footer_parts, 1)
        self.assertTrue(any("word/footer1.xml" in item for item in audit.details))

    def test_footer_audit_does_not_fall_back_to_unreferenced_valid_footer(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "unreferenced-footer.docx"
            write_footer_audit_fixture(path, inherited_absolute_footer=False)

            audit = audit_section_footer_pagination_fields(path)

        self.assertFalse(audit.valid)
        self.assertEqual(audit.footer_parts, 0)
        self.assertTrue(
            any("no resolvable inherited/default footer" in item for item in audit.details),
            audit.details,
        )

    def test_ensure_repairs_visible_format_when_fields_are_already_correct(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "legacy-static-pagination.docx"
            output = Path(temp_dir) / "repaired-format.docx"
            write_legacy_guanbing_footer_fixture(template)
            document = Document(template)
            ensure_section_footer_pagination_fields(document)
            for section in document.sections:
                for paragraph in section.footer._element.iter(qn("w:p")):
                    instructions = "".join(
                        node.text or "" for node in paragraph.iter(qn("w:instrText"))
                    )
                    if "PAGE" not in instructions or "SECTIONPAGES" not in instructions:
                        continue
                    for run in paragraph.findall(".//" + qn("w:r")):
                        if run.find(".//" + qn("w:t")) is None:
                            continue
                        r_pr = run.find(qn("w:rPr"))
                        if r_pr is not None:
                            run.remove(r_pr)

            self.assertGreaterEqual(ensure_section_footer_pagination_fields(document), 1)
            document.save(output)
            audit = audit_section_footer_pagination_fields(output)

        self.assertTrue(audit.valid, (audit.details, audit.formatting_errors))
        self.assertEqual(audit.formatting_errors, ())

    def test_footer_audit_ignores_hidden_field_code_formatting(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "legacy-static-pagination.docx"
            output = Path(temp_dir) / "hidden-field-format.docx"
            write_legacy_guanbing_footer_fixture(template)
            document = Document(template)
            ensure_section_footer_pagination_fields(document)
            for section in document.sections:
                for paragraph in section.footer._element.iter(qn("w:p")):
                    for run in paragraph.findall(".//" + qn("w:r")):
                        if run.find(".//" + qn("w:t")) is not None:
                            continue
                        if (
                            run.find(".//" + qn("w:fldChar")) is None
                            and run.find(".//" + qn("w:instrText")) is None
                        ):
                            continue
                        r_pr = run.find(qn("w:rPr"))
                        if r_pr is not None:
                            run.remove(r_pr)
            document.save(output)
            audit = audit_section_footer_pagination_fields(output)

        self.assertTrue(audit.valid, (audit.details, audit.formatting_errors))
        self.assertEqual(audit.formatting_errors, ())

    def test_footer_audit_rejects_extra_static_page_number_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "legacy-static-pagination.docx"
            output = Path(temp_dir) / "extra-static-page-number.docx"
            write_legacy_guanbing_footer_fixture(template)
            document = Document(template)
            ensure_section_footer_pagination_fields(document)
            appended = 0
            for section in document.sections:
                for paragraph in section.footer._element.iter(qn("w:p")):
                    instructions = "".join(
                        node.text or "" for node in paragraph.iter(qn("w:instrText"))
                    )
                    if "PAGE" not in instructions or "SECTIONPAGES" not in instructions:
                        continue
                    visible_run = next(
                        run
                        for run in paragraph.findall(".//" + qn("w:r"))
                        if run.find(".//" + qn("w:t")) is not None
                    )
                    extra = deepcopy(visible_run)
                    extra.find(".//" + qn("w:t")).text = " 99"
                    paragraph.append(extra)
                    appended += 1
            self.assertGreaterEqual(appended, 1)
            document.save(output)
            audit = audit_section_footer_pagination_fields(output)

        self.assertFalse(audit.valid)
        self.assertGreaterEqual(audit.static_total_paragraphs, 1)

    def test_canonical_hongtang_template_has_current_header_and_wind_contract(self):
        repo_root = Path(__file__).resolve().parents[1]
        template = repo_root / "reports" / "洪塘大桥健康监测周期报模板-自动报告.docx"

        audit = audit_header_pagination_fields(template)
        document = Document(template)
        wind_table = next(
            table
            for table in document.tables
            if any("平均风向" in cell.text for cell in table.rows[0].cells)
        )

        self.assertTrue(audit.valid, (audit.details, audit.formatting_errors))
        self.assertEqual(audit.front_matter_pages, (3,))
        self.assertEqual(len(wind_table.columns), 7)
        self.assertEqual(tuple(cell.text for cell in wind_table.rows[0].cells), (
            "测点",
            "平均风向（°）",
            "主导风向（°）",
            "平均风速（m/s）",
            "瞬时最大风速（m/s）",
            "10min平均风速最大值（m/s）",
            "主要风速等级（m/s）",
        ))

    def test_audit_rejects_unadjusted_numpages_field(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "word-normalized.docx"
            header = (
                '<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:p><w:r><w:t>报告编号：BG02FQJC2600002-J2</w:t></w:r></w:p>'
                '<w:p><w:r><w:t xml:space="preserve">第 </w:t></w:r>'
                '<w:r><w:instrText> PAGE </w:instrText></w:r><w:r><w:t>1</w:t></w:r>'
                '<w:r><w:t xml:space="preserve"> 页 共 </w:t></w:r>'
                '<w:fldSimple w:instr=" NUMPAGES "><w:r><w:t>106</w:t></w:r></w:fldSimple>'
                '<w:r><w:t xml:space="preserve"> 页</w:t></w:r></w:p></w:hdr>'
            )
            with zipfile.ZipFile(path, "w") as archive:
                archive.writestr("word/header1.xml", header)

            audit = audit_header_pagination_fields(path)

            self.assertFalse(audit.valid, audit.details)
            self.assertEqual(audit.page_fields, 1)
            self.assertEqual(audit.numpages_fields, 1)
            self.assertEqual(audit.total_page_formula_fields, 0)

    def test_replaces_duplicate_static_pagination_with_adjusted_fields_and_fonts(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "header.docx"
            doc = Document()
            add_hongtang_header(doc)

            self.assertEqual(ensure_header_pagination_fields(doc), 1)
            doc.save(path)

            audit = audit_header_pagination_fields(path)
            self.assertTrue(audit.valid, (audit.details, audit.formatting_errors))
            self.assertEqual(audit.page_fields, 1)
            self.assertEqual(audit.numpages_fields, 1)
            self.assertEqual(audit.total_page_formula_fields, 1)
            self.assertEqual(audit.front_matter_pages, (HONGTANG_FRONT_MATTER_PAGES,))
            self.assertEqual(audit.duplicate_page_phrases, 0)
            self.assertEqual(audit.formatting_errors, ())
            calibrated = Document(path)
            widths = tuple(
                int(cell._tc.tcPr.tcW.w)
                for cell in calibrated.sections[0].header.tables[0].rows[0].cells
            )
            self.assertEqual(widths, HONGTANG_HEADER_WIDTHS_TWIPS)

            with zipfile.ZipFile(path) as archive:
                header_xml = archive.read("word/header1.xml").decode("utf-8")
            self.assertIn(" = ", header_xml)
            self.assertIn(" NUMPAGES ", header_xml)
            self.assertIn(f" - {HONGTANG_FRONT_MATTER_PAGES} ", header_xml)
            self.assertIn(f'w:ascii="{HEADER_LATIN_FONT}"', header_xml)
            self.assertIn(f'w:hAnsi="{HEADER_LATIN_FONT}"', header_xml)
            self.assertIn(f'w:eastAsia="{HEADER_CJK_FONT}"', header_xml)
            self.assertIn(f'w:sz w:val="{HEADER_FONT_SIZE_PT * 2}"', header_xml)

            namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            root = ET.fromstring(header_xml)
            page_paragraph = next(
                paragraph
                for paragraph in root.findall(f".//{namespace}p")
                if any(
                    (node.text or "").strip().upper() == "PAGE"
                    for node in paragraph.findall(f".//{namespace}instrText")
                )
            )
            field_tokens = []
            for run in page_paragraph.findall(f"./{namespace}r"):
                field_char = run.find(f"{namespace}fldChar")
                instruction = run.find(f"{namespace}instrText")
                text = run.find(f"{namespace}t")
                if field_char is not None:
                    field_tokens.append(("field", field_char.get(f"{namespace}fldCharType")))
                elif instruction is not None:
                    field_tokens.append(("instruction", instruction.text))
                elif text is not None:
                    field_tokens.append(("text", text.text))
            self.assertEqual(field_tokens, [
                ("text", "第 "),
                ("field", "begin"),
                ("instruction", " PAGE "),
                ("field", "separate"),
                ("text", "1"),
                ("field", "end"),
                ("text", " 页 共 "),
                ("field", "begin"),
                ("instruction", " = "),
                ("field", "begin"),
                ("instruction", " NUMPAGES "),
                ("field", "separate"),
                ("text", str(HONGTANG_FRONT_MATTER_PAGES + 1)),
                ("field", "end"),
                ("instruction", f" - {HONGTANG_FRONT_MATTER_PAGES} "),
                ("field", "separate"),
                ("text", "1"),
                ("field", "end"),
                ("text", " 页"),
            ])

    def test_audit_accepts_word_normalized_nested_result_and_inherited_fonts(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "word-normalized.docx"
            write_word_normalized_header_fixture(path)

            audit = audit_header_pagination_fields(path)

            self.assertTrue(audit.valid, (audit.details, audit.formatting_errors))
            self.assertEqual(audit.page_fields, 1)
            self.assertEqual(audit.numpages_fields, 1)
            self.assertEqual(audit.total_page_formula_fields, 1)
            self.assertEqual(audit.front_matter_pages, (HONGTANG_FRONT_MATTER_PAGES,))
            self.assertEqual(audit.formatting_errors, ())

    def test_audit_rejects_wrong_effective_inherited_cjk_font(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "wrong-inherited-font.docx"
            write_word_normalized_header_fixture(path, east_asia_font="宋体")

            audit = audit_header_pagination_fields(path)

            self.assertFalse(audit.valid)
            self.assertTrue(
                any("effective eastAsia font" in item for item in audit.formatting_errors),
                audit.formatting_errors,
            )

    def test_front_matter_page_offset_is_configurable(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "header-offset.docx"
            doc = Document()
            add_hongtang_header(doc)
            self.assertEqual(ensure_header_pagination_fields(doc, front_matter_pages=2), 1)
            doc.save(path)

            audit = audit_header_pagination_fields(path, front_matter_pages=2)

            self.assertTrue(audit.valid, (audit.details, audit.formatting_errors))
            self.assertEqual(audit.front_matter_pages, (2,))

    def test_template_calibration_applies_proofreading_header_and_wind_contract(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source.docx"
            output = root / "output.docx"
            doc = Document()
            add_hongtang_header(doc)
            doc.add_paragraph("主塔倾斜采用倾角仪实时测量，通过倾角值的变化来反应主塔的倾斜情况。")
            doc.add_paragraph("图 2-12 主梁加速度传感纵断面布置图（单位：m）")
            doc.add_paragraph("监测结果如表4-12所示。监测结果表明，桥面10min平均风速最大值为5.46m/s。")
            add_legacy_wind_table(doc)
            doc.save(source)

            result = calibrate_template(source, output)

            self.assertEqual(result["header_cells"], 1)
            self.assertEqual(result["wind_table_upgraded"], 1)
            calibrated = Document(output)
            text = "\n".join(paragraph.text for paragraph in calibrated.paragraphs)
            self.assertIn("反映主塔的倾斜情况", text)
            self.assertIn("主梁加速度传感器纵断面布置图", text)
            self.assertIn("瞬时最大风速", text)
            self.assertTrue(audit_header_pagination_fields(output).valid)
            wind_table = next(
                table
                for table in calibrated.tables
                if any("平均风向" in cell.text for cell in table.rows[0].cells)
            )
            self.assertEqual(len(wind_table.columns), 7)
            self.assertEqual(wind_table.cell(0, 5).text, "10min平均风速最大值（m/s）")
            self.assertEqual(wind_table.cell(1, 5).text, "")
            self.assertEqual(wind_table.cell(1, 6).text, "2-4 m/s")


if __name__ == "__main__":
    unittest.main()
