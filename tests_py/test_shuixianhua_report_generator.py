import json
import re
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch
from zipfile import ZipFile

from docx import Document
from openpyxl import Workbook
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from build_shuixianhua_monthly_report import (  # noqa: E402
    _sxh_accel_rms_limits_mps2,
    _sxh_accel_rms_status,
    _sxh_context,
    _sxh_parse_wind_summaries,
    _sxh_summary_payload,
    _sxh_update_docx_package,
    _sxh_validate_generated_content,
    build_report,
    report_acquisition_rows,
    scaled_accel_rows,
)
from artifact_lookup import ArtifactLookupResult  # noqa: E402
from report_qc import ReportQcIssue, ReportQcResult  # noqa: E402
from shuixianhua_table_anchors import required_result_tables  # noqa: E402
from lxml import etree  # noqa: E402


def _write_rows(path: Path, headers: list[str], rows: list[list[object]]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(path)


class ShuixianhuaReportGeneratorTest(unittest.TestCase):
    @staticmethod
    def _write_wind_summary(
        root: Path,
        point_id: str,
        start_date: str,
        end_date: str,
        *,
        mean_speed: float = 1.25,
    ) -> Path:
        folder = root / "风速风向结果" / "风玫瑰"
        folder.mkdir(parents=True, exist_ok=True)
        path = folder / (
            f"{point_id}_windrose_{start_date}_{end_date}_summary.txt"
        )
        path.write_text(
            "\n".join(
                [
                    f"风玫瑰简要结论（{point_id}）",
                    "平均风向: 92.1°",
                    "主导风向: 90.0°-112.5°，占比 21.5%",
                    f"平均风速: {mean_speed:.2f} m/s",
                    "最大风速: 7.90 m/s",
                    "主要风速等级: 0-2 m/s（依据：全样本风速分级占比最高）",
                ]
            ),
            encoding="utf-8",
        )
        return path

    @staticmethod
    def _minimal_build_fixture(root: Path) -> tuple[Path, Path, Path]:
        template = root / "template.docx"
        image = root / "pixel.png"
        Image.new("RGB", (20, 20), "blue").save(image)
        doc = Document()
        doc.add_paragraph("水仙花大桥监测结果")
        doc.add_paragraph("监测时间：2026年3月")
        doc.add_picture(str(image))
        doc.save(template)
        config = root / "config.json"
        config.write_text("{}", encoding="utf-8")
        return template, config, root / "reports"

    @staticmethod
    def _minimal_context() -> dict[str, object]:
        return {
            "cfg": {},
            "date_span": "2026-05-01~2026-05-31",
            "report_rows": [],
            "temp_rows": [
                {"PointID": "WD-01", "Min": -2.0, "Max": 3.0, "Mean": 1.0},
                {"PointID": "WSD-01-11#-S11", "Min": 20.0, "Max": 30.0, "Mean": 25.0},
            ],
            "humidity_rows": [
                {"PointID": "WSD-01-11#-S11", "Min": 45.0, "Max": 55.0, "Mean": 50.0}
            ],
            "wind_rows": [],
            "earthquake_rows": [],
            "deflection_rows": [],
            "bearing_rows": [],
            "accel_rows": [],
            "accel_freq_map": {},
            "strain_rows": [],
            "cable_rows": [],
        }

    def test_builder_writes_real_report_manifest_for_strict_runner(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            image = root / "pixel.png"
            Image.new("RGB", (20, 20), "blue").save(image)
            doc = Document()
            doc.add_paragraph("水仙花大桥监测结果")
            doc.add_paragraph("监测时间：2026年3月")
            doc.add_picture(str(image))
            doc.save(template)
            config = root / "config.json"
            config.write_text("{}", encoding="utf-8")
            output = root / "reports"

            with patch(
                "build_shuixianhua_monthly_report._sxh_context",
                return_value={"report_rows": []},
            ), patch(
                "build_shuixianhua_monthly_report._sxh_update_docx_package"
            ), patch(
                "build_shuixianhua_monthly_report._sxh_resolve_report_images",
                return_value=({}, [], []),
            ), patch(
                "build_shuixianhua_monthly_report._sxh_replace_report_image_blocks",
                return_value=([], []),
            ), patch(
                "build_shuixianhua_monthly_report._sxh_validate_generated_content",
                return_value=[],
            ):
                report, pdf = build_report(
                    template=template,
                    config_path=config,
                    result_root=root,
                    output_dir=output,
                    update_word=False,
                )

            manifests = list(output.glob("shuixianhua_report_build_manifest_*.json"))
            self.assertTrue(report.is_file())
            self.assertIsNone(pdf)
            self.assertEqual(len(manifests), 1)
            payload = json.loads(manifests[0].read_text(encoding="utf-8"))
            self.assertEqual(payload["manifest_type"], "report_build")
            self.assertEqual(payload["report_type"], "shuixianhua_monthly")
            self.assertEqual(payload["missing_count"], 0)
            self.assertNotIn("Legacy builder did not return", "\n".join(payload["warnings"]))

    def test_builder_fails_closed_after_warning_qc_and_writes_evidence(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template, config, output = self._minimal_build_fixture(root)
            qc_result = ReportQcResult(
                kind="shuixianhua_monthly",
                docx_path="candidate.docx",
                checked_at="2026-07-14 12:00:00",
                status="warning",
                issue_count=1,
                warning_count=1,
                issues=[
                    ReportQcIssue(
                        code="synthetic-warning",
                        severity="warning",
                        message="synthetic QC warning",
                    )
                ],
                summary={"exists": True, "image_count": 1},
            )

            with patch(
                "build_shuixianhua_monthly_report._sxh_context",
                return_value={"report_rows": []},
            ), patch(
                "build_shuixianhua_monthly_report._sxh_update_docx_package"
            ), patch(
                "build_shuixianhua_monthly_report._sxh_resolve_report_images",
                return_value=({}, [], []),
            ), patch(
                "build_shuixianhua_monthly_report._sxh_replace_report_image_blocks",
                return_value=([], []),
            ), patch(
                "build_shuixianhua_monthly_report._sxh_validate_generated_content",
                return_value=[],
            ), patch(
                "build_shuixianhua_monthly_report.check_shuixianhua_report",
                return_value=qc_result,
            ):
                with self.assertRaisesRegex(
                    RuntimeError,
                    r"QC did not pass \(status=warning\)",
                ):
                    build_report(
                        template=template,
                        config_path=config,
                        result_root=root,
                        output_dir=output,
                        update_word=False,
                    )

            manifests = list(output.glob("shuixianhua_report_build_manifest_*.json"))
            self.assertEqual(len(manifests), 1)
            payload = json.loads(manifests[0].read_text(encoding="utf-8"))
            self.assertEqual(payload["status"], "failed")
            self.assertEqual(payload["report_qc_status"], "warning")
            self.assertIn("synthetic-warning", "\n".join(payload["warnings"]))
            self.assertTrue(Path(payload["report_qc_txt"]).is_file())
            qc_payload = json.loads(Path(payload["report_qc_json"]).read_text(encoding="utf-8"))
            self.assertEqual(qc_payload["status"], "warning")

    def test_builder_records_qc_exception_before_failing_closed(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template, config, output = self._minimal_build_fixture(root)

            with patch(
                "build_shuixianhua_monthly_report._sxh_context",
                return_value={"report_rows": []},
            ), patch(
                "build_shuixianhua_monthly_report._sxh_update_docx_package"
            ), patch(
                "build_shuixianhua_monthly_report._sxh_resolve_report_images",
                return_value=({}, [], []),
            ), patch(
                "build_shuixianhua_monthly_report._sxh_replace_report_image_blocks",
                return_value=([], []),
            ), patch(
                "build_shuixianhua_monthly_report._sxh_validate_generated_content",
                return_value=[],
            ), patch(
                "build_shuixianhua_monthly_report.check_shuixianhua_report",
                side_effect=ValueError("synthetic QC crash"),
            ):
                with self.assertRaisesRegex(
                    RuntimeError,
                    r"QC did not pass \(status=failed\)",
                ):
                    build_report(
                        template=template,
                        config_path=config,
                        result_root=root,
                        output_dir=output,
                        update_word=False,
                    )

            manifests = list(output.glob("shuixianhua_report_build_manifest_*.json"))
            self.assertEqual(len(manifests), 1)
            payload = json.loads(manifests[0].read_text(encoding="utf-8"))
            self.assertEqual(payload["status"], "failed")
            self.assertEqual(payload["report_qc_status"], "failed")
            self.assertTrue(Path(payload["report_qc_txt"]).is_file())
            qc_path = Path(payload["report_qc_json"])
            self.assertTrue(qc_path.is_file())
            qc_payload = json.loads(qc_path.read_text(encoding="utf-8"))
            self.assertEqual(qc_payload["status"], "failed")
            self.assertEqual(qc_payload["issues"][0]["code"], "report-qc-exception")
            self.assertIn("ValueError: synthetic QC crash", qc_payload["issues"][0]["detail"])
            self.assertTrue(
                any(
                    item["label"] == "report-qc-exception"
                    for item in payload["missing_items"]
                )
            )

    def test_acceleration_rows_convert_cm_s2_to_m_s2(self):
        rows = scaled_accel_rows([
            {"PointID": "ZLZD-01", "Min": -68.3, "Max": 12.5, "Mean": 0, "RMS10minMax": 68.3}
        ])

        self.assertAlmostEqual(rows[0]["Min"], -0.683)
        self.assertAlmostEqual(rows[0]["Max"], 0.125)
        self.assertAlmostEqual(rows[0]["Mean"], 0.0)
        self.assertAlmostEqual(rows[0]["RMS10minMax"], 0.683)

    def test_wind_summary_normal_may_period_populates_expected_point(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            self._write_wind_summary(
                root,
                "FSFX-01",
                "2026-05-01",
                "2026-05-31",
                mean_speed=2.35,
            )

            rows = _sxh_parse_wind_summaries(
                root,
                [{"PointID": "FSFX-01", "Mean10minMax": 3.4}],
                "2026-05-01",
                "2026-05-31",
            )

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["PointID"], "FSFX-01")
        self.assertEqual(rows[0]["平均风速"], "2.35 m/s")
        self.assertEqual(rows[0]["主导风向"], "90.0°-112.5°，占比21.5%")

    def test_wind_summary_ignores_other_month_and_selects_unique_may_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            self._write_wind_summary(
                root,
                "FSFX-01",
                "2026-03-23",
                "2026-03-31",
                mean_speed=0.86,
            )
            self._write_wind_summary(
                root,
                "FSFX-01",
                "2026-05-01",
                "2026-05-31",
                mean_speed=2.35,
            )

            rows = _sxh_parse_wind_summaries(
                root,
                [{"PointID": "FSFX-01"}],
                "2026-05-01",
                "2026-05-31",
            )

        self.assertEqual(rows[0]["平均风速"], "2.35 m/s")

    def test_wind_summary_missing_or_ambiguous_current_period_fails_closed(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            with self.assertRaisesRegex(FileNotFoundError, "未找到水仙花本报告期"):
                _sxh_parse_wind_summaries(
                    root,
                    [{"PointID": "FSFX-01"}],
                    "2026-05-01",
                    "2026-05-31",
                )

            self._write_wind_summary(
                root,
                "FSFX-01",
                "2026-03-23",
                "2026-03-31",
            )
            with self.assertRaisesRegex(ValueError, "均不属于请求的报告期"):
                _sxh_parse_wind_summaries(
                    root,
                    [{"PointID": "FSFX-01"}],
                    "2026-05-01",
                    "2026-05-31",
                )

            self._write_wind_summary(
                root,
                "FSFX-01",
                "2026-05-01",
                "2026-05-31",
            )
            duplicate_root = root / "duplicate"
            self._write_wind_summary(
                duplicate_root,
                "FSFX-01",
                "2026-05-01",
                "2026-05-31",
                mean_speed=9.99,
            )
            with self.assertRaisesRegex(ValueError, "同一报告期存在多份"):
                _sxh_parse_wind_summaries(
                    root,
                    [{"PointID": "FSFX-01"}],
                    "2026-05-01",
                    "2026-05-31",
                )

    def test_wind_summary_wrong_period_bound_manifest_fails_closed(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            stale = self._write_wind_summary(
                root,
                "FSFX-01",
                "2026-03-23",
                "2026-03-31",
            )
            lookup = ArtifactLookupResult(
                stale,
                {
                    "source": "analysis_manifest",
                    "manifest": "analysis_manifest_march.json",
                },
            )
            with patch(
                "build_shuixianhua_monthly_report.lookup_latest_file_patterns",
                return_value=lookup,
            ):
                with self.assertRaisesRegex(
                    ValueError, "已绑定分析结果清单"
                ):
                    _sxh_parse_wind_summaries(
                        root,
                        [{"PointID": "FSFX-01"}],
                        "2026-05-01",
                        "2026-05-31",
                    )

    def test_report_acquisition_rows_preserves_actual_temperature_counts(self):
        source = [{
            "模块代码": "temperature",
            "模块": "温度",
            "配置测点数": 2,
            "实际获取测点数": 2,
            "缺失测点数": 0,
            "获取率": 1.0,
            "缺失说明": "无",
        }]

        rows = report_acquisition_rows(source)

        self.assertEqual(rows, source)

    def test_summary_uses_actual_temperature_and_humidity_rows(self):
        payload = _sxh_summary_payload(self._minimal_context())

        has_overall_range = re.search(r"-2(?:\.0)?℃~30(?:\.0)?℃", payload["temp"])
        has_per_point_ranges = (
            re.search(r"WD-01[^。]*-2(?:\.0)?℃~3(?:\.0)?℃", payload["temp"])
            and re.search(r"WSD-01[^。]*20(?:\.0)?℃~30(?:\.0)?℃", payload["temp"])
        )
        self.assertTrue(has_overall_range or has_per_point_ranges, payload["temp"])
        self.assertNotRegex(payload["temp"], r"WSD-01[^。]*-2(?:\.0)?℃~30(?:\.0)?℃")
        self.assertNotIn("WD-01~WD-09温度测点本月未获取有效数据", payload["temp"])
        self.assertRegex(payload["humidity"], r"45(?:\.0)?%~55(?:\.0)?%")

    def test_real_template_replaces_period_and_temperature_humidity_summary(self):
        repo_root = Path(__file__).resolve().parents[1]
        template = repo_root / "reports" / "水仙花大桥健康监测月报模板.docx"
        with tempfile.TemporaryDirectory() as tmp:
            report = Path(tmp) / "report.docx"
            report.write_bytes(template.read_bytes())

            _sxh_update_docx_package(
                report,
                self._minimal_context(),
                "2026年5月1日~2026年5月31日",
                "2026年6月5日",
                period_label="2026年5月",
            )

            document = Document(report)
            text = "\n".join(
                [paragraph.text for paragraph in document.paragraphs]
                + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            )
            with ZipFile(report) as archive:
                root = etree.fromstring(archive.read("word/document.xml"))
            temperature_table = required_result_tables(root)["temperature"]
            temperature_row_count = len(
                temperature_table.findall(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"
                )
            )
            content_issues = _sxh_validate_generated_content(
                report, self._minimal_context()
            )

        self.assertIn("监测时间：2026年5月", text)
        self.assertNotIn("监测时间：2026年03月", text)
        self.assertNotIn("61.3%~73.9%", text)
        self.assertRegex(text, r"45(?:\.0)?%~55(?:\.0)?%")
        self.assertNotIn("WD-01~WD-09温度测点本月未获取有效数据", text)
        self.assertNotIn("温度监测中9个结构温度测点本月未获取数据", text)
        self.assertIn("本月温度监测2个测点均获取到有效数据", text)
        self.assertEqual(temperature_row_count, 3)
        self.assertEqual(content_issues, [])

    def test_acceleration_rms_limits_are_read_from_cm_s2_config(self):
        cfg = {
            "groups": {"acceleration": {"ZG": ["ZGZD-01"], "ZL": ["ZLZD-01"]}},
            "plot_styles": {
                "acceleration": {
                    "rms_warn_lines": {
                        "ZG": [{"y": 31.5}, {"y": 50}],
                        "ZL": [{"y": 31.5}, {"y": 50}],
                    }
                }
            },
        }

        first, second = _sxh_accel_rms_limits_mps2(cfg, "ZGZD-01")

        self.assertAlmostEqual(first, 0.315)
        self.assertAlmostEqual(second, 0.5)

    def test_acceleration_rms_status_reports_exceedance(self):
        level1_status = _sxh_accel_rms_status(0.4, 0.315, 0.5)
        level2_status = _sxh_accel_rms_status(0.683, 0.315, 0.5)

        self.assertIn("超过一级阈值0.315m/s²", level1_status)
        self.assertIn("未超过二级阈值0.5m/s²", level1_status)
        self.assertEqual(level2_status, "超过二级阈值0.5m/s²。")

    def test_template_tables_are_found_by_caption(self):
        template = Path(__file__).resolve().parents[1] / "reports" / "水仙花大桥健康监测月报模板.docx"
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "template.docx"
            docx.write_bytes(template.read_bytes())
            from zipfile import ZipFile

            with ZipFile(docx) as zf:
                root = etree.fromstring(zf.read("word/document.xml"))

        tables = required_result_tables(root)

        self.assertIn("acquisition", tables)
        self.assertIn("strain", tables)
        self.assertIn("cable_accel", tables)

    def test_context_falls_back_to_config_and_stats_when_acquisition_summary_is_absent(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            stats = root / "stats"
            stats.mkdir()
            cfg = {
                "points": {
                    "temperature": ["WD-01", "WSD-01"],
                    "humidity": ["WSD-01"],
                    "wind_speed": ["FSFX-01"],
                    "eq": ["JSD-01"],
                    "deflection": ["GSNDY-01"],
                    "bearing_displacement": ["ZZWY-01"],
                    "acceleration": ["ZLZD-01"],
                    "strain": ["GLYB-01"],
                    "cable_accel": ["SL-01"],
                },
                "design_points_pending": {"gnss": ["GNSS-01"]},
            }
            cfg_path = root / "config.json"
            cfg_path.write_text(json.dumps(cfg, ensure_ascii=False), encoding="utf-8")

            _write_rows(stats / "temp_stats.xlsx", ["PointID", "Min", "Max", "Mean"], [["WSD-01", 20, 30, 25]])
            _write_rows(stats / "humidity_stats.xlsx", ["PointID", "Min", "Max", "Mean"], [["WSD-01", 50, 60, 55]])
            _write_rows(stats / "wind_stats.xlsx", ["PointID", "MeanSpeed", "MaxSpeed", "Mean10minMax"], [["FSFX-01", 1, 2, 1.5]])
            self._write_wind_summary(
                root,
                "FSFX-01",
                "2026-03-23",
                "2026-03-31",
            )
            _write_rows(stats / "eq_stats.xlsx", ["PointID", "Component", "Min", "Max"], [["JSD-01", "X", -0.1, 0.1]])
            _write_rows(stats / "deflection_stats.xlsx", ["PointID", "OrigMin_mm", "OrigMax_mm", "FiltMin_mm", "FiltMax_mm"], [["GSNDY-01", -1, 2, -0.5, 1.5]])
            _write_rows(stats / "bearing_displacement_stats.xlsx", ["PointID", "OrigMin_mm", "OrigMax_mm", "FiltMin_mm", "FiltMax_mm"], [["ZZWY-01", -1, 1, -0.8, 0.8]])
            _write_rows(stats / "accel_stats.xlsx", ["PointID", "Min", "Max", "RMS10minMax"], [["ZLZD-01", -0.01, 0.02, 0.01]])
            _write_rows(stats / "strain_stats.xlsx", ["PointID", "Min", "Max", "Mean"], [["GLYB-01", -10, 20, 1]])
            _write_rows(stats / "cable_accel_stats.xlsx", ["PointID", "Min", "Max", "RMS10minMax"], [["SL-01", -0.01, 0.02, 0.01]])

            context = _sxh_context(cfg_path, root, "2026年03月23日~2026年03月31日")

        rows = {row["模块"]: row for row in context["report_rows"]}
        self.assertEqual(rows["温度"]["配置测点数"], 2)
        self.assertEqual(rows["温度"]["实际获取测点数"], 1)
        self.assertEqual(rows["拱顶、拱脚位移（GNSS）"]["缺失说明"], "本月未获取有效数据")


if __name__ == "__main__":
    unittest.main()
