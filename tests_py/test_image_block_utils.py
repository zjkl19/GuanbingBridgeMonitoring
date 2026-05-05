import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from docx import Document  # noqa: E402
from PIL import Image  # noqa: E402

from image_block_utils import count_docx_images, stack_images_vertical  # noqa: E402


class TestImageBlockUtils(unittest.TestCase):
    def test_stack_images_vertical_resizes_and_reports_missing(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            img1 = root / "a.png"
            img2 = root / "b.png"
            output = root / "out.jpg"
            Image.new("RGB", (20, 10), "red").save(img1)
            Image.new("RGB", (10, 10), "blue").save(img2)

            result = stack_images_vertical([img1, root / "missing.png", img2], output, gap=5)
            self.assertEqual(result.path, output)
            self.assertEqual(result.source_count, 2)
            self.assertEqual(result.missing_count, 1)
            self.assertTrue(output.exists())
            with Image.open(output) as stacked:
                self.assertEqual(stacked.width, 20)
                self.assertEqual(stacked.height, 35)

    def test_count_docx_images(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            img = root / "a.png"
            docx = root / "x.docx"
            Image.new("RGB", (10, 10), "white").save(img)
            doc = Document()
            doc.add_paragraph().add_run().add_picture(str(img))
            doc.save(docx)
            self.assertEqual(count_docx_images(docx), 1)
            self.assertEqual(count_docx_images(root / "missing.docx"), 0)


if __name__ == "__main__":
    unittest.main()
