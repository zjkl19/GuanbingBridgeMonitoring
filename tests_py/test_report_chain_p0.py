from __future__ import annotations

import hashlib
import json
import os
import sys
import tempfile
import types
import unittest
from pathlib import Path
from unittest.mock import Mock, patch

from docx import Document
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
REPORTING = ROOT / "reporting"
for candidate in (ROOT, REPORTING):
    if str(candidate) not in sys.path:
        sys.path.insert(0, str(candidate))

from analysis_manifest import (  # noqa: E402
    manifest_artifact_paths,
    pinned_analysis_manifest_scope,
)
from report_job import ReportJobRequest, execute_report_job  # noqa: E402
from report_job_cli import run_context  # noqa: E402
from report_qc import check_shuixianhua_report  # noqa: E402
from template_precheck import TemplatePrecheckError  # noqa: E402
from tests_py.locked_docx_media_test_utils import create_minimal_docx  # noqa: E402
from word_pdf_export import export_authoritative_word_pdf  # noqa: E402


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def _strict_request(root: Path, template: Path, config: Path, manifest: Path) -> ReportJobRequest:
    output = root / "output"
    output.mkdir(exist_ok=True)
    return ReportJobRequest(
        report_type="guanbing_monthly",
        template=template,
        config_path=config,
        result_root=root,
        analysis_root=ROOT,
        output_dir=output,
        period_label="2026-05",
        monitoring_range="2026-05-01 to 2026-05-31",
        report_date="2026-06-01",
        start_date="2026-05-01",
        end_date="2026-05-31",
        analysis_manifest_path=manifest,
        analysis_manifest_sha256=_sha256(manifest),
        require_source_provenance=True,
    )


class ReportChainP0Tests(unittest.TestCase):
    def test_strict_job_rechecks_manifest_hash_before_builder(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            template = create_minimal_docx(root / "template.docx")
            config = root / "config.json"
            config.write_text("{}", encoding="utf-8")
            manifest = root / "analysis.json"
            manifest.write_text('{"status":"ok"}', encoding="utf-8")
            request = _strict_request(root, template, config, manifest)
            manifest.write_text('{"status":"changed"}', encoding="utf-8")

            with patch("report_job.build_guanbing_monthly_report") as builder:
                with self.assertRaisesRegex(ValueError, "SHA-256 mismatch"):
                    execute_report_job(request)

            builder.assert_not_called()

    def test_strict_job_runs_template_precheck_before_builder(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            template = create_minimal_docx(root / "incomplete-template.docx")
            config = root / "config.json"
            config.write_text("{}", encoding="utf-8")
            manifest = root / "analysis.json"
            manifest.write_text('{"status":"ok"}', encoding="utf-8")
            request = _strict_request(root, template, config, manifest)

            with patch("report_job.build_guanbing_monthly_report") as builder:
                with self.assertRaises(TemplatePrecheckError):
                    execute_report_job(request)

            builder.assert_not_called()

    def test_strict_artifact_list_rejects_same_size_hash_tampering(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image = root / "plots" / "A1.jpg"
            image.parent.mkdir()
            image.write_bytes(b"approved")
            record = {
                "kind": "figure",
                "path": str(image),
                "exists": True,
                "bytes": image.stat().st_size,
                "sha256": _sha256(image),
            }
            manifest_payload = {
                "module_results": [
                    {"key": "acceleration", "status": "ok", "artifacts": [record]}
                ]
            }
            manifest = root / "analysis.json"
            manifest.write_text(json.dumps(manifest_payload), encoding="utf-8")
            image.write_bytes(b"tampered")  # same byte length as the approved payload

            with pinned_analysis_manifest_scope(
                manifest,
                _sha256(manifest),
                require_source_provenance=True,
                result_root=root,
            ):
                with self.assertRaisesRegex(ValueError, "artifact SHA-256 mismatch"):
                    manifest_artifact_paths(
                        manifest_payload,
                        "acceleration",
                        kind="figure",
                        suffixes=(".jpg",),
                    )

    def test_strict_artifact_list_rejects_record_without_sha256(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image = root / "plots" / "A1.jpg"
            image.parent.mkdir()
            image.write_bytes(b"approved")
            record = {
                "kind": "figure",
                "path": str(image),
                "exists": True,
                "bytes": image.stat().st_size,
            }
            manifest_payload = {
                "module_results": [
                    {"key": "acceleration", "status": "ok", "artifacts": [record]}
                ]
            }
            manifest = root / "analysis.json"
            manifest.write_text(json.dumps(manifest_payload), encoding="utf-8")

            with pinned_analysis_manifest_scope(
                manifest,
                _sha256(manifest),
                require_source_provenance=True,
                result_root=root,
            ):
                with self.assertRaisesRegex(ValueError, "missing its SHA-256"):
                    manifest_artifact_paths(
                        manifest_payload,
                        "acceleration",
                        kind="figure",
                        suffixes=(".jpg",),
                    )

    def test_strict_job_rejects_missing_or_external_builder_output(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            template = create_minimal_docx(root / "template.docx")
            config = root / "config.json"
            config.write_text("{}", encoding="utf-8")
            manifest = root / "analysis.json"
            manifest.write_text('{"status":"ok"}', encoding="utf-8")
            request = _strict_request(root, template, config, manifest)
            report_manifest = request.output_dir / "report_build_manifest.json"
            report_manifest.write_text(
                json.dumps({"status": "ok", "missing_count": 0, "warnings": []}),
                encoding="utf-8",
            )

            with patch("report_job.raise_for_template"), patch(
                "report_job.build_guanbing_monthly_report",
                return_value=(request.output_dir / "missing.docx", report_manifest),
            ), patch("report_job.export_authoritative_word_pdf") as exporter:
                with self.assertRaisesRegex(FileNotFoundError, "report DOCX does not exist"):
                    execute_report_job(request)
            exporter.assert_not_called()

            outside = create_minimal_docx(root / "stale-report.docx")
            with patch("report_job.raise_for_template"), patch(
                "report_job.build_guanbing_monthly_report",
                return_value=(outside, report_manifest),
            ), patch("report_job.export_authoritative_word_pdf") as exporter:
                with self.assertRaisesRegex(ValueError, "outside output_dir"):
                    execute_report_job(request)
            exporter.assert_not_called()

    def test_wrong_month_and_stale_media_are_blocking_qc_errors(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            stale = root / "march.png"
            current = root / "may.png"
            Image.new("RGB", (16, 16), "red").save(stale)
            Image.new("RGB", (16, 16), "green").save(current)
            report = root / "report.docx"
            doc = Document()
            doc.add_paragraph("水仙花大桥监测结果")
            doc.add_paragraph("监测时间：2026年3月")
            doc.add_picture(str(stale))
            doc.save(report)

            result = check_shuixianhua_report(
                report,
                expected_period_label="2026年5月",
                expected_image_paths=[current],
            )

            self.assertEqual(result.status, "failed")
            self.assertTrue({"period-mismatch", "missing-expected-image"}.issubset(
                {issue.code for issue in result.issues}
            ))

    def test_word_export_failure_never_publishes_partial_pdf(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            docx = root / "report.docx"
            target = root / "report.pdf"
            docx.write_bytes(b"docx")
            target.write_bytes(b"previous-authoritative-pdf")

            fields = types.SimpleNamespace(Update=Mock())
            document = types.SimpleNamespace(
                StoryRanges=[],
                TablesOfContents=[],
                TablesOfFigures=[],
                TablesOfAuthorities=[],
                Fields=fields,
                Save=Mock(),
                Close=Mock(),
            )

            def fail_after_partial_write(path: str, _format: int) -> None:
                Path(path).write_bytes(b"partial")
                raise RuntimeError("synthetic Word export failure")

            document.ExportAsFixedFormat = Mock(side_effect=fail_after_partial_write)
            word = types.SimpleNamespace(
                Visible=True,
                DisplayAlerts=1,
                Documents=types.SimpleNamespace(Open=Mock(return_value=document)),
                Quit=Mock(),
            )
            client = types.ModuleType("win32com.client")
            client.DispatchEx = Mock(return_value=word)
            win32com = types.ModuleType("win32com")
            win32com.client = client
            pythoncom = types.ModuleType("pythoncom")
            pythoncom.CoInitialize = Mock()
            pythoncom.CoUninitialize = Mock()

            with patch.object(os, "name", "nt"), patch.dict(
                os.environ, {"BMS_NO_WORD": "0"}
            ), patch.dict(
                sys.modules,
                {
                    "pythoncom": pythoncom,
                    "win32com": win32com,
                    "win32com.client": client,
                },
            ):
                result = export_authoritative_word_pdf(docx, target)

            self.assertEqual(result.status, "failed")
            self.assertFalse(result.authoritative)
            self.assertIsNone(result.path)
            self.assertEqual(target.read_bytes(), b"previous-authoritative-pdf")
            self.assertEqual(list(root.glob(".*.word-export.pdf")), [])
            document.Close.assert_called_once_with(SaveChanges=False)
            word.Quit.assert_called_once_with()
            pythoncom.CoUninitialize.assert_called_once_with()

    def test_report_worker_records_failure_in_both_status_files(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            status = root / "status.json"
            result = root / "result.json"
            request = Mock()
            context = Mock()
            context.report.launch_id = ""
            with patch("report_job_cli.JobContext.read", return_value=context), patch(
                "report_job_cli.request_from_context", return_value=request
            ), patch(
                "report_job_cli.execute_report_job",
                side_effect=RuntimeError("synthetic report gate failure"),
            ):
                code = run_context(root / "context.json", status, result)

            self.assertEqual(code, 1)
            status_payload = json.loads(status.read_text(encoding="utf-8"))
            result_payload = json.loads(result.read_text(encoding="utf-8"))
            self.assertEqual(status_payload["state"], "failed")
            self.assertEqual(result_payload["state"], "failed")
            self.assertIn("synthetic report gate failure", status_payload["message"])
            self.assertIn("RuntimeError", result_payload["traceback"])

    def test_report_worker_preserves_launch_id_when_context_is_unreadable(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            context = root / "context.json"
            context.write_text("not-json", encoding="utf-8")
            status = root / "status.json"
            result = root / "result.json"

            code = run_context(
                context, status, result, expected_launch_id="launch-from-command"
            )

            self.assertEqual(code, 1)
            self.assertEqual(
                json.loads(status.read_text(encoding="utf-8"))["launch_id"],
                "launch-from-command",
            )
            self.assertEqual(
                json.loads(result.read_text(encoding="utf-8"))["launch_id"],
                "launch-from-command",
            )

    def test_report_worker_commits_success_result_before_completed_status(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            status = root / "status.json"
            result = root / "result.json"
            context = Mock()
            context.report.launch_id = "report-order"

            def execute(_request, progress):
                progress("completed", 1.0, "builder callback finished")
                observed = json.loads(status.read_text(encoding="utf-8"))
                self.assertEqual(observed["state"], "running")
                self.assertFalse(result.exists())
                return types.SimpleNamespace(
                    report_path=root / "report.docx",
                    pdf_path=None,
                    manifest_path=root / "report_build_manifest.json",
                    missing=(),
                    summary_files=(),
                    qc={"status": "ok"},
                )

            with patch("report_job_cli.JobContext.read", return_value=context), patch(
                "report_job_cli.request_from_context", return_value=Mock()
            ), patch("report_job_cli.execute_report_job", side_effect=execute):
                code = run_context(
                    root / "context.json",
                    status,
                    result,
                    expected_launch_id="report-order",
                )

            self.assertEqual(code, 0)
            self.assertEqual(
                json.loads(result.read_text(encoding="utf-8"))["state"],
                "completed",
            )
            self.assertEqual(
                json.loads(status.read_text(encoding="utf-8"))["state"],
                "completed",
            )


if __name__ == "__main__":
    unittest.main()
