import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from docx import Document  # noqa: E402
from PIL import Image  # noqa: E402
from docx.oxml import OxmlElement, parse_xml  # noqa: E402
from docx.oxml.ns import nsdecls, qn  # noqa: E402
from docx_header_fields import ensure_section_footer_pagination_fields  # noqa: E402
from report_qc import (  # noqa: E402
    check_guanbing_report,
    check_hongtang_report,
    check_jlj_report,
    check_report,
    check_shuixianhua_report,
    write_report_qc_report,
)


def _append_complex_field(paragraph, instruction: str, result: str = "1") -> None:
    def append_field_char(field_type: str) -> None:
        node = OxmlElement("w:fldChar")
        node.set(qn("w:fldCharType"), field_type)
        paragraph.add_run()._r.append(node)

    append_field_char("begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = f" {instruction} "
    paragraph.add_run()._r.append(instruction_node)
    append_field_char("separate")
    paragraph.add_run(result)
    append_field_char("end")


def write_legacy_guanbing_footer_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph(
        "G104 \u7ba1\u67c4\u5927\u6865 "
        "\u62a5\u544a\u7f16\u53f7\uff1aBG02FQJC2400001-M18"
    )
    section = document.sections[0]
    page_numbering = OxmlElement("w:pgNumType")
    page_numbering.set(qn("w:start"), "1")
    section._sectPr.append(page_numbering)
    paragraph = section.footer.paragraphs[0]
    paragraph.add_run("\u7b2c ")
    _append_complex_field(paragraph, "PAGE")
    paragraph.add_run(" \u9875 \u5171 76 \u9875")
    document.save(path)


class TestReportQc(unittest.TestCase):
    @staticmethod
    def _append_report_number_text_box(header, report_number: str) -> None:
        prefix, suffix = report_number[:-2], report_number[-2:]
        header._element.append(
            parse_xml(
                f"""
                <w:p {nsdecls('w')} xmlns:v="urn:schemas-microsoft-com:vml">
                  <w:r><w:pict><v:shape><v:textbox><w:txbxContent>
                    <w:p>
                      <w:r><w:t>报告编号：</w:t></w:r>
                      <w:r><w:t>{prefix}</w:t></w:r>
                      <w:r><w:t>{suffix}</w:t></w:r>
                    </w:p>
                  </w:txbxContent></v:textbox></v:shape></w:pict></w:r>
                </w:p>
                """
            )
        )

    def test_guanbing_qc_rejects_static_total_and_accepts_sectionpages(self):
        with tempfile.TemporaryDirectory() as tmp:
            legacy = Path(tmp) / "guanbing-static-pagination.docx"
            output = Path(tmp) / "guanbing-dynamic-pagination.docx"
            write_legacy_guanbing_footer_fixture(legacy)
            invalid = check_guanbing_report(legacy)
            self.assertIn(
                "invalid-footer-pagination",
                {issue.code for issue in invalid.issues},
            )

            document = Document(legacy)
            self.assertGreaterEqual(
                ensure_section_footer_pagination_fields(document),
                1,
            )
            document.save(output)
            valid = check_guanbing_report(output)

        self.assertNotIn("invalid-footer-pagination", {issue.code for issue in valid.issues})

    def test_guanbing_qc_rejects_mixed_report_number_in_header_text_box(self):
        with tempfile.TemporaryDirectory() as tmp:
            report = Path(tmp) / "mixed-report-number.docx"
            doc = Document()
            doc.add_paragraph("G104 管柄大桥　报告编号：BG02FQJC2400001-M18")
            self._append_report_number_text_box(
                doc.sections[0].first_page_header,
                "BG02FQJC2400001-M16",
            )
            doc.save(report)

            result = check_guanbing_report(report)

        codes = {issue.code for issue in result.issues}
        self.assertIn("inconsistent-report-number", codes)
        self.assertEqual(
            result.summary["report_numbers"],
            ["BG02FQJC2400001-M16", "BG02FQJC2400001-M18"],
        )
        self.assertEqual(result.status, "failed")

    def test_guanbing_qc_rejects_missing_report_number(self):
        with tempfile.TemporaryDirectory() as tmp:
            report = Path(tmp) / "missing-report-number.docx"
            doc = Document()
            doc.add_paragraph("G104 管柄大桥月报")
            doc.save(report)

            result = check_guanbing_report(report)

        codes = {issue.code for issue in result.issues}
        self.assertIn("missing-report-number", codes)
        self.assertEqual(result.summary["report_numbers"], [])
        self.assertEqual(result.summary["report_number_occurrence_count"], 0)
        self.assertEqual(result.status, "failed")

    def test_guanbing_qc_accepts_one_report_number_across_text_boxes(self):
        with tempfile.TemporaryDirectory() as tmp:
            report = Path(tmp) / "one-report-number.docx"
            doc = Document()
            doc.add_paragraph("G104 管柄大桥　报告编号：BG02FQJC2400001-M18")
            self._append_report_number_text_box(
                doc.sections[0].even_page_header,
                "BG02FQJC2400001-M18",
            )
            doc.save(report)

            result = check_guanbing_report(report)

        codes = {issue.code for issue in result.issues}
        self.assertNotIn("inconsistent-report-number", codes)
        self.assertEqual(result.summary["report_numbers"], ["BG02FQJC2400001-M18"])
        self.assertEqual(result.summary["report_number_occurrence_count"], 2)

    def test_jlj_qc_detects_forbidden_phrase_and_front_summary(self):
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "report.docx"
            doc = Document()
            cover = doc.add_table(rows=2, cols=2)
            cover.cell(0, 0).text = "委托单位"
            cover.cell(1, 0).text = "监测结果"
            cover.cell(1, 1).text = "建议结合原始数据进一步复核\n（转下页）"
            cont = doc.add_table(rows=1, cols=2)
            cont.cell(0, 0).text = "监测结果"
            cont.cell(0, 1).text = "（续上页）\n后续内容"
            doc.save(docx)

            result = check_jlj_report(docx)

            self.assertEqual(result.summary["front_summary_table_indices"], [0, 1])
            self.assertTrue(any(issue.code == "forbidden-review-phrase" for issue in result.issues))

            txt_path, json_path = write_report_qc_report(result, tmp, timestamp="20260101_000000")
            self.assertTrue(txt_path.exists())
            self.assertTrue(json_path.exists())

    def test_jlj_qc_reports_summary_table_outside_front_block(self):
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "report.docx"
            doc = Document()
            cover = doc.add_table(rows=2, cols=2)
            cover.cell(0, 0).text = "委托单位"
            cover.cell(1, 0).text = "监测结果"
            cover.cell(1, 1).text = "首页"
            doc.add_table(rows=1, cols=1).cell(0, 0).text = "正文表"
            stale = doc.add_table(rows=1, cols=2)
            stale.cell(0, 0).text = "监测结果"
            stale.cell(0, 1).text = "错位"
            doc.save(docx)

            result = check_jlj_report(docx)

            self.assertTrue(any(issue.code == "summary-table-outside-front-block" for issue in result.issues))

    def test_jlj_qc_rejects_stale_period_and_missing_current_image(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            stale_image = root / "march-result.png"
            expected_image = root / "may-result.png"
            Image.new("RGB", (24, 24), "red").save(stale_image)
            Image.new("RGB", (24, 24), "green").save(expected_image)
            report = root / "jlj.docx"
            doc = Document()
            cover = doc.add_table(rows=2, cols=2)
            cover.cell(0, 0).text = "委托单位"
            cover.cell(1, 0).text = "监测结果"
            doc.add_paragraph("监测时间：2026年03月")
            doc.add_picture(str(stale_image))
            doc.save(report)

            result = check_jlj_report(
                report,
                expected_period_label="2026年5月",
                expected_image_paths=[expected_image],
            )

            codes = {issue.code for issue in result.issues}
            self.assertIn("period-mismatch", codes)
            self.assertIn("missing-expected-image", codes)
            self.assertEqual(result.status, "failed")

    def test_generic_qc_supports_hongtang_and_guanbing(self):
        with tempfile.TemporaryDirectory() as tmp:
            hongtang = Path(tmp) / "hongtang.docx"
            doc = Document()
            doc.add_paragraph("监测结果 交通状况监测 结构应变监测")
            doc.save(hongtang)
            ht_result = check_hongtang_report(hongtang)
            self.assertEqual(ht_result.kind, "hongtang_period")

            guanbing = Path(tmp) / "guanbing.docx"
            doc = Document()
            doc.add_paragraph("G104 管柄大桥 m/s2")
            doc.save(guanbing)
            gb_result = check_guanbing_report(guanbing)
            self.assertEqual(gb_result.kind, "guanbing_monthly")
            self.assertTrue(any(issue.code == "unit-superscript-risk" for issue in gb_result.issues))
            self.assertEqual(check_report("guanbing_monthly", guanbing).kind, "guanbing_monthly")

    def test_shuixianhua_qc_requires_bridge_and_monitoring_text(self):
        with tempfile.TemporaryDirectory() as tmp:
            report = Path(tmp) / "shuixianhua.docx"
            doc = Document()
            doc.add_paragraph("水仙花大桥监测结果")
            doc.save(report)

            result = check_shuixianhua_report(report)

            self.assertEqual(result.kind, "shuixianhua_monthly")
            self.assertFalse([issue for issue in result.issues if issue.code == "missing-expected-text"])
            self.assertEqual(check_report("shuixianhua_monthly", report).kind, "shuixianhua_monthly")

    def test_shuixianhua_qc_accepts_expected_period_and_embedded_source_image(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            expected_image = root / "may-result.png"
            Image.new("RGB", (24, 24), "green").save(expected_image)
            report = root / "shuixianhua.docx"
            doc = Document()
            doc.add_paragraph("水仙花大桥监测结果")
            doc.add_paragraph("监测时间：2026年5月")
            doc.add_paragraph("监测时间：2026年5月")
            doc.add_picture(str(expected_image))
            doc.save(report)

            result = check_shuixianhua_report(
                report,
                expected_period_label="2026年5月",
                expected_image_paths=[expected_image],
            )

            codes = {issue.code for issue in result.issues}
            self.assertNotIn("period-mismatch", codes)
            self.assertNotIn("missing-expected-image", codes)

    def test_shuixianhua_qc_rejects_stale_period_and_missing_source_image(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            stale_image = root / "march-result.png"
            expected_image = root / "may-result.png"
            Image.new("RGB", (24, 24), "red").save(stale_image)
            Image.new("RGB", (24, 24), "green").save(expected_image)
            report = root / "shuixianhua.docx"
            doc = Document()
            doc.add_paragraph("水仙花大桥监测结果")
            doc.add_paragraph("监测时间：2026年03月")
            doc.add_picture(str(stale_image))
            doc.save(report)

            result = check_shuixianhua_report(
                report,
                expected_period_label="2026年5月",
                expected_image_paths=[expected_image],
            )

            codes = {issue.code for issue in result.issues}
            self.assertIn("period-mismatch", codes)
            self.assertIn("missing-expected-image", codes)
            self.assertEqual(result.status, "failed")


if __name__ == "__main__":
    unittest.main()
