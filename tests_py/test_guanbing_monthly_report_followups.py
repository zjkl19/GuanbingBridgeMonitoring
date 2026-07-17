from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls


ROOT = Path(__file__).resolve().parents[1]
REPORTING = ROOT / "reporting"
if str(REPORTING) not in sys.path:
    sys.path.insert(0, str(REPORTING))

from build_guanbing_monthly_report import (  # noqa: E402
    apply_image_updates,
    apply_text_updates,
    is_sample_monitoring_period,
    iter_all_paragraphs,
    normalize_acceleration_units,
    normalize_report_number,
    resolve_monitoring_range,
)


class TestGuanbingMonthlyReportFollowups(unittest.TestCase):
    @staticmethod
    def _append_report_number_text_box(header, report_number: str) -> None:
        prefix, suffix = report_number[:-2], report_number[-2:]
        header._element.append(
            parse_xml(
                f"""
                <w:p {nsdecls('w')} xmlns:v="urn:schemas-microsoft-com:vml">
                  <w:r><w:pict><v:shape><v:textbox><w:txbxContent>
                    <w:p>
                      <w:r><w:t>管柄月报　报告编号：</w:t></w:r>
                      <w:r><w:rPr><w:b/></w:rPr><w:t>{prefix}</w:t></w:r>
                      <w:r><w:rPr><w:i/></w:rPr><w:t>{suffix}</w:t></w:r>
                    </w:p>
                  </w:txbxContent></v:textbox></v:shape></w:pict></w:r>
                </w:p>
                """
            )
        )

    @staticmethod
    def _package_xml_text(doc: Document) -> str:
        roots = [doc.element]
        roots.extend(
            part.element
            for part in doc.part.package.parts
            if str(part.partname).startswith(("/word/header", "/word/footer"))
            and getattr(part, "element", None) is not None
        )
        return "\n".join(
            "".join(node.text or "" for node in root.xpath(".//w:t"))
            for root in roots
        )

    def test_report_replaces_template_march_operation_and_maintenance_text(self):
        doc = Document()
        doc.add_paragraph("（时间范围：2026年02月26日~2026年03月25日）")
        doc.add_paragraph("报告日期：2026年04月10日")
        doc.add_paragraph(
            "G104线管柄大桥健康监测系统于2024年9月26日交工验收后进入运维服务期，"
            "至2026年03月25日监测系统已运行18个月，本月监测周期内健康监测系统运行状况如表 3所示。"
            "本月系统运行过程中，2026年03月09日5:20~03月16日20:55由于监测系统软件故障，数据传输中断。"
        )
        doc.add_paragraph(
            "软件线上检查维护：本月每日对系统运行情况进行检查，"
            "2026年03月09日5:20~03月16日20:55监测系统软件故障，数据传输中断，经维护后恢复正常。"
        )
        doc.add_paragraph(
            "（1）本月月度监测周期内，2026年03月09日5:20~03月16日20:55由于监测系统软件故障，"
            "数据传输中断，其余时间监测传感器实时在线。"
        )
        calendar = doc.add_table(rows=2, cols=7)
        for cell, value in zip(
            calendar.rows[0].cells,
            ["一", "二", "三", "四", "五", "六", "日"],
        ):
            cell.text = value
        for cell, value in zip(
            calendar.rows[1].cells,
            ["2月26日", "2月27日", "2月28日", "3月1日", "3月2日", "3月3日", "3月4日"],
        ):
            cell.text = value
        legend = doc.add_table(rows=1, cols=2)
        legend.cell(0, 0).text = "监测系统运行正常"
        legend.cell(0, 1).text = "监测系统运行故障"

        maintenance_picture = doc.add_paragraph()
        maintenance_picture.add_run()._r.append(OxmlElement("w:drawing"))
        doc.add_paragraph("图 2 系统维护情况")
        doc.add_paragraph(
            "预警信息处理：本月系统共计发送1114条报警信息，"
            "主要原因为2026年3月21日17:30~3月22日23:23数据异常。"
        )
        for _ in range(3):
            picture = doc.add_paragraph()
            picture.add_run()._r.append(OxmlElement("w:drawing"))
        doc.add_paragraph("图 3 部分预警信息处理情况")
        doc.add_paragraph(
            "现场硬件维护：本月21日17:30~22日23:23拟对现场硬件设备进行排查。"
        )

        monitoring_range = "2026年5月26日至2026年5月28日（源数据不完整）"
        updated = apply_text_updates(
            doc,
            {},
            monitoring_range,
            "2026年7月17日",
            start_date="2026-05-26",
            end_date="2026-05-28",
            sample_period=True,
        )

        all_text = "\n".join(paragraph.text for paragraph in iter_all_paragraphs(doc))
        self.assertNotIn("2026年03月", all_text)
        self.assertNotIn("03月09日", all_text)
        self.assertNotIn("03月16日", all_text)
        self.assertNotIn("其余时间监测传感器实时在线", all_text)
        self.assertNotIn("1114", all_text)
        self.assertNotIn("2026年3月21日", all_text)
        self.assertNotIn("本月21日", all_text)
        self.assertNotIn("本月", all_text)
        self.assertNotIn("监测系统运行故障", all_text)
        self.assertFalse(doc.element.body.xpath(".//w:drawing"))
        self.assertIn(f"本期报告监测范围为{monitoring_range}", all_text)
        self.assertIn("不根据数据缺口推断系统故障", all_text)
        self.assertIn("不沿用模板中的历史报警数量、日期或原因", all_text)
        self.assertIn("历史截图已移除", all_text)
        self.assertIn("图 2 系统维护情况（本期未提供对应记录）", all_text)
        self.assertIn("图 3 部分预警信息处理情况（本期未提供对应记录）", all_text)
        self.assertIn("system operation scope", updated)
        self.assertIn("software maintenance scope", updated)
        self.assertIn("conclusion operation scope", updated)
        self.assertIn("warning handling scope", updated)
        self.assertIn("hardware maintenance scope", updated)
        self.assertIn("operation calendar tables (2)", updated)
        self.assertIn("maintenance screenshots (1)", updated)
        self.assertIn("warning screenshots (3)", updated)

    def test_empty_monitoring_range_uses_requested_dates_not_template_dates(self):
        self.assertEqual(
            resolve_monitoring_range("", "2026-05-26", "2026-05-28"),
            "2026年05月26日~2026年05月28日",
        )

    def test_normalize_acceleration_units_handles_split_legacy_runs(self):
        doc = Document()
        for index in range(200):
            doc.add_paragraph(f"前置段落 {index}")
        table = doc.add_table(rows=1, cols=1)
        paragraph = table.cell(0, 0).paragraphs[0]
        paragraph.add_run("阈值31.5cm/s")
        exponent = paragraph.add_run("2")
        exponent.font.superscript = True
        paragraph.add_run("，另一个为1m/s2。")

        self.assertEqual(normalize_acceleration_units(doc), 2)
        self.assertEqual(paragraph.text, "阈值31.5cm/s²，另一个为1m/s²。")
        self.assertNotIn("m/s2", paragraph.text)

    def test_report_number_is_derived_from_cover_and_applied_to_split_header_runs(self):
        doc = Document()
        cover = doc.add_paragraph()
        cover.add_run("报告编号：")
        cover.add_run("BG02FQJC2400001-M")
        cover.add_run("18")
        header = doc.sections[0].header.paragraphs[0]
        header.add_run("管柄大桥月报     报告编号：")
        header_prefix = header.add_run("BG02FQJC2400001-M")
        header_prefix.bold = True
        header.add_run("16")

        report_number, changed = normalize_report_number(doc)

        self.assertEqual(report_number, "BG02FQJC2400001-M18")
        self.assertEqual(changed, 1)
        self.assertIn("报告编号：BG02FQJC2400001-M18", cover.text)
        self.assertIn("报告编号：BG02FQJC2400001-M18", header.text)
        self.assertTrue(header_prefix.bold)
        self.assertNotIn("M16", header.text)

    def test_report_number_fails_closed_when_template_has_none(self):
        doc = Document()
        doc.add_paragraph("G104 管柄大桥月报")

        with self.assertRaisesRegex(ValueError, "Report number is required"):
            normalize_report_number(doc)

    def test_explicit_report_number_overrides_cover_and_all_headers(self):
        doc = Document()
        doc.add_paragraph("报告编号：BG02FQJC2400001-M18")
        doc.sections[0].header.paragraphs[0].add_run(
            "报告编号：BG02FQJC2400001-M16"
        )

        report_number, changed = normalize_report_number(
            doc,
            "BG02FQJC2600001-M21",
        )

        self.assertEqual(report_number, "BG02FQJC2600001-M21")
        self.assertEqual(changed, 2)
        all_text = "\n".join(paragraph.text for paragraph in iter_all_paragraphs(doc))
        self.assertNotIn("M18", all_text)
        self.assertNotIn("M16", all_text)
        self.assertEqual(all_text.count("BG02FQJC2600001-M21"), 2)

    def test_report_number_normalizes_split_text_box_in_first_page_header(self):
        doc = Document()
        doc.add_paragraph("报告编号：BG02FQJC2400001-M18")
        first_page_header = doc.sections[0].first_page_header
        self._append_report_number_text_box(
            first_page_header,
            "BG02FQJC2400001-M16",
        )

        report_number, changed = normalize_report_number(doc)

        self.assertEqual(report_number, "BG02FQJC2400001-M18")
        self.assertEqual(changed, 1)
        package_text = self._package_xml_text(doc)
        self.assertNotIn("BG02FQJC2400001-M16", package_text)
        self.assertEqual(package_text.count("BG02FQJC2400001-M18"), 2)
        bold_number_nodes = first_page_header._element.xpath(
            ".//w:r[w:rPr/w:b]/w:t"
        )
        self.assertEqual(len(bold_number_nodes), 1)
        self.assertEqual(bold_number_nodes[0].text, "BG02FQJC2400001-M18")

    def test_three_day_range_uses_sample_wording_but_full_month_does_not(self):
        self.assertTrue(
            is_sample_monitoring_period(
                "2026年5月样本",
                "2026年5月26日至5月28日",
                "2026-05-26",
                "2026-05-28",
            )
        )
        self.assertTrue(
            is_sample_monitoring_period(
                "2026年5月",
                "2026年5月26日至5月28日",
                "2026-05-26",
                "2026-05-28",
            )
        )
        self.assertFalse(
            is_sample_monitoring_period(
                "2026年5月",
                "2026年5月1日至5月31日",
                "2026-05-01",
                "2026-05-31",
            )
        )

    def test_real_template_content_gate_removes_fixed_march_material_in_memory(self):
        template = ROOT / "reports" / "G104线管柄大桥监测月报模板-自动报告.docx"
        doc = Document(str(template))
        drawings_before = len(doc.element.body.xpath(".//w:drawing"))

        updated = apply_text_updates(
            doc,
            {},
            "2026年5月26日至2026年5月28日（源数据不完整）",
            "2026年7月17日",
            start_date="2026-05-26",
            end_date="2026-05-28",
            sample_period=True,
        )
        report_number, changed = normalize_report_number(doc)

        all_text = "\n".join(paragraph.text for paragraph in iter_all_paragraphs(doc))
        self.assertNotIn("1114", all_text)
        self.assertNotIn("2026年3月21日", all_text)
        self.assertNotIn("本月21日", all_text)
        self.assertNotIn("监测系统运行故障", all_text)
        self.assertNotIn("2月26日", all_text)
        self.assertNotIn("本月", all_text)
        self.assertIn("warning handling scope", updated)
        self.assertIn("hardware maintenance scope", updated)
        self.assertEqual(drawings_before - len(doc.element.body.xpath(".//w:drawing")), 4)
        self.assertEqual(report_number, "BG02FQJC2400001-M18")
        self.assertGreaterEqual(changed, 1)
        self.assertNotIn("BG02FQJC2400001-M16", all_text)
        package_text = self._package_xml_text(doc)
        self.assertNotIn("BG02FQJC2400001-M16", package_text)
        self.assertGreaterEqual(package_text.count("BG02FQJC2400001-M18"), 4)

    def test_report_uses_group_directories_for_tilt_and_lowpass_strain(self):
        calls: list[tuple[str, str]] = []

        def fake_find(_root: Path, configured_dir: str, prefix: str) -> Path:
            calls.append((configured_dir, prefix))
            return Path(configured_dir) / f"{prefix}.jpg"

        with tempfile.TemporaryDirectory() as tmp, patch(
            "build_guanbing_monthly_report.find_latest_image", side_effect=fake_find
        ), patch(
            "build_guanbing_monthly_report.build_accel_combined_image", return_value=Path(tmp) / "accel.jpg"
        ), patch(
            "build_guanbing_monthly_report.replace_picture_before_anchor",
            return_value=(True, "mock.jpg"),
        ):
            apply_image_updates(Document(), Path(tmp), Path(tmp) / "assets")

        self.assertIn(("时程曲线_倾角_组图", "Tilt_X"), calls)
        self.assertIn(("时程曲线_倾角_组图", "Tilt_Y"), calls)
        self.assertIn(("时程曲线_动应变_低通滤波_组图", "dynstrain_lp_G05"), calls)
        self.assertIn(("时程曲线_动应变_低通滤波_组图", "dynstrain_lp_G06"), calls)


if __name__ == "__main__":
    unittest.main()
