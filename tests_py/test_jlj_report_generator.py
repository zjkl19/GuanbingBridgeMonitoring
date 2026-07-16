import hashlib
import json
import os
import shutil
import sys
import unittest
from unittest.mock import patch
from datetime import date
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from openpyxl import Workbook

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from artifact_lookup import ArtifactLookupResult  # noqa: E402
from analysis_manifest import (  # noqa: E402
    pinned_analysis_manifest_scope,
    pinned_derived_artifact_manifest_scope,
)
from build_jlj_monthly_report import (  # noqa: E402
    JLJ_CABLE_FORCE_UNVERIFIED_DISCLOSURE,
    build_cable_section,
    build_main_bearing_section,
    build_report,
    build_wind_section,
    clean_jlj_report_xml_text,
    collect_jlj_data_acquisition_rows,
    find_latest_two_deflection_images,
    find_latest_point_image_patterns,
    find_wind_summary_file,
    jlj_image_matches_report_period,
    jlj_cable_force_report_policy,
    jlj_report_period_scope,
    normalize_cover_monitoring_time,
    read_stats_rows,
    summarize_jlj_pinned_source_coverage,
    update_jlj_warning_threshold_table,
    validate_jlj_analysis_profile,
    validate_jlj_analysis_period,
)


def write_plot_provenance(path: Path, incomplete_days: list[str]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(
            {
                "schema_version": 1,
                "series": [
                    {
                        "sampling_mode": "full",
                        "reduction_applied": False,
                        "input_count": 10,
                        "finite_count": 9,
                        "plotted_finite_count": 9,
                        "source": {
                            "source_sample_count": 10,
                            "finite_source_sample_count": 9,
                            "completeness_scope": "required_export_contribution",
                            "internal_gap_coverage_assessed": True,
                            "calendar_day_count_requested": 31,
                            "complete_day_count": 31 - len(incomplete_days),
                            "incomplete_day_count": len(incomplete_days),
                            "incomplete_days": incomplete_days,
                            "missing_required_sources": [
                                f"rolling_export:{value}" for value in incomplete_days
                            ],
                        },
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return path


class TestJljReportGenerator(unittest.TestCase):
    def setUp(self):
        self.tmp = Path("tmp") / "test_jlj_report_generator"
        shutil.rmtree(self.tmp, ignore_errors=True)
        self.tmp.mkdir(parents=True)

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_cover_monitoring_time_uses_report_month(self):
        self.assertEqual(
            normalize_cover_monitoring_time("2026年3月份", "2026.03.23~2026.03.31"),
            "2026年3月",
        )
        self.assertEqual(
            normalize_cover_monitoring_time("", "2026.03.23~2026.03.31"),
            "2026年3月",
        )

    def test_data_acquisition_includes_pending_expansion_joint_points(self):
        rows = collect_jlj_data_acquisition_rows({}, self.tmp, date(2026, 3, 23), date(2026, 3, 31))
        expansion_rows = [row for row in rows if row["module"] == "伸缩缝位移监测"]

        self.assertEqual(len(expansion_rows), 1)
        row = expansion_rows[0]
        self.assertEqual(row["design_count"], "4")
        self.assertEqual(row["acquired_count"], "0")
        self.assertEqual(row["rate"], "0.0%")
        self.assertEqual(row["date_range"], "/")
        self.assertIn("新版竣工图新增测点SSFWYJ-01~04", row["remarks"])

    def test_warning_threshold_table_fixed_text_and_arch_lower_bound(self):
        doc = Document()
        table = doc.add_table(rows=4, cols=6)
        for col_idx, value in enumerate(["报警类别", "监测项", "报警规则", "统计值", "上限", "下限"]):
            table.cell(0, col_idx).text = value
        table.cell(1, 1).text = "主梁跨中静应变"
        table.cell(1, 2).text = "超过设计值"
        table.cell(2, 1).text = "主拱拱顶静应变"
        table.cell(2, 2).text = "超过设计值"
        table.cell(3, 1).text = "拱桥主拱拱顶位移"
        table.cell(3, 2).text = "超过0.8倍设计值"
        table.cell(3, 5).text = "-57.4mm"

        update_jlj_warning_threshold_table(doc)

        self.assertEqual(table.cell(3, 5).text, "-59.4mm")
        self.assertIn("MPa", table.cell(1, 4).text)
        self.assertIn("με", table.cell(2, 5).text)

    def test_cleanup_normalizes_accepted_report_text_typos(self):
        path = self.tmp / "cleanup.docx"
        doc = Document()
        doc.add_paragraph("10min加速度速度均方根；桥墩沉降°；DYBCQG-24-K16-ZGD-A20。")
        doc.save(path)

        clean_jlj_report_xml_text(path)

        cleaned = "\n".join(p.text for p in Document(path).paragraphs)
        self.assertIn("10min加速度均方根", cleaned)
        self.assertIn("桥墩沉降", cleaned)
        self.assertIn("DYBCGQ-24-K16-ZGD-A20", cleaned)
        self.assertNotIn("DYBCQG", cleaned)

    def test_missing_optional_stats_returns_empty_rows(self):
        rows = read_stats_rows(self.tmp / "stats", "strain_stats.xlsx", self.tmp)
        self.assertEqual(rows, [])

    def test_deflection_images_use_split_raw_filtered_dirs(self):
        raw_dir = self.tmp / "时程曲线_挠度_原始"
        filt_dir = self.tmp / "时程曲线_挠度_滤波"
        raw_dir.mkdir()
        filt_dir.mkdir()
        raw = raw_dir / "Defl_NDY-01_Orig_20260301_20260331.jpg"
        filt = filt_dir / "Defl_NDY-01_Filt_20260301_20260331.jpg"
        raw.write_bytes(b"raw")
        filt.write_bytes(b"filt")

        self.assertEqual(find_latest_two_deflection_images(self.tmp, "NDY-01"), (raw.resolve(), filt.resolve()))

    def test_period_filter_prefers_current_month_over_newer_stale_file(self):
        folder = self.tmp / "figures"
        folder.mkdir()
        stale = folder / "P-01_20260301_20260331.jpg"
        current = folder / "P-01_20260501_20260531.jpg"
        stale.write_bytes(b"stale")
        current.write_bytes(b"current")
        os.utime(current, (1, 1))
        os.utime(stale, (2, 2))

        with jlj_report_period_scope(date(2026, 5, 1), date(2026, 5, 31)):
            selected = find_latest_point_image_patterns(
                self.tmp,
                "figures",
                "P-01",
                ["P-01_*.jpg"],
            )

        self.assertEqual(selected, current.resolve())
        self.assertTrue(
            jlj_image_matches_report_period(current, date(2026, 5, 1), date(2026, 5, 31))
        )
        self.assertFalse(
            jlj_image_matches_report_period(stale, date(2026, 5, 1), date(2026, 5, 31))
        )

    def test_strict_pinned_manifest_accepts_in_month_partial_range_only(self):
        folder = self.tmp / "时程曲线_加速度"
        folder.mkdir()

        for filename, expected in (
            ("A1_20260509_20260531.jpg", True),
            ("A1_20260309_20260331.jpg", False),
            ("A1_20260509_20260601.jpg", False),
        ):
            with self.subTest(filename=filename):
                image = folder / filename
                image.write_bytes(filename.encode("ascii"))
                artifact = {
                    "kind": "figure",
                    "role": "time_history",
                    "path": str(image.resolve()),
                    "exists": True,
                    "bytes": image.stat().st_size,
                    "sha256": hashlib.sha256(image.read_bytes()).hexdigest().upper(),
                }
                manifest = self.tmp / f"analysis_{image.stem}.json"
                manifest.write_text(
                    json.dumps(
                        {
                            "module_results": [
                                {
                                    "key": "acceleration",
                                    "status": "ok",
                                    "artifacts": [artifact],
                                }
                            ]
                        }
                    ),
                    encoding="utf-8",
                )
                manifest_hash = hashlib.sha256(manifest.read_bytes()).hexdigest().upper()
                with pinned_analysis_manifest_scope(
                    manifest,
                    manifest_hash,
                    require_source_provenance=True,
                    result_root=self.tmp,
                ), jlj_report_period_scope(date(2026, 5, 1), date(2026, 5, 31)):
                    selected = find_latest_point_image_patterns(
                        self.tmp,
                        "时程曲线_加速度",
                        "A1",
                        ["A1_*.jpg"],
                    )

                self.assertEqual(selected, image.resolve() if expected else None)
                image.unlink()

    def test_unpinned_manifest_does_not_relax_partial_range_gate(self):
        folder = self.tmp / "时程曲线_加速度"
        folder.mkdir()
        image = folder / "A1_20260509_20260531.jpg"
        image.write_bytes(b"unpinned-partial")
        run_logs = self.tmp / "run_logs"
        run_logs.mkdir()
        (run_logs / "analysis_manifest_1.json").write_text(
            json.dumps(
                {
                    "module_results": [
                        {
                            "key": "acceleration",
                            "status": "ok",
                            "artifacts": [
                                {
                                    "kind": "figure",
                                    "role": "time_history",
                                    "path": str(image.resolve()),
                                }
                            ],
                        }
                    ]
                }
            ),
            encoding="utf-8",
        )

        with jlj_report_period_scope(date(2026, 5, 1), date(2026, 5, 31)):
            selected = find_latest_point_image_patterns(
                self.tmp,
                "时程曲线_加速度",
                "A1",
                ["A1_*.jpg"],
            )

        self.assertIsNone(selected)

    def test_strict_derived_manifest_accepts_in_month_partial_range(self):
        folder = self.tmp / "时程曲线_加速度"
        folder.mkdir()
        image = folder / "A1_20260509_20260531.jpg"
        image.write_bytes(b"partial-derived")

        analysis = self.tmp / "analysis.json"
        analysis.write_text(json.dumps({"module_results": []}), encoding="utf-8")
        analysis_hash = hashlib.sha256(analysis.read_bytes()).hexdigest().upper()
        derived = self.tmp / "derived.json"
        derived.write_text(
            json.dumps(
                {
                    "manifest_type": "derived_artifact_manifest",
                    "analysis_manifest": {
                        "path": str(analysis.resolve()),
                        "sha256": analysis_hash,
                    },
                    "result_root": str(self.tmp.resolve()),
                    "artifacts": [
                        {
                            "kind": "figure",
                            "role": "time_history",
                            "path": str(image.resolve()),
                            "bytes": image.stat().st_size,
                            "sha256": hashlib.sha256(image.read_bytes()).hexdigest().upper(),
                        }
                    ],
                }
            ),
            encoding="utf-8",
        )
        derived_hash = hashlib.sha256(derived.read_bytes()).hexdigest().upper()

        with pinned_analysis_manifest_scope(
            analysis,
            analysis_hash,
            require_source_provenance=True,
            result_root=self.tmp,
        ), pinned_derived_artifact_manifest_scope(
            derived,
            derived_hash,
            require_source_provenance=True,
        ), jlj_report_period_scope(date(2026, 5, 1), date(2026, 5, 31)):
            selected = find_latest_point_image_patterns(
                self.tmp,
                "时程曲线_加速度",
                "A1",
                ["A1_*.jpg"],
            )

        self.assertEqual(selected, image.resolve())

    def test_main_bearing_section_uses_strict_filtered_directory(self):
        stats = self.tmp / "stats"
        figures = self.tmp / "时程曲线_支座位移_滤波"
        stats.mkdir()
        figures.mkdir()
        point_id = "WYJ-01-K15-ZF-G14"

        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["PointID", "FiltMin_mm", "FiltMax_mm", "FiltMean_mm"])
        sheet.append([point_id, -0.5, 0.6, 0.1])
        stats_file = stats / "bearing_displacement_stats.xlsx"
        workbook.save(stats_file)
        image = figures / f"BearingDisp_{point_id}_20260501_20260531_Filt.jpg"
        image.write_bytes(b"filtered-bearing")

        def record(path: Path, kind: str, role: str) -> dict:
            return {
                "kind": kind,
                "role": role,
                "path": str(path.resolve()),
                "exists": True,
                "bytes": path.stat().st_size,
                "sha256": hashlib.sha256(path.read_bytes()).hexdigest().upper(),
            }

        manifest = self.tmp / "analysis_bearing.json"
        manifest.write_text(
            json.dumps(
                {
                    "module_results": [
                        {
                            "key": "bearing_displacement",
                            "status": "ok",
                            "artifacts": [
                                record(stats_file, "stats", "stats"),
                                record(image, "figure", "filtered"),
                            ],
                        }
                    ]
                }
            ),
            encoding="utf-8",
        )
        manifest_hash = hashlib.sha256(manifest.read_bytes()).hexdigest().upper()

        with pinned_analysis_manifest_scope(
            manifest,
            manifest_hash,
            require_source_provenance=True,
            result_root=self.tmp,
        ), jlj_report_period_scope(date(2026, 5, 1), date(2026, 5, 31)):
            section = build_main_bearing_section({}, self.tmp, stats, None, self.tmp)

        self.assertEqual(len(section.image_items or []), 1)
        self.assertEqual(section.image_items[0].path, image.resolve())

    def test_analysis_manifest_period_mismatch_is_blocked(self):
        context = {
            "available": True,
            "path": "analysis_manifest_march.json",
            "run_request": {"start_date": "2026-03-23", "end_date": "2026-03-31"},
        }
        with self.assertRaisesRegex(ValueError, "manifest period does not match"):
            validate_jlj_analysis_period(context, date(2026, 5, 1), date(2026, 5, 31))

    def test_strict_composite_uses_top_level_jiulongjiang_bridge_profile(self):
        context = {
            "available": True,
            "strict_source_provenance": True,
            "path": "analysis_manifest_composite.json",
            "bridge_profile": {"bridge_id": "jiulongjiang"},
            "run_request": {"bridge_profile": {"bridge_id": "stale_source_profile"}},
        }
        validate_jlj_analysis_profile(context)

        context["bridge_profile"] = {"bridge_id": "guanbing"}
        with self.assertRaisesRegex(ValueError, "bridge profile does not match"):
            validate_jlj_analysis_profile(context)

        context["bridge_profile"] = {}
        with self.assertRaisesRegex(ValueError, "does not declare bridge_profile.bridge_id"):
            validate_jlj_analysis_profile(context)

    def test_strict_analysis_run_cable_force_policy_is_fail_closed(self):
        for raw_valid, expected_valid, expected_include in (
            (None, None, False),
            (False, False, False),
            (True, True, True),
            ("true", None, False),
        ):
            manifest = {"manifest_type": "analysis_run"}
            if raw_valid is not None:
                manifest["cable_force_engineering_valid"] = raw_valid
            context = {
                "strict_source_provenance": True,
                "manifest": manifest,
            }

            with self.subTest(raw_valid=raw_valid):
                policy = jlj_cable_force_report_policy(context)
                self.assertIs(policy["engineering_valid"], expected_valid)
                self.assertIs(
                    policy["include_cable_force_figures"], expected_include
                )
                self.assertIs(
                    policy["include_engineering_conclusions"], expected_include
                )
                self.assertEqual(
                    bool(policy["disclosure"]), not expected_include
                )

    def test_non_strict_missing_cable_force_flag_retains_legacy_behaviour(self):
        policy = jlj_cable_force_report_policy(
            {"manifest": {"manifest_type": "analysis_run"}}
        )

        self.assertIsNone(policy["engineering_valid"])
        self.assertTrue(policy["include_cable_force_figures"])
        self.assertTrue(policy["include_engineering_conclusions"])
        self.assertEqual(policy["disclosure"], "")

    def test_strict_composite_source_coverage_without_gap_is_complete(self):
        provenance = write_plot_provenance(self.tmp / "complete.plot.json", [])
        context = {
            "strict_source_provenance": True,
            "manifest": {
                "manifest_type": "composite_analysis_recovery",
                "module_results": [
                    {
                        "key": "acceleration",
                        "artifacts": [
                            {"kind": "plot_provenance", "path": str(provenance)}
                        ],
                    }
                ],
            },
        }

        coverage = summarize_jlj_pinned_source_coverage(context)

        self.assertEqual(coverage["status"], "complete")
        self.assertEqual(coverage["plot_provenance_source_count"], 1)
        self.assertEqual(coverage["incomplete_source_record_count"], 0)
        self.assertEqual(coverage["incomplete_source_days"], [])
        self.assertFalse(coverage["disclosure_required"])
        self.assertEqual(coverage["disclosure"], "")

    def test_strict_composite_source_coverage_requires_plot_provenance(self):
        context = {
            "strict_source_provenance": True,
            "manifest": {
                "manifest_type": "composite_analysis_recovery",
                "module_results": [
                    {"key": "acceleration", "status": "ok", "artifacts": []}
                ],
            },
        }

        with self.assertRaisesRegex(ValueError, "contains no plot_provenance artifacts"):
            summarize_jlj_pinned_source_coverage(context)

    def test_strict_analysis_run_source_coverage_requires_plot_provenance(self):
        context = {
            "strict_source_provenance": True,
            "manifest": {
                "manifest_type": "analysis_run",
                "module_results": [
                    {"key": "acceleration", "status": "ok", "artifacts": []}
                ],
            },
        }

        with self.assertRaisesRegex(ValueError, "contains no plot_provenance artifacts"):
            summarize_jlj_pinned_source_coverage(context)

    def test_strict_analysis_run_source_coverage_discloses_incomplete_days(self):
        provenance = write_plot_provenance(
            self.tmp / "analysis_run_incomplete.plot.json",
            ["2026-05-03", "2026-05-04"],
        )
        context = {
            "strict_source_provenance": True,
            "manifest": {
                "manifest_type": "analysis_run",
                "module_results": [
                    {
                        "key": "acceleration",
                        "status": "ok",
                        "artifacts": [
                            {"kind": "plot_provenance", "path": str(provenance)}
                        ],
                    }
                ],
            },
        }

        coverage = summarize_jlj_pinned_source_coverage(context)

        self.assertEqual(coverage["manifest_type"], "analysis_run")
        self.assertEqual(coverage["status"], "disclosed")
        self.assertEqual(coverage["plot_provenance_source_count"], 1)
        self.assertEqual(
            coverage["incomplete_source_days"],
            ["2026-05-03", "2026-05-04"],
        )
        self.assertEqual(coverage["affected_modules"], ["acceleration"])
        self.assertTrue(coverage["disclosure_required"])
        self.assertIn("2026-05-03~2026-05-04", coverage["disclosure"])

    def test_strict_composite_source_coverage_deduplicates_days(self):
        first = write_plot_provenance(
            self.tmp / "first.plot.json", ["2026-05-01", "2026-05-02"]
        )
        second = write_plot_provenance(
            self.tmp / "second.plot.json", ["2026-05-02", "2026-05-08"]
        )
        context = {
            "strict_source_provenance": True,
            "manifest": {
                "manifest_type": "composite_analysis_recovery",
                "module_results": [
                    {
                        "key": "acceleration",
                        "artifacts": [
                            {"kind": "plot_provenance", "path": str(first)}
                        ],
                    },
                    {
                        "key": "cable_accel",
                        "artifacts": [
                            {"kind": "plot_provenance", "path": str(second)}
                        ],
                    },
                ],
            },
        }

        coverage = summarize_jlj_pinned_source_coverage(context)

        self.assertEqual(coverage["status"], "disclosed")
        self.assertEqual(coverage["plot_provenance_source_count"], 2)
        self.assertEqual(coverage["incomplete_source_record_count"], 2)
        self.assertEqual(coverage["incomplete_source_day_occurrence_count"], 4)
        self.assertEqual(
            coverage["incomplete_source_days"],
            ["2026-05-01", "2026-05-02", "2026-05-08"],
        )
        self.assertEqual(coverage["affected_modules"], ["acceleration", "cable_accel"])
        self.assertIn("2026-05-01~2026-05-02；2026-05-08", coverage["disclosure"])
        self.assertIn("未对缺失时段补造数据", coverage["disclosure"])

    def test_unverified_cable_force_keeps_vibration_but_omits_force_figures(self):
        stats = self.tmp / "stats"
        stats.mkdir()
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["PointID", "Min", "Max", "RMS10minMax"])
        sheet.append(["SLCGQ-01-K16-ZDG1", -1.2, 0.8, 0.25])
        workbook.save(stats / "cable_accel_stats.xlsx")

        unverified = build_cable_section(
            {},
            self.tmp,
            stats,
            None,
            self.tmp,
            cable_force_engineering_valid=False,
            cable_force_disclosure=JLJ_CABLE_FORCE_UNVERIFIED_DISCLOSURE,
        )
        self.assertIn("吊杆振动加速度", unverified.narrative)
        self.assertIn(JLJ_CABLE_FORCE_UNVERIFIED_DISCLOSURE, unverified.narrative)
        self.assertEqual(unverified.figure_title, "吊杆振动监测典型图")
        self.assertFalse(any("索力时程图" in item.label for item in unverified.image_items or []))

        verified = build_cable_section(
            {}, self.tmp, stats, None, self.tmp, cable_force_engineering_valid=True
        )
        self.assertEqual(verified.figure_title, "吊杆振动与索力换算典型图")
        self.assertTrue(any("索力时程图" in item.label for item in verified.image_items or []))

        empty = build_cable_section(
            {},
            self.tmp,
            self.tmp / "missing_stats",
            None,
            self.tmp,
            cable_force_engineering_valid=False,
        )
        self.assertIn(JLJ_CABLE_FORCE_UNVERIFIED_DISCLOSURE, empty.narrative)

    def test_report_manifest_records_unverified_force_policy_and_disclosure(self):
        repo_root = Path(__file__).resolve().parents[1]
        template = repo_root / "reports" / "九龙江大桥健康监测2026年3月份月报_0508.docx"
        result_root = self.tmp / "unverified_force_result"
        output_dir = self.tmp / "unverified_force_report"
        result_root.mkdir()
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["PointID", "Min", "Max", "RMS10minMax"])
        sheet.append(["SLCGQ-01-K16-ZDG1", -1.2, 0.8, 0.25])
        workbook.save(result_root / "cable_accel_stats.xlsx")
        incomplete_provenance = write_plot_provenance(
            result_root / "run_logs" / "SLCGQ-01.plot.json",
            ["2026-05-01", "2026-05-02", "2026-05-08"],
        )
        analysis_context = {
            "available": True,
            "strict_source_provenance": True,
            "path": str(result_root / "run_logs" / "analysis_manifest_composite.json"),
            "bridge_profile": {"bridge_id": "jiulongjiang"},
            "run_request": {"start_date": "2026-05-01", "end_date": "2026-05-31"},
            "missing_modules": [],
            "manifest": {
                "manifest_type": "composite_analysis_recovery",
                "cable_force_engineering_valid": False,
                "cable_force_engineering_status": "placeholder_parameters",
                "module_results": [
                    {
                        "key": "cable_accel",
                        "artifacts": [
                            {
                                "kind": "plot_provenance",
                                "path": str(incomplete_provenance),
                            }
                        ],
                    }
                ],
            },
        }

        with patch(
            "build_jlj_monthly_report.ReportBuildContext.analysis_context",
            return_value=analysis_context,
        ):
            output = build_report(
                template=template,
                config_path=repo_root / "config" / "jiulongjiang_config.json",
                result_root=result_root,
                output_dir=output_dir,
                period_label="2026年5月份",
                monitoring_range="2026.05.01~2026.05.31",
                report_date="2026年06月05日",
                update_word=False,
            )

        doc = Document(str(output))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        text += "\n" + "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        self.assertIn(JLJ_CABLE_FORCE_UNVERIFIED_DISCLOSURE, text)
        self.assertIn("滚动导出来源覆盖不完整", text)
        self.assertIn("未对缺失时段补造数据", text)
        self.assertNotIn("吊杆振动与索力换算典型图", text)

        manifests = sorted(output_dir.glob("jlj_report_build_manifest_*.json"))
        manifest = json.loads(manifests[-1].read_text(encoding="utf-8"))
        self.assertFalse(manifest["cable_force_engineering_valid"])
        self.assertEqual(manifest["cable_force_engineering_status"], "placeholder_parameters")
        self.assertFalse(manifest["cable_force_report_policy"]["include_cable_force_figures"])
        self.assertFalse(manifest["cable_force_report_policy"]["include_engineering_conclusions"])
        self.assertEqual(
            manifest["incomplete_source_days"],
            ["2026-05-01", "2026-05-02", "2026-05-08"],
        )
        self.assertEqual(manifest["plot_provenance_source_count"], 1)
        self.assertEqual(manifest["incomplete_source_record_count"], 1)
        self.assertEqual(manifest["source_coverage"]["status"], "disclosed")
        self.assertFalse(
            any("CableForce_" in record["path"] for record in manifest["report_image_sources"])
        )

    def test_wind_summary_prefers_requested_period_and_rejects_stale_only(self):
        folder = self.tmp / "wind"
        folder.mkdir()
        stale = folder / "W1_windrose_2026-03-01_2026-03-31_summary.txt"
        current = folder / "W1_windrose_2026-05-01_2026-05-31_summary.txt"
        stale.write_text("stale", encoding="utf-8")
        current.write_text("current", encoding="utf-8")
        os.utime(current, (1, 1))
        os.utime(stale, (2, 2))

        with jlj_report_period_scope(date(2026, 5, 1), date(2026, 5, 31)):
            self.assertEqual(find_wind_summary_file(self.tmp, "wind", "W1"), current.resolve())
            current.unlink()
            with self.assertRaisesRegex(ValueError, "do not match the requested report period"):
                find_wind_summary_file(self.tmp, "wind", "W1")

    def test_wind_summary_rejects_wrong_period_manifest_record(self):
        stale = self.tmp / "W1_windrose_2026-03-01_2026-03-31_summary.txt"
        stale.write_text("stale", encoding="utf-8")
        lookup = ArtifactLookupResult(
            stale,
            {
                "source": "analysis_manifest",
                "manifest": "analysis_run_manifest.json",
            },
        )
        with patch("build_jlj_monthly_report.lookup_latest_file_patterns", return_value=lookup):
            with jlj_report_period_scope(date(2026, 5, 1), date(2026, 5, 31)):
                with self.assertRaisesRegex(ValueError, "bound analysis manifest period"):
                    find_wind_summary_file(self.tmp, "wind", "W1")

    def test_wind_narrative_explicitly_reports_bridge_deck_point(self):
        stats = self.tmp / "stats"
        stats.mkdir()
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["PointID", "Mean10minMax", "Mean10minTime"])
        sheet.append(["CSFSY-01-K16-GD-A20", 17.69, "2026-03-30 16:20:33"])
        sheet.append(["CSFSY-02-K16-QM-G20", 14.72, "2026-03-25 15:47:53"])
        workbook.save(stats / "wind_stats.xlsx")

        summary_dir = self.tmp / "风速风向结果" / "风玫瑰"
        summary_dir.mkdir(parents=True)
        for point_id in ("CSFSY-01-K16-GD-A20", "CSFSY-02-K16-QM-G20"):
            (summary_dir / f"{point_id}_windrose_2026-03-23_2026-03-31_summary.txt").write_text(
                "平均风速: 1.00 m/s\n最大风速: 2.00 m/s\n主导风向: 292.5°-315.0°\n主要风速等级: 0-2 m/s",
                encoding="utf-8",
            )

        with jlj_report_period_scope(date(2026, 3, 23), date(2026, 3, 31)):
            section = build_wind_section({}, self.tmp, stats, None, self.tmp)

        self.assertIn("主桥桥面10min平均风速最大值为14.72m/s", section.narrative)
        self.assertNotIn("主桥10min平均风速最大值", section.narrative)

    def test_real_template_future_month_clears_stale_cover_patrol_and_media(self):
        repo_root = Path(__file__).resolve().parents[1]
        template = repo_root / "reports" / "九龙江大桥健康监测2026年3月份月报_0508.docx"
        result_root = self.tmp / "may_result"
        output_dir = self.tmp / "may_report"
        result_root.mkdir()

        template_section_count = len(Document(str(template)).sections)
        output = build_report(
            template=template,
            config_path=repo_root / "config" / "jiulongjiang_config.json",
            result_root=result_root,
            output_dir=output_dir,
            period_label="2026年5月份",
            monitoring_range="2026.05.01~2026.05.31",
            report_date="2026年06月05日",
            precheck_template=True,
            update_word=False,
        )

        doc = Document(str(output))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        text += "\n" + "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        self.assertIn("监测时间：2026年5月", text)
        self.assertIn("报告日期：2026年06月05日", text)
        self.assertIn("本期巡查资料未提供", text)
        self.assertNotIn("监测时间：2026年03月", text)
        self.assertNotIn("2026年03月09日上午", text)
        self.assertNotIn("2026年3月份", output.name)
        self.assertEqual(len(doc.sections), template_section_count)
        self.assertEqual(doc.element.body[-1].tag, qn("w:sectPr"))

        manifests = sorted(output_dir.glob("jlj_report_build_manifest_*.json"))
        self.assertTrue(manifests)
        manifest = json.loads(manifests[-1].read_text(encoding="utf-8"))
        self.assertIn("main_traffic", manifest["not_applicable_sections"])
        self.assertFalse(
            any(item.get("label") == "main_traffic" for item in manifest["missing_items"])
        )
        patrol = manifest["source_availability"]["patrol"]
        self.assertFalse(patrol["required"])
        self.assertEqual(patrol["status"], "not_available")
        self.assertEqual(patrol["target_period"], "2026-05")
        self.assertEqual(patrol["source"], "")
        self.assertEqual(patrol["source_sha256"], "")
        self.assertEqual(patrol["action"], "template_content_cleared_and_note_inserted")
        self.assertFalse(
            any(item.get("category") == "巡查资料缺失" for item in manifest["missing_items"])
        )

        used_rel_ids = set(doc.element.body.xpath(".//a:blip/@r:embed"))
        relationship_id_attr = (
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        )
        for element in doc.element.body.iter():
            if str(element.tag).endswith("}imagedata"):
                rel_id = element.get(relationship_id_attr)
                if rel_id:
                    used_rel_ids.add(rel_id)
        image_rel_ids = {
            rel_id
            for rel_id, relationship in doc.part.rels.items()
            if relationship.reltype == RT.IMAGE
        }
        self.assertEqual(image_rel_ids, used_rel_ids)

    def test_real_template_records_verified_period_matched_patrol_source(self):
        repo_root = Path(__file__).resolve().parents[1]
        template = repo_root / "reports" / "九龙江大桥健康监测2026年3月份月报_0508.docx"
        result_root = self.tmp / "matched_patrol_result"
        output_dir = self.tmp / "matched_patrol_report"
        result_root.mkdir()
        patrol_docx = self.tmp / "九龙江大桥巡查报告-2026年05月.docx"
        patrol = Document()
        patrol.add_paragraph("2026年05月09日巡查：本期已完成现场巡查。")
        patrol.save(patrol_docx)

        output = build_report(
            template=template,
            config_path=repo_root / "config" / "jiulongjiang_config.json",
            result_root=result_root,
            output_dir=output_dir,
            period_label="2026年5月份",
            monitoring_range="2026.05.01~2026.05.31",
            report_date="2026年6月15日",
            patrol_docx=patrol_docx,
            precheck_template=True,
            update_word=False,
        )

        text = "\n".join(paragraph.text for paragraph in Document(str(output)).paragraphs)
        self.assertIn("2026年05月09日巡查", text)
        manifests = sorted(output_dir.glob("jlj_report_build_manifest_*.json"))
        manifest = json.loads(manifests[-1].read_text(encoding="utf-8"))
        record = manifest["source_availability"]["patrol"]
        self.assertEqual(record["status"], "available")
        self.assertEqual(record["source"], str(patrol_docx.resolve()))
        self.assertEqual(record["source_period"], "2026-05")
        self.assertEqual(len(record["source_sha256"]), 64)
        self.assertEqual(record["action"], "verified_source_inserted")

    def test_real_template_required_missing_patrol_remains_blocking(self):
        repo_root = Path(__file__).resolve().parents[1]
        template = repo_root / "reports" / "九龙江大桥健康监测2026年3月份月报_0508.docx"
        result_root = self.tmp / "required_patrol_result"
        output_dir = self.tmp / "required_patrol_report"
        result_root.mkdir()
        config = json.loads(
            (repo_root / "config" / "jiulongjiang_config.json").read_text(encoding="utf-8-sig")
        )
        config["reporting"]["patrol"]["required"] = True
        config_path = self.tmp / "required_patrol_config.json"
        config_path.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")

        build_report(
            template=template,
            config_path=config_path,
            result_root=result_root,
            output_dir=output_dir,
            period_label="2026年5月份",
            monitoring_range="2026.05.01~2026.05.31",
            report_date="2026年6月15日",
            precheck_template=True,
            update_word=False,
        )

        manifests = sorted(output_dir.glob("jlj_report_build_manifest_*.json"))
        manifest = json.loads(manifests[-1].read_text(encoding="utf-8"))
        record = manifest["source_availability"]["patrol"]
        self.assertTrue(record["required"])
        self.assertEqual(record["status"], "not_available")
        self.assertTrue(
            any(item.get("category") == "巡查资料缺失" for item in manifest["missing_items"])
        )

    def test_report_qc_exception_writes_manifest_then_fails_closed(self):
        repo_root = Path(__file__).resolve().parents[1]
        template = repo_root / "reports" / "九龙江大桥健康监测2026年3月份月报_0508.docx"
        result_root = self.tmp / "qc_failure_result"
        output_dir = self.tmp / "qc_failure_report"
        result_root.mkdir()

        with patch(
            "build_jlj_monthly_report.check_jlj_report",
            side_effect=RuntimeError("synthetic QC crash"),
        ):
            with self.assertRaisesRegex(RuntimeError, "QC execution failed"):
                build_report(
                    template=template,
                    config_path=repo_root / "config" / "jiulongjiang_config.json",
                    result_root=result_root,
                    output_dir=output_dir,
                    period_label="2026年5月份",
                    monitoring_range="2026.05.01~2026.05.31",
                    report_date="2026年06月05日",
                    precheck_template=True,
                    update_word=False,
                )

        manifests = sorted(output_dir.glob("jlj_report_build_manifest_*.json"))
        self.assertTrue(manifests)
        payload = manifests[-1].read_text(encoding="utf-8")
        self.assertIn("report_qc_failed", payload)
        self.assertIn("synthetic QC crash", payload)


if __name__ == "__main__":
    unittest.main()
