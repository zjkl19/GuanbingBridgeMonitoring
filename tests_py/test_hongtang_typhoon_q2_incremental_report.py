from __future__ import annotations

import base64
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from reporting.build_hongtang_typhoon_q2_incremental_report import (
    build_incremental_report,
    element_text,
    find_heading,
    load_manifest,
    locate_anchors,
    validate_refreshed_base,
)


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _body_payload_between(doc, start, end):
    children = list(doc._element.body.iterchildren())
    start_index = children.index(start)
    end_index = children.index(end)
    return [
        element
        for element in children[start_index + 1 : end_index]
        if not (
            element.tag == qn("w:p")
            and element.find(".//" + qn("w:sectPr")) is not None
            and not element_text(element)
        )
    ]


class HongtangTyphoonQ2IncrementalReportTests(unittest.TestCase):
    def _make_base(self, path: Path, picture: Path) -> None:
        doc = Document()
        doc.add_paragraph("G316线洪塘大桥桥梁健康监测")
        doc.add_paragraph("周期报告")
        doc.add_paragraph("（监测时间：2026年7月10日~2026年7月12日）")
        signatures = doc.add_table(rows=3, cols=2)
        signatures.cell(0, 0).text = "项目负责："

        metadata = doc.add_table(rows=4, cols=5)
        metadata.cell(3, 0).text = "监测结果"
        first_page_cell = metadata.cell(3, 2)
        first_page_cell.text = "首页导语保留。"
        for text in (
            "1、交通状况监测",
            "Q2旧交通数据5906370辆。",
            "2、结构应变监测",
            "Q2旧应变152.3με。",
            "3、主塔倾斜监测",
            "Q2旧倾角。",
            "4、支座变位监测",
            "Q2旧变位。",
            "5、吊索索力监测",
            "Q2旧索力保留为底座第5项。",
        ):
            first_page_cell.add_paragraph(text)
        continuation = doc.add_table(rows=2, cols=2)
        continuation.cell(0, 0).text = "监测结果"
        result_cell = continuation.cell(0, 1)
        result_cell.text = ""
        for text in (
            "6、主梁、主塔振动监测",
            "Q2原有第6项结论必须保留。",
            "7、风向风速监测",
            "标准底座台风第7项：W1的10min平均风速最大值为5.21m/s。",
            "8、地震动监测",
            "Q2原有第8项结论必须保留。",
        ):
            result_cell.add_paragraph(text)
        continuation.cell(1, 0).text = "建  议"
        continuation.cell(1, 1).text = (
            "针对目前的监测状况，建议如下：\n"
            "洪塘大桥存在超重车辆通行现象，管养单位应采取有效的限载措施，对过往超载车辆进行必要管控。\n"
            "（本栏以下空白）"
        )

        doc.add_heading("监测概况", level=1)
        doc.add_heading("工程概况", level=2)
        doc.add_paragraph("工程概况保留。")
        doc.add_heading("监测内容", level=2)
        doc.add_paragraph("监测内容保留。")
        doc.add_heading("健康监测系统运行状况", level=2)
        doc.add_paragraph("Q2旧运行状况必须删除。")
        doc.add_table(rows=2, cols=2).cell(0, 0).text = "Q2旧缺失表必须删除"
        doc.add_paragraph().add_run().add_picture(str(picture))
        doc.add_heading("软硬件维护状况", level=2)
        doc.add_paragraph("Q2旧维护状况必须删除。")
        doc.add_table(rows=2, cols=2).cell(0, 0).text = "Q2旧维护表必须删除"
        doc.add_paragraph().add_run().add_picture(str(picture))

        doc.add_heading("监测项目及内容", level=1)
        for title in (
            "交通状况监测",
            "结构应变监测",
            "主塔倾斜监测",
            "支座变位监测",
            "吊索索力监测",
            "主梁、主塔振动监测",
            "风向风速监测",
            "地震动监测",
        ):
            doc.add_heading(title, level=2)
            doc.add_paragraph(f"第2章{title}布置说明保留。")

        doc.add_heading("报警阈值设置", level=1)
        doc.add_heading("报警阈值设置", level=2)
        doc.add_paragraph("报警阈值保留。")
        doc.add_heading("监测结果", level=1)
        stale = {
            "交通状况监测": "Q2旧交通结果5906370辆必须删除。",
            "结构应变监测": "Q2旧应变结果152.3με必须删除。",
            "主塔倾斜监测": "Q2旧倾斜结果必须删除。",
            "支座变位监测": "Q2旧支座变位结果必须删除。",
        }
        for title in (
            "交通状况监测",
            "结构应变监测",
            "主塔倾斜监测",
            "支座变位监测",
        ):
            doc.add_heading(title, level=2)
            doc.add_paragraph(stale[title])
            doc.add_table(rows=2, cols=2).cell(0, 0).text = stale[title]
            doc.add_paragraph().add_run().add_picture(str(picture))
        for title, original in (
            ("吊索索力监测", "4.5原有Q2索力正文必须保留。"),
            ("主梁、主塔振动监测", "4.6原有Q2振动正文必须保留。"),
            ("风向风速监测", "4.7原有Q2风正文必须保留。"),
            ("地震动监测", "4.8原有Q2地震正文必须保留。"),
        ):
            doc.add_heading(title, level=2)
            doc.add_paragraph(original)
        doc.add_paragraph("(以下无正文)")
        doc.save(path)

    def _make_manifest(self, path: Path) -> None:
        data = {
            "status": "ok",
            "window": {
                "start": "2026-07-10T23:20:00",
                "landfall": "2026-07-11T23:20:00",
                "end": "2026-07-12T09:00:05.066000",
            },
            "wind_summary": {
                "W1": {
                    "raw_max": 14.74487,
                    "raw_max_time": "2026-07-11T16:09:29.416000",
                    "raw_max_direction": 259.99,
                    "max_10min": 5.21159,
                    "max_10min_time": "2026-07-11T14:10:00",
                    "max_10min_direction": 270.54,
                    "pre_mean": 3.2136,
                    "pre_max_10min": 5.21159,
                    "post_mean": 2.5453,
                    "post_max_10min": 3.2748,
                    "bins": 202,
                },
                "W2": {
                    "raw_max": 17.4194,
                    "raw_max_time": "2026-07-11T17:02:49.966000",
                    "raw_max_direction": 76.37,
                    "max_10min": 7.6013,
                    "max_10min_time": "2026-07-11T16:00:00",
                    "max_10min_direction": 88.32,
                    "pre_mean": 4.2103,
                    "pre_max_10min": 7.6013,
                    "post_mean": 1.0539,
                    "post_max_10min": 2.0112,
                    "bins": 202,
                },
            },
            "structure_summary": {
                "主梁加速度": {
                    "pre_median": 0.069545,
                    "post_median": 0.071385,
                    "median_ratio": 1.0265,
                    "maximum": 0.29619,
                    "maximum_point": "A4",
                    "maximum_time": "2026-07-10T23:20:00",
                    "bins": 202,
                },
                "主塔加速度": {
                    "pre_median": 0.009195,
                    "post_median": 0.00953,
                    "median_ratio": 1.0364,
                    "maximum": 0.11362,
                    "maximum_point": "A10-X",
                    "maximum_time": "2026-07-10T23:20:00",
                    "bins": 202,
                },
                "南侧索振动": {
                    "pre_median": 0.81004,
                    "post_median": 0.6886,
                    "median_ratio": 0.8501,
                    "maximum": 2.23327,
                    "maximum_point": "CS5",
                    "maximum_time": "2026-07-11T22:00:00",
                    "bins": 202,
                },
                "北侧索振动": {
                    "pre_median": 0.649385,
                    "post_median": 0.568725,
                    "median_ratio": 0.8758,
                    "maximum": 2.04174,
                    "maximum_point": "CX7",
                    "maximum_time": "2026-07-10T23:20:00",
                    "bins": 202,
                },
            },
            "quality": {
                "expected_complete_10min_bins": 202,
                "speed_rows": 6910720,
                "speed_rejected_rows": 0,
                "direction_rejected_rate": 0.01618,
            },
            "missing_entries": [],
            "charts": [],
        }
        path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

    def test_build_preserves_q2_structure_and_clears_only_unavailable_sections(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            picture = root / "pixel.png"
            picture.write_bytes(PNG_1X1)
            base = root / "base.docx"
            self._make_base(base, picture)
            manifest = root / "manifest.json"
            self._make_manifest(manifest)
            for basename in (
                "wind_speed_window.png",
                "wind_maximum_comparison.png",
                "wind_direction_window.png",
                "structure_response_window.png",
            ):
                (root / basename).write_bytes(PNG_1X1)
            output = root / "incremental.docx"

            _, audit = build_incremental_report(
                base,
                manifest,
                output,
                charts_dir=root,
            )

            self.assertEqual("ok", audit["status"])
            self.assertTrue(audit["base_refresh_gate"]["period_title"])
            self.assertTrue(audit["base_refresh_gate"]["window_dates"])
            self.assertTrue(audit["base_refresh_gate"]["front_item7_w1_10min"])
            self.assertEqual(["1.3", "1.4", "4.1", "4.2", "4.3", "4.4"], audit["blank_sections_verified"])
            self.assertFalse(audit["front_summary"]["whole_cell_replaced"])
            doc = Document(output)

            chapter_one = find_heading(doc, "监测概况", level=1)
            section_13 = find_heading(doc, "健康监测系统运行状况", level=2, after=chapter_one)
            section_14 = find_heading(doc, "软硬件维护状况", level=2, after=section_13)
            chapter_two = find_heading(doc, "监测项目及内容", level=1, after=section_14)
            self.assertEqual([], _body_payload_between(doc, section_13, section_14))
            self.assertEqual([], _body_payload_between(doc, section_14, chapter_two))

            chapter_four = find_heading(doc, "监测结果", level=1, after=chapter_two)
            titles = (
                "交通状况监测",
                "结构应变监测",
                "主塔倾斜监测",
                "支座变位监测",
                "吊索索力监测",
                "主梁、主塔振动监测",
                "风向风速监测",
                "地震动监测",
            )
            headings = []
            cursor = chapter_four
            for title in titles:
                cursor = find_heading(doc, title, level=2, after=cursor)
                headings.append(cursor)
            for index in range(4):
                self.assertEqual([], _body_payload_between(doc, headings[index], headings[index + 1]))

            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            for original in (
                "4.5原有Q2索力正文必须保留。",
                "4.6原有Q2振动正文必须保留。",
                "4.7原有Q2风正文必须保留。",
                "4.8原有Q2地震正文必须保留。",
            ):
                self.assertIn(original, text)
            self.assertIn("台风期间吊索振动增量分析", text)
            self.assertIn("台风期间主梁、主塔振动增量分析", text)
            self.assertIn("台风期间风速最大值及登陆前后增量分析", text)
            self.assertIn("台风影响综合分析、运营建议与数据限制", text)
            self.assertIn("https://news.weather.com.cn/2026/07/4711229.shtml", text)
            self.assertIn("引用日期：2026年7月12日", text)
            self.assertEqual(4, len(doc.inline_shapes))

            front_text = "\n".join(paragraph.text for paragraph in doc.tables[2].cell(0, 1).paragraphs)
            self.assertNotIn("Q2原有第6项结论必须保留。", front_text)
            self.assertIn("标准底座台风第7项：W1的10min平均风速最大值为5.21m/s。", front_text)
            self.assertIn("Q2原有第8项结论必须保留。", front_text)
            self.assertEqual(2, front_text.count("台风窗口增补："))
            self.assertIn("0.2962m/s²", front_text)
            self.assertIn("0.1136m/s²", front_text)
            self.assertIn("未见登陆后持续、多测点同步放大", front_text)
            for number in (6, 7, 8):
                self.assertEqual(1, front_text.count(f"{number}、"))

            advice_text = doc.tables[2].cell(1, 1).text
            for stale in ("超重车辆", "超载车辆", "限载措施", "本栏以下空白"):
                self.assertNotIn(stale, advice_text)
            self.assertIn("针对台风影响期监测情况", advice_text)
            self.assertIn("持续关注W1/W2风速峰值", advice_text)

            first_page_text = "\n".join(
                paragraph.text for paragraph in doc.tables[1].cell(3, 2).paragraphs
            )
            for value in ("5906370", "152.3με", "Q2旧倾角", "Q2旧变位"):
                self.assertNotIn(value, first_page_text)
            self.assertIn("Q2旧索力保留为底座第5项。", first_page_text)
            for number in (1, 2, 3, 4, 5):
                self.assertEqual(1, first_page_text.count(f"{number}、"))
            self.assertNotIn("(以下无正文)", text)
            self.assertNotIn("第二季度报告", text)
            self.assertNotIn("周期报告", text)
            self.assertIn("洪塘大桥台风巴威影响监测专题报告", text)
            self.assertIn("2026年07月10日23:20~2026年07月12日09:00", text)
            self.assertEqual(
                "2026年07月10日23:20~2026年07月12日09:00",
                doc.tables[1].cell(1, 4).text,
            )

            xml = doc._element.xml
            self.assertIn("SEQ 图", xml)
            self.assertIn("SEQ 表", xml)
            self.assertGreaterEqual(audit["sequence_fields"]["图"], 4)
            self.assertGreaterEqual(audit["sequence_fields"]["表"], 5)

    def test_manifest_missing_entries_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            manifest = root / "manifest.json"
            self._make_manifest(manifest)
            data = json.loads(manifest.read_text(encoding="utf-8"))
            data["missing_entries"] = ["W2_speed"]
            manifest.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "missing source entries"):
                load_manifest(manifest)

    def test_raw_q2_delivery_is_rejected_by_production_gate(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            picture = root / "pixel.png"
            picture.write_bytes(PNG_1X1)
            base = root / "raw_q2.docx"
            self._make_base(base, picture)
            doc = Document(base)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() == "周期报告":
                    paragraph.text = "第二季度报告"
                if "2026年7月10日~2026年7月12日" in paragraph.text:
                    paragraph.text = "（监测时间：2026年4月1日~2026年6月30日）"
            cell = doc.tables[2].cell(0, 1)
            for paragraph in cell.paragraphs:
                if paragraph.text.startswith("标准底座台风第7项"):
                    paragraph.text = "Q2旧风摘要：W1最大10min平均风速6.89m/s。"
            doc.save(base)
            manifest_path = root / "manifest.json"
            self._make_manifest(manifest_path)
            raw_doc = Document(base)
            with self.assertRaisesRegex(RuntimeError, "not a refreshed typhoon-window"):
                validate_refreshed_base(raw_doc, locate_anchors(raw_doc), load_manifest(manifest_path))

    def test_quick_mode_removes_all_q2_result_payloads(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            picture = root / "pixel.png"
            picture.write_bytes(PNG_1X1)
            base = root / "raw_q2.docx"
            self._make_base(base, picture)
            doc = Document(base)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() == "周期报告":
                    paragraph.text = "第二季度报告"
            doc.save(base)
            manifest = root / "manifest.json"
            self._make_manifest(manifest)
            for basename in (
                "wind_speed_window.png",
                "wind_maximum_comparison.png",
                "wind_direction_window.png",
                "structure_response_window.png",
            ):
                (root / basename).write_bytes(PNG_1X1)
            output = root / "quick.docx"
            _, audit = build_incremental_report(
                base,
                manifest,
                output,
                charts_dir=root,
                quick_from_template=True,
            )
            self.assertEqual("quick_from_template", audit["mode"])
            self.assertIn("4.8", audit["blank_sections_verified"])
            result = Document(output)
            text = "\n".join(paragraph.text for paragraph in result.paragraphs)
            for stale in (
                "4.5原有Q2索力正文必须保留。",
                "4.6原有Q2振动正文必须保留。",
                "4.7原有Q2风正文必须保留。",
                "4.8原有Q2地震正文必须保留。",
            ):
                self.assertNotIn(stale, text)
            self.assertIn("未取得该窗口的索力识别结果", text)
            self.assertIn("台风期间主梁、主塔振动增量分析", text)
            self.assertIn("台风期间风速最大值及登陆前后增量分析", text)
            front = result.tables[1].cell(3, 2).text + "\n" + result.tables[2].cell(0, 1).text
            self.assertNotIn("Q2旧索力保留", front)
            self.assertNotIn("Q2原有第6项", front)
            self.assertNotIn("Q2原有第8项", front)
            self.assertIn("仅分析吊索振动，不分析索力", front)
            self.assertIn("台风窗口轻量结论：主梁、主塔", front)
            self.assertIn("台风窗口轻量结论：W1桥面、W2塔顶", front)
            chapter_two = find_heading(result, "监测项目及内容", level=1)
            chapter_four = find_heading(result, "监测结果", level=1, after=chapter_two)
            eq = find_heading(result, "地震动监测", level=2, after=chapter_four)
            conclusion = find_heading(
                result,
                "台风影响综合分析、运营建议与数据限制",
                level=2,
                after=eq,
            )
            self.assertEqual([], _body_payload_between(result, eq, conclusion))


if __name__ == "__main__":
    unittest.main()
