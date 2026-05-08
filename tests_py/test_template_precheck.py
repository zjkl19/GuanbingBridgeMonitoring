from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from reporting.template_precheck import (
    TemplateIssue,
    build_precheck_payload,
    check_template,
    summarize_template,
    write_precheck_report,
)


class TemplatePrecheckTests(unittest.TestCase):
    def test_missing_template_file_returns_missing_file_issue(self) -> None:
        missing = Path(tempfile.gettempdir()) / "template_that_does_not_exist.docx"

        issues = check_template("hongtang_period", missing)

        self.assertEqual(len(issues), 1)
        self.assertEqual(issues[0].code, "missing-file")
        self.assertIn(str(missing), issues[0].message)

    def test_incomplete_hongtang_template_reports_missing_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "incomplete.docx"
            doc = Document()
            doc.add_paragraph("健康监测系统运行状况")
            doc.save(template)

            issues = check_template("hongtang_period", template)

            codes = {issue.code for issue in issues}
            messages = "\n".join(issue.message for issue in issues)
            self.assertIn("missing-text", codes)
            self.assertIn("交通状况监测", messages)
            self.assertIn("季度交通状况分月统计表", messages)

    def test_jlj_template_precheck_reports_missing_summary_table(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "jlj_incomplete.docx"
            doc = Document()
            doc.add_paragraph("温度监测")
            doc.save(template)

            issues = check_template("jlj_monthly", template)

            self.assertIn("missing-table", {issue.code for issue in issues})
            self.assertIn("Jiulongjiang summary table", "\n".join(issue.message for issue in issues))

    def test_precheck_payload_status(self) -> None:
        ok_payload = build_precheck_payload("kind", Path("a.docx"), [])
        warn_payload = build_precheck_payload("kind", Path("a.docx"), [], warnings=["warn"])
        fail_payload = build_precheck_payload("kind", Path("a.docx"), [TemplateIssue("x", "bad")])

        self.assertEqual(ok_payload["status"], "ok")
        self.assertEqual(warn_payload["status"], "warning")
        self.assertEqual(fail_payload["status"], "failed")
        self.assertEqual(fail_payload["issues"][0]["code"], "x")
        self.assertIn("template_summary", ok_payload["context"])

    def test_write_precheck_report_outputs_txt_and_json(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out_dir = Path(tmp)
            txt_path, json_path = write_precheck_report(
                "hongtang_period",
                Path("demo.docx"),
                [TemplateIssue("missing-text", "缺少锚点")],
                out_dir,
                warnings=["缺少 stats"],
                context={"result_root": "E:/data"},
            )

            self.assertTrue(txt_path.exists())
            self.assertTrue(json_path.exists())
            self.assertIn("缺少锚点", txt_path.read_text(encoding="utf-8"))
            payload = json.loads(json_path.read_text(encoding="utf-8"))
            self.assertEqual(payload["status"], "failed")
            self.assertEqual(payload["warning_count"], 1)
            self.assertEqual(payload["context"]["result_root"], "E:/data")
            self.assertIn("template_summary", payload["context"])

    def test_summarize_template_counts_docx_structure(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            doc = Document()
            doc.add_paragraph("anchor")
            doc.add_table(rows=1, cols=1).cell(0, 0).text = "cell"
            doc.save(template)

            summary = summarize_template(template)
            self.assertEqual(summary["exists"], 1)
            self.assertEqual(summary["paragraph_count"], 1)
            self.assertEqual(summary["table_count"], 1)
            self.assertEqual(summary["table_text_count"], 1)


if __name__ == "__main__":
    unittest.main()
