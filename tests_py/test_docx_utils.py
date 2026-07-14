import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "reporting"))

from docx import Document  # noqa: E402
from docx.opc.constants import RELATIONSHIP_TYPE as RT  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from PIL import Image  # noqa: E402

from docx_utils import (  # noqa: E402
    find_paragraph_contains,
    paragraph_has_image,
    prune_unused_document_image_relationships,
    replace_picture_before_anchor,
    set_cell_paragraphs,
    set_cell_text_preserve,
)


class TestDocxUtils(unittest.TestCase):
    def test_set_cell_text_preserve_reuses_first_paragraph(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "old"
        cell.add_paragraph("extra")
        set_cell_text_preserve(cell, "new")
        self.assertEqual(cell.paragraphs[0].text, "new")
        self.assertEqual(cell.paragraphs[1].text, "")

    def test_set_cell_paragraphs_applies_bold_lines(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        set_cell_paragraphs(cell, ["title", "body"], bold_indices={0})
        self.assertEqual([p.text for p in cell.paragraphs], ["title", "body"])
        self.assertTrue(cell.paragraphs[0].runs[0].bold)
        self.assertFalse(cell.paragraphs[1].runs[0].bold)

    def test_replace_picture_before_anchor_replaces_existing_picture(self):
        with tempfile.TemporaryDirectory() as td:
            img_path = Path(td) / "x.png"
            Image.new("RGB", (10, 10), "white").save(img_path)
            doc = Document()
            anchor = doc.add_paragraph("caption anchor")
            ok, _ = replace_picture_before_anchor(doc, "caption anchor", img_path)
            self.assertTrue(ok)
            pictures = [p for p in doc.paragraphs if paragraph_has_image(p)]
            self.assertEqual(len(pictures), 1)
            ok, _ = replace_picture_before_anchor(doc, "caption anchor", img_path)
            self.assertTrue(ok)
            pictures = [p for p in doc.paragraphs if paragraph_has_image(p)]
            self.assertEqual(len(pictures), 1)

    def test_prune_image_relationships_preserves_linked_blip(self):
        doc = Document()
        used_rel_id = doc.part.relate_to(
            "https://example.invalid/used.png", RT.IMAGE, is_external=True
        )
        unused_rel_id = doc.part.relate_to(
            "https://example.invalid/unused.png", RT.IMAGE, is_external=True
        )
        paragraph = doc.add_paragraph()
        drawing = OxmlElement("w:drawing")
        blip = OxmlElement("a:blip")
        blip.set(qn("r:link"), used_rel_id)
        drawing.append(blip)
        paragraph._p.append(drawing)

        dropped = prune_unused_document_image_relationships(doc)

        self.assertIn(used_rel_id, doc.part.rels)
        self.assertNotIn(unused_rel_id, doc.part.rels)
        self.assertEqual(dropped, [unused_rel_id])


if __name__ == "__main__":
    unittest.main()
