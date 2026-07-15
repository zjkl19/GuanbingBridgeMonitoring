from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path
from typing import Any
from unittest.mock import patch


def _make_reporting_importable() -> None:
    """Expose source-tree report modules without changing the frozen layout."""
    if getattr(sys, "frozen", False):
        return
    reporting_root = Path(__file__).resolve().parents[1] / "reporting"
    value = str(reporting_root)
    if value not in sys.path:
        sys.path.insert(0, value)


_make_reporting_importable()

# These imports are intentionally static.  Besides being clearer than loading
# the old report GUI dynamically, they let PyInstaller collect the complete
# headless report runtime into the workbench executable.
from report_job import REPORT_TYPE_NAMES  # noqa: E402
from report_job_cli import request_from_context, run_context  # noqa: E402
from word_pdf_export import WordPdfExportResult  # noqa: E402


def run_embedded_report_job(
    context_path: Path,
    status_path: Path,
    result_path: Path,
    launch_id: str = "",
) -> int:
    """Run one report in the workbench's hidden worker process."""
    return run_context(
        context_path.expanduser().resolve(),
        status_path.expanduser().resolve(),
        result_path.expanduser().resolve(),
        launch_id,
    )


def report_runtime_contract() -> dict[str, Any]:
    report_types = sorted(REPORT_TYPE_NAMES)
    return {
        "ok": bool(report_types) and callable(run_context) and callable(request_from_context),
        "runtime": "embedded_headless_worker",
        "standalone_report_window": False,
        "report_type_count": len(report_types),
        "report_types": report_types,
    }


def _report_gate_contract_smoke() -> None:
    from .config_layers import config_dependency_sha256
    from .models import JobContext, file_sha256

    with tempfile.TemporaryDirectory(prefix="bms_embedded_report_gate_") as folder:
        root = Path(folder)
        data_root = root / "data"
        data_root.mkdir()
        config = root / "config.json"
        template = root / "template.docx"
        provenance = data_root / "temperature.plot.json"
        manifest = data_root / "analysis_manifest.json"
        config.write_text("{}", encoding="utf-8")
        template.write_bytes(b"embedded report gate smoke")
        provenance.write_text(json.dumps({
            "series": [{
                "sampling_mode": "full",
                "reduction_applied": False,
                "input_count": 10,
                "finite_count": 9,
                "plotted_finite_count": 9,
                "source": {
                    "source_sample_count": 10,
                    "finite_source_sample_count": 9,
                    "completeness_scope": "required_export_contribution",
                    "internal_gap_coverage_assessed": True,
                    "calendar_day_count_requested": 1,
                    "complete_day_count": 1,
                    "incomplete_day_count": 0,
                    "incomplete_days": [],
                    "missing_required_sources": [],
                },
            }],
        }), encoding="utf-8")
        manifest_payload = {
            "status": "ok",
            "bridge_profile": {"bridge_id": "guanbing"},
            "run_request": {
                "data_root": str(data_root),
                "start_date": "2026-01-01",
                "end_date": "2026-01-01",
                "config_path": str(config.resolve()),
                "config_sha256": config_dependency_sha256(config),
            },
            "module_results": [{
                "key": "temperature",
                "status": "ok",
                "artifacts": [{"kind": "plot_provenance", "path": str(provenance)}],
            }],
        }
        manifest.write_text(json.dumps(manifest_payload), encoding="utf-8")
        context = JobContext.create(
            project_root=Path.cwd(),
            bridge_id="guanbing",
            bridge_name="guanbing",
            data_root=data_root,
            start_date="2026-01-01",
            end_date="2026-01-01",
            config_path=config,
            selected_modules=["temperature"],
            options={},
            report_type="guanbing_monthly",
            template_path=template,
            output_dir=root / "output",
        )
        context.analysis.state = "completed"
        context.analysis.manifest_path = str(manifest)
        context.analysis.manifest_sha256 = file_sha256(manifest)
        context.report.plots_approved = True
        context_path = context.write(root / "job_context.json")
        request_from_context(context_path)

        manifest_payload["module_results"][0]["artifacts"] = []
        manifest.write_text(json.dumps(manifest_payload), encoding="utf-8")
        context.analysis.manifest_sha256 = file_sha256(manifest)
        context.write(context_path)
        try:
            request_from_context(context_path)
        except RuntimeError:
            return
        raise RuntimeError("report gate accepted an analysis manifest without plot provenance")


def _visual_qc_contract_smoke() -> None:
    from PIL import Image
    from report_visual_qc import analyze_page_image, create_contact_sheet

    with tempfile.TemporaryDirectory(prefix="bms_embedded_visual_qc_") as folder:
        root = Path(folder)
        page = root / "page-1.png"
        Image.new("RGB", (100, 140), "white").save(page)
        analysis = analyze_page_image(page)
        contact = create_contact_sheet([page], root / "contact.png")
        if not analysis["blank"] or not contact.is_file():
            raise RuntimeError("embedded visual-QC contract smoke failed")


def _embedded_report_job_smoke() -> None:
    """Exercise the frozen worker, status/result protocol and a real DOCX build.

    The regular bridge builders need large bridge-specific templates and result
    trees, so the packaging smoke substitutes only the builder call with a
    one-page deterministic DOCX fixture.  Everything around it is production
    code: the report gate, pinned analysis manifest, ``run_context`` worker,
    atomic status/result files, structural DOCX QC and visual-QC invocation.
    """
    from docx import Document
    from report_job import build_qc

    from .config_layers import config_dependency_sha256
    from .models import JobContext, file_sha256

    with tempfile.TemporaryDirectory(prefix="bms_embedded_report_job_") as folder:
        root = Path(folder)
        data_root = root / "data"
        output_dir = root / "output"
        data_root.mkdir()
        output_dir.mkdir()
        config = root / "config.json"
        template = root / "template.docx"
        provenance = data_root / "temperature.plot.json"
        manifest = data_root / "analysis_manifest.json"
        config.write_text("{}\n", encoding="utf-8")
        template_doc = Document()
        template_doc.add_heading("Embedded report worker smoke", level=1)
        template_doc.save(template)
        provenance.write_text(json.dumps({
            "series": [{
                "sampling_mode": "full",
                "reduction_applied": False,
                "input_count": 10,
                "finite_count": 10,
                "plotted_finite_count": 10,
                "source": {
                    "source_sample_count": 10,
                    "finite_source_sample_count": 10,
                    "completeness_scope": "required_export_contribution",
                    "internal_gap_coverage_assessed": True,
                    "calendar_day_count_requested": 1,
                    "complete_day_count": 1,
                    "incomplete_day_count": 0,
                    "incomplete_days": [],
                    "missing_required_sources": [],
                },
            }],
        }), encoding="utf-8")
        manifest.write_text(json.dumps({
            "status": "ok",
            "bridge_profile": {"bridge_id": "guanbing"},
            "run_request": {
                "data_root": str(data_root),
                "start_date": "2026-01-01",
                "end_date": "2026-01-01",
                "config_path": str(config.resolve()),
                "config_sha256": config_dependency_sha256(config),
            },
            "module_results": [{
                "key": "temperature",
                "status": "ok",
                "artifacts": [{"kind": "plot_provenance", "path": str(provenance)}],
            }],
        }), encoding="utf-8")
        context = JobContext.create(
            project_root=Path.cwd(),
            bridge_id="guanbing",
            bridge_name="guanbing",
            data_root=data_root,
            start_date="2026-01-01",
            end_date="2026-01-01",
            config_path=config,
            selected_modules=["temperature"],
            options={},
            report_type="guanbing_monthly",
            template_path=template,
            output_dir=output_dir,
        )
        context.analysis.state = "completed"
        context.analysis.manifest_path = str(manifest)
        context.analysis.manifest_sha256 = file_sha256(manifest)
        context.report.plots_approved = True
        context_path = context.write(root / "job_context.json")
        status_path = root / "report_status.json"
        result_path = root / "report_result.json"

        def build_smoke_report(**_kwargs: Any) -> tuple[Path, Path]:
            report_path = output_dir / "embedded_report_worker_smoke.docx"
            document = Document(str(template))
            document.add_paragraph("The embedded report worker completed a real DOCX build.")
            document.save(report_path)
            report_manifest = output_dir / "report_build_manifest_smoke.json"
            report_manifest.write_text(json.dumps({
                "schema_version": 1,
                "manifest_type": "report_build",
                "report_type": "guanbing_monthly",
                "status": "ok",
                "output_docx": str(report_path),
                "output_docx_image_count": 0,
                "missing_count": 0,
                "missing_items": [],
                "warnings": [],
            }), encoding="utf-8")
            return report_path, report_manifest

        with patch(
            "report_job.build_guanbing_monthly_report", side_effect=build_smoke_report
        ), patch(
            "report_job.raise_for_template"
        ), patch(
            "report_job.export_authoritative_word_pdf",
            return_value=WordPdfExportResult(
                None,
                "skipped",
                "embedded runtime smoke uses LibreOffice preview only",
            ),
        ):
            exit_code = run_context(context_path, status_path, result_path)
        if exit_code != 0:
            detail = result_path.read_text(encoding="utf-8-sig") if result_path.is_file() else ""
            raise RuntimeError(f"embedded report worker smoke failed: {detail}")
        status = json.loads(status_path.read_text(encoding="utf-8-sig"))
        result = json.loads(result_path.read_text(encoding="utf-8-sig"))
        report_path = Path(result.get("report_path") or "")
        report_manifest = Path(result.get("manifest_path") or "")
        if status.get("state") != "completed" or result.get("state") != "completed":
            raise RuntimeError("embedded report worker did not close its status/result protocol")
        if not report_path.is_file() or not report_manifest.is_file():
            raise RuntimeError("embedded report worker did not create DOCX and report manifest")
        qc = build_qc(report_path, report_manifest, None)
        if qc.get("status") == "failed" or result.get("qc", {}).get("status") == "failed":
            raise RuntimeError("embedded report worker output failed structural QC")


def write_report_runtime_smoke(output_path: Path | None = None) -> int:
    payload = report_runtime_contract()
    try:
        _report_gate_contract_smoke()
        payload["report_gate_contract"] = True
        _embedded_report_job_smoke()
        payload["embedded_report_job"] = True
        _visual_qc_contract_smoke()
        payload["visual_qc_contract"] = True
    except Exception as exc:  # noqa: BLE001
        payload["ok"] = False
        payload["error"] = str(exc)
    text = json.dumps(payload, ensure_ascii=False, indent=2)
    if output_path is not None:
        output_path = output_path.expanduser().resolve()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(text + "\n", encoding="utf-8")
    if sys.stdout is not None:
        print(text)
    return 0 if payload["ok"] else 1
